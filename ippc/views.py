import autocomplete_light
#autocomplete_light.autodiscover()

from django.shortcuts import render
from django.utils.translation import ugettext_lazy as _
from django.contrib.messages import info, error
from django.utils import timezone
from django.core import mail
from django.conf import settings

from django.contrib.auth.models import User,Group
from .models import ReminderMessage,ContactType,PublicationLibrary,Publication,EppoCode,EmailUtilityMessage, EmailUtilityMessageFile, Poll_Choice, Poll,PollVotes, IppcUserProfile,\
CountryPage,PartnersPage, PestStatus, PestReport, IS_PUBLIC, IS_HIDDEN, Publication,PestReportFile,PestReportUrl,\
PublicationFile,PublicationUrl,ReportingObligation_File,ReportingObligationUrl, EventreportingFile,EventreportingUrl,    ImplementationISPMFile,ImplementationISPMUrl, PestFreeAreaFile,PestFreeAreaUrl,\
DraftProtocol,DraftProtocolComments,NotificationMessageRelate,CommentFile,AnswerVotes,\
ReportingObligation, BASIC_REP_TYPE_CHOICES, EventReporting, EVT_REP_TYPE_CHOICES,Website,CnPublication,PartnersPublication,PartnersNews, PartnersWebsite,CountryNews, \
PestFreeArea,ImplementationISPM,REGIONS, IssueKeywordsRelate,CommodityKeywordsRelate,EventreportingFile,ReportingObligation_File,\
ContactUsEmailMessage,FAQsItem,FAQsCategory,QAQuestion, QAAnswer,UserAutoRegistration,IRSSActivity,IRSSActivityFile,IRSS_ACT_TYPE_CHOICES,\
TransFAQsCategory,TransFAQsItem,MassEmailUtilityMessage,MassEmailUtilityMessageFile,\
OCPHistory, PartnersContactPointHistory,CnEditorsHistory,PartnersEditorHistory,UserMembershipHistory,MediaKitDocument,MyTool,\
PhytosanitaryTreatment,PhytosanitaryTreatmentPestsIdentity,PhytosanitaryTreatmentCommodityIdentity,CertificatesTool,WorkshopCertificatesTool,CPMS,TOPIC_PRIORITY_CHOICES,\
Topic,TopicAssistants,TopicLeads,TransTopic,TOPIC_STATUS_CHOICES,SC_TYPE_CHOICES,B_CertificatesTool,NROStats,MyTool2,ContributedResource,CommitteeMeeting
#TransReportingObligation,#COOPTYPE_CHOICES,

from mezzanine.core.models import Displayable, CONTENT_STATUS_DRAFT, CONTENT_STATUS_PUBLISHED
from mezzanine.pages.models import Page, RichTextPage

from .forms import PestReportForm,PublicationUrlFormSet,PublicationForm, PublicationFileFormSet, ReportingObligationForm, EventReportingForm, PestFreeAreaForm,\
ImplementationISPMForm,IssueKeywordsRelateForm,CommodityKeywordsRelateForm,EventreportingFileFormSet,ReportingoblicationFileFormSet,\
ImplementationISPMFileFormSet,PestFreeAreaFileFormSet, PestReportFileFormSet,WebsiteUrlFormSet,WebsiteForm, \
EventreportingUrlFormSet, ReportingObligationUrlFormSet ,PestFreeAreaUrlFormSet,ImplementationISPMUrlFormSet,PestReportUrlFormSet,\
CnPublicationUrlFormSet,CnPublicationForm, CnPublicationFileFormSet,\
PartnersPublicationUrlFormSet,PartnersPublicationForm, PartnersPublicationFileFormSet,PartnerPageForm,\
PollForm,Poll_ChoiceFormSet,\
PartnersNewsUrlFormSet,PartnersNewsForm, PartnersNewsFileFormSet,PartnersWebsiteUrlFormSet,PartnersWebsiteForm,\
EmailUtilityMessageForm,EmailUtilityMessageFileFormSet,MassEmailUtilityMessageForm,MassEmailUtilityMessageFileFormSet,\
CountryNewsUrlFormSet,CountryNewsForm, CountryNewsFileFormSet,NotificationMessageRelateForm,\
DraftProtocolForm,  DraftProtocolFileFormSet,DraftProtocolCommentsForm,IppcUserProfileForm,\
ContactUsEmailMessageForm,FAQsItemForm,FAQsCategoryForm,QAQuestionForm, QAAnswerForm,UserAutoRegistrationForm,IRSSActivityForm,IRSSActivityFileFormSet,\
UserMembershipHistoryForm,PhytosanitaryTreatmentForm,PhytosanitaryTreatmentPestsIdentityFormSet,PhytosanitaryTreatmentCommodityIdentityFormSet,\
CertificatesToolForm,WorkshopCertificatesToolForm, TopicForm ,TransTopicForm, TopicLeadsFormSet,TopicAssistantsFormSet,B_CertificatesToolForm  ,MyToolForm,MyTool2Form,NROStatsForm,\
ContributedResourceForm, ContributedResourceUrlFormSet, ContributedResourceFileFormSet, ContributedResourcePhotoFormSet
   
##TansReportingObligationForm , UserForm,

from schedule.models import Event, EventParticipants
from django.views.generic import ListView, MonthArchiveView, YearArchiveView, DetailView, TemplateView, CreateView
from django.core.urlresolvers import reverse
from django.core.mail import send_mail,send_mass_mail

from django.template.defaultfilters import slugify, lower
from django.template import RequestContext
from django.shortcuts import render_to_response, get_object_or_404, redirect

from django.contrib.auth.decorators import login_required, permission_required
from django.utils.decorators import method_decorator
from django.forms.models import inlineformset_factory
from django.contrib.contenttypes.generic import generic_inlineformset_factory 
from django.forms.formsets import formset_factory
from compiler.pyassem import order_blocks
import time
from django.http import HttpResponse
from datetime import datetime
from django.contrib.contenttypes.models import ContentType
from mezzanine.generic import views as myview
from mezzanine.generic import models
from t_eppo.models  import Names,Links

import os
import shutil

import zipfile
import StringIO
from settings import PROJECT_ROOT, PROJECT_DIRNAME, MEDIA_ROOT,DATABASES,ALLOWED_HOSTS
from django.core.files.storage import default_storage

import getpass, imaplib, email
from xml.dom import minidom

def get_profile():
    return IppcUserProfile.objects.all()
# def pest_report_country():
#     return PestReport.objects.all()



def reporting_trough_eppo(request):
    eppo_tmp_dir =MEDIA_ROOT+'/eppo_tmp'
    eppo_done_dir = MEDIA_ROOT+'/eppo_done'
    
    user_obj_reportEmail=User.objects.get(username='ippctest@gmail.com')
    password=   user_obj_reportEmail.password
    imap_server = imaplib.IMAP4_SSL("imap.gmail.com",993)
    imap_server.login("ippctest@gmail.com",password)
      
    imap_server.select("[Gmail]/All Mail") # here you a can choose a mail box like INBOX instead
    resp, items = imap_server.search(None, '(UNSEEN)') # you could filter using the IMAP rules here (check http://www.example-code.com/csharp/imap-search-critera.asp)
    items = items[0].split() # getting the mails id
 
    for emailid in items:
        resp, data = imap_server.fetch(emailid, "(RFC822)") # fetching the mail, "`(RFC822)`" means "get the whole stuff", but you can ask for headers only, etc
        email_body = data[0][1] # getting the mail content
        mailmsg = email.message_from_string(email_body) # parsing the mail content to get a mail object

        #Check if any attachments at all
        if mailmsg.get_content_maintype() != 'multipart':
            continue
        # we use walk to create a generator so we can iterate on the parts and forget about the recursive headach
        for part in mailmsg.walk():
            # multipart are just containers, so we skip them
            if part.get_content_maintype() == 'multipart':
                continue
            # is this part an attachment ?
            if part.get('Content-Disposition') is None:
                continue
            filename = part.get_filename()
            counter = 1
            # if there is no filename, we create one with a counter to avoid duplicates
            if not filename:
                filename = 'part-%03d%s' % (counter, 'bin')
                counter += 1
            att_path = os.path.join(eppo_tmp_dir, filename)
            #Check if its already there
            if not os.path.isfile(att_path) :
                # finally write the stuff
                fp = open(att_path, 'wb')
                fp.write(part.get_payload(decode=True))
                fp.close()
    xml_files = os.listdir(eppo_tmp_dir)
    pest_reports=[]
    #create lof file with report of data uploaded
    log_report =  open(os.path.join(eppo_done_dir, "eppo_reporting_"+timezone.now().strftime('%Y%m%d%H%M%S')+".log"), 'wb')
    log_report.write("List of uploaded pest report from Eppo:\n\n")
            
    for file_name in xml_files:
        if file_name.endswith('.xml'):
            title =''
            slug=''
            status =2
            author =1
            publish_date=''
            modify_date=''
            country = ''
            report_status = ''
            report_number = ''
            pest_status = ''
            pest_identity =''
            summary = ''
            hosts = ''
            geographical_distribution = ''
            nature_of_danger = ''
            contact_for_more_information = ''
           
            xml_file = open(os.path.join(eppo_tmp_dir, file_name),'r')
            xmldoc = minidom.parse(xml_file)
            xml_file.close()
           
            doc_element = xmldoc.documentElement
            reportIdentity  = doc_element.getElementsByTagName("ReportIdentity")[0]
            reportdata  = doc_element.getElementsByTagName("ReportData")[0]
           
            countryelement= reportIdentity.getElementsByTagName("CountryIdentity")[0]
            countryname=''
            countryslug=''
            if countryelement.hasAttribute("ISO3"):
                countryo= CountryPage.objects.filter(iso3=countryelement.getAttribute("ISO3"))
                if countryo:
                    if countryo[0].accepted_epporeport :
                        country =countryo[0].id 
                        countryname=countryo[0].name
                        countryslug=countryo[0].slug

                        title= reportIdentity.getElementsByTagName("Title")[0].childNodes[0].data  
                        report_status=0
                        slug = lower(slugify(title))

                        numberR=PestReport.objects.filter(country_id=country).count()
                        numberR=numberR+1
                        pestnumber=str(numberR)
                        if numberR<10 :
                            pestnumber='0'+pestnumber
                        report_number=countryelement.getAttribute("ISO3")+'-'+pestnumber+'/1'

                        if reportIdentity.hasAttribute("DateCreate"):
                            publish_date=reportIdentity.getAttribute("DateCreate")
                            modify_date=reportIdentity.getAttribute("DateCreate")
                        if reportIdentity.hasAttribute("UID"):
                            eppouid=reportIdentity.getAttribute("UID")
                        if reportIdentity.hasAttribute("numReport"):
                            epponumreport=reportIdentity.getAttribute("numReport")
                        if reportIdentity.hasAttribute("DateValidation"):
                            eppovalidationdate=reportIdentity.getAttribute("DateValidation")

                        eppoPublisher = reportIdentity.getElementsByTagName("Publisher")[0]
                        eppoPublishername = eppoPublisher.getElementsByTagName("fullname")[0].childNodes[0].data  
                        eppoPublisheremail = ''
                        if eppoPublisher.getElementsByTagName("email"):
                            eppoPublisheremail = eppoPublisher.getElementsByTagName("email")[0].childNodes[0].data  
                        if eppoPublisher.getElementsByTagName("date"):
                            eppoPublisherdate = eppoPublisher.getElementsByTagName("date")[0].childNodes[0].data  

                        pest_identity = reportIdentity.getElementsByTagName("EppoCode")[0].childNodes[0].data  
                        pestidentityFinal=Names.objects.get(eppocode=pest_identity, preferred=1)

                        if reportdata.getElementsByTagName("GeogDistrib")[0].childNodes:
                            geographical_distribution= reportdata.getElementsByTagName("GeogDistrib")[0].childNodes[0].data  
                        if reportdata.getElementsByTagName("Context")[0].childNodes:
                            summary= reportdata.getElementsByTagName("Context")[0].childNodes[0].data  

                        hosts= reportdata.getElementsByTagName("HostName")[0].childNodes[0].data  
                        peststatuselement = reportdata.getElementsByTagName("PestStatus")[0]
                        pest_status_label=''
                        if peststatuselement.getElementsByTagName("libelle")[0].childNodes:
                            pest_status_label = peststatuselement.getElementsByTagName("libelle")[0].childNodes[0].data  
                        if pest_status_label!='':
                            ps=PestStatus.objects.get(status=pest_status_label)
                            pest_status=ps.id
                        else:    
                            ps=PestStatus.objects.get(status='Other')
                            pest_status=ps.id#Other

                        new_pest_report = PestReport()
                        new_pest_report.country_id=country
                        new_pest_report.title=title
                        new_pest_report.publish_date=  publish_date
                        new_pest_report.country_id=country
                        new_pest_report.report_number=report_number
                        new_pest_report.pest_identity_id=pestidentityFinal.id
                        new_pest_report.geographical_distribution=geographical_distribution
                        
                        safe_str = summary.encode('ascii', 'ignore')
                        new_pest_report.summary=str(safe_str.encode('utf-8'))#TO DO: #problme with encoded summary
                        new_pest_report.author_id=1
                        new_pest_report.hosts=hosts
                        new_pest_report.importedfromeppo = True
                        new_pest_report.save()
                        
                        new_pest_report.pest_status.add(ps)
                        pest_reports.append(new_pest_report)

                        #move xml processed in 'eppo_done' dir
                        os.rename(os.path.join(eppo_tmp_dir, file_name),os.path.join(eppo_done_dir, timezone.now().strftime('%Y%m%d%H%M%S')+'_'+file_name))
                        #create log and email messages notifications
                        year=new_pest_report.publish_date.strftime("%Y")
                        month=new_pest_report.publish_date.strftime("%m")
                        pest_url="https://www.ippc.int/en/"+countryslug+"/pestreports/"+year+"/"+month+"/"+slug

                        log_report.write("["+ timezone.now().strftime('%Y%m%d%H%M%S')+"] "+countryname+" ["+report_number+"] '"+title+" "+pest_url+"\n\n")

                        msgtpeppo="Dear EPPO,<br><br>the Pest report<br><br><strong>UID</strong>: "+str(eppouid) +" <br><strong>Numreport:</strong> "+str(epponumreport)+" <br><strong>Publish date:</strong>"+str(eppoPublisherdate)+"<br><br>has been successefully uploaded in the IPPC website<br><br><Strong>URL</strong>: "+pest_url+""
                        msgtoCP="Dear "+str(eppoPublishername)+",<br><hr><br>the Pest report published in EPPO with:<br><br><strong>UID</strong>: "+str(eppouid) +" <br><strong>Numreport:</strong> "+str(epponumreport)+" <br><strong>Publish date:</strong>"+str(eppoPublisherdate)+"<br><br>has been successefully uploaded in the IPPC website<br><br><Strong>URL</strong>: "+pest_url+""

                        subject='EPPO Pest Report successefully uploaded in IPPC'  
                        #TO DO: #SEND TO roy@eppo.int
                        notifificationmessageeppo = mail.EmailMessage(subject,msgtpeppo,'ippc@fao.org', ['paola.sentinelli@fao.org'], ['paola.sentinelli@fao.org'])
                        notifificationmessageeppo.content_subtype = "html"
                        #print('test-sending')
                        sent =notifificationmessageeppo.send()
                        if eppoPublisheremail:
                                notifificationmessageCp = mail.EmailMessage(subject,msgtoCP,'ippc@fao.org', ['paola.sentinelli@fao.org'], ['paola.sentinelli@fao.org'])
                                notifificationmessageCp.content_subtype = "html"
                                #print('test-sending')
                                sent =notifificationmessageCp.send()
                                
                    else:
                        #move xml processed in 'eppo_done' dir & create LOG
                        os.rename(os.path.join(eppo_tmp_dir, file_name),os.path.join(eppo_done_dir, timezone.now().strftime('%Y%m%d%H%M%S')+'_'+file_name))
                        log_report.write("["+ timezone.now().strftime('%Y%m%d%H%M%S')+"] "+file_name+" [NOT IMPORTED, NOT ACCETTED TO REPORT TROUH EPPO]\n\n")

    log_report.close()        
#    context = {"pest_reports":pest_reports,}
    context = {}
    #TO DO:        
    #SET COUNTIES in Country page 'allow eppo to report automatically'
    #SET proper email ROY and Publisher, send real emails'
    response = render(request, "countries/eppo_reporting.html", context)
    return response

#---------------------NEW
from email.parser import Parser

def process_multipart_message(message):
    rtn = ''
    if message.is_multipart():
        for m in message.get_payload():
            rtn += process_multipart_message(m)
    else:
        rtn += message.get_payload()
    return rtn

def reporting_trough_eppo1(request):
    eppo_tmp_dir =MEDIA_ROOT+'/eppo_tmp'
    eppo_done_dir = MEDIA_ROOT+'/eppo_done'
    user_obj_reportEmail=User.objects.get(username='ippctest@gmail.com')
    password=   user_obj_reportEmail.password
   
    username = 'paola.sentinelli@fao.org'
    password = 'Poldina69'
    conn = imaplib.IMAP4_SSL("outlook.office365.com",993)
    conn.login(username, password)
   # print('COOOOOOOONNNNNN')
    conn.select('INBOX')

    resp, items = conn.search(None, '(UNSEEN)')#conn.uid('search', '(UNSEEN)', 'ALL')
    
  #  print(items)
    uids = items[0].split()
    
    for emailid in uids:
 
        resp, data = conn.fetch(emailid, "(RFC822)") # fetching the mail, "`(RFC822)`" means "get the whole stuff", but you can ask for headers only, etc
     
        email_body = data[0][1] # getting the mail content
        mailmsg = email.message_from_string(email_body) # parsing the mail content to get a mail object
        
        #Check if any attachments at all
        if mailmsg.get_content_maintype() != 'multipart':
            continue
        # we use walk to create a generator so we can iterate on the parts and forget about the recursive headach
        for part in mailmsg.walk():
            # multipart are just containers, so we skip them
            if part.get_content_maintype() == 'multipart':
                continue
            # is this part an attachment ?
            if part.get('Content-Disposition') is None:
                continue
            filename = part.get_filename()
            counter = 1
            # if there is no filename, we create one with a counter to avoid duplicates
            if not filename:
                filename = 'part-%03d%s' % (counter, 'bin')
                counter += 1
            #print(filename)
            #print(filename.endswith('.xml'))
            # print(mailmsg.get('From'))
            if filename.endswith('.xml') and  mailmsg.get('From').endswith('ippctest@gmail.com>'):#pestreporting@eppo.int
            
                att_path = os.path.join(eppo_tmp_dir, filename)
                #Check if its already there
                if not os.path.isfile(att_path) :
                    # finally write the stuff
                    fp = open(att_path, 'wb')
                    fp.write(part.get_payload(decode=True))
                    fp.close()
        xml_files = os.listdir(eppo_tmp_dir)
        pest_reports=[]
        #create lof file with report of data uploaded
        log_report =  open(os.path.join(eppo_done_dir, "eppo_reporting_"+timezone.now().strftime('%Y%m%d%H%M%S')+".log"), 'wb')
        log_report.write("List of uploaded pest report from Eppo:\n\n")
    #           
        for file_name in xml_files:
          if file_name.endswith('.xml'):
              title =''
              slug=''
              status =2
              author =1
              publish_date=''
              modify_date=''
              country = ''
              report_status = ''
              report_number = ''
              pest_status = ''
              pest_identity =''
              summary = ''
              hosts = ''
              geographical_distribution = ''
              nature_of_danger = ''
              contact_for_more_information = ''

              xml_file = open(os.path.join(eppo_tmp_dir, file_name),'r')
              xmldoc = minidom.parse(xml_file)
              xml_file.close()

              doc_element = xmldoc.documentElement
              reportIdentity  = doc_element.getElementsByTagName("ReportIdentity")[0]
              reportdata  = doc_element.getElementsByTagName("ReportData")[0]

              countryelement= reportIdentity.getElementsByTagName("CountryIdentity")[0]
              countryname=''
              countryslug=''
              if countryelement.hasAttribute("ISO3"):
                  iso31=countryelement.getAttribute("ISO3")
                  if iso31 == 'ZZZ':
                      iso31='ITA'
                  countryo= CountryPage.objects.filter(iso3=iso31)
                  if countryo:
                      if True:#countryo[0].accepted_epporeport :
                          country =countryo[0].id 
                          countryname=countryo[0].name
                          countryslug=countryo[0].slug
    #
                          title= reportIdentity.getElementsByTagName("Title")[0].childNodes[0].data  
                          report_status=0
                          slug = lower(slugify(title))
    #
                          numberR=PestReport.objects.filter(country_id=country).count()
                          numberR=numberR+1
                          pestnumber=str(numberR)
                          if numberR<10 :
                              pestnumber='0'+pestnumber
                          report_number=countryelement.getAttribute("ISO3")+'-'+pestnumber+'/1'
    #
                          if reportIdentity.hasAttribute("DateCreate"):
                              publish_date=reportIdentity.getAttribute("DateCreate")
                              modify_date=reportIdentity.getAttribute("DateCreate")
                          if reportIdentity.hasAttribute("UID"):
                              eppouid=reportIdentity.getAttribute("UID")
                          if reportIdentity.hasAttribute("numReport"):
                              epponumreport=reportIdentity.getAttribute("numReport")
                          if reportIdentity.hasAttribute("DateValidation"):
                              eppovalidationdate=reportIdentity.getAttribute("DateValidation")
    #
                          eppoPublisher = reportIdentity.getElementsByTagName("Publisher")[0]
                          eppoPublishername = eppoPublisher.getElementsByTagName("fullname")[0].childNodes[0].data  
                          eppoPublisheremail = ''
                          if eppoPublisher.getElementsByTagName("email"):
                              eppoPublisheremail = eppoPublisher.getElementsByTagName("email")[0].childNodes[0].data  
                          if eppoPublisher.getElementsByTagName("date"):
                              eppoPublisherdate = eppoPublisher.getElementsByTagName("date")[0].childNodes[0].data  
    #
                          pest_identity = reportIdentity.getElementsByTagName("EppoCode")[0].childNodes[0].data  
                          pestidentityFinal=Names.objects.get(eppocode=pest_identity, preferred=1)
    #
                          if reportdata.getElementsByTagName("GeogDistrib")[0].childNodes:
                              geographical_distribution= reportdata.getElementsByTagName("GeogDistrib")[0].childNodes[0].data  
                          if reportdata.getElementsByTagName("Context")[0].childNodes:
                              summary= reportdata.getElementsByTagName("Context")[0].childNodes[0].data  
    #
                          hosts= reportdata.getElementsByTagName("HostName")[0].childNodes[0].data  
                          peststatuselement = reportdata.getElementsByTagName("PestStatus")[0]
                          pest_status_label=''
                          if peststatuselement.getElementsByTagName("libelle")[0].childNodes:
                              pest_status_label = peststatuselement.getElementsByTagName("libelle")[0].childNodes[0].data  
                          if pest_status_label!='':
                              ps=PestStatus.objects.get(status=pest_status_label)
                              pest_status=ps.id
                          else:    
                              ps=PestStatus.objects.get(status='Other')
                              pest_status=ps.id#Other
    #
                          new_pest_report = PestReport()
                          new_pest_report.country_id=country
                          new_pest_report.title=title
                          new_pest_report.publish_date=  publish_date
                          new_pest_report.country_id=country
                          new_pest_report.report_number=report_number
                          new_pest_report.pest_identity_id=pestidentityFinal.id
                          new_pest_report.geographical_distribution=geographical_distribution

                          safe_str = summary.encode('ascii', 'ignore')
                          new_pest_report.summary=str(safe_str.encode('utf-8'))#TO DO: #problme with encoded summary
                          new_pest_report.author_id=1
                          new_pest_report.hosts=hosts
                          new_pest_report.importedfromeppo = True
                          new_pest_report.save()

                          new_pest_report.pest_status.add(ps)
                          pest_reports.append(new_pest_report)
    #
                          #move xml processed in 'eppo_done' dir
                          os.rename(os.path.join(eppo_tmp_dir, file_name),os.path.join(eppo_done_dir, timezone.now().strftime('%Y%m%d%H%M%S')+'_'+file_name))
                          #create log and email messages notifications
                          year=new_pest_report.publish_date.strftime("%Y")
                          month=new_pest_report.publish_date.strftime("%m")
                          pest_url="https://www.ippc.int/en/"+countryslug+"/pestreports/"+year+"/"+month+"/"+slug
    #
                          log_report.write("["+ timezone.now().strftime('%Y%m%d%H%M%S')+"] "+countryname+" ["+report_number+"] '"+title+" "+pest_url+"\n\n")
    #
                          msgtpeppo="Dear EPPO,<br><br>the Pest report<br><br><strong>UID</strong>: "+str(eppouid) +" <br><strong>Numreport:</strong> "+str(epponumreport)+" <br><strong>Publish date:</strong>"+str(eppoPublisherdate)+"<br><br>has been successefully uploaded in the IPPC website<br><br><Strong>URL</strong>: "+pest_url+""
                          msgtoCP="Dear "+str(eppoPublishername)+",<br><hr><br>the Pest report published in EPPO with:<br><br><strong>UID</strong>: "+str(eppouid) +" <br><strong>Numreport:</strong> "+str(epponumreport)+" <br><strong>Publish date:</strong>"+str(eppoPublisherdate)+"<br><br>has been successefully uploaded in the IPPC website<br><br><Strong>URL</strong>: "+pest_url+""
    #
                          subject='EPPO Pest Report successefully uploaded in IPPC'  
                          #TO DO: #SEND TO roy@eppo.int
                          notifificationmessageeppo = mail.EmailMessage(subject,msgtpeppo,'ippc@fao.org', ['paola.sentinelli@fao.org'], ['paola.sentinelli@fao.org'])
                          notifificationmessageeppo.content_subtype = "html"
                          #print('test-sending')
                          sent =notifificationmessageeppo.send()
                          if eppoPublisheremail:
                              notifificationmessageCp = mail.EmailMessage(subject,msgtoCP,'ippc@fao.org', ['paola.sentinelli@fao.org'], ['paola.sentinelli@fao.org'])
                              notifificationmessageCp.content_subtype = "html"
                              #print('test-sending')
                              sent =notifificationmessageCp.send()
                      else:
                          #move xml processed in 'eppo_done' dir & create LOG
                          os.rename(os.path.join(eppo_tmp_dir, file_name),os.path.join(eppo_done_dir, timezone.now().strftime('%Y%m%d%H%M%S')+'_'+file_name))
                          log_report.write("["+ timezone.now().strftime('%Y%m%d%H%M%S')+"] "+file_name+" [NOT IMPORTED, NOT ACCETTED TO REPORT TROUH EPPO]\n\n")
#
    log_report.close()        
    #  context = {"pest_reports":pest_reports,}
    context = {}
      #TO DO:        
      #SET COUNTIES in Country page 'allow eppo to report automatically'
      #SET proper email ROY and Publisher, send real emails'
    response = render(request, "countries/eppo_reporting.html", context)
    return response



#EPPO REP import imaplib
#EPPO REP from email.parser import Parser

#EPPO REP def process_multipart_message(message):
#    rtn = ''
#    if message.is_multipart():
#        for m in message.get_payload():
#            rtn += process_multipart_message(m)
#    else:
#        rtn += message.get_payload()
#    return rtn

#def reporting_trough_eppo(request):
#    print('aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa')
#    eppo_tmp_dir =MEDIA_ROOT+'/eppo_tmp'
#    eppo_done_dir = MEDIA_ROOT+'/eppo_done'
#    user_obj_reportEmail=User.objects.get(username='ippctest@gmail.com')
#    password=   user_obj_reportEmail.password
#    #imap_server = imaplib.IMAP4_SSL("imap.gmail.com",993)
#    #imap_server.login("ippctest@gmail.com",password)
#      
#    #imap_server.select("[Gmail]/All Mail") # here you a can choose a mail box like INBOX instead
#    #resp, items = imap_server.search(None, '(UNSEEN)') # you could filter using the IMAP rules here (check http://www.example-code.com/csharp/imap-search-critera.asp)
#    #items = items[0].split() # getting the mails id
# 
# 
# 
#
#    url = "faohqmail.fao.org"
#    conn = imaplib.IMAP4_SSL(url,993)
#    user,password = ("Sentinelli","Poldana57")
#    conn.login(user,password)
#    conn.select('INBOX')
#    #results,data = conn.search(None,'ALL')
#    resp,items = conn.search(None,'(UNSEEN)')
#    #msg_ids = data[0]
#    #msg_id_list = msg_ids.split()
#
#    #latest_email_id = msg_id_list[-1]
#    #result,data = conn.fetch(latest_email_id,"(RFC822)")
#   # resp,items = conn.fetch(latest_email_id,"(RFC822)")
#  #  raw_email = data[0][1]
#    items = items[0].split() # getting the mails id
# 
#
##    p = Parser()
##    msg = p.parsestr(raw_email)
##
##    msg.get('From')
##    msg.get('Subject')
##    print( msg.get('From'))
##    print( msg.get('Subject'))
##    msg.get_payload()
##    print( msg.get_payload())
##  
##    msg_contant = process_multipart_message(msg)
##    print(msg_contant)
#  
#    for emailid in items:
#        resp, data = conn.fetch(emailid, "(RFC822)") # fetching the mail, "`(RFC822)`" means "get the whole stuff", but you can ask for headers only, etc
#        email_body = data[0][1] # getting the mail content
#        mailmsg = email.message_from_string(email_body) # parsing the mail content to get a mail object
##
#        #Check if any attachments at all
#        if mailmsg.get_content_maintype() != 'multipart':
#            continue
#        # we use walk to create a generator so we can iterate on the parts and forget about the recursive headach
#        for part in mailmsg.walk():
#            # multipart are just containers, so we skip them
#            if part.get_content_maintype() == 'multipart':
#                continue
#            # is this part an attachment ?
#            if part.get('Content-Disposition') is None:
#                continue
#            filename = part.get_filename()
#            counter = 1
#            # if there is no filename, we create one with a counter to avoid duplicates
#            if not filename:
#                filename = 'part-%03d%s' % (counter, 'bin')
#                counter += 1
#            #print(filename)
#            #print(filename.endswith('.xml'))
#            # print(mailmsg.get('From'))
#            if filename.endswith('.xml') and  mailmsg.get('From').endswith('paola.sentinelli@gmail.com>'):#pestreporting@eppo.int
#                att_path = os.path.join(eppo_tmp_dir, filename)
#                #Check if its already there
#                if not os.path.isfile(att_path) :
#                    # finally write the stuff
#                    fp = open(att_path, 'wb')
#                    fp.write(part.get_payload(decode=True))
#                    fp.close()
#    xml_files = os.listdir(eppo_tmp_dir)
#    pest_reports=[]
#    #create lof file with report of data uploaded
#    log_report =  open(os.path.join(eppo_done_dir, "eppo_reporting_"+timezone.now().strftime('%Y%m%d%H%M%S')+".log"), 'wb')
#    log_report.write("List of uploaded pest report from Eppo:\n\n")
#            
#    for file_name in xml_files:
#        if file_name.endswith('.xml'):
#            title =''
#            slug=''
#            status =2
#            author =1
#            publish_date=''
#            modify_date=''
#            country = ''
#            report_status = ''
#            report_number = ''
#            pest_status = ''
#            pest_identity =''
#            summary = ''
#            hosts = ''
#            geographical_distribution = ''
#            nature_of_danger = ''
#            contact_for_more_information = ''
#           
#            xml_file = open(os.path.join(eppo_tmp_dir, file_name),'r')
#            xmldoc = minidom.parse(xml_file)
#            xml_file.close()
#           
#            doc_element = xmldoc.documentElement
#            reportIdentity  = doc_element.getElementsByTagName("ReportIdentity")[0]
#            reportdata  = doc_element.getElementsByTagName("ReportData")[0]
#           
#            countryelement= reportIdentity.getElementsByTagName("CountryIdentity")[0]
#            countryname=''
#            countryslug=''
#            if countryelement.hasAttribute("ISO3"):
#                countryo= CountryPage.objects.filter(iso3=countryelement.getAttribute("ISO3"))
#                if countryo:
#                    if countryo[0].accepted_epporeport :
#                        country =countryo[0].id 
#                        countryname=countryo[0].name
#                        countryslug=countryo[0].slug
#
#                        title= reportIdentity.getElementsByTagName("Title")[0].childNodes[0].data  
#                        report_status=0
#                        slug = lower(slugify(title))
#
#                        numberR=PestReport.objects.filter(country_id=country).count()
#                        numberR=numberR+1
#                        pestnumber=str(numberR)
#                        if numberR<10 :
#                            pestnumber='0'+pestnumber
#                        report_number=countryelement.getAttribute("ISO3")+'-'+pestnumber+'/1'
#
#                        if reportIdentity.hasAttribute("DateCreate"):
#                            publish_date=reportIdentity.getAttribute("DateCreate")
#                            modify_date=reportIdentity.getAttribute("DateCreate")
#                        if reportIdentity.hasAttribute("UID"):
#                            eppouid=reportIdentity.getAttribute("UID")
#                        if reportIdentity.hasAttribute("numReport"):
#                            epponumreport=reportIdentity.getAttribute("numReport")
#                        if reportIdentity.hasAttribute("DateValidation"):
#                            eppovalidationdate=reportIdentity.getAttribute("DateValidation")
#
#                        eppoPublisher = reportIdentity.getElementsByTagName("Publisher")[0]
#                        eppoPublishername = eppoPublisher.getElementsByTagName("fullname")[0].childNodes[0].data  
#                        eppoPublisheremail = ''
#                        if eppoPublisher.getElementsByTagName("email"):
#                            eppoPublisheremail = eppoPublisher.getElementsByTagName("email")[0].childNodes[0].data  
#                        if eppoPublisher.getElementsByTagName("date"):
#                            eppoPublisherdate = eppoPublisher.getElementsByTagName("date")[0].childNodes[0].data  
#
#                        pest_identity = reportIdentity.getElementsByTagName("EppoCode")[0].childNodes[0].data  
#                        pestidentityFinal=Names.objects.get(eppocode=pest_identity, preferred=1)
#
#                        if reportdata.getElementsByTagName("GeogDistrib")[0].childNodes:
#                            geographical_distribution= reportdata.getElementsByTagName("GeogDistrib")[0].childNodes[0].data  
#                        if reportdata.getElementsByTagName("Context")[0].childNodes:
#                            summary= reportdata.getElementsByTagName("Context")[0].childNodes[0].data  
#
#                        hosts= reportdata.getElementsByTagName("HostName")[0].childNodes[0].data  
#                        peststatuselement = reportdata.getElementsByTagName("PestStatus")[0]
#                        pest_status_label=''
#                        if peststatuselement.getElementsByTagName("libelle")[0].childNodes:
#                            pest_status_label = peststatuselement.getElementsByTagName("libelle")[0].childNodes[0].data  
#                        if pest_status_label!='':
#                            ps=PestStatus.objects.get(status=pest_status_label)
#                            pest_status=ps.id
#                        else:    
#                            ps=PestStatus.objects.get(status='Other')
#                            pest_status=ps.id#Other
#
#                        new_pest_report = PestReport()
#                        new_pest_report.country_id=country
#                        new_pest_report.title=title
#                        new_pest_report.publish_date=  publish_date
#                        new_pest_report.country_id=country
#                        new_pest_report.report_number=report_number
#                        new_pest_report.pest_identity_id=pestidentityFinal.id
#                        new_pest_report.geographical_distribution=geographical_distribution
#                        
#                        safe_str = summary.encode('ascii', 'ignore')
#                        new_pest_report.summary=str(safe_str.encode('utf-8'))#TO DO: #problme with encoded summary
#                        new_pest_report.author_id=1
#                        new_pest_report.hosts=hosts
#                        new_pest_report.importedfromeppo = True
#                        new_pest_report.save()
#                        
#                        new_pest_report.pest_status.add(ps)
#                        pest_reports.append(new_pest_report)
#
#                        #move xml processed in 'eppo_done' dir
#                        os.rename(os.path.join(eppo_tmp_dir, file_name),os.path.join(eppo_done_dir, timezone.now().strftime('%Y%m%d%H%M%S')+'_'+file_name))
#                        #create log and email messages notifications
#                        year=new_pest_report.publish_date.strftime("%Y")
#                        month=new_pest_report.publish_date.strftime("%m")
#                        pest_url="https://www.ippc.int/en/"+countryslug+"/pestreports/"+year+"/"+month+"/"+slug
#
#                        log_report.write("["+ timezone.now().strftime('%Y%m%d%H%M%S')+"] "+countryname+" ["+report_number+"] '"+title+" "+pest_url+"\n\n")
#
#                        msgtpeppo="Dear EPPO,<br><br>the Pest report<br><br><strong>UID</strong>: "+str(eppouid) +" <br><strong>Numreport:</strong> "+str(epponumreport)+" <br><strong>Publish date:</strong>"+str(eppoPublisherdate)+"<br><br>has been successefully uploaded in the IPPC website<br><br><Strong>URL</strong>: "+pest_url+""
#                        msgtoCP="Dear "+str(eppoPublishername)+",<br><hr><br>the Pest report published in EPPO with:<br><br><strong>UID</strong>: "+str(eppouid) +" <br><strong>Numreport:</strong> "+str(epponumreport)+" <br><strong>Publish date:</strong>"+str(eppoPublisherdate)+"<br><br>has been successefully uploaded in the IPPC website<br><br><Strong>URL</strong>: "+pest_url+""
#
#                        subject='EPPO Pest Report successefully uploaded in IPPC'  
#                        #TO DO: #SEND TO roy@eppo.int
#                        notifificationmessageeppo = mail.EmailMessage(subject,msgtpeppo,'ippc@fao.org', ['paola.sentinelli@fao.org'], ['paola.sentinelli@fao.org'])
#                        notifificationmessageeppo.content_subtype = "html"
#                        if eppoPublisheremail:
#                            notifificationmessageCp = mail.EmailMessage(subject,msgtoCP,'ippc@fao.org', ['paola.sentinelli@fao.org'], ['paola.sentinelli@fao.org'])
#                            notifificationmessageCp.content_subtype = "html"
#                    else:
#                        #move xml processed in 'eppo_done' dir & create LOG
#                        os.rename(os.path.join(eppo_tmp_dir, file_name),os.path.join(eppo_done_dir, timezone.now().strftime('%Y%m%d%H%M%S')+'_'+file_name))
#                        log_report.write("["+ timezone.now().strftime('%Y%m%d%H%M%S')+"] "+file_name+" [NOT IMPORTED, NOT ACCETTED TO REPORT TROUH EPPO]\n\n")
#
#    log_report.close()        
##    context = {"pest_reports":pest_reports,}
#    context = {}
#    #TO DO:        
#    #SET COUNTIES in Country page 'allow eppo to report automatically'
#    #SET proper email ROY and Publisher, send real emails'
#    response = render(request, "countries/eppo_reporting.html", context)
#    return response


import datetime as dt
import urllib2
import socket
def reminder_getlink(cn,type,obj):
    return "https://www.ippc.int/countries/"+cn+"/"+type+"/"+str(obj.publish_date.strftime("%Y"))+'/'+str(obj.publish_date.strftime("%m"))+'/'+obj.slug+'/'
    
def reminder_getemptylink(cn,type,rep):
    return "https://www.ippc.int/countries/"+cn+"/"+type+"/"+rep
         
def reminder_to_cn(request,id=None):
    # OCP 3 months: 90 days
    # reporting obligation 12 months
    # Pestreports and Emergency Actions 6 months
    # LINKS/Files 3 months
    
    BASIC_REP_TYPE_CHOICES_LABELS = (
        (1, ("Description of the NPPO")), 
        (2, ("Entry Points")),
        (3, ("List of Regulated Pests")),
        (4, ("Legislation: Phytosanitary Requirements/Restrictions/Prohibitions")),
    )
    EVT_REP_TYPE_CHOICES_LABELS = (
        (1, ("Emergency Actions")), 
        #(2, ("Non-compliance")),
        #(3, ("Organizational Arrangements of Plant Protection")),
        #(4, ("Pest status")),
        #(5, ("Rationale for Phytosanitary Requirements")),
    )
    GENDER_CHOICES = (
        (1, ("Mr.")),
        (2, ("Ms.")),
        (3, ("Mrs.")),
        (4, ("Professor.")),
        (5, ("M.")),
        (6, ("Mme.")),
        (7, ("Dr.")),
        (8, ("Sr.")),
        (9, ("Sra.")),
    )
    ocp_range=dt.timedelta(days=90)
    ro_range=dt.timedelta(days=365)
    pestreport_range=dt.timedelta(days=182)
    ev_range=dt.timedelta(days=182)
    
    reminder_done_dir = MEDIA_ROOT+'/reminder_done'
  
    textmessages=[]
  
    date_day=28
    date12month=[11]
    date12monthEditors=[5]
    date6month=[2,8]
    date3month=[1,4,7,10]
    date3monthFiles=[3,6,9,12]
    
#    reminder_log_report =  open(os.path.join(reminder_done_dir, "reminder_done_"+timezone.now().strftime('%Y%m%d%H%M%S')+".log"), 'wb')
 #   reminder_log_report.write("List of sent reminder to NPPOs:\n\n")
    
  
    id_range_min=0
    id_range_max=0
    
    id_range_min1=4
    id_range_min2=41
    id_range_min3=77
    id_range_min4=117
    id_range_min5=153
    id_range_min6=189
    id_range_min7=228
    
    id_range_max1=40
    id_range_max2=76
    id_range_max3=116
    id_range_max4=152
    id_range_max5=187
    id_range_max6=226
    id_range_max7=1002
    
    
    if id == '1':
        id_range_min=id_range_min1
        id_range_max=id_range_max1
    elif id == '2':
        id_range_min=id_range_min2
        id_range_max=id_range_max2
    elif id == '3':
            id_range_min=id_range_min3
            id_range_max=id_range_max3
    elif id == '4':
            id_range_min=id_range_min4
            id_range_max=id_range_max4
    elif id == '5':
            id_range_min=id_range_min5
            id_range_max=id_range_max5
    elif id == '6':
            id_range_min=id_range_min6
            id_range_max=id_range_max6
    elif id == '7':
            id_range_min=id_range_min7
            id_range_max=id_range_max7
    aa=''
    if timezone.now().day == date_day:
        db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
        cursor = db.cursor()
       # reminder_log_report.write("date_day="+ str(date_day)+" \n\n")
       # reminder_log_report.write("id cron="+ str(id)+" \n\n")
      #  reminder_log_report.write("id_range_min="+ str(id_range_min)+", id_range_max="+str(id_range_max)+" \n\n")
              
        countriesList=CountryPage.objects.filter().exclude(id='-1')
        for cn in countriesList :
	 #   reminder_log_report.write("cn="+lower(slugify(cn))+" \n\n")
            if cn.send_reminder and cn.id >=id_range_min and cn.id <=id_range_max :          
		aa+=lower(slugify(cn)) +'<br>'
                country_slug = lower(slugify(cn))   
              #  reminder_log_report.write("["+ timezone.now().strftime('%Y%m%d%H%M%S')+" - send reminder] "+country_slug+" \n\n")
                
                countryo = get_object_or_404(CountryPage, page_ptr_id=cn.id)
                cp = countryo.contact_point_id
                editors = countryo.editors
                emails=[]
                emailstest=[]
                emailstest.append('paola.sentinelli@gmail.com')
                cp_email=''
                e_email=''
                if cp != None:
                    user_obj=User.objects.get(id=cp)
                    cp_email=user_obj.email
                    emails.append(cp_email)
              
                #12 months reportingobligations
                if timezone.now().month in date12month :
                    reportingobligation_to_notify=''
                    ro_to_notify=''
                    
                    for i in range(1,5):
                        ro=0
                        reportingobligations = ReportingObligation.objects.filter(country_id=cn.id,reporting_obligation_type=i, is_version=False)
                        for r in reportingobligations:
                            if timezone.now()- r.modify_date > ro_range :
                                sql = "UPDATE ippc_reportingobligation set to_verify=True where id="+str(r.id)
                                try:
                                    cursor.execute(sql)
                                    db.commit()
                                except:
                                    db.rollback()
                                ro = ro + 1
                                if ro == 1:
                                    ro_to_notify=ro_to_notify+'<tr><td><b>'+str(dict(BASIC_REP_TYPE_CHOICES_LABELS)[i])+'</b></td><td><strong>Last Updated</strong></td></tr>'
                                ro_to_notify=ro_to_notify+'<tr><td><a href="'+reminder_getlink(country_slug,'reportingobligation',r)+'">'+r.title+'</a></td><td>'+str(r.modify_date.strftime("%d-%m-%Y"))+'</td></tr>'
                        if ro == 0:
                            ro_to_notify=ro_to_notify+'<tr><td><b>'+str(dict(BASIC_REP_TYPE_CHOICES_LABELS)[i])+'</b></td><td><strong>Last Updated</strong></td></tr>'
                            ro_to_notify=ro_to_notify+'<tr><td><a href="'+reminder_getemptylink(country_slug,'reportingobligation',str(i))+'">link to empty folder</a></td><td>empty</td></tr>'
                    
                    textmessage=''
                   
                    if ro_to_notify!='':
                        textmessage ='<table bgcolor="#FFFFFF" cellspacing="2" cellpadding="2" valign="top" width="100%" border=1 style="border: 1px solid #10501F;><tr><td width="100%" bgcolor="#FFFFFF" colspan=2><p>Dear Sir/Madam,<br><br>This is an automatic reminder sent every 12 months for National Reporting Obligations: Description of the NPPO, Entry points, List of Regulated Pests, Legislation:  Phytosanitary Requirements/Restrictions/Prohibitions.<br><br>Some reports have not been updated by your country in the <b><a href="https://www.ippc.int/countries/'+country_slug+'">International Phytosanitary Portal</a></b> for the last 12 months or more. Please check this information by following links in the table below and update when needed.<br><br>In the table below, you can also find the NRO categories which have never had any information uploaded so far. If that information is available, please upload it as soon as possible.<br><br><u><a href="http://www.ippc.int/en/accounts/login">You will need to log in to the IPP to confirm and/or update the reports.</a></u><br><br></td></tr>'
                        reportingobligation_to_notify='<!--tr><td><strong>NRO Category</strong></td><td><strong>Report Title</strong></td><td><strong>Last Updated</strong></td></tr-->'+ro_to_notify
                        textmessage = textmessage +reportingobligation_to_notify +'</table>'
                  
       
                        textmessage=textmessage+'<br>The following resources are available to assist you:<br><ol><li> The manual on editing on the IPP [<a href="https://www.ippc.int/en/publications/80405/">Guide to the IPP</a>]</li><li> <a href="https://www.ippc.int/en/faq/">FAQs</a> including forgotten password.</li></ol><br>Should you require any assistance, please contact IPPC Secretariat at <a href="mailto:IPPC-IT@fao.org">IPPC-IT@fao.org</a>.<br><br>Best regards<br><br>IPPC Secretariat'
                        remider_message= ReminderMessage()
                        remider_message.pk = None
                        remider_message.emailfrom = "ippc@fao.org"
                        remider_message.emailto = emails
                        remider_message.subject = "IPPC NRO reminder for "+str(cn)+": Description of NPPO, Entry points, List of Regulated Pests, Legislation"
                        remider_message.messagebody = textmessage.encode('utf-8')
                        remider_message.date = timezone.now()
   
                        messages=[]
                        message = mail.EmailMessage(remider_message.subject,remider_message.messagebody,remider_message.emailfrom,
                        emails, ['paola.sentinelli@fao.org'])
                        message.content_subtype = "html"
                        messages.append(message)
                            
                        connection = mail.get_connection()
                        connection.open()
                        #print('test-sending')
                        sent=connection.send_messages(messages)#
                        connection.close()
                        remider_message.sent = True
                        remider_message.save()
                        
                #CN editors        
                if timezone.now().month in date12monthEditors :
                    textmessage =''
                    editors_text=''
                    
                    for e in editors.all():
                        user_obj=User.objects.get(id=e.id)
                        editors_text +='<li><b>'+user_obj.first_name+' '+user_obj.last_name+'</b> ('+user_obj.email+')</li> '
                    
                    if editors_text !='':
                        textmessage+='<p>Dear Sir/Madam,<br><br>This is a routine reminder sent every 12 months in an attempt to ensure that the IPP editors remain accurate and active.<br><br>The IPPC team would like to ask you to <b>check the Contact details</b> of your <b>IPP editors</b>:<br><br>'
                        textmessage+= '<ul>'+editors_text+'</ul>'
                        textmessage+= '<br><br>In case any detail has changed please kindly inform us [<a href="mailto:IPPC-IT@fao.org">IPPC-IT@fao.org</a>]. <br><br>In case you would like to nominate a new editor provide us with a completed form. You can find the form <a href="https://www.ippc.int/en/publications/ipp-editor-nomination-request-nppos/">here</a>. <br><br><br>Best regards<br><br>IPPC Secretariat'
                        remider_message= ReminderMessage()
                        remider_message.pk = None
                        remider_message.emailfrom = "ippc@fao.org"
                        remider_message.emailto = emails
                        remider_message.subject = "IPPC NRO reminder for "+str(cn)+": IPP editors "
                        remider_message.messagebody = textmessage
                        remider_message.date = timezone.now()

                        messages=[]
                        message = mail.EmailMessage(remider_message.subject,remider_message.messagebody,remider_message.emailfrom,
                        emails, ['paola.sentinelli@fao.org'])
                        message.content_subtype = "html"
                        messages.append(message)
                        connection = mail.get_connection()
                        connection.open()
                        #print('test-sending')
                        sent=connection.send_messages(messages)#
                        connection.close()
			remider_message.sent = True
                        remider_message.save()  
                #every 3 months  CPs                          
                if timezone.now().month in date3month :
                    textmessage=''
                   
                    cps=IppcUserProfile.objects.filter(contact_type='1' , country=cn.id)|IppcUserProfile.objects.filter(contact_type='2' , country=cn.id)|IppcUserProfile.objects.filter(contact_type='3' , country=cn.id)|IppcUserProfile.objects.filter(contact_type='4' , country=cn.id)
                    if cps.count() >0:
                        textmessage ='<p>Dear Sir/Madam,<br><br><br>This is a routine reminder sent every 3 months in an attempt to ensure IPPC Contact Points remain more accurate and active.<br><br>The IPPC team kindly requests you to <b>check your national IPPC Contact Point details</b> specified on the <b><a href="https://www.ippc.int/countries/'+country_slug+'">IPP Country page</a></b>:<br><br><br>'
                        textmessage+='<b>'
                        if cps[0].gender!='' and  cps[0].gender!=None:
                            textmessage+= str(dict(GENDER_CHOICES)[cps[0].gender]) 
                        textmessage+=' '+  cps[0].first_name+ ' '+ cps[0].last_name+ '</b><br>'
                        textmessage+= '<br><i>'+cps[0].title+ '</i><br>'
                        textmessage+= cps[0].address1 + '<br>'+ cps[0].address2+ '<br>'
                        textmessage+= '<b>Phone:</b> '+ cps[0].phone+ '<br>'
                        textmessage+= '<b>Email:</b> '+cp_email+ '<br>'
                        textmessage+= '<b>Preferred languages:</b> '
                        for lang in cps[0].preferredlanguage.all():
                            textmessage+= ''+ lang.preferredlanguage 
                        if cps[0].website!='' :
                         textmessage+= '<br><b>Website:</b> '+ cps[0].website+ '<br>'
                        
                        textmessage+='<br><br>You can edit your profile yourself <a href="https://www.ippc.int/en/accounts/login/">after login into the IPP</a> (you can change all contact details except your name and job title). Please find the "Guide to the IPP" <a href="https://www.ippc.int/en/publications/80405/">here</a> for guidance on editing and uploading your national information on the IPP.<br>If you need assistance please contact the IPPC Secretariat at <a href="mailto:IPPC-IT@fao.org">IPPC-IT@fao.org</a>. '
                        textmessage+='<br><br>In case you cannot remember your password, please go to the FAQs: <a href="https://www.ippc.int/en/faq/#LostPassword">lost password</a>.'
                        textmessage+='<br><br>In case the IPPC contact point person has changed, please get <a href="https://www.ippc.int/en/publications/23/">the relevant form</a> duly completed and sign appropriately, and then return it to the IPPC Secretariat at <a href="mailto:ippc@fao.org">ippc@fao.org</a>.<br><br>Best regards<br><br>IPPC Secretariat'
                    
                        remider_message= ReminderMessage()
                        remider_message.pk = None
                        remider_message.emailfrom = "ippc@fao.org"
                        remider_message.emailto = emails
                        remider_message.subject = "IPPC NRO Reminder for "+str(cn)+": Contact Details"
                        remider_message.messagebody = textmessage
                        remider_message.date = timezone.now()
    
                        messages=[]
                        message = mail.EmailMessage(remider_message.subject,remider_message.messagebody,remider_message.emailfrom,
                        emails, ['paola.sentinelli@fao.org'])
                        message.content_subtype = "html"
                        messages.append(message)
                        connection = mail.get_connection()
                        connection.open()
                        #print('test-sending')
                        sent=connection.send_messages(messages)#
                        connection.close()
                        remider_message.sent = True
                        remider_message.save()
                
                #every 6 months pestreporting, eventreporting            
                if timezone.now().month in date6month :
                    countpest=0
                    pests_to_notify=''
                    pestreports = PestReport.objects.filter(country_id=cn.id,status=CONTENT_STATUS_PUBLISHED, is_version=False)
                    for p in pestreports:
                         if timezone.now()- p.modify_date > pestreport_range and p.report_status < 3:
                            countpest+=1
                            if countpest == 1:
                                pests_to_notify=pests_to_notify+'<tr><td><b>Pest report</b></td><td><strong>Last Updated</strong></td></tr>'
                            pests_to_notify+='<tr><td><a href="'+reminder_getlink(country_slug,'pestreports',p)+'">'+p.title+'</a></td><td>'+str(p.modify_date.strftime("%d-%m-%Y"))+'</td></tr>'
                            sql = "UPDATE  ippc_pestreport set to_verify=True where id="+str(p.id)
                            #print(sql)
                            try:
                                cursor.execute(sql)
                                db.commit()
                            except:
                                db.rollback()
                    if countpest==0:
                        pests_to_notify=pests_to_notify+'<tr><td><b>Pest report</b></td><td><strong>Last Updated</strong></td></tr>'
                        pests_to_notify+='<tr><td><a href="'+reminder_getemptylink(country_slug,'pestreports','')+'">link to empty folder</a></td><td>empty folder</td></tr>'
                       
                   
                    eventreporting_to_notify=''
                    evr_to_notify=''
                    evr=0
                    eventreportings = EventReporting.objects.filter(country_id=cn.id,event_rep_type=1, is_version=False)
                    for e in eventreportings:
                      if timezone.now()- e.modify_date > ev_range :
                        evr = evr +1
                        if evr == 1:
                            evr_to_notify=evr_to_notify+'<tr><td><b>'+str(dict(EVT_REP_TYPE_CHOICES_LABELS)[1])+'</b></td><td><strong>Last Updated</strong></td></tr>'
                        evr_to_notify+='<tr><td><a href="'+reminder_getlink(country_slug,'eventreporting',e)+'">'+e.title+'</a></td><td>'+str(p.modify_date.strftime("%d-%m-%Y"))+'</td></tr>'
                        sql = "UPDATE  ippc_eventreporting set to_verify=True where id="+str(e.id)
                        try:
                            cursor.execute(sql)
                            db.commit()
                        except:
                            db.rollback()
                    if evr == 0:
                        evr_to_notify=evr_to_notify+'<tr><td><b>'+str(dict(EVT_REP_TYPE_CHOICES_LABELS)[1])+'</b></td><td><strong>Last Updated</strong></td></tr>'
                        evr_to_notify+='<tr><td><a href="'+reminder_getemptylink(country_slug,'eventreporting','1')+'">link to empty folder</a></td><td>empty folder</td></tr>'

                        
                    textmessage=''
                    if evr_to_notify!='' or pests_to_notify!='':
                        textmessage ='<table bgcolor="#FFFFFF" cellspacing="2" cellpadding="2" valign="top" width="100%" border=1  style="border: 1px solid #10501F;><tr><td width="100%" bgcolor="#FFFFFF" colspan=2>Dear Sir/Madam,<br><br>This is an automatic reminder sent every 6 months for National Reporting Obligations: pest reporting and emergency action.<br><br>Some of the pest reports uploaded by your country in the <a href="https://www.ippc.int/countries/'+country_slug+'"><b>International Phytosanitary Portal</b></a> has got a "Draft" or "Preliminary" status and have not been updated for the last 6 months or more. Please update draft reports or finalize them if necessary. If required please also upload new reports if they are available.  <br><br> Some of emergency action reports have not been updated by your country in the <a href="https://www.ippc.int/countries/'+country_slug+'"><b>International Phytosanitary Portal</b></a> for the last 6 months or more. Please check this information by following links in the table below and update when needed.<br><br>In the table below, you can also find folders which have never had any information uploaded so far. If that information is available, please upload it as soon as possible.<br><br><u><a href="https://www.ippc.int/en/accounts/login/">You will need to log in to the IPP to confirm and/or update the reports.</a></u><br><br></td></tr>'
                        textmessage =textmessage +pests_to_notify
                        textmessage =textmessage +evr_to_notify
                        textmessage =textmessage +'</table>' 
                  
                    if textmessage!='':
                        textmessage=textmessage+'<br>The following resources are available to assist you:<br><ol><li> The manual on editing on the IPP [<a href="https://www.ippc.int/en/publications/80405/">Guide to the IPP</a>]</li><li> <a href="https://www.ippc.int/en/faq/#LostPassword">FAQs</a> including forgotten password </li></ol><br>Should you require any assistance, please contact contact the IPPC Secretariat at <a href="mailto:IPPC-IT@fao.org">IPPC-IT@fao.org</a>.<br><br>Best regards<br><br>IPPC Secretariat'
                   
                        remider_message= ReminderMessage()
                        remider_message.pk = None
                        remider_message.emailfrom = "ippc@fao.org"
                        remider_message.emailto = emails
                        remider_message.subject = "IPPC NRO reminder for "+str(cn)+":  Pest reporting and Emergency action"
                        remider_message.messagebody = textmessage
                        remider_message.date = timezone.now()
    
                        messages=[]
                        message = mail.EmailMessage(remider_message.subject,remider_message.messagebody,remider_message.emailfrom,
                        emails, ['paola.sentinelli@fao.org'])
                        message.content_subtype = "html"
                        messages.append(message)
                        
			connection = mail.get_connection()
                        connection.open()
                        #print('test-sending')
                        sent=connection.send_messages(messages)#
                        connection.close()
    
                        remider_message.sent = True
                        remider_message.save()        

                 #every 3 months       LINKS                    
                if timezone.now().month in date3monthFiles :
                    links_to_notify=''
                    url_to_notify=''
                    r_links=''
                    #timeout in seconds
                    socket.setdefaulttimeout(20)
                    reportingobligations = ReportingObligation.objects.filter(country_id=cn.id, is_version=False)
                    for r in reportingobligations:
                       links_to_notify=''
                       url_to_notify='' 
                       reportingobligations_files= ReportingObligation_File.objects.filter(reportingobligation_id=r.id)
                       reportingobligations_urls= ReportingObligationUrl.objects.filter(reportingobligation_id=r.id)
                       for f in reportingobligations_files:
                          link=''
                          link='https://www.ippc.int/static/media/'+str(f.file)
                          try:
                             urllib2.urlopen(link)
#                             req = urllib2.Request(link)
#                             response = urllib2.urlopen(req)
                          except:
                             print "invalid: ", link
                             links_to_notify=links_to_notify+link+'<br>'
                       for u in reportingobligations_urls:
                          url=''
                          url=str(u.url_for_more_information)
                          url_to_notify=url_to_notify+url+'<br>'
                          #aa+=url+'<br>'
#                          try:
#                             urllib2.urlopen(url)
#                          except:
#                             print "invalid: ", url
#                             url_to_notify=url_to_notify+url+'<br>'
                       if links_to_notify!='' or url_to_notify!='':
                         r_links= r_links+'<tr><td><a href="'+reminder_getlink(country_slug,'reportingobligation',r)+'">'+r.title+'</a></td><td>'+links_to_notify+url_to_notify+'</td></tr>'
                   
                    eventreportings = EventReporting.objects.filter(country_id=cn.id,event_rep_type=1, is_version=False)
                    for e in eventreportings:
                       links_to_notifye=''
                       url_to_notifye='' 
                       ev_files= EventreportingFile.objects.filter(eventreporting_id=e.id)
                       ev_urls= EventreportingUrl.objects.filter(eventreporting_id=e.id)
                       for f in ev_files:
                          link=''
                          link='https://www.ippc.int/static/media/'+str(f.file)
                          try:
                             urllib2.urlopen(link)
#                             req = urllib2.Request(link)
#                             response = urllib2.urlopen(req)
                          except:
                             print "invalid: ", link
                             links_to_notifye=links_to_notifye+link+'<br>'
                       for u in ev_urls:
                          url=''
                          url=str(u.url_for_more_information)
                          url_to_notifye=url_to_notifye+url+'<br>'
                          #aa+=url+'<br>'
#                          try:
#                             urllib2.urlopen(url)
#                          except:
#                             print "invalid: ", url
#                             url_to_notify=url_to_notify+url+'<br>'
                       if links_to_notifye!='' or url_to_notifye!='':
                         r_links= r_links+'<tr><td><a href="'+reminder_getlink(country_slug,'eventreporting',e)+'">'+e.title+'</a></td><td>'+links_to_notifye+url_to_notifye+'</td></tr>'
#                
                    pestreports = PestReport.objects.filter(country_id=cn.id,status=CONTENT_STATUS_PUBLISHED, is_version=False)
                    for p in pestreports:
                       links_to_notifyp=''
                       url_to_notifyp='' 
                       p_files= PestReportFile.objects.filter(pestreport_id=p.id)
                       p_urls=PestReportUrl.objects.filter(pestreport_id=p.id)
                       for f in p_files:
                          link=''
                          link='https://www.ippc.int/static/media/'+str(f.file)
                          try:
                             urllib2.urlopen(link)
#                             req = urllib2.Request(link)
#                             response = urllib2.urlopen(req)
                          except:
                             print "invalid: ", link
                             links_to_notifyp=links_to_notifyp+link+'<br>'
                       for u in p_urls:
                          url=''
                          url=str(u.url_for_more_information)
                          url_to_notifyp=url_to_notifyp+url+'<br>'
                          #aa+=url+'<br>'
#                          try:
#                             urllib2.urlopen(url)
#                          except:
#                             print "invalid: ", url
#                             url_to_notify=url_to_notify+url+'<br>'
                       if links_to_notifyp!='' or url_to_notifyp!='':
                         r_links= r_links+'<tr><td><a href="'+reminder_getlink(country_slug,'pestreports',p)+'">'+p.title+'</a></td><td>'+links_to_notifyp+url_to_notifyp+'</td></tr>'
#                         


                    if r_links!='':
                        textmessage=''
                        textmessage ='<p>Dear Sir/Madam,<br><br>This is an automatic reminder sent every 3 months for National Reporting Obligations that some of the uploaded information by your country in the <a href="https://www.ippc.int">International Phytosanitary Portal</a> have Files that are no longer valid and URLs or links that might need your attention.<br><br>Please check if these files or links still work by following the links in the table below and update if required.<br><br><u><a href="https://www.ippc.int/en/accounts/login/">You will need to log into the IPP to do this action.</a> </u>'
                        
                        textmessage =textmessage+'<table bgcolor="#FFFFFF" cellspacing="2" cellpadding="2" valign="top" width="100%" border=1  style="border: 1px solid #10501F;"><tr><td><strong>Report title</strong></td><td><strong>LINKS/FILES</strong></td></tr>'
                        textmessage=textmessage+ r_links
                        textmessage=textmessage+'</table>'
                        textmessage=textmessage+'<br>The following resources are available to assist you:<br><ol><li> The manual on editing on the IPP [<a href="https://www.ippc.int/en/publications/80405/">Guide to the IPP</a>]</li><li> <a href="https://www.ippc.int/en/faq/#LostPassword">FAQs</a> including forgotten password </li></ol><br>Should you require any assistance, please contact [<a href="mailto:IPPC-IT@fao.org">IPPC-IT@fao.org</a>]<br><br>Best regards<br><br>IPPC Secretariat'
                  
                        remider_message= ReminderMessage()
                        remider_message.pk = None
                        remider_message.emailfrom = "ippc@fao.org"
                        remider_message.emailto = emails
                        remider_message.subject = "IPPC NRO remider for "+str(cn)+": LINKS and FILES"
                        remider_message.messagebody = textmessage
                        remider_message.date = timezone.now()

                        messages=[]
                        message = mail.EmailMessage(remider_message.subject,remider_message.messagebody,remider_message.emailfrom,
                        emails, ['paola.sentinelli@fao.org'])
                        message.content_subtype = "html"
                        messages.append(message)
                        connection = mail.get_connection()
                        connection.open()
                        #print('test-sending')
                        sent=connection.send_messages(messages)#
                        connection.close()
                        remider_message.sent = True
                        remider_message.save()           
                        
        db.close()
   # reminder_log_report.close()     
    msg=notification_to_cn(id)
    
    context = {'textmessages':textmessages,'aa':aa+msg}
    
    response = render(request, "countries/reminder_system.html", context)
    return response



def notification_to_cn(id=None):
    id_range_min=0
    id_range_max=0
    
    id_range_min1=4
    id_range_min2=41
    id_range_min3=77
    id_range_min4=117
    id_range_min5=153
    id_range_min6=189
    id_range_min7=228
    
    id_range_max1=40
    id_range_max2=76
    id_range_max3=116
    id_range_max4=152
    id_range_max5=187
    id_range_max6=226
    id_range_max7=1002
    
   
    if id == '1':
        id_range_min=id_range_min1
        id_range_max=id_range_max1
    elif id == '2':
        id_range_min=id_range_min2
        id_range_max=id_range_max2
    elif id == '3':
            id_range_min=id_range_min3
            id_range_max=id_range_max3
    elif id == '4':
            id_range_min=id_range_min4
            id_range_max=id_range_max4
    elif id == '5':
            id_range_min=id_range_min5
            id_range_max=id_range_max5
    elif id == '6':
            id_range_min=id_range_min6
            id_range_max=id_range_max6
    elif id == '7':
            id_range_min=id_range_min7
            id_range_max=id_range_max7
                 
    msg=""  
    db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
    cursor = db.cursor()
                   
    notify_instances=NotificationMessageRelate.objects.filter()
    for notify_instance in notify_instances:
        if notify_instance.notify:
             
            
            if (notify_instance.sent_date== None and notify_instance.updated_last!=None) or (notify_instance.sent_date!= None and notify_instance.updated_last!=None and notify_instance.updated_last > notify_instance.sent_date):
                subject=''
                url=notify_instance.link
                if notify_instance.new_or_updated=='NEW':
                    subject='IPPC Notification: Country added new content'
                else:    
                    subject='IPPC Notification: Country updated content'
                itemllink="https://www.ippc.int/countries/"+url
                textmessage ='<table bgcolor="#FFFFFF" cellspacing="2" cellpadding="2" valign="top" width="100%" style="border-bottom: 1px solid #10501F;border-top: 1px solid #10501F;border-left: 1px solid #10501F;border-right: 1px solid #10501F"> <tr><td width="100%" bgcolor="#FFFFFF">Please be informed that the following information has been added/updated on the <b>International Phytosanitary Portal:</b><br>'+itemllink+'</td></tr><tr bgcolor="#FFFFFF"><td bgcolor="#FFFFFF"></td></tr><tr><td width="100%" bgcolor="#FFFFFF">If you no longer wish to receive these notifications, please notify this country\'s IPPC Contact Point.</td></tr></table>'
               
                emailto_all = ['']
                for cn in notify_instance.countries.all():
                    countryo = get_object_or_404(CountryPage, page_ptr_id=cn.id)
                    if  cn.id >=id_range_min and cn.id <=id_range_max :          
                        #print(countryo)
                        if countryo.contact_point_id:
                            user_obj=User.objects.get(id=countryo.contact_point_id)
                            emailto_all.append(str(user_obj.email))
                if id == '7':
                    for partner in notify_instance.partners.all():
                        partnerp = get_object_or_404(PartnersPage, page_ptr_id=partner.id)
                        if partnerp.contact_point_id:
                            user_obj=User.objects.get(id=partnerp.contact_point_id)
                            emailto_all.append(str(user_obj.email))
                    if notify_instance.notifysecretariat :
                        emailto_all.append(str('ippc@fao.org'))
                    sql = "UPDATE ippc_notificationmessagerelate set sent_date='"+str(datetime.today())+"' where id="+str(notify_instance.id)+";"
                    #print(sql)
                    try:
                        cursor.execute(sql)
                        db.commit()
                    except:
                        db.rollback()
                    msg+=" __SENT__ "
                    
                       
                messages=[]
               
                message = mail.EmailMessage(subject,textmessage,'ippc@fao.org',#from
                    emailto_all, ['paola.sentinelli@fao.org'])#emailto_all for PROD, in TEST all to paola#
                message.content_subtype = "html"
                messages.append(message)
                msg+="SENDNING TO:"+str(emailto_all)
                
                connection = mail.get_connection()
                connection.open()
                #print('test-sending')
                sent=connection.send_messages(messages)#
                connection.close()
    db.close()
    return msg
                
                        
   

class ReminderMessageListView(ListView):
    """    ReminderMessage List view """
    context_object_name = 'latest'
    model = ReminderMessage
    date_field = 'date'
    template_name = 'emailutility/remindermessage_list.html'
    queryset = ReminderMessage.objects.all().order_by('-date', 'subject')
   
       
class ReminderMessageDetailView(DetailView):
    """ ReminderMessage detail page """
    model = ReminderMessage
    context_object_name = 'remindermessage'
    template_name = 'emailutility/remindermessage_detail.html'
    queryset = ReminderMessage.objects.filter()
from django.utils.translation import ugettext

def commenta(request, template="generic/comments.html"):
    """
    Handle a ``ThreadedCommentForm`` submission and redirect back to its
    related object.
    """
  
    response = myview.initial_validation(request, "comment")
    if isinstance(response, HttpResponse):
        return response
    obj, post_data = response
    #print(obj.categories)
    emailto_all=[]
    #notification to Secretariat of comments
    #for g in obj.groups.all():
    for g in obj.notification_groups.all():
    
       group=Group.objects.get(id=g.id)
       users = group.user_set.all()
       for u in users:
            user_obj=User.objects.get(username=u)
            user_email=user_obj.email
            #print(user_email)   
            emailto_all.append(str(user_email))
           # print("-----------------------------")
    
    import sys;
    reload(sys);
    sys.setdefaultencoding("utf8")
    error_msg=''
    form = myview.ThreadedCommentForm(request, obj, post_data)
    if form.is_valid():
        url = obj.get_absolute_url()
        #print(url)
        if myview.is_spam(request, form, url):
            return redirect(url)
        comment = form.save(request)
        if request.FILES:
            commentfile = CommentFile(comment=comment, file=request.FILES['id_commentfile'])
            commentfile.save()
    
        subject='IPPC new comment on: '+str(obj)  
        commenttext=request.POST['comment']
        splitcommenttext=commenttext.splitlines()
        comment_final=''
        for tt in splitcommenttext:
            comment_final+=tt+'<br>'
        text=request.POST['name']+' has commented on: '+str(obj) +'<br><hr><br>'+str(ugettext(comment_final))+'<br><hr><br>Forum discussion link: <a href="'+str(ugettext('https://www.ippc.int'+url))+'">'+str(ugettext('https://www.ippc.int'+url))+'</a>'
      
        notifificationmessage = mail.EmailMessage(subject,text,'ippc@fao.org', emailto_all, ['paola.sentinelli@fao.org'])
        notifificationmessage.content_subtype = "html"
        sent =notifificationmessage.send()
            
        response = redirect(myview.add_cache_bypass(comment.get_absolute_url()))
        # Store commenter's details in a cookie for 90 days.
        for field in myview.ThreadedCommentForm.cookie_fields:
            cookie_name = myview.ThreadedCommentForm.cookie_prefix + field
            cookie_value = post_data.get(field, "")
            myview.set_cookie(response, cookie_name, cookie_value)
        return response
    elif request.is_ajax() and form.errors:
        return HttpResponse(dumps({"errors": form.errors}))
    else :
        if len(request.POST['comment']) > 3000:
             error_msg='ERROR: the message in the comment exceed the limit of 3000 characters.'
    
    # Show errors with stand-alone comment form.
    context = {"obj": obj, "posted_comment_form": form,"error_msg":error_msg}
    response = render(request, template, context)
    
    return response

class CountryView(TemplateView):
    """ 
    Individual country homepage 
    """
    template_name = 'countries/country_page.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(TemplateView, self).get_context_data(**kwargs)
        context['event_types'] =EVT_REP_TYPE_CHOICES
        context['basic_types'] =BASIC_REP_TYPE_CHOICES 
        self.country = self.kwargs['country']
#CN REMINDER       
        ro=ReportingObligation.objects.filter(country__country_slug=self.country,is_version=False,to_verify=True).count()
        ev=EventReporting.objects.filter(country__country_slug=self.country,is_version=False,to_verify=True).count()
        p=PestReport.objects.filter(country__country_slug=self.country,is_version=False,to_verify=True).count()
        if ro>0 or ev>0 or p>0:
            context['ro_to_verify_1'] ='verify' 
        else:
            context['ro_to_verify_1'] ='noverify'
        context.update({
            'country': self.kwargs['country']
            # 'editors': self.kwargs['editors']
            # 'profile_user': self.kwargs['profile_user']
        })
        
     
        reporting_array = []
        for i in range(1,5):
            reps=ReportingObligation.objects.filter(country__country_slug=self.country,reporting_obligation_type=i,is_version=False)
            reporting_array.append(reps.count())
        
        eventreporting_array = []
        for i in range(1,6):
            evrep=EventReporting.objects.filter(country__country_slug=self.country,event_rep_type=i,is_version=False)
            eventreporting_array.append(evrep.count())
        
     
        pestreporting_array = []
        pests=PestReport.objects.filter(country__country_slug=self.country,is_version=False)
        pestreporting_array.append(pests.count())
        
        datachart= '{label: "Description of NPPO", y: '+str(reporting_array[0])+'},{label: "Legislation: phytosanitary requirements/ restrictions/ prohibitions", y:  '+str(reporting_array[3])+'},{label: "Entry points", y:  '+str(reporting_array[1])+'},{label: "List of regulated pests", y: '+str(reporting_array[2])+'},{label: "Pest reports", y: '+str(pestreporting_array[0])+'},{label: "Organizational arrangements of plant protection", y: '+str(eventreporting_array[2])+'},{label: "Rationale for phytosanitary requirements", y: '+str(eventreporting_array[4])+'},{label: "Non-compliance", y:  '+str(eventreporting_array[1])+'},{label: "Pest status", y:'+str(eventreporting_array[3])+'},{label: "Emergency action", y: '+str(eventreporting_array[0])+'},'

        context['datachart']=datachart
        
        
        
        return context


class PublicationLibraryView(ListView):
    """ 

    """
    template_name = 'pages/publicationlibrary.html'
    #queryset = DraftProtocol.objects.all()
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PublicationLibraryView, self).get_context_data(**kwargs)
#        print('--------------------------------------------------')
#        print('--------------------------------------------------')
#        tot_array = []
#   
#        for k,v in COOPTYPE_CHOICES:
#            print(k)
#            single_array = []
#            single_array.append(k)
#            single_array.append(str(ugettext((v))))
#            cooperations= PartnersPage.objects.filter(coop_type=k).order_by('table_id')
#            single_array3=[]
#            for coop in cooperations:
#                single_array2=[]
#                single_array2.append(coop.table_id)
#                single_array2.append(coop.title)
#                single_array2.append(coop.acronym)
#                responsable_sec=''
#                i=0
#                for cc in coop.responsable_sec.all():
#                    if i>0:
#                        responsable_sec+=' / '
#                    if cc !=None:
#                        user_obj=User.objects.get(username=cc)
#                        userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
#                        val=userippc.gender
#                        gender=''
#                        if val!= None:
#                            if val == 1:
#                                gender= "Mr."
#                            elif val==2: 
#                                gender= "Ms."
#                            elif val==3:    
#                                gender= "Mrs."
#                            elif val==4:    
#                                gender= "Professor."
#                            elif val==5:    
#                                gender= "M."
#                            elif val==6:    
#                                gender= "Mme."
#                            elif val==7:    
#                                gender= "Dr."
#                            elif val==8:    
#                                gender= "Sr."
#                            elif val==9:    
#                                gender= "Sra."
#
#                        responsable_sec+=(unicode(gender))+' '+(unicode(userippc.first_name))+' '+(unicode(userippc.last_name))  
#                    i=i+1
#                  
#                single_array2.append(responsable_sec)
#                cp=''
#                if coop.contact_point != None:
#                    user_obj2=User.objects.get(username=coop.contact_point)
#                    userippc2 = get_object_or_404(IppcUserProfile, user_id=user_obj2.id)
#                    val=userippc2.gender
#                    gender=''
#                    if val!= None:
#                        if val == 1:
#                            gender= "Mr."
#                        elif val==2: 
#                            gender= "Ms."
#                        elif val==3:    
#                            gender= "Mrs."
#                        elif val==4:    
#                            gender= "Professor."
#                        elif val==5:    
#                            gender= "M."
#                        elif val==6:    
#                            gender= "Mme."
#                        elif val==7:    
#                            gender= "Dr."
#                        elif val==8:    
#                            gender= "Sr."
#                        elif val==9:    
#                            gender= "Sra."
#
#                    cp=(unicode(gender))+' '+(unicode(userippc2.first_name))+' '+(unicode(userippc2.last_name))
#
#                single_array2.append(cp)
#                single_array3.append(single_array2)
#            single_array.append(single_array3)
#            tot_array.append(single_array)  
#            
#        print(tot_array)
#        context['coop']=tot_array
#        print('--------------------------------------------------')
#        print('--------------------------------------------------')o
#       
        context['polls']=Poll.objects.order_by('-pub_date').all
        
        
        queryset = DraftProtocol.objects.all()
        dps=[]
        for dp in queryset:
            if (timezone.now() < dp.closing_date):
               dps.append(dp)
        queryset= dps  
        user = self.request.user  
       
        if user.groups.filter(name='IPPC Secretariat'):
           queryset = DraftProtocol.objects.all()
       
        elif user.groups.filter(name='TPDP'):
            dps=[]
            for dp in queryset:
                if (timezone.now() < dp.closing_date):
                   dps.append(dp)
            queryset= dps      
                   
        elif user.groups.filter(name='TPDPc'):
           dps=[]
           for dp in queryset:
               #print(dp.users.all())
               if (timezone.now() < dp.closing_date) and user in dp.users.all():
                   dps.append(dp)
           queryset= dps     
            
    
        
        context['latest1']=queryset
       
        users_sec=[]
        users_sc=[] 
        users_bureau=[]
        users_tpfq=[]
        users_tpdp=[]
        users_tpff=[]
        users_tpg=[]
        users_tppt=[]
        users_esg=[]
        for g in Group.objects.filter():
            if g.name == 'IPPC Secretariat': 
                users = g.user_set.all()
                for u in users:
                   users_u=[]
                   user_obj=User.objects.get(username=u)
                   userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                   users_u.append((unicode(userippc.last_name)))
                   users_u.append((unicode(userippc.first_name)))
                   users_u.append((userippc.profile_photo))
                   users_u.append((userippc.title))
                   users_u.append((user_obj.username))
                   users_sec.append(users_u)
            if g.name == 'Standards committee':
                users = g.user_set.all()
                for u in users:
                   users_u=[]
                   user_obj=User.objects.get(username=u)
                   userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                   users_u.append((unicode(userippc.last_name)))
                   users_u.append((unicode(userippc.first_name)))
                   users_u.append((userippc.profile_photo))
                   users_u.append((userippc.title))
                   users_u.append((user_obj.username))
                   users_sc.append(users_u)
          
            if g.name == 'Bureau': 
                users = g.user_set.all()
                for u in users:
                   users_u=[]
                   user_obj=User.objects.get(username=u)
                   userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                   users_u.append((unicode(userippc.last_name)))
                   users_u.append((unicode(userippc.first_name)))
                   users_u.append((userippc.profile_photo))
                   users_u.append((userippc.title))
                   users_u.append((user_obj.username))
                   users_bureau.append(users_u)
            if g.name == 'TPDP': 
                users = g.user_set.all()
                for u in users:
                   users_u=[]
                   user_obj=User.objects.get(username=u)
                   userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                   users_u.append((unicode(userippc.last_name)))
                   users_u.append((unicode(userippc.first_name)))
                   users_u.append((userippc.profile_photo))
                   users_u.append((userippc.title))
                   users_u.append((user_obj.username))
                   users_tpdp.append(users_u)
            if g.name == 'TPFQ': 
                users = g.user_set.all()
                for u in users:
                   users_u=[]
                   user_obj=User.objects.get(username=u)
                   userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                   users_u.append((unicode(userippc.last_name)))
                   users_u.append((unicode(userippc.first_name)))
                   users_u.append((userippc.profile_photo))
                   users_u.append((userippc.title))
                   users_u.append((user_obj.username))
                   users_tpfq.append(users_u)
            if g.name == 'TPFF': 
                users = g.user_set.all()
                for u in users:
                   users_u=[]
                   user_obj=User.objects.get(username=u)
                   userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                   users_u.append((unicode(userippc.last_name)))
                   users_u.append((unicode(userippc.first_name)))
                   users_u.append((userippc.profile_photo))
                   users_u.append((userippc.title))
                   users_u.append((user_obj.username))
                   users_tpff.append(users_u)
            if g.name == 'TPG': 
                users = g.user_set.all()
                for u in users:
                   users_u=[]
                   user_obj=User.objects.get(username=u)
                   userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                   users_u.append((unicode(userippc.last_name)))
                   users_u.append((unicode(userippc.first_name)))
                   users_u.append((userippc.profile_photo))
                   users_u.append((userippc.title))
                   users_u.append((user_obj.username))
                   users_tpg.append(users_u)
            if g.name == 'TPPT': 
                users = g.user_set.all()
                for u in users:
                   users_u=[]
                   user_obj=User.objects.get(username=u)
                   #print(user_obj.groups)
                   gg=''
                   for h in user_obj.groups.all():
                        gg=gg+str(h)+', '
                   #print(gg)
                   userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                   users_u.append((unicode(userippc.last_name)))
                   users_u.append((unicode(userippc.first_name)))
                   users_u.append((userippc.profile_photo))
                   users_u.append((userippc.title))
                   users_u.append((user_obj.username))
                   users_u.append((gg))
                   users_tppt.append(users_u)
            if g.name == 'ePhyto Steering Group (ESG)': 
                users = g.user_set.all()
                for u in users:
                   users_u=[]
                   user_obj=User.objects.get(username=u)
                   gg=''
                   for h in user_obj.groups.all():
                        gg=gg+str(h)+', '
                   userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                   users_u.append((unicode(userippc.last_name)))
                   users_u.append((unicode(userippc.first_name)))
                   users_u.append((userippc.profile_photo))
                   users_u.append((userippc.title))
                   users_u.append((user_obj.username))
                   users_u.append((gg))
                   users_esg.append(users_u)       
            
        
        context['secretariat']=users_sec
        context['users_sc']=users_sc
        context['users_bureau']=users_bureau
        context['users_tpdp']=users_tpdp
        context['users_tpff']=users_tpff
        context['users_tppt']=users_tppt
        context['users_tpg']=users_tpg
        context['users_tpfq']=users_tpfq
        context['users_esg']=users_esg
        
        secretariats_final=[]
        for i in range(1,6):
            unit=''
            if i ==  1:
                unit='IPPC Secretary'
            if i ==  2:
                unit='Standard Setting Unit'
            if i ==  3:
                unit='Implementation Facilitation Unit'
            if i ==  4:
                unit='Integration and Support Team'
            if i ==  5:
                unit='ePhyto Group'
                    
            secretariats_ss=[]
            secretariats=IppcUserProfile.objects.filter(contact_type='18',unit_team=i).order_by('staff_oder',)
            secretariats_ss.append(unit)
            secretariats_ss.append(secretariats)
            secretariats_final.append(secretariats_ss) 
      
  



      
        table_sec='<table class="table table-condensed  table-bordered dataTable">'
        table_sec+='<tr>'
        table_sec+='<th style="background-color: #0f405b;color: white" colspan="2">Organizational Structure of the IPPC Secretariat</th>'
        table_sec+='</tr>'

        
           
    
        i=0
        for sec in secretariats_final:
            table_sec+='<tr> <th style="background-color: #E0E0E0" colspan="2">'+sec[0]+'</th></tr>'
            for userippc in sec[1]:
                table_sec+='<tr>'
                user_obj=User.objects.get(id=userippc.user_id)
                if userippc.profile_photo:
                     table_sec+='<td><img class="profile-picture" style="margin:0 0 20px 10px;" src="/static/media/'+str(userippc.profile_photo)+'" width="80px"></td>'
                else:
                    table_sec+='<td><i class="fas fa-user"> </i></td>'
                table_sec+='<td>'+(unicode(userippc.first_name))+' '+(unicode(userippc.last_name))+'<br>'+userippc.title
                table_sec+='<input class="toggle-box" id="identifier-'+str(i)+'" type="checkbox" style="display:none">'
                table_sec+='<label for="identifier-'+str(i)+'"></label><div>'
                table_sec+='<i class="far fa-envelope"></i> E-mail <a href="mailto:'+user_obj.email+'">'+user_obj.email+'</a><br>'+unicode(userippc.bio)
                table_sec+='</div></td>'
                table_sec+='</tr>'
                i+=1
        table_sec+='</table>'
        context['table_sec']=table_sec
         
            
        return context
    def get_queryset(self):
        queryset = DraftProtocol.objects.all()
        return  queryset 
 

class CountryRelatedView(TemplateView):
    """ 
    Individual country related info 
    """
    template_name = 'countries/country_related_info.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(TemplateView, self).get_context_data(**kwargs)
        #print(self.kwargs['country'])
        countryo= CountryPage.objects.filter(country_slug=self.kwargs['country'])
        context['iso3'] =countryo[0].iso3 
        context['cn_map'] =countryo[0].cn_map 
        context.update({
            'country': self.kwargs['country']
        })
        return context



class PestReportListView(ListView):
    """
    Pest reports
        http://stackoverflow.com/questions/8547880/listing-object-with-specific-tag-using-django-taggit
        http://stackoverflow.com/a/7382708/412329
    """
    context_object_name = 'latest'
    model = PestReport
    date_field = 'publish_date'
    template_name = 'countries/pest_report_list.html'
    queryset = PestReport.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return pest reports from the specific country """
        # self.country = get_object_or_404(CountryPage, country=self.kwargs['country'])
        self.country = self.kwargs['country']
        # CountryPage country_slug == country URL parameter keyword argument
        return PestReport.objects.filter(country__country_slug=self.country, is_version=False)# status=CONTENT_STATUS_PUBLISHED,
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PestReportListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        # context.update({
        #     'country': self.kwargs['country']
        # })
        return context

# @login_required
# @permission_required('ippc.add_pestreport', login_url="/accounts/login/")

class PestReportHiddenListView(ListView):
    """
    Hidden Pest reports so editors can still edit them
    """
    context_object_name = 'latest'
    model = PestReport
    date_field = 'publish_date'
    template_name = 'countries/pest_report_hidden_list.html'
    queryset = PestReport.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return pest reports from the specific country """
        # self.country = get_object_or_404(CountryPage, country=self.kwargs['country'])
        self.country = self.kwargs['country']
        # CountryPage country_slug == country URL parameter keyword argument
        return PestReport.objects.filter(country__country_slug=self.country, status=CONTENT_STATUS_DRAFT, is_version=False)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PestReportHiddenListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        # context.update({
        #     'country': self.kwargs['country']
        # })
        return context

    # put class-based generic view behind login
    # https://docs.djangoproject.com/en/dev/topics/class-based-views/intro/#decorating-the-class
    @method_decorator(login_required)
    @method_decorator(permission_required('ippc.add_pestreport', login_url="/accounts/login/"))
    def dispatch(self, *args, **kwargs):
        return super(PestReportHiddenListView, self).dispatch(*args, **kwargs)



class PestReportDetailView(DetailView):
    """ Pest report detail page """
    model = PestReport
    context_object_name = 'report'
    template_name = 'countries/pest_report_detail.html'
    queryset = PestReport.objects.filter(status=CONTENT_STATUS_PUBLISHED)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PestReportDetailView, self).get_context_data(**kwargs)
      
        p = get_object_or_404(PestReport, slug=self.kwargs['slug'])
        context['8col'] = 1
        versions= PestReport.objects.filter(country__country_slug=self.kwargs['country'], status=CONTENT_STATUS_PUBLISHED, is_version=True, parent_id=p.id).order_by('-modify_date')
        context['versions'] = versions
        return context
    
class PublicationListView(ListView):
    """
    Publications List
    """
    context_object_name = 'latest'
    model = Publication
    date_field = 'modify_date'
    template_name = 'pages/publication_list.html'
    queryset = Publication.objects.filter(status=IS_PUBLIC,is_version=False).order_by('-publication_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

class PublicationDetailView(DetailView):
    """ Publication detail page """
    model = Publication
    context_object_name = 'publication'
    template_name = 'pages/publication_detail.html'
    queryset = Publication.objects.filter(status=IS_PUBLIC)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PublicationDetailView, self).get_context_data(**kwargs)
        
        publication = get_object_or_404(Publication, id=self.kwargs['pk'])
    
        publicationlibrary = get_object_or_404(PublicationLibrary, id=publication.library_id)
        
       
        context['users'] =publicationlibrary.users
        context['groups'] = publicationlibrary.groups
        context['login_required'] = publicationlibrary.login_required
        
      
        restrictedmessage=''    
        versions=''
        if publicationlibrary.login_required:
            context['publication'] = ''
        
            if self.request.user.is_authenticated():
                if self.request.user.groups.filter(name='IPPC Secretariat') or self.request.user.groups.filter(name='Admin'):
                    context['publication'] = publication
                    versions= Publication.objects.filter( is_version=True, parent_id=publication.id).order_by('-modify_date')
                    restrictedmessage= ''
                else:
                   for g in self.request.user.groups.all():
                     
                        if g in publicationlibrary.groups.all() and g not in publication.groups.all():
                              context['publication'] = publication
                              versions= Publication.objects.filter( is_version=True, parent_id=publication.id).order_by('-modify_date')
                              restrictedmessage= ''
                        else:      
                            restrictedmessage= 'true'
            else: 
                context['publication'] = ''
                restrictedmessage= 'true'
            
        else:
            context['publication'] =publication
            versions= Publication.objects.filter( is_version=True, parent_id=publication.id).order_by('-modify_date')
      
       
     
        context['versions'] = versions
        context['restrictedmessage'] = restrictedmessage
      
      
        
        return context
    
class PublicationDetail2View(DetailView):
    """ Publication detail page """
    model = Publication
    context_object_name = 'publication'
    template_name = 'pages/publication_detail.html'
    queryset = Publication.objects.filter(status=IS_PUBLIC)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PublicationDetail2View, self).get_context_data(**kwargs)
        publication = get_object_or_404(Publication, slug=self.kwargs['slug'])
    
        publicationlibrary = get_object_or_404(PublicationLibrary, id=publication.library_id)
        
       
        context['users'] =publicationlibrary.users
        context['groups'] = publicationlibrary.groups
        context['login_required'] = publicationlibrary.login_required
        
      
        restrictedmessage=''    
        versions=''
        if publicationlibrary.login_required:
            context['publication'] = ''
        
            if self.request.user.is_authenticated():
                if self.request.user.groups.filter(name='IPPC Secretariat') or self.request.user.groups.filter(name='Admin'):
                    context['publication'] = publication
                    versions= Publication.objects.filter( is_version=True, parent_id=publication.id).order_by('-modify_date')
                    restrictedmessage= ''
                else:
                   for g in self.request.user.groups.all():
                     
                        if g in publicationlibrary.groups.all() and g not in publication.groups.all():
                              context['publication'] = publication
                              versions= Publication.objects.filter( is_version=True, parent_id=publication.id).order_by('-modify_date')
                              restrictedmessage= ''
                        else:      
                            restrictedmessage= 'true'
            else: 
                context['publication'] = ''
                restrictedmessage= 'true'
            
        else:
            context['publication'] =publication
            versions= Publication.objects.filter( is_version=True, parent_id=publication.id).order_by('-modify_date')
      
       
     
        context['versions'] = versions
        context['restrictedmessage'] = restrictedmessage
        return context
    


class PublicationFilesListView(ListView):
    """
    Publications Files List
    """
    context_object_name = 'latest'
    model = Publication
    date_field = 'modify_date'
    template_name = 'pages/publicationfilestable.html'
    queryset = Publication.objects.filter(status=IS_PUBLIC,is_version=False).order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500
 
    def get_context_data(self, **kwargs):
        context = super(PublicationFilesListView, self).get_context_data(**kwargs)
        queryset = Publication.objects.filter(status=IS_PUBLIC,library_id=self.kwargs['id'],is_version=False).order_by('-modify_date', 'title')
        queryset2 = PublicationLibrary.objects.filter(id=self.kwargs['id'])
        for p in queryset2:
            context['titlepage']=p.title
        filenames_all=[]
        filenames_en=[]
        filenames_es=[]
        filenames_fr=[]
        filenames_ar=[]
        filenames_ru=[]
        filenames_zh=[]
        
        langs=[]
        langs.append(["en",filenames_en])
        langs.append(["es",filenames_es])
        langs.append(["fr",filenames_fr])
        langs.append(["ar",filenames_ar])
        langs.append(["ru",filenames_ru])
        langs.append(["zh",filenames_zh])
        
        for p in queryset:
            user_g=self.request.user.groups
            pub_g=p.groups
            if pub_g.all():
                for pg in pub_g.all():
                    if pg in user_g.all():
                        break;
                    else:
                        filenames_all.append(p.file_en)
                        filenames_all.append(p.file_es)
                        filenames_all.append(p.file_fr)
                        filenames_all.append(p.file_ru)
                        filenames_all.append(p.file_zh)
                        filenames_all.append(p.file_ar)
                        filenames_en.append(p.file_en)
                        filenames_es.append(p.file_es)
                        filenames_fr.append(p.file_fr)
                        filenames_ar.append(p.file_ar)
                        filenames_ru.append(p.file_ru)
                        filenames_zh.append(p.file_zh)
            else:
                filenames_all.append(p.file_en)
                filenames_all.append(p.file_es)
                filenames_all.append(p.file_fr)
                filenames_all.append(p.file_ru)
                filenames_all.append(p.file_zh)
                filenames_all.append(p.file_ar)
                filenames_en.append(p.file_en)
                filenames_es.append(p.file_es)
                filenames_fr.append(p.file_fr)
                filenames_ar.append(p.file_ar)
                filenames_ru.append(p.file_ru)
                filenames_zh.append(p.file_zh)             

            
        
        # The zip compressor
        date = timezone.now().strftime('%Y%m%d%H%M%S')+"_"+str(self.kwargs['id'])
        zip_all1 ="/static/media/tmp/"+"archive_all_"+ date+".zip"
        zip_all = zipfile.ZipFile(MEDIA_ROOT+"/tmp/"+"archive_all_"+ date+".zip", "w")
        for lang in langs:
            zip_lang1 = "/static/media/tmp/"+"archive_"+str(lang[0])+"_"+ date+".zip"
            zip_lang = zipfile.ZipFile(MEDIA_ROOT+"/tmp/"+"archive_"+str(lang[0])+"_"+ date+".zip", "w")
            for file_path in lang[1]:
                strfpath=os.path.join('/work/projects/ippcdj-env/public/', '/work/projects/ippcdj-env/public/static/media/')+str(file_path)
                filename = strfpath.split('/');
                fname=filename[len(filename)-1]
                zip_lang.write(strfpath, fname)
                zip_all.write(strfpath, fname)
            
            zip_lang.close()
            context['zip_'+str(lang[0])]=zip_lang1
            size=os.path.getsize(zip_lang.filename)
            if size >182:
                context['zip_'+str(lang[0])+'_s']=os.path.getsize(zip_lang.filename)
        
        zip_all.close()
        
        context['zip_all']=zip_all1
        context['zip_all_s']=os.path.getsize(zip_all.filename)
        
        destination = '/work/projects/ippcdj-env/public/static/media/tmp/'
        src_files = os.listdir(MEDIA_ROOT+"/tmp/")
        for file_name in src_files:
            full_file_name = os.path.join(MEDIA_ROOT+"/tmp/", file_name)
            #if (os.path.isfile(full_file_name)):
            #     shutil.move(full_file_name, destination)
        source = MEDIA_ROOT+"/tmp/"
        
        return context
            
class PublicationMeetingFilesListView(ListView):
    """
    PublicationMeeting Files List
    """
    context_object_name = 'latest'
    model = Publication
    date_field = 'modify_date'
    template_name = 'pages/publicationmeetingfilestable.html'
    queryset = CommitteeMeeting.objects.filter().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500
 
    def get_context_data(self, **kwargs):
        context = super(PublicationMeetingFilesListView, self).get_context_data(**kwargs)
        queryset = CommitteeMeeting.objects.filter(library_id=self.kwargs['id']).order_by('-modify_date', 'title')
        queryset2 = PublicationLibrary.objects.filter(id=self.kwargs['id'])
        for p in queryset2:
            context['titlepage']=p.title
        filenames_all=[]
        filenames_en=[]
        filenames_es=[]
        filenames_fr=[]
        filenames_ar=[]
        filenames_ru=[]
        filenames_zh=[]
        
        langs=[]
        langs.append(["en",filenames_en])
        langs.append(["es",filenames_es])
        langs.append(["fr",filenames_fr])
        langs.append(["ar",filenames_ar])
        langs.append(["ru",filenames_ru])
        langs.append(["zh",filenames_zh])
        
        for p in queryset:
            filenames_all.append(p.report_en)
            filenames_all.append(p.report_es)
            filenames_all.append(p.report_fr)
            filenames_all.append(p.report_ru)
            filenames_all.append(p.report_zh)
            filenames_all.append(p.report_ar)
            filenames_en.append(p.report_en)
            filenames_es.append(p.report_es)
            filenames_fr.append(p.report_fr)
            filenames_ar.append(p.report_ar)
            filenames_ru.append(p.report_ru)
            filenames_zh.append(p.report_zh)             

            
        
        # The zip compressor
        date = timezone.now().strftime('%Y%m%d%H%M%S')+"_"+str(self.kwargs['id'])
        zip_all1 ="/static/media/tmp/"+"archive_all_"+ date+".zip"
        zip_all = zipfile.ZipFile(MEDIA_ROOT+"/tmp/"+"archive_all_"+ date+".zip", "w")
        for lang in langs:
            zip_lang1 = "/static/media/tmp/"+"archive_"+str(lang[0])+"_"+ date+".zip"
            zip_lang = zipfile.ZipFile(MEDIA_ROOT+"/tmp/"+"archive_"+str(lang[0])+"_"+ date+".zip", "w")
            for file_path in lang[1]:
                strfpath=os.path.join('/work/projects/ippcdj-env/public/', '/work/projects/ippcdj-env/public/static/media/')+str(file_path)
                filename = strfpath.split('/');
                fname=filename[len(filename)-1]
                zip_lang.write(strfpath, fname)
                zip_all.write(strfpath, fname)
            
            zip_lang.close()
            context['zip_'+str(lang[0])]=zip_lang1
            size=os.path.getsize(zip_lang.filename)
            if size >182:
                context['zip_'+str(lang[0])+'_s']=os.path.getsize(zip_lang.filename)
        
        zip_all.close()
        
        context['zip_all']=zip_all1
        context['zip_all_s']=os.path.getsize(zip_all.filename)
        
       
        destination = '/work/projects/ippcdj-env/public/static/media/tmp/'
        src_files = os.listdir(MEDIA_ROOT+"/tmp/")
        for file_name in src_files:
            full_file_name = os.path.join(MEDIA_ROOT+"/tmp/", file_name)
            #if (os.path.isfile(full_file_name)):
            #     shutil.move(full_file_name, destination)
        source = MEDIA_ROOT+"/tmp/"
        
        return context
    
            
            
            
            
def send_notification_message(newitem,id,content_type,title,url):
    """ send_notification_message """
    notify_instance = get_object_or_404(NotificationMessageRelate, object_id=id,content_type__pk=content_type.id,)
#    
#    if notify_instance.notify:
#        emailto_all = ['']
#        for cn in notify_instance.countries.all():
#            countryo = get_object_or_404(CountryPage, page_ptr_id=cn.id)
#            if countryo.contact_point_id:
#                user_obj=User.objects.get(id=countryo.contact_point_id)
#                emailto_all.append(str(user_obj.email))
#        for partner in notify_instance.partners.all():
#            countryo = get_object_or_404(PartnersPage, page_ptr_id=partner.id)
#            if countryo.contact_point_id:
#                user_obj=User.objects.get(id=countryo.contact_point_id)
#                emailto_all.append(str(user_obj.email))
#        if notify_instance.notifysecretariat :
#         
#            emailto_all.append(str('ippc@fao.org'))
#        subject=''
#        if newitem==1:
#            subject='ADDED new content: ' +title
#        else:    
#            subject='UPDATE to: ' +title
#        itemllink="https://www.ippc.int/countries/"+url
#        textmessage ='<table bgcolor="#FFFFFF" cellspacing="2" cellpadding="2" valign="top" width="100%" style="border-bottom: 1px solid #10501F;border-top: 1px solid #10501F;border-left: 1px solid #10501F;border-right: 1px solid #10501F"> <tr><td width="100%" bgcolor="#FFFFFF">Please be informed that the following information has been added/updated on the <b>International Phytosanitary Portal:</b><br>'+title+' ('+itemllink+')</td></tr><tr bgcolor="#FFFFFF"><td bgcolor="#FFFFFF"></td></tr><tr><td width="100%" bgcolor="#FFFFFF">If you no longer wish to receive these notifications, please notify this country\'s IPPC Contact Point.</td></tr></table>'
#        message = mail.EmailMessage(subject,textmessage,'ippc@fao.org',#from
#            emailto_all, ['paola.sentinelli@fao.org'])#emailto_all for PROD, in TEST all to paola#
#        print(textmessage)
#        message.content_subtype = "html"
#        sent =message.send()
   
#def send_pest_notification_message(newitem,id,content_type,title,url):
#    """ send_notification_message """
#    #print("send!!!")
#    emailto_all = ['roy@eppo.int','qingpo.yang@fao.org']#'dave.nowell@fao.org',
#    
#    subject=''
#    if newitem==1:
#        subject='ADDED new Pest report on IPPC: ' +title
#    else:    
#        subject='UPDATED Pest report on IPPC: ' +title
#    itemllink="https://www.ippc.int/countries/"+url
#    textmessage ='<table bgcolor="#FFFFFF" cellspacing="2" cellpadding="2" valign="top" width="100%" style="border-bottom: 1px solid #10501F;border-top: 1px solid #10501F;border-left: 1px solid #10501F;border-right: 1px solid #10501F"> <tr><td width="100%" bgcolor="#FFFFFF">Please be informed that the following information has been added/updated on the <b>International Phytosanitary Portal:</b><br>'+title+' ('+itemllink+')</td></tr><tr bgcolor="#FFFFFF"><td bgcolor="#FFFFFF"></td></tr></table>'
#
#    message = mail.EmailMessage(subject,textmessage,'ippc@fao.org',#from
#        emailto_all, ['paola.sentinelli@fao.org'])#emailto_all for PROD, in TEST all to paola#
#    #print(textmessage)
#    message.content_subtype = "html"
#    #print('test-sending')
#    sent =message.send()
        
def send_report_notification_message(newitem,type,id,content_type,title,url):
    """ send_report_notification_message """
    #print("send!!!")
    emailto_all = ['qingpo.yang@fao.org']
    if type == 0:
        emailto_all = ['roy@eppo.int','qingpo.yang@fao.org','arop.deng@fao.org','mirko.montuori@fao.org']#'dave.nowell@fao.org',
    
    subject=''
    typename=''
    
    if type==0:
        typename='Pest report'
    elif type==1:
        typename='Reporting Obligation'
    elif type==2:    
        typename='Event Reporting'
    
    
    if newitem==1:
        subject='ADDED new '+str(typename)+' on IPP: ' +title
    else:    
        subject='UPDATED '+str(typename)+' on IPP: ' +title
        
    itemllink="https://www.ippc.int/countries/"+url
    textmessage ='<table bgcolor="#FFFFFF" cellspacing="2" cellpadding="2" valign="top" width="100%" style="border-bottom: 1px solid #10501F;border-top: 1px solid #10501F;border-left: 1px solid #10501F;border-right: 1px solid #10501F"> <tr><td width="100%" bgcolor="#FFFFFF">Please be informed that the following information has been added/updated on the <b>International Phytosanitary Portal:</b><br>'+title+' ('+itemllink+')</td></tr><tr bgcolor="#FFFFFF"><td bgcolor="#FFFFFF"></td></tr></table>'

    message = mail.EmailMessage(subject,textmessage,'ippc@fao.org',#from
        emailto_all, ['paola.sentinelli@fao.org'])#emailto_all for PROD, in TEST all to paola#
    #print(textmessage)
    message.content_subtype = "html"
    #print('test-sending')
    sent =message.send()
 
@login_required
@permission_required('ippc.add_pestreport', login_url="/accounts/login/")
def pest_report_create(request, country):
    """ Create Pest Report """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))

    form = PestReportForm(request.POST, request.FILES)
    issueform =IssueKeywordsRelateForm(request.POST)
    commodityform =CommodityKeywordsRelateForm(request.POST)
    notifyrelateform =NotificationMessageRelateForm(request.POST)
        
    countryo = get_object_or_404(CountryPage, name=country)
    numberR=PestReport.objects.filter(country_id=country.id,is_version=False).count()
    numberR=numberR+1
    pestnumber=str(numberR)
    if numberR<10 :
        pestnumber='0'+pestnumber
    report_number_val=countryo.iso3+'-'+pestnumber+'/1'
    #print   (report_number_val)     
   
    if request.method == "POST":
         f_form = PestReportFileFormSet(request.POST, request.FILES)
         u_form = PestReportUrlFormSet(request.POST)
         if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            new_pest_report = form.save(commit=False)
            new_pest_report.author = request.user
            new_pest_report.report_number=report_number_val
            new_pest_report.author_id = author.id
            
            form.save()
           
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = new_pest_report
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = new_pest_report
            commodity_instance.save()
            commodityform.save_m2m()
            
            notify_instance = notifyrelateform.save(commit=False)
            notify_instance.content_object = new_pest_report
            notify_instance.new_or_updated = 'NEW'
            notify_instance.link = user_country_slug+'/pestreports/'+str(new_pest_report.publish_date.strftime("%Y"))+'/'+str(new_pest_report.publish_date.strftime("%m"))+'/'+new_pest_report.slug+'/'
            notify_instance.updated_last = new_pest_report.publish_date
            
            notify_instance.save()
            notifyrelateform.save_m2m()
            
            f_form.instance = new_pest_report
            f_form.save()
            u_form.instance = new_pest_report
            u_form.save()
            content_type = ContentType.objects.get_for_model(new_pest_report)
       
            send_notification_message(1,new_pest_report.id,content_type,new_pest_report.title,user_country_slug+'/pestreports/'+str(new_pest_report.publish_date.strftime("%Y"))+'/'+str(new_pest_report.publish_date.strftime("%m"))+'/'+new_pest_report.slug+'/')
            send_report_notification_message(1,0,new_pest_report.id,content_type,new_pest_report.title,user_country_slug+'/pestreports/'+str(new_pest_report.publish_date.strftime("%Y"))+'/'+str(new_pest_report.publish_date.strftime("%m"))+'/'+new_pest_report.slug+'/')
            info(request, _("Successfully created pest report."))
            
            if new_pest_report.status == CONTENT_STATUS_DRAFT:
                return redirect("pest-report-hidden-list", country=user_country_slug)
            else:
                return redirect("pest-report-detail", country=user_country_slug, year=new_pest_report.publish_date.strftime("%Y"), month=new_pest_report.publish_date.strftime("%m"), slug=new_pest_report.slug)
         else:
             return render_to_response('countries/pest_report_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform,'notifyrelateform':notifyrelateform},
             context_instance=RequestContext(request))
       
    else:
        form = PestReportForm(initial={'country': country}, instance=PestReport())
        issueform =IssueKeywordsRelateForm(request.POST)
        commodityform =CommodityKeywordsRelateForm(request.POST)
        notifyrelateform =NotificationMessageRelateForm(request.POST)
   
        f_form =PestReportFileFormSet()
        u_form =PestReportUrlFormSet()
    return render_to_response('countries/pest_report_create.html', {'form': form,'f_form': f_form,'u_form':u_form,'issueform':issueform, 'commodityform':commodityform,'notifyrelateform':notifyrelateform},
        context_instance=RequestContext(request))

import MySQLdb
# http://stackoverflow.com/a/1854453/412329
@login_required
@permission_required('ippc.change_pestreport', login_url="/accounts/login/")
def pest_report_edit(request, country, id=None, template_name='countries/pest_report_edit.html'):
    """ Edit Pest Report """
    user = request.user
    author = user
    country = user.get_profile().country
    # country_id = PestReport.objects.filter(country__country_id=country.id)
    user_country_slug = lower(slugify(country))
    
    
    
    if id:
        pest_report = get_object_or_404(PestReport, country=country, pk=id)
        old_pest_report=get_object_or_404(PestReport, country=country, pk=id)
        
        content_type = ContentType.objects.get_for_model(pest_report)
        try: 
            notifications = get_object_or_404(NotificationMessageRelate, object_id=id,content_type__pk=content_type.id)
        except:
            notifications = None
        
        rep_num=pest_report.report_number
        indexof=rep_num.rfind('/')
        numberRep_part=rep_num[:indexof+1]
        numberRep=int(rep_num[indexof+1:])+1
        pest_report.report_number=numberRep_part+str(numberRep)
       # if pest_report.author != request.user:
        #     return HttpResponseForbidden()
    else:
        pest_report = PestReport(author=request.user)
        

    
    old_issue=[]
    if pest_report.issuename.count()>0:
        for e in pest_report.issuename.all():
                obj_i=e.content_object.issuename
                for o in obj_i.all():
                    for iss in o.issuename.all():
                        old_issue.append(iss.id)
    old_comm=[]
    if pest_report.commname.count()>0:
        for e in pest_report.commname.all():
                obj_c=e.content_object.commname
                for o in obj_c.all():
                    for com in o.commname.all():
                        old_comm.append(com.id)                    
  
    if request.POST:
        form = PestReportForm(request.POST,  request.FILES, instance=pest_report)
        if pest_report.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=pest_report.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
             
        if pest_report.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=pest_report.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm(request.POST,instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm(request.POST)
        notifyrelateform =NotificationMessageRelateForm(request.POST,instance=notifications)
        f_form = PestReportFileFormSet(request.POST,  request.FILES,instance=pest_report)
        u_form =PestReportUrlFormSet(request.POST,  instance=pest_report)
         
      
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            """old pestreport save"""
            
            old_pest_report.pk = None
            old_pest_report.is_version = True
            old_pest_report.parent_id= id
            versions= PestReport.objects.filter( is_version=True, parent_id=id).count()
            slug1 = versions+1
      
            old_pest_report.slug= pest_report.slug+'-'+str(slug1)
        

            old_pest_report.save()
            for e in pest_report.pest_status.all():
                ps=PestStatus.objects.get(status=e)
                old_pest_report.pest_status.add(ps)
           
            issueformold =IssueKeywordsRelateForm()
            issue_instanceold=issueformold.save(commit=False)
            issue_instanceold.content_object = old_pest_report
            issue_instanceold.save()
            
            commformold =CommodityKeywordsRelateForm()
            comm_instanceold=commformold.save(commit=False)
            comm_instanceold.content_object = old_pest_report
            comm_instanceold.save()

            db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
            cursor = db.cursor()
            
            files=PestReportFile.objects.filter(pestreport_id=pest_report.id)
            urls=PestReportUrl.objects.filter(pestreport_id=pest_report.id)
            for f in files:
                sql = "INSERT INTO ippc_pestreportfile(pestreport_id,description,file) VALUES ("+str(old_pest_report.id)+", '"+ugettext(f.description)+"', '"+PestReportFile.name(f)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for u in urls:
                sql = "INSERT INTO ippc_pestreporturl(pestreport_id,url_for_more_information) VALUES ("+str(old_pest_report.id)+", '"+str(u)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()   
            for iss in old_issue:
                sql = """INSERT INTO ippc_issuekeywordsrelate_issuename(issuekeywordsrelate_id,issuekeyword_id) VALUES ("""+str(issue_instanceold.id)+""", """+str(iss)+""")"""
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for com in old_comm:
                sql = """INSERT INTO ippc_commoditykeywordsrelate_commname(commoditykeywordsrelate_id,commoditykeyword_id) VALUES ("""+str(comm_instanceold.id)+""", """+str(com)+""")"""
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    #print("###################error ")
                    db.rollback()        
            db.close()
        
          
            
            form.save()
          
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = pest_report
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = pest_report
            commodity_instance.save()
            commodityform.save_m2m() 
            
            notify_instance = notifyrelateform.save(commit=False)
            notify_instance.content_object = pest_report
            notify_instance.new_or_updated = 'UPDATE'
            notify_instance.link = user_country_slug+'/pestreports/'+str(pest_report.publish_date.strftime("%Y"))+'/'+str(pest_report.publish_date.strftime("%m"))+'/'+pest_report.slug+'/'
            notify_instance.updated_last = datetime.today()
            notify_instance.save()
            notifyrelateform.save_m2m()
            
            f_form.instance = pest_report
            f_form.save()
            u_form.instance = pest_report
            u_form.save()
            db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
            cursor = db.cursor()
            sql = "UPDATE ippc_pestreport set to_verify=False, verified_date='"+str(datetime.today())+"',modify_date='"+str(datetime.today())+"' where  id="+str(id)
            aaa=''
            try:
                cursor.execute(sql)
                db.commit()
            except:
                db.rollback()
        
            db.close()
            send_notification_message(0,id,content_type,pest_report.title,user_country_slug+'/pestreports/'+str(pest_report.publish_date.strftime("%Y"))+'/'+str(pest_report.publish_date.strftime("%m"))+'/'+pest_report.slug+'/')
            send_report_notification_message(0,0,id,content_type,pest_report.title,user_country_slug+'/pestreports/'+str(pest_report.publish_date.strftime("%Y"))+'/'+str(pest_report.publish_date.strftime("%m"))+'/'+pest_report.slug+'/')
            
            # If the save was successful, success message and redirect to another page
            # info(request, _("Successfully updated pest report."))
            if pest_report.status == CONTENT_STATUS_DRAFT:
                return redirect("pest-report-hidden-list", country=user_country_slug)
            else:
                return redirect("pest-report-detail",country=user_country_slug, year=pest_report.publish_date.strftime("%Y"), month=pest_report.publish_date.strftime("%m"), slug=pest_report.slug)

    else:
        form = PestReportForm(instance=pest_report)
        if pest_report.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=pest_report.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm( instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
        if pest_report.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=pest_report.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm( instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm()
        notifyrelateform =NotificationMessageRelateForm(instance=notifications)
        f_form = PestReportFileFormSet(instance=pest_report)
        u_form = PestReportUrlFormSet(instance=pest_report)
    return render_to_response(template_name, {
        'form': form,'f_form':f_form,'u_form':u_form,'issueform': issueform,'commodityform': commodityform,  "pest_report": pest_report,'notifyrelateform':notifyrelateform
    }, context_instance=RequestContext(request))

#
## http://stackoverflow.com/a/1854453/412329
@login_required
@permission_required('ippc.change_pestreport', login_url="/accounts/login/")
def pest_report_validate(request, country, id=None):
    """ VALIDATE PestReport  """
    user = request.user
    author = user
    country = user.get_profile().country
    # country_id = PestReport.objects.filter(country__country_id=country.id)
    user_country_slug = lower(slugify(country))
    if id:
        pest_report = get_object_or_404(PestReport, country=country, pk=id)
        db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
        cursor = db.cursor()
        sql = "UPDATE ippc_pestreport set to_verify=False, verified_date='"+str(datetime.today())+"',modify_date='"+str(datetime.today())+"' where  id="+str(pest_report.id)
        #print(sql)
        try:
            cursor.execute(sql)
            db.commit()
        except:
           # print('dddddddddddddddddddddddddddddddddddddddddddddddddddddddd')
            db.rollback()
        db.close()            
        info(request, _("Successfully Validated Pest report."))
        return redirect("pest-report-detail", country=user_country_slug, year=pest_report.publish_date.strftime("%Y"), month=pest_report.publish_date.strftime("%m"), slug=pest_report.slug)


class ReportingObligationListView(ListView):
    """    Reporting Obligation """
    context_object_name = 'latest'
    model = ReportingObligation
    date_field = 'publish_date'
    template_name = 'countries/reporting_obligation_list.html'
    queryset = ReportingObligation.objects.all().order_by('-modify_date', 'title')
    
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return pest reports from the specific country """
        # self.country = get_object_or_404(CountryPage, country=self.kwargs['country'])
        self.country = self.kwargs['country']
        self.type = self.kwargs['type']
        # CountryPage country_slug == country URL parameter keyword argument
        return ReportingObligation.objects.filter(country__country_slug=self.country,reporting_obligation_type=self.kwargs['type'],is_version=False)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(ReportingObligationListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['basic_types'] =BASIC_REP_TYPE_CHOICES
        context['current_type'] =int(self.kwargs['type'])
        return context
   
   
 

   
   
   
   
class IppcUserProfileDetailView(DetailView):
    """  Profile """
    model = IppcUserProfile
    context_object_name = 'user'
    template_name = 'accounts/account_profile.html'
    queryset = IppcUserProfile.objects.filter()


@login_required
def profile_update(request, template="accounts/account_profile_update.html"):
    """
    Profile update form.
    """
    profile_form = get_profile_form()
    form = profile_form(request.POST or None, request.FILES or None,
                        instance=request.user)
    if request.method == "POST" and form.is_valid():
        user = form.save()
        info(request, _("Profile updated"))
        try:
            return redirect("profile", username=user.username)
        except NoReverseMatch:
            return redirect("profile_update")
    context = {"form": form, "title": _("Update Profile")}
    return render(request, template, context)

# http://stackoverflow.com/a/1854453/412329
# @login_required
# def profile_update(request ,id=None, template_name='accounts/account_profile_update.html'):
#     """ Edit Profile """
#     user = request.user
#
#     if id:
#         profile = get_object_or_404(IppcUserProfile, pk=id)
#         userprofile = get_object_or_404(User, pk=profile.user_id)
#     if request.POST:
#         form = IppcUserProfileForm(request.POST,   instance=profile)
#         userform = UserForm(request.POST, instance=request.user)
#         if form.is_valid() and userform.is_valid():
#             form.save()
#             userform.save()
#             return redirect("user-detail",id)
#     else:
#         form = IppcUserProfileForm(instance=profile)
#         userform = UserForm(request.POST, instance=request.user)
#
#     return render_to_response(template_name, {
#         'form': form, 'userform': userform,'email':request.user.email,
#     }, context_instance=RequestContext(request))
       
   
class ReportingObligationDetailView(DetailView):
    """  Reporting Obligation detail page """
    model = ReportingObligation
    context_object_name = 'reportingobligation'
    template_name = 'countries/reporting_obligation_detail.html'
    queryset = ReportingObligation.objects.filter()
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(ReportingObligationDetailView, self).get_context_data(**kwargs)
        p = get_object_or_404(ReportingObligation, slug=self.kwargs['slug'])
        context['8col'] = 1
        
        versions= ReportingObligation.objects.filter(country__country_slug=self.kwargs['country'], is_version=True, parent_id=p.id).order_by('-modify_date')
        context['versions'] = versions
      
        return context

@login_required
@permission_required('ippc.add_reportingobligation', login_url="/accounts/login/")
def reporting_obligation_create(request, country,type):
    """ Create Reporting Obligation """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))

    form = ReportingObligationForm(request.POST, request.FILES)
    issueform =IssueKeywordsRelateForm(request.POST)
    commodityform =CommodityKeywordsRelateForm(request.POST)
    notifyrelateform =NotificationMessageRelateForm(request.POST)
     
    if request.method == "POST":
         f_form = ReportingoblicationFileFormSet(request.POST, request.FILES)
         u_form = ReportingObligationUrlFormSet(request.POST)
         if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            new_reporting_obligation = form.save(commit=False)
            new_reporting_obligation.author = request.user
            new_reporting_obligation.author_id = author.id
            new_reporting_obligation.report_obligation_type = type
            form.save()
            
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = new_reporting_obligation
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = new_reporting_obligation
            commodity_instance.save()
            commodityform.save_m2m()
            
            notify_instance = notifyrelateform.save(commit=False)
            notify_instance.content_object = new_reporting_obligation
            notify_instance.new_or_updated = 'NEW'
            notify_instance.link = user_country_slug+'/reportingobligation/'+str(new_reporting_obligation.publish_date.strftime("%Y"))+'/'+str(new_reporting_obligation.publish_date.strftime("%m"))+'/'+new_reporting_obligation.slug+'/'
            notify_instance.updated_last = new_reporting_obligation.publish_date
            notify_instance.save()
            notifyrelateform.save_m2m()
            
            f_form.instance = new_reporting_obligation
            f_form.save()
            u_form.instance = new_reporting_obligation
            u_form.save()
            
            content_type = ContentType.objects.get_for_model(new_reporting_obligation)
       
            send_notification_message(1,new_reporting_obligation.id,content_type,new_reporting_obligation.title,user_country_slug+'/reportingobligation/'+str(new_reporting_obligation.publish_date.strftime("%Y"))+'/'+str(new_reporting_obligation.publish_date.strftime("%m"))+'/'+new_reporting_obligation.slug+'/')
            send_report_notification_message(1,1,new_reporting_obligation.id,content_type,new_reporting_obligation.title,user_country_slug+'/reportingobligation/'+str(new_reporting_obligation.publish_date.strftime("%Y"))+'/'+str(new_reporting_obligation.publish_date.strftime("%m"))+'/'+new_reporting_obligation.slug+'/')
            
            info(request, _("Successfully created Reporting obligation."))
            return redirect("reporting-obligation-detail", country=user_country_slug, year=new_reporting_obligation.publish_date.strftime("%Y"), month=new_reporting_obligation.publish_date.strftime("%m"), slug=new_reporting_obligation.slug)
         else:
            return render_to_response('countries/reporting_obligation_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform,'notifyrelateform':notifyrelateform},
             context_instance=RequestContext(request))
    else:
        form = ReportingObligationForm(initial={'country': country,'reporting_obligation_type': type}, instance=ReportingObligation())
        issueform =IssueKeywordsRelateForm(request.POST)
        commodityform =CommodityKeywordsRelateForm(request.POST)
        notifyrelateform =NotificationMessageRelateForm(request.POST)
        f_form =ReportingoblicationFileFormSet()
        u_form =ReportingObligationUrlFormSet()

    return render_to_response('countries/reporting_obligation_create.html', {'form': form  ,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform,'notifyrelateform':notifyrelateform},
        context_instance=RequestContext(request))
        
        
        
# http://stackoverflow.com/a/1854453/412329
@login_required
@permission_required('ippc.change_reportingobligation', login_url="/accounts/login/")
def reporting_obligation_edit(request, country, id=None, template_name='countries/reporting_obligation_edit.html'):
    """ Edit Reporting Obligation """
    user = request.user
    author = user
    country = user.get_profile().country
    # country_id = PestReport.objects.filter(country__country_id=country.id)
    user_country_slug = lower(slugify(country))
    if id:
        reporting_obligation = get_object_or_404(ReportingObligation, country=country, pk=id)
        old_reporting_obligation=get_object_or_404(ReportingObligation, country=country, pk=id)
       
        content_type = ContentType.objects.get_for_model(reporting_obligation)
        try:
            notifications = get_object_or_404(NotificationMessageRelate, object_id=id,content_type__pk=content_type.id)
        except:
            notifications = None
    else:
        reporting_obligation = ReportingObligation(author=request.user)
    old_issue=[]
    if reporting_obligation.issuename.count()>0:
        for e in reporting_obligation.issuename.all():
                obj_i=e.content_object.issuename
                for o in obj_i.all():
                    for iss in o.issuename.all():
                        old_issue.append(iss.id)
    old_comm=[]
    if reporting_obligation.commname.count()>0:
        for e in reporting_obligation.commname.all():
                obj_c=e.content_object.commname
                for o in obj_c.all():
                    for com in o.commname.all():
                        old_comm.append(com.id)    
    if request.POST:
        form = ReportingObligationForm(request.POST, request.FILES, instance=reporting_obligation)
        if reporting_obligation.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=reporting_obligation.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
        if reporting_obligation.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=reporting_obligation.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm(request.POST,instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm(request.POST)
        notifyrelateform =NotificationMessageRelateForm(request.POST,instance=notifications)
        f_form = ReportingoblicationFileFormSet(request.POST,  request.FILES,instance=reporting_obligation)
        u_form =ReportingObligationUrlFormSet(request.POST,instance=reporting_obligation)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            old_reporting_obligation.pk = None
            old_reporting_obligation.is_version = True
            old_reporting_obligation.parent_id= id
            versions= ReportingObligation.objects.filter( is_version=True, parent_id=id).count()
            slug1 = versions+1
      
            old_reporting_obligation.slug= reporting_obligation.slug+'-'+str(slug1)
            old_reporting_obligation.save()
            
            issueformold =IssueKeywordsRelateForm()
            issue_instanceold=issueformold.save(commit=False)
            issue_instanceold.content_object = old_reporting_obligation
            issue_instanceold.save()
            
            commformold =CommodityKeywordsRelateForm()
            comm_instanceold=commformold.save(commit=False)
            comm_instanceold.content_object = old_reporting_obligation
            comm_instanceold.save()

            db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
            cursor = db.cursor()
            
            files=ReportingObligation_File.objects.filter(reportingobligation_id=reporting_obligation.id)
            urls=ReportingObligationUrl.objects.filter(reportingobligation_id=reporting_obligation.id)
            for f in files:
                sql = "INSERT INTO ippc_reportingobligation_file(reportingobligation_id,description,file) VALUES ("+str(old_reporting_obligation.id)+", '"+ugettext(f.description)+"', '"+ReportingObligation_File.name(f)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for u in urls:
                sql = "INSERT INTO ippc_reportingobligationurl(reportingobligation_id,url_for_more_information) VALUES ("+str(old_reporting_obligation.id)+", '"+str(u)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()   
            for iss in old_issue:
                sql = """INSERT INTO ippc_issuekeywordsrelate_issuename(issuekeywordsrelate_id,issuekeyword_id) VALUES ("""+str(issue_instanceold.id)+""", """+str(iss)+""")"""
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for com in old_comm:
                sql = """INSERT INTO ippc_commoditykeywordsrelate_commname(commoditykeywordsrelate_id,commoditykeyword_id) VALUES ("""+str(comm_instanceold.id)+""", """+str(com)+""")"""
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    #print("###################error ")
                    db.rollback() 
         
        
            db.close()
            
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = reporting_obligation
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = reporting_obligation
            commodity_instance.save()
            commodityform.save_m2m() 
            
            notify_instance = notifyrelateform.save(commit=False)
            notify_instance.content_object = reporting_obligation
            notify_instance.new_or_updated = 'UPDATE'
            notify_instance.link = user_country_slug+'/reportingobligation/'+str(reporting_obligation.publish_date.strftime("%Y"))+'/'+str(reporting_obligation.publish_date.strftime("%m"))+'/'+reporting_obligation.slug+'/'
            
            notify_instance.updated_last = datetime.today()
            notify_instance.save()
            notifyrelateform.save_m2m()
            
            f_form.instance = reporting_obligation
            f_form.save()
            u_form.instance = reporting_obligation
            u_form.save()
            # If the save was successful, success message and redirect to another page
            db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
            cursor = db.cursor()
            sql = "UPDATE ippc_reportingobligation set to_verify=False, verified_date='"+str(datetime.today())+"',modify_date='"+str(datetime.today())+"' where  id="+str(id)
            aaa=''
            try:
                cursor.execute(sql)
                db.commit()
            except:
                db.rollback()
        
            db.close()
            send_notification_message(0,id,content_type,reporting_obligation.title,user_country_slug+'/reportingobligation/'+str(reporting_obligation.publish_date.strftime("%Y"))+'/'+str(reporting_obligation.publish_date.strftime("%m"))+'/'+reporting_obligation.slug+'/')
            send_report_notification_message(0,1,reporting_obligation.id,content_type,reporting_obligation.title,user_country_slug+'/reportingobligation/'+str(reporting_obligation.publish_date.strftime("%Y"))+'/'+str(reporting_obligation.publish_date.strftime("%m"))+'/'+reporting_obligation.slug+'/')
            
            info(request, _("Successfully updated Reporting obligation."))
            return redirect("reporting-obligation-detail", country=user_country_slug, year=reporting_obligation.publish_date.strftime("%Y"), month=reporting_obligation.publish_date.strftime("%m"), slug=reporting_obligation.slug)

    else:
        form = ReportingObligationForm(instance=reporting_obligation)
        if reporting_obligation.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=reporting_obligation.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm( instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
        if reporting_obligation.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=reporting_obligation.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm( instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm()
        notifyrelateform =NotificationMessageRelateForm(instance=notifications)
        f_form = ReportingoblicationFileFormSet(instance=reporting_obligation)
        u_form = ReportingObligationUrlFormSet(instance=reporting_obligation)
    return render_to_response(template_name, {
        'form': form,'f_form':f_form,'u_form': u_form,'issueform': issueform,'commodityform': commodityform,  "reporting_obligation": reporting_obligation,'notifyrelateform':notifyrelateform
    }, context_instance=RequestContext(request))

## http://stackoverflow.com/a/1854453/412329
#@login_required
#@permission_required('ippc.change_reportingobligation', login_url="/accounts/login/")
#def reporting_obligation_translate(request,lang, country, id=None, template_name='countries/reporting_obligation_translate.html'):
#    """ translate Reporting Obligation """
#    user = request.user
#    author = user
#    country = user.get_profile().country
#    # country_id = PestReport.objects.filter(country__country_id=country.id)
#    user_country_slug = lower(slugify(country))
#    if id:
#        reporting_obligation = get_object_or_404(ReportingObligation, id=id)
#        try:
#            treporting_obligation = get_object_or_404(TransReportingObligation, translation_id=id,lang=lang)
#        except:
#            treporting_obligation = TransReportingObligation(lang=lang)        
#    else:
#            treporting_obligation = TransReportingObligation()
#            reporting_obligation = ReportingObligation()
#
#    if request.POST:
#        transform = TransReportingObligationForm(request.POST, instance=treporting_obligation)
#        if transform.is_valid():
#            transform.save()
#            info(request, _("Successfully translated Reporting obligation."))
#            return redirect("reporting-obligation-detail", country=user_country_slug, year=reporting_obligation.publish_date.strftime("%Y"), month=reporting_obligation.publish_date.strftime("%m"), slug=reporting_obligation.slug)
#    else:
#        transform = TransReportingObligationForm(instance=treporting_obligation)
#        
#    return render_to_response(template_name, {
#        'transform':transform,'lang':lang,"reporting_obligation":reporting_obligation,"treporting_obligation": treporting_obligation,
#    }, context_instance=RequestContext(request))
#
#
#
## http://stackoverflow.com/a/1854453/412329
@login_required
@permission_required('ippc.change_reportingobligation', login_url="/accounts/login/")
def reporting_obligation_validate(request, country, id=None):
    """ VALIDATE Reporting Obligation """
    user = request.user
    author = user
    country = user.get_profile().country
    # country_id = PestReport.objects.filter(country__country_id=country.id)
    user_country_slug = lower(slugify(country))
    #print(id)
    if id:
        reporting_obligation = get_object_or_404(ReportingObligation, country=country, pk=id)
        db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
        cursor = db.cursor()
        sql = "UPDATE ippc_reportingobligation set to_verify=False, verified_date='"+str(datetime.today())+"',modify_date='"+str(datetime.today())+"' where  id="+str(reporting_obligation.id)
       # OR MAYBE SHOULD SET THE LASTUPDATED DATE!!!!
        try:
            cursor.execute(sql)
            db.commit()
        except:
            db.rollback()
        db.close()            
        info(request, _("Successfully Validated Reporting obligation."))
        return redirect("reporting-obligation-detail", country=user_country_slug, year=reporting_obligation.publish_date.strftime("%Y"), month=reporting_obligation.publish_date.strftime("%m"), slug=reporting_obligation.slug)

class EventReportingListView(ListView):
    """    Event Reporting """
    context_object_name = 'latest'
    model = EventReporting
    date_field = 'publish_date'
    template_name = 'countries/event_reporting_list.html'
    queryset = EventReporting.objects.all().order_by('-modify_date', 'title')
    
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return pest reports from the specific country """
        # self.country = get_object_or_404(CountryPage, country=self.kwargs['country'])
        self.country = self.kwargs['country']
        self.type = self.kwargs['type']
        # CountryPage country_slug == country URL parameter keyword argument
        return EventReporting.objects.filter(country__country_slug=self.country,event_rep_type=self.kwargs['type'],is_version=False)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(EventReportingListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['event_types'] =EVT_REP_TYPE_CHOICES
        context['current_type'] =int(self.kwargs['type'])
        return context
   
       
   
class EventReportingDetailView(DetailView):
    """ EventReporting detail page """
    model = EventReporting
    context_object_name = 'eventreporting'
    template_name = 'countries/eventreporting_detail.html'
    queryset = EventReporting.objects.filter()
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(EventReportingDetailView, self).get_context_data(**kwargs)
        p = get_object_or_404(EventReporting, slug=self.kwargs['slug'])
        context['8col'] = 1
        
        versions= EventReporting.objects.filter(country__country_slug=self.kwargs['country'], is_version=True, parent_id=p.id).order_by('-modify_date')
        context['versions'] = versions
      
        return context
      

@login_required
@permission_required('ippc.add_eventreporting', login_url="/accounts/login/")
def event_reporting_create(request, country,type):
    """ Create Event Reporting """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))


    form = EventReportingForm(request.POST or None, request.FILES)
    issueform =IssueKeywordsRelateForm(request.POST)
    commodityform =CommodityKeywordsRelateForm(request.POST)
    notifyrelateform =NotificationMessageRelateForm(request.POST)
    
    if request.method == "POST":
        f_form = EventreportingFileFormSet(request.POST, request.FILES)
        u_form = EventreportingUrlFormSet(request.POST)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            new_event_reporting = form.save(commit=False)
            new_event_reporting.author = request.user
            new_event_reporting.author_id = author.id
            new_event_reporting.event_rep_type = type
            form.save()
            
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = new_event_reporting
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = new_event_reporting
            commodity_instance.save()
            commodityform.save_m2m()      
            
            notification_instance = notifyrelateform.save(commit=False)
            notification_instance.content_object = new_event_reporting
            notification_instance.new_or_updated = 'NEW'
            notification_instance.link = user_country_slug+'/eventreporting/'+str(new_event_reporting.publish_date.strftime("%Y"))+'/'+str(new_event_reporting.publish_date.strftime("%m"))+'/'+new_event_reporting.slug+'/'
          
            notification_instance.updated_last = new_event_reporting.publish_date
            notification_instance.save()
            notifyrelateform.save_m2m()
            
            f_form.instance = new_event_reporting
            f_form.save()
            u_form.instance = new_event_reporting
            u_form.save()
            content_type = ContentType.objects.get_for_model(new_event_reporting)
            send_notification_message(1,new_event_reporting.id,content_type,new_event_reporting.title,user_country_slug+'/eventreporting/'+str(new_event_reporting.publish_date.strftime("%Y"))+'/'+str(new_event_reporting.publish_date.strftime("%m"))+'/'+new_event_reporting.slug+'/')
            send_report_notification_message(1,2,new_event_reporting.id,content_type,new_event_reporting.title,user_country_slug+'/eventreporting/'+str(new_event_reporting.publish_date.strftime("%Y"))+'/'+str(new_event_reporting.publish_date.strftime("%m"))+'/'+new_event_reporting.slug+'/')
            info(request, _("Successfully added Event reporting."))
            return redirect("event-reporting-detail", country=user_country_slug, year=new_event_reporting.publish_date.strftime("%Y"), month=new_event_reporting.publish_date.strftime("%m"), slug=new_event_reporting.slug)
        else:
            return render_to_response('countries/event_reporting_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform,'notifyrelateform': notifyrelateform,},#'docform':myformset,
             context_instance=RequestContext(request))

          
        
    else:
        form = EventReportingForm(initial={'country': country,'event_rep_type': type}, instance=EventReporting())
        issueform =IssueKeywordsRelateForm(request.POST)
        commodityform =CommodityKeywordsRelateForm(request.POST)
        notifyrelateform =NotificationMessageRelateForm(request.POST)
        f_form = EventreportingFileFormSet()
        u_form = EventreportingUrlFormSet()
    
    return render_to_response('countries/event_reporting_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform,'notifyrelateform': notifyrelateform,},
        context_instance=RequestContext(request))

        

@login_required
@permission_required('ippc.change_eventreporting', login_url="/accounts/login/")
def event_reporting_edit(request, country, id=None, template_name='countries/event_reporting_edit.html'):
    """ Edit  Reporting """
    user = request.user
    author = user
    country = user.get_profile().country
    # country_id = PestReport.objects.filter(country__country_id=country.id)
    user_country_slug = lower(slugify(country))
    if id:
        event_reporting = get_object_or_404(EventReporting, country=country, pk=id)
        old_event_reporting=get_object_or_404(EventReporting, country=country, pk=id)
        

        content_type = ContentType.objects.get_for_model(event_reporting)
        try:
            notifications = get_object_or_404(NotificationMessageRelate, object_id=id,content_type__pk=content_type.id)
        except:
            notifications = None
    else:
        event_reporting = EventReporting(author=request.user)
    old_issue=[]
    if event_reporting.issuename.count()>0:
        for e in event_reporting.issuename.all():
                obj_i=e.content_object.issuename
                for o in obj_i.all():
                    for iss in o.issuename.all():
                        old_issue.append(iss.id)
    old_comm=[]
    if event_reporting.commname.count()>0:
        for e in event_reporting.commname.all():
                obj_c=e.content_object.commname
                for o in obj_c.all():
                    for com in o.commname.all():
                        old_comm.append(com.id)    
                        
    if request.POST:
        form = EventReportingForm(request.POST,  request.FILES, instance=event_reporting)
        if event_reporting.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=event_reporting.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
        if event_reporting.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=event_reporting.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm(request.POST,instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm(request.POST)
            
        notifyrelateform =NotificationMessageRelateForm(request.POST,instance=notifications)
        f_form = EventreportingFileFormSet(request.POST,  request.FILES,instance=event_reporting)
        u_form = EventreportingUrlFormSet(request.POST,  instance=event_reporting)
      
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            old_event_reporting.pk = None
            old_event_reporting.is_version = True
            old_event_reporting.parent_id= id
            versions= EventReporting.objects.filter( is_version=True, parent_id=id).count()
            slug1 = versions+1
      
            old_event_reporting.slug= event_reporting.slug+'-'+str(slug1)
            
            old_event_reporting.save()
            
            issueformold =IssueKeywordsRelateForm()
            issue_instanceold=issueformold.save(commit=False)
            issue_instanceold.content_object = old_event_reporting
            issue_instanceold.save()
            
            commformold =CommodityKeywordsRelateForm()
            comm_instanceold=commformold.save(commit=False)
            comm_instanceold.content_object = old_event_reporting
            comm_instanceold.save()

            db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
            cursor = db.cursor()
            
            files=EventreportingFile.objects.filter(eventreporting_id=event_reporting.id)
            urls=EventreportingUrl.objects.filter(eventreporting_id=event_reporting.id)
            for f in files:
                sql = "INSERT INTO ippc_eventreportingfile(eventreporting_id,description,file) VALUES ("+str(old_event_reporting.id)+", '"+ugettext(f.description)+"', '"+EventreportingFile.name(f)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for u in urls:
                sql = "INSERT INTO ippc_eventreportingurl(eventreporting_id,url_for_more_information) VALUES ("+str(old_event_reporting.id)+", '"+str(u)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()   
            for iss in old_issue:
                sql = """INSERT INTO ippc_issuekeywordsrelate_issuename(issuekeywordsrelate_id,issuekeyword_id) VALUES ("""+str(issue_instanceold.id)+""", """+str(iss)+""")"""
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for com in old_comm:
                sql = """INSERT INTO ippc_commoditykeywordsrelate_commname(commoditykeywordsrelate_id,commoditykeyword_id) VALUES ("""+str(comm_instanceold.id)+""", """+str(com)+""")"""
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    #print("###################error ")
                    db.rollback()        
            db.close()
            
            
            
            
            
            
            
            
            
            
            
            
            
            
            
            
            
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = event_reporting
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = event_reporting
            commodity_instance.save()
            commodityform.save_m2m() 
       
            notification_instance = notifyrelateform.save(commit=False)
            notification_instance.content_object = event_reporting
            notification_instance.new_or_updated = 'UPDATE'
            notification_instance.link = user_country_slug+'/eventreporting/'+str(event_reporting.publish_date.strftime("%Y"))+'/'+str(event_reporting.publish_date.strftime("%m"))+'/'+event_reporting.slug+'/'
          
            notification_instance.updated_last = datetime.today()
            notification_instance.save()
            notifyrelateform.save_m2m()
            
            f_form.instance = event_reporting
            f_form.save()
            u_form.instance = event_reporting
            u_form.save()
            db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
            cursor = db.cursor()
            sql = "UPDATE ippc_eventreporting set to_verify=False, verified_date='"+str(datetime.today())+"',modify_date='"+str(datetime.today())+"' where  id="+str(id)
            aaa=''
            try:
                cursor.execute(sql)
                db.commit()
            except:
                db.rollback()
        
            db.close()
            
            send_notification_message(0,id,content_type,event_reporting.title,user_country_slug+'/eventreporting/'+str(event_reporting.publish_date.strftime("%Y"))+'/'+str(event_reporting.publish_date.strftime("%m"))+'/'+event_reporting.slug+'/')
            send_report_notification_message(0,2,event_reporting.id,content_type,event_reporting.title,user_country_slug+'/eventreporting/'+str(event_reporting.publish_date.strftime("%Y"))+'/'+str(event_reporting.publish_date.strftime("%m"))+'/'+event_reporting.slug+'/')
            info(request, _("Successfully updated Event reporting."))
            return redirect("event-reporting-detail", country=user_country_slug, year=event_reporting.publish_date.strftime("%Y"), month=event_reporting.publish_date.strftime("%m"), slug=event_reporting.slug)

    else:
        form = EventReportingForm(instance=event_reporting)
        if event_reporting.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=event_reporting.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm( instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
        if event_reporting.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=event_reporting.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm( instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm()
        notifyrelateform =NotificationMessageRelateForm(instance=notifications)
        f_form = EventreportingFileFormSet(instance=event_reporting)
        u_form = EventreportingUrlFormSet( instance=event_reporting)
      
    return render_to_response(template_name, {
        'form': form, 'f_form':f_form,'u_form': u_form,'issueform': issueform,  'commodityform': commodityform, "event_reporting": event_reporting,'notifyrelateform':notifyrelateform
    }, context_instance=RequestContext(request))
    

## http://stackoverflow.com/a/1854453/412329
@login_required
@permission_required('ippc.change_eventreporting', login_url="/accounts/login/")
def event_reporting_validate(request, country, id=None):
    """ VALIDATE Event reporting  """
    user = request.user
    author = user
    country = user.get_profile().country
    # country_id = PestReport.objects.filter(country__country_id=country.id)
    user_country_slug = lower(slugify(country))
    #print(id)
    if id:
        eventreporting = get_object_or_404(EventReporting, country=country, pk=id)
        db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
        cursor = db.cursor()
        sql = "UPDATE ippc_eventreporting set to_verify=False, verified_date='"+str(datetime.today())+"',modify_date='"+str(datetime.today())+"' where  id="+str(eventreporting.id)
        #print(sql)
        try:
            cursor.execute(sql)
            db.commit()
        except:
            #print('dddddddddddddddddddddddddddddddddddddddddddddddddddddddd')
            db.rollback()
        db.close()            
        info(request, _("Successfully Validated Event reporting."))
        return redirect("event-reporting-detail", country=user_country_slug, year=eventreporting.publish_date.strftime("%Y"), month=eventreporting.publish_date.strftime("%m"), slug=eventreporting.slug)

class DraftProtocolListView(ListView):
    """    DraftProtocol """
    context_object_name = 'latest'
    model = DraftProtocol
    date_field = 'publish_date'
    template_name = 'dp/dp_list.html'
    queryset = DraftProtocol.objects.all().order_by('-publish_date', 'title')
    
    allow_future = False
    allow_empty = True
    paginate_by = 50

class DraftProtocolDetailView(DetailView):
    """ DraftProtocol detail page """
    model = DraftProtocol
    context_object_name = 'draftprotocol'
    template_name = 'dp/dp_detail.html'
    queryset = DraftProtocol.objects.filter()

    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(DraftProtocolDetailView, self).get_context_data(**kwargs)
        draftprotocol = get_object_or_404(DraftProtocol,  slug=self.kwargs['slug'])
        queryset = DraftProtocolComments.objects.filter(draftprotocol_id=draftprotocol.id)
        add_comment= 1
        see_all_comment= 0
        
        for obj in queryset:
           if self.request.user == obj.author:
                  add_comment=0
        if self.request.user.groups.filter(name='IPPC Secretariat') or self.request.user.groups.filter(name='TPDPc') :
            see_all_comment=1
        
         
        context['comments'] =queryset
        context['see_all_comment'] =see_all_comment
        context['add_comment'] =add_comment
        return context

@login_required
@permission_required('ippc.add_publication', login_url="/accounts/login/")
def draftprotocol_create(request):
    """ Create DraftProtocol """
    user = request.user
    author = user


    form = DraftProtocolForm(request.POST or None, request.FILES)
     
    if request.method == "POST":
        f_form = DraftProtocolFileFormSet(request.POST, request.FILES)
        if form.is_valid() and f_form.is_valid() :
            new_draftprotocol = form.save(commit=False)
            new_draftprotocol.author = request.user
            new_draftprotocol.author_id = author.id
            form.save()
            
            f_form.instance = new_draftprotocol
            f_form.save()
           
            info(request, _("Successfully added Draft Protocol."))
            return redirect("draftprotocol-detail",  year=new_draftprotocol.publish_date.strftime("%Y"), month=new_draftprotocol.publish_date.strftime("%m"), slug=new_draftprotocol.slug)
        else:
            return render_to_response('dp/dp_create.html', {'form': form,'f_form': f_form,},#'entryform': entryform,'docform':myformset,
             context_instance=RequestContext(request))

          
        
    else:
        form = DraftProtocolForm(instance=DraftProtocol())
        f_form =DraftProtocolFileFormSet()
      
    return render_to_response('dp/dp_create.html', {'form': form,'f_form': f_form,},
        context_instance=RequestContext(request))

        

@login_required
@permission_required('ippc.add_publication', login_url="/accounts/login/")
def draftprotocol_edit(request, id=None, template_name='dp/dp_edit.html'):
    """ DraftProtocol """
    user = request.user
    author = user
    if id:
        draftprotocol = get_object_or_404(DraftProtocol,  pk=id)
    else:
        draftprotocol = DraftProtocol(author=request.user)
      
    if request.POST:
        form = DraftProtocolForm(request.POST,  request.FILES, instance=draftprotocol)
        f_form = DraftProtocolFileFormSet(request.POST,  request.FILES,instance=draftprotocol)

      
        if form.is_valid() and f_form.is_valid():
            form.save()
            f_form.instance = draftprotocol
            f_form.save()
            info(request, _("Successfully updated DraftProtocol."))
            return redirect("draftprotocol-detail", year=draftprotocol.publish_date.strftime("%Y"), month=draftprotocol.publish_date.strftime("%m"), slug=draftprotocol.slug)

    else:
        form = DraftProtocolForm(instance=draftprotocol)
        f_form = DraftProtocolFileFormSet(instance=draftprotocol)
      
      
    return render_to_response(template_name, {
        'form': form, 'f_form':f_form, "draftprotocol": draftprotocol
    }, context_instance=RequestContext(request))
    



@login_required
@permission_required('ippc.add_draftcomment', login_url="/accounts/login/")
def draftprotocol_comment_create(request, id=None):
    """ Create  draftprotocol comment"""
    user = request.user
    author = user
    draftprotocol = get_object_or_404(DraftProtocol,  pk=id)
    form = DraftProtocolCommentsForm(request.POST or None, request.FILES)
     
    if request.method == "POST":
        if form.is_valid() :
            new_draftprotocolComment = form.save(commit=False)
            new_draftprotocolComment.author = request.user
            new_draftprotocolComment.author_id = author.id
            new_draftprotocolComment.title = request.user
            new_draftprotocolComment.draftprotocol_id = id
          
            form.save()
            
            emailto_all=['Martin.Farren@fao.org','IPPC-DP@fao.org']
            subject='IPPC new comment on ECDP: '+str(draftprotocol.title)
            text=''
            text=str(request.user)+' has commented on: '+str(draftprotocol.title)  +'<br><hr><br>'+str(new_draftprotocolComment.comment)
            notifificationmessage = mail.EmailMessage(subject,text,'ippc@fao.org', emailto_all, ['paola.sentinelli@fao.org'])
            notifificationmessage.content_subtype = "html"
            #print('test-sending')
            sent =notifificationmessage.send()
            
            info(request, _("Successfully added Comment."))
            return redirect("draftprotocol-detail",  year=draftprotocol.publish_date.strftime("%Y"), month=draftprotocol.publish_date.strftime("%m"), slug=draftprotocol.slug)
        else:
            return render_to_response('dp/dp_comment_create.html', {'form': form,},
             context_instance=RequestContext(request))
    else:
        form = DraftProtocolCommentsForm(initial={'draftprotocol': id},instance=DraftProtocolComments())
      
    return render_to_response('dp/dp_comment_create.html', {'form': form,'draftprotocol': id,},
        context_instance=RequestContext(request))

        

@login_required
@permission_required('ippc.add_draftcomment', login_url="/accounts/login/")
def draftprotocol_comment_edit(request, id=None, dp_id=None, template_name='dp/dp_comment_edit.html'):
    """ DraftProtocol comment edit"""
    user = request.user
    author = user
    if id:
        draftprotocolcomment = get_object_or_404(DraftProtocolComments,  pk=id)
    else:
        draftprotocolcomment = DraftProtocolComments(author=request.user)
    
    draftprotocol = get_object_or_404(DraftProtocol,  pk=dp_id)
        
      
    if request.POST:
        form = DraftProtocolCommentsForm(request.POST,  request.FILES, instance=draftprotocolcomment)
        if form.is_valid() :
            form.save()
           # f_form.instance = draftprotocolcomment
          #  f_form.save()
            info(request, _("Successfully updated DraftProtocol."))
            return redirect("draftprotocol-detail", year=draftprotocol.publish_date.strftime("%Y"), month=draftprotocol.publish_date.strftime("%m"), slug=draftprotocol.slug)

    else:
        form = DraftProtocolCommentsForm(instance=draftprotocolcomment)
        
      
      
    return render_to_response(template_name, {
        'form': form,  "draftprotocolcomment": draftprotocolcomment
    }, context_instance=RequestContext(request))
    



class WebsiteListView(ListView):
    """ Website """
    context_object_name = 'latest'
    model = Website
    date_field = 'publish_date'
    template_name = 'countries/website_list.html'
    queryset = Website.objects.all().order_by('-modify_date', 'title')
    
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Website from the specific country """
        # self.country = get_object_or_404(CountryPage, country=self.kwargs['country'])
        self.country = self.kwargs['country']
        # CountryPage country_slug == country URL parameter keyword argument
        return Website.objects.filter(country__country_slug=self.country)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(WebsiteListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        return context
   
       
   
class WebsiteDetailView(DetailView):
    """ EventReporting detail page """
    model = Website
    context_object_name = 'website'
    template_name = 'countries/website_detail.html'
    queryset = Website.objects.filter()



@login_required
@permission_required('ippc.add_website', login_url="/accounts/login/")
def website_create(request, country):
    """ Create website """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))

    form = WebsiteForm(request.POST or None, request.FILES)
    issueform =IssueKeywordsRelateForm(request.POST)
    commodityform =CommodityKeywordsRelateForm(request.POST)
    
         
    if request.method == "POST":
        u_form =WebsiteUrlFormSet(request.POST)
        if form.is_valid() and u_form.is_valid():
            new_website = form.save(commit=False)
            new_website.author = request.user
            new_website.author_id = author.id
            form.save()
            
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = new_website
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = new_website
            commodity_instance.save()
            commodityform.save_m2m()      
              
          
            u_form.instance = new_website
            u_form.save()
            info(request, _("Successfully added Website."))
            return redirect("website-detail", country=user_country_slug, year=new_website.publish_date.strftime("%Y"), month=new_website.publish_date.strftime("%m"), slug=new_website.slug)
        else:
            return render_to_response('countries/website_create.html', {'form': form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform},#'entryform': entryform,'docform':myformset,
             context_instance=RequestContext(request))

          
        
    else:
        form = WebsiteForm(initial={'country': country}, instance=Website())
        issueform =IssueKeywordsRelateForm(request.POST)
        commodityform =CommodityKeywordsRelateForm(request.POST)
        u_form = WebsiteUrlFormSet()
    
    return render_to_response('countries/website_create.html', {'form': form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform},
        context_instance=RequestContext(request))

        

@login_required
@permission_required('ippc.change_website', login_url="/accounts/login/")
def website_edit(request, country, id=None, template_name='countries/website_edit.html'):
    """ Edit  website """
    user = request.user
    author = user
    country = user.get_profile().country
    # country_id = PestReport.objects.filter(country__country_id=country.id)
    user_country_slug = lower(slugify(country))
    if id:
        website = get_object_or_404(Website, country=country, pk=id)
     
    else:
        website = Website(author=request.user)
      
    if request.POST:
        form = WebsiteForm(request.POST,  request.FILES, instance=website)
        if website.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=website.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
        if website.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=website.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm(request.POST,instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm(request.POST)
        u_form = WebsiteUrlFormSet(request.POST,  instance=website)
      
        if form.is_valid()  and u_form.is_valid():
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = website
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = website
            commodity_instance.save()
            commodityform.save_m2m() 
    
           
            u_form.instance = website
            u_form.save()
            info(request, _("Successfully updated Website."))
            return redirect("website-detail", country=user_country_slug, year=website.publish_date.strftime("%Y"), month=website.publish_date.strftime("%m"), slug=website.slug)

    else:
        form = WebsiteForm(instance=website)
        if website.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=website.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm( instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
        if website.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=website.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm( instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm()
        u_form = WebsiteUrlFormSet( instance=website)
      
    return render_to_response(template_name, {
        'form': form, 'u_form': u_form,'issueform': issueform,  'commodityform': commodityform, "website": website
    }, context_instance=RequestContext(request))
    







class PartnersWebsiteDetailView(DetailView):
    """ EventReporting detail page """
    model = PartnersWebsite
    context_object_name = 'website'
    template_name = 'partners/website_detail.html'
    queryset = PartnersWebsite.objects.filter()

    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PartnersWebsiteDetailView, self).get_context_data(**kwargs)
        page = get_object_or_404(PartnersPage, name=self.kwargs['partners'])
        context['pagetitle'] =  page.name
        context['pageslug'] =  page.slug
       # context['page'] =  page.partner_slug
        return context
     

@login_required
@permission_required('ippc.add_partnerswebsite', login_url="/accounts/login/")
def partner_websites_create(request, partners):
    """ Create website """
    user = request.user
    author = user
    if user.get_profile().partner:
        partners=user.get_profile().partner
         
    user_partner_slug = lower(slugify(partners))
    #partners=user.get_profile().partner
    page = get_object_or_404(PartnersPage, name=user_partner_slug)
    partners=page.id 

    form = PartnersWebsiteForm(request.POST or None, request.FILES)
    issueform =IssueKeywordsRelateForm(request.POST)
    commodityform =CommodityKeywordsRelateForm(request.POST)
    
         
    if request.method == "POST":
        u_form =PartnersWebsiteUrlFormSet(request.POST)
        if form.is_valid() and u_form.is_valid():
            new_website = form.save(commit=False)
            new_website.author = request.user
            new_website.author_id = author.id
            form.save()
            
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = new_website
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = new_website
            commodity_instance.save()
            commodityform.save_m2m()      
              
          
            u_form.instance = new_website
            u_form.save()
            info(request, _("Successfully added Website."))
            return redirect("partner-websites-detail", partners=user_partner_slug, year=new_website.publish_date.strftime("%Y"), month=new_website.publish_date.strftime("%m"), slug=new_website.slug)
        else:
            return render_to_response('partners/website_create.html', {'form': form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform},#'entryform': entryform,'docform':myformset,
             context_instance=RequestContext(request))

          
        
    else:
        form = PartnersWebsiteForm(initial={'partners': partners}, instance=PartnersWebsite())
        issueform =IssueKeywordsRelateForm(request.POST)
        commodityform =CommodityKeywordsRelateForm(request.POST)
        u_form = PartnersWebsiteUrlFormSet()
    
    return render_to_response('partners/website_create.html', {'form': form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform},
        context_instance=RequestContext(request))

        

@login_required
@permission_required('ippc.change_partnerswebsite', login_url="/accounts/login/")
def partner_websites_edit(request, partners, id=None, template_name='partners/website_edit.html'):
    """ Edit  website """
    user = request.user
    author = user
    if user.get_profile().partner:
        partners=user.get_profile().partner
         
    user_partner_slug = lower(slugify(partners))
    #partners=user.get_profile().partner
    page = get_object_or_404(PartnersPage, name=user_partner_slug)
    partners=page.id 
    if id:
        website = get_object_or_404(PartnersWebsite, partners=partners, pk=id)
    else:
        website = PartnersWebsite(author=request.user)
      
    if request.POST:
        form = PartnersWebsiteForm(request.POST,  request.FILES, instance=website)
        if website.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=website.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
        if website.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=website.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm(request.POST,instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm(request.POST)
        
        u_form = PartnersWebsiteUrlFormSet(request.POST,  instance=website)
        
        if form.is_valid()  and u_form.is_valid():
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = website
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = website
            commodity_instance.save()
            commodityform.save_m2m() 
    
           
            u_form.instance = website
            u_form.save()
            info(request, _("Successfully updated Website."))
            return redirect("partner-websites-detail", partners=user_partner_slug, year=website.publish_date.strftime("%Y"), month=website.publish_date.strftime("%m"), slug=website.slug)

    else:
        form = PartnersWebsiteForm(instance=website)
        if website.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=website.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm( instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
        if website.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=website.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm( instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm()
        u_form =PartnersWebsiteUrlFormSet( instance=website)
      
    return render_to_response(template_name, {
        'form': form, 'u_form': u_form,'issueform': issueform,  'commodityform': commodityform, "website": website
    }, context_instance=RequestContext(request))
    








            
class CnPublicationListView(ListView):
    """   Contry Publication """
    context_object_name = 'latest'
    model = CnPublication
    date_field = 'publish_date'
    template_name = 'countries/cnpublication_list.html'
    queryset = CnPublication.objects.all().order_by('-modify_date', 'title')
    
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return pest reports from the specific country """
        # self.country = get_object_or_404(CountryPage, country=self.kwargs['country'])
        self.country = self.kwargs['country']
        # CountryPage country_slug == country URL parameter keyword argument
        return CnPublication.objects.filter(country__country_slug=self.country)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CnPublicationListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        return context
   
       
   
class CnPublicationDetailView(DetailView):
    """ Country Publication detail page """
    model = CnPublication
    context_object_name = 'cnpublication'
    template_name = 'countries/cnpublication_detail.html'
    queryset = CnPublication.objects.filter()



@login_required
@permission_required('ippc.add_cnpublication', login_url="/accounts/login/")
def country_publication_create(request, country):
    """ Create  Country Publication """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))

    form = CnPublicationForm(request.POST or None, request.FILES)
    issueform =IssueKeywordsRelateForm(request.POST)
    commodityform =CommodityKeywordsRelateForm(request.POST)
    
         
    if request.method == "POST":
        f_form = CnPublicationFileFormSet(request.POST, request.FILES)
        u_form = CnPublicationUrlFormSet(request.POST)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            new_cnpublication = form.save(commit=False)
            new_cnpublication.author = request.user
            new_cnpublication.author_id = author.id
            form.save()
            
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = new_cnpublication
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = new_cnpublication
            commodity_instance.save()
            commodityform.save_m2m()      
              
            f_form.instance = new_cnpublication
            f_form.save()
            u_form.instance = new_cnpublication
            u_form.save()
            info(request, _("Successfully added publication."))
            return redirect("country-publication-detail", country=user_country_slug, year=new_cnpublication.publish_date.strftime("%Y"), month=new_cnpublication.publish_date.strftime("%m"), slug=new_cnpublication.slug)
        else:
            return render_to_response('countries/cnpublication_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform},#'entryform': entryform,'docform':myformset,
             context_instance=RequestContext(request))

          
        
    else:
        form = CnPublicationForm(initial={'country': country}, instance=CnPublication())
        issueform =IssueKeywordsRelateForm(request.POST)
        commodityform =CommodityKeywordsRelateForm(request.POST)
        f_form = CnPublicationFileFormSet()
        u_form = CnPublicationUrlFormSet()
    
    return render_to_response('countries/cnpublication_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform},
        context_instance=RequestContext(request))

        

@login_required
@permission_required('ippc.change_cnpublication', login_url="/accounts/login/")
def country_publication_edit(request, country, id=None, template_name='countries/cnpublication_edit.html'):
    """ Edit   Country Publication """
    user = request.user
    author = user
    country = user.get_profile().country
    # country_id = PestReport.objects.filter(country__country_id=country.id)
    user_country_slug = lower(slugify(country))
    if id:
        cnpublication = get_object_or_404(CnPublication, country=country, pk=id)
    else:
        cnpublication = CnPublication(author=request.user)
      
    if request.POST:
        form = CnPublicationForm(request.POST,  request.FILES, instance=cnpublication)
        if cnpublication.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=cnpublication.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
        if cnpublication.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=cnpublication.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm(request.POST,instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm(request.POST)
        
        f_form = CnPublicationFileFormSet(request.POST,  request.FILES,instance=cnpublication)
        u_form = CnPublicationUrlFormSet(request.POST,  instance=cnpublication)
      
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = cnpublication
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = cnpublication
            commodity_instance.save()
            commodityform.save_m2m() 
    
            f_form.instance = cnpublication
            f_form.save()
            u_form.instance = cnpublication
            u_form.save()
            info(request, _("Successfully updated publication."))
            return redirect("country-publication-detail", country=user_country_slug, year=cnpublication.publish_date.strftime("%Y"), month=cnpublication.publish_date.strftime("%m"), slug=cnpublication.slug)

    else:
        form = CnPublicationForm(instance=cnpublication)
        if cnpublication.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=cnpublication.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm( instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
        if cnpublication.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=cnpublication.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm( instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm()
        f_form = CnPublicationFileFormSet(instance=cnpublication)
        u_form = CnPublicationUrlFormSet( instance=cnpublication)
      
    return render_to_response(template_name, {
        'form': form, 'f_form':f_form,'u_form': u_form,'issueform': issueform,  'commodityform': commodityform, "cnpublication": cnpublication
    }, context_instance=RequestContext(request))            
            



class PartnersView(TemplateView):
    """ 
    Individual Partners homepage 
    """
    template_name = 'partners/partners_page.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(TemplateView, self).get_context_data(**kwargs)
        context.update({
            'partner': self.kwargs['partner']
            # 'editors': self.kwargs['editors']
            # 'profile_user': self.kwargs['profile_user']
        })
        page = get_object_or_404(PartnersPage, name=self.kwargs['partner'])
        pageparent = get_object_or_404(PublicationLibrary, id=page.parent_id)
        titleparent=pageparent.title
        titleparent = titleparent.replace(" ", "-").lower()
        context['content']  =page.content
        context['edituser']  =page.edituser
        context['modify_date']  =page.modify_date
       
          
        context['titleparent']  =pageparent.title
        context['titleparentslug'] = pageparent.slug
        context['pageparentid'] = pageparent.id
       
        context['publications'] = PartnersPublication.objects.filter(partners__partner_slug=self.kwargs['partner'],status=2)
        context['news'] = PartnersNews.objects.filter(partners__partner_slug=self.kwargs['partner'],status=2)
        context['websites'] = PartnersWebsite.objects.filter(partners__partner_slug=self.kwargs['partner'],status=2)
       
        return context
    
     
    
       
   
class PartnersPublicationDetailView(DetailView):
    """ Partner Publication detail page """
    model = PartnersPublication
    context_object_name = 'partnerspublication'
    template_name = 'partners/p_publication_detail.html'
    queryset = PartnersPublication.objects.filter()
   
      
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PartnersPublicationDetailView, self).get_context_data(**kwargs)
        page = get_object_or_404(PartnersPage, name=self.kwargs['partners'])
        context['pagetitle'] =  page.name
        context['pageslug'] =  page.slug
       # context['page'] =  page.partner_slug
        return context
     

            
@login_required
@permission_required('ippc.add_partnerspublication', login_url="/accounts/login/")
def partner_publication_create(request, partners):
    """ Create  partner Publication """
    user = request.user
    author = user
    if user.get_profile().partner:
        partners=user.get_profile().partner
         
    user_partner_slug = lower(slugify(partners))
    #partners=user.get_profile().partner
    page = get_object_or_404(PartnersPage, name=user_partner_slug)
    partners=page.id  
    #print("-----------------------------------------")
    #print(user_partner_slug)
    #print("-----------------------------------------")
    form = PartnersPublicationForm(request.POST or None, request.FILES)
    issueform =IssueKeywordsRelateForm(request.POST)
    commodityform =CommodityKeywordsRelateForm(request.POST)
    
         
    if request.method == "POST":
        f_form = PartnersPublicationFileFormSet(request.POST, request.FILES)
        u_form = PartnersPublicationUrlFormSet(request.POST)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            new_partnerpublication = form.save(commit=False)
            new_partnerpublication.author = request.user
            new_partnerpublication.author_id = author.id
            form.save()
            
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = new_partnerpublication
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = new_partnerpublication
            commodity_instance.save()
            commodityform.save_m2m()      
              
            f_form.instance = new_partnerpublication
            f_form.save()
            u_form.instance = new_partnerpublication
            u_form.save()
            info(request, _("Successfully added publication."))
            return redirect("partner-publication-detail", partners=user_partner_slug, year=new_partnerpublication.publish_date.strftime("%Y"), month=new_partnerpublication.publish_date.strftime("%m"), slug=new_partnerpublication.slug)
        else:
            return render_to_response('partners/p_publication_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform},#'entryform': entryform,'docform':myformset,
             context_instance=RequestContext(request))

          
        
    else:
        form = PartnersPublicationForm(initial={'partners': partners}, instance=PartnersPublication())
        issueform =IssueKeywordsRelateForm(request.POST)
        commodityform =CommodityKeywordsRelateForm(request.POST)
        f_form = PartnersPublicationFileFormSet()
        u_form = PartnersPublicationUrlFormSet()
    
    return render_to_response('partners/p_publication_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform},
        context_instance=RequestContext(request))

        

@login_required
@permission_required('ippc.change_partnerspublication', login_url="/accounts/login/")
def partner_publication_edit(request, partners, id=None, template_name='partners/p_publication_edit.html'):
    """ Edit   partners Publication """
    user = request.user
    author = user
    if user.get_profile().partner:
        partners=user.get_profile().partner
         
    user_partner_slug = lower(slugify(partners))
    #partners=user.get_profile().partner
    page = get_object_or_404(PartnersPage, name=user_partner_slug)
    partners=page.id  
    if id:
        partnerspublication = get_object_or_404(PartnersPublication, partners=partners, pk=id)
    else:
        partnerspublication = PartnersPublication(author=request.user)
      
    if request.POST:
        form =PartnersPublicationForm(request.POST,  request.FILES, instance=partnerspublication)
        if partnerspublication.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=partnerspublication.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
        if partnerspublication.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=partnerspublication.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm(request.POST,instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm(request.POST) 
        f_form = PartnersPublicationFileFormSet(request.POST,  request.FILES,instance=partnerspublication)
        u_form = PartnersPublicationUrlFormSet(request.POST,  instance=partnerspublication)
      
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = partnerspublication
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = partnerspublication
            commodity_instance.save()
            commodityform.save_m2m() 
    
            f_form.instance = partnerspublication
            f_form.save()
            u_form.instance = partnerspublication
            u_form.save()
            info(request, _("Successfully updated publication."))
            return redirect("partner-publication-detail", partners=user_partner_slug, year=partnerspublication.publish_date.strftime("%Y"), month=partnerspublication.publish_date.strftime("%m"), slug=partnerspublication.slug)

    else:
        form = PartnersPublicationForm(instance=partnerspublication)
        if partnerspublication.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=partnerspublication.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm( instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
        if partnerspublication.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=partnerspublication.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm( instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm()
        f_form = PartnersPublicationFileFormSet(instance=partnerspublication)
        u_form = PartnersPublicationUrlFormSet( instance=partnerspublication)
      
    return render_to_response(template_name, {
        'form': form, 'f_form':f_form,'u_form': u_form,'issueform': issueform,  'commodityform': commodityform, "partnerspublication": partnerspublication
    }, context_instance=RequestContext(request))            
                        
 
@login_required
@permission_required('ippc.change_partnerspublication', login_url="/accounts/login/")
def partner_page_edit(request, id=None, template_name='partners/partners_page_form.html'):
    """ Edit   partners page """
    ppage = get_object_or_404(PartnersPage, page_ptr_id=id)
    page= get_object_or_404(Page, id=id)
  
    
    
    if request.POST:
        form =PartnerPageForm(request.POST,   instance=ppage)
      
        if form.is_valid():
            ppage = form.save(commit=False)
            ppage.edituser = request.user.username
            ppage.modify_date=datetime.now()
    
            form.save()
        
            info(request, _("Successfully updated page."))
           
            return redirect('https://www.ippc.int/'+str(page.slug)  )
   
    else:
        form = PartnerPageForm(instance=ppage)
     
      
    return render_to_response(template_name, {
        'form': form, "partnerspage": ppage
    }, context_instance=RequestContext(request))            
                        
            
                       
            
class PestFreeAreaListView(ListView):
    """    Event Reporting """
    context_object_name = 'latest'
    model = PestFreeArea
    date_field = 'publish_date'
    template_name = 'countries/pfa_list.html'
    queryset = PestFreeArea.objects.all().order_by('-modify_date', 'title')
    
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return PestFreeArea from the specific country """
        # self.country = get_object_or_404(CountryPage, country=self.kwargs['country'])
        self.country = self.kwargs['country']
        # CountryPage country_slug == country URL parameter keyword argument
        return PestFreeArea.objects.filter(country__country_slug=self.country,is_version=False)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PestFreeAreaListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        return context
   
       
   
class PestFreeAreaDetailView(DetailView):
    """ PestFreeArea Detail page """
    model = PestFreeArea
    context_object_name = 'pfa'
    template_name = 'countries/pfa_detail.html'
    queryset = PestFreeArea.objects.filter()

    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PestFreeAreaDetailView, self).get_context_data(**kwargs)
        p = get_object_or_404(PestFreeArea, slug=self.kwargs['slug'])
        
        versions= PestFreeArea.objects.filter(country__country_slug=self.kwargs['country'], is_version=True, parent_id=p.id).order_by('-modify_date')
        context['versions'] = versions
      
        return context

@login_required
@permission_required('ippc.add_pestfreearea', login_url="/accounts/login/")
def pfa_create(request, country):
    """ Create PestFreeArea """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))


    form = PestFreeAreaForm(request.POST)
    issueform =IssueKeywordsRelateForm(request.POST)
    commodityform =CommodityKeywordsRelateForm(request.POST)
    notifyrelateform =NotificationMessageRelateForm(request.POST)
     
    if request.method == "POST":
         f_form = PestFreeAreaFileFormSet(request.POST, request.FILES)
         u_form = PestFreeAreaUrlFormSet(request.POST)

         if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            new_pfa = form.save(commit=False)
            new_pfa.author = request.user
            new_pfa.author_id = author.id
            form.save()
            
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = new_pfa
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = new_pfa
            commodity_instance.save()
            commodityform.save_m2m() 
            
            notify_instance = notifyrelateform.save(commit=False)
            notify_instance.content_object = new_pfa
            notify_instance.new_or_updated = 'NEW'
            notify_instance.link = user_country_slug+'/pestfreeareas/'+str(new_pfa.publish_date.strftime("%Y"))+'/'+str(new_pfa.publish_date.strftime("%m"))+'/'+new_pfa.slug+'/'
          
            notify_instance.updated_last = new_pfa.publish_date
            notify_instance.save()
            notifyrelateform.save_m2m()
            
            f_form.instance = new_pfa
            f_form.save()
            
            u_form.instance = new_pfa
            u_form.save()
            content_type = ContentType.objects.get_for_model(new_pfa)
            send_notification_message(1,new_pfa.id,content_type,new_pfa.title,user_country_slug+'/pestfreeareas/'+str(new_pfa.publish_date.strftime("%Y"))+'/'+str(new_pfa.publish_date.strftime("%m"))+'/'+new_pfa.slug+'/')
           
            info(request, _("Successfully created PestFreeArea."))
            
            return redirect("pfa-detail", country=user_country_slug, year=new_pfa.publish_date.strftime("%Y"), month=new_pfa.publish_date.strftime("%m"), slug=new_pfa.slug)
         else:
             return render_to_response('countries/pfa_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform,'notifyrelateform':notifyrelateform},
             context_instance=RequestContext(request))
    else:
        form = PestFreeAreaForm(initial={'country': country}, instance=PestFreeArea())
        issueform =IssueKeywordsRelateForm(request.POST)
        commodityform =CommodityKeywordsRelateForm(request.POST)
        notifyrelateform =NotificationMessageRelateForm(request.POST)
        f_form =PestFreeAreaFileFormSet()
        u_form = PestFreeAreaUrlFormSet()

    return render_to_response('countries/pfa_create.html', {'form': form  ,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform,'notifyrelateform':notifyrelateform},
        context_instance=RequestContext(request))


        
# http://stackoverflow.com/a/1854453/412329
@login_required
@permission_required('ippc.change_pestfreearea', login_url="/accounts/login/")
def pfa_edit(request, country, id=None, template_name='countries/pfa_edit.html'):
    """ Edit PestFreeArea """
    user = request.user
    author = user
    country = user.get_profile().country
    # country_id = PestReport.objects.filter(country__country_id=country.id)
    user_country_slug = lower(slugify(country))
    if id:
        pfa = get_object_or_404(PestFreeArea, country=country, pk=id)
        old_pfa=get_object_or_404(PestFreeArea, country=country, pk=id)
       

        content_type = ContentType.objects.get_for_model(pfa)
        try:
            notifications = get_object_or_404(NotificationMessageRelate, object_id=id,content_type__pk=content_type.id)
        except:
            notifications = None
        # if pest_report.author != request.user:
        #     return HttpResponseForbidden()
    else:
        pfa = PestFreeArea(author=request.user)
    old_issue=[]
    if pfa.issuename.count()>0:
        for e in pfa.issuename.all():
                obj_i=e.content_object.issuename
                for o in obj_i.all():
                    for iss in o.issuename.all():
                        old_issue.append(iss.id)
    old_comm=[]
    if pfa.commname.count()>0:
        for e in pfa.commname.all():
                obj_c=e.content_object.commname
                for o in obj_c.all():
                    for com in o.commname.all():
                        old_comm.append(com.id)    
                              
    if request.POST:

        form = PestFreeAreaForm(request.POST,  request.FILES, instance=pfa)
        if pfa.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=pfa.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
        if pfa.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=pfa.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm(request.POST,instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm(request.POST)
        notifyrelateform =NotificationMessageRelateForm(request.POST,instance=notifications)
        f_form = PestFreeAreaFileFormSet(request.POST,  request.FILES,instance=pfa)
        u_form = PestFreeAreaUrlFormSet(request.POST,  instance=pfa)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            old_pfa.pk = None
            old_pfa.is_version = True
            old_pfa.parent_id= id
            versions= PestFreeArea.objects.filter( is_version=True, parent_id=id).count()
            slug1 = versions+1
      
            old_pfa.slug= pfa.slug+'-'+str(slug1)
            old_pfa.save()

            
            issueformold =IssueKeywordsRelateForm()
            issue_instanceold=issueformold.save(commit=False)
            issue_instanceold.content_object = old_pfa
            issue_instanceold.save()
            
            commformold =CommodityKeywordsRelateForm()
            comm_instanceold=commformold.save(commit=False)
            comm_instanceold.content_object = old_pfa
            comm_instanceold.save()

            db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
            cursor = db.cursor()
            
            files=PestFreeAreaFile.objects.filter(pfa_id=pfa.id)
            urls=PestFreeAreaUrl.objects.filter(pfa_id=pfa.id)
            for f in files:
                sql = "INSERT INTO ippc_pestfreeareafile(pfa_id,description,file) VALUES ("+str(old_pfa.id)+", '"+ugettext(f.description)+"', '"+str(f)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for u in urls:
                sql = "INSERT INTO ippc_pestfreeareaurl(pfa_id,url_for_more_information) VALUES ("+str(old_pfa.id)+", '"+str(u)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()   
            for iss in old_issue:
                sql = """INSERT INTO ippc_issuekeywordsrelate_issuename(issuekeywordsrelate_id,issuekeyword_id) VALUES ("""+str(issue_instanceold.id)+""", """+str(iss)+""")"""
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for com in old_comm:
                sql = """INSERT INTO ippc_commoditykeywordsrelate_commname(commoditykeywordsrelate_id,commoditykeyword_id) VALUES ("""+str(comm_instanceold.id)+""", """+str(com)+""")"""
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                   #print("###################error ")
                    db.rollback()        
            db.close()
            
            
            
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = pfa
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = pfa
            commodity_instance.save()
            commodityform.save_m2m() 
            
            notify_instance = notifyrelateform.save(commit=False)
            notify_instance.content_object = pfa
            notify_instance.new_or_updated = 'UPDATE'
            
            notify_instance.link = user_country_slug+'/pestfreeareas/'+str(pfa.publish_date.strftime("%Y"))+'/'+str(pfa.publish_date.strftime("%m"))+'/'+pfa.slug+'/'
            notify_instance.updated_last = datetime.today()
            notify_instance.save()
            notifyrelateform.save_m2m()
            
            f_form.instance = pfa
            f_form.save()
            u_form.instance = pfa
            u_form.save()
            
            send_notification_message(0,id,content_type,pfa.title,user_country_slug+'/pestfreeareas/'+str(pfa.publish_date.strftime("%Y"))+'/'+str(pfa.publish_date.strftime("%m"))+'/'+pfa.slug+'/')
            # If the save was successful, success message and redirect to another page
            # info(request, _("Successfully updated pest report."))
            return redirect("pfa-detail", country=user_country_slug, year=pfa.publish_date.strftime("%Y"), month=pfa.publish_date.strftime("%m"), slug=pfa.slug)

    else:
        form = PestFreeAreaForm(instance=pfa)
        if pfa.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=pfa.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm( instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
        if pfa.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=pfa.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm( instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm()
        notifyrelateform =NotificationMessageRelateForm(instance=notifications)
        f_form = PestFreeAreaFileFormSet(instance=pfa)
        u_form = PestFreeAreaUrlFormSet(instance=pfa)
        
    return render_to_response(template_name, {
        'form': form,'f_form':f_form,'u_form':u_form,'issueform': issueform,'commodityform': commodityform,  "pfa": pfa,'notifyrelateform':notifyrelateform
    }, context_instance=RequestContext(request))
    

class ImplementationISPMListView(ListView):
    """    ImplementationISPM """
    context_object_name = 'latest'
    model = ImplementationISPM
    date_field = 'publish_date'
    template_name = 'countries/implementationispm_list.html'
    queryset = ImplementationISPM.objects.all().order_by('-publish_date', 'title')
    
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return pest reports from the specific country """
        # self.country = get_object_or_404(CountryPage, country=self.kwargs['country'])
        self.country = self.kwargs['country']
        # CountryPage country_slug == country URL parameter keyword argument
        return ImplementationISPM.objects.filter(country__country_slug=self.country,is_version=False)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(ImplementationISPMListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        return context
   
       
   
class ImplementationISPMDetailView(DetailView):
    """ ImplementationISPM detail page """
    model = ImplementationISPM
    context_object_name = 'implementationispm'
    template_name = 'countries/implementationispm_detail.html'
    queryset = ImplementationISPM.objects.filter()

    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(ImplementationISPMDetailView, self).get_context_data(**kwargs)
        p = get_object_or_404(ImplementationISPM, slug=self.kwargs['slug'])
        
        versions= ImplementationISPM.objects.filter(country__country_slug=self.kwargs['country'], is_version=True, parent_id=p.id).order_by('-modify_date')
        context['versions'] = versions
      
        return context

@login_required
@permission_required('ippc.add_implementationispm', login_url="/accounts/login/")
def implementationispm_create(request, country):
    """ Create ImplementationISPM """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))


    form = ImplementationISPMForm(request.POST)
    issueform =IssueKeywordsRelateForm(request.POST)
    commodityform =CommodityKeywordsRelateForm(request.POST)
    notifyrelateform =NotificationMessageRelateForm(request.POST)
     
    if request.method == "POST":
        f_form =ImplementationISPMFileFormSet(request.POST, request.FILES)
        u_form =ImplementationISPMUrlFormSet(request.POST)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            new_implementationispm = form.save(commit=False)
            new_implementationispm.author = request.user
            new_implementationispm.author_id = author.id
            form.save()
            
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = new_implementationispm
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = new_implementationispm
            commodity_instance.save()
            commodityform.save_m2m()
            
            notify_instance = notifyrelateform.save(commit=False)
            notify_instance.content_object = new_implementationispm
            notify_instance.new_or_updated = 'NEW'
            notify_instance.link = user_country_slug+'/implementationispm/'+str(new_implementationispm.publish_date.strftime("%Y"))+'/'+str(new_implementationispm.publish_date.strftime("%m"))+'/'+new_implementationispm.slug+'/'
     
            notify_instance.updated_last = new_implementationispm.publish_date
            notify_instance.save()
            notifyrelateform.save_m2m()
            
            f_form.instance = new_implementationispm
            f_form.save()
            u_form.instance = new_implementationispm
            u_form.save()
            content_type = ContentType.objects.get_for_model(new_implementationispm)
            send_notification_message(1,new_implementationispm.id,content_type,new_implementationispm.title,user_country_slug+'/implementationispm/'+str(new_implementationispm.publish_date.strftime("%Y"))+'/'+str(new_implementationispm.publish_date.strftime("%m"))+'/'+new_implementationispm.slug+'/')
            info(request, _("Successfully created implementationispm."))
            
            return redirect("implementationispm-detail", country=user_country_slug, year=new_implementationispm.publish_date.strftime("%Y"), month=new_implementationispm.publish_date.strftime("%m"), slug=new_implementationispm.slug)
        else:
             return render_to_response('countries/implementationispm_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform,'notifyrelateform':notifyrelateform},
             context_instance=RequestContext(request))
    else:
        form = ImplementationISPMForm(initial={'country': country}, instance=ImplementationISPM())
        issueform =IssueKeywordsRelateForm(request.POST)
        commodityform =CommodityKeywordsRelateForm(request.POST)
        notifyrelateform =NotificationMessageRelateForm(request.POST)
        f_form =ImplementationISPMFileFormSet()
        u_form =ImplementationISPMUrlFormSet()

    return render_to_response('countries/implementationispm_create.html', {'form': form  ,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform,'notifyrelateform':notifyrelateform},
        context_instance=RequestContext(request))


        
# http://stackoverflow.com/a/1854453/412329
@login_required
@permission_required('ippc.change_implementationispm', login_url="/accounts/login/")
def implementationispm_edit(request, country, id=None, template_name='countries/implementationispm_edit.html'):
    """ Edit implementationispm """
    user = request.user
    author = user
    country = user.get_profile().country
    # country_id = PestReport.objects.filter(country__country_id=country.id)
    user_country_slug = lower(slugify(country))
    if id:
        implementationispm = get_object_or_404(ImplementationISPM, country=country, pk=id)
        old_implementationispm=get_object_or_404(ImplementationISPM, country=country, pk=id)
       

        content_type = ContentType.objects.get_for_model(implementationispm)
        try:
            notifications = get_object_or_404(NotificationMessageRelate, object_id=id,content_type__pk=content_type.id)
        except:
            notifications = None
       # if pest_report.author != request.user:
        #     return HttpResponseForbidden()
    else:
        implementationispm = ImplementationISPM(author=request.user)
    old_issue=[]
    if implementationispm.issuename.count()>0:
        for e in implementationispm.issuename.all():
                obj_i=e.content_object.issuename
                for o in obj_i.all():
                    for iss in o.issuename.all():
                        old_issue.append(iss.id)
    old_comm=[]
    if implementationispm.commname.count()>0:
        for e in implementationispm.commname.all():
                obj_c=e.content_object.commname
                for o in obj_c.all():
                    for com in o.commname.all():
                        old_comm.append(com.id)     
    if request.POST:
        form = ImplementationISPMForm(request.POST,  request.FILES, instance=implementationispm)
        if implementationispm.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=implementationispm.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
        if implementationispm.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=implementationispm.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm(request.POST,instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm(request.POST)
        
        notifyrelateform =NotificationMessageRelateForm(request.POST,instance=notifications)
        f_form = ImplementationISPMFileFormSet(request.POST,  request.FILES,instance=implementationispm)
        u_form = ImplementationISPMUrlFormSet(request.POST,  instance=implementationispm)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            old_implementationispm.pk = None
            old_implementationispm.is_version = True
            old_implementationispm.parent_id= id
            versions= ImplementationISPM.objects.filter( is_version=True, parent_id=id).count()
            slug1 = versions+1
      
            old_implementationispm.slug=implementationispm.slug+'-'+str(slug1)
            old_implementationispm.save()

            
            issueformold =IssueKeywordsRelateForm()
            issue_instanceold=issueformold.save(commit=False)
            issue_instanceold.content_object = old_implementationispm
            issue_instanceold.save()
            
            commformold =CommodityKeywordsRelateForm()
            comm_instanceold=commformold.save(commit=False)
            comm_instanceold.content_object = old_implementationispm
            comm_instanceold.save()

            db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
            cursor = db.cursor()
            
            files=ImplementationISPMFile.objects.filter(implementationispm_id=implementationispm.id)
            urls=ImplementationISPMUrl.objects.filter(implementationispm_id=implementationispm.id)
            for f in files:
                sql = "INSERT INTO ippc_implementationispmfile(implementationispm_id,description,file) VALUES ("+str(old_implementationispm.id)+", '"+ugettext(f.description)+"', '"+str(f)+"')"
                ##print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for u in urls:
                sql = "INSERT INTO ippc_implementationispmurl(implementationispm_id,url_for_more_information) VALUES ("+str(old_implementationispm.id)+", '"+str(u)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()   
            for iss in old_issue:
                sql = """INSERT INTO ippc_issuekeywordsrelate_issuename(issuekeywordsrelate_id,issuekeyword_id) VALUES ("""+str(issue_instanceold.id)+""", """+str(iss)+""")"""
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for com in old_comm:
                sql = """INSERT INTO ippc_commoditykeywordsrelate_commname(commoditykeywordsrelate_id,commoditykeyword_id) VALUES ("""+str(comm_instanceold.id)+""", """+str(com)+""")"""
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    print("###################error ")
                    db.rollback()        
            db.close()
            
            
            
            
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = implementationispm
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = implementationispm
            commodity_instance.save()
            commodityform.save_m2m() 
            
            notify_instance = notifyrelateform.save(commit=False)
            notify_instance.content_object = implementationispm
            notify_instance.new_or_updated = 'UPDATE'
            notify_instance.link = user_country_slug+'/implementationispm/'+str(implementationispm.publish_date.strftime("%Y"))+'/'+str(implementationispm.publish_date.strftime("%m"))+'/'+implementationispm.slug+'/'
     
            notify_instance.updated_last = datetime.today()
            notify_instance.save()
            notifyrelateform.save_m2m()
            
            f_form.instance = implementationispm
            f_form.save()
            u_form.instance = implementationispm
            u_form.save()
            send_notification_message(0,id,content_type,implementationispm.title,user_country_slug+'/implementationispm/'+str(implementationispm.publish_date.strftime("%Y"))+'/'+str(implementationispm.publish_date.strftime("%m"))+'/'+implementationispm.slug+'/')
            
            # If the save was successful, success message and redirect to another page
            # info(request, _("Successfully updated pest report."))
            return redirect("implementationispm-detail", country=user_country_slug, year=implementationispm.publish_date.strftime("%Y"), month=implementationispm.publish_date.strftime("%m"), slug=implementationispm.slug)

    else:
        form = ImplementationISPMForm(instance=implementationispm)
        if implementationispm.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=implementationispm.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm( instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
        if implementationispm.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=implementationispm.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm( instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm()
        notifyrelateform =NotificationMessageRelateForm(instance=notifications)
        f_form = ImplementationISPMFileFormSet(instance=implementationispm)
        u_form = ImplementationISPMUrlFormSet(instance=implementationispm)
    return render_to_response(template_name, {
        'form': form,'f_form':f_form,'u_form': u_form,'issueform': issueform,'commodityform': commodityform,  "implementationispm": implementationispm,'notifyrelateform':notifyrelateform
    }, context_instance=RequestContext(request))

    #/**************************************************************/
class CountryNewsListView(ListView):
    """    CountryNews """
    context_object_name = 'latest'
    model = CountryNews
    date_field = 'publish_date'
    template_name = 'countries/countrynews_list.html'
    queryset = CountryNews.objects.all().order_by('-publish_date', 'title')
    
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return CountryNews from the specific country """
        # self.country = get_object_or_404(CountryPage, country=self.kwargs['country'])
        self.country = self.kwargs['country']
        # CountryPage country_slug == country URL parameter keyword argument
        return CountryNews.objects.filter(country__country_slug=self.country)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryNewsListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        return context
   
       
   
class CountryNewsDetailView(DetailView):
    """ CountryNews detail page """
    model = CountryNews
    context_object_name = 'countrynews'
    template_name = 'countries/countrynews_detail.html'
    queryset = CountryNews.objects.filter()



@login_required
@permission_required('ippc.add_countrynews', login_url="/accounts/login/")
def countrynews_create(request, country):
    """ Create countrynews """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))


    form = CountryNewsForm(request.POST)
    issueform =IssueKeywordsRelateForm(request.POST)
    commodityform =CommodityKeywordsRelateForm(request.POST)
    
    if request.method == "POST":
        f_form =CountryNewsFileFormSet(request.POST, request.FILES)
        u_form =CountryNewsUrlFormSet(request.POST)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            new_countrynews = form.save(commit=False)
            new_countrynews.author = request.user
            new_countrynews.author_id = author.id
            form.save()
            
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = new_countrynews
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = new_countrynews
            commodity_instance.save()
            commodityform.save_m2m()
            
            f_form.instance = new_countrynews
            f_form.save()
            u_form.instance = new_countrynews
            u_form.save()
            
            info(request, _("Successfully created news."))
            
            return redirect("country-news-detail", country=user_country_slug, year=new_countrynews.publish_date.strftime("%Y"), month=new_countrynews.publish_date.strftime("%m"), slug=new_countrynews.slug)
        else:
             return render_to_response('countries/countrynews_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform},
             context_instance=RequestContext(request))
    else:
        form = CountryNewsForm(initial={'country': country}, instance=CountryNews())
        issueform =IssueKeywordsRelateForm(request.POST)
        commodityform =CommodityKeywordsRelateForm(request.POST)
        f_form =CountryNewsFileFormSet()
        u_form =CountryNewsUrlFormSet()

    return render_to_response('countries/countrynews_create.html', {'form': form  ,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform},
        context_instance=RequestContext(request))


        
# http://stackoverflow.com/a/1854453/412329
@login_required
@permission_required('ippc.change_countrynews', login_url="/accounts/login/")
def countrynews_edit(request, country, id=None, template_name='countries/countrynews_edit.html'):
    """ Edit countrynews """
    user = request.user
    author = user
    country = user.get_profile().country
    # country_id = PestReport.objects.filter(country__country_id=country.id)
    user_country_slug = lower(slugify(country))
    if id:
        countrynews = get_object_or_404(CountryNews, country=country, pk=id)
       # if pest_report.author != request.user:
        #     return HttpResponseForbidden()
    else:
        countrynews = CountryNews(author=request.user)
      
    if request.POST:
        form = CountryNewsForm(request.POST,  request.FILES, instance=countrynews)
        if countrynews.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=countrynews.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
        if countrynews.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=countrynews.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm(request.POST,instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm(request.POST)
        f_form = CountryNewsFileFormSet(request.POST,  request.FILES,instance=countrynews)
        u_form = CountryNewsUrlFormSet(request.POST,  instance=countrynews)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = countrynews
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = countrynews
            commodity_instance.save()
            commodityform.save_m2m() 
            
            f_form.instance = countrynews
            f_form.save()
            u_form.instance = countrynews
            u_form.save()
            # If the save was successful, success message and redirect to another page
            # info(request, _("Successfully updated pest report."))
            return redirect("country-news-detail", country=user_country_slug, year=countrynews.publish_date.strftime("%Y"), month=countrynews.publish_date.strftime("%m"), slug=countrynews.slug)

    else:
        form = CountryNewsForm(instance=countrynews)
        if countrynews.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=countrynews.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
        if countrynews.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=countrynews.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm( instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm( )
      
        f_form = CountryNewsFileFormSet(instance=countrynews)
        u_form = CountryNewsUrlFormSet(instance=countrynews)
    return render_to_response(template_name, {
        'form': form,'f_form':f_form,'u_form': u_form,'issueform': issueform,'commodityform': commodityform,  "countrynews": countrynews
    }, context_instance=RequestContext(request))
   
   
   
   
   
   
   
   
class PartnersNewsDetailView(DetailView):
    """ Partners News detail page """
    model = PartnersNews
    context_object_name = 'partnersnews'
    template_name = 'partners/partnersnews_detail.html'
    queryset = PartnersNews.objects.filter()
      
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PartnersNewsDetailView, self).get_context_data(**kwargs)
        page = get_object_or_404(PartnersPage, name=self.kwargs['partners'])
        context['pagetitle'] =  page.name
        context['pageslug'] =  page.slug
       # context['page'] =  page.partner_slug
        return context
     


@login_required
@permission_required('ippc.add_partnersnews', login_url="/accounts/login/")
def partners_news_create(request, partners):
    """ Create partnersnews """
    user = request.user
    author = user
    if user.get_profile().partner:
        partners=user.get_profile().partner
         
    user_partner_slug = lower(slugify(partners))
    #partners=user.get_profile().partner
    #print("--------------------------------")
    #print(user_partner_slug)
    page = get_object_or_404(PartnersPage, name=user_partner_slug)
    partners=page.id  
    #print(partners)
    #print("--------------------------------")
    
    form = PartnersNewsForm(request.POST)
    issueform =IssueKeywordsRelateForm(request.POST)
    commodityform =CommodityKeywordsRelateForm(request.POST)
    
    if request.method == "POST":
        f_form =PartnersNewsFileFormSet(request.POST, request.FILES)
        u_form =PartnersNewsUrlFormSet(request.POST)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            new_partnersnews = form.save(commit=False)
            new_partnersnews.author = request.user
            new_partnersnews.author_id = author.id
            form.save()
            
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = new_partnersnews
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = new_partnersnews
            commodity_instance.save()
            commodityform.save_m2m()
            
            f_form.instance = new_partnersnews
            f_form.save()
            u_form.instance = new_partnersnews
            u_form.save()
            
            info(request, _("Successfully created news."))
            
            return redirect("partner-news-detail", partners=user_partner_slug, year=new_partnersnews.publish_date.strftime("%Y"), month=new_partnersnews.publish_date.strftime("%m"), slug=new_partnersnews.slug)
        else:
             return render_to_response('partners/partnersnews_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform},
             context_instance=RequestContext(request))
    else:
        form = PartnersNewsForm(initial={'partners': partners}, instance=PartnersNews())
        issueform =IssueKeywordsRelateForm(request.POST)
        commodityform =CommodityKeywordsRelateForm(request.POST)
        f_form = PartnersNewsFileFormSet()
        u_form = PartnersNewsUrlFormSet()

    return render_to_response('partners/partnersnews_create.html', {'form': form  ,'f_form': f_form,'u_form': u_form,'issueform':issueform, 'commodityform':commodityform},
        context_instance=RequestContext(request))


        
# http://stackoverflow.com/a/1854453/412329
@login_required
@permission_required('ippc.change_partnersnews', login_url="/accounts/login/")
def partners_news_edit(request, partners, id=None, template_name='partners/partnersnews_edit.html'):
    """ Edit partner news """
    user = request.user
    author = user
    if user.get_profile().partner:
        partners=user.get_profile().partner
         
    user_partner_slug = lower(slugify(partners))
    #partners=user.get_profile().partner
    page = get_object_or_404(PartnersPage, name=user_partner_slug)
    partners=page.id 
    if id:
        partnernews = get_object_or_404( PartnersNews,  partners= partners, pk=id)
    else:
        partnernews =  PartnersNews(author=request.user)
      
    if request.POST:
        form =  PartnersNewsForm(request.POST,  request.FILES, instance=partnernews)
        if partnernews.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=partnernews.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
        if partnernews.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=partnernews.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm(request.POST,instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm(request.POST)
        f_form =  PartnersNewsFileFormSet(request.POST,  request.FILES,instance=partnernews)
        u_form =  PartnersNewsUrlFormSet(request.POST,  instance=partnernews)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = partnernews
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = partnernews
            commodity_instance.save()
            commodityform.save_m2m() 
            
            f_form.instance = partnernews
            f_form.save()
            u_form.instance = partnernews
            u_form.save()
            # If the save was successful, success message and redirect to another page
            # info(request, _("Successfully updated pest report."))
            return redirect("partner-news-detail", partners=user_partner_slug, year=partnernews.publish_date.strftime("%Y"), month=partnernews.publish_date.strftime("%m"), slug=partnernews.slug)

    else:
        form = PartnersNewsForm(instance=partnernews)
        if partnernews.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=partnernews.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm( instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
        if partnernews.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=partnernews.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm( instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm()
        f_form = PartnersNewsFileFormSet(instance=partnernews)
        u_form = PartnersNewsUrlFormSet(instance=partnernews)
    return render_to_response(template_name, {
        'form': form,'f_form':f_form,'u_form': u_form,'issueform': issueform,'commodityform': commodityform,  "partnernews": partnernews
    }, context_instance=RequestContext(request))
   
    
   
   
   
@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def publication_edit(request, id=None, template_name='pages/publication_edit.html'):
    """ Edit  Publication """
    user = request.user
    author = user
    if id:
        publication = get_object_or_404(Publication, pk=id)
        old_publication=get_object_or_404(Publication, pk=id)
     
    #    if publication.issuename:
    #        print(publication.issuename.all[0])
    #    issues = get_object_or_404(IssueKeywordsRelate, pk=publication.issuename.all()[0].id)
    #    commodities = get_object_or_404(CommodityKeywordsRelate, pk=publication.commname.all()[0].id)
    else:
        publication = Publication(author=request.user)
    old_issue=[]
    if publication.issuename.count()>0:
        for e in publication.issuename.all():
                obj_i=e.content_object.issuename
                for o in obj_i.all():
                    for iss in o.issuename.all():
                        old_issue.append(iss.id)
    old_comm=[]
    if publication.commname.count()>0:
        for e in publication.commname.all():
                obj_c=e.content_object.commname
                for o in obj_c.all():
                    for com in o.commname.all():
                        old_comm.append(com.id)     
    if request.POST:
        form = PublicationForm(request.POST,  request.FILES, instance=publication)
        if publication.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=publication.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
        if publication.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=publication.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm(request.POST,instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm(request.POST)
        f_form = PublicationFileFormSet(request.POST,  request.FILES,instance=publication)
        u_form = PublicationUrlFormSet(request.POST,  instance=publication)
      
        if form.is_valid() and f_form.is_valid() and u_form.is_valid():
            old_publication.pk = None
            old_publication.is_version = True
            old_publication.parent_id= id
            versions= Publication.objects.filter( is_version=True, parent_id=id).count()
            slug1 = versions+1
      
            old_publication.slug= publication.slug+'-'+str(slug1)
            old_publication.save()

            
            issueformold =IssueKeywordsRelateForm()
            issue_instanceold=issueformold.save(commit=False)
            issue_instanceold.content_object = old_publication
            issue_instanceold.save()
            
            commformold =CommodityKeywordsRelateForm()
            comm_instanceold=commformold.save(commit=False)
            comm_instanceold.content_object = old_publication
            comm_instanceold.save()

            db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
            cursor = db.cursor()
            
            files=PublicationFile.objects.filter(publication_id=publication.id)
            urls=PublicationUrl.objects.filter(publication_id=publication.id)
            for f in files:
                sql = "INSERT INTO ippc_publicationfile(publication_id,description,file) VALUES ("+str(old_publication.id)+", '"+ugettext(f.description)+"', '"+str(f)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for u in urls:
                sql = "INSERT INTO ippc_ippc_publicationurl(publication_id,url_for_more_information) VALUES ("+str(old_publication.id)+", '"+str(u)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()   
            for iss in old_issue:
                sql = """INSERT INTO ippc_issuekeywordsrelate_issuename(issuekeywordsrelate_id,issuekeyword_id) VALUES ("""+str(issue_instanceold.id)+""", """+str(iss)+""")"""
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for com in old_comm:
                sql = """INSERT INTO ippc_commoditykeywordsrelate_commname(commoditykeywordsrelate_id,commoditykeyword_id) VALUES ("""+str(comm_instanceold.id)+""", """+str(com)+""")"""
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    #print("###################error ")
                    db.rollback()        
            db.close()
            
            
            
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = publication
            issue_instance.save()
            issueform.save_m2m()
            
            commodity_instance = commodityform.save(commit=False)
            commodity_instance.content_object = publication
            commodity_instance.save()
            commodityform.save_m2m() 
    
            f_form.instance = publication
            f_form.save()
            u_form.instance = publication
            u_form.save()
            info(request, _("Successfully updated publication."))
            return redirect("publication-detail", pk=publication.id)

    else:
        form = PublicationForm(instance=publication)
        
        if publication.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=publication.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm( instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
        if publication.commname.count()>0:
            commodities = get_object_or_404(CommodityKeywordsRelate, pk=publication.commname.all()[0].id)
            commodityform =CommodityKeywordsRelateForm( instance=commodities)
        else:
            commodityform =CommodityKeywordsRelateForm()
        f_form = PublicationFileFormSet(instance=publication)
        u_form = PublicationUrlFormSet( instance=publication)
      
    return render_to_response(template_name, {
        'form': form, 'f_form':f_form,'u_form': u_form,'issueform': issueform,  'commodityform': commodityform, "publication": publication
    }, context_instance=RequestContext(request))       
    
    
class CountryListView(ListView):
    """   alphabetic countries list  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_list.html'
    queryset = CountryPage.objects.all().order_by('title')
    #region_name=self.kwargs['region']
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryListView, self).get_context_data(**kwargs)
        context['number_of_cp']= CountryPage.objects.filter(cp_ncp_t_type='CP').count()
        if self.kwargs['region'] == 'all':
            context['countries']= CountryPage.objects.all()
        elif self.kwargs['region'] == 'nppos':
            context['countries']= CountryPage.objects.filter(cp_ncp_t_type='CP')
            context['region_name']= 'NPPOs'
        else:
            for k,v in REGIONS:
                reg = v.lower()
                reg = reg.replace(" ", "-");
                if reg == self.kwargs['region']:
                    kindex=k
                    context['region_name']=v
            context['countries']= CountryPage.objects.filter(region=kindex)
        return context

    
class CountryStatsTotalreportsListView(ListView):
    """   Statistics reports  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_statstotalreports.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): 
        context = super(CountryStatsTotalreportsListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
        datachart1=''
        datachart2=''
        datachart3=''
        tot_rep_count=0
        tot_ev_count1=0
        tot_ev_count2=0
        results_list = []
        tot = []
        p_count=PestReport.objects.filter(is_version=False).count()
        tot.append(p_count)
        tot_ev_count1+=p_count
        for i in range(1,6):
           rep_count=ReportingObligation.objects.filter(reporting_obligation_type=i,is_version=False).count()
           tot_rep_count+=rep_count
           ev_count=EventReporting.objects.filter(event_rep_type=i,is_version=False).count()
           if i<=2:
               tot_ev_count1+=ev_count
           else:    
               tot_ev_count2+=ev_count
           tot.append(rep_count)
           tot.append(ev_count)
        results_list.append(tot)
        context['results_list']=results_list
        
        for r in results_list:
             datachart1 += ' {  y: '+str(r[1]*100/tot_rep_count)+', legendText:"Description of the NPPO", label: "Description of the NPPO: '+str(r[1]*100/tot_rep_count)+'%" },'
             datachart1 += ' {  y: '+str(r[3]*100/tot_rep_count)+', legendText:"Entry points", label: "Entry points: '+str(r[3]*100/tot_rep_count)+'%" },'
             datachart1 += ' {  y: '+str(r[5]*100/tot_rep_count)+', legendText:"List of regulated pests", label: "List of regulated pests: '+str(r[5]*100/tot_rep_count)+'%" },'
             datachart1 += ' {  y: '+str(r[7]*100/tot_rep_count)+', legendText:"Phytosanitary restrictions", label: "Phytosanitary restrictions: '+str(r[7]*100/tot_rep_count)+'%" },'
	
             datachart2 += ' {  y: '+str(r[4]*100/tot_ev_count1)+', legendText:"Non compliance", label: "Non compliance: '+str(r[4]*100/tot_rep_count)+'%" },'
             datachart2 += ' {  y: '+str(r[2]*100/tot_ev_count1)+', legendText:"Emergency actions", label: "Emergency actions: '+str(r[2]*100/tot_rep_count)+'%" },'
             datachart2 += ' {  y: '+str(r[0]*100/tot_ev_count1)+', legendText:"Pest report", label: "Pest report: '+str(r[0]*100/tot_rep_count)+'%" },'
            
             if r[6]>0:
                datachart3 += ' {  y: '+str(r[6]*100/tot_ev_count2)+', legendText:"Organizational (NPPO info)", label: "Organizational (NPPO info): '+str(r[6]*100/tot_rep_count)+'%" },'
             if r[10]>0:
                datachart3 += ' {  y: '+str(r[10]*100/tot_ev_count2)+', legendText:"PRA (rationale phytosanitary requirements) ", label: "PRA (rationale phytosanitary requirements): '+str(r[10]*100/tot_rep_count)+'%" },'
             if r[8]>0:
                datachart3 += ' {  y: '+str(r[8]*100/tot_ev_count2)+', legendText:"Pest status", label: "Pest status: '+str(r[8]*100/tot_rep_count)+'%" },'
        context['datachart1']=datachart1
        context['datachart2']=datachart2
        context['datachart3']=datachart3
        return context   
class CountryStatsTotalreports1ListView(ListView):
    """   Statistics reports  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_statstotalreports_1.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): 
        context = super(CountryStatsTotalreports1ListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
        context['selyear_range']=range(2010,timezone.now().year+1)
       
        tot_rep_count=0
       
        prevyear=0
        #print('>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>')
        if 'year' in self.kwargs:
            #print(self.kwargs['year'])
            prevyear=int(self.kwargs['year'])
        else :   
             prevyear=timezone.now().year -1
        
        
        
        rep_array=[]
        ev_array=[]
        pest_array=[]
        for i in range(1,6):
            reporting_array = []
            eventreporting_array = []
            rep_count=0
            evrep_count=0 
            reps=ReportingObligation.objects.filter(reporting_obligation_type=i,is_version=False)
            for r in reps:
                if r.publication_date != None and r.publication_date.year == prevyear :
                    rep_count=rep_count+1
                   # print(rep_count)
            reporting_array.append(rep_count)
            rep_array.append(reporting_array)
            evrep=EventReporting.objects.filter(event_rep_type=i,is_version=False)
            for e in evrep:
                if e.publication_date != None and e.publication_date.year  == prevyear :
                    evrep_count=evrep_count+1
            eventreporting_array.append(evrep_count)
            ev_array.append(eventreporting_array)
        
        pestreporting_array = []
        pests=PestReport.objects.filter(is_version=False)
        p_count=0
        for p in pests:
            if p.publish_date != None and p.publish_date.year == prevyear :
                p_count=p_count+1
        
        pestreporting_array.append(p_count)

        pest_array.append(pestreporting_array)       
        datachart=''
        
        
#        datachart += '{type: "column", name: "'+str(prevyear)+'", legendText: "'+str(prevyear)+'",showInLegend: true, dataPoints:['
#        datachart += '{label: "Description of NPPO", y: '+str(rep_array[0][0])+'},	'
#        datachart += '{label: "Pest reports", y: '+str(pest_array[0][0])+'},'
#        datachart += '{label: "Emergency action", y: '+str(ev_array[0][0])+'},'
#        datachart += '{label: "List of regulated pests", y: '+str(rep_array[3][0])+'},'
#        datachart += '{label: "Entry points", y:  '+str(rep_array[2][0])+'},'
#        datachart += '{label: "Legislation: phytosanitary requirements/ restrictions/ prohibitions", y:  '+str(rep_array[1][0])+'},'
#        datachart += '{label: "Non-compliance", y:  '+str(ev_array[1][0])+'},'
#        datachart += '{label: "Organizational arrangements of plant protection", y: '+str(ev_array[2][0])+'},'
#        datachart += '{label: "Pest status", y:'+str(ev_array[3][0])+'},'
#        datachart += '{label: "Rationale for phytosanitary requirements", y: '+str(ev_array[4][0])+'}]},'
#        
        tot_rep_count=rep_array[0][0]+rep_array[1][0]+rep_array[2][0]+rep_array[3][0]+pest_array[0][0]+ev_array[0][0]+ev_array[1][0]+ev_array[2][0]+ev_array[3][0]+ev_array[4][0]
        
        #print(tot_rep_count)
        datachart0=''
        datachart0 += ' {  y: '+str(rep_array[0][0]*100/tot_rep_count)+', legendText:"Description of the NPPO", label: "Description of the NPPO: '+str(rep_array[0][0]*100/tot_rep_count)+'%" },'
        datachart0 += ' {  y: '+str(pest_array[0][0]*100/tot_rep_count)+', legendText:"Pest report", label: "Pest report: '+str(pest_array[0][0]*100/tot_rep_count)+'%" },'
        datachart0 += ' {  y: '+str(ev_array[0][0]*100/tot_rep_count)+', legendText:"Emergency actions", label: "Emergency actions: '+str(ev_array[0][0]*100/tot_rep_count)+'%" },'
        datachart0 += ' {  y: '+str(rep_array[2][0]*100/tot_rep_count)+', legendText:"List of regulated pests", label: "List of regulated pests: '+str(rep_array[2][0]*100/tot_rep_count)+'%" },'
        datachart0 += ' {  y: '+str(rep_array[1][0]*100/tot_rep_count)+', legendText:"Entry points", label: "Entry points: '+str(rep_array[1][0]*100/tot_rep_count)+'%" },'
        datachart0 += ' {  y: '+str(rep_array[3][0]*100/tot_rep_count)+', legendText:"Legislation: phytosanitary requirements/ restrictions/ prohibitions", label: "Legislation: phytosanitary requirements/ restrictions/ prohibitions: '+str(rep_array[3][0]*100/tot_rep_count)+'%" },'
        datachart0 += ' {  y: '+str(ev_array[1][0]*100/tot_rep_count)+', legendText:"Non compliance", label: "Non compliance: '+str(ev_array[1][0]*100/tot_rep_count)+'%" },'
        datachart0 += ' {  y: '+str(ev_array[2][0]*100/tot_rep_count)+', legendText:"Organizational arrangements of plant protection", label: "Organizational arrangements of plant protection: '+str(ev_array[2][0]*100/tot_rep_count)+'%" },'
        datachart0 += ' {  y: '+str(ev_array[3][0]*100/tot_rep_count)+', legendText:"Pest status", label: "Pest status: '+str(ev_array[3][0]*100/tot_rep_count)+'%" },'
        datachart0 += ' {  y: '+str(ev_array[4][0]*100/tot_rep_count)+', legendText:"Rationale for phytosanitary requirements", label: "Rationale for phytosanitary requirements: '+str(ev_array[4][0]*100/tot_rep_count)+'%" },'
       
   
        
        context['prevyear']=prevyear
        context['rep_array']=rep_array
        context['ev_array']=ev_array
        context['pest_array']=pest_array
            
        context['datachart']=datachart
        context['datachart0']=datachart0
        context['tot_rep_count']=tot_rep_count
        
        
        return context
    
       
		

class CountryStatsreportsListView(ListView):
    """   Statistics reports  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_statsreports.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): 
        context = super(CountryStatsreportsListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
        results_list = []
        countriesList=CountryPage.objects.filter().exclude(id='-1')
        for c in countriesList:
             totcn = []
             totcn.append(c)    
             p_count=PestReport.objects.filter(country_id=c.id,is_version=False).count()
             totcn.append(p_count)
             for i in range(1,6):
                rep_count=ReportingObligation.objects.filter(country_id=c.id,reporting_obligation_type=i,is_version=False).count()
                ev_count=EventReporting.objects.filter(country_id=c.id,event_rep_type=i,is_version=False).count()
                totcn.append(rep_count)
                totcn.append(ev_count)
             totcn.append((slugify(c)))
             results_list.append(totcn)
        context['results_list']=results_list
    
        return context
  
class CountryRegionsPercentageListView(ListView):
    """   stat  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_regionspercentage.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryRegionsPercentageListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
        regionCN = []
        regionCNcp = []
        regionCNncp = []
        
        regionpest = []
        regionrep1 = []
        regionrep2 = []
        regionrep3 = []
        regionrep4 = []
        regionev1 = []
        regionev2 = []
        regionev3 = []
        regionev4 = []
        regionev5 = []
        regionpestcp = []
        regionrep1cp = []
        regionrep2cp = []
        regionrep3cp= []
        regionrep4cp = []
        regionev1cp = []
        regionev2cp = []
        regionev3cp = []
        regionev4cp = []
        regionev5cp = []
        regionpestncp = []
        regionrep1ncp= []
        regionrep2ncp = []
        regionrep3ncp = []
        regionrep4ncp = []
        regionev1ncp = []
        regionev2ncp = []
        regionev3ncp = []
        regionev4ncp = []
        regionev5ncp = []
        for k,v in REGIONS:
                reg = v.lower()
                numCN = []
                countriesperregion=CountryPage.objects.filter(region=k).exclude(cp_ncp_t_type='T')
                numb_countriesperregion=countriesperregion.count()
                numCN.append(reg)
                numCN.append(numb_countriesperregion)
                regionCN.append(numCN)
                context['region_cn']=regionCN
                
                numCNcp = []
                countriesperregioncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='CP')
                numb_countriesperregioncp=countriesperregioncp.count()
                numCNcp.append(reg)
                numCNcp.append(numb_countriesperregioncp)
                regionCNcp.append(numCNcp)
                context['region_cp']=regionCNcp
                
                numCNncp = []
                countriesperregionncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='NCP')
                numb_countriesperregionncp=countriesperregionncp.count()
                numCNncp.append(reg)
                numCNncp.append(numb_countriesperregionncp)
                regionCNncp.append(numCNncp)
                context['region_ncp']=regionCNncp
                
                pests = []
                p_count=0
                for c in countriesperregion:
                    p=PestReport.objects.filter(country_id=c.id,is_version=False)
                    p_count+=p.count()
                pestC=(int)((p_count * 100)/numb_countriesperregion)
                pests.append(pestC)
                regionpest.append(pests)   
                context['region_pest']=regionpest
                
                for i in range(1,6):
                    rep_count=0
                    ev_count=0
                    reporting_array = []
                    evreporting_array=[]
                    for c in countriesperregion:
                        r=ReportingObligation.objects.filter(country_id=c.id,reporting_obligation_type=i,is_version=False)
                        rep_count+=r.count()
                        r1=EventReporting.objects.filter(country_id=c.id,event_rep_type=i,is_version=False)
                        ev_count+=r1.count()
                    repC=(int)((rep_count * 100)/numb_countriesperregion)
                    reporting_array.append(repC)
                    repE=(int)((ev_count * 100)/numb_countriesperregion)
                    evreporting_array.append(repE)
                    if i==1:
                        regionrep1.append(reporting_array)   
                        context['region_rep'+str(i)]=regionrep1
                        regionev1.append(evreporting_array)   
                        context['region_ev'+str(i)]=regionev1
                    elif i==2:
                        regionrep2.append(reporting_array)   
                        context['region_rep'+str(i)]=regionrep2
                        regionev2.append(evreporting_array)   
                        context['region_ev'+str(i)]=regionev2
                    elif i==3:
                        regionrep3.append(reporting_array)   
                        context['region_rep'+str(i)]=regionrep3
                        regionev3.append(evreporting_array)   
                        context['region_ev'+str(i)]=regionev3
                    elif i==4:
                        regionrep4.append(reporting_array)   
                        context['region_rep'+str(i)]=regionrep4
                        regionev4.append(evreporting_array)   
                        context['region_ev'+str(i)]=regionev4
                    elif i==5:
                        regionev5.append(evreporting_array)   
                        context['region_ev'+str(i)]=regionev5    
                  #CP
                pests = []
                p_count=0
                for c in countriesperregioncp:
                    p=PestReport.objects.filter(country_id=c.id,is_version=False)
                    p_count+=p.count()
                pestC=(int)((p_count * 100)/numb_countriesperregioncp)
                pests.append(pestC)
                regionpestcp.append(pests)   
                context['region_pestcp']=regionpestcp

                for i in range(1,6):
                    rep_count=0
                    ev_count=0
                    reporting_array = []
                    evreporting_array=[]
                    for c in countriesperregioncp:
                        r=ReportingObligation.objects.filter(country_id=c.id,reporting_obligation_type=i,is_version=False)
                        rep_count+=r.count()
                        r1=EventReporting.objects.filter(country_id=c.id,event_rep_type=i,is_version=False)
                        ev_count+=r1.count()
                    repC=(int)((rep_count * 100)/numb_countriesperregioncp)
                    reporting_array.append(repC)
                    repE=(int)((ev_count * 100)/numb_countriesperregioncp)
                    evreporting_array.append(repE)
                    if i==1:
                        regionrep1cp.append(reporting_array)   
                        context['region_rep'+str(i)+'cp']=regionrep1cp
                        regionev1cp.append(evreporting_array)   
                        context['region_ev'+str(i)+'cp']=regionev1cp
                    elif i==2:
                        regionrep2cp.append(reporting_array)   
                        context['region_rep'+str(i)+'cp']=regionrep2cp
                        regionev2cp.append(evreporting_array)   
                        context['region_ev'+str(i)+'cp']=regionev2cp
                    elif i==3:
                        regionrep3cp.append(reporting_array)   
                        context['region_rep'+str(i)+'cp']=regionrep3cp
                        regionev3cp.append(evreporting_array)   
                        context['region_ev'+str(i)+'cp']=regionev3cp
                    elif i==4:
                        regionrep4cp.append(reporting_array)   
                        context['region_rep'+str(i)+'cp']=regionrep4cp
                        regionev4cp.append(evreporting_array)   
                        context['region_ev'+str(i)+'cp']=regionev4cp
                    elif i==5:
                        regionev5cp.append(evreporting_array)   
                        context['region_ev'+str(i)+'cp']=regionev5cp
                   #NCP
                pests = []
                p_count=0
                for c in countriesperregionncp:
                    p=PestReport.objects.filter(country_id=c.id,is_version=False)
                    p_count+=p.count()
                pestC=0
                if numb_countriesperregionncp>0:
                    pestC=(int)((p_count * 100)/numb_countriesperregionncp)
                pests.append(pestC)
                regionpestncp.append(pests)   
                context['region_pestncp']=regionpestncp

                for i in range(1,6):
                    rep_count=0
                    ev_count=0
                    reporting_array = []
                    evreporting_array=[]
                    for c in countriesperregionncp:
                        r=ReportingObligation.objects.filter(country_id=c.id,reporting_obligation_type=i,is_version=False)
                        rep_count+=r.count()
                        r1=EventReporting.objects.filter(country_id=c.id,event_rep_type=i,is_version=False)
                        ev_count+=r1.count()
                    repC=0
                    repE=0
                    if numb_countriesperregionncp>0:
                        repC=(int)((rep_count * 100)/numb_countriesperregionncp)
                        repE=(int)((ev_count * 100)/numb_countriesperregionncp)
              
                    reporting_array.append(repC)
                    evreporting_array.append(repE)
                    if i==1:
                        regionrep1ncp.append(reporting_array)   
                        context['region_rep'+str(i)+'ncp']=regionrep1ncp
                        regionev1ncp.append(evreporting_array)   
                        context['region_ev'+str(i)+'ncp']=regionev1ncp
                    elif i==2:
                        regionrep2ncp.append(reporting_array)   
                        context['region_rep'+str(i)+'ncp']=regionrep2ncp
                        regionev2ncp.append(evreporting_array)   
                        context['region_ev'+str(i)+'ncp']=regionev2ncp
                    elif i==3:
                        regionrep3ncp.append(reporting_array)   
                        context['region_rep'+str(i)+'ncp']=regionrep3ncp
                        regionev3ncp.append(evreporting_array)   
                        context['region_ev'+str(i)+'ncp']=regionev3ncp
                    elif i==4:
                        regionrep4ncp.append(reporting_array)   
                        context['region_rep'+str(i)+'ncp']=regionrep4ncp
                        regionev4ncp.append(evreporting_array)   
                        context['region_ev'+str(i)+'ncp']=regionev4ncp
                    elif i==5:
                        regionev5ncp.append(evreporting_array)   
                        context['region_ev'+str(i)+'ncp']=regionev5ncp
  
        return context

from datetime import date



class CountryStatsSingleListOfRegulatesPestsListView(ListView):
    """   stat  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_year_of_listofregulatedpests.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryStatsSingleListOfRegulatesPestsListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
        context['selyear_range']=range(2010,timezone.now().year+1)
       
        curryear=0
        prevyear=0
        num_years=0
        
        #if 'year' in self.kwargs:
        #    curryear=int(self.kwargs['year'])
        #else:   
            #curryear=timezone.now().year+1
        curryear=2019   
        prevyear=curryear-1
        num_years=curryear-2005
        
        startstartdate = datetime(1970, 1, 1, 00, 01,00)
        startdate = datetime(prevyear, 4, 1, 00, 01,00)
        enddate = datetime(curryear, 3, 31, 23, 59,00)
      
        
        regionsPCP = []
        totNumReg=CountryPage.objects.filter(cp_ncp_t_type='CP').count()
        region_cp_p = []
        numRepP=0
           
        for k,v in REGIONS:
            reg = v+''
            numRepP=0
            countriesperregioncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='CP')
            numCP_P = []
            numCP_P.append(reg)
            numCP_P.append(countriesperregioncp.count())
            countP=0
            cNewP=0
            cUpP=0
            cns=''
            
            cns_4=''
            cns_5=''
            cns_6=''
            cns_7=''
            cns_8=''
            cns_9=''
            cns_10=''
            cns_11=''
            cns_12=''
            cns_1=''
            cns_2=''
            cns_3=''
            tot_4=''
            tot_5=''
            tot_6=''
            tot_7=''
            tot_8=''
            tot_9=''
            tot_10=''
            tot_11=''
            tot_12=''
            tot_1=''
            tot_2=''
            tot_3=''
            p_count=0
            
            cNewP_4=0
            cNewP_5=0
            cNewP_6=0
            cNewP_7=0
            cNewP_8=0
            cNewP_9=0
            cNewP_10=0
            cNewP_11=0
            cNewP_12=0
            cNewP_1=0
            cNewP_2=0
            cNewP_3=0
            cUpP_4=0
            cUpP_5=0
            cUpP_6=0
            cUpP_7=0
            cUpP_8=0
            cUpP_9=0
            cUpP_10=0
            cUpP_11=0
            cUpP_12=0
            cUpP_1=0
            cUpP_2=0
            cUpP_3=0
           
            cnsP_4=""
            cnsP_5=""
            cnsP_6=""
            cnsP_7=""
            cnsP_8=""
            cnsP_9=""
            cnsP_10=""
            cnsP_11=""
            cnsP_12=""
            cnsP_1=""
            cnsP_2=""
            cnsP_3=""
            cnsUpP_4=""
            cnsUpP_5=""
            cnsUpP_6=""
            cnsUpP_7=""
            cnsUpP_8=""
            cnsUpP_9=""
            cnsUpP_10=""
            cnsUpP_11=""
            cnsUpP_12=""
            cnsUpP_1=""
            cnsUpP_2=""
            cnsUpP_3=""
           
            array_cns_p=[]
            for c in countriesperregioncp:
                array_cns_p_up_new=[]
                array_cns_pn_=[]
                array_cns_pu_=[]
                array_cns_n_=[]
                array_cns_u_=[]
                pests= ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=3,is_version=False,publication_date__gte=startstartdate,publication_date__lte=enddate)
                p_count=pests.count()
                
                pests1= ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=3,is_version=False,publication_date__gte=startdate,publication_date__lte=enddate)
                pests2= ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=3,is_version=False,modify_date__gte=startdate,modify_date__lte=enddate)
           
           
                if p_count>0:
                    numRepP+=1    
                if pests1.count()>0 or pests2.count()>0:
                    cns+=c.title+', '
                    
                countP+=p_count
                #cNewP+=pests1.count()
                #cUpP+=pests2.count()
                
                if pests1.count()>0:
                    for p in pests1:
                        p_date=p.publication_date.month
                       
                        if p_date==4:
                          #print(p_date)  
                          cNewP_4=cNewP_4+1
                          cNewP+=1
                          #print(c.title)  
                          if c.title in cnsP_4:
                              print("NO")
                          else:    
                             cnsP_4=cnsP_4+c.title+', '
                    
                        elif   p_date==5:
                          cNewP_5=cNewP_5+1
                          cNewP+=1
                          if c.title in cnsP_5:
                              print("NO")
                          else:    
                             cnsP_5=cnsP_5+c.title+', '
                        elif   p_date==6:
                              cNewP_6=cNewP_6+1
                              cNewP+=1
                              if c.title in cnsP_6:
                                print("NO")
                              else:    
                                cnsP_6=cnsP_6+c.title+', '
                        elif   p_date==5:
                              cNewP_7=cNewP_7+1
                              cNewP+=1
                              if c.title in cnsP_7:
                                  print("NO")
                              else:    
                               cnsP_7=cnsP_7+c.title+', '
                        elif   p_date==8:
                              cNewP_8=cNewP_8+1
                              cNewP+=1
                              if c.title in cnsP_8:
                                  print("NO")
                              else:    
                                cnsP_8=cnsP_8+c.title+', '
                        elif   p_date==9:
                              cNewP_9=cNewP_9+1
                              cNewP+=1
                              if c.title in cnsP_9:
                                print("NO")
                              else:    
                                cnsP_9=cnsP_9+c.title+', '
                        elif   p_date==10:
                              cNewP_10=cNewP_10+1
                              cNewP+=1
                              if c.title in cnsP_10:
                                print("NO")
                              else:    
                                cnsP_10=cnsP_10+c.title+', '
                        elif   p_date==11:
                              cNewP_11=cNewP_11+1
                              cNewP+=1
                              if c.title in cnsP_11:
                                 print("NO")
                              else:    
                                cnsP_11=cnsP_11+c.title+', '
                        elif   p_date==12:
                              cNewP_12=cNewP_12+1
                              cNewP+=1
                              if c.title in cnsP_12:
                                  print("NO")
                              else:    
                                cnsP_12=cnsP_12+c.title+', '
                        elif   p_date==1:
                              cNewP_1=cNewP_1+1
                              cNewP+=1
                              if c.title in cnsP_1:
                                 print("NO")
                              else:    
                                 cnsP_1=cnsP_1+c.title+', '
                        elif   p_date==2:
                              cNewP_2=cNewP_2+1
                              cNewP+=1
                              if c.title in cnsP_2:
                                 print("NO")
                              else:    
                                cnsP_2=cnsP_2+c.title+', '
                        elif   p_date==3:
                              cNewP_3=cNewP_3+1
                              cNewP+=1
                              if c.title in cnsP_3:
                                print("NO")
                              else:    
                                  cnsP_3=cnsP_3+c.title+', '
                if pests2.count()>0:
                    for p in pests2:
                        # @type p 
                        p_date=p.modify_date.month
                        if   p_date==4 and p.modify_date != p.publication_date:
                           cUpP_4= cUpP_4+1
                           cUpP+=1
                           if c.title in cnsUpP_4:
                              print("NO")
                           else:    
                            cnsUpP_4=cnsUpP_4+c.title+', '
                        elif   p_date==5 and p.modify_date != p.publication_date:
                           cUpP_5= cUpP_5+1
                           cUpP+=1
                           if c.title in cnsUpP_5:
                              print("NO")
                           else:    
                            cnsUpP_5=cnsUpP_5+c.title+', '

                        elif   p_date==6 and p.modify_date != p.publication_date:
                               cUpP_6= cUpP_6+1
                               cUpP+=1
                               if c.title in cnsUpP_6:
                                 print("NO")
                               else:    
                                  cnsUpP_6=cnsUpP_6+c.title+', '
                        elif   p_date==7 and p.modify_date != p.publication_date:
                               cUpP_7= cUpP_7+1
                               cUpP+=1
                               if c.title in cnsUpP_7:
                                 print("NO")
                               else:    
                                 cnsUpP_7=cnsUpP_7+c.title+', '
                        elif   p_date==8 and p.modify_date != p.publication_date:
                               cUpP_8= cUpP_8+1
                               cUpP+=1
                               #print('****************************')
                               #print('****************************cUpP_8='+str(cUpP_8))
                               if c.title in cnsUpP_8:
                                 print("NO")
                               else:    
                                  cnsUpP_8=cnsUpP_8+c.title+', '
                        elif   p_date==9 and p.modify_date != p.publication_date:
                               cUpP_9= cUpP_9+1
                               cUpP+=1
                               if c.title in cnsUpP_9:
                                 print("NO")
                               else:    
                                 cnsUpP_9=cnsUpP_9+c.title+', '
                        elif   p_date==10 and p.modify_date != p.publication_date:
                               cUpP_10= cUpP_10+1
                               cUpP+=1
                               if c.title in cnsUpP_10:
                                 print("NO")
                               else:    
                                  cnsUpP_10=cnsUpP_10+c.title+', '
                        elif   p_date==11 and p.modify_date != p.publication_date:
                               cUpP_11= cUpP_11+1
                               cUpP+=1
                               if c.title in cnsUpP_11:
                                print("NO")
                               else:    
                                 cnsUpP_11=cnsUpP_11+c.title+', '
                        elif   p_date==12 and p.modify_date != p.publication_date:
                               cUpP_12= cUpP_12+1
                               cUpP+=1
                               if c.title in cnsUpP_12:
                                print("NO")
                               else:    
                                 cnsUpP_12=cnsUpP_12+c.title+', '
                        elif   p_date==1 and p.modify_date != p.publication_date:
                               cUpP_1= cUpP_1+1
                               cUpP+=1
                               if c.title in cnsUpP_1:
                                print("NO")
                               else:    
                                 cnsUpP_1=cnsUpP_1+c.title+', '
                        elif   p_date==2 and p.modify_date != p.publication_date:
                               cUpP_2= cUpP_2+1
                               cUpP+=1
                               if c.title in cnsUpP_2:
                                    print("NO")
                               else:    
                                     cnsUpP_2=cnsUpP_2+c.title+', '
                        elif   p_date==3 and p.modify_date != p.publication_date:
                               cUpP_3= cUpP_3+1 
                               cUpP+=1
                               if c.title in cnsUpP_3:
                                print("NO")
                               else:    
                                cnsUpP_3=cnsUpP_3+c.title+', '

                
                array_cns_pn_.append(cNewP_4)
                array_cns_pn_.append(cNewP_5)
                array_cns_pn_.append(cNewP_6)
                array_cns_pn_.append(cNewP_7)
                array_cns_pn_.append(cNewP_8)
                array_cns_pn_.append(cNewP_9)
                array_cns_pn_.append(cNewP_10)
                array_cns_pn_.append(cNewP_11)
                array_cns_pn_.append(cNewP_12)
                array_cns_pn_.append(cNewP_1)
                array_cns_pn_.append(cNewP_2)
                array_cns_pn_.append(cNewP_3)

                array_cns_pu_.append( cUpP_4)
                array_cns_pu_.append( cUpP_5)
                array_cns_pu_.append( cUpP_6)
                array_cns_pu_.append( cUpP_7)
                array_cns_pu_.append( cUpP_8)
                array_cns_pu_.append( cUpP_9)
                array_cns_pu_.append( cUpP_10)
                array_cns_pu_.append( cUpP_11)
                array_cns_pu_.append( cUpP_12)
                array_cns_pu_.append( cUpP_1)
                array_cns_pu_.append( cUpP_2)
                array_cns_pu_.append( cUpP_3)

                array_cns_n_.append(cnsP_4)
                array_cns_n_.append(cnsP_5)
                array_cns_n_.append(cnsP_6)
                array_cns_n_.append(cnsP_7)
                array_cns_n_.append(cnsP_8)
                array_cns_n_.append(cnsP_9)
                array_cns_n_.append(cnsP_10)
                array_cns_n_.append(cnsP_11)
                array_cns_n_.append(cnsP_12)
                array_cns_n_.append(cnsP_1)
                array_cns_n_.append(cnsP_2)
                array_cns_n_.append(cnsP_3)


                array_cns_u_.append(cnsUpP_4)
                array_cns_u_.append(cnsUpP_5)
                array_cns_u_.append(cnsUpP_6)
                array_cns_u_.append(cnsUpP_7)
                array_cns_u_.append(cnsUpP_8)
                array_cns_u_.append(cnsUpP_9)
                array_cns_u_.append(cnsUpP_10)
                array_cns_u_.append(cnsUpP_11)
                array_cns_u_.append(cnsUpP_12)
                array_cns_u_.append(cnsUpP_1)
                array_cns_u_.append(cnsUpP_2)
                array_cns_u_.append(cnsUpP_3)

                array_cns_p_up_new.append(array_cns_pn_)#0
                
                array_cns_p_up_new.append(array_cns_pu_)#1

                array_cns_p_up_new.append(array_cns_n_)#2
                array_cns_p_up_new.append(array_cns_u_)#3
 
            numCP_P.append(numRepP)
            numCP_P.append(countP)
            numCP_P.append(cNewP)
            numCP_P.append(cUpP)
            numCP_P.append(cns)
            numCP_P.append(array_cns_p_up_new)#7
            region_cp_p.append(numCP_P)
        regionsPCP.append(region_cp_p)

        #-----------------ALL COUNTRIES---------
        regionsAll = []
        totALLcn=countriesperregioncp=CountryPage.objects.filter().count()
        region_all_p = []
        numRepPAll=0
        for k,v in REGIONS:
            reg = v+''
            numRepPAll=0
            countriesperregioncp=CountryPage.objects.filter(region=k)
            numCP_P = []
            numCP_P.append(reg)
            numCP_P.append(countriesperregioncp.count())
            countP=0
            cnsP=0
            cUpP=0
            cNewP=0
            cns=''
            
        

            p_count=0
            
            cNewP_4=0
            cNewP_5=0
            cNewP_6=0
            cNewP_7=0
            cNewP_8=0
            cNewP_9=0
            cNewP_10=0
            cNewP_11=0
            cNewP_12=0
            cNewP_1=0
            cNewP_2=0
            cNewP_3=0
            cUpP_4=0
            cUpP_5=0
            cUpP_6=0
            cUpP_7=0
            cUpP_8=0
            cUpP_9=0
            cUpP_10=0
            cUpP_11=0
            cUpP_12=0
            cUpP_1=0
            cUpP_2=0
            cUpP_3=0
            
            cnsP_4=""
            cnsP_5=""
            cnsP_6=""
            cnsP_7=""
            cnsP_8=""
            cnsP_9=""
            cnsP_10=""
            cnsP_11=""
            cnsP_12=""
            cnsP_1=""
            cnsP_2=""
            cnsP_3=""
            cnsUpP_4=""
            cnsUpP_5=""
            cnsUpP_6=""
            cnsUpP_7=""
            cnsUpP_8=""
            cnsUpP_9=""
            cnsUpP_10=""
            cnsUpP_11=""
            cnsUpP_12=""
            cnsUpP_1=""
            cnsUpP_2=""
            cnsUpP_3=""
           
            array_cns_p=[]
            for c in countriesperregioncp:
                array_cns_p_up_new=[]
                array_cns_pn_=[]
                array_cns_pu_=[]
                array_cns_n_=[]
                array_cns_u_=[]
                if c.id !=199:
                    pests= ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=3,is_version=False,publication_date__gte=startstartdate,publication_date__lte=enddate)
                    p_count=pests.count()

                    pests1= ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=3,is_version=False,publication_date__gte=startdate,publication_date__lte=enddate)
                    pests2= ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=3,is_version=False,modify_date__gte=startdate,modify_date__lte=enddate)


                    if p_count>0:
                        numRepPAll+=1    
                    if pests1.count()>0 or pests2.count()>0:
                        cns+=c.title+', '
                    countP+=p_count
                    #cNewP+=pests1.count()
                    #cUpP+=pests2.count()
                
                if pests1.count()>0:
                    for p in pests1:
                        p_date=p.publication_date.month
                       
                        if  p_date==4:
                          cNewP_4=cNewP_4+1
                          cNewP+=1
                          if c.title in cnsP_4:
                              print("NO")
                          else:    
                             cnsP_4=cnsP_4+c.title+', '
                    
                        elif   p_date==5:
                          cNewP_5=cNewP_5+1
                          cNewP+=1
                          if c.title in cnsP_5:
                              print("NO")
                          else:    
                             cnsP_5=cnsP_5+c.title+', '
                        elif   p_date==6:
                              cNewP_6=cNewP_6+1
                              cNewP+=1
                              if c.title in cnsP_6:
                                print("NO")
                              else:    
                                cnsP_6=cnsP_6+c.title+', '
                        elif   p_date==7:
                            
                              cNewP_7=cNewP_7+1
                              cNewP+=1
                              if c.title in cnsP_7:
                                  print("NO")
                              else:    
                               cnsP_7=cnsP_7+c.title+', '
                        elif   p_date==8:
                              print("")
                              print("")
                              #print("p_date"+str(p_date))
                              print("")
                              cNewP_8=cNewP_8+1
                              cNewP+=1
                              if c.title in cnsP_8:
                                  print("NO")
                              else:    
                                cnsP_8=cnsP_8+c.title+', '
                              print("")
                              #print(cnsP_8)
                              #print(cnsP_8)
                              print("")
                        elif   p_date==9:
                              cNewP_9=cNewP_9+1
                              cNewP+=1
                              if c.title in cnsP_9:
                                print("NO")
                              else:    
                                cnsP_9=cnsP_9+c.title+', '
                        elif   p_date==10:
                              cNewP_10=cNewP_10+1
                              cNewP+=1
                              if c.title in cnsP_10:
                                print("NO")
                              else:    
                                cnsP_10=cnsP_10+c.title+', '
                        elif   p_date==11:
                              cNewP_11=cNewP_11+1
                              cNewP+=1
                              if c.title in cnsP_11:
                                 print("NO")
                              else:    
                                cnsP_11=cnsP_11+c.title+', '
                        elif   p_date==12:
                              cNewP_12=cNewP_12+1
                              cNewP+=1
                              if c.title in cnsP_12:
                                  print("NO")
                              else:    
                                cnsP_12=cnsP_12+c.title+', '
                        elif   p_date==1:
                              cNewP_1=cNewP_1+1
                              cNewP+=1
                              if c.title in cnsP_1:
                                 print("NO")
                              else:    
                                 cnsP_1=cnsP_1+c.title+', '
                        elif   p_date==2:
                              cNewP_2=cNewP_2+1
                              cNewP+=1
                              if c.title in cnsP_2:
                                 print("NO")
                              else:    
                                cnsP_2=cnsP_2+c.title+', '
                        elif   p_date==3:
                              cNewP_3=cNewP_3+1
                              cNewP+=1
                              if c.title in cnsP_3:
                                print("NO")
                              else:    
                                  cnsP_3=cnsP_3+c.title+', '
                if pests2.count()>0:
                    for p in pests2:
                        # @type p 
                        p_date=p.modify_date.month
                        if   p_date==4 and p.modify_date != p.publication_date:
                           cUpP_4= cUpP_4+1
                           cUpP+=1
                           if c.title in cnsUpP_4:
                                print("NO")
                           else:    
                                cnsUpP_4=cnsUpP_4+c.title+', '

                        elif   p_date==5 and p.modify_date != p.publication_date:
                           cUpP_5= cUpP_5+1
                           cUpP+=1
                           if c.title in cnsUpP_5:
                                print("NO")
                           else:    
                                cnsUpP_5=cnsUpP_5+c.title+', '

                        elif   p_date==6 and p.modify_date != p.publication_date:
                               cUpP_6= cUpP_6+1
                               cUpP+=1
                               if c.title in cnsUpP_6:
                                print("NO")
                               else:    
                                cnsUpP_6=cnsUpP_6+c.title+', '

                        elif   p_date==7 and p.modify_date != p.publication_date:
                               cUpP_7= cUpP_7+1
                               cUpP+=1
                               if c.title in cnsUpP_7:
                                print("NO")
                               else:    
                                cnsUpP_7=cnsUpP_7+c.title+', '

                        elif   p_date==8 and p.modify_date != p.publication_date:
                               cUpP_8= cUpP_8+1
                               cUpP+=1
                               if c.title in cnsUpP_8:
                                print("NO")
                               else:    
                                cnsUpP_8=cnsUpP_8+c.title+', '

                        elif   p_date==9 and p.modify_date != p.publication_date:
                               cUpP_9= cUpP_9+1
                               cUpP+=1
                               if c.title in cnsUpP_9:
                                print("NO")
                               else:    
                                cnsUpP_9=cnsUpP_9+c.title+', '

                        elif   p_date==10 and p.modify_date != p.publication_date:
                               cUpP_10= cUpP_10+1
                               cUpP+=1
                               if c.title in cnsUpP_10:
                                print("NO")
                               else:    
                                cnsUpP_10=cnsUpP_10+c.title+', '

                        elif   p_date==11 and p.modify_date != p.publication_date:
                               cUpP_11= cUpP_11+1
                               cUpP+=1
                               if c.title in cnsUpP_11:
                                print("NO")
                               else:    
                                cnsUpP_11=cnsUpP_11+c.title+', '

                        elif   p_date==12 and p.modify_date != p.publication_date:
                               cUpP_12= cUpP_12+1
                               cUpP+=1
                               if c.title in cnsUpP_12:
                                print("NO")
                               else:    
                                cnsUpP_12=cnsUpP_12+c.title+', '

                        elif   p_date==1 and p.modify_date != p.publication_date:
                               cUpP_1= cUpP_1+1
                               cUpP+=1
                               if c.title in cnsUpP_1:
                                print("NO")
                               else:    
                                cnsUpP_1=cnsUpP_1+c.title+', '

                        elif   p_date==2 and p.modify_date != p.publication_date:
                               cUpP_2= cUpP_2+1
                               cUpP+=1
                               if c.title in cnsUpP_2:
                                print("NO")
                               else:    
                                cnsUpP_2=cnsUpP_2+c.title+', '

                        elif   p_date==3 and p.modify_date != p.publication_date:
                               cUpP_3= cUpP_3+1 
                               cUpP+=1
                               if c.title in cnsUpP_3:
                                 print("NO")
                               else:    
                                 cnsUpP_3=cnsUpP_3+c.title+', '

                
                array_cns_pn_.append(cNewP_4)
                array_cns_pn_.append(cNewP_5)
                array_cns_pn_.append(cNewP_6)
                array_cns_pn_.append(cNewP_7)
                array_cns_pn_.append(cNewP_8)
                array_cns_pn_.append(cNewP_9)
                array_cns_pn_.append(cNewP_10)
                array_cns_pn_.append(cNewP_11)
                array_cns_pn_.append(cNewP_12)
                array_cns_pn_.append(cNewP_1)
                array_cns_pn_.append(cNewP_2)
                array_cns_pn_.append(cNewP_3)

                array_cns_pu_.append( cUpP_4)
                array_cns_pu_.append( cUpP_5)
                array_cns_pu_.append( cUpP_6)
                array_cns_pu_.append( cUpP_7)
                array_cns_pu_.append( cUpP_8)
                array_cns_pu_.append( cUpP_9)
                array_cns_pu_.append( cUpP_10)
                array_cns_pu_.append( cUpP_11)
                array_cns_pu_.append( cUpP_12)
                array_cns_pu_.append( cUpP_1)
                array_cns_pu_.append( cUpP_2)
                array_cns_pu_.append( cUpP_3)

                array_cns_n_.append(cnsP_4)
                array_cns_n_.append(cnsP_5)
                array_cns_n_.append(cnsP_6)
                array_cns_n_.append(cnsP_7)
                array_cns_n_.append(cnsP_8)
                array_cns_n_.append(cnsP_9)
                array_cns_n_.append(cnsP_10)
                array_cns_n_.append(cnsP_11)
                array_cns_n_.append(cnsP_12)
                array_cns_n_.append(cnsP_1)
                array_cns_n_.append(cnsP_2)
                array_cns_n_.append(cnsP_3)


                array_cns_u_.append(cnsUpP_4)
                array_cns_u_.append(cnsUpP_5)
                array_cns_u_.append(cnsUpP_6)
                array_cns_u_.append(cnsUpP_7)
                array_cns_u_.append(cnsUpP_8)
                array_cns_u_.append(cnsUpP_9)
                array_cns_u_.append(cnsUpP_10)
                array_cns_u_.append(cnsUpP_11)
                array_cns_u_.append(cnsUpP_12)
                array_cns_u_.append(cnsUpP_1)
                array_cns_u_.append(cnsUpP_2)
                array_cns_u_.append(cnsUpP_3)

                
                array_cns_p_up_new.append(array_cns_pn_)
                
                array_cns_p_up_new.append(array_cns_pu_)

                array_cns_p_up_new.append(array_cns_n_)
                array_cns_p_up_new.append(array_cns_u_)
                
            numCP_P.append(numRepPAll)
            numCP_P.append(countP)
            numCP_P.append(cNewP)
            numCP_P.append(cUpP)
            numCP_P.append(cns)
            numCP_P.append(array_cns_p_up_new)#7
            region_all_p.append(numCP_P)
        regionsAll.append(region_all_p)
        
        
        #----------------
        
        regionsPCPTot=[]
        totarray=[]
        tot=0
        tot2=0
        tot4=0
        tot5=0
        
        for x in  regionsPCP[0]:
            tot+= x[2]
            tot2+= x[3]
            tot4+= x[4]
            tot5+= x[5]
        totarray.append(tot)
        totarray.append(tot2)
        totarray.append(tot4)
        totarray.append(tot5)

        regionsPCPTot.append(totarray)     
        #----------------
        totarray_months_n=[]
        totarray_months_u=[]
        tot_n_4=0
        tot_n_5=0
        tot_n_6=0
        tot_n_7=0
        tot_n_8=0
        tot_n_9=0
        tot_n_10=0
        tot_n_11=0
        tot_n_12=0
        tot_n_1=0
        tot_n_2=0
        tot_n_3=0
       
        tot_u_4=0
        tot_u_5=0
        tot_u_6=0
        tot_u_7=0
        tot_u_8=0
        tot_u_9=0
        tot_u_10=0
        tot_u_11=0
        tot_u_12=0
        tot_u_1=0
        tot_u_2=0
        tot_u_3=0
        for x in  regionsPCP[0]:
            #print('TOT-------------------')
            #print(x[7][0])
            tot_n_4+=x[7][0][0]
            tot_n_5+=x[7][0][1]
            tot_n_6+=x[7][0][2]
            tot_n_7+=x[7][0][3]
            tot_n_8+=x[7][0][4]
            tot_n_9+=x[7][0][5]
            tot_n_10+=x[7][0][6]
            tot_n_11+=x[7][0][7]
            tot_n_12+=x[7][0][8]
            tot_n_1+=x[7][0][9]
            tot_n_2+=x[7][0][10]
            tot_n_3+=x[7][0][11]
            
            #print(x[7][1])
            tot_u_4+=x[7][1][0]
            tot_u_5+=x[7][1][1]
            tot_u_6+=x[7][1][2]
            tot_u_7+=x[7][1][3]
            tot_u_8+=x[7][1][4]
            tot_u_9+=x[7][1][5]
            tot_u_10+=x[7][1][6]
            tot_u_11+=x[7][1][7]
            tot_u_12+=x[7][1][8]
            tot_u_1+=x[7][1][9]
            tot_u_2+=x[7][1][10]
            tot_u_3+=x[7][1][11]
         
        totarray_months_n.append(tot_n_4)
        totarray_months_n.append(tot_n_5)
        totarray_months_n.append(tot_n_6)
        totarray_months_n.append(tot_n_7)
        totarray_months_n.append(tot_n_8)
        totarray_months_n.append(tot_n_9)
        totarray_months_n.append(tot_n_10)
        totarray_months_n.append(tot_n_11)
        totarray_months_n.append(tot_n_12)
        totarray_months_n.append(tot_n_1)
        totarray_months_n.append(tot_n_2)
        totarray_months_n.append(tot_n_3)
        
        totarray_months_u.append(tot_u_4)
        totarray_months_u.append(tot_u_5)
        totarray_months_u.append(tot_u_6)
        totarray_months_u.append(tot_u_7)
        totarray_months_u.append(tot_u_8)
        totarray_months_u.append(tot_u_9)
        totarray_months_u.append(tot_u_10)
        totarray_months_u.append(tot_u_11)
        totarray_months_u.append(tot_u_12)
        totarray_months_u.append(tot_u_1)
        totarray_months_u.append(tot_u_2)
        totarray_months_u.append(tot_u_3)
        
        #----------------
        
        regionsALLTot=[]
        totarray1=[]
        tot=0
        tot2=0
        tot4=0
        tot5=0
        for x in  regionsAll[0]:
            tot+= x[2]
            tot2+= x[3]
            tot4+= x[4]
            tot5+= x[5]
        totarray1.append(tot)
        totarray1.append(tot2)
        totarray1.append(tot4)
        totarray1.append(tot5)

        regionsALLTot.append(totarray1)     



  #----------------
        totarray_months_n_all=[]
        totarray_months_u_all=[]
        tot_n_4=0
        tot_n_5=0
        tot_n_6=0
        tot_n_7=0
        tot_n_8=0
        tot_n_9=0
        tot_n_10=0
        tot_n_11=0
        tot_n_12=0
        tot_n_1=0
        tot_n_2=0
        tot_n_3=0
       
        tot_u_4=0
        tot_u_5=0
        tot_u_6=0
        tot_u_7=0
        tot_u_8=0
        tot_u_9=0
        tot_u_10=0
        tot_u_11=0
        tot_u_12=0
        tot_u_1=0
        tot_u_2=0
        tot_u_3=0
        for x in  regionsAll[0]:
            #print('TOT----AOLL---------------')
            #print(x[7][0])
            tot_n_4+=x[7][0][0]
            tot_n_5+=x[7][0][1]
            tot_n_6+=x[7][0][2]
            tot_n_7+=x[7][0][3]
            tot_n_8+=x[7][0][4]
            tot_n_9+=x[7][0][5]
            tot_n_10+=x[7][0][6]
            tot_n_11+=x[7][0][7]
            tot_n_12+=x[7][0][8]
            tot_n_1+=x[7][0][9]
            tot_n_2+=x[7][0][10]
            tot_n_3+=x[7][0][11]
            
            #print(x[7][1])
            tot_u_4+=x[7][1][0]
            tot_u_5+=x[7][1][1]
            tot_u_6+=x[7][1][2]
            tot_u_7+=x[7][1][3]
            tot_u_8+=x[7][1][4]
            tot_u_9+=x[7][1][5]
            tot_u_10+=x[7][1][6]
            tot_u_11+=x[7][1][7]
            tot_u_12+=x[7][1][8]
            tot_u_1+=x[7][1][9]
            tot_u_2+=x[7][1][10]
            tot_u_3+=x[7][1][11]
         
        totarray_months_n_all.append(tot_n_4)
        totarray_months_n_all.append(tot_n_5)
        totarray_months_n_all.append(tot_n_6)
        totarray_months_n_all.append(tot_n_7)
        totarray_months_n_all.append(tot_n_8)
        totarray_months_n_all.append(tot_n_9)
        totarray_months_n_all.append(tot_n_10)
        totarray_months_n_all.append(tot_n_11)
        totarray_months_n_all.append(tot_n_12)
        totarray_months_n_all.append(tot_n_1)
        totarray_months_n_all.append(tot_n_2)
        totarray_months_n_all.append(tot_n_3)
        
        totarray_months_u_all.append(tot_u_4)
        totarray_months_u_all.append(tot_u_5)
        totarray_months_u_all.append(tot_u_6)
        totarray_months_u_all.append(tot_u_7)
        totarray_months_u_all.append(tot_u_8)
        totarray_months_u_all.append(tot_u_9)
        totarray_months_u_all.append(tot_u_10)
        totarray_months_u_all.append(tot_u_11)
        totarray_months_u_all.append(tot_u_12)
        totarray_months_u_all.append(tot_u_1)
        totarray_months_u_all.append(tot_u_2)
        totarray_months_u_all.append(tot_u_3)







        context['curryear']=curryear
        context['prevyear']=prevyear
        context['regionsPCP']=regionsPCP
        context['totNumReg']=totNumReg
        context['regionsPCPTot']=regionsPCPTot
        
        context['regionsAll']=regionsAll
        context['totALLcn']=totALLcn
        context['regionsALLTot']=regionsALLTot
        context['totarray_months_n']=totarray_months_n
        context['totarray_months_u']=totarray_months_u
        context['totarray_months_n_all']=totarray_months_n_all
        context['totarray_months_u_all']=totarray_months_u_all
 
       
        pest_array=[]
        pest_array1=[]
     
        pestreporting_array = []
        pestreporting_array1 = []
        p_count=0
        for y in range(2005,curryear+1):
            pests= ReportingObligation.objects.filter(reporting_obligation_type=3,is_version=False)
              
            p_count1=0
            for p in pests:
                if p.publication_date != None and p.publication_date.year == y:
                    p_count=p_count+1
                    p_count1=p_count1+1
            pestreporting_array.append(p_count)
            pestreporting_array1.append(p_count1)
            
        pest_array.append(pestreporting_array)       
        pest_array1.append(pestreporting_array1)       
        datachartbis=''
        datachart1bis=''
        datachart2=''
        datachart3=''
        i=0
     
        for y in range(2005,curryear+1 ):
            datachartbis += '{type: "column", name: "'+str(y)+'", legendText: "'+str(y)+'",showInLegend: true, dataPoints:[{label: "List of Regulated Pests", y: '+str(pest_array[0][i])+'},]},'
            datachart1bis +='{type: "column", name: "'+str(y)+'", legendText: "'+str(y)+'",showInLegend: true, dataPoints:[{label: "List of Regulated Pests", y: '+str(pest_array1[0][i])+'},]},'
            
            datachart2+= '{label: "'+str(y)+'", y: '+str(pest_array[0][i])+'}, '
            datachart3+= '{label: "'+str(y)+'", y: '+str(pest_array1[0][i])+'}, '
            i=i+1
        context['pest_array']=pest_array
        context['pest_array1']=pest_array1
        context['num_years_range']=range(2005,curryear+1)
        context['datachartbis']=datachartbis
        context['datachart1bis']=datachart1bis
        context['num_years']=num_years
        context['datachart2']=datachart2
        context['datachart3']=datachart3
        
        return context    



class CountryStatsSingleLegislationsListView(ListView):
    """   stat  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_year_of_legislations.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryStatsSingleLegislationsListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
        context['selyear_range']=range(2010,timezone.now().year+1)
       
        curryear=0
        prevyear=0
        num_years=0
        
        #if 'year' in self.kwargs:
        #    curryear=int(self.kwargs['year'])
        #else:   
            #curryear=timezone.now().year+1
        curryear=2018   
        prevyear=curryear-1
        num_years=curryear-2005
        
        startstartdate = datetime(1972, 1, 1, 00, 01,00)
        startdate = datetime(prevyear, 4, 1, 00, 01,00)
        enddate = datetime(curryear, 3, 31, 23, 59,00)
      
        
        regionsPCP = []
        totNumReg=CountryPage.objects.filter(cp_ncp_t_type='CP').count()
        region_cp_p = []
        numRepP=0
           
        for k,v in REGIONS:
            reg = v+''
            numRepP=0
            countriesperregioncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='CP')
            numCP_P = []
            numCP_P.append(reg)
            numCP_P.append(countriesperregioncp.count())
            countP=0
            cNewP=0
            cUpP=0
            cns=''
            
            cns_4=''
            cns_5=''
            cns_6=''
            cns_7=''
            cns_8=''
            cns_9=''
            cns_10=''
            cns_11=''
            cns_12=''
            cns_1=''
            cns_2=''
            cns_3=''
            tot_4=''
            tot_5=''
            tot_6=''
            tot_7=''
            tot_8=''
            tot_9=''
            tot_10=''
            tot_11=''
            tot_12=''
            tot_1=''
            tot_2=''
            tot_3=''
            p_count=0
            
            cNewP_4=0
            cNewP_5=0
            cNewP_6=0
            cNewP_7=0
            cNewP_8=0
            cNewP_9=0
            cNewP_10=0
            cNewP_11=0
            cNewP_12=0
            cNewP_1=0
            cNewP_2=0
            cNewP_3=0
            cUpP_4=0
            cUpP_5=0
            cUpP_6=0
            cUpP_7=0
            cUpP_8=0
            cUpP_9=0
            cUpP_10=0
            cUpP_11=0
            cUpP_12=0
            cUpP_1=0
            cUpP_2=0
            cUpP_3=0
           
            cnsP_4=""
            cnsP_5=""
            cnsP_6=""
            cnsP_7=""
            cnsP_8=""
            cnsP_9=""
            cnsP_10=""
            cnsP_11=""
            cnsP_12=""
            cnsP_1=""
            cnsP_2=""
            cnsP_3=""
            cnsUpP_4=""
            cnsUpP_5=""
            cnsUpP_6=""
            cnsUpP_7=""
            cnsUpP_8=""
            cnsUpP_9=""
            cnsUpP_10=""
            cnsUpP_11=""
            cnsUpP_12=""
            cnsUpP_1=""
            cnsUpP_2=""
            cnsUpP_3=""
           
            array_cns_p=[]
            for c in countriesperregioncp:
                array_cns_p_up_new=[]
                array_cns_pn_=[]
                array_cns_pu_=[]
                array_cns_n_=[]
                array_cns_u_=[]
                pests= ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=4,is_version=False,publication_date__gte=startstartdate,publication_date__lte=enddate)
                p_count=pests.count()
                
                pests1= ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=4,is_version=False,publication_date__gte=startdate,publication_date__lte=enddate)
                pests2= ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=4,is_version=False,modify_date__gte=startdate,modify_date__lte=enddate)
           
           
                if p_count>0:
                    numRepP+=1    
                if pests1.count()>0 or pests2.count()>0:
                    cns+=c.title+', '
                    
                countP+=p_count
                #cNewP+=pests1.count()
                #cUpP+=pests2.count()
                
                if pests1.count()>0:
                    for p in pests1:
                        p_date=p.publication_date.month
                       
                        if p_date==4:
                          #print(p_date)  
                          cNewP_4=cNewP_4+1
                          cNewP+=1
                          #print(c.title)  
                          if c.title in cnsP_4:
                              print("NO")
                          else:    
                             cnsP_4=cnsP_4+c.title+', '
                    
                        elif   p_date==5:
                          cNewP_5=cNewP_5+1
                          cNewP+=1
                          if c.title in cnsP_5:
                              print("NO")
                          else:    
                             cnsP_5=cnsP_5+c.title+', '
                        elif   p_date==6:
                              cNewP_6=cNewP_6+1
                              cNewP+=1
                              if c.title in cnsP_6:
                                print("NO")
                              else:    
                                cnsP_6=cnsP_6+c.title+', '
                        elif   p_date==5:
                              cNewP_7=cNewP_7+1
                              cNewP+=1
                              if c.title in cnsP_7:
                                  print("NO")
                              else:    
                               cnsP_7=cnsP_7+c.title+', '
                        elif   p_date==8:
                              cNewP_8=cNewP_8+1
                              cNewP+=1
                              if c.title in cnsP_8:
                                  print("NO")
                              else:    
                                cnsP_8=cnsP_8+c.title+', '
                        elif   p_date==9:
                              cNewP_9=cNewP_9+1
                              cNewP+=1
                              if c.title in cnsP_9:
                                print("NO")
                              else:    
                                cnsP_9=cnsP_9+c.title+', '
                        elif   p_date==10:
                              cNewP_10=cNewP_10+1
                              cNewP+=1
                              if c.title in cnsP_10:
                                print("NO")
                              else:    
                                cnsP_10=cnsP_10+c.title+', '
                        elif   p_date==11:
                              cNewP_11=cNewP_11+1
                              cNewP+=1
                              if c.title in cnsP_11:
                                 print("NO")
                              else:    
                                cnsP_11=cnsP_11+c.title+', '
                        elif   p_date==12:
                              cNewP_12=cNewP_12+1
                              cNewP+=1
                              if c.title in cnsP_12:
                                  print("NO")
                              else:    
                                cnsP_12=cnsP_12+c.title+', '
                        elif   p_date==1:
                              cNewP_1=cNewP_1+1
                              cNewP+=1
                              if c.title in cnsP_1:
                                 print("NO")
                              else:    
                                 cnsP_1=cnsP_1+c.title+', '
                        elif   p_date==2:
                              cNewP_2=cNewP_2+1
                              cNewP+=1
                              if c.title in cnsP_2:
                                 print("NO")
                              else:    
                                cnsP_2=cnsP_2+c.title+', '
                        elif   p_date==3:
                              cNewP_3=cNewP_3+1
                              cNewP+=1
                              if c.title in cnsP_3:
                                print("NO")
                              else:    
                                  cnsP_3=cnsP_3+c.title+', '
                if pests2.count()>0:
                    for p in pests2:
                        # @type p 
                        p_date=p.modify_date.month
                        if   p_date==4 and p.modify_date != p.publication_date:
                           cUpP_4= cUpP_4+1
                           cUpP+=1
                           if c.title in cnsUpP_4:
                              print("NO")
                           else:    
                            cnsUpP_4=cnsUpP_4+c.title+', '
                        elif   p_date==5 and p.modify_date != p.publication_date:
                           cUpP_5= cUpP_5+1
                           cUpP+=1
                           if c.title in cnsUpP_5:
                              print("NO")
                           else:    
                            cnsUpP_5=cnsUpP_5+c.title+', '

                        elif   p_date==6 and p.modify_date != p.publication_date:
                               cUpP_6= cUpP_6+1
                               cUpP+=1
                               if c.title in cnsUpP_6:
                                 print("NO")
                               else:    
                                  cnsUpP_6=cnsUpP_6+c.title+', '
                        elif   p_date==7 and p.modify_date != p.publication_date:
                               cUpP_7= cUpP_7+1
                               cUpP+=1
                               if c.title in cnsUpP_7:
                                 print("NO")
                               else:    
                                 cnsUpP_7=cnsUpP_7+c.title+', '
                        elif   p_date==8 and p.modify_date != p.publication_date:
                               cUpP_8= cUpP_8+1
                               cUpP+=1
                               #print('****************************')
                               #print('****************************cUpP_8='+str(cUpP_8))
                               if c.title in cnsUpP_8:
                                 print("NO")
                               else:    
                                  cnsUpP_8=cnsUpP_8+c.title+', '
                        elif   p_date==9 and p.modify_date != p.publication_date:
                               cUpP_9= cUpP_9+1
                               cUpP+=1
                               if c.title in cnsUpP_9:
                                 print("NO")
                               else:    
                                 cnsUpP_9=cnsUpP_9+c.title+', '
                        elif   p_date==10 and p.modify_date != p.publication_date:
                               cUpP_10= cUpP_10+1
                               cUpP+=1
                               if c.title in cnsUpP_10:
                                 print("NO")
                               else:    
                                  cnsUpP_10=cnsUpP_10+c.title+', '
                        elif   p_date==11 and p.modify_date != p.publication_date:
                               cUpP_11= cUpP_11+1
                               cUpP+=1
                               if c.title in cnsUpP_11:
                                print("NO")
                               else:    
                                 cnsUpP_11=cnsUpP_11+c.title+', '
                        elif   p_date==12 and p.modify_date != p.publication_date:
                               cUpP_12= cUpP_12+1
                               cUpP+=1
                               if c.title in cnsUpP_12:
                                print("NO")
                               else:    
                                 cnsUpP_12=cnsUpP_12+c.title+', '
                        elif   p_date==1 and p.modify_date != p.publication_date:
                               cUpP_1= cUpP_1+1
                               cUpP+=1
                               if c.title in cnsUpP_1:
                                print("NO")
                               else:    
                                 cnsUpP_1=cnsUpP_1+c.title+', '
                        elif   p_date==2 and p.modify_date != p.publication_date:
                               cUpP_2= cUpP_2+1
                               cUpP+=1
                               if c.title in cnsUpP_2:
                                    print("NO")
                               else:    
                                     cnsUpP_2=cnsUpP_2+c.title+', '
                        elif   p_date==3 and p.modify_date != p.publication_date:
                               cUpP_3= cUpP_3+1 
                               cUpP+=1
                               if c.title in cnsUpP_3:
                                print("NO")
                               else:    
                                cnsUpP_3=cnsUpP_3+c.title+', '

                
                array_cns_pn_.append(cNewP_4)
                array_cns_pn_.append(cNewP_5)
                array_cns_pn_.append(cNewP_6)
                array_cns_pn_.append(cNewP_7)
                array_cns_pn_.append(cNewP_8)
                array_cns_pn_.append(cNewP_9)
                array_cns_pn_.append(cNewP_10)
                array_cns_pn_.append(cNewP_11)
                array_cns_pn_.append(cNewP_12)
                array_cns_pn_.append(cNewP_1)
                array_cns_pn_.append(cNewP_2)
                array_cns_pn_.append(cNewP_3)

                array_cns_pu_.append( cUpP_4)
                array_cns_pu_.append( cUpP_5)
                array_cns_pu_.append( cUpP_6)
                array_cns_pu_.append( cUpP_7)
                array_cns_pu_.append( cUpP_8)
                array_cns_pu_.append( cUpP_9)
                array_cns_pu_.append( cUpP_10)
                array_cns_pu_.append( cUpP_11)
                array_cns_pu_.append( cUpP_12)
                array_cns_pu_.append( cUpP_1)
                array_cns_pu_.append( cUpP_2)
                array_cns_pu_.append( cUpP_3)

                array_cns_n_.append(cnsP_4)
                array_cns_n_.append(cnsP_5)
                array_cns_n_.append(cnsP_6)
                array_cns_n_.append(cnsP_7)
                array_cns_n_.append(cnsP_8)
                array_cns_n_.append(cnsP_9)
                array_cns_n_.append(cnsP_10)
                array_cns_n_.append(cnsP_11)
                array_cns_n_.append(cnsP_12)
                array_cns_n_.append(cnsP_1)
                array_cns_n_.append(cnsP_2)
                array_cns_n_.append(cnsP_3)


                array_cns_u_.append(cnsUpP_4)
                array_cns_u_.append(cnsUpP_5)
                array_cns_u_.append(cnsUpP_6)
                array_cns_u_.append(cnsUpP_7)
                array_cns_u_.append(cnsUpP_8)
                array_cns_u_.append(cnsUpP_9)
                array_cns_u_.append(cnsUpP_10)
                array_cns_u_.append(cnsUpP_11)
                array_cns_u_.append(cnsUpP_12)
                array_cns_u_.append(cnsUpP_1)
                array_cns_u_.append(cnsUpP_2)
                array_cns_u_.append(cnsUpP_3)

                array_cns_p_up_new.append(array_cns_pn_)#0
                
                array_cns_p_up_new.append(array_cns_pu_)#1

                array_cns_p_up_new.append(array_cns_n_)#2
                array_cns_p_up_new.append(array_cns_u_)#3
 
            numCP_P.append(numRepP)
            numCP_P.append(countP)
            numCP_P.append(cNewP)
            numCP_P.append(cUpP)
            numCP_P.append(cns)
            numCP_P.append(array_cns_p_up_new)#7
            region_cp_p.append(numCP_P)
        regionsPCP.append(region_cp_p)

        #-----------------ALL COUNTRIES---------
        regionsAll = []
        totALLcn=countriesperregioncp=CountryPage.objects.filter().count()
        region_all_p = []
        numRepPAll=0
        for k,v in REGIONS:
            reg = v+''
            numRepPAll=0
            countriesperregioncp=CountryPage.objects.filter(region=k)
            numCP_P = []
            numCP_P.append(reg)
            numCP_P.append(countriesperregioncp.count())
            countP=0
            cnsP=0
            cUpP=0
            cNewP=0
            cns=''
            
        

            p_count=0
            
            cNewP_4=0
            cNewP_5=0
            cNewP_6=0
            cNewP_7=0
            cNewP_8=0
            cNewP_9=0
            cNewP_10=0
            cNewP_11=0
            cNewP_12=0
            cNewP_1=0
            cNewP_2=0
            cNewP_3=0
            cUpP_4=0
            cUpP_5=0
            cUpP_6=0
            cUpP_7=0
            cUpP_8=0
            cUpP_9=0
            cUpP_10=0
            cUpP_11=0
            cUpP_12=0
            cUpP_1=0
            cUpP_2=0
            cUpP_3=0
            
            cnsP_4=""
            cnsP_5=""
            cnsP_6=""
            cnsP_7=""
            cnsP_8=""
            cnsP_9=""
            cnsP_10=""
            cnsP_11=""
            cnsP_12=""
            cnsP_1=""
            cnsP_2=""
            cnsP_3=""
            cnsUpP_4=""
            cnsUpP_5=""
            cnsUpP_6=""
            cnsUpP_7=""
            cnsUpP_8=""
            cnsUpP_9=""
            cnsUpP_10=""
            cnsUpP_11=""
            cnsUpP_12=""
            cnsUpP_1=""
            cnsUpP_2=""
            cnsUpP_3=""
           
            array_cns_p=[]
            for c in countriesperregioncp:
                array_cns_p_up_new=[]
                array_cns_pn_=[]
                array_cns_pu_=[]
                array_cns_n_=[]
                array_cns_u_=[]
                if c.id !=199:
                    pests= ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=4,is_version=False,publication_date__gte=startstartdate,publication_date__lte=enddate)
                    p_count=pests.count()

                    pests1= ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=4,is_version=False,publication_date__gte=startdate,publication_date__lte=enddate)
                    pests2= ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=4,is_version=False,modify_date__gte=startdate,modify_date__lte=enddate)


                    if p_count>0:
                        numRepPAll+=1    
                    if pests1.count()>0 or pests2.count()>0:
                        cns+=c.title+', '
                    countP+=p_count
                    #cNewP+=pests1.count()
                    #cUpP+=pests2.count()
                
                if pests1.count()>0:
                    for p in pests1:
                        p_date=p.publication_date.month
                       
                        if  p_date==4:
                          cNewP_4=cNewP_4+1
                          cNewP+=1
                          if c.title in cnsP_4:
                              print("NO")
                          else:    
                             cnsP_4=cnsP_4+c.title+', '
                    
                        elif   p_date==5:
                          cNewP_5=cNewP_5+1
                          cNewP+=1
                          if c.title in cnsP_5:
                              print("NO")
                          else:    
                             cnsP_5=cnsP_5+c.title+', '
                        elif   p_date==6:
                              cNewP_6=cNewP_6+1
                              cNewP+=1
                              if c.title in cnsP_6:
                                print("NO")
                              else:    
                                cnsP_6=cnsP_6+c.title+', '
                        elif   p_date==7:
                            
                              cNewP_7=cNewP_7+1
                              cNewP+=1
                              if c.title in cnsP_7:
                                  print("NO")
                              else:    
                               cnsP_7=cnsP_7+c.title+', '
                        elif   p_date==8:
                              print("")
                              print("")
                              #print("p_date"+str(p_date))
                              print("")
                              cNewP_8=cNewP_8+1
                              cNewP+=1
                              if c.title in cnsP_8:
                                  print("NO")
                              else:    
                                cnsP_8=cnsP_8+c.title+', '
                              print("")
                              #print(cnsP_8)
                              #print(cnsP_8)
                              print("")
                        elif   p_date==9:
                              cNewP_9=cNewP_9+1
                              cNewP+=1
                              if c.title in cnsP_9:
                                print("NO")
                              else:    
                                cnsP_9=cnsP_9+c.title+', '
                        elif   p_date==10:
                              cNewP_10=cNewP_10+1
                              cNewP+=1
                              if c.title in cnsP_10:
                                print("NO")
                              else:    
                                cnsP_10=cnsP_10+c.title+', '
                        elif   p_date==11:
                              cNewP_11=cNewP_11+1
                              cNewP+=1
                              if c.title in cnsP_11:
                                 print("NO")
                              else:    
                                cnsP_11=cnsP_11+c.title+', '
                        elif   p_date==12:
                              cNewP_12=cNewP_12+1
                              cNewP+=1
                              if c.title in cnsP_12:
                                  print("NO")
                              else:    
                                cnsP_12=cnsP_12+c.title+', '
                        elif   p_date==1:
                              cNewP_1=cNewP_1+1
                              cNewP+=1
                              if c.title in cnsP_1:
                                 print("NO")
                              else:    
                                 cnsP_1=cnsP_1+c.title+', '
                        elif   p_date==2:
                              cNewP_2=cNewP_2+1
                              cNewP+=1
                              if c.title in cnsP_2:
                                 print("NO")
                              else:    
                                cnsP_2=cnsP_2+c.title+', '
                        elif   p_date==3:
                              cNewP_3=cNewP_3+1
                              cNewP+=1
                              if c.title in cnsP_3:
                                print("NO")
                              else:    
                                  cnsP_3=cnsP_3+c.title+', '
                if pests2.count()>0:
                    for p in pests2:
                        # @type p 
                        p_date=p.modify_date.month
                        if   p_date==4 and p.modify_date != p.publication_date:
                           cUpP_4= cUpP_4+1
                           cUpP+=1
                           if c.title in cnsUpP_4:
                                print("NO")
                           else:    
                                cnsUpP_4=cnsUpP_4+c.title+', '

                        elif   p_date==5 and p.modify_date != p.publication_date:
                           cUpP_5= cUpP_5+1
                           cUpP+=1
                           if c.title in cnsUpP_5:
                                print("NO")
                           else:    
                                cnsUpP_5=cnsUpP_5+c.title+', '

                        elif   p_date==6 and p.modify_date != p.publication_date:
                               cUpP_6= cUpP_6+1
                               cUpP+=1
                               if c.title in cnsUpP_6:
                                print("NO")
                               else:    
                                cnsUpP_6=cnsUpP_6+c.title+', '

                        elif   p_date==7 and p.modify_date != p.publication_date:
                               cUpP_7= cUpP_7+1
                               cUpP+=1
                               if c.title in cnsUpP_7:
                                print("NO")
                               else:    
                                cnsUpP_7=cnsUpP_7+c.title+', '

                        elif   p_date==8 and p.modify_date != p.publication_date:
                               cUpP_8= cUpP_8+1
                               cUpP+=1
                               if c.title in cnsUpP_8:
                                print("NO")
                               else:    
                                cnsUpP_8=cnsUpP_8+c.title+', '

                        elif   p_date==9 and p.modify_date != p.publication_date:
                               cUpP_9= cUpP_9+1
                               cUpP+=1
                               if c.title in cnsUpP_9:
                                print("NO")
                               else:    
                                cnsUpP_9=cnsUpP_9+c.title+', '

                        elif   p_date==10 and p.modify_date != p.publication_date:
                               cUpP_10= cUpP_10+1
                               cUpP+=1
                               if c.title in cnsUpP_10:
                                print("NO")
                               else:    
                                cnsUpP_10=cnsUpP_10+c.title+', '

                        elif   p_date==11 and p.modify_date != p.publication_date:
                               cUpP_11= cUpP_11+1
                               cUpP+=1
                               if c.title in cnsUpP_11:
                                print("NO")
                               else:    
                                cnsUpP_11=cnsUpP_11+c.title+', '

                        elif   p_date==12 and p.modify_date != p.publication_date:
                               cUpP_12= cUpP_12+1
                               cUpP+=1
                               if c.title in cnsUpP_12:
                                print("NO")
                               else:    
                                cnsUpP_12=cnsUpP_12+c.title+', '

                        elif   p_date==1 and p.modify_date != p.publication_date:
                               cUpP_1= cUpP_1+1
                               cUpP+=1
                               if c.title in cnsUpP_1:
                                print("NO")
                               else:    
                                cnsUpP_1=cnsUpP_1+c.title+', '

                        elif   p_date==2 and p.modify_date != p.publication_date:
                               cUpP_2= cUpP_2+1
                               cUpP+=1
                               if c.title in cnsUpP_2:
                                print("NO")
                               else:    
                                cnsUpP_2=cnsUpP_2+c.title+', '

                        elif   p_date==3 and p.modify_date != p.publication_date:
                               cUpP_3= cUpP_3+1 
                               cUpP+=1
                               if c.title in cnsUpP_3:
                                 print("NO")
                               else:    
                                 cnsUpP_3=cnsUpP_3+c.title+', '

                
                array_cns_pn_.append(cNewP_4)
                array_cns_pn_.append(cNewP_5)
                array_cns_pn_.append(cNewP_6)
                array_cns_pn_.append(cNewP_7)
                array_cns_pn_.append(cNewP_8)
                array_cns_pn_.append(cNewP_9)
                array_cns_pn_.append(cNewP_10)
                array_cns_pn_.append(cNewP_11)
                array_cns_pn_.append(cNewP_12)
                array_cns_pn_.append(cNewP_1)
                array_cns_pn_.append(cNewP_2)
                array_cns_pn_.append(cNewP_3)

                array_cns_pu_.append( cUpP_4)
                array_cns_pu_.append( cUpP_5)
                array_cns_pu_.append( cUpP_6)
                array_cns_pu_.append( cUpP_7)
                array_cns_pu_.append( cUpP_8)
                array_cns_pu_.append( cUpP_9)
                array_cns_pu_.append( cUpP_10)
                array_cns_pu_.append( cUpP_11)
                array_cns_pu_.append( cUpP_12)
                array_cns_pu_.append( cUpP_1)
                array_cns_pu_.append( cUpP_2)
                array_cns_pu_.append( cUpP_3)

                array_cns_n_.append(cnsP_4)
                array_cns_n_.append(cnsP_5)
                array_cns_n_.append(cnsP_6)
                array_cns_n_.append(cnsP_7)
                array_cns_n_.append(cnsP_8)
                array_cns_n_.append(cnsP_9)
                array_cns_n_.append(cnsP_10)
                array_cns_n_.append(cnsP_11)
                array_cns_n_.append(cnsP_12)
                array_cns_n_.append(cnsP_1)
                array_cns_n_.append(cnsP_2)
                array_cns_n_.append(cnsP_3)


                array_cns_u_.append(cnsUpP_4)
                array_cns_u_.append(cnsUpP_5)
                array_cns_u_.append(cnsUpP_6)
                array_cns_u_.append(cnsUpP_7)
                array_cns_u_.append(cnsUpP_8)
                array_cns_u_.append(cnsUpP_9)
                array_cns_u_.append(cnsUpP_10)
                array_cns_u_.append(cnsUpP_11)
                array_cns_u_.append(cnsUpP_12)
                array_cns_u_.append(cnsUpP_1)
                array_cns_u_.append(cnsUpP_2)
                array_cns_u_.append(cnsUpP_3)

                
                array_cns_p_up_new.append(array_cns_pn_)
                
                array_cns_p_up_new.append(array_cns_pu_)

                array_cns_p_up_new.append(array_cns_n_)
                array_cns_p_up_new.append(array_cns_u_)
                
            numCP_P.append(numRepPAll)
            numCP_P.append(countP)
            numCP_P.append(cNewP)
            numCP_P.append(cUpP)
            numCP_P.append(cns)
            numCP_P.append(array_cns_p_up_new)#7
            region_all_p.append(numCP_P)
        regionsAll.append(region_all_p)
        
        
        #----------------
        
        regionsPCPTot=[]
        totarray=[]
        tot=0
        tot2=0
        tot4=0
        tot5=0
        
        for x in  regionsPCP[0]:
            tot+= x[2]
            tot2+= x[3]
            tot4+= x[4]
            tot5+= x[5]
        totarray.append(tot)
        totarray.append(tot2)
        totarray.append(tot4)
        totarray.append(tot5)

        regionsPCPTot.append(totarray)     
        #----------------
        totarray_months_n=[]
        totarray_months_u=[]
        tot_n_4=0
        tot_n_5=0
        tot_n_6=0
        tot_n_7=0
        tot_n_8=0
        tot_n_9=0
        tot_n_10=0
        tot_n_11=0
        tot_n_12=0
        tot_n_1=0
        tot_n_2=0
        tot_n_3=0
       
        tot_u_4=0
        tot_u_5=0
        tot_u_6=0
        tot_u_7=0
        tot_u_8=0
        tot_u_9=0
        tot_u_10=0
        tot_u_11=0
        tot_u_12=0
        tot_u_1=0
        tot_u_2=0
        tot_u_3=0
        for x in  regionsPCP[0]:
            #print('TOT-------------------')
            #print(x[7][0])
            tot_n_4+=x[7][0][0]
            tot_n_5+=x[7][0][1]
            tot_n_6+=x[7][0][2]
            tot_n_7+=x[7][0][3]
            tot_n_8+=x[7][0][4]
            tot_n_9+=x[7][0][5]
            tot_n_10+=x[7][0][6]
            tot_n_11+=x[7][0][7]
            tot_n_12+=x[7][0][8]
            tot_n_1+=x[7][0][9]
            tot_n_2+=x[7][0][10]
            tot_n_3+=x[7][0][11]
            
            #print(x[7][1])
            tot_u_4+=x[7][1][0]
            tot_u_5+=x[7][1][1]
            tot_u_6+=x[7][1][2]
            tot_u_7+=x[7][1][3]
            tot_u_8+=x[7][1][4]
            tot_u_9+=x[7][1][5]
            tot_u_10+=x[7][1][6]
            tot_u_11+=x[7][1][7]
            tot_u_12+=x[7][1][8]
            tot_u_1+=x[7][1][9]
            tot_u_2+=x[7][1][10]
            tot_u_3+=x[7][1][11]
         
        totarray_months_n.append(tot_n_4)
        totarray_months_n.append(tot_n_5)
        totarray_months_n.append(tot_n_6)
        totarray_months_n.append(tot_n_7)
        totarray_months_n.append(tot_n_8)
        totarray_months_n.append(tot_n_9)
        totarray_months_n.append(tot_n_10)
        totarray_months_n.append(tot_n_11)
        totarray_months_n.append(tot_n_12)
        totarray_months_n.append(tot_n_1)
        totarray_months_n.append(tot_n_2)
        totarray_months_n.append(tot_n_3)
        
        totarray_months_u.append(tot_u_4)
        totarray_months_u.append(tot_u_5)
        totarray_months_u.append(tot_u_6)
        totarray_months_u.append(tot_u_7)
        totarray_months_u.append(tot_u_8)
        totarray_months_u.append(tot_u_9)
        totarray_months_u.append(tot_u_10)
        totarray_months_u.append(tot_u_11)
        totarray_months_u.append(tot_u_12)
        totarray_months_u.append(tot_u_1)
        totarray_months_u.append(tot_u_2)
        totarray_months_u.append(tot_u_3)
        
        #----------------
        
        regionsALLTot=[]
        totarray1=[]
        tot=0
        tot2=0
        tot4=0
        tot5=0
        for x in  regionsAll[0]:
            tot+= x[2]
            tot2+= x[3]
            tot4+= x[4]
            tot5+= x[5]
        totarray1.append(tot)
        totarray1.append(tot2)
        totarray1.append(tot4)
        totarray1.append(tot5)

        regionsALLTot.append(totarray1)     



  #----------------
        totarray_months_n_all=[]
        totarray_months_u_all=[]
        tot_n_4=0
        tot_n_5=0
        tot_n_6=0
        tot_n_7=0
        tot_n_8=0
        tot_n_9=0
        tot_n_10=0
        tot_n_11=0
        tot_n_12=0
        tot_n_1=0
        tot_n_2=0
        tot_n_3=0
       
        tot_u_4=0
        tot_u_5=0
        tot_u_6=0
        tot_u_7=0
        tot_u_8=0
        tot_u_9=0
        tot_u_10=0
        tot_u_11=0
        tot_u_12=0
        tot_u_1=0
        tot_u_2=0
        tot_u_3=0
        for x in  regionsAll[0]:
            #print('TOT----AOLL---------------')
            #print(x[7][0])
            tot_n_4+=x[7][0][0]
            tot_n_5+=x[7][0][1]
            tot_n_6+=x[7][0][2]
            tot_n_7+=x[7][0][3]
            tot_n_8+=x[7][0][4]
            tot_n_9+=x[7][0][5]
            tot_n_10+=x[7][0][6]
            tot_n_11+=x[7][0][7]
            tot_n_12+=x[7][0][8]
            tot_n_1+=x[7][0][9]
            tot_n_2+=x[7][0][10]
            tot_n_3+=x[7][0][11]
            
            #print(x[7][1])
            tot_u_4+=x[7][1][0]
            tot_u_5+=x[7][1][1]
            tot_u_6+=x[7][1][2]
            tot_u_7+=x[7][1][3]
            tot_u_8+=x[7][1][4]
            tot_u_9+=x[7][1][5]
            tot_u_10+=x[7][1][6]
            tot_u_11+=x[7][1][7]
            tot_u_12+=x[7][1][8]
            tot_u_1+=x[7][1][9]
            tot_u_2+=x[7][1][10]
            tot_u_3+=x[7][1][11]
         
        totarray_months_n_all.append(tot_n_4)
        totarray_months_n_all.append(tot_n_5)
        totarray_months_n_all.append(tot_n_6)
        totarray_months_n_all.append(tot_n_7)
        totarray_months_n_all.append(tot_n_8)
        totarray_months_n_all.append(tot_n_9)
        totarray_months_n_all.append(tot_n_10)
        totarray_months_n_all.append(tot_n_11)
        totarray_months_n_all.append(tot_n_12)
        totarray_months_n_all.append(tot_n_1)
        totarray_months_n_all.append(tot_n_2)
        totarray_months_n_all.append(tot_n_3)
        
        totarray_months_u_all.append(tot_u_4)
        totarray_months_u_all.append(tot_u_5)
        totarray_months_u_all.append(tot_u_6)
        totarray_months_u_all.append(tot_u_7)
        totarray_months_u_all.append(tot_u_8)
        totarray_months_u_all.append(tot_u_9)
        totarray_months_u_all.append(tot_u_10)
        totarray_months_u_all.append(tot_u_11)
        totarray_months_u_all.append(tot_u_12)
        totarray_months_u_all.append(tot_u_1)
        totarray_months_u_all.append(tot_u_2)
        totarray_months_u_all.append(tot_u_3)







        context['curryear']=curryear
        context['prevyear']=prevyear
        context['regionsPCP']=regionsPCP
        context['totNumReg']=totNumReg
        context['regionsPCPTot']=regionsPCPTot
        
        context['regionsAll']=regionsAll
        context['totALLcn']=totALLcn
        context['regionsALLTot']=regionsALLTot
        context['totarray_months_n']=totarray_months_n
        context['totarray_months_u']=totarray_months_u
        context['totarray_months_n_all']=totarray_months_n_all
        context['totarray_months_u_all']=totarray_months_u_all
 
       
        pest_array=[]
        pest_array1=[]
     
        pestreporting_array = []
        pestreporting_array1 = []
        p_count=0
        for y in range(2005,curryear+1):
            pests= ReportingObligation.objects.filter(reporting_obligation_type=4,is_version=False)
              
            p_count1=0
            for p in pests:
                if p.publication_date != None and p.publication_date.year == y:
                    p_count=p_count+1
                    p_count1=p_count1+1
            pestreporting_array.append(p_count)
            pestreporting_array1.append(p_count1)
            
        pest_array.append(pestreporting_array)       
        pest_array1.append(pestreporting_array1)       
        datachartbis=''
        datachart1bis=''
        datachart2=''
        datachart3=''
        i=0
     
        for y in range(2005,curryear+1 ):
            datachartbis += '{type: "column", name: "'+str(y)+'", legendText: "'+str(y)+'",showInLegend: true, dataPoints:[{label: "Legislations", y: '+str(pest_array[0][i])+'},]},'
            datachart1bis +='{type: "column", name: "'+str(y)+'", legendText: "'+str(y)+'",showInLegend: true, dataPoints:[{label: "Legislations", y: '+str(pest_array1[0][i])+'},]},'
            
            datachart2+= '{label: "'+str(y)+'", y: '+str(pest_array[0][i])+'}, '
            datachart3+= '{label: "'+str(y)+'", y: '+str(pest_array1[0][i])+'}, '
            i=i+1
        context['pest_array']=pest_array
        context['pest_array1']=pest_array1
        context['num_years_range']=range(2005,curryear+1)
        context['datachartbis']=datachartbis
        context['datachart1bis']=datachart1bis
        context['num_years']=num_years
        context['datachart2']=datachart2
        context['datachart3']=datachart3
        
        return context    

class CountryStatsSinglePestReportsListView(ListView):
    """   stat  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_year_of_pestreport.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryStatsSinglePestReportsListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
        context['selyear_range']=range(2010,timezone.now().year+1)
       
        curryear=0
        prevyear=0
        num_years=0
        
#        if 'year' in self.kwargs:
#            curryear=int(self.kwargs['year'])
#        else :   
#            curryear=timezone.now().year
        curryear=2017
        prevyear=curryear-1  
        num_years=curryear-2005
        
        startstartdate = datetime(1972, 1, 1, 00, 01,00)
        startdate = datetime(prevyear, 4, 1, 00, 01,00)
        enddate = datetime(curryear, 3, 31, 23, 59,00)
        
        regionsPCP = []
     #   totNumReg=countriesperregioncp=CountryPage.objects.filter(cp_ncp_t_type='CP').count()
        totNumReg=CountryPage.objects.filter(cp_ncp_t_type='CP').count()
        region_cp_p = []
        numRepP=0
           
        for k,v in REGIONS:
            reg = v+''
            numRepP=0
            countriesperregioncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='CP')
            numCP_P = []
            numCP_P.append(reg)
            numCP_P.append(countriesperregioncp.count())
            countP=0
            cNewP=0
            cUpP=0
            cns=''
            p_count=0
            cNewP=0
            cUpP=0
            for c in countriesperregioncp:
                pests= PestReport.objects.filter(country=c.id,is_version=False,publish_date__gte=startstartdate,publish_date__lte=enddate)
                p_count=pests.count()
                pests1= PestReport.objects.filter(country=c.id,is_version=False,publish_date__gte=startdate,publish_date__lte=enddate)
                pests2= PestReport.objects.filter(country=c.id,is_version=False,modify_date__gte=startdate,modify_date__lte=enddate)
               
                if p_count>0:
                    numRepP+=1    
                if pests1.count()>0 or pests2.count()>0:
                    cns+=c.title+', '
                countP+=p_count
                cNewP+=pests1.count()
                cUpP+=pests2.count()
                    
            numCP_P.append(numRepP)
            numCP_P.append(countP)
            numCP_P.append(cNewP)
            numCP_P.append(cUpP)
            numCP_P.append(cns)
            region_cp_p.append(numCP_P)
        regionsPCP.append(region_cp_p)
        
        #-----------------ALL COUNTRIES---------
        regionsAll = []
        totALLcn=countriesperregioncp=CountryPage.objects.filter().count()
        region_all_p = []
        numRepPAll=0
        for k,v in REGIONS:
            reg = v+''
            numRepPAll=0
            countriesperregioncp=CountryPage.objects.filter(region=k)
            numCP_P = []
            numCP_P.append(reg)
            numCP_P.append(countriesperregioncp.count())
            countP=0
            cNewP=0
            cUpP=0
            cns=''
            p_count=0
            cNewP=0
            cUpP=0
            for c in countriesperregioncp:
                pests= PestReport.objects.filter(country=c.id,is_version=False,publish_date__gte=startstartdate,publish_date__lte=enddate)
                p_count=pests.count()
                pests1= PestReport.objects.filter(country=c.id,is_version=False,publish_date__gte=startdate,publish_date__lte=enddate)
                pests2= PestReport.objects.filter(country=c.id,is_version=False,modify_date__gte=startdate,modify_date__lte=enddate)
               
                if p_count>0:
                    numRepPAll+=1    
                if pests1.count()>0 or pests2.count()>0:
                    cns+=c.title+', '
                countP+=p_count
                cNewP+=pests1.count()
                cUpP+=pests2.count()
                       
                    
            numCP_P.append(numRepPAll)
            numCP_P.append(countP)
            numCP_P.append(cNewP)
            numCP_P.append(cUpP)
            numCP_P.append(cns)
            region_all_p.append(numCP_P)
        regionsAll.append(region_all_p)
        
        
        
        #----------------
        
        regionsPCPTot=[]
        totarray=[]
        tot=0
        tot2=0
        tot4=0
        tot5=0
        for x in  regionsPCP[0]:
            tot+= x[2]
            tot2+= x[3]
            tot4+= x[4]
            tot5+= x[5]
        totarray.append(tot)
        totarray.append(tot2)
        totarray.append(tot4)
        totarray.append(tot5)

        regionsPCPTot.append(totarray)     
        
        #----------------
        
        regionsALLTot=[]
        totarray1=[]
        tot=0
        tot2=0
        tot4=0
        tot5=0
        for x in  regionsAll[0]:
            tot+= x[2]
            tot2+= x[3]
            tot4+= x[4]
            tot5+= x[5]
        totarray1.append(tot)
        totarray1.append(tot2)
        totarray1.append(tot4)
        totarray1.append(tot5)

        regionsALLTot.append(totarray1)     

        context['curryear']=curryear
        context['prevyear']=prevyear
        context['regionsPCP']=regionsPCP
        context['totNumReg']=totNumReg
        context['regionsPCPTot']=regionsPCPTot
        
        context['regionsAll']=regionsAll
        context['totALLcn']=totALLcn
        context['regionsALLTot']=regionsALLTot
 
 
        pest_array=[]
        pest_array1=[]
     
        pestreporting_array = []
        pestreporting_array1 = []
        p_count=0
        for y in range(2005,curryear):
            pests=PestReport.objects.filter(is_version=False)
            p_count1=0
            for p in pests:
                if p.publish_date != None and p.publish_date.year == y:
                    p_count=p_count+1
                    p_count1=p_count1+1
            pestreporting_array.append(p_count)
            pestreporting_array1.append(p_count1)
            
        pest_array.append(pestreporting_array)       
        pest_array1.append(pestreporting_array1)       
        datachartbis=''
        datachart1bis=''
        datachart2=''
        datachart3=''
        i=0
     
        for y in range(2005,curryear ):
            datachartbis += '{type: "column", name: "'+str(y)+'", legendText: "'+str(y)+'",showInLegend: true, dataPoints:[{label: "Pest reports", y: '+str(pest_array[0][i])+'},]},'
            datachart1bis +='{type: "column", name: "'+str(y)+'", legendText: "'+str(y)+'",showInLegend: true, dataPoints:[{label: "Pest reports", y: '+str(pest_array1[0][i])+'},]},'
            
            datachart2+= '{label: "'+str(y)+'", y: '+str(pest_array[0][i])+'}, '
            datachart3+= '{label: "'+str(y)+'", y: '+str(pest_array1[0][i])+'}, '
            i=i+1
        context['pest_array']=pest_array
        context['pest_array1']=pest_array1
        context['num_years_range']=range(2005,curryear)
        context['datachartbis']=datachartbis
        context['datachart1bis']=datachart1bis
        context['num_years']=num_years
        context['datachart2']=datachart2
        context['datachart3']=datachart3
        
        return context   

class CountryStatsNROsDetailView(DetailView):
    """ CountryStatsNROsDetailView detail page """
    model = NROStats
    context_object_name = 'nrostats'
    template_name = 'countries/countries_stats_nros_new.html'
    queryset = NROStats.objects.filter()

    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryStatsNROsDetailView, self).get_context_data(**kwargs)
        nrostat = get_object_or_404(NROStats, id=self.kwargs['pk'])
        
        date=nrostat.date
        datetraining=nrostat.datetraining
        datetraining_checked=nrostat.datetraining_checked
        cns=nrostat.selcns
        
        context['datetraining_checked']=datetraining_checked
        context['datetraining']=datetraining
        context['date1']=date
        context['cns']=cns
        
        startdate = datetime(1972, 1, 1, 00, 01,00)
        enddate   = datetime(date.year, date.month, date.day+1, 23, 59,00)
        trainingdate=''
        if datetraining_checked:
            trainingdate   = datetime(datetraining.year, datetraining.month, datetraining.day+1, 23, 59,00)
            
        
        cnns=cns.split(",")
     
        country_array_pest_tot=[]
        country_array_rep_tot=[]
        country_array_ev_tot=[]
        for cn in cnns :
            if cn!='':
                country= get_object_or_404(CountryPage, page_ptr_id=cn )
                        
                array_cn_pest_tot=[]
                array_cn_rep_tot=[]
                array_cn_ev_tot=[]
                
                array_cn_pest_tot.append(country.title)#0
                array_cn_rep_tot.append(country.title)#0
                array_cn_ev_tot.append(country.title)#0
               
                #PEST
                if datetraining_checked:
                    pests2= PestReport.objects.filter(country=cn,is_version=False,publish_date__gte=startdate,publish_date__lte=trainingdate)
                    array_cn_pest_tot.append(pests2.count())
                else:
                    array_cn_pest_tot.append('-')
                pests1= PestReport.objects.filter(country=cn,is_version=False,publish_date__gte=startdate,publish_date__lte=enddate)
                array_cn_pest_tot.append(pests1.count())
                if datetraining_checked:
                    pests3= PestReport.objects.filter(country=cn,is_version=False,publish_date__gte=trainingdate,publish_date__lte=enddate)
                    array_cn_pest_tot.append(pests3.count())
                else:
                    array_cn_pest_tot.append('-')
                  
                if datetraining_checked:
                    pests4_count=0
                    pests4= PestReport.objects.filter(country=cn,is_version=False,modify_date__gte=trainingdate,modify_date__lte=enddate)
                    for p in pests4:
                        if p.modify_date != p.publish_date:
                            pests4_count=pests4_count+1
                    array_cn_pest_tot.append(pests4_count)
                else:
                    array_cn_pest_tot.append('-')
                 
                #REP       
                for i in range(1,5):
                    if datetraining_checked:
                        rep_2= ReportingObligation.objects.filter(country=cn,reporting_obligation_type=i,is_version=False,publish_date__gte=startdate,publish_date__lte=trainingdate)
                        array_cn_rep_tot.append(rep_2.count())
                    else:    
                        array_cn_rep_tot.append('-')
                    rep_1= ReportingObligation.objects.filter(country=cn,reporting_obligation_type=i,is_version=False,publish_date__gte=startdate,publish_date__lte=enddate)
                    array_cn_rep_tot.append(rep_1.count())
                    if datetraining_checked:
                        rep_3= ReportingObligation.objects.filter(country=cn,reporting_obligation_type=i,is_version=False,publish_date__gte=trainingdate,publish_date__lte=enddate)
                        array_cn_rep_tot.append(rep_3.count())
                    else:    
                        array_cn_rep_tot.append('-')
                    if datetraining_checked:
                        rep4_count=0
                        rep4= ReportingObligation.objects.filter(country=cn,reporting_obligation_type=i,is_version=False,publish_date__gte=trainingdate,publish_date__lte=trainingdate)
                        for r in rep4:
                            if r.modify_date != r.publish_date:
                                rep4_count=rep4_count+1
                        array_cn_rep_tot.append(rep4_count)        
                    else:    
                        array_cn_rep_tot.append('-')    
                #EV        
                for i in range(1,6):
                    if datetraining_checked:
                        ev_2= EventReporting.objects.filter(country=cn,event_rep_type=i,is_version=False,publish_date__gte=startdate,publish_date__lte=trainingdate)
                        array_cn_ev_tot.append(ev_2.count())
                    else:    
                        array_cn_ev_tot.append('-')
                    ev_1= EventReporting.objects.filter(country=cn,event_rep_type=i,is_version=False,publish_date__gte=startdate,publish_date__lte=enddate)
                    array_cn_ev_tot.append(ev_1.count())
                
                    if datetraining_checked:
                        ev_3= EventReporting.objects.filter(country=cn,event_rep_type=i,is_version=False,publish_date__gte=trainingdate,publish_date__lte=enddate)
                        array_cn_ev_tot.append(ev_3.count())
                    else:    
                        array_cn_ev_tot.append('-')
                    if datetraining_checked:
                        ev4_count=0
                        ev4= EventReporting.objects.filter(country=cn,event_rep_type=i,is_version=False,publish_date__gte=trainingdate,publish_date__lte=trainingdate)
                        for e in ev4:
                            if e.modify_date != e.publish_date:
                                ev4_count=ev4_count+1
                        array_cn_ev_tot.append(ev4_count)        
                    else:    
                        array_cn_ev_tot.append('-')    
                
                country_array_pest_tot.append(array_cn_pest_tot)
                country_array_rep_tot.append(array_cn_rep_tot)
                country_array_ev_tot.append(array_cn_ev_tot)
        
        context['country_array_pest_tot']=country_array_pest_tot
        context['country_array_rep_tot']=country_array_rep_tot
        context['country_array_ev_tot']=country_array_ev_tot
       
        return context

@login_required
@permission_required('ippc.add_nrostats', login_url="/accounts/login/")
def select_cns_nros_stats(request):
    """ Create nor stats """
    form = NROStatsForm(request.POST)
    countries=CountryPage.objects.all().order_by('title')
    print(PROJECT_ROOT)
    if request.method == "POST":
        if form.is_valid():
            new_nrostats = form.save(commit=False)
            new_nrostats.selcns=request.POST['selcns1']
  
            form.save()
           
            info(request, _("NROs Statistics created."))
            return redirect("nro-stats-detail",new_nrostats.id)
        else:
             return render_to_response('countries/countries_stats_nros_select.html', {'form': form,'countries':countries},
             context_instance=RequestContext(request))
    else:
        form = NROStatsForm(instance=NROStats())
     
    return render_to_response('countries/countries_stats_nros_select.html', {'form': form ,'countries':countries},
        context_instance=RequestContext(request))



   
  

class CountryStatsSingleReportsListView(ListView):
    """   stat  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_singlereport.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryStatsSingleReportsListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
        context['selyear_range']=range(2010,timezone.now().year+1)
       
        curryear=0
       #print('>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>')
        if 'year' in self.kwargs:
           #print(self.kwargs['year'])
            curryear=int(self.kwargs['year'])
        else :   
             curryear=timezone.now().year-1
        #curryear=2017#timezone.now().year-1
          
        #curryear=timezone.now().year-1
        
        regionsRepCP = []
        regionsEvCP = []
        regionsPCP = []
        totNumReg=countriesperregioncp=CountryPage.objects.filter(cp_ncp_t_type='CP').count()
        totRepREG=0   
        for i in range(1,6):
            region_cp_r = []
            region_cp_e = []
            numRepCN=0
            numEvsCN=0 
                
            for k,v in REGIONS:
                reg = v+''
                numRepCN=0
                numEvsCN=0 
                countriesperregioncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='CP')
                numCP_rep = []
                numCP_ev = []
                numCP_rep.append(reg)
                numCP_rep.append(countriesperregioncp.count())
                numCP_ev.append(reg)
                numCP_ev.append(countriesperregioncp.count())
                countRep=0
                countEv=0
                cNewR=0
                cUpR=0
                cNewE=0
                cUpE=0
                for c in countriesperregioncp:
                    reps_count=0
                    evs_count=0
                    reps=ReportingObligation.objects.filter(country=c.id,reporting_obligation_type=i,is_version=False)
                    #PAOLAPAOLA
                    for r in reps:
                        if r.publication_date != None and r.publication_date.year <= curryear:
                            reps_count= reps_count+1
                        if r.publication_date != None and r.publication_date.year == curryear:
                            cNewR+=1
                        if r.modify_date != None and r.modify_date.year == curryear:
                            cUpR+=1
                    evs=EventReporting.objects.filter(country=c.id,event_rep_type=i,is_version=False)
                    for e in evs:
                        if e.publication_date != None and e.publication_date.year <= curryear:
                            evs_count=evs_count+1
                        if e.publication_date != None and e.publication_date.year == curryear:
                             cNewE+=1
                        if e.modify_date != None and e.modify_date.year == curryear:
                            cUpE+=1
                    #countRep+=reps.count()
                    #countEv+=evs.count()
                    countRep+=reps_count
                    countEv+=evs_count
                    
                    if reps_count>0:#reps.count()>0:
                        numRepCN+=1
                    if evs_count>0:#evs.count()>0:
                        numEvsCN+=1
             
                            
                numCP_rep.append(numRepCN)
                numCP_rep.append(countRep)
                numCP_rep.append(cNewR)
                numCP_rep.append(cUpR)
                
                region_cp_r.append(numCP_rep)
                numCP_ev.append(numEvsCN)
                numCP_ev.append(countEv)
                numCP_ev.append(cNewE)
                numCP_ev.append(cUpE)
                region_cp_e.append(numCP_ev)
            regionsRepCP.append(region_cp_r)
            regionsEvCP.append(region_cp_e)
            
        region_cp_p = []
        numRepP=0
        for k,v in REGIONS:
            reg = v+''
            numRepP=0
            countriesperregioncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='CP')
            numCP_P = []
            numCP_P.append(reg)
            numCP_P.append(countriesperregioncp.count())
            countP=0
            cNewP=0
            cUpP=0
            for c in countriesperregioncp:
                pests= PestReport.objects.filter(country=c.id,is_version=False)
                p_count=0
                for p in pests:
                    if p.publish_date != None and p.publish_date.year <= curryear:
                        p_count=p_count+1
                    if p.publish_date != None and p.publish_date.year == curryear:
                         cNewP+=1
                    if p.modify_date != None and p.modify_date.year == curryear:
                        cUpP+=1
                countP+=p_count#pests.count()
                if p_count>0:#pests.count()>0:
                    numRepP+=1    
            numCP_P.append(numRepP)
            numCP_P.append(countP)
            numCP_P.append(cNewP)
            numCP_P.append(cUpP)
            region_cp_p.append(numCP_P)
        regionsPCP.append(region_cp_p)
        regionsRepCPTot=[]
       
        for i in range(0,4):
            totarray=[]
            tot=0
            tot2=0
            tot4=0
            tot5=0
            for x in  regionsRepCP[i]:
                tot+= x[2]
                tot2+= x[3]
                tot4+= x[4]
                tot5+= x[5]
            totarray.append(tot)
            totarray.append(tot2)
            totarray.append(tot4)
            totarray.append(tot5)
            
            regionsRepCPTot.append(totarray)  
        regionsEvCPTot=[]
        for i in range(0,5):
            totarray=[]
            tot=0
            tot2=0
            tot4=0
            tot5=0
            for x in  regionsEvCP[i]:
                tot+= x[2]
                tot2+= x[3]
                tot4+= x[4]
                tot5+= x[5]
            totarray.append(tot)
            totarray.append(tot2)
            totarray.append(tot4)
            totarray.append(tot5)
            
            regionsEvCPTot.append(totarray)  
        regionsPCPTot=[]
        totarray=[]
        tot=0
        tot2=0
        tot4=0
        tot5=0
        for x in  regionsPCP[0]:
            tot+= x[2]
            tot2+= x[3]
            tot4+= x[4]
            tot5+= x[5]
        totarray.append(tot)
        totarray.append(tot2)
        totarray.append(tot4)
        totarray.append(tot5)

        regionsPCPTot.append(totarray)      

        context['curryear']=curryear
        context['regionsRepCP']=regionsRepCP
        context['regionsEvCP']=regionsEvCP
        context['regionsPCP']=regionsPCP
        context['totNumReg']=totNumReg
        context['regionsRepCPTot']=regionsRepCPTot
        context['regionsEvCPTot']=regionsEvCPTot
        context['regionsPCPTot']=regionsPCPTot
        

 
        return context   


class CountryStatsTotalReportsIncreaseListView(ListView):
    """   stat  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_totalreport_increase.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryStatsTotalReportsIncreaseListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
        context['selyear_range']=range(2010,timezone.now().year+1)
       
        
        curryear=0
        prevyear=0
       #print('>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>')
        if 'year' in self.kwargs:
           #print(self.kwargs['year'])
            curryear=int(self.kwargs['year'])
            prevyear=curryear-1
        else :   
            curryear=timezone.now().year-1
            prevyear=timezone.now().year-2
        
        #curryear=timezone.now().year-1
        #prevyear=timezone.now().year-2
   

        rep_array=[]
        ev_array=[]
        pest_array=[]
        
             
        for i in range(1,6):
            reporting_array = []
            eventreporting_array = []
            for y in range(prevyear,curryear +1):
                rep_count=0
                evrep_count=0 
                reps=ReportingObligation.objects.filter(reporting_obligation_type=i,is_version=False)
                evrep=EventReporting.objects.filter(event_rep_type=i,is_version=False)
                for r in reps:
                    if r.publication_date != None and r.publication_date.year == y:
                        rep_count=rep_count+1
                for e in evrep:
                    if e.publication_date != None and e.publication_date.year == y:
                        evrep_count=evrep_count+1
                
            
                reporting_array.append(rep_count)
                eventreporting_array.append(evrep_count)
            ev_array.append(eventreporting_array)
            rep_array.append(reporting_array)
            
      
        
        pestreporting_array = []
        tot_p_count=0
        for y in range(prevyear,curryear +1):
            pests=PestReport.objects.filter(is_version=False)
            p_count=0
            for p in pests:
                if p.publish_date != None and p.publish_date.year == y:
                    p_count=p_count+1
                  
            pestreporting_array.append(p_count)
            tot_p_count+=p_count
        pest_array.append(pestreporting_array)       
        
        numIncrease=0
        increase=0
        rep_array1=[]
        ev_array1=[]
        pest_array1=[]
        tot1=0
        tot2=0
        for x in rep_array:
            increaseRep=[]
            a=float(x[0])
            b=float(x[1])
            tot1+=a
            tot2+=b
            numIncrease=b-a
            ok=False
            if x[0]!=0:
                increase=((b-a)/a)*100
                increase= round(increase, 2)
                ok=True
            else:
                ok=False
                
            nn=''
            nn1=''
          
                
            if numIncrease > 0:
                nn='increase by '+str(abs(numIncrease))+' report'
                if numIncrease>1:
                    nn+='s' 
                if ok:
                    nn1='+'+str(increase)+' %'
                else:
                    nn1='*'    
                
            elif numIncrease == 0:
                nn='N/A'
                nn1=' '
                   
                
            else:
                nn='<span style="color: red;">decrease by '+str(abs(numIncrease))+' report'
                if numIncrease<1:
                    nn+='s' 
                nn+='<span>'
                if ok:
                    nn1='<span style="color: red;">'+str(increase)+' %<span>'
                else:
                    nn1='*' 
              
                 
            increaseRep.append(nn)
            increaseRep.append(nn1)
            rep_array1.append(increaseRep)     
        numIncrease=0
        increase=0     
        for x in ev_array:
            increaseEv=[]
            
            a=float(x[0])
            b=float(x[1])
            tot1+=a
            tot2+=b
            numIncrease=b-a
            ok=False
            if x[0]!=0:
                increase=((b-a)/a)*100
                increase= round(increase, 2)
                ok=True
            else:
                ok=False
                
            nn=''
            nn1=''
          
                
            if numIncrease > 0:
                nn='increase by '+str(abs(numIncrease))+' report'
                if numIncrease>1:
                    nn+='s' 
                if ok:
                    nn1='+'+str(increase)+' %'
                else:
                    nn1='*'    
            elif numIncrease == 0:
                nn='N/A'
                nn1=' '    
            else:
                nn='<span style="color: red;">decrease by '+str(abs(numIncrease))+' report'
                if numIncrease<1:
                    nn+='s' 
                nn+='<span>'
                if ok:
                    nn1='<span style="color: red;">'+str(increase)+' %<span>'
                else:
                    nn1='*' 
              
                 
            increaseEv.append(nn)
            increaseEv.append(nn1)
            ev_array1.append(increaseEv)
        
        numIncrease=0
        increase=0     
        for x in pest_array:
            increaseP=[]
            a=float(x[0])
            b=float(x[1])
            tot1+=a
            tot2+=b
            numIncrease=b-a
            ok=False
            if x[0]!=0:
                increase=((b-a)/a)*100
                increase= round(increase, 2)
                ok=True
            else:
                ok=False
                
            nn=''
            nn1=''
          
                
            if numIncrease > 0:
                nn='increase by '+str(abs(numIncrease))+' report'
                if numIncrease>1:
                    nn+='s' 
                if ok:
                    nn1='+'+str(increase)+' %'
                else:
                    nn1='*'    
            elif numIncrease == 0:
                nn='N/A'
                nn1=' '    
            else:
                nn='<span style="color: red;">decrease by '+str(abs(numIncrease))+' report'
                if numIncrease<1:
                    nn+='s' 
                nn+='<span>'
                if ok:
                    nn1='<span style="color: red;">'+str(increase)+' %<span>'
                else:
                    nn1='*' 
              
            increaseP.append(nn)
            increaseP.append(nn1)
           
            pest_array1.append(increaseP)
       #print('tot1: '+str(tot1))                    
        context['curryear']=curryear
        context['prevyear']=prevyear
        context['tot1']=int(tot1)
        context['tot2']=int(tot2)
     
        context['rep_array']=rep_array
        context['ev_array']=ev_array
        context['pest_array']=pest_array
       #print(rep_array1)
        context['rep_array1']=rep_array1
        context['ev_array1']=ev_array1
        context['pest_array1']=pest_array1
      
        return context   
    
class CountryStatisticsTotalNroByYearListView(ListView):
    """   stat  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_totalnrobyyear.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryStatisticsTotalNroByYearListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
      
        context['selyear_range']=range(2010,timezone.now().year+1)
       
        
        #curryear=timezone.now().year   
        num_years=0


        curryear=0
       #print('>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>')
        if 'year' in self.kwargs:
           #print(self.kwargs['year'])
            curryear=int(self.kwargs['year'])
            num_years=curryear-2005
        else :   
            curryear=timezone.now().year #-1  
            num_years=curryear-2005

        rep_array=[]
        ev_array=[]
        pest_array=[]
        rep_array1=[]
        ev_array1=[]
        pest_array1=[]
       
        for i in range(1,5):
            reporting_array = []
            reporting_array1 = []
            rep_count=0
            for y in range(1970,curryear +1):
                reps=ReportingObligation.objects.filter(reporting_obligation_type=i,is_version=False)
                rep_count1=0
                for r in reps:##*****check date before 2005  and sum

                    if r.publication_date != None and r.publication_date.year == y:
                        rep_count=rep_count+1
                        rep_count1=rep_count1+1
                if y>=2005:
                    reporting_array.append(rep_count)
                    reporting_array1.append(rep_count1)
            rep_array.append(reporting_array)
            rep_array1.append(reporting_array1)
        
        for i in range(1,6):
            eventreporting_array = []
            eventreporting_array1 = []
            evrep_count=0
            for y in range(1970,curryear +1):
                evrep=EventReporting.objects.filter(event_rep_type=i,is_version=False)
                evrep_count1=0 
                for e in evrep:

                    if e.publication_date != None and e.publication_date.year == y:
                        evrep_count=evrep_count+1
                        evrep_count1=evrep_count1+1

                if y>=2005:
                    eventreporting_array.append(evrep_count)
                    eventreporting_array1.append(evrep_count1)
            ev_array.append(eventreporting_array)
            ev_array1.append(eventreporting_array1)
        
     
        pestreporting_array = []
        pestreporting_array1 = []
        p_count=0
        for y in range(2005,curryear +1):
            pests=PestReport.objects.filter(is_version=False)
            p_count1=0
            for p in pests:
                if p.publish_date != None and p.publish_date.year == y:
                    p_count=p_count+1
                    p_count1=p_count1+1
            pestreporting_array.append(p_count)
            pestreporting_array1.append(p_count1)
            
        pest_array.append(pestreporting_array)       
        pest_array1.append(pestreporting_array1)       
        datachart=''
        datachart1=''
        i=0
        totyearsarray=[]
        totyears1array=[]
        for y in range(2005,curryear +1):
            totyear=rep_array[0][i]+pest_array[0][i]+ev_array[0][i]+rep_array[2][i]+rep_array[1][i]+rep_array[3][i]+ev_array[1][i]+ev_array[2][i]+ev_array[3][i]+ev_array[4][i]
            totyear1=rep_array1[0][i]+pest_array1[0][i]+ev_array1[0][i]+rep_array1[2][i]+rep_array1[1][i]+rep_array1[3][i]+ev_array1[1][i]+ev_array1[2][i]+ev_array1[3][i]+ev_array1[4][i]
            datachart += '{type: "column", name: "'+str(y)+'", legendText: "'+str(y)+'",showInLegend: true, dataPoints:[{label: "Description of NPPO", y: '+str(rep_array[0][i])+'},	{label: "Pest reports", y: '+str(pest_array[0][i])+'},	{label: "Emergency action", y: '+str(ev_array[0][i])+'},{label: "List of regulated pests", y: '+str(rep_array[2][i])+'},{label: "Entry points", y:  '+str(rep_array[1][i])+'},{label: "Legislation: phytosanitary requirements/ restrictions/ prohibitions", y:  '+str(rep_array[3][i])+'},	{label: "Non-compliance", y:  '+str(ev_array[1][i])+'},{label: "Organizational arrangements of plant protection", y: '+str(ev_array[2][i])+'},{label: "Pest status", y:'+str(ev_array[3][i])+'},{label: "Rationale for phytosanitary requirements", y: '+str(ev_array[4][i])+'}]},'
            datachart1 += '{type: "column", name: "'+str(y)+'", legendText: "'+str(y)+'",showInLegend: true, dataPoints:[{label: "Description of NPPO", y: '+str(rep_array1[0][i])+'},	{label: "Pest reports", y: '+str(pest_array1[0][i])+'},	{label: "Emergency action", y: '+str(ev_array1[0][i])+'},{label: "List of regulated pests", y: '+str(rep_array1[2][i])+'},{label: "Entry points", y:  '+str(rep_array1[1][i])+'},{label: "Legislation: phytosanitary requirements/ restrictions/ prohibitions", y:  '+str(rep_array1[3][i])+'},	{label: "Non-compliance", y:  '+str(ev_array1[1][i])+'},{label: "Organizational arrangements of plant protection", y: '+str(ev_array1[2][i])+'},{label: "Pest status", y:'+str(ev_array1[3][i])+'},{label: "Rationale for phytosanitary requirements", y: '+str(ev_array1[4][i])+'}]},'
            i=i+1
            totyearsarray.append(totyear)
            totyears1array.append(totyear1)
      
        datachart2=''
        datachart3=''
        i=0
        for y in range(2005,curryear +1):
           datachart2+='{ y:'+ str(totyears1array[i])+', label: '+str(y)+'},'
           datachart3+='{ y:'+ str(totyearsarray[i])+', label: '+str(y)+'},'
           i=i+1
           
        context['totyearsarray']=totyearsarray
        context['totyears1array']=totyears1array
       
        context['curryear']=curryear
        context['num_years']=num_years
        context['num_years_range']=range(2005,curryear +1)
        context['rep_array']=rep_array
        context['ev_array']=ev_array
        context['pest_array']=pest_array
        context['rep_array1']=rep_array1
        context['ev_array1']=ev_array1
        context['pest_array1']=pest_array1
            
        context['datachart']=datachart
        context['datachart2']=datachart2
        context['datachart3']=datachart3
       
        
            
        context['datachart1']=datachart1
        return context   
    

class CountryRegionsUsersListView(ListView):
    """   Statistic users per regions  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_regionsusers.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryRegionsUsersListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
    
        regionCNcp = []
        regionCNncp = []
        regionall= []
        regionOffcp = []
        regionUnOffcp = []
        regionInfoncp = []
        regionLocalncp = []
        regionEditors = []
        
        
        tot_o_count=1       
        tot_u_count=1       
        tot_i_count=1       
        tot_l_count=1       
        tot_e_count=1       
        for k,v in REGIONS:
            reg = v.lower()
            numCNcp = []
            countriesperregioncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='CP')
            numb_countriesperregioncp=countriesperregioncp.count()
            numCNcp.append(reg)
            numCNcp.append(numb_countriesperregioncp)
            regionCNcp.append(numCNcp)
            context['region_cp']=regionCNcp

            numCNncp = []
            countriesperregionncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='NCP')
            numb_countriesperregionncp=countriesperregionncp.count()
            numCNncp.append(reg)
            numCNncp.append(numb_countriesperregionncp)
            regionCNncp.append(numCNncp)
            context['region_ncp']=regionCNncp
            
            numAll = []
            numAll.append(reg)
            numAll.append(numb_countriesperregioncp+numb_countriesperregionncp)
            regionall.append(numAll)
            context['regions']=regionall

            official = []
            unofficial = []
            infopoint = []
            local = []
            editors = []
            
            o_count=0
            u_count=0
            i_count=0
            l_count=0
            e_count=0
            #CP
            for c in countriesperregioncp:
                o_count+=IppcUserProfile.objects.filter(country=c.id,contact_type='1').count()
                u_count+=IppcUserProfile.objects.filter(country=c.id,contact_type='2').count()
                e_count+=IppcUserProfile.objects.filter(country=c.id,contact_type='5').count()
            official.append(o_count)
            regionOffcp.append(official)   
            context['region_off_cp']=regionOffcp
            unofficial.append(u_count)
            regionUnOffcp.append(unofficial)   
            context['region_unoff_cp']=regionUnOffcp
            tot_o_count+=o_count
            tot_e_count+=e_count
            tot_u_count+=u_count
            
            #NCP
            for c in countriesperregionncp:
                i_count+=IppcUserProfile.objects.filter(country=c.id,contact_type='3').count()
                l_count+=IppcUserProfile.objects.filter(country=c.id,contact_type='4').count()
                e_count+=IppcUserProfile.objects.filter(country=c.id,contact_type='5').count()
            infopoint.append(i_count)
            regionInfoncp.append(infopoint)   
            context['region_info_ncp']=regionInfoncp
            local.append(l_count)
            regionLocalncp.append(local)   
            context['region_local_ncp']=regionLocalncp
            editors.append(e_count)
            regionEditors.append(editors)   
            context['region_editors']=regionEditors
            tot_i_count+=i_count
            tot_l_count+=l_count
            tot_e_count+=e_count
            
        context['tot_o_count']=tot_o_count       
        context['tot_u_count']=tot_u_count       
        context['tot_i_count']=tot_i_count       
        context['tot_l_count']=tot_l_count       
        context['tot_e_count']=tot_e_count       
        
        datachart1=''
        datachart2=''
        datachart3=''
        datachart4=''
        datachart5=''
        for k,v in REGIONS:
           datachart1 += ' {  y: '+str(regionOffcp[k-1][0]*100/tot_o_count)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(regionOffcp[k-1][0]*100/tot_o_count)+'%" },'
           datachart2 += ' {  y: '+str(regionUnOffcp[k-1][0]*100/tot_u_count)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(regionUnOffcp[k-1][0]*100/tot_u_count)+'%" },'
           datachart3 += ' {  y: '+str(regionInfoncp[k-1][0]*100/tot_i_count)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(regionInfoncp[k-1][0]*100/tot_i_count)+'%" },'
           datachart4 += ' {  y: '+str(regionLocalncp[k-1][0]*100/tot_l_count)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(regionLocalncp[k-1][0]*100/tot_l_count)+'%" },'
           datachart5 += ' {  y: '+str(regionEditors[k-1][0]*100/tot_e_count)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(regionEditors[k-1][0]*100/tot_e_count)+'%" },'

        context['datachart1']=datachart1       
        context['datachart2']=datachart2       
        context['datachart3']=datachart3       
        context['datachart4']=datachart4       
        context['datachart5']=datachart5       
        return context

class CountryRegionsUsersNeverLoggedNewListView(ListView):
    """   Statistic users per regions that never logged in  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_regionsusersneverlogged_new.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryRegionsUsersNeverLoggedNewListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
    
        regionCNcp = []
        regionCNncp = []
        regionCNterr = []
        regionall= []
        
        regionOffcp = []
        regionInfoncp = []
        regionLocal = []
        
        regionEditorsInfo = []
        regionEditorsTerr = []
        
        #prev_year = datetime.now().year -1 
        context['selyear_range']=range(2010,timezone.now().year+1)
       
        prev_year=0
       #print('>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>')
        if 'year' in self.kwargs:
           #print(self.kwargs['year'])
            prev_year=int(self.kwargs['year'])
        else :   
             prev_year= datetime.now().year-1
             
        context['prev_year']=prev_year
        context['tot_num_CP']=CountryPage.objects.filter(cp_ncp_t_type='CP').count()
        context['tot_num_NCP']=CountryPage.objects.filter( cp_ncp_t_type='NCP').count()
        context['tot_num_T']=CountryPage.objects.filter( cp_ncp_t_type='T').count()
            
        tot_o_count=0   
        tot_o_2015_count=0
        tot_e_count=0
        tot_e_2015_count=0
        
        tot_i_count=0
        tot_i_2015_count=0
        tot_l_count=0
        tot_l_2015_count=0
        tot_encp_count=0
        tot_encp_2015_count=0
        tot_eterr_count=0
        tot_eterr_2015_count=0
        
        
        t=''   
        for k,v in REGIONS:
            
            reg = v.lower()
            t=t+reg+': '
            numCNcp = []
            countriesperregioncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='CP')
            numb_countriesperregioncp=countriesperregioncp.count()
            numCNcp.append(reg)
            numCNcp.append(numb_countriesperregioncp)
            regionCNcp.append(numCNcp)
            context['region_cp']=regionCNcp

            numCNncp = []
            countriesperregionncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='NCP')
            numb_countriesperregionncp=countriesperregionncp.count()
            infop_count1=0
            for c in countriesperregionncp:
                infop_count1=infop_count1+IppcUserProfile.objects.filter(country=c.id,contact_type='3').count()
      
            newnum=infop_count1
            
            numCNncp.append(reg)
            numCNncp.append(newnum)
            regionCNncp.append(numCNncp)
            context['region_ncp']=regionCNncp
            
            
            numCNterr = []
            countriesperregioterr=CountryPage.objects.filter(region=k,cp_ncp_t_type='T')
            numb_countriesperregioterr=countriesperregioterr.count()
            numCNterr.append(reg)
            numCNterr.append(numb_countriesperregioterr)
            regionCNterr.append(numCNterr)
            context['region_terr']=regionCNterr
            
            
            numAll = []
            numAll.append(reg)
            numAll.append(numb_countriesperregioncp+numb_countriesperregionncp)
            regionall.append(numAll)
            context['regions']=regionall

            official = []
            infopoint = []
            local = []
            editorsncp = []
            editorsterr = []
            
            o_count=0
            o_2015_count=0
            e_count=0
            e_2015_count=0
            editorCPcount=0
            editorCPcount_reg=0
            #CP
            for c in countriesperregioncp:
                offneverlogg=IppcUserProfile.objects.filter(country=c.id,contact_type='1')
                editorneverlogg=IppcUserProfile.objects.filter(country=c.id,contact_type='5')
                editorCPcount=editorneverlogg.count()
                editorCPcount_reg=editorCPcount_reg+     editorCPcount     
                for o in offneverlogg:
                    u= User.objects.get(id=o.user_id)
                    if u.last_login.year == 1970:
                        #print(u.last_login.year)
                        o_count=o_count+1
                    if u.last_login.year > 1970 and u.last_login.year!=prev_year:
                        #print(u.last_login.year)
                        o_2015_count=o_2015_count+1
                for o in editorneverlogg:
                    u= User.objects.get(id=o.user_id)
                    if u.last_login.year == 1970:
                        #print(u.last_login.year)
                        e_count=e_count+1
                    if u.last_login.year > 1970 and u.last_login.year!=prev_year:
                        #print(u.last_login.year)
                        e_2015_count=e_2015_count+1        
            official.append(o_count)
            official.append(o_2015_count)
            if editorCPcount_reg>0:
                official.append(e_count)
            else: 
                official.append('-')
            official.append(e_2015_count)
            regionOffcp.append(official)   
            context['region_off_cp']=regionOffcp
  
            tot_o_count+=o_count
            tot_o_2015_count+=o_2015_count
            tot_e_count+=e_count
            tot_e_2015_count+=e_2015_count
            
           
            i_count=0
            i_2015_count=0
            encp_count=0
            encp_2015_count=0
            #NCP
            editorInfocount_reg=0
            for c in countriesperregionncp:
                infop_neverlogg=IppcUserProfile.objects.filter(country=c.id,contact_type='3')
                
                editorneverlogg=IppcUserProfile.objects.filter(country=c.id,contact_type='5')
                editorInfocount_reg=editorInfocount_reg+editorneverlogg.count()
                for o in infop_neverlogg:
                    u= User.objects.get(id=o.user_id)
                    if u.last_login.year == 1970:
                        i_count=i_count+1
                    if u.last_login.year > 1970 and u.last_login.year!=prev_year:
                        i_2015_count=i_2015_count+1
               
                for o in editorneverlogg:
                    u= User.objects.get(id=o.user_id)
                    if u.last_login.year == 1970:
                        encp_count=encp_count+1
                    if u.last_login.year > 1970 and u.last_login.year!=prev_year:
                        encp_2015_count=encp_2015_count+1           
            infopoint.append(i_count)
            infopoint.append(i_2015_count)
            regionInfoncp.append(infopoint)   
            context['region_info_ncp']=regionInfoncp
            if editorInfocount_reg>0:
                 editorsncp.append(encp_count)
            else :
                  editorsncp.append('-')
            editorsncp.append(encp_2015_count)
            regionEditorsInfo.append(editorsncp)   
            context['region_editors']=regionEditorsInfo
            
            tot_i_count+=i_count
            tot_i_2015_count+=i_2015_count
            
            tot_encp_count+=encp_count
            tot_encp_2015_count+=encp_2015_count
            
            
            l_count=0
            l_2015_count=0
            eterr_count=0
            eterr_2015_count=0
            
          #TERR
            editorTerrcount_reg=0
            for c in countriesperregioterr:
                localp_neverlogg=IppcUserProfile.objects.filter(country=c.id,contact_type='4')
                editorneverlogg=IppcUserProfile.objects.filter(country=c.id,contact_type='5')
                editorTerrcount_reg=editorTerrcount_reg+editorneverlogg.count()
                for o in localp_neverlogg:
                    u= User.objects.get(id=o.user_id)
                    if u.last_login.year == 1970:
                        l_count=l_count+1
                    if u.last_login.year > 1970 and u.last_login.year!=prev_year:
                        l_2015_count=l_2015_count+1   
                for o in editorneverlogg:
                    u= User.objects.get(id=o.user_id)
                    if u.last_login.year == 1970:
                        t=t+ '1970 - '+ u.last_name+', '
                        #print(u.last_login.year)
                        eterr_count=eterr_count+1
                    if u.last_login.year > 1970 and u.last_login.year!=prev_year:
                        #print(u.last_login.year)
                        t=t+ '2015 - '+ u.last_name+', '
                        eterr_2015_count=eterr_2015_count+1           

            local.append(l_count)
            local.append(l_2015_count)
            regionLocal.append(local)   
            context['region_local_terr']=regionLocal
            if editorTerrcount_reg>0:
                editorsterr.append(eterr_count)
            else:
                editorsterr.append('-')
            editorsterr.append(eterr_2015_count)
            regionEditorsTerr.append(editorsterr)   
            context['region_terr_editors']=regionEditorsTerr
            
            tot_l_count+=l_count
            tot_l_2015_count+=l_2015_count
            
            tot_eterr_count+=eterr_count
            tot_eterr_2015_count+=eterr_2015_count
            
        infop_count=0
        t=''
        for k,v in REGIONS:
            countriesperregioNcp=CountryPage.objects.filter(region=k,cp_ncp_t_type='NCP')
        
            
            for c in countriesperregioNcp:
                    infop_count=infop_count+IppcUserProfile.objects.filter(country=c.id,contact_type='3').count()
                    
            
     
        context['tot_num_NCP']=CountryPage.objects.filter( cp_ncp_t_type='NCP')
        
        context['infop_count']=infop_count
           
        context['tot_o_count']=tot_o_count       
        context['tot_o_2015_count']=tot_o_2015_count       
        context['tot_e_count']=tot_e_count       
        context['tot_e_2015_count']=tot_e_2015_count       
        context['tot_i_count']=tot_i_count       
        context['tot_i_2015_count']=tot_i_2015_count       
        context['tot_encp_count']=tot_encp_count       
        context['tot_encp_2015_count']=tot_encp_2015_count       
        context['tot_l_count']=tot_l_count       
        context['tot_l_2015_count']=tot_l_2015_count       
        context['tot_eterr_count']=tot_eterr_count       
        context['tot_eterr_2015_count']=tot_eterr_2015_count       
        context['t']=t       
         
        
        datachartCPnever=''
        datachartCPnever += '{type: "stackedColumn",  color: "#ff6666",       name: "Never logged in", showInLegend: "true",  dataPoints:['
        datachartCPnever_year=''
        datachartCPnever_year += '{type: "stackedColumn",  color: "#ff6666",       name: "Never logged in '+str(prev_year)+'", showInLegend: "true",  dataPoints:['
        datachartEnever=''
        datachartEnever += '{type: "stackedColumn",  color: "#ff6666",       name: "Never logged in", showInLegend: "true",  dataPoints:['
        datachartEnever_year=''
        datachartEnever_year += '{type: "stackedColumn", color: "#ff6666",        name: "Never logged in '+str(prev_year)+'", showInLegend: "true",  dataPoints:['
       
        datachartNCPnever=''
        datachartNCPnever += '{type: "stackedColumn", color: "#ff6666",        name: "Never logged in", showInLegend: "true",  dataPoints:['
        datachartNCPnever_year=''
        datachartNCPnever_year += '{type: "stackedColumn",  color: "#ff6666",       name: "Never logged in '+str(prev_year)+'", showInLegend: "true",  dataPoints:['
        datachartNCPEnever=''
        datachartNCPEnever += '{type: "stackedColumn",  color: "#ff6666",       name: "Never logged in", showInLegend: "true",  dataPoints:['
        datachartNCPEnever_year=''
        datachartNCPEnever_year += '{type: "stackedColumn",  color: "#ff6666",       name: "Never logged in '+str(prev_year)+'", showInLegend: "true",  dataPoints:['
      
        datachartTnever=''
        datachartTnever += '{type: "stackedColumn",   color: "#ff6666",      name: "Never logged in", showInLegend: "true",  dataPoints:['
        datachartTnever_year=''
        datachartTnever_year += '{type: "stackedColumn",   color: "#ff6666",      name: "Never logged in '+str(prev_year)+'", showInLegend: "true",  dataPoints:['
        datachartTEnever=''
        datachartTEnever += '{type: "stackedColumn",  color: "#ff6666",       name: "Never logged in", showInLegend: "true",  dataPoints:['
        datachartTEnever_year=''
        datachartTEnever_year += '{type: "stackedColumn",   color: "#ff6666",      name: "Never logged in '+str(prev_year)+'", showInLegend: "true",  dataPoints:['
		
        datachartTnever_1 =''
        datachartTnever_year_1=''
        datachartTEnever_1 =''
        datachartTEnever_year_1 =''
	   
        datachartCPnever_1 =''
        datachartCPnever_year_1=''
        datachartEnever_1 =''
        datachartEnever_year_1 =''
        
        datachartNCPnever_1 =''
        datachartNCPnever_year_1 =''
        datachartNCPEnever_1 =''
        datachartNCPEnever_year_1=''
            
		
        for k,v in REGIONS:
            reg = str(ugettext((v)))#v.lower()
            numb_countriesperregio_cp=CountryPage.objects.filter(region=k,cp_ncp_t_type='CP').count()
       	    numb_countriesperregionncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='NCP').count()
            numb_countriesperregioterr=CountryPage.objects.filter(region=k,cp_ncp_t_type='T').count()
         
    	    cp=0
            cp_currYear=0
            cp_1=0
            cp_currYear_1=0
            eCPcount=0
            cpEditor=0
            cpEditor_currYear=0
            cpEditor_1=0
            cpEditor_currYear_1=0

            ncp=0
            ncp_currYear=0
            ncp_1=0
            ncp_currYear_1=0
            eNCPcount=0
            ncpEditor=0
            ncpEditor_currYear=0
            ncpEditor_1=0
            ncpEditor_currYear_1=0
			
            terr=0
            terr_currYear=0
            terr_1=0
            terr_currYear_1=0
            eTerrCount=0
            terrEditor=0
            terrEditor_currYear=0
            terrEditor_1=0
            terrEditor_currYear_1=0
		
           
			
            
           #-------------------------------CP----------------------------------------
            cp=regionOffcp[k-1][0]*100/numb_countriesperregio_cp
            cp_currYear=regionOffcp[k-1][1]*100/numb_countriesperregio_cp
            cp_1=(numb_countriesperregio_cp-regionOffcp[k-1][0])*100/numb_countriesperregio_cp
            cp_currYear_1=(numb_countriesperregio_cp-regionOffcp[k-1][1])*100/numb_countriesperregio_cp
            
            for c in CountryPage.objects.filter(region=k,cp_ncp_t_type='CP'):
                eCPcount=eCPcount+IppcUserProfile.objects.filter(country=c.id,contact_type='5').count()
            if editorCPcount>0:
                if regionOffcp[k-1][2]!='-':
                   #print(regionOffcp[k-1][2])
                    cpEditor=regionOffcp[k-1][2]*100/eCPcount
                else:
                   #print('???')
                   #print(regionOffcp[k-1][2])
                    cpEditor=0
                cpEditor_currYear=regionOffcp[k-1][3]*100/eCPcount
                if regionOffcp[k-1][2]!='-':
                    cpEditor_1=(eCPcount-regionOffcp[k-1][2])*100/eCPcount
                else:    
                    cpEditor_1=(eCPcount-0)*100/eCPcount
                cpEditor_currYear_1=(eCPcount-regionOffcp[k-1][3])*100/eCPcount	
           #-------------------------------NCP----------------------------------------
            for c in CountryPage.objects.filter(region=k,cp_ncp_t_type='NCP'):
                eNCPcount=eNCPcount+IppcUserProfile.objects.filter(country=c.id,contact_type='5').count()
               #print()
               #print(str(reg)+': '+str(eNCPcount))
            if numb_countriesperregionncp>0:
                ncp=regionInfoncp[k-1][0]*100/numb_countriesperregionncp
                ncp_currYear=regionInfoncp[k-1][1]*100/numb_countriesperregionncp
                ncp_1=(numb_countriesperregionncp-regionInfoncp[k-1][0])*100/numb_countriesperregionncp
                ncp_currYear_1=(numb_countriesperregionncp-regionInfoncp[k-1][1])*100/numb_countriesperregionncp
                if eNCPcount>0:
                    if regionEditorsInfo[k-1][0] !='-':
                       #print( regionEditorsInfo[k-1][0])
                        ncpEditor=regionEditorsInfo[k-1][0]*100/eNCPcount
                    else:
                        ncpEditor=0
                    ncpEditor_currYear=regionEditorsInfo[k-1][1]*100/eNCPcount
                    if regionEditorsInfo[k-1][0] !='-':
                        ncpEditor_1=(eNCPcount-0)*100/eNCPcount
                    else:
                        ncpEditor=0
                  
                    ncpEditor_currYear_1=(eNCPcount-regionEditorsInfo[k-1][1])*100/eNCPcount

			#-----------------------------TERR------------------------------------------  
            for c in CountryPage.objects.filter(region=k,cp_ncp_t_type='T'):
                eTerrCount=eTerrCount+IppcUserProfile.objects.filter(country=c.id,contact_type='5').count()
            if numb_countriesperregioterr>0:
                terr=regionLocal[k-1][0]*100/numb_countriesperregioterr
                terr_currYear=regionLocal[k-1][1]*100/numb_countriesperregioterr
                terr_1=(numb_countriesperregioterr-regionLocal[k-1][0])*100/numb_countriesperregioterr
                terr_currYear_1=(numb_countriesperregioterr-regionLocal[k-1][1])*100/numb_countriesperregioterr
                if eTerrCount>0:
                    if regionEditorsTerr[k-1][0] !='-':
                         terrEditor=regionEditorsTerr[k-1][0]*100/eTerrCount
                    else:
                        terrEditor=0
                    terrEditor_currYear=regionEditorsTerr[k-1][1]*100/eTerrCount		
                    if regionEditorsTerr[k-1][0] !='-':
                        terrEditor_1=(eTerrCount-regionEditorsTerr[k-1][0])*100/eTerrCount
                    else:
                        terrEditor_1=(eTerrCount-0)*100/eTerrCount
                    terrEditor_currYear_1=(eTerrCount-regionEditorsTerr[k-1][1])*100/eTerrCount
         #------------------------------------------------------------------------------------------ 
		
        
            datachartCPnever += '{  y: '+str(cp)+' , label: "'+str(reg)+'"},'
            datachartCPnever_year += '{  y: '+str(cp_currYear)+' , label: "'+str(reg)+'"},'
            datachartEnever += '{  y: '+str(cpEditor)+' , label: "'+str(reg)+'"},'
            datachartEnever_year += '{  y: '+str(cpEditor_currYear)+' , label: "'+str(reg)+'"},'
			
            datachartNCPnever += '{  y: '+str(ncp)+' , label: "'+str(reg)+'"},'
            datachartNCPnever_year += '{  y: '+str(ncp_currYear)+' , label: "'+str(reg)+'"},'
            datachartNCPEnever += '{  y: '+str(ncpEditor)+' , label: "'+str(reg)+'"},'
            datachartNCPEnever_year += '{  y: '+str(ncpEditor_currYear)+' , label: "'+str(reg)+'"},'
          
            datachartTnever += '{  y: '+str(terr)+' , label: "'+str(reg)+'"},'
            datachartTnever_year += '{  y: '+str(terr_currYear)+' , label: "'+str(reg)+'"},'
            datachartTEnever += '{  y: '+str(terrEditor)+' , label: "'+str(reg)+'"},'
            datachartTEnever_year += '{  y: '+str(terrEditor_currYear)+' , label: "'+str(reg)+'"},'
     # ------------------------------------------------------------------------------------------------------
		
                
               
            
              
            datachartCPnever_1+= '{  y: '+str(cp_1)+' , label: "'+str(reg)+'"},'
            datachartCPnever_year_1+= '{  y: '+str(cp_currYear_1)+' , label: "'+str(reg)+'"},'
            datachartEnever_1 += '{  y: '+str(cpEditor_1)+' , label: "'+str(reg)+'"},'
            datachartEnever_year_1 += '{  y: '+str(cpEditor_currYear_1)+' , label: "'+str(reg)+'"},'
			
            datachartNCPnever_1 += '{  y: '+str(ncp_1)+' , label: "'+str(reg)+'"},'
            datachartNCPnever_year_1 += '{  y: '+str(ncp_currYear_1)+' , label: "'+str(reg)+'"},'
            datachartNCPEnever_1 += '{  y: '+str(ncpEditor_1)+' , label: "'+str(reg)+'"},'
            datachartNCPEnever_year_1 += '{  y: '+str(ncpEditor_currYear_1)+' , label: "'+str(reg)+'"},'
			
            datachartTnever_1 += '{  y: '+str(terr_1)+' , label: "'+str(reg)+'"},'
            datachartTnever_year_1 += '{  y: '+str(terr_currYear_1)+' , label: "'+str(reg)+'"},'
            datachartTEnever_1 += '{  y: '+str(terrEditor_1)+' , label: "'+str(reg)+'"},'
            datachartTEnever_year_1 += '{  y: '+str(terrEditor_currYear_1)+' , label: "'+str(reg)+'"},'

			
	datachartCPnever += ' ]},'
        datachartCPnever_year += ' ]},'
        datachartEnever += ' ]},'
        datachartEnever_year += ' ]},'
		
	datachartNCPnever += ' ]},'
        datachartNCPnever_year += ' ]},'
        datachartNCPEnever += ' ]},'
        datachartNCPEnever_year += ' ]},'
		
	datachartTnever += ' ]},'
        datachartTnever_year += ' ]},'
        datachartTEnever += ' ]},'
        datachartTEnever_year += ' ]},'
		
		
	datachartTnever_1 += ' ]},'
        datachartTnever_year_1 += ' ]},'
        datachartTEnever_1 += ' ]},'
        datachartTEnever_year_1 += ' ]},'
	   
        datachartCPnever_1 += ' ]},'
        datachartCPnever_year_1 += ' ]},'
        datachartEnever_1 += ' ]},'
        datachartEnever_year_1 += ' ]},'
        
        datachartNCPnever_1 += ' ]},'
        datachartNCPnever_year_1 += ' ]},'
        datachartNCPEnever_1 += ' ]},'
        datachartNCPEnever_year_1 += ' ]},'
            
			
        datachartCPnever += '{type: "stackedColumn",   color:   "#00cc66" ,   name: "Logged in", showInLegend: "true",  dataPoints:['
        datachartCPnever_year += '{type: "stackedColumn",   color:   "#00cc66" ,   name: "Logged in in '+str(prev_year)+'", showInLegend: "true",  dataPoints:['
        datachartEnever += '{type: "stackedColumn",   color:   "#00cc66" ,   name: "Logged in", showInLegend: "true",  dataPoints:['
        datachartEnever_year += '{type: "stackedColumn",   color:   "#00cc66" ,   name: "Logged in in '+str(prev_year)+'", showInLegend: "true",  dataPoints:['
		
		
        datachartNCPnever += '{type: "stackedColumn",    color:   "#00cc66" ,  name: "Logged in", showInLegend: "true",  dataPoints:['
        datachartNCPnever_year += '{type: "stackedColumn",  color:   "#00cc66" ,    name: "Logged in in '+str(prev_year)+'", showInLegend: "true",  dataPoints:['
        datachartNCPEnever += '{type: "stackedColumn", color:   "#00cc66" ,     name: "Logged in", showInLegend: "true",  dataPoints:['
        datachartNCPEnever_year += '{type: "stackedColumn",   color:   "#00cc66" ,   name: "Logged in in '+str(prev_year)+'", showInLegend: "true",  dataPoints:['
        
            
       
        datachartTnever += '{type: "stackedColumn",  color:   "#00cc66" ,    name: "Logged in", showInLegend: "true",  dataPoints:['
        datachartTnever_year += '{type: "stackedColumn",  color:   "#00cc66" ,    name: "Logged in in '+str(prev_year)+'", showInLegend: "true",  dataPoints:['
        datachartTEnever += '{type: "stackedColumn",  color:   "#00cc66" ,    name: "Logged in", showInLegend: "true",  dataPoints:['
        datachartTEnever_year += '{type: "stackedColumn", color:   "#00cc66" ,     name: "Logged in in '+str(prev_year)+'", showInLegend: "true",  dataPoints:['
		
		
	datachartCPnever += datachartCPnever_1
        datachartCPnever_year += datachartCPnever_year_1
        datachartEnever += datachartEnever_1
        datachartEnever_year+= datachartEnever_year_1
		
	datachartNCPnever += datachartNCPnever_1
        datachartNCPnever_year+= datachartNCPnever_year_1
        datachartNCPEnever += datachartNCPEnever_1
        datachartNCPEnever_year += datachartNCPEnever_year_1
		
	datachartTnever += datachartTnever_1
        datachartTnever_year += datachartTnever_year_1
        datachartTEnever += datachartTEnever_1
        datachartTEnever_year += datachartTEnever_year_1
       
            
       
   	  
           
        context['datachartCPnever']=datachartCPnever      
        context['datachartCPnever_year']=datachartCPnever_year       
        context['datachartEnever']=datachartEnever       
        context['datachartEnever_year']=datachartEnever_year       

        context['datachartNCPnever']=datachartNCPnever
        context['datachartNCPnever_year']=datachartNCPnever_year       
        context['datachartNCPEnever']=datachartNCPEnever       
        context['datachartNCPEnever_year']=datachartNCPEnever_year  


        context['datachartTnever']=datachartTnever
        context['datachartTnever_year']=datachartTnever_year       
        context['datachartTEnever']=datachartTEnever       
        context['datachartTEnever_year']=datachartTEnever_year  


        return context
class CountryRegionsUsersNeverLoggedListView(ListView):

    """   Statistic users per regions that never logged in  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_regionsusersneverlogged.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryRegionsUsersNeverLoggedListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
    
        regionCNcp = []
        regionCNncp = []
        regionCNterr = []
        regionall= []
        
        regionOffcp = []
        regionInfoncp = []
        regionLocal = []
        
        regionEditorsInfo = []
        regionEditorsTerr = []
        
        prev_year = datetime.now().year -1 
        context['prev_year']=prev_year
        context['tot_num_CP']=CountryPage.objects.filter(cp_ncp_t_type='CP').count()
        context['tot_num_NCP']=CountryPage.objects.filter( cp_ncp_t_type='NCP').count()
        context['tot_num_T']=CountryPage.objects.filter( cp_ncp_t_type='T').count()
            
        tot_o_count=0   
        tot_o_2015_count=0
        tot_e_count=0
        tot_e_2015_count=0
        tot_i_count=0
        tot_i_2015_count=0
        tot_l_count=0
        tot_l_2015_count=0
        tot_encp_count=0
        tot_encp_2015_count=0
        tot_eterr_count=0
        tot_eterr_2015_count=0
            
        t=''   
        for k,v in REGIONS:
            reg = v.lower()
            t=t+reg+': '
            numCNcp = []
            countriesperregioncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='CP')
            numb_countriesperregioncp=countriesperregioncp.count()
            numCNcp.append(reg)
            numCNcp.append(numb_countriesperregioncp)
            regionCNcp.append(numCNcp)
            context['region_cp']=regionCNcp

            numCNncp = []
            countriesperregionncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='NCP')
            numb_countriesperregionncp=countriesperregionncp.count()
            infop_count1=0
            for c in countriesperregionncp:
                infop_count1=infop_count1+IppcUserProfile.objects.filter(country=c.id,contact_type='3').count()
      
            newnum=infop_count1
           #print("-------------------"+reg+"--------------------------")
           #print(numb_countriesperregionncp)
           #print(infop_count1)
           #print(newnum)
           
            numCNncp.append(reg)
            numCNncp.append(newnum)
            regionCNncp.append(numCNncp)
            context['region_ncp']=regionCNncp
            
            
            numCNterr = []
            countriesperregioterr=CountryPage.objects.filter(region=k,cp_ncp_t_type='T')
            numb_countriesperregioterr=countriesperregioterr.count()
            numCNterr.append(reg)
            numCNterr.append(numb_countriesperregioterr)
            regionCNterr.append(numCNterr)
            context['region_terr']=regionCNterr
            
            
            numAll = []
            numAll.append(reg)
            numAll.append(numb_countriesperregioncp+numb_countriesperregionncp)
            regionall.append(numAll)
            context['regions']=regionall

            official = []
            infopoint = []
            local = []
            editorsncp = []
            editorsterr = []
            
            o_count=0
            o_2015_count=0
            e_count=0
            e_2015_count=0
            #CP
            for c in countriesperregioncp:
                offneverlogg=IppcUserProfile.objects.filter(country=c.id,contact_type='1')
                editorneverlogg=IppcUserProfile.objects.filter(country=c.id,contact_type='5')
                editorCPcount=editorneverlogg.count()
              
                
                for o in offneverlogg:
                    u= User.objects.get(id=o.user_id)
                    if u.last_login.year == 1970:
                        #print(u.last_login.year)
                        o_count=o_count+1
                    if u.last_login.year > 1970 and u.last_login.year!=prev_year:
                       #print(u.last_login.year)
                        o_2015_count=o_2015_count+1
                for o in editorneverlogg:
                    u= User.objects.get(id=o.user_id)
                    if u.last_login.year == 1970:
                        #print(u.last_login.year)
                        e_count=e_count+1
                    if u.last_login.year > 1970 and u.last_login.year!=prev_year:
                       #print(u.last_login.year)
                        e_2015_count=e_2015_count+1        
            official.append(o_count)
            official.append(o_2015_count)
            official.append(e_count)
            official.append(e_2015_count)
            regionOffcp.append(official)   
            context['region_off_cp']=regionOffcp
            
        
  
            tot_o_count+=o_count
            tot_o_2015_count+=o_2015_count
            tot_e_count+=e_count
            tot_e_2015_count+=e_2015_count
            
           
            i_count=0
            i_2015_count=0
            encp_count=0
            encp_2015_count=0
            #NCP
            for c in countriesperregionncp:
                infop_neverlogg=IppcUserProfile.objects.filter(country=c.id,contact_type='3')
                
                editorneverlogg=IppcUserProfile.objects.filter(country=c.id,contact_type='5')
                for o in infop_neverlogg:
                    u= User.objects.get(id=o.user_id)
                    if u.last_login.year == 1970:
                        i_count=i_count+1
                    if u.last_login.year > 1970 and u.last_login.year!=prev_year:
                        i_2015_count=i_2015_count+1
                for o in editorneverlogg:
                    u= User.objects.get(id=o.user_id)
                    if u.last_login.year == 1970:
                        encp_count=encp_count+1
                    if u.last_login.year > 1970 and u.last_login.year!=prev_year:
                        encp_2015_count=encp_2015_count+1           
            infopoint.append(i_count)
            infopoint.append(i_2015_count)
            regionInfoncp.append(infopoint)   
            context['region_info_ncp']=regionInfoncp
            editorsncp.append(encp_count)
            editorsncp.append(encp_2015_count)
            regionEditorsInfo.append(editorsncp)   
            context['region_editors']=regionEditorsInfo
            
            tot_i_count+=i_count
            tot_i_2015_count+=i_2015_count
            
            tot_encp_count+=encp_count
            tot_encp_2015_count+=encp_2015_count
            
            
            l_count=0
            l_2015_count=0
            eterr_count=0
            eterr_2015_count=0
            
          #TERR
            for c in countriesperregioterr:
                localp_neverlogg=IppcUserProfile.objects.filter(country=c.id,contact_type='4')
                editorneverlogg=IppcUserProfile.objects.filter(country=c.id,contact_type='5')
              
                for o in localp_neverlogg:
                    u= User.objects.get(id=o.user_id)
                    if u.last_login.year == 1970:
                        l_count=l_count+1
                    if u.last_login.year > 1970 and u.last_login.year!=prev_year:
                        l_2015_count=l_2015_count+1   
                for o in editorneverlogg:
                    u= User.objects.get(id=o.user_id)
                    if u.last_login.year == 1970:
                        t=t+ '1970 - '+ u.last_name+', '
                        #print(u.last_login.year)
                        eterr_count=eterr_count+1
                    if u.last_login.year > 1970 and u.last_login.year!=prev_year:
                        #print(u.last_login.year)
                        t=t+ '2015 - '+ u.last_name+', '
                        eterr_2015_count=eterr_2015_count+1           
           
            local.append(l_count)
            local.append(l_2015_count)
            regionLocal.append(local)   
            context['region_local_terr']=regionLocal
            editorsterr.append(eterr_count)
            editorsterr.append(eterr_2015_count)
            regionEditorsTerr.append(editorsterr)   
            context['region_terr_editors']=regionEditorsTerr
            
            tot_l_count+=l_count
            tot_l_2015_count+=l_2015_count
            
            tot_eterr_count+=eterr_count
            tot_eterr_2015_count+=eterr_2015_count
            
        infop_count=0
        t=''
        for k,v in REGIONS:
            countriesperregioNcp=CountryPage.objects.filter(region=k,cp_ncp_t_type='NCP')
        
            
            for c in countriesperregioNcp:
                    infop_count=infop_count+IppcUserProfile.objects.filter(country=c.id,contact_type='3').count()
                    
            
     
        context['tot_num_NCP']=CountryPage.objects.filter( cp_ncp_t_type='NCP')
        
        context['infop_count']=infop_count
           
        context['tot_o_count']=tot_o_count       
        context['tot_o_2015_count']=tot_o_2015_count       
        context['tot_e_count']=tot_e_count       
        context['tot_e_2015_count']=tot_e_2015_count       
        context['tot_i_count']=tot_i_count       
        context['tot_i_2015_count']=tot_i_2015_count       
        context['tot_encp_count']=tot_encp_count       
        context['tot_encp_2015_count']=tot_encp_2015_count       
        context['tot_l_count']=tot_l_count       
        context['tot_l_2015_count']=tot_l_2015_count       
        context['tot_eterr_count']=tot_eterr_count       
        context['tot_eterr_2015_count']=tot_eterr_2015_count       
        context['t']=t       
         
        
        datachart1=''
        datachart2=''
        datachart3=''
        datachart4=''
        datachart_1=''
        datachart_2=''
        datachart_3=''
        datachart_4=''
        datachart_5=''
        datachart_6=''
        datachart_7=''
        datachart_8=''
       
        for k,v in REGIONS:
            reg = v.lower()
            t=t+''+reg+': '
            numb_countriesperregioncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='CP').count()
              
            percoff=regionOffcp[k-1][0]*100/numb_countriesperregioncp
            percoff2015=regionOffcp[k-1][1]*100/numb_countriesperregioncp
            editorCPcount=0
            percoffedit=0
            percoffedit2015=0
            for c in CountryPage.objects.filter(region=k,cp_ncp_t_type='CP'):
                editorsCP=IppcUserProfile.objects.filter(country=c.id,contact_type='5').count()
                editorCPcount=editorCPcount+editorsCP
            
            
            
            editorneverloggCount=editorneverlogg.count()
            if editorCPcount>0:
                percoffedit=regionOffcp[k-1][2]*100/editorCPcount
                percoffedit2015=regionOffcp[k-1][3]*100/editorCPcount
         
           
            
            
            datachart1 += ' {  y: '+str(percoff)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(percoff)+'%" },'
            datachart2 += ' {  y: '+str(percoff2015)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(percoff2015)+'%" },'
            datachart3 += ' {  y: '+str(percoffedit)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(percoffedit)+'%" },'
            datachart4 += ' {  y: '+str(percoffedit2015)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(percoffedit2015)+'%" },'
            percInfo=0
            percoInfo2015=0
            percoLocal=0
            percoLocal2015=0
            percoeditncp=0
            percoeditncp2015=0
            editorNCPcount=0
       
            
            for c in CountryPage.objects.filter(region=k,cp_ncp_t_type='NCP'):
                editorsNCP=IppcUserProfile.objects.filter(country=c.id,contact_type='5').count()
                editorNCPcount=editorNCPcount+editorsNCP
          
            numb_countriesperregionncp=CountryPage.objects.filter(region=k,cp_ncp_t_type='NCP').count()
            if numb_countriesperregionncp>0:
                percInfo=regionInfoncp[k-1][0]*100/numb_countriesperregionncp
                percoInfo2015=regionInfoncp[k-1][1]*100/numb_countriesperregionncp
                if editorNCPcount>0:
                    percoeditncp=regionEditorsInfo[k-1][0]*100/editorNCPcount
                    percoeditncp2015=regionEditorsInfo[k-1][1]*100/editorNCPcount
            
            percoeditterr=0
            percoeditterr2015=0   
            editorTcount=0
            
            for c in CountryPage.objects.filter(region=k,cp_ncp_t_type='T'):
                editorsT=IppcUserProfile.objects.filter(country=c.id,contact_type='5').count()
                editorTcount=editorTcount+editorsT
                
            numb_countriesperregioterr=CountryPage.objects.filter(region=k,cp_ncp_t_type='T').count()
            if numb_countriesperregioterr>0:
                percoLocal=regionLocal[k-1][0]*100/numb_countriesperregioterr
                percoLocal2015=regionLocal[k-1][1]*100/numb_countriesperregioterr
                if editorTcount>0:
                    percoeditterr=regionEditorsTerr[k-1][0]*100/editorTcount
                    percoeditterr2015=regionEditorsTerr[k-1][1]*100/editorTcount
               
         
            datachart_1 += ' {  y: '+str(percInfo)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(percInfo)+'%" },'
            datachart_2 += ' {  y: '+str(percoInfo2015)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(percoInfo2015)+'%" },'
            datachart_5 += ' {  y: '+str(percoeditncp)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(percoeditncp)+'%" },'
            datachart_6 += ' {  y: '+str(percoeditncp2015)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(percoeditncp2015)+'%" },'
            datachart_3 += ' {  y: '+str(percoLocal)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(percoLocal)+'%" },'
            datachart_4 += ' {  y: '+str(percoLocal2015)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(percoLocal2015)+'%" },'
            datachart_7 += ' {  y: '+str(percoeditterr)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(percoeditterr)+'%" },'
            datachart_8 += ' {  y: '+str(percoeditterr2015)+', legendText:"'+str(v.__unicode__())+'", label: "'+str(v.__unicode__())+': '+str(percoeditterr2015)+'%" },'
        
      
        
        context['datachart1']=datachart1       
        context['datachart2']=datachart2       
        context['datachart3']=datachart3       
        context['datachart4']=datachart4       
        context['datachart_1']=datachart_1       
        context['datachart_2']=datachart_2       
        context['datachart_3']=datachart_3       
        context['datachart_4']=datachart_4       
        context['datachart_5']=datachart_5       
        context['datachart_6']=datachart_6       
        context['datachart_7']=datachart_7       
        context['datachart_8']=datachart_8       
            
        return context


class CountryStatsChangeInCPsListView(ListView):
    """    Changes in CPs/Local/Infopoints  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_changes_cp.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryStatsChangeInCPsListView, self).get_context_data(**kwargs)
        
        context['dategenerate']=timezone.now()
        context['selyear_range']=range(2015,timezone.now().year+1)
       
        prevyear=0
       #print('>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>')
        if 'year' in self.kwargs:
           #print(self.kwargs['year'])
            prevyear=int(self.kwargs['year'])
        else :   
             prevyear=timezone.now().year-1
     
        months=[1,2,3,4,5,6,7,8,9,10,11,12]
        months_days=[31,28,31,30,31,30,31,31,30,31,30,31]
        
        i=0
        ocp_change=[]
        for m in months:
            mm=[]
            datex=datetime(prevyear,m, 01,00,01,00)
            datey=datetime(prevyear,m, months_days[i],23,59,00)
            ocp=OCPHistory.objects.filter(start_date__gte=datex,start_date__lte=datey).count()
            rppo=PartnersContactPointHistory.objects.filter(start_date__gte=datex,start_date__lte=datey).count()
            ocpEd=CnEditorsHistory.objects.filter(start_date__gte=datex,start_date__lte=datey).count()
            rppoEd=PartnersEditorHistory.objects.filter(start_date__gte=datex,start_date__lte=datey).count()
            
            mm.append(m)
            mm.append(ocp+rppo)
            mm.append(ocpEd+rppoEd)
            ocp_change.append(mm)
            i+=1
        datachart=''
        datachart1=''
        for  xx in ocp_change:
            datachart+=' { x: new Date('+str(prevyear)+', '+str(int(xx[0])-1)+', 1), y: '+str(xx[1])+' },'
            datachart1+=' { x: new Date('+str(prevyear)+', '+str(int(xx[0])-1)+', 1), y: '+str(xx[2])+' },'
  
       
        context['prevyear']=prevyear
        context['datachart1']=datachart1
        context['datachart']=datachart
        context['ocp_change']=ocp_change
        return context

class CountryTotalUsersListView(ListView):
    """    Statistic status of ippc contact points,editors,users  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_totalusers.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(CountryTotalUsersListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
        
        context['tot_o_count']=IppcUserProfile.objects.filter(contact_type='1').count()
        context['tot_u_count']=IppcUserProfile.objects.filter(contact_type='2').count()
        context['tot_i_count']=IppcUserProfile.objects.filter(contact_type='3').count()
        context['tot_l_count']=IppcUserProfile.objects.filter(contact_type='4').count()
        context['tot_e_count']=IppcUserProfile.objects.filter(contact_type='5').count()
        context['tot_users']=IppcUserProfile.objects.filter().count()
        context['tot_users']=IppcUserProfile.objects.filter().count()
        context['tot_users']=IppcUserProfile.objects.filter().count()
        timezone.now()
       
        curryear=timezone.now().year
        date1=datetime(curryear-2, 12, 31,23,59,00)
        date2=datetime(curryear-1, 12, 31,23,59,00)
        context['date1']=curryear-2
        context['date2']=curryear-1
      
        u_date1=IppcUserProfile.objects.filter(date_account_created__lte=date1).count()
        u_date2=IppcUserProfile.objects.filter(date_account_created__gte=date1,date_account_created__lte=date2).count()

        u_percentage=0
        if u_date1>0:
            u_percentage=u_date2*100/u_date1
            
       #print('--------------------------------')    
       #print(u_date1)    
       #print(u_date2)    
       #print(u_percentage)  
        context['u_date1']=u_date1
        context['u_date2']=u_date2
        context['u_percentage']=u_percentage
        
        new_content1=0#modify_date
        new_content1+=EventReporting.objects.filter(publish_date__lte=date1,is_version=False).count()
        new_content1+=ReportingObligation.objects.filter(publish_date__lte=date1,is_version=False).count()
        new_content1+=PestReport.objects.filter(publish_date__lte=date1,is_version=False).count()
        new_content1+=ImplementationISPM.objects.filter(publish_date__lte=date1,is_version=False).count()
        new_content1+=PestFreeArea.objects.filter(publish_date__lte=date1,is_version=False).count()
        new_content1+=Website.objects.filter(publish_date__lte=date1).count()
        new_content1+=CnPublication.objects.filter(publish_date__lte=date1).count()
        new_content1+=CountryNews.objects.filter(publish_date__lte=date1).count()
      
        new_content2=0#modify_date
        new_content2+=EventReporting.objects.filter(publish_date__lte=date2,publish_date__gte=date1,is_version=False).count()
        new_content2+=ReportingObligation.objects.filter(publish_date__lte=date2,publish_date__gte=date1,is_version=False).count()
        new_content2+=PestReport.objects.filter(publish_date__lte=date2,publish_date__gte=date1,is_version=False).count()
        new_content2+=ImplementationISPM.objects.filter(publish_date__lte=date2,publish_date__gte=date1,is_version=False).count()
        new_content2+=PestFreeArea.objects.filter(publish_date__lte=date2,publish_date__gte=date1,is_version=False).count()
        new_content2+=Website.objects.filter(publish_date__lte=date2,publish_date__gte=date1).count()
        new_content2+=CnPublication.objects.filter(publish_date__lte=date2,publish_date__gte=date1).count()
        new_content2+=CountryNews.objects.filter(publish_date__lte=date2,publish_date__gte=date1).count()
        
        new_content_percentage=0
        if new_content1>0:
            new_content_percentage=new_content2*100/new_content1
        context['new_content1']=new_content1
        context['new_content2']=new_content2
        context['new_content_percentage']=new_content_percentage
        
        up_content1=0#modify_date
        up_content1+=EventReporting.objects.filter(modify_date__lte=date1,is_version=False).count()
        up_content1+=ReportingObligation.objects.filter(modify_date__lte=date1,is_version=False).count()
        up_content1+=PestReport.objects.filter(modify_date__lte=date1,is_version=False).count()
        up_content1+=ImplementationISPM.objects.filter(modify_date__lte=date1,is_version=False).count()
        up_content1+=PestFreeArea.objects.filter(modify_date__lte=date1,is_version=False).count()
        up_content1+=Website.objects.filter(modify_date__lte=date1).count()
        up_content1+=CnPublication.objects.filter(modify_date__lte=date1).count()
        up_content1+=CountryNews.objects.filter(modify_date__lte=date1).count()
      
        up_content2=0#modify_date
        up_content2+=EventReporting.objects.filter(modify_date__lte=date2,modify_date__gte=date1,is_version=False).count()
        up_content2+=ReportingObligation.objects.filter(modify_date__lte=date2,modify_date__gte=date1,is_version=False).count()
        up_content2+=PestReport.objects.filter(modify_date__lte=date2,modify_date__gte=date1,is_version=False).count()
        up_content2+=ImplementationISPM.objects.filter(modify_date__lte=date2,modify_date__gte=date1,is_version=False).count()
        up_content2+=PestFreeArea.objects.filter(modify_date__lte=date2,modify_date__gte=date1,is_version=False).count()
        up_content2+=Website.objects.filter(modify_date__lte=date2,modify_date__gte=date1).count()
        up_content2+=CnPublication.objects.filter(modify_date__lte=date2,modify_date__gte=date1).count()
        up_content2+=CountryNews.objects.filter(modify_date__lte=date2,modify_date__gte=date1).count()
        
        up_content_percentage=0
        if up_content1>0:
            up_content_percentage=up_content2*100/up_content1
        context['up_content1']=up_content1
        context['up_content2']=up_content2
        context['up_content_percentage']=up_content_percentage
       
            
        return context
from news.models import NewsPost, NewsCategory

class NewsStatisticsByYearListView(ListView):
    """   stat  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'news/news_totalbyyear.html'
    queryset = CountryPage.objects.all().order_by('title')
   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(NewsStatisticsByYearListView, self).get_context_data(**kwargs)
        context['dategenerate']=timezone.now()
        context['selyear_range']=range(2003,timezone.now().year+1)
       
        num_years=0
        curryear=0
        if 'year' in self.kwargs:
            curryear=int(self.kwargs['year'])
            num_years=curryear-2003
        else :   
            curryear=timezone.now().year   
            num_years=curryear-2003
        y_array=[]    
        n_array=[]
        news_posts = NewsPost.objects.published()
        news_posts = news_posts.filter(categories=1)

        for y in range(2003,curryear +1):
    
            
            num=0
            for n in news_posts:
                if n.publish_date.year == y:
                 num =num+1
            n_array.append(num)
            y_array.append(y)
        lastY=[]  
    
        for y in range(1,13):
            num=0
            for n in news_posts:
                if n.publish_date.year == curryear:
                   #print(n.publish_date.month)
                    if n.publish_date.month ==  y:
                        num =num+1
                       #print('..................................................')
                        
            lastY.append(num)

        datachart=''
        datachart3=''
        i=0
        
        for y in range(2003,curryear +1):
           datachart+=' { x: new Date('+str(y)+',1,1), y: '+str(n_array[i])+' },'
           datachart3+='{ y:'+ str(n_array[i])+', label: '+str(y)+'},'
       
            
           i=i+1
     
        datachart4=''
        datachart5=''
        i=0
        arrayMOnths=['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec',]
        for y in range(1,13):
           datachart4+=' { x: new Date('+str(curryear)+','+str(y-1)+',1), y: '+str(lastY[i])+' },'
           datachart5+='{ y:'+ str(lastY[i])+', label: "'+str(arrayMOnths[y-1])+'"},'
       
            
           i=i+1   
        context['n_array']=n_array
        context['curryear']=curryear
        context['num_years']=num_years
        context['num_years_range']=range(2003,curryear +1)
        context['num_month_range']=arrayMOnths#range(1,13)
        context['y_array']=y_array
        context['lastY']=lastY
        context['datachart']=datachart
        context['datachart3']=datachart3
        context['datachart4']=datachart4
        context['datachart5']=datachart5
        
        return context   
    
class PollListView(ListView):
    template_name = 'polls/index.html'
    context_object_name = 'latest_poll_list'
    def get_queryset(self):
        """Return the last five published polls."""
        return Poll.objects.order_by('-pub_date').all


class PollDetailView(DetailView):
    model = Poll
    template_name = 'polls/detail.html'
    def get_queryset(self):
        """Return the last five published polls."""
        return Poll.objects.filter(pub_date__lte=timezone.now())
        

class PollResultsView(DetailView):
    model = Poll
    template_name = 'polls/results.html'
    
    def get_context_data(self, **kwargs):
        context = super(PollResultsView, self).get_context_data(**kwargs)
        pollid=self.kwargs['pk']
        votes=PollVotes.objects.filter(poll_id=pollid)
        context['votes']= votes
        return context

def send_pollnotification_message(id):
    """ send_pollnotification_message """
    #send notification to SC
    poll = get_object_or_404(Poll,  pk=id)
    emailto_all = ['']
    for g in Group.objects.filter(id=4):
        users = g.user_set.all()
        for u in users:
           user_obj=User.objects.get(username=u)
           emailto_all.append(str(user_obj.email))
    for g in Group.objects.filter(id=36):
        users = g.user_set.all()
        for u in users:
           user_obj=User.objects.get(username=u)
           emailto_all.append(str(user_obj.email))
           
    subject='IPPC POLL:  new poll: '+poll.question
    textmessage='<p>Dear IPPC user,<br><br>a new poll has been posted and it is open for your answer ( selecting YES or NO) and comments:<br>    <br>Poll: '+poll.question+'<br><br>'+poll.polltext+'<br><br>You can view it at the following url: https://www.ippc.int/poll/'+str(id)+'<br><br>International Plant Protection Convention team </p>'

    #message = mail.EmailMessage(subject,textmessage,'paola.sentinelli@fao.org',#from
    #    ['paola.sentinelli@fao.org',], ['paola.sentinelli@fao.org'])#emailto_all for PROD, in TEST all to paola#
    message = mail.EmailMessage(subject,textmessage,'ippc@fao.org',#from
        [emailto_all], ['paola.sentinelli@fao.org'])#emailto_all for PROD, in TEST all to paola#
    
    message.content_subtype = "html"
    #print('test-sending')#
    sent =message.send()
        
        


@login_required
@permission_required('ippc.add_poll', login_url="/accounts/login/")
def poll_create(request):
    """ Create Poll """
    user = request.user
    author = user

    form = PollForm(request.POST)
    if request.method == "POST":
         c_form = Poll_ChoiceFormSet(request.POST)
         if form.is_valid() and c_form.is_valid():
            new_poll = form.save(commit=False)
            form.save()
           
            c_form.instance = new_poll
            c_form.save()
            send_pollnotification_message(new_poll.id)
            
            info(request, _("Successfully created Poll."))
            return redirect("detail", pk=new_poll.id)
         else:
             return render_to_response('polls/poll_create.html', {'form': form,'c_form': c_form,},
             context_instance=RequestContext(request))
       
    else:
        form = PollForm( instance=Poll())
        c_form =Poll_ChoiceFormSet()
    return render_to_response('polls/poll_create.html', {'form': form,'c_form': c_form},
        context_instance=RequestContext(request))

@login_required
@permission_required('ippc.change_poll', login_url="/accounts/login/")
def poll_edit(request, id=None, template_name='polls/poll_edit.html'):
    """ Edit Poll """
    if id:
        poll = get_object_or_404(Poll,  pk=id)
    else:
        poll = Poll()
      
    if request.POST:

        form =PollForm(request.POST, instance=poll)
        c_form = Poll_ChoiceFormSet(request.POST,  instance=poll)
        if form.is_valid() and c_form.is_valid():
            form.save()
            c_form.instance = poll
            c_form.save()
            

            return redirect("detail", pk=id)
    else:
        form = PollForm(instance=poll)
        c_form = Poll_ChoiceFormSet(instance=poll)
        
    return render_to_response(template_name, {
        'form': form,'c_form':c_form, "poll": poll
    }, context_instance=RequestContext(request))
    
        
def vote_poll(request, poll_id):
    p = get_object_or_404(Poll, pk=poll_id)
    if PollVotes.objects.filter(poll_id=poll_id, user_id=request.user.id).exists():
        return render(request, 'polls/detail.html', {
        'poll': p,
        'error_message': "Sorry, but you have already voted."
        })
    try:
        selected_choice = p.poll_choice_set.get(pk=request.POST['choice'])
    except (KeyError, Poll_Choice.DoesNotExist):
        # Redisplay the poll voting form.
        return render(request, 'polls/detail.html', {
            'poll': p,
            'error_message': "You didn't select a choice.",
        })
    else:
        selected_choice.votes += 1
        selected_choice.save()
        v = PollVotes(user=request.user, poll=p,choice=selected_choice,comment=request.POST['comment'])
        v.save()
        return redirect("results", pk=p.id)

class EmailUtilityMessageListView(ListView):
    """    EmailUtilityMessage List view """
    context_object_name = 'latest'
    model = EmailUtilityMessage
    date_field = 'date'
    template_name = 'emailutility/emailutility_list.html'
    queryset = EmailUtilityMessage.objects.all().order_by('-date', 'subject')
   
       
class EmailUtilityMessageDetailView(DetailView):
    """ EmailUtilityMessage detail page """
    model = EmailUtilityMessage
    context_object_name = 'emailmessage'
    template_name = 'emailutility/emailutility_detail.html'
    queryset = EmailUtilityMessage.objects.filter()

def split(arr, size):
     arrs = []
     while len(arr) > size:
         pice = arr[:size]
         arrs.append(pice)
         arr   = arr[size:]
         
     arrs.append(arr)
        
     return arrs

@login_required
@permission_required('ippc.add_emailutilitymessage', login_url="/accounts/login/")
def email_send(request):
    """ Create email to send """
    form = EmailUtilityMessageForm(request.POST)
    g_set=[]
    for g in Group.objects.filter():
        users = g.user_set.all()
        users_all=[]
        users_all.append(str(g))
        users_all.append(str(g.id))
        for u in users:
           users_u=[]
           user_obj=User.objects.get(username=u)
           if user_obj.is_active:
            userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
            users_u.append((unicode(userippc.first_name)))
            users_u.append((unicode(userippc.last_name)))
            users_u.append((user_obj.email))
            users_all.append(users_u)
        g_set.append(users_all)
    
    users_all=[]
    cp_set0=[]
    users_all=[]
    cp=IppcUserProfile.objects.filter(contact_type=1)
    cpname = get_object_or_404(ContactType,id=1)
    for u in cp:
           users_u=[]
           user_obj=User.objects.get(id=u.user_id)
           cn = get_object_or_404(CountryPage,id=u.country_id)
           users_u.append(str(cn))
           users_u.append(' ('+(unicode(u.first_name))+' '+(unicode(u.last_name))+' - '+str(user_obj.email)+') ')
           users_u.append(str(user_obj.email))
           users_all.append(users_u)
           
    users_all2= split(users_all,30)   
    j=0
    users_all_2=[]
    for xx in users_all2:
        users_all_2=[]
       
        k=j+1
        users_all_2.append(str(cpname)+" - Group "+str(k))
        users_all_2.append(str(j))
        j=j+1
        for x in xx:
            users_all_2.append(x)
    
        cp_set0.append(users_all_2)
        
    
   
    cp_set=[]      
    for h in range(2,5):
        users_all=[]
        cp=IppcUserProfile.objects.filter(contact_type=h)
        cpname = get_object_or_404(ContactType,id=h)
        users_all.append(str(cpname))
        users_all.append(str(h))
        cp=IppcUserProfile.objects.filter(contact_type=h)
        for u in cp:
               users_u=[]
               user_obj=User.objects.get(id=u.user_id)
               cn = get_object_or_404(CountryPage,id=u.country_id)
               users_u.append(str(cn))
               users_u.append(' ('+(unicode(u.first_name))+' '+(unicode(u.last_name))+') ')
               users_u.append(str(user_obj.email))
               users_all.append(users_u)

        cp_set.append(users_all)
   
    
    emaile2=[]
    users_all_e=[]
    for g in Group.objects.filter():
       
        if g.name == 'Country editor':
            users = g.user_set.all()
           
            for u in users:
                users_u=[]
                #user_obj=User.objects.get(id=u.user_id)
                
                user_obj=User.objects.get(username=u)
                if user_obj.is_active:
                    #print('$%$%$%$%$%$%$%$%$%$')
                    #print(user_obj.id)
                    userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                    #print(userippc)
                    
                
                
                    #cn = get_object_or_404(CountryPage,id=userippc.country_id)
                    #print(cn)
                    #users_u.append(str(cn))
                    users_u.append(' ('+(unicode(userippc.first_name))+' '+(unicode(userippc.last_name))+') ')
                    users_u.append(str(user_obj.email))
                    users_all_e.append(users_u)
    users_all_e2= split(users_all_e,30)   
    j=0
    users_all_e_2=[]
    for xx in users_all_e2:
        users_all_e_2=[]
        k=j+1
        users_all_e_2.append("Country editors - Group "+str(k))
        users_all_e_2.append(str(j))
        j=j+1
        for x in xx:
            users_all_e_2.append(x)
    
        emaile2.append(users_all_e_2)
            
   
           
    if request.method == "POST":
        f_form =EmailUtilityMessageFileFormSet(request.POST, request.FILES)
        if form.is_valid() and f_form.is_valid():
            emailto_all1 = str(request.POST['emailto'])
            emailto_all2=emailto_all1[3:-4]
            emailto_all=[emailto_all2]
            for u in request.POST.getlist('users'):
                user_obj=User.objects.get(id=u)
                user_email=user_obj.email
                emailto_all.append(str(user_email))
            for g in Group.objects.filter():
                for uemail in request.POST.getlist('user_'+str(g.id)+'_0'):
                    emailto_all.append(str(uemail))
                    
            for h in range(2,5):
                  for uemail in request.POST.getlist('usercp_'+str(h)+'_0'):
                     emailto_all.append(str(uemail))
            for h in range(0,6):
                  for uemail in request.POST.getlist('usercp1_'+str(h)+'_0'):
                     emailto_all.append(str(uemail)) 
            for h in range(0,7):
                  for uemail in request.POST.getlist('usere1_'+str(h)+'_0'):
                     emailto_all.append(str(uemail))                        
            #print(emailto_all)
            new_emailmessage = form.save(commit=False)
            new_emailmessage.date=timezone.now()
            
            new_emailmessage.emailto=emailto_all
            form.save()
            #save file to message in db
            f_form.instance = new_emailmessage
            f_form.save()
            #EmailMessage('Hello', 'Body goes here', 'from@example.com', ['to1@example.com', 'to2@example.com'], ['bcc@example.com'],  headers = {'Reply-To': 'another@example.com'})
            #send email message
            #message = mail.EmailMessage(request.POST['subject'],request.POST['messagebody'],request.POST['emailfrom'],
            #['paola.sentinelli@fao.org',], ['paola.sentinelli@fao.org'])#emailto_all for PROD, in TEST all to paola#
            emailto_all_split=[]
            #print('================================')
            #print(emailto_all)
            #print('================================')
            #if len(emailto_all) >30 :
            emailto_all_split = split(emailto_all,30)
            sent =0
            for emails_arr in emailto_all_split:
                messages=[]
                for emails_a in emails_arr:
                    emails_to=emails_a
                    
                    message = mail.EmailMessage(request.POST['subject'],request.POST['messagebody'],request.POST['emailfrom'],[emails_to], ['paola.sentinelli@fao.org'])#emailto_all for PROD, in TEST all to paola#
                    #print('===*******SENDING**********===')
                    #print (emails_arr)
                    #print('====******************************===')
                    # Attach a files to message
                    fileset= EmailUtilityMessageFile.objects.filter(emailmessage_id=new_emailmessage.id)
                    for f in fileset:
                        pf=MEDIA_ROOT+str(f.file)
                        message.attach_file(pf) 
                    message.content_subtype = "html"
                    messages.append(message)
                    
                    
                    #timeout, so changed to send_messages
                  
                # Manually open the connection
                #sends a list of EmailMessage objects. If the connection is not open, this call will implicitly open the connection, and close the connection afterwards. If the connection is already open, it will be left open after mail has been sent.
                connection = mail.get_connection()
                connection.open()
                #print('test-sending')
                sent=connection.send_messages(messages)#
                connection.close()
               
            #update status SENT/NOT SENT mail message in db
            new_emailmessage.sent=sent
          
            
            form.save()
           
            info(request, _("Email  sent."))
            return redirect("email-detail",new_emailmessage.id)
        else:
             return render_to_response('emailutility/emailutility_send.html', {'form': form,'f_form': f_form,'emailgroups':g_set,'emailcp':cp_set,'emailcp2':cp_set0,'emaile2':emaile2},
             context_instance=RequestContext(request))
    else:
        form = EmailUtilityMessageForm(instance=EmailUtilityMessage())
        f_form =EmailUtilityMessageFileFormSet()
      
    return render_to_response('emailutility/emailutility_send.html', {'form': form  ,'f_form': f_form,'emailgroups':g_set,'emailcp':cp_set,'emailcp2':cp_set0,'emaile2':emaile2},#'emailcpu':cpu_set,'emailcpi':cpi_set,'emailcpl':cpl_set
        context_instance=RequestContext(request))


class MassEmailUtilityMessageListView(ListView):
    """    MassEmailUtilityMessage List view """
    context_object_name = 'latest'
    model = MassEmailUtilityMessage
    date_field = 'date'
    template_name = 'emailutility/massemailutility_list.html'
    queryset = MassEmailUtilityMessage.objects.all().order_by('-date', 'subject')
   
       
class MassEmailUtilityMessageDetailView(DetailView):
    """ MassEmailUtilityMessage detail page """
    model = MassEmailUtilityMessage
    context_object_name = 'massemailmessage'
    template_name = 'emailutility/massemailutility_detail.html'
    queryset = MassEmailUtilityMessage.objects.filter()

    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(MassEmailUtilityMessageDetailView, self).get_context_data(**kwargs)
        mail = get_object_or_404(MassEmailUtilityMessage, id=self.kwargs['pk'])
        
        aaa=''
        emailto=[]
        sentto=[]
        not_sentto=[]
        if mail.massmerge ==0 or mail.massmerge ==2:
            if mail.emailto != None and mail.emailto!='':
                emailto=mail.emailto.split(",") 
            if mail.sentto != None and mail.sentto!='':
                sentto=mail.sentto.split(",") 
            if mail.not_sentto != None and mail.not_sentto!='':
                not_sentto=mail.not_sentto.split(",") 
            for m in emailto:
                if m in not_sentto:
                    aaa+='<span style="color: red;">'+m+'</span>,  '
                if m in  sentto:   
                    aaa+='<span style="color: green;">'+m+'</span>,  '
        elif mail.massmerge ==1: 
            if mail.emailtoISO3 != None and mail.emailtoISO3!='':
                emailto=mail.emailtoISO3.split(",") 
            if mail.senttoISO3 != None and mail.senttoISO3!='':
                sentto=mail.senttoISO3.split(",") 
            if mail.not_senttoISO3 != None and mail.not_senttoISO3!='':
                not_sentto=mail.not_senttoISO3.split(",") 
            for m in emailto:
                if m in not_sentto:
                    aaa+='<span style="color: red;">'+m+'</span>,  '
                if m in  sentto:   
                    aaa+='<span style="color: green;">'+m+'</span>,  '            
        context['aaa']=aaa
       #print(aaa)
        
        return context
 
@login_required
@permission_required('ippc.add_massemailutilitymessage', login_url="/accounts/login/")
def massemailutility_setstatus(request, pk=None , status=None):
    email=get_object_or_404(MassEmailUtilityMessage,id=pk)
    if email:
       #print("++++++++++++++++++++++++++++++++++++++++++")
       #print(status)
        if status == "1":
           #print("!!!!!!! 1")
            email.status=1
            email.save()
     
            info(request, _("Successfully set mass email message to 'TO BE SENT'."))
            return redirect("mass-email-list")
        elif status == "0":
           #print("!!!!!!! 0")
            email.status=0
            email.save()
            info(request, _("Successfully set mass email message to 'DRAFT'."))
            return redirect("mass-email-list")
        
    else: 
        warning(request, _("Error setting the status of mass email message."))
        return redirect("mass-email-list")
   
def massemailutility_to_send(request):
    
    emailutility_to_send_done_dir = MEDIA_ROOT+'/massemailutility_sent'
    log_report =  open(os.path.join(emailutility_to_send_done_dir, "massemailutility_sent_"+timezone.now().strftime('%Y%m%d%H%M%S')+".log"), 'wb')
    log_report.write("List of sent emails:\n\n")
    
    text_=''
    emails_to_send = MassEmailUtilityMessage.objects.filter(sent=0,status=1)
    for email in emails_to_send :
        log_report.write("["+ timezone.now().strftime('%Y%m%d%H%M%S')+" - email id  [ID:"+str(email.id)+"] \n")
        text_+="["+ timezone.now().strftime('%Y%m%d%H%M%S')+" - email id  [ID:"+str(email.id)+"] \n"
        if email.massmerge == 0 or email.massmerge == 2:
            email_not_sentto=email.not_sentto
            email_sent=email.sentto
            emailfrom=email.emailfrom
            subject=email.subject
            messagebody=email.messagebody
            email_not_sentto1=email_not_sentto.split(",") 
            email_sent1=''  
            eee=''
            if email_sent!= None and email_sent!='':
                 aaa=email_sent
                 email_sent1=email_sent 
            for y in range(0,len(email_not_sentto1)):
                if y > 5:
                    eee+=email_not_sentto1[y]+','
                else:
                   message = mail.EmailMessage(subject,messagebody,emailfrom,[email_not_sentto1[y]], ['paola.sentinelli@fao.org'])#/**/
                   fileset= MassEmailUtilityMessageFile.objects.filter(emailmessage_id=email.id)
                   for f in fileset:
                       pf=MEDIA_ROOT+str(f.file)
                       message.attach_file(pf) 
                   message.content_subtype = "html"
                   sent =message.send()

                   email_sent1+=email_not_sentto1[y]+','
                   text_+='sent to: '+email_not_sentto1[y]+'<br>'
                   log_report.write("sent to: "+email_not_sentto1[y]+":\n")
                    #COmmentato 24-3-2020
                    #if sent:
                    #   email_sent1+=email_not_sentto1[y]+','
                    #   text_+='sent to: '+email_not_sentto1[y]+'<br>'
                    #   log_report.write("sent to: "+email_not_sentto1[y]+":\n")
                    #else:
                    #   eee+=email_not_sentto1[y]+','

            email.not_sentto = eee
            email.sentto = email_sent1
            text_+='Sent email:'+subject+'<br>'

            if email_not_sentto1 ==  [u''] or email_not_sentto1 == '' :
                email.sent = True
                messages=[]
                bodytext='Please note that the mass email message has been sent out to all recipients: https://www.ippc.int/massemailutility/'+str(email.id)
                message = mail.EmailMessage('Notification mass email sent: '+ugettext(subject),bodytext,emailfrom,[email.author.email], ['paola.sentinelli@fao.org'])
                message.content_subtype = "html"
                messages.append(message) 
                connection = mail.get_connection()
                connection.open()
                #print('test-sending')
                sent=connection.send_messages(messages)#
                connection.close()
            email.save()
            
        elif email.massmerge ==1:    
            subject=email.subject
            emailfrom=email.emailfrom
            emailcc=email.emailcc
            email_cc=emailcc.split(",")
            email_not_sentto=email.not_senttoISO3
            email_sent=email.senttoISO3
            if email_sent == None:
                finalemailsent=''
            
            else:
                finalemailsent=email_sent
            
            #CSV
            csv_file_name=str(email.csv_file).split("/")[2]
            csv_path = os.path.join(MEDIA_ROOT, 'files')
            csv_path1 = os.path.join(csv_path, 'email')
            csv_path_final = os.path.join(csv_path1, csv_file_name)

            csv_dictionary = {}
            with open(csv_path_final, "r") as ins:
                for line in ins:
                    line_0 = line.split(",")
                    #text_+='line:'+str(line_0)
                    cn_0=line_0[1]
                    #text_+='cn_0:'+str(cn_0)
                    if len(cn_0)>3:
                       cn_0= cn_0[0:3]
                    csv_dictionary[cn_0]=line_0[0]
            
            email_not_sentto1=email_not_sentto.split(",") ##AFG,ALB,....
            email_sent1=''  
            eee=''
            if email_sent!= None and email_sent!='':
                 #aaa=email_sent.split(",") 
                 email_sent1=email_sent.split(",")  

            for y in range(0,len(email_not_sentto1)):
                if y > 5:
                    eee+=email_not_sentto1[y]+','
                else:
                   email_to_cn=email_not_sentto1[y]#AFG
                   
                   if email_to_cn !='':
                        cn = get_object_or_404(CountryPage,iso3=email_to_cn)
                        cn_cpid=cn.contact_point_id
                        user_obj=User.objects.get(id=cn_cpid)
                        cp=get_object_or_404(IppcUserProfile,user_id=cn_cpid)
                        username=get_gender(cp.gender)+' '+(unicode(cp.first_name))+' '+(unicode(cp.last_name))
                        cp_email=user_obj.email
                        cp_alternateemail=cp.email_address_alt
                        cplink_to_replace=''
                        cplink_to_replace=csv_dictionary.get(email_to_cn, "none")
                        text_+='<br>cp_email<br>:'+cp_email+'<br>cp_alternateemail:'+cp_alternateemail
                        #xcxc='<br>cp_email<br>:'+cp_email+'<br>cp_alternateemail:'+cp_alternateemail
                        singlelink='<a href="'+str(cplink_to_replace)+'">'+str(cplink_to_replace)+'</a>'
                        #MSG
                        messagebody=email.messagebody
                        message_split=messagebody.split("XXXXXXXXXX")
                        message_part_0=message_split[0]
                        message_split2=message_split[1]
                        message_split2=message_split2.replace("YYYYYYYYYY",singlelink)
                        #messagebodyfinal=message_part_0+' '+username+',<br>'+message_part_1+ '<a href="'+str(cplink_to_replace)+'">'+str(cplink_to_replace)+'</a><br>'+message_part_2
                        messagebodyfinal=message_part_0+' '+username+',<br>'+message_split2
                        text_+=messagebodyfinal
                        emailtTo=[]
                        #emailtTo.append('paola.sentinelli@fao.org')
                        #emailtTo.append('s.hess@kpnplanet.nl')
                        emailtTo.append(cp_email)
                        if cp_alternateemail!='':
                           emailtTo.append(cp_alternateemail)
                        
                        message = mail.EmailMessage(subject,messagebodyfinal,emailfrom,emailtTo, email_cc)
                        fileset= MassEmailUtilityMessageFile.objects.filter(emailmessage_id=email.id)
                        for f in fileset:
                            pf=MEDIA_ROOT+str(f.file)
                            message.attach_file(pf) 
                        message.content_subtype = "html"
                        sent =message.send()
                    
                        if sent:
                            email_sent1+=email_not_sentto1[y]+','
                            finalemailsent+=email_not_sentto1[y]+','
                            text_+='sent to: '+email_not_sentto1[y]+'<br>'
                            log_report.write("sent to: "+email_not_sentto1[y]+":\n")
                        #else:
                            # eee+=email_not_sentto1[y]+','

                email.not_senttoISO3 = eee
                email.senttoISO3 = finalemailsent
                text_+='Sent email:'+subject+'<br>'
           
          
          
            
            if email_not_sentto1 ==  [u''] or email_not_sentto1 == '' :
                
                email.sent = True
                messages=[]
                bodytext='Please note that the Merge email messages have been sent out to all recipients: https://www.ippc.int/massemailutility/'+str(email.id)
                text_+=bodytext
                message = mail.EmailMessage('Notification Merge email messages sent: '+ugettext(subject),bodytext,emailfrom,[email.author.email], ['paola.sentinelli@fao.org'])
                message.content_subtype = "html"
                messages.append(message) 
                connection = mail.get_connection()
                connection.open()

                sent=connection.send_messages(messages)#
                connection.close()
            email.save()
                        
        
    log_report.close()     
    context = {'aa':text_}
    response = render(request, "emailutility/sendemail_system.html", context)
   
    return response


@login_required
@permission_required('ippc.add_massemailutilitymessage', login_url="/accounts/login/")
def massemail_send(request):
    """ Create  mass email to send """
    form = MassEmailUtilityMessageForm(request.POST)
    g_set=[]
    for g in Group.objects.filter():
        users = g.user_set.all()
        users_all=[]
        users_all.append(str(g))
        users_all.append(str(g.id))
        for u in users:
           users_u=[]
           user_obj=User.objects.get(username=u)
           if user_obj.is_active:
            userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
            users_u.append((unicode(userippc.first_name)))
            users_u.append((unicode(userippc.last_name)))
            users_u.append((user_obj.email))
            users_all.append(users_u)
        g_set.append(users_all)
    
    users_all=[]
    cp_set0=[]
    users_all=[]
    cp=IppcUserProfile.objects.filter(contact_type=1)
    cpname = get_object_or_404(ContactType,id=1)
    for u in cp:
           users_u=[]
           user_obj=User.objects.get(id=u.user_id)
           cn = get_object_or_404(CountryPage,id=u.country_id)
           users_u.append(str(cn))
           users_u.append(' ('+(unicode(u.first_name))+' '+(unicode(u.last_name))+' - '+str(user_obj.email)+') ')
           users_u.append(str(user_obj.email))
           users_all.append(users_u)
           
    j=0
    users_all_2=[]
    k=j+1
    users_all_2.append(str(cpname))
    users_all_2.append(str(j))
    j=j+1
    for x in users_all:
         users_all_2.append(x)
    cp_set0.append(users_all_2)    
   
    cp_set=[]      
    for h in range(2,5):
        users_all=[]
        users_all_iso=[]
        cp=IppcUserProfile.objects.filter(contact_type=h)
        cpname = get_object_or_404(ContactType,id=h)
        users_all.append(str(cpname))
        users_all.append(str(h))
        cp=IppcUserProfile.objects.filter(contact_type=h)
        for u in cp:
               users_u=[]
               user_obj=User.objects.get(id=u.user_id)
               cn = get_object_or_404(CountryPage,id=u.country_id)
               users_u.append(str(cn))
               users_u.append(' ('+(unicode(u.first_name))+' '+(unicode(u.last_name))+') ')
               users_u.append(str(user_obj.email))
               users_all.append(users_u)
        cp_set.append(users_all)
    
    emaile2=[]
    users_all_e=[]
    for g in Group.objects.filter():
        if g.name == 'Country editor':
            users = g.user_set.all()
            for u in users:
                users_u=[]
                user_obj=User.objects.get(username=u)
                if user_obj.is_active:
                    userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                    users_u.append(' ('+(unicode(userippc.first_name))+' '+(unicode(userippc.last_name))+') ')
                    users_u.append(str(user_obj.email))
                    users_all_e.append(users_u)
    users_all_e2= split(users_all_e,30)   
    j=0
    users_all_e_2=[]
    k=j+1
    users_all_e_2.append("Country editors")
    users_all_e_2.append(str(j))
    j=j+1
    for x in users_all_e:
        users_all_e_2.append(x)
    emaile2.append(users_all_e_2)
            
    if request.method == "POST":
        f_form =MassEmailUtilityMessageFileFormSet(request.POST, request.FILES)
        if form.is_valid() and f_form.is_valid():
            
            emailto_all_last=''
            emailto_cc_last=''
            emailto_to_iso=''
            emailto_notsentto_iso=''
            emailto_sent_iso=''
          
            emailto_all1 = str(request.POST['emailto'])
            emailto_all=[emailto_all1]
            for u in request.POST.getlist('users'):
                user_obj=User.objects.get(id=u)
                user_email=user_obj.email
                emailto_all.append(str(user_email))
            for g in Group.objects.filter():
                for uemail in request.POST.getlist('user_'+str(g.id)+'_0'):
                    emailto_all.append(str(uemail))
            for h in range(2,5):
                  for uemail in request.POST.getlist('usercp_'+str(h)+'_0'):
                     emailto_all.append(str(uemail))
            for h in range(0,6):
                  for uemail in request.POST.getlist('usercp1_'+str(h)+'_0'):
                     emailto_all.append(str(uemail)) 
            for h in range(0,7):
                  for uemail in request.POST.getlist('usere1_'+str(h)+'_0'):
                     emailto_all.append(str(uemail))                        

            emailto_all_clean=[]
            for ee in emailto_all:
                if ee in emailto_all_clean:
                    print('no')
                else:    
                    emailto_all_clean.append(ee)    

            for ee in emailto_all_clean:
                    emailto_all_last=emailto_all_last+ee+','
                
           
            new_emailmessage = form.save(commit=False)
            new_emailmessage.author = request.user
            new_emailmessage.date=timezone.now()
            new_emailmessage.sent=0
            new_emailmessage.status=0
            
            new_emailmessage.emailto=emailto_all_last
            new_emailmessage.not_sentto=emailto_all_last
            
            new_emailmessage.emailcc=emailto_cc_last
            new_emailmessage.emailtoISO3=emailto_to_iso
            new_emailmessage.not_senttoISO3=  emailto_notsentto_iso

            form.save()
            #save file to message in db
            f_form.instance = new_emailmessage
            f_form.save()
            
            
            messages=[]
            message = mail.EmailMessage('COPY MASS email: '+ugettext(request.POST['subject']),request.POST['messagebody'],request.POST['emailfrom'],[request.user.email], ['paola.sentinelli@fao.org'])
            fileset= MassEmailUtilityMessageFile.objects.filter(emailmessage_id=new_emailmessage.id)
            for f in fileset:
                pf=MEDIA_ROOT+str(f.file)
                message.attach_file(pf) 
            message.content_subtype = "html"
            messages.append(message) 
            connection = mail.get_connection()
            connection.open()
            #print('test-sending')
            sent=connection.send_messages(messages)#
            connection.close()
            form.save()
            
            info(request, _("MASS Email stored as draft, check the copy message in your inbox, change the status in 'TO BE SENT' and it would be sent out in the next hours."))
           
            return redirect("mass-email-detail",new_emailmessage.id)
        else:
             return render_to_response('emailutility/massemailutility_send.html', {'form': form,'f_form': f_form,'emailgroups':g_set,'emailcp':cp_set,'emailcp2':cp_set0,'emaile2':emaile2,},
             context_instance=RequestContext(request))
    else:
        form = MassEmailUtilityMessageForm(instance=MassEmailUtilityMessage())
        f_form =MassEmailUtilityMessageFileFormSet()
      
    return render_to_response('emailutility/massemailutility_send.html', {'form': form  ,'f_form': f_form,'emailgroups':g_set,'emailcp':cp_set,'emailcp2':cp_set0,'emaile2':emaile2,},#'emailcpu':cpu_set,'emailcpi':cpi_set,'emailcpl':cpl_set
        context_instance=RequestContext(request))


@login_required
@permission_required('ippc.add_massemailutilitymessage', login_url="/accounts/login/")
def mergemassemail_send(request):
    """ Create  mass email to send """
    form = MassEmailUtilityMessageForm(request.POST, request.FILES)
    cp_setiso=[] 
    for h in range(2,5):
        users_all_iso=[]
        cp=IppcUserProfile.objects.filter(contact_type=h)
        cpname = get_object_or_404(ContactType,id=h)
        users_all_iso.append(str(cpname))
        users_all_iso.append(str(h))
        cp=IppcUserProfile.objects.filter(contact_type=h)
        for u in cp:
               users_u_iso=[]
               user_obj=User.objects.get(id=u.user_id)
               cn = get_object_or_404(CountryPage,id=u.country_id)
               users_u_iso.append(str(cn.iso3))
               users_u_iso.append(' - '+(unicode(cn)+' ['+(unicode(u.first_name))+' '+(unicode(u.last_name))+' - '+str(user_obj.email)+'] '))
               users_all_iso.append(users_u_iso)
        cp_setiso.append(users_all_iso)
    ##ISO CP
    cpname = get_object_or_404(ContactType,id=1)
    cp_set0iso=[]
    j=0
    users_alliso=[]
    users_alliso.append(str(cpname))
    users_alliso.append(str(j))
    countriesList=CountryPage.objects.filter(cp_ncp_t_type='CP').exclude(id='-1').order_by('name')
    for cn in countriesList :
	users_u=[]
        cn_iso3=cn.iso3
        cn_cpid=cn.contact_point_id
        user_obj=User.objects.get(id=cn_cpid)
        cp=get_object_or_404(IppcUserProfile,user_id=cn_cpid)
        users_u.append(str(cn_iso3))
        users_u.append(' - '+(unicode(cn.name)+' ['+(unicode(cp.first_name))+' '+(unicode(cp.last_name))+' - '+str(user_obj.email)+']'))
        users_alliso.append(users_u)     
    cp_set0iso.append(users_alliso)    
    
    if request.method == "POST":
        f_form =MassEmailUtilityMessageFileFormSet(request.POST, request.FILES)
        if form.is_valid() and f_form.is_valid():
            
            emailto_all_last=''
            emailto_cc_last=''
            emailto_to_iso=''
            emailto_notsentto_iso=''
            emailto_sent_iso=''
            

            #CC
            emailto_cc_clean=[]
            email_cc = str(request.POST['emailcc']).split(',');
            for e_cc in email_cc:
                if e_cc !='':
                    if e_cc in emailto_cc_clean:
                        print('no')
                    else:    
                        emailto_cc_clean.append(e_cc)    
            emailto_cc_clean.append('paola.sentinelli@fao.org')
            for e_cc1 in emailto_cc_clean:
                emailto_cc_last=emailto_cc_last+e_cc1+','

            #TO
            emailto_all=[]
            for h in range(2,5):
                  for uemail in request.POST.getlist('usercp_'+str(h)+'_0'):
                     emailto_all.append(str(uemail))
            for h in range(0,6):
                  for uemail in request.POST.getlist('usercp1_'+str(h)+'_0'):
                     emailto_all.append(str(uemail)) 
            emailto_all_last=''
            emailto_all_clean=[]
            for ee in emailto_all:
                if ee!='':
                    if ee in emailto_all_clean:
                        print('no')
                    else:    
                        emailto_all_clean.append(ee)    
            for ee in emailto_all_clean:
                emailto_all_last=emailto_all_last+ee+','

            emailto_to_iso =emailto_all_last
            emailto_notsentto_iso =emailto_all_last
            emailto_all_last=''
            new_emailmessage = form.save(commit=False)
            new_emailmessage.author = request.user
            new_emailmessage.date=timezone.now()
            new_emailmessage.sent=0
            new_emailmessage.status=0
            
            new_emailmessage.emailto=emailto_all_last
            new_emailmessage.not_sentto=emailto_all_last
            
            new_emailmessage.emailcc=emailto_cc_last
            new_emailmessage.emailtoISO3=emailto_to_iso
            new_emailmessage.not_senttoISO3=  emailto_notsentto_iso

            form.save()
            #save file to message in db
            f_form.instance = new_emailmessage
            f_form.save()
            
            
            messages=[]
            message = mail.EmailMessage('COPY email: '+ugettext(request.POST['subject']),request.POST['messagebody'],request.POST['emailfrom'],[request.user.email], ['paola.sentinelli@fao.org'])
            fileset= MassEmailUtilityMessageFile.objects.filter(emailmessage_id=new_emailmessage.id)
            for f in fileset:
                pf=MEDIA_ROOT+str(f.file)
                message.attach_file(pf) 
            message.content_subtype = "html"
            messages.append(message) 
            connection = mail.get_connection()
            connection.open()
            #print('test-sending')
            sent=connection.send_messages(messages)#
            connection.close()
            form.save()
            
            info(request, _("MERGE Email stored as draft, check the copy message in your inbox, change the status in 'TO BE SENT' and it would be sent out in the next hours."))
           
            return redirect("mass-email-detail",new_emailmessage.id)
        else:
             return render_to_response('emailutility/mergemassemailutility_send.html', {'form': form,'f_form': f_form,'emailcpiso':cp_set0iso,'emailcp2iso':cp_setiso,},
             context_instance=RequestContext(request))
    else:
        form = MassEmailUtilityMessageForm(instance=MassEmailUtilityMessage())
        f_form =MassEmailUtilityMessageFileFormSet()
      
    return render_to_response('emailutility/mergemassemailutility_send.html', {'form': form  ,'f_form': f_form,'emailcpiso':cp_set0iso,'emailcp2iso':cp_setiso,},#'emailcpu':cpu_set,'emailcpi':cpi_set,'emailcpl':cpl_set
        context_instance=RequestContext(request))

            
#
class ContactUsEmailMessageListView(ListView):
    """    ContactUsEmailMessageListView List view """
    context_object_name = 'latest'
    model = ContactUsEmailMessage
    date_field = 'date'
    template_name = 'emailcontact_us/emailcontact_us_list.html'
    queryset = ContactUsEmailMessage.objects.all().order_by('-date', 'subject')
   
       
class ContactUsEmailMessageDetailView(DetailView):
    """ ContactUsEmailMessage detail page """
    model = ContactUsEmailMessage
    context_object_name = 'emailmessage'
    template_name = 'emailcontact_us/emailcontact_us_detail.html'
    queryset = ContactUsEmailMessage.objects.filter()


import random

#@login_required
#@permission_required('ippc.add_contactusemailmessage', login_url="/accounts/login/")
def contactus_email_send(request):
    """ Create contactus email to send """
    form = ContactUsEmailMessageForm(request.POST)
    emailfrom =''
   #print(request.user)
    if request.user.is_anonymous():
      emailfrom=''
    else:
      emailfrom=request.user.email
    if request.method == "POST" :
        if form.is_valid() and request.POST['captcha'] ==  request.POST['result_element']  and request.POST['captcha2'] ==  request.POST['result2_element'] :
             emails_a=''
             subj1='Contact IPPC: '
             if request.POST['contact_us_type'] == "1":
                emails_a='ippc@fao.org'
                subj1='Contact IPPC: General enquiries - '
             elif request.POST['contact_us_type']== "2":
                emails_a='brent.larson@fao.org'
                subj1='Contact IPPC: Implementation / Capacity Development - '
             elif request.POST['contact_us_type']== "3":
                 emails_a='qingpo.yang@fao.org'
                 subj1='Contact IPPC: Registration of ISPM 15 symbol - '
             elif request.POST['contact_us_type']== "4":
                 emails_a='paola.sentinelli@fao.org'
                 subj1='Contact IPPC: National Reporting Obligations (NROs) - '
             elif request.POST['contact_us_type']== "5":
                 emails_a='Craig.Fedchock@fao.org'
                 subj1='Contact IPPC: News / Communications - '
             elif request.POST['contact_us_type']== "6":
                 emails_a='Craig.Fedchock@fao.org'
                 subj1='Contact IPPC: ePhyto - '
             elif request.POST['contact_us_type']== "7":
                 emails_a='IPPC-OCS@fao.org'
                 subj1='Contact IPPC: Online Comment System (OCS) -'
             elif request.POST['contact_us_type']== "8":
                 emails_a='Craig.Fedchock@fao.org'
                 subj1='Contact IPPC: Resource Mobilization -'
             elif request.POST['contact_us_type']== "9":
                 emails_a='avetik.nersisyan@fao.org'
                 subj1='Contact IPPC: Standard Setting - '
             elif request.POST['contact_us_type']== "10":
                 emails_a='IPPC-IT@fao.org'
                 subj1='Contact IPPC: Technical assistance (IT/bugs) - '
             elif request.POST['contact_us_type']== "11":
                 subj1='Contact IPPC: International Year of Plant Health (IYPH) -'
                 emails_a='IPPC-IYPH@fao.org'
     
     
             new_contactusemailmessage = form.save(commit=False)
             new_contactusemailmessage.date=timezone.now()
             form.save()
             sent = 0
             messages=[]
             #print(request.POST['contact_us_type'])
             subject=subj1+': '+ugettext(request.POST['subject'])
             message= mail.EmailMessage(subject,request.POST['messagebody'],request.POST['emailfrom'],[emails_a], ['paola.sentinelli@fao.org'])#emailto_all for PROD, in TEST all to paola#
             message.content_subtype = "html"
             messages.append(message)
             if "loans bad" in request.POST['messagebody']:
                new_contactusemailmessage.sent=False
             else:
                connection = mail.get_connection()
                connection.open()
                #print('test-sending')
                sent=connection.send_messages(messages)#
                connection.close()
                new_contactusemailmessage.sent=sent
                form.save()
#
             #update status SENT/NOT SENT mail message in db
             
            
             info(request, _("Email sent."))
             return redirect("contactus-email-detail",new_contactusemailmessage.id)
        else:
             error_captcha=''
             if not(request.POST['captcha'] == request.POST['result_element'] ) and not(request.POST['captcha2'] == request.POST['result2_element'] ) :
                   error_captcha='error'
                  
             return render_to_response('emailcontact_us/emailcontact_us_send.html', {'form': form,'x_element': request.POST['x_element'],'y_element': request.POST['y_element'],'result_element': request.POST['result_element'] ,'z_element':request.POST['z_element'],'t_element':request.POST['t_element'],'result2_element':request.POST['result2_element'],'error_captcha':error_captcha},
             context_instance=RequestContext(request))
    else:
         x_element=random.randint(1,50)   
         y_element=random.randint(1,50)
         z_element=random.randint(1,15)   
         t_element=random.randint(1,15)
        
         result_element=x_element+y_element
         result2_element=z_element*t_element
     
         form = ContactUsEmailMessageForm(instance=ContactUsEmailMessage(emailfrom=emailfrom))
#
       
    return render_to_response('emailcontact_us/emailcontact_us_send.html', {'form': form  ,'x_element':x_element,'y_element':y_element,'result_element':result_element,'z_element':z_element,'t_element':t_element,'result2_element':result2_element},
         context_instance=RequestContext(request))
#
#   
class AdvancesSearchCNListView(ListView):
    """  AdvancesSearchCNListView list  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_advsearchresults.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(AdvancesSearchCNListView, self).get_context_data(**kwargs)
        if self.kwargs['type'] == 'pestreport':
            context['type_label'] = 'Official pest report (Art. VIII.1a)'
            context['link_to_item'] = 'pest-report-detail'
            context['items']= PestReport.objects.filter(is_version=False,status=CONTENT_STATUS_PUBLISHED)#PestReport.objects.all()
            context['counttotal'] =context['items'].count() 
           
            cns= CountryPage.objects.all()
            maparray=[]
            maparray1=""
            tot_p=0
            for cn in cns:
              pests=PestReport.objects.filter(country_id=cn.id,is_version=False)
              p=pests.count()
              tot_p+=p
              if p>0:
                  if cn>0:
                 
                    maparray.append([str('<a href="/'+cn.country_slug+'/pestreports/">'+cn.name.encode('utf-8'))+': '+str(p)+'</a>',str(cn.cn_lat),str(cn.cn_long)])
                    maparray1+='citymap[\''+str(cn.country_slug)+'\'] = {center: new google.maps.LatLng('+str(cn.cn_lat)+','+str(cn.cn_long)+'), text:\''+str(cn.name)+': '+str(p)+''+'\', html:\''+str('<a href="/countries/'+cn.country_slug+'/pestreports/">'+cn.name)+': '+str(p)+'</a>'+'\',  population:' +str(p)+'};'
              
            context['map']=maparray
            context['map1']=maparray1
            
        if self.kwargs['type'] == 'pestreportstat':
            context['type_label'] = 'Official pest report (Art. VIII.1a)'
            context['item'] = 'pestreportstat'
            context['link_to_item'] = 'pest-report-detail'
            arrayGen={'1ANIMK':'Animalia;',
                      '1ARCAK':'Archaea;',
                      '1BACTK':'Bacteria;',
                      '1CHROK':'Chromista;',
                      '1FUNGK':'Fungi;',
                      '1PLAK':'Plantae;',
                      '1PROTK':'Protozoa;',
                      '1VIRUK':'Viruses and viroids;'}
            cns= CountryPage.objects.all()
          
            tot_p=0
            for cn in cns:
              pests=PestReport.objects.filter(country_id=cn.id,is_version=False)
              p=pests.count()
              tot_p+=p
              for pp in pests:
                  e=EppoCode.objects.filter(codename=pp.pest_identity)
                  if e:
                    ecode=e[0].code
                    codeparent=e[0].codeparent
                    for h in range(1,10):
                          e1=EppoCode.objects.filter(code=codeparent)
                          if(e1.count()>0):
                             if(e1[0].codeparent=='null'):
                                   break
                             else:
                                  ecode=e1[0].code
                                  codeparent=e1[0].codeparent
                                  h=h+1
                    
                    aaa=arrayGen[codeparent]
                    aaa+=str(pp.id)+'*'
                    arrayGen[codeparent]=aaa
                   
            
            datachart=''
               
            for h in arrayGen:
                s=arrayGen[h].split(';');
                values=s[1].split('*');
                val=len(values)-1
                perc=0
                if tot_p>0:
                    perc=(val*100/ tot_p)
                datachart+= ' {  y: '+str(perc)+', legendText:"'+s[0]+'", label: "'+s[0]+' '+str(perc)+'%" },'
            context['datachart']=datachart

            
        elif self.kwargs['type'] == 'contactpoints':
            context['type_label'] = 'Contact points'
            context['users']=User.objects.all()
            context['cns']=CountryPage.objects.all()
         
            context['items']=IppcUserProfile.objects.filter(contact_type='1')|IppcUserProfile.objects.filter(contact_type='2')|IppcUserProfile.objects.filter(contact_type='3')|IppcUserProfile.objects.filter(contact_type='4')
            context['counttotal'] =context['items'].count() 
            context['link_to_item'] = 'contactpoint'
                 
        elif self.kwargs['type'] == 'nppo':
            context['type_label'] = dict(BASIC_REP_TYPE_CHOICES)[1]
            context['link_to_item'] = 'reporting-obligation-detail'
            context['items']= ReportingObligation.objects.filter(reporting_obligation_type=1,is_version=False)
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'entrypoints':
            context['type_label'] = dict(BASIC_REP_TYPE_CHOICES)[2]
            context['link_to_item'] = 'reporting-obligation-detail'
            context['items']= ReportingObligation.objects.filter(reporting_obligation_type=2,is_version=False)
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'regulatedpests':
            context['type_label'] = dict(BASIC_REP_TYPE_CHOICES)[3]
            context['link_to_item'] = 'reporting-obligation-detail'
            context['items']= ReportingObligation.objects.filter(reporting_obligation_type=3,is_version=False)
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'legislation':
            context['type_label'] = dict(BASIC_REP_TYPE_CHOICES)[4]
            context['link_to_item'] = 'reporting-obligation-detail'
            context['items']= ReportingObligation.objects.filter(reporting_obligation_type=4,is_version=False)
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'emergencyactions':
            context['type_label'] = dict(EVT_REP_TYPE_CHOICES)[1]
            context['link_to_item'] = 'event-reporting-detail'
            context['items']= EventReporting.objects.filter(event_rep_type=1,is_version=False)
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'noncompliance':
            context['type_label'] = dict(EVT_REP_TYPE_CHOICES)[2]
            context['link_to_item'] = 'event-reporting-detail'
            context['items']= EventReporting.objects.filter(event_rep_type=2,is_version=False)
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'plantprotection':
            context['type_label'] = dict(EVT_REP_TYPE_CHOICES)[3]
            context['link_to_item'] = 'event-reporting-detail'
            context['items']= EventReporting.objects.filter(event_rep_type=3,is_version=False)
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'peststatus':
            context['type_label'] = dict(EVT_REP_TYPE_CHOICES)[4]
            context['link_to_item'] = 'event-reporting-detail'
            context['items']= EventReporting.objects.filter(event_rep_type=4)
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'phytosanitaryrequirements':
            context['type_label'] = dict(EVT_REP_TYPE_CHOICES)[5]
            context['link_to_item'] = 'event-reporting-detail'
            context['items']= EventReporting.objects.filter(event_rep_type=5,is_version=False)
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'pfa':
            context['type_label'] = 'Pest free areas'
            context['link_to_item'] = 'pfa-detail'
            context['items']= PestFreeArea.objects.filter(is_version=False)
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'ispm15':
            context['type_label'] = 'Implementation of ISPM 15'
            context['link_to_item'] = 'implementationispm-detail'
            context['items']= ImplementationISPM.objects.filter(is_version=False)
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'countrynews':
            context['type_label'] = 'Country news'
            context['link_to_item'] = 'country-news-detail'
            context['items']= CountryNews.objects.all()
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'countrywebsites':
            context['type_label'] = 'Country websites'
            context['link_to_item'] = 'website-detail'
            context['items']= Website.objects.all()
            context['counttotal'] =context['items'].count() 
        elif self.kwargs['type'] == 'cnpublication':
            context['type_label'] = 'Country publication'
            context['link_to_item'] = 'country-publication-detail'
            context['items']= CnPublication.objects.all()
            context['counttotal'] =context['items'].count() 
         
        return context
    

class ReportingSystemSummaryCNListView(ListView):
    """  ReportingSystemSummaryCNListView list  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_reportingsystem_summary.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(ReportingSystemSummaryCNListView, self).get_context_data(**kwargs)
        context['type'] = self.kwargs['type'] 
          
        summary=''
        curryear=timezone.now().year
        currmonth=timezone.now().month
        months=['January','February','March','April','May','June','July','August','September','October','November','December']
        rangeyear=range(2012,curryear+1)

        for yy in list(reversed(rangeyear)):
            m_i=1
            month=''
            summary+=' <div class="time"><div class="tidate">'+str(yy)+'</div><div class="timatter"><div class="row"></div><div class="row">'
            for ss in months:
                summary+='<div class="col-sm-3">'
                if m_i <10:
                    month='0'+str(m_i)
                else:
                  month=str(m_i)
                if m_i > currmonth and yy == curryear:
                    summary+='&#160;'
                else:  
                    summary+='<a href="/countries/reportingsystem/all/'+str(yy)+'/'+month+'/">'+str(ss)+'</a>'
                summary+='</div>'
                m_i=m_i+1    
            summary+='</div></div> <div class="clearfix"></div></div>'
        context['summary'] = summary
    
        return context
    
    
class ReportingSystemCNListView(ListView):
    """  ReportingSystemCNListView list  """
    context_object_name = 'latest'
    model = CountryPage
    template_name = 'countries/countries_reportingsystem.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(ReportingSystemCNListView, self).get_context_data(**kwargs)
        context['type'] = self.kwargs['type'] 
          
        if self.kwargs['type'] == 'all':
            sel_year=self.kwargs['year']
            sel_months=self.kwargs['month']
          
            context['sel_months'] = sel_months
            context['sel_year'] = sel_year
            selmonth=''
            selnmonth=''
            if sel_months == '01' :
                selmonth='1'
                selnmonth='January'
            if sel_months == '02' :
                selmonth='2'
                selnmonth='February'
            if sel_months == '03' :
                selmonth='3'
                selnmonth='March'
            if sel_months == '04' :
                selmonth='4'
                selnmonth='April'
            if sel_months == '05' :
                selmonth='5'
                selnmonth='May'
            if sel_months == '06' :
                selmonth='6'
                selnmonth='June'
            if sel_months == '07' :
                selmonth='7'
                selnmonth='July'
            if sel_months == '08' :
                selmonth='8'
                selnmonth='August'
            if sel_months == '09' :
                selmonth='9'
                selnmonth='September'
            if sel_months == '10' :
                selmonth='10'
                selnmonth='October'
            if sel_months == '11' :
                selmonth='11'
                selnmonth='November'
            if sel_months == '12' :
                selmonth='12'
                selnmonth='December'
         
      
                
            context['namemonth'] = selnmonth
            new_pests=[]
            up_pests=[]
            tot_pests=[]
           
            pests= PestReport.objects.filter(is_version=False,status=CONTENT_STATUS_PUBLISHED)
            for p in pests:
                pest=[]
                if p.country_id == 199:
                    print("no")
                else:
                    if p.publish_date != None and str(p.publish_date.year) == sel_year :
                        if str(p.publish_date.month) == selmonth:
                           new_pests.append(p)
                           pest.append('new')
                           pest.append('Pest report')
                           pest.append(p)
                           pest.append('pest-report-detail')
                           tot_pests.append(pest)
                    if p.modify_date != None and str(p.modify_date.year) == sel_year :
                        if str(p.modify_date.month) == selmonth:
                            if p in new_pests:
                                print('aaaa')
                            else:      
                                up_pests.append(p)
                                pest.append('up')
                                pest.append('Pest report')
                                pest.append(p)
                                pest.append('pest-report-detail')
                                tot_pests.append(pest)
        #    for i in range(1,6):
        #        reps=ReportingObligation.objects.filter(reporting_obligation_type=i,is_version=False)
        #        for r in reps:
        #            rep=[]
        #            if r.publication_date != None and (r.publication_date.year) == sel_year :
        #                if str(r.publication_date.month) == selmonth:
        #                   new_pests.append(r)
        #                   rep.append('new')
        #                   rep.append(dict(BASIC_REP_TYPE_CHOICES)[i])
        #                   rep.append(r)
        #                   rep.append('reporting-obligation-detail')
        #                   tot_pests.append(rep)
        #            if r.modify_date != None and str(r.modify_date.year) == sel_year :
        #                if str(r.modify_date.month) == selmonth:
        #                    if r in new_pests:
        #                        print('aaaa')
        #                    else:      
        #                        up_pests.append(r)
        #                        rep.append('up')
        #                        rep.append(dict(BASIC_REP_TYPE_CHOICES)[i])
        #                        rep.append(r)
        #                        rep.append('reporting-obligation-detail')
        #                        tot_pests.append(rep)
        #        evrep=EventReporting.objects.filter(event_rep_type=i,is_version=False)
        #        for e in evrep:
        #            ev=[]
        #            if e.publication_date != None and (e.publication_date.year) == sel_year :
        #                if str(e.publication_date.month) == selmonth:
        #                   new_pests.append(r)
        #                   ev.append('new')
        #                   ev.append(dict(EVT_REP_TYPE_CHOICES)[i])
        #                   ev.append(e)
        #                   ev.append('event-reporting-detail')
        #                   tot_pests.append(ev)
        #            if e.modify_date != None and str(e.modify_date.year) == sel_year :
        #                if str(e.modify_date.month) == selmonth:
        #                    if e in new_pests:
        #                        print('aaaa')
        #                    else:      
        #                        up_pests.append(e)
        #                        ev.append('up')
        #                        ev.append(dict(EVT_REP_TYPE_CHOICES)[i])
        #                        ev.append(e)
        #                        ev.append('event-reporting-detail')
        #                        tot_pests.append(ev)
            context['tot_pests']  = tot_pests
        return context
    
    
import csv
from django.http import HttpResponse	

def contactPointExtractor(request):
    # Create the HttpResponse object with the appropriate CSV header.
    contacts=IppcUserProfile.objects.filter(contact_type='1')|IppcUserProfile.objects.filter(contact_type='2')|IppcUserProfile.objects.filter(contact_type='3')|IppcUserProfile.objects.filter(contact_type='4')
    users=User.objects.all()
    cns=CountryPage.objects.all()
             
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="contactpoints.csv"'

    writer = csv.writer(response)
    writer.writerow(['Country', 'Contact Type', 'Prefix', 'First Name','Last Name','Email','Alternate E-mail','Organization','Address'])
    GENDER_CHOICES = (
        (1, "Mr."),
        (2, "Ms."),
        (3, "Mrs."),
        (4, "Professor."),
        (5, "M."),
        (6, "Mme."),
        (7, "Dr."),
        (8, "Sr."),
        (9, "Sra."),
    )
    for c in contacts:
        country=''
        c_type=''
        for cn in cns:          
            if cn.id == c.country_id:
               country = cn
        for o in c.contact_type.all():
            if o.id==1 or o.id==2 or o.id==3 or o.id==4:
                c_type=o
    
        c_gender=''
        for k,v in GENDER_CHOICES:
            if c.gender == k:
               c_gender=v
               break
        c_first_name= c.first_name
        c_last_name= c.last_name
        c_email=''
        for u in users:
            if u.id == c.user_id:
             c_email= u.email
        c_emailalt =c.email_address_alt
        c_organization=c.address1
        c_address=c.address2
        
        writer.writerow([country,c_type,c_gender, c_first_name.encode('utf-8'), c_last_name.encode('utf-8'),c_email,c_emailalt,c_organization.encode('utf-8'),c_address.encode('utf-8')])

    return response	

def draftprotocol_compilecomments(request,id=None):
    # Create the HttpResponse object with the appropriate CSV header.
       response = HttpResponse(content_type='text/csv')
       response['Content-Disposition'] = 'attachment; filename="compiled_comments_DP_'+timezone.now().strftime('%Y%m%d%H%M%S')+'.csv"'
       writer = csv.writer(response)
       writer.writerow(['Author','Expertise On This Pest','Institution','Comment','Attachments'])
       draftprotocol = get_object_or_404(DraftProtocol,  pk=id)
       queryset = DraftProtocolComments.objects.filter(draftprotocol_id=draftprotocol.id)
     
         
       for obj in queryset:
            author=''
            expertise=''
            intitution=''
            comment=''
            attch=''
            attch2=''
            lastdate=''
            
            author=obj.author
            expertise=obj.expertise
            intitution=obj.institution
            comment=obj.comment
            attch=obj.filetext
            attch2=obj.filefig
            writer.writerow([author,expertise.encode('utf-8'),intitution.encode('utf-8'), comment.encode('utf-8'), attch,attch2])

       return response	

#class QuestionListView(ListView):
#    template_name = 'question/index.html'
#    context_object_name = 'latest_question_list'
#    def get_queryset(self):
#        """Return the last five published questions."""
#        return Question.objects.order_by('-pub_date').all
#
#    def get_context_data(self, **kwargs):
#        context = super(QuestionListView, self).get_context_data(**kwargs)
#        arrayquestions=[]
#        questions=Question.objects.all()
#        
#        for q in questions:
#            array_q=[]
#            answers=Answer.objects.filter(question_id=q.id).count()
#            answersbest=Answer.objects.filter(question_id=q.id,bestanswer='1').count()
#           
#            array_q.append(q.id)
#            array_q.append(q.question_title)
#            array_q.append(answers)
#            array_q.append(answersbest)
#            array_q.append(q.pub_date)
#            arrayquestions.append(array_q)
#       
#        context['questions']= arrayquestions
#        return context
#
#
#
#
#class QuestionDetailView(DetailView):
#    model = Question
#    template_name = 'question/detail.html'
#    def get_queryset(self):
#        """Return the last five published question."""
#        return Question.objects.filter(pub_date__lte=timezone.now())
#        
#
#class QuestionAnswersView(DetailView):
#    model = Question
#    template_name = 'question/answers.html'
#    
#    def get_context_data(self, **kwargs):
#        context = super(QuestionAnswersView, self).get_context_data(**kwargs)
#        q_id=self.kwargs['pk']
#        answers=Answer.objects.filter(question_id=q_id)
#        
#        arrayanswers=[]
#        for a in answers:
#            array_a=[]
#            answervote=AnswerVotes.objects.filter(answer=a).count()
#            answervoteup=AnswerVotes.objects.filter(answer=a,up='1').count()
#            answervotedown=AnswerVotes.objects.filter(answer=a,up='-1').count()
#            upval=0
#            downval=0
#            if answervoteup >0:
#                upval=answervoteup/answervote*100
#            if answervotedown >0:
#                downval=answervotedown/answervote*100
#            
#            array_a.append(a.answertext)
#            array_a.append(a.id)
#            array_a.append(upval)
#            array_a.append(downval)
#            arrayanswers.append(array_a)
#       
#        context['answers']= arrayanswers
#        return context
#
#
#@login_required
#@permission_required('ippc.add_question', login_url="/accounts/login/")
#def question_create(request):
#    """ Create question """
#    user = request.user
#    author = user
#
#    form = QuestionForm(request.POST)
#    if request.method == "POST":
#         if form.is_valid():
#            new_question = form.save(commit=False)
#            new_question.user = request.user
#            new_question.user_id = author.id
#            form.save()
#            info(request, _("Successfully created Question."))
#            return redirect("answers", pk=new_question.id)
#         else:
#             return render_to_response('question/question_create.html', {'form': form,},
#             context_instance=RequestContext(request))
#       
#    else:
#        form = QuestionForm( instance=Question())
#    
#    return render_to_response('question/question_create.html', {'form': form,},
#        context_instance=RequestContext(request))
#
#@login_required
#@permission_required('ippc.change_question', login_url="/accounts/login/")
#def question_edit(request, id=None, template_name='question/question_edit.html'):
#    """ Edit question """
#    if id:
#        question = get_object_or_404(Question,  pk=id)
#    else:
#        question = Question()
#      
#    if request.POST:
#
#        form =QuestionForm(request.POST, instance=question)
#        if form.is_valid():
#            form.save()
#          
#            return redirect("answers", pk=id)
#    else:
#        form = QuestionForm(instance=question)
#      
#        
#    return render_to_response(template_name, {
#        'form': form, "question": question
#    }, context_instance=RequestContext(request))
#    
#        
#@login_required
#@permission_required('ippc.add_answer', login_url="/accounts/login/")
#def answer_create(request, question_id):
#    """ Create answer """
#    q = get_object_or_404(Question, pk=question_id)
#    
#    user = request.user
#    author = user
#
#    form =AnswerForm(request.POST)
#    if request.method == "POST":
#         if form.is_valid():
#            new_answer = form.save(commit=False)
#            new_answer.user = request.user
#            new_answer.user_id = author.id
#            new_answer.question_id = q.id
#            form.save()
#            info(request, _("Successfully added Answer."))
#            return redirect("answers", pk=q.id)
#         else:
#            return render_to_response('question/answer_create.html', {'form': form,"question":q,},
#            context_instance=RequestContext(request))
#       
#    else:
#        form = AnswerForm( instance=Question())
#    
#    return render_to_response('question/answer_create.html', {'form': form,"question":q,},
#        context_instance=RequestContext(request))
#    
#@login_required
#@permission_required('ippc.change_answer', login_url="/accounts/login/")
##def answer_edit(request, id=None , question_id ,template_name='question/answer_edit.html'):
#def answer_edit(request, question_id,id=None,template_name='question/answer_edit.html'):
#    """ edit answer """
#    print(question_id)
#    q = get_object_or_404(Question, pk=question_id)
#    """ Edit question """
#    if id:
#        answer = get_object_or_404(Answer,  pk=id)
#    else:
#        answer = Answer()
#      
#    if request.POST:
#        form =AnswerForm(request.POST, instance=answer)
#        if form.is_valid():
#            form.save()
#            return redirect("answers", pk=question_id)
#    else:
#        form = AnswerForm(instance=answer)
#      
#        
#    return render_to_response(template_name, {
#        'form': form, "answer": answer,"question":q,
#    }, context_instance=RequestContext(request))
#    
#        
#def vote_answer_up(request,question_id,id=None,):
#    answer = get_object_or_404(Answer, pk=id)
#    q = get_object_or_404(Question, pk=question_id)
#    answers=Answer.objects.filter(question_id=question_id)
#        
#    if AnswerVotes.objects.filter(answer_id=id, user_id=request.user.id).exists():
#        return render(request, 'question/answers.html', {
#        'question': q,
#        'answers':answers,
#        'error_message': "Sorry, but you have already voted."
#        })
#
#    try:
#        up = 1
#    except (KeyError):
#        # Redisplay the poll voting form.
#        return render(request, 'question/answers.html', {
#            'question': q,
#            'answers':answers,
#            'error_message': "You didn't select a rate.",
#        })
#    else:
#       # selected_choice.votes += 1
#        #selected_choice.save()
#        
#        v = AnswerVotes(user=request.user, answer=answer,up='1')
#        v.save()
#        return redirect("answers", pk=q.id)
#
#def vote_answer_down(request,question_id,id=None,):
#    answer = get_object_or_404(Answer, pk=id)
#    q = get_object_or_404(Question, pk=question_id)
#    answers=Answer.objects.filter(question_id=question_id)
#    if AnswerVotes.objects.filter(answer_id=id, user_id=request.user.id).exists():
#        return render(request, 'question/answers.html', {
#        'question': q,
#        'answers':answers,
#        'error_message': "Sorry, but you have already voted."
#        })
#    try:
#        down = -1
#    except (KeyError):
#        # Redisplay the poll voting form.
#        return render(request, 'question/answers.html', {
#            'question': q,
#            'answers':answers,
#            'error_message': "You didn't select a rate.",
#        })
#    else:
#       # selected_choice.votes += 1
#        #selected_choice.save()
#        
#        v = AnswerVotes(user=request.user, answer=answer,up='-1')
#        v.save()
#        return redirect("answers", pk=q.id)

 
# 
# 
# 
# 
# 
# 
# 
# 
class QAQuestionListView(ListView):
    template_name = 'question/index.html'
    context_object_name = 'latest_question_list'
    def get_queryset(self):
        """Return the last five published questions."""
        return QAQuestion.objects.order_by('-pub_date').all
#
    def get_context_data(self, **kwargs):
        context = super(QAQuestionListView, self).get_context_data(**kwargs)
        arrayquestions=[]
        questions=QAQuestion.objects.all()
        
        for q in questions:
            array_q=[]
            answers=QAAnswer.objects.filter(question_id=q.id).count()
            answersbest=QAAnswer.objects.filter(question_id=q.id,bestanswer='1').count()
            answers_count=QAAnswer.objects.filter(question_id=q.id,status=2,answerdiscard=2).count()
           #print('>>>>>>>>>>'+str(answers_count))
            array_q.append(q.id)
            array_q.append(q.title)
            array_q.append(q.status)
            array_q.append(answers_count)
            array_q.append(answersbest)
            array_q.append(q.publish_date)
            array_q.append(q.questionopen)
            array_q.append(q.questiondiscard)
          
            arrayquestions.append(array_q)
       
        context['questions']= arrayquestions
        return context

class QAQuestionDetailView(DetailView):
    model = QAQuestion
    template_name = 'question/detail.html'
    def get_queryset(self):
        """Return the last five published question."""
        return QAQuestion.objects.filter(pub_date__lte=timezone.now())
        

class QAQuestionAnswersView(DetailView):
    model = QAQuestion
    template_name = 'question/answers.html'
    
    def get_context_data(self, **kwargs):
        context = super(QAQuestionAnswersView, self).get_context_data(**kwargs)
        q_id=self.kwargs['pk']
        answers=QAAnswer.objects.filter(question_id=q_id)
        arrayanswers=[]
        for a in answers:
            array_a=[]
            answervote=AnswerVotes.objects.filter(answer=a).count()
            answervoteup=AnswerVotes.objects.filter(answer=a,up='1').count()
            answervotedown=AnswerVotes.objects.filter(answer=a,up='-1').count()
            upval=0
            downval=0
            if answervoteup >0:
                upval=answervoteup/answervote*100
            if answervotedown >0:
                downval=answervotedown/answervote*100
            
            array_a.append(a.answertext)
            array_a.append(a.id)
            array_a.append(upval)
            array_a.append(downval)
            array_a.append(a.status)
            array_a.append(a.bestanswer)
            array_a.append(a.answerdiscard)
            
            arrayanswers.append(array_a)
       
        context['question']= get_object_or_404(QAQuestion,  pk=q_id)
        context['answers']= arrayanswers
        return context


def send_qa_notification_message(type,id):
    """ send_qa_notification_message """
    #send notification to QA moderator
    subject=''
    textmessage_moderator=''
    textmessage_user=''
    question=None
    answer =None
    if type == 'QAQuestion':
        question = get_object_or_404(QAQuestion,  pk=id)
        subject='IRSS - Q&A Forum automatic notification: new question has been posted'
        textmessage_moderator='<p>Dear Q&A moderator,<br><br>A new question has been posted in the Q&A forum in draft mode and requires your action.<br><br>You can view it at the following url: <a href="https://www.ippc.int/en/qa/'+str(id)+'/answers/">https://www.ippc.int/en/qa/'+str(id)+'/answers/</a><br><br>International Plant Protection Convention team </p>'
    else:    
        answer = get_object_or_404(QAAnswer,  pk=id)
        qid=answer.question.id
        subject='IRSS - Q&A Forum automatic notification: new answer has been posted'
        textmessage_moderator='<p>Dear Q&A moderator,<br><br>A new answer has been posted in the Q&A forum in draft mode and requires your action.<br><br>You can view it at the following url: <a href="https://www.ippc.int/en/qa/'+str(qid)+'/answers/">https://www.ippc.int/en/qa/'+str(qid)+'/answers/</a><br><br>International Plant Protection Convention team </p>'
     
    emailto_moderator = ['']
    for g in Group.objects.filter(id=68):## mettere ID GROUP moderator QA
            users = g.user_set.all()
            for u in users:
               user_obj=User.objects.get(username=u)
               if str(user_obj.email)!='':
                   emailto_moderator.append(str(user_obj.email))
    message = mail.EmailMessage(subject,textmessage_moderator,'ippc@fao.org', emailto_moderator, ['paola.sentinelli@fao.org'])
    message.content_subtype = "html"
    #print('test-sending')
    sent =message.send()
        
from django.http import HttpResponseRedirect

@login_required
def question_create(request):
    """ Create question """
    user = request.user
    author = user

    form = QAQuestionForm(request.POST)
    if request.method == "POST":
         if form.is_valid():
            new_question = form.save(commit=False)
            new_question.user = request.user
            new_question.user_id = author.id
            new_question.status = 1
            form.save()
            info(request, _("Successfully added Question in the Q&A forum and it in is draft mode. A moderator will revise it. You will be shortly notified with action taken. "))
           #print('dddddddddddddddddddddddd:'+str(new_question.id))
            send_qa_notification_message('QAQuestion',new_question.id)
            #return redirect("answers", pk=new_question.id)
            return HttpResponseRedirect("/qa")
         else:
             return render_to_response('question/question_create.html', {'form': form,},
             context_instance=RequestContext(request))
       
    else:
        form = QAQuestionForm( instance=QAQuestion())
    
    return render_to_response('question/question_create.html', {'form': form,},
        context_instance=RequestContext(request))

        
@login_required
def answer_create(request, question_id):
    """ Create answer """
    q = get_object_or_404(QAQuestion, pk=question_id)
    
    user = request.user
    author = user

    form =QAAnswerForm(request.POST)
    if request.method == "POST":
         if form.is_valid():
            new_answer = form.save(commit=False)
            new_answer.user = request.user
            new_answer.user_id = author.id
            new_answer.status = 1
            new_answer.question_id = q.id
            form.save()
           #print('>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>'+str(new_answer.question_id ))
            
            info(request, _("Successfully added Answer in the Q&A forum and it in is draft mode. A moderator will revise it. You will be shortly notified with action taken. "))
            send_qa_notification_message('QAAnswer',new_answer.id)
            return redirect("answers", pk=q.id)
         else:
            return render_to_response('question/answer_create.html', {'form': form,"question":q,},
            context_instance=RequestContext(request))
       
    else:
        form = QAAnswerForm( instance=QAQuestion())
    
    return render_to_response('question/answer_create.html', {'form': form,"question":q,},
        context_instance=RequestContext(request))


def vote_answer_up(request,question_id,id=None,):
    answer = get_object_or_404(QAAnswer, pk=id)
    q = get_object_or_404(QAQuestion, pk=question_id)
    answers=QAAnswer.objects.filter(question_id=question_id)
        
    if AnswerVotes.objects.filter(answer_id=id, user_id=request.user.id).exists():
        return render(request, 'question/answers.html', {
        'question': q,
        'answers':answers,
        'error_message': "Sorry, but you have already voted."
        })

    try:
        up = 1
    except (KeyError):
        #Redisplay the poll voting form.
        return render(request, 'question/answers.html', {
            'question': q,
            'answers':answers,
            'error_message': "You didn't select a rate.",
        })
    else:
        v = AnswerVotes(user=request.user, answer=answer,up='1')
        v.save()
        return redirect("answers", pk=q.id)

def vote_answer_down(request,question_id,id=None,):
    answer = get_object_or_404(QAAnswer, pk=id)
    q = get_object_or_404(QAQuestion, pk=question_id)
    answers=QAAnswer.objects.filter(question_id=question_id)
    if AnswerVotes.objects.filter(answer_id=id, user_id=request.user.id).exists():
        return render(request, 'question/answers.html', {
        'question': q,
        'answers':answers,
        'error_message': "Sorry, but you have already voted."
        })
    try:
        down = -1
    except (KeyError):
        return render(request, 'question/answers.html', {
            'question': q,
            'answers':answers,
            'error_message': "You didn't select a rate.",
        })
    else:
        v = AnswerVotes(user=request.user, answer=answer,up='-1')
        v.save()
        return redirect("answers", pk=q.id)

class FAQsListView(ListView):
    template_name = 'faq/faqlist.html'
    context_object_name = 'latest_faq_list'
    def get_queryset(self):
        """Return the last five published questions."""
        return FAQsCategory.objects.order_by('-pub_date').all

    def get_context_data(self, **kwargs):
        context = super(FAQsListView, self).get_context_data(**kwargs)
        array_faqcategories=[]
        
        langsel=self.request.LANGUAGE_CODE
       #print(langsel)
        faqcategories= FAQsCategory.objects.order_by('faqcat_oder').all()

        for fc in faqcategories:
            array_faqsitems=[]
            faqs_array=[]
            faqs=FAQsItem.objects.filter(faqcategory_id=fc.id)
           
            #print(faqs)
            array_faqsitems.append(fc.id)
            fc_title=fc.title
            if  langsel!='en' and langsel!='':
                faqcategory_tra=TransFAQsCategory.objects.filter(translation_id=fc.id,lang=langsel)
                if faqcategory_tra:
                    fc_title=faqcategory_tra[0].title
                
                #faqs=TransFAQsItem.objects.filter(faqcategory_id=fc.id)
                for faq_item in faqs:
                    faq= TransFAQsItem.objects.filter(translation_id=faq_item.id,lang=langsel)
                    if faq:
                       # print(faq_item)
                       # print(faq)
                       # print(faq[0].faq_description)
                        faqs_array.append(faq[0])
                    else :
                        faqs_array.append(faq_item)
            else:
                faqs_array=faqs        
            #print(fc_title)
            array_faqsitems.append(fc_title)
            array_faqsitems.append(faqs_array)
            
            array_faqcategories.append(array_faqsitems)
        context['array_faqcategories']= array_faqcategories
        return context
       
class FAQsCategoryDetailView(DetailView):
    """ FAQsCategory detail page """
    model = FAQsCategory
    context_object_name = 'faqcategory'
    template_name = 'faq/faqcategory_detail.html'
    queryset = FAQsCategory.objects.filter()

 
   
class FAQsItemDetailView(DetailView):
    """ FAQsItems detail page """
    model = FAQsItem
    context_object_name = 'faq'
    template_name = 'faq/faq_detail.html'
    queryset = FAQsItem.objects.filter()

@login_required
@permission_required('ippc.add_faqcategory', login_url="/accounts/login/")
def faqcategory_create(request):
    """ faqcategory create  """
    user = request.user
    author = user
#
    form = FAQsCategoryForm(request.POST)
    if request.method == "POST":
         if form.is_valid():
            new_faqcat = form.save(commit=False)
            form.save()
            info(request, _("Successfully created FAQs Category."))
            return redirect("faqcategory-detail", pk=new_faqcat.id)
         else:
             return render_to_response('faq/faqcategory_create.html', {'form': form,},
             context_instance=RequestContext(request))
       
    else:
        form = FAQsCategoryForm( instance=FAQsCategory())
    
    return render_to_response('faq/faqcategory_create.html', {'form': form,},
        context_instance=RequestContext(request))

@login_required
@permission_required('ippc.change_faqcategory', login_url="/accounts/login/")
def faqcategory_edit(request, id=None, template_name='faq/faqcategory_edit.html'):
    """ Edit faqcategory """
    if id:
        faqcategory = get_object_or_404(FAQsCategory,  pk=id)
    else:
        faqcategory = FAQsCategory()
      
    if request.POST:
#
        form =FAQsCategoryForm(request.POST, instance=faqcategory)
        if form.is_valid():
            form.save()
          
            return redirect("faqcategory-detail", pk=id)
    else:
        form = FAQsCategoryForm(instance=faqcategory)
      
        
    return render_to_response(template_name, {
        'form': form, "faqcategory": faqcategory
    }, context_instance=RequestContext(request))
    
        
@login_required
@permission_required('ippc.add_faqitem', login_url="/accounts/login/")
def faq_create(request):
    """ Create faq """
    
    form =FAQsItemForm(request.POST)
    if request.method == "POST":
         if form.is_valid():
            new_faq = form.save(commit=False)
            form.save()
            info(request, _("Successfully added new faq."))
            return redirect("faq-detail", pk=new_faq.id)
         else:
            return render_to_response('faq/faq_create.html', {'form': form,},
            context_instance=RequestContext(request))
       
    else:
        form = FAQsItemForm( instance=FAQsItem())
    
    return render_to_response('faq/faq_create.html', {'form': form,},
        context_instance=RequestContext(request))
    
@login_required
@permission_required('ippc.change_faqitem', login_url="/accounts/login/")
def faq_edit(request, id=None,template_name='faq/faq_edit.html'):
    """ edit faq """
    if id:
        faq = get_object_or_404(FAQsItem,  pk=id)
    else:
        faq = FAQsItem()
      
    if request.POST:
        form =FAQsItemForm(request.POST, instance=faq)
        if form.is_valid():
            form.save()
            return redirect("faq-detail", pk=id)
    else:
        form = FAQsItemForm(instance=faq)
      
        
    return render_to_response(template_name, {
        'form': form, "faq": faq, 
    }, context_instance=RequestContext(request))


class UserAutoRegistrationListView(ListView):
    """    UserAutoRegistration List view """
    context_object_name = 'latest'
    model = UserAutoRegistration
    date_field = 'date'
    template_name = 'accounts_auto/accounts_list.html'
    queryset = UserAutoRegistration.objects.all().order_by('-publish_date')

from django.http import HttpResponseRedirect
#@login_required
def auto_register(request):
    """ auto_register """
    
    form =UserAutoRegistrationForm(request.POST)
    if request.method == "POST" :
         if form.is_valid() and request.POST['captcha'] ==  request.POST['result_element']:
            new_user = form.save(commit=False)
            form.save()
            info(request, _("Successfully registered to subscribe to IPPC News, Announcements or Calls, the IPPC Team will revise your subscription."))
            subject='A new user has self-subscribed for IPPC News, Announcements or Calls'  
            msg='<p>A new user has self-subscribed for IPPC News, Announcements or Calls.<br><br>Please use the link below to view the list of users pending approval:<br><br><a href="https://www.ippc.int/accounts/pendingapproval/">https://www.ippc.int/accounts/pendingapproval/</a>.'
            message = mail.EmailMessage(subject,msg,'ippc@fao.org', ['ippc-it@fao.org'], ['paola.sentinelli@fao.org'])
            message.content_subtype = "html"
            #print('test-sending')
            sent =message.send()
            
            return HttpResponseRedirect("/")
         else:
            error_captcha=''
            if not(request.POST['captcha'] == request.POST['result_element'] ) :
                error_captcha='error'
                  
            return render_to_response('accounts_auto/autoregister_create.html', {'form': form,'x_element': request.POST['x_element'],'y_element': request.POST['y_element'],'result_element': request.POST['result_element'] ,'error_captcha':error_captcha},
            context_instance=RequestContext(request))
    else:
         x_element=random.randint(1,10)   
         y_element=random.randint(1,10)
         result_element=x_element+y_element
     
         form = UserAutoRegistrationForm()
# 
    return render_to_response('accounts_auto/autoregister_create.html', {'form': form ,'x_element':x_element,'y_element':y_element,'result_element':result_element},
        context_instance=RequestContext(request))
     
    
@login_required
@permission_required('ippc.change_userautoregistration', login_url="/accounts/login/")
def auto_register_approve(request, id=None):
    """  auto registered User approve  """
    
    if id:
       #print('APPROVED')
        newuser = get_object_or_404(UserAutoRegistration,  pk=id)
        user_obj=User.objects.filter(email=newuser.email)
        if user_obj.count()>0:
            newuser.status=3
            newuser.save()
            error(request, _("An user with the same email address is alredy registered in the system."))
            return HttpResponseRedirect("/accounts/pendingapproval/")
        else:
            #create new user
            user1=User()
            user1.username=slugify(newuser.firstname+"-"+newuser.lastname).lower()
            user1.first_name=newuser.firstname
            user1.last_name=newuser.lastname
            user1.email=newuser.email
            user1.save()
            if newuser.subscribe_news == 1:
                g1=Group.objects.get(name="News Notification group")
                user1.groups.add(g1)
            if newuser.subscribe_announcement == 1:
                g2=Group.objects.get(name="Announcement Notification group")
                user1.groups.add(g2)
            if newuser.subscribe_calls == 1:
                g3=Group.objects.get(name="Calls Notification group")
                user1.groups.add(g3)
            
            
            #set profile
            userp = get_object_or_404(IppcUserProfile, user_id=user1.id)
            userp.first_name=newuser.firstname
            userp.last_name=newuser.lastname
            userp.country=newuser.country
            userp.address1=newuser.organisation
            userp.save()
              
            #sendmessage to new user
            user_email = []
            user_email.append(newuser.email)
            subtext=''
            if newuser.subscribe_news == 1:
               subtext+=" News"
            if newuser.subscribe_announcement == 1:
                subtext+="Announcements"
	    if newuser.subscribe_calls == 1:
                subtext+=" Calls"
	
            subject='Your subscription to IPPC '+subtext+' has been approved.' 
            msg='<p>Your subscription to '+subtext+' has been approved. From now on you will receive an automatic notification when one of this item will be posted on our website.<br><br>A related IPPC account has been created for you, please activate it! <br><br>You will need this in case you want to un-subscribe from these notifications. <br><br>To active the account click on the link below to set your password.<br><br><a href="https://www.ippc.int/en/account/password/reset/?next=/en/account/update/">https://www.ippc.int/en/account/password/reset/?next=/en/account/update/</a><br><br>Insert your email address, click on "Password Reset" button and follow instructions to create your password.<br><br>After setting your password, you will be able to log to the IPPC website.<br><br>IPPC Secretariat<br>'

            message = mail.EmailMessage(subject,msg,'ippc@fao.org', user_email, ['paola.sentinelli@fao.org'])
            message.content_subtype = "html"
            sent =message.send()
            #delete temporary user
            newuser.delete()
           
            info(request, _("Successfully approved user."))
            return HttpResponseRedirect("/accounts/pendingapproval/")
       

@login_required
@permission_required('ippc.delete_userautoregistration', login_url="/accounts/login/")
def auto_register_delete(request, id=None):
    """ auto registered User delete   """
    if id:
       #print('delete')
        newuser = get_object_or_404(UserAutoRegistration,  pk=id)
        newuser.delete()
     
        info(request, _("Successfully deleted user."))
        return HttpResponseRedirect("/accounts/pendingapproval/")
    


            
class IRSSActivityListView(ListView):
    """   IRSSActivity """
    context_object_name = 'latest'
    model = IRSSActivity
    date_field = 'publish_date'
    template_name = 'irss/irss_activities_list.html'
    queryset = IRSSActivity.objects.all().order_by('-publish_date', 'title')
    
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return pest reports from the specific country """
        # self.country = get_object_or_404(CountryPage, country=self.kwargs['country'])
        # CountryPage country_slug == country URL parameter keyword argument
        return IRSSActivity.objects.all().order_by('-publish_date', 'title')
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(IRSSActivityListView, self).get_context_data(**kwargs)
        context['activity_types'] =IRSS_ACT_TYPE_CHOICES
       
        #context['country'] = self.kwargs['country']
        return context
   
       
   
class IRSSActivityDetailView(DetailView):
    """ IRSSActivity detail page """
    model = IRSSActivity
    context_object_name = 'irssactivity'
    template_name = 'irss/irss_activities_detail.html'
    queryset = IRSSActivity.objects.filter()



@login_required
@permission_required('ippc.add_irssactivity', login_url="/accounts/login/")
def irss_activity_create(request):
    """ Create  IRSSActivity """
    user = request.user
    author = user
  
    form = IRSSActivityForm(request.POST or None, request.FILES)
    
         
    if request.method == "POST":
        f_form = IRSSActivityFileFormSet(request.POST, request.FILES)
        
        if form.is_valid() and f_form.is_valid() :
            new_irssactivity = form.save(commit=False)
            new_irssactivity.author = request.user
            new_irssactivity.author_id = author.id
            form.save()
            
         
           
            f_form.instance = new_irssactivity
            f_form.save()
           
            info(request, _("Successfully added IRSS Activity."))
            return redirect("irss-activity-detail", pk=new_irssactivity.id)
        else:
            return render_to_response('irss/irss_activities_create.html', {'form': form,'f_form': f_form,},
             context_instance=RequestContext(request))

          
        
    else:
        form = IRSSActivityForm(instance=IRSSActivity())
        f_form = IRSSActivityFileFormSet()
    
    return render_to_response('irss/irss_activities_create.html', {'form': form,'f_form': f_form},
        context_instance=RequestContext(request))

        

@login_required
@permission_required('ippc.change_irssactivity', login_url="/accounts/login/")
def irss_activity_edit(request,  id=None, template_name='irss/irss_activities_edit.html'):
    """ Edit   Country IRSSActivity """
    user = request.user
    author = user
    if id:
        irssactivity = get_object_or_404(IRSSActivity, pk=id)
    else:
        irssactivity = IRSSActivity()
      
    if request.POST:
        form = IRSSActivityForm(request.POST,  request.FILES, instance=irssactivity)
        
        f_form = IRSSActivityFileFormSet(request.POST,  request.FILES,instance=irssactivity)
       
        if form.is_valid() and f_form.is_valid():
            form.save()
          
    
            f_form.instance = irssactivity
            f_form.save()
            info(request, _("Successfully updated IRSS Activity."))
            return redirect("irss-activity-detail", pk=irssactivity.id)

    else:
        form = IRSSActivityForm(instance=irssactivity)
       
        f_form = IRSSActivityFileFormSet(instance=irssactivity)
       
      
    return render_to_response(template_name, {
        'form': form, 'f_form':f_form,"irssactivity": irssactivity
    }, context_instance=RequestContext(request))            
            


class UserMembershipHistoryListView(ListView):
    """
    UserMembershipHistory
    
    """
    context_object_name = 'latest'
    model = UserMembershipHistory
    date_field = 'start_date'
    template_name = 'accounts/user_membership_list.html'
    queryset = UserMembershipHistory.objects.all().order_by('-start_date')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only  """
        # self.country = get_object_or_404(CountryPage, country=self.kwargs['country'])
        # CountryPage country_slug == country URL parameter keyword argument
        return UserMembershipHistory.objects.all().order_by('-start_date')
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(UserMembershipHistoryListView, self).get_context_data(**kwargs)
       # context['activity_types'] =IRSS_ACT_TYPE_CHOICES
       
        #context['country'] = self.kwargs['country']
        return context


class UserMembershipHistoryDetailView(DetailView):
    """ Pest report detail page """
    model = UserMembershipHistory
    context_object_name = 'user_membership'
    template_name = 'accounts/user_membership_detail.html'
    queryset = UserMembershipHistory.objects.filter()
    
  
    
@login_required
@permission_required('ippc.add_usermembershiphistory', login_url="/accounts/login/")
def usermembershiphistory_create(request):
    """ Create  usermembershiphistory """
    
    form = UserMembershipHistoryForm(request.POST)
    
    if request.method == "POST":
        if form.is_valid():
            new_usermembershiphistory = form.save(commit=False)
            form.save()
            
            info(request, _("Successfully created user membership entry."))
            
            return redirect("usermembershiphistory-detail", pk=new_usermembershiphistory.id)
        else:
             return render_to_response('accounts/user_membership_create.html', {'form': form,},
             context_instance=RequestContext(request))
    else:
        form = UserMembershipHistoryForm(instance= UserMembershipHistory())
     

    return render_to_response('accounts/user_membership_create.html', {'form': form },
        context_instance=RequestContext(request))

  
@login_required
@permission_required('ippc.change_usermembershiphistory', login_url="/accounts/login/")
def usermembershiphistory_edit(request, id=None, template_name='accounts/user_membership_edit.html'):
    """ Edit  usermembershiphistory """
    if id:
        usermembership = get_object_or_404(UserMembershipHistory,  pk=id)
       #print(id)
       #print(usermembership)
        
    if request.POST:
        form = UserMembershipHistoryForm(request.POST, instance=usermembership)
        
        if form.is_valid():
           form.save()
          
           return redirect("usermembership-detail",id=usermembership.id)

    else:
        form = UserMembershipHistoryForm(instance=usermembership)
       #print(usermembership.user)
    return render_to_response(template_name, {
        'form': form, "usermembership": usermembership,"username":usermembership.user,
    }, context_instance=RequestContext(request))



class MediaKitDocumentListView(ListView):
    """
    MediaKitDocument
    
    """
    context_object_name = 'latest'
    model = MediaKitDocument
    date_field = 'start_date'
    template_name = 'pages/mediakit_list.html'
    queryset = MediaKitDocument.objects.all().order_by('-_order')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only  """
        # self.country = get_object_or_404(CountryPage, country=self.kwargs['country'])
        # CountryPage country_slug == country URL parameter keyword argument
        return MediaKitDocument.objects.all().order_by('-_order')
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(MediaKitDocumentListView, self).get_context_data(**kwargs)
        #context['types'] =MEDIAKIT_TYPE_CHOICES
        context['12col'] = 1
       
        arraymain=[]
        innerarray1=[]
        innerarray1.append("Annual Reports and Strategies")
        innerarray1.append(MediaKitDocument.objects.filter(mediakit_type=1).order_by('-_order'))
        arraymain.append(innerarray1)
        
        innerarray2=[]
        innerarray2.append("Meeting reports")
        innerarray2.append(MediaKitDocument.objects.filter(mediakit_type=7).order_by('-_order'))
        arraymain.append(innerarray2)
        
        innerarray3=[]
        innerarray3.append("Brochures")
        innerarray3.append(MediaKitDocument.objects.filter(mediakit_type=2).order_by('-_order'))
        arraymain.append(innerarray3)
        
        innerarray4=[]
        innerarray4.append("Factsheets")
        innerarray4.append(MediaKitDocument.objects.filter(mediakit_type=3).order_by('-_order'))
        arraymain.append(innerarray4)
        
        innerarray5=[]
        innerarray5.append("Advocacy materials")
        innerarray5.append(MediaKitDocument.objects.filter(mediakit_type=4).order_by('-_order'))
        arraymain.append(innerarray5)
      
        innerarray6=[]
        innerarray6.append("Guides and Training resources")
        innerarray6.append(MediaKitDocument.objects.filter(mediakit_type=6).order_by('-_order'))
        arraymain.append(innerarray6)
      
        innerarray7=[]
        innerarray7.append("Training kits")
        innerarray7.append(MediaKitDocument.objects.filter(mediakit_type=5).order_by('-_order'))
        arraymain.append(innerarray7)
      
        
        
        context['latest'] =arraymain
       
        return context  
  
class PhytosanitaryTreatmentListView(ListView):
    """
    PhytosanitaryTreatment List
    """
    context_object_name = 'latest'
    model = PhytosanitaryTreatment
    date_field = 'modify_date'
    template_name = 'phytotreatment/phytotreatment_list.html'
    queryset = PhytosanitaryTreatment.objects.filter(status=IS_PUBLIC).order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 10

    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PhytosanitaryTreatmentListView, self).get_context_data(**kwargs)
        phyto_treatments_array=[]
        phytos = PhytosanitaryTreatment.objects.filter(status=IS_PUBLIC).order_by('-modify_date', 'title')
        for phyto in phytos:
            phytos_array=[]
            phytos_array.append(phyto)
            pestidentity=''
            otherpests=phyto.treatment_pestidentity_other
            pests =  PhytosanitaryTreatmentPestsIdentity.objects.filter(phytosanitarytreatment=phyto.id)
            if pests.count()>0:
                pestidentity='' 
                for p in pests.all():
                   if p.pestidentitydescr!=None:
                        pestidentity=pestidentity+ p.pestidentitydescr
            pestidentity=pestidentity+otherpests
            
            commodityidentity=''
            othercommodity=phyto.treatment_commodityidentity_other
            commodities =  PhytosanitaryTreatmentCommodityIdentity.objects.filter(phytosanitarytreatment=phyto.id)
            if  commodities.count()>0:
                commodityidentity='' 
                for c in commodities.all():
                    if c.commoditydescr!=None:
                        commodityidentity=commodityidentity+ c.commoditydescr
                

            commodityidentity=commodityidentity+othercommodity
            
            phytos_array.append(pestidentity)
            phytos_array.append(commodityidentity)
            phyto_treatments_array.append(phytos_array)

        context['phyto_treatments_array'] = phyto_treatments_array
        return context
    
class PhytosanitaryTreatmentDetailView(DetailView):
    """ PhytosanitaryTreatment detail page """
    model = PhytosanitaryTreatment
    context_object_name = 'phytotreatment'
    template_name = 'phytotreatment/phytotreatment_detail.html'
    queryset = PhytosanitaryTreatment.objects.filter()
#   
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PhytosanitaryTreatmentDetailView, self).get_context_data(**kwargs)
        
        phyto = get_object_or_404(PhytosanitaryTreatment, slug=self.kwargs['slug'])
        pestidentity=''
        otherpests=phyto.treatment_pestidentity_other
        pests =  PhytosanitaryTreatmentPestsIdentity.objects.filter(phytosanitarytreatment=phyto.id)
        if pests.count()>0:
            pestidentity='' 
            for p in pests.all():
               if p.pestidentitydescr!=None:
                    pestidentity=pestidentity+ p.pestidentitydescr
        pestidentity=pestidentity+otherpests
       
        commodityidentity=''
        othercommodity=phyto.treatment_commodityidentity_other
        commodities =  PhytosanitaryTreatmentCommodityIdentity.objects.filter(phytosanitarytreatment=phyto.id)
        commodities =  PhytosanitaryTreatmentCommodityIdentity.objects.filter(phytosanitarytreatment=phyto.id)
        if  commodities.count()>0:
            commodityidentity='' 
            for c in commodities.all():
                if c.commoditydescr!=None:
                    commodityidentity=commodityidentity+ c.commoditydescr
        commodityidentity=commodityidentity+othercommodity
      
        context['phytotreatment'] = phyto
        context['pestidentity'] = pestidentity
        context['commodityidentity'] = commodityidentity
       
        return context





@login_required
@permission_required('ippc.add_phytosanitarytreatment', login_url="/accounts/login/")
def phytosanitarytreatment_create(request):
    """ Create phytosanitraytreatment """
    user = request.user
    author = user
    form = PhytosanitaryTreatmentForm(request.POST, request.FILES)
  
    if request.method == "POST":
        pestform = PhytosanitaryTreatmentPestsIdentityFormSet(request.POST or None)
        commodityform = PhytosanitaryTreatmentCommodityIdentityFormSet(request.POST or None)
       #print(pestform)
        if form.is_valid() and pestform.is_valid() and commodityform.is_valid():
            new_phytotreatment = form.save(commit=False)
            new_phytotreatment.author = request.user
            new_phytotreatment.author_id = author.id
            
            form.save()
           
            pestform.instance = new_phytotreatment
            pestform.save()
           
            commodityform.instance = new_phytotreatment
            commodityform.save()
    
            
            info(request, _("Successfully created Phytosanitary Treatment."))
            return redirect("phytosanitary-treatment-detail",  slug=new_phytotreatment.slug)
        else:
             return render_to_response('phytotreatment/phytotreatment_create.html', {'form': form,'pestform': pestform,'commodityform':commodityform,},
             context_instance=RequestContext(request))
       
    else:
        form = PhytosanitaryTreatmentForm(instance=PhytosanitaryTreatment())
        pestform = PhytosanitaryTreatmentPestsIdentityFormSet()
        commodityform = PhytosanitaryTreatmentCommodityIdentityFormSet()
     
      
    return render_to_response('phytotreatment/phytotreatment_create.html', {'form': form,'pestform': pestform,'commodityform':commodityform,},
        context_instance=RequestContext(request))
		


@login_required
@permission_required('ippc.change_phytosanitarytreatment', login_url="/accounts/login/")
def phytosanitarytreatment_edit(request,  id=None, template_name='phytotreatment/phytotreatment_edit.html'):
    """ Edit phytotreatment """
    user = request.user
    author = user
    if id:
        phytotreatment = get_object_or_404(PhytosanitaryTreatment,  pk=id)
    else:
        phytotreatment = PhytosanitaryTreatment(author=request.user)
        

    if request.POST:
        form = PhytosanitaryTreatmentForm(request.POST,  request.FILES, instance=phytotreatment)
        pestform = PhytosanitaryTreatmentPestsIdentityFormSet(request.POST, instance=phytotreatment)
        commodityform = PhytosanitaryTreatmentCommodityIdentityFormSet(request.POST, instance=phytotreatment)
     
        if form.is_valid() and pestform.is_valid() and commodityform.is_valid():
            form.save()
            
            pestform.instance = phytotreatment
            pestform.save()
           
            commodityform.instance = phytotreatment
            commodityform.save()
          
            # If the save was successful, success message and redirect to another page
            info(request, _("Successfully updated phytosanitary treatment."))
            return redirect("phytosanitary-treatment-detail", slug=phytotreatment.slug)
                             
    else:
        form = PhytosanitaryTreatmentForm(instance=phytotreatment)
        pestform = PhytosanitaryTreatmentPestsIdentityFormSet( instance=phytotreatment)
        commodityform = PhytosanitaryTreatmentCommodityIdentityFormSet(instance=phytotreatment)
     
       
    return render_to_response(template_name, {
        'form': form, "phytotreatment": phytotreatment,"pestform":pestform,"commodityform":commodityform
    }, context_instance=RequestContext(request))
		
    
    
@login_required
def subscribe_to_news(request,type=None):
    """ subscribe to news """
    #previous_page = request.get('return_url')
    previous_page = request.GET.get('return_url', None)
   #print(previous_page)
    
    #data = {'previous_page': previous_page,}
    user = request.user
    g1=None
    c=0
    typename=''
    
    if type == '1':
        g1=Group.objects.get(name="News Notification group")
        c=user.groups.filter(name="News Notification group").count()
        typename='News'
    elif type == '3':
        print('dsadadas')
        g1=Group.objects.get(name="Announcement Notification group")
        
        c=user.groups.filter(name="Announcement Notification group").count()
        typename='Announcements'
    elif type == '5':
        g1=Group.objects.get(name="Calls Notification group")
        c=user.groups.filter(name="Calls Notification group").count()
        typename='Calls'
    if c==0:
        user.groups.add(g1)
        info(request, _("Successfully Subcribed to "+str(typename)+"."))
    else:
        error(request, _("You are already subcribed to "+str(typename)+"."))
    

    return redirect('https://www.ippc.int/'+str(previous_page)  )
   # return render(request, "news/news_post_list.html", data)
   # return render_to_response(context_instance=RequestContext(request)) 
   

@login_required
def unsubscribe_to_news(request,type=None):
    """ subscribe to news """
    #previous_page = request.get('return_url')
    previous_page = request.GET.get('return_url', None)
   #print(previous_page)
    
    #data = {'previous_page': previous_page,}
    user = request.user
    g1=None
    c=0
    typename=''
    
    if type ==  '1':
        g1=Group.objects.get(name="News Notification group")
        c=user.groups.filter(name="News Notification group").count()
        typename='News'
    elif type == '3':
        g1=Group.objects.get(name="Announcement Notification group")
        c=user.groups.filter(name="Announcement Notification group").count()
        typename='Announcements'
    elif type == '5':
        g1=Group.objects.get(name="Calls Notification group")
        c=user.groups.filter(name="Calls Notification group").count()
        typename='Calls'
    if c>0:
        user.groups.remove(g1)
        info(request, _("Successfully UN-Subcribed from "+str(typename)+"."))
    else:
        error(request, _("You are already UN_Subcribed from "+str(typename)+"."))
    

    return redirect('https://www.ippc.int/'+str(previous_page)  )
    ##return redirect('https://www.ippc.int/'+str(previous_page)  )
   # return render(request, "news/news_post_list.html", data)
   # return render_to_response(context_instance=RequestContext(request)) 
   

def send_notificationevent_message(id):
    """ send_notification_message """
    event = get_object_or_404(Event, id=id)

    emailto_all = ['']
    participant_already_registered=[]
    participant_not_registered=[] 
   
    eventParticipants=EventParticipants.objects.filter(event_id=id)
    for u in eventParticipants:
        if u.registered:
            participant_already_registered.append(u.user)
        else:
            participant_not_registered.append(u.user)
    for g in event.groups.all():
       #print(g)
        group=Group.objects.get(id=g.id)
        users = group.user_set.all()
        for u in users:
           if u in participant_already_registered:
               print("already registered")
           else:    
                participant_not_registered.append(u)

    for u in participant_not_registered:
         user_registered_obj=User.objects.get(username=u)
         email=user_registered_obj.email
         emailto_all.append(email)

    

    subject='Register to IPPC Meeting: '+event.title
    itemllink="https://www.ippc.int/en/events/event/"+str(id)
    textmessage =textmessage ='<table bgcolor="#FFFFFF" cellspacing="2" cellpadding="2" valign="top" width="100%" style="border-bottom: 1px solid #10501F;border-top: 1px solid #10501F;border-left: 1px solid #10501F;border-right: 1px solid #10501F"> <tr><td width="100%" bgcolor="#FFFFFF">Dear Sir/Madam, <br><br>Please be informed that the following meeting:<br><br><a href="'+itemllink+'">'+event.title+'</a> ('+itemllink+') <br><br>is now open for registration on the <b>International Phytosanitary Portal</b>.<br><br>Please LOGIN to IPPC and <b><a href="'+itemllink+'">REGISTER</a></b> before <b>'+str(event.end_register_date)+'</b><br><br>Please check the data in your profile, you can edit your profile yourself <a href="https://www.ippc.int/en/accounts/login/">after login into the IPP</a> (you can change all contact details except your name and job title). <br><br><p>-- International Plant Protection Convention team </p></td></tr></table>'

    message = mail.EmailMessage(subject,textmessage,'ippc@fao.org',#from
        emailto_all, ['paola.sentinelli@fao.org'])#emailto_all for PROD, in TEST all to paola#
   # print(textmessage)
    message.content_subtype = "html"
    if event.can_register:
        #print('test-sending')
        sent =message.send()
        
def send_notification_register_event_message(register,id,user):
    """ send_notification_register_event_message """
    event = get_object_or_404(Event, id=id)
    emailto_all = ['']
         
    user_registered_obj=User.objects.get(username=user)
    user_registered_email=user_registered_obj.email
    user_registered_first_name=user_registered_obj.first_name
    user_registered_last_name=user_registered_obj.last_name
    
    author_id=event.creator_id
    auth_event_obj=User.objects.get(id=author_id)
    auth_email=auth_event_obj.email
    emailto_all.append(str(auth_email))
    
    registerlabel=''
    if register:
        registerlabel='REGISTERED'
    else:
        registerlabel='UN-REGISTERED'
        
    subject='User has '+registerlabel+' to IPPC Meeting: '+event.title
    itemllink="https://www.ippc.int/en/events/event/"+str(id)
    textmessage ='<table bgcolor="#FFFFFF" cellspacing="2" cellpadding="2" valign="top" width="100%" style="border-bottom: 1px solid #10501F;border-top: 1px solid #10501F;border-left: 1px solid #10501F;border-right: 1px solid #10501F"> <tr><td width="100%" bgcolor="#FFFFFF">Dear IPPC Secretariat, <br><br>Please be informed that the following user <b>'+user_registered_first_name+' '+user_registered_last_name+' ('+user_registered_email+')</b> has <b>'+registerlabel+'</b> to the meeting: <a href="'+itemllink+'">'+event.title+'</a> ('+itemllink+').<br><br><p>-- International Plant Protection Convention team </p></td></tr></table>'

    message = mail.EmailMessage(subject,textmessage,'ippc@fao.org',#from
        emailto_all, ['paola.sentinelli@fao.org'])#emailto_all for PROD, in TEST all to paola#
    #print(textmessage)
    message.content_subtype = "html"
    #print('test-sending')
    sent =message.send()
    
    
from PIL import Image
from PIL import ImageFont
from PIL import ImageDraw
from PIL import ImageFile


import os

class CertificatesToolListView(ListView):
    """    CertificatesTool List view """
    context_object_name = 'latest'
    model = CertificatesTool
    date_field = 'date'
    template_name = 'certificates/certificatestool_list.html'
    queryset = CertificatesTool.objects.all().order_by('-date')
   
       
class CertificatesToolDetailView(DetailView):
    """ CertificatesTool detail page """
    model = CertificatesTool
    context_object_name = 'certificatestool'
    template_name = 'certificates/certificatestool_detail.html'
    queryset = CertificatesTool.objects.filter()

import shutil


@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def generate_certificatesnew(request):
    """ Create generate_certificates """
    form = CertificatesToolForm(request.POST)
    certificatedir_template = MEDIA_ROOT+'/certificate_template'
    template_path = MEDIA_ROOT+'/certificate_template/TemplateCert2017_XIA_BRENT.docx'
    
    g_set=[]
    for g in Group.objects.filter():
        users = g.user_set.all()
        users_all=[]
        users_all.append(str(g))
        users_all.append(str(g.id))
        for u in users:
           users_u=[]
           user_obj=User.objects.get(username=u)
           if user_obj.is_active:
            userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
            users_u.append((unicode(userippc.first_name)))
            users_u.append((unicode(userippc.last_name)))
            users_u.append((user_obj.id))
            users_all.append(users_u)
        g_set.append(users_all)
        

    zip_all=''
    zip_all_s=''
    users_participants=[]       
    users_address=[]
    if request.method == "POST":
        if form.is_valid():
            topicnumber = str(request.POST['topicnumber'])
            topic  = get_object_or_404(Topic, id= topicnumber)
            
            usersfrom_g=[]
           #print(request.POST)
            for g in Group.objects.filter():
                for u_id in request.POST.getlist('user_'+str(g.id)+'_0'):
                    if u_id not in usersfrom_g:
                        usersfrom_g.append(str(u_id))
          
             
            new_certificatestool = form.save(commit=False)
            new_certificatestool.creation_date = timezone.now()
            new_certificatestool.author =  request.user
            
            date = timezone.now().strftime('%Y%m%d%H%M%S')
            certificatedir_new = MEDIA_ROOT+'files/certificates/'+date
            zip_all1 ="/static/media/files/certificates/certificates_"+ date+".zip"
            zip_all = zipfile.ZipFile(MEDIA_ROOT+"/files/certificates/"+"/certificates_"+ date+".zip", "w")
            new_certificatestool.filezip =zip_all1
            
            form.save()
            try: 
                os.makedirs(certificatedir_new)
            except OSError:
                if not os.path.isdir(certificatedir_new):
                    raise
                
            if topic!=None:
                topictitle=  request.POST['title']
               
                events = Event.objects.filter(topic_numbers=topic)
                for ev in events:
                    eventParticipants=EventParticipants.objects.filter(event_id=ev.id)
                    for u in eventParticipants:
                        if u not in users_participants:
                            users_participants.append(u)
                
                for u in users_participants:
                    user_obj=User.objects.get(username=u.user)
                    usersfrom_g.append(str(user_obj.id))
                    userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                    name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
                    role=u.role
                    roletoprint=''
                    for r in u.role.all():
                        role_=str(r)
                        roletoprint=roletoprint+role_

                    doc_title="cert_"+userippc.last_name+".docx"
                    document = Document(template_path)
                
                    obj_styles = document.styles
                    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
                    obj_font = obj_charstyle.font
                    obj_font.size = Pt(18)
                    obj_font.color.rgb = RGBColor(0,119,67)
                    obj_font.name = 'Times New Roman'

                    obj_styles1 = document.styles
                    obj_charstyle1 = obj_styles1.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
                    obj_font1 = obj_charstyle1.font
                    obj_font1.size = Pt(11)
                    obj_font1.name = 'Times New Roman'
                    p= document.add_paragraph("")

                    p.add_run("\n\n\n\n\n\n"+name, style = 'CommentsStyle').bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p= document.add_paragraph("")

                    p.add_run("\n"+roletoprint, style = 'CommentsStyle').bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p= document.add_paragraph("")

                    p.add_run("\n"+topictitle, style = 'CommentsStyle').bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p= document.add_paragraph("")

                    p.add_run("\n\n\n\n\n\n\n\n\n\n\nRome, Italy, "+timezone.now().strftime('%Y'), style = 'CommentsStyle1').bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    #p= document.add_paragraph("")

                    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = 'attachment; filename=' + doc_title
                    strfpath1=os.path.join(certificatedir_new,doc_title+'.docx')
                    document.save(strfpath1)
                    zip_all.write(strfpath1, doc_title)
                  
            
          
            document = Document()
             
            docx_title="Address_labels.docx"
            strfpath1=os.path.join(certificatedir_new,docx_title+'.jpg')
            sections = document.sections
            for section in sections:
                section.top_margin = Cm(0.01)
                section.bottom_margin = Cm(0.01)
                section.left_margin = Cm(0.01)
                section.right_margin = Cm(0.01)

            style = document.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(12)

            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(10)
            obj_font.name = 'Times New Roman'
          
            singleletter=[]
            doubleletter=[]
            for user in usersfrom_g:
               if user in singleletter:
                   doubleletter.append(user)
               else:
                   singleletter.append(user)
            
            users_split = split(singleletter,12)
            userdone=[]
            
            for users in users_split:
                table = document.add_table(rows=1, cols=2)
                table.style = document.styles['Table Grid']
              
                font.size = Pt(9)

                h=0
                row_cells=None
                for u in users:
                   
                        
                        user_obj=User.objects.get(id=u)
                        userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                      
                        
                        name_title_address=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
                       #print(name_title_address)
                        address2f=''
                        splitaddress=userippc.address2.splitlines()
                        for s in splitaddress:
                            if s!='':
                                address2f+='\r'+s  
                        if userippc.address1!='':
                            name_title_address=name_title_address+'\r'+userippc.address1
                        if address2f!='':
                           name_title_address=name_title_address+address2f
                        if u in doubleletter:
                           name_title_address=name_title_address+'\r\r\t\t\t\t\t\t\t*'
                       
                            
                        if  h==0 or h==1:
                            hdr_cells = table.rows[0].cells
                            hdr_cells[h].text =name_title_address
                            hdr_cells[h].text =name_title_address

                        if  h>1 and (h % 2) == 0:
                           row_cells = table.add_row().cells

                        if h>1 :
                            if (h % 2) == 0:
                                run = row_cells[0].paragraphs[0].add_run(name_title_address) 
                                row_cells[0].style = "borderColor:red;background-color:gray"

                            else :
                                run = row_cells[1].paragraphs[0].add_run(name_title_address) 

                        h+=1    
                        userdone.append(u)
                
                range1=range(0,2)
                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "2500")
                    trHeight.set(qn('w:hRule'), "atLeast")
                    trPr.append(trHeight)
                    for i in range1 :
                        tc = row.cells[i]._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcVAlign = OxmlElement('w:vAlign')
                        tcVAlign.set(qn('w:val'), "center")
                        tcPr.append(tcVAlign)
           
            set_column_width(table.columns[0], Cm(10.6))
            set_column_width(table.columns[1], Cm(10.6))
             
            document.add_page_break()
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=' + docx_title
            document.save(strfpath1)
            zip_all.write(strfpath1, docx_title)
    
            zip_all.close()
            zip_all=zip_all1
            shutil.rmtree(certificatedir_new)
            
            info(request, _("Certificates generated."))
            return redirect("certificatestool-detail",new_certificatestool.id)
        else:
             return render_to_response('certificates/certificatestool_create.html', {'form': form,'emailgroups':g_set,'img':"TemplateCertificate2017_XiaBrent_example.jpg",'zip_all':zip_all,},
             context_instance=RequestContext(request))
    else:
        form = CertificatesToolForm(instance=CertificatesTool())
        
      
    return render_to_response('certificates/certificatestool_create.html', {'form': form,'emailgroups':g_set,'zip_all':zip_all,'img':"TemplateCertificate2017_XiaBrent_example.jpg" ,},
        context_instance=RequestContext(request))
        
        

@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def generate_certificates(request):
    """ Create generate_certificates """
    form = CertificatesToolForm(request.POST)
    certificatedir_template = MEDIA_ROOT+'/certificate_template'
    img_template=MEDIA_ROOT+'certificate_template/TemplateCertificate2017_XiaBrent.png'
    fontfile =MEDIA_ROOT+"certificate_template/times.ttf"
  
    g_set=[]
    for g in Group.objects.filter():
        users = g.user_set.all()
        users_all=[]
        users_all.append(str(g))
        users_all.append(str(g.id))
        for u in users:
           users_u=[]
           user_obj=User.objects.get(username=u)
           if user_obj.is_active:
            userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
            users_u.append((unicode(userippc.first_name)))
            users_u.append((unicode(userippc.last_name)))
            users_u.append((user_obj.id))
            users_all.append(users_u)
        g_set.append(users_all)
        

    zip_all=''
    zip_all_s=''
    users_participants=[]       
    users_address=[]
    if request.method == "POST":
        if form.is_valid():
            topicnumber = str(request.POST['topicnumber'])
            topic  = get_object_or_404(Topic, id= topicnumber)
            
            usersfrom_g=[]
           #print(request.POST)
            for g in Group.objects.filter():
                for u_id in request.POST.getlist('user_'+str(g.id)+'_0'):
                    if u_id not in usersfrom_g:
                        usersfrom_g.append(str(u_id))
          
             
            new_certificatestool = form.save(commit=False)
            new_certificatestool.creation_date = timezone.now()
            new_certificatestool.author =  request.user
            
            date = timezone.now().strftime('%Y%m%d%H%M%S')
            certificatedir_new = MEDIA_ROOT+'files/certificates/'+date
            zip_all1 ="/static/media/files/certificates/certificates_"+ date+".zip"
            zip_all = zipfile.ZipFile(MEDIA_ROOT+"/files/certificates/"+"/certificates_"+ date+".zip", "w")
            new_certificatestool.filezip =zip_all1
            
            form.save()
            try: 
                os.makedirs(certificatedir_new)
            except OSError:
                if not os.path.isdir(certificatedir_new):
                    raise
                
            if topic!=None:
                topictitle=  request.POST['title']
               
                events = Event.objects.filter(topic_numbers=topic)
                for ev in events:
                    eventParticipants=EventParticipants.objects.filter(event_id=ev.id)
                    for u in eventParticipants:
                        if u not in users_participants:
                            users_participants.append(u)
                
                for u in users_participants:
                  user_obj=User.objects.get(username=u.user)
                  usersfrom_g.append(str(user_obj.id))
                  userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                  name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
                  role=u.role
                  roletoprint=''
                  for r in u.role.all():
                        role_=str(r)
                       #print(role_)
                        roletoprint=roletoprint+role_
                 

                  
                  lenghtofname=len(name)
                  xcoord_name=1760-(lenghtofname/2)*50

                  lenghtofrole=len(roletoprint)
                  xcoordrole=1760-(lenghtofrole/2)*50

                  lenghtoftopic=len(topictitle)
                  xcoordtopic=1760-(lenghtoftopic/2)*50

                  font = ImageFont.truetype(fontfile, 100) #Arial.ttf", 48)
                  font1= ImageFont.truetype(fontfile, 50) #Arial.ttf", 48)

                  img = Image.open(img_template)
                  draw = ImageDraw.Draw(img)
                  draw.text((xcoord_name, 929),name,(0,119,67),font=font)
                  draw.text((xcoordrole, 1165),roletoprint,(0,119,67),font=font)
                  draw.text((xcoordtopic, 1400),topictitle,(0,119,67),font=font)
                  draw.text((1605, 2134),"Rome, Italy, "+timezone.now().strftime('%Y')+" ",(0,0,0),font=font1)
    #                  
                  img.save(os.path.join(certificatedir_new,str(u.user)+'.png'))
                  strfpath=os.path.join(certificatedir_new,str(u.user)+'.png')
                  zip_all.write(strfpath, str(u.user)+'.png')

    
          
            document = Document()
             
            docx_title="Address_labels.docx"
            strfpath1=os.path.join(certificatedir_new,docx_title+'.jpg')
            sections = document.sections
            for section in sections:
                section.top_margin = Cm(0.01)
                section.bottom_margin = Cm(0.01)
                section.left_margin = Cm(0.01)
                section.right_margin = Cm(0.01)

            style = document.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(12)

            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(10)
            obj_font.name = 'Times New Roman'
          
            singleletter=[]
            doubleletter=[]
            for user in usersfrom_g:
               if user in singleletter:
                   doubleletter.append(user)
               else:
                   singleletter.append(user)
            
            users_split = split(singleletter,12)
            userdone=[]
            
            for users in users_split:
               #print("-------")
               #print(users)
                #p= document.add_paragraph("")
                table = document.add_table(rows=1, cols=2)
                table.style = document.styles['Table Grid']
              
                font.size = Pt(9)

                h=0
                row_cells=None
                for u in users:
                   
                        
                        user_obj=User.objects.get(id=u)
                        userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                      
                        
                        name_title_address=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
                       #print(name_title_address)
                        address2f=''
                        splitaddress=userippc.address2.splitlines()
                        for s in splitaddress:
                            if s!='':
                                address2f+='\r'+s  
                        if userippc.address1!='':
                            name_title_address=name_title_address+'\r'+userippc.address1
                        if address2f!='':
                           name_title_address=name_title_address+address2f
                        if u in doubleletter:
                           name_title_address=name_title_address+'\r\r\t\t\t\t\t\t\t*'
                       
                            
                        if  h==0 or h==1:
                            hdr_cells = table.rows[0].cells
                            hdr_cells[h].text =name_title_address
                            hdr_cells[h].text =name_title_address

                        if  h>1 and (h % 2) == 0:
                           row_cells = table.add_row().cells

                        if h>1 :
                            if (h % 2) == 0:
                                run = row_cells[0].paragraphs[0].add_run(name_title_address) 
                                row_cells[0].style = "borderColor:red;background-color:gray"

                            else :
                                run = row_cells[1].paragraphs[0].add_run(name_title_address) 

                        h+=1    
                        userdone.append(u)
                
                range1=range(0,2)
                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "2500")
                    trHeight.set(qn('w:hRule'), "atLeast")
                    trPr.append(trHeight)
                    for i in range1 :
                        tc = row.cells[i]._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcVAlign = OxmlElement('w:vAlign')
                        tcVAlign.set(qn('w:val'), "center")
                        tcPr.append(tcVAlign)
           
            set_column_width(table.columns[0], Cm(10.6))
            set_column_width(table.columns[1], Cm(10.6))
             
            document.add_page_break()
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=' + docx_title
            document.save(strfpath1)
            zip_all.write(strfpath1, docx_title)
    
            zip_all.close()
            zip_all=zip_all1
            shutil.rmtree(certificatedir_new)
            
            info(request, _("Certificates generated."))
            return redirect("certificatestool-detail",new_certificatestool.id)
        else:
             return render_to_response('certificates/certificatestool_create.html', {'form': form,'fontfile':img_template,'emailgroups':g_set,'img':"TemplateCertificate2017_XiaBrent_example.jpg",'zip_all':zip_all,},
             context_instance=RequestContext(request))
    else:
        form = CertificatesToolForm(instance=CertificatesTool())
        
      
    return render_to_response('certificates/certificatestool_create.html', {'form': form,'fontfile':img_template,'emailgroups':g_set,'zip_all':zip_all,'img':"TemplateCertificate2017_XiaBrent_example.jpg" ,},
        context_instance=RequestContext(request))

class B_CertificatesToolListView(ListView):
    """   B_CertificatesTool List view """
    context_object_name = 'latest'
    model = B_CertificatesTool
    date_field = 'date'
    template_name = 'certificates/b_certificatestool_list.html'
    queryset = B_CertificatesTool.objects.all().order_by('-date')
   
       
class B_CertificatesToolDetailView(DetailView):
    """ B_CertificatesTool detail page """
    model = B_CertificatesTool
    context_object_name = 'certificatestool'
    template_name = 'certificates/b_certificatestool_detail.html'
    queryset = B_CertificatesTool.objects.filter()


@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def generate_b_certificatesnew(request):
    """ Create generate_certificates """
    form =B_CertificatesToolForm(request.POST)
    img_template_example="TemplateCertificate2017_Xia_example.jpg"
    template_path = MEDIA_ROOT+'certificate_template/TemplateCert2017_XIA.docx'
             
    g_set=[]
    for g in Group.objects.filter():
        users = g.user_set.all()
        users_all=[]
        users_all.append(str(g))
        users_all.append(str(g.id))
        for u in users:
           users_u=[]
           user_obj=User.objects.get(username=u)
           if user_obj.is_active:
            userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
            users_u.append((unicode(userippc.first_name)))
            users_u.append((unicode(userippc.last_name)))
            users_u.append((user_obj.id))
            users_all.append(users_u)
        g_set.append(users_all)
        

    zip_all=''
    zip_all_s=''
    users_participants=[]       
    users_address=[]
    if request.method == "POST":
        if form.is_valid():
            committee = str(request.POST['text3'])
            role = str(request.POST['role'])
            username = str(request.POST['user_name'])
            
            usersfrom_g=[]
            for g in Group.objects.filter():
                for u_id in request.POST.getlist('user_'+str(g.id)+'_0'):
                    usersfrom_g.append(str(u_id))
                    
            new_certificatestool = form.save(commit=False)
            new_certificatestool.creation_date = timezone.now()
            new_certificatestool.author =  request.user
            
            date = timezone.now().strftime('%Y%m%d%H%M%S')
            certificatedir_new = MEDIA_ROOT+'files/b_certificates/'+date
            zip_all1 ="/static/media/files/b_certificates/certificates_"+ date+".zip"
            zip_all = zipfile.ZipFile(MEDIA_ROOT+"/files/b_certificates/certificates_"+ date+".zip", "w")
            new_certificatestool.filezip =zip_all1
            form.save()
            try: 
                os.makedirs(certificatedir_new)
            except OSError:
                if not os.path.isdir(certificatedir_new):
                    raise
                
            if username!='' :
                name= username
                roletoprint=role
                doc_title="cert_"+username+".docx"
             
                document = Document(template_path)
                
                obj_styles = document.styles
                obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
                obj_font = obj_charstyle.font
                obj_font.size = Pt(18)
                obj_font.color.rgb = RGBColor(0,119,67)
                obj_font.name = 'Times New Roman'
                
                obj_styles1 = document.styles
                obj_charstyle1 = obj_styles1.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
                obj_font1 = obj_charstyle1.font
                obj_font1.size = Pt(11)
                obj_font1.name = 'Times New Roman'
                p= document.add_paragraph("")
                  
                p.add_run("\n\n\n\n\n\n"+name, style = 'CommentsStyle').bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p= document.add_paragraph("")
             
               
                p.add_run("\n"+roletoprint, style = 'CommentsStyle').bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p= document.add_paragraph("")
               
                p.add_run("\n"+committee, style = 'CommentsStyle').bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p= document.add_paragraph("")
                
                p.add_run("\n\n\n\n\n\n\n\n\n\n\nRome, Italy, "+timezone.now().strftime('%Y'), style = 'CommentsStyle1').bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #p= document.add_paragraph("")
            
                response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = 'attachment; filename=' + doc_title
                strfpath1=os.path.join(certificatedir_new,doc_title+'.docx')
                document.save(strfpath1)
                zip_all.write(strfpath1, doc_title)
        
            else:
                for u in usersfrom_g:
                    user_obj=User.objects.get(id=u)
                    userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                    
                    name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
                    roletoprint=role
                    doc_title="cert_"+userippc.last_name+".docx"
                    document = Document(template_path)
                
                    obj_styles = document.styles
                    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
                    obj_font = obj_charstyle.font
                    obj_font.size = Pt(18)
                    obj_font.color.rgb = RGBColor(0,119,67)
                    obj_font.name = 'Times New Roman'

                    obj_styles1 = document.styles
                    obj_charstyle1 = obj_styles1.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
                    obj_font1 = obj_charstyle1.font
                    obj_font1.size = Pt(11)
                    obj_font1.name = 'Times New Roman'
                    p= document.add_paragraph("")

                    p.add_run("\n\n\n\n\n\n"+name, style = 'CommentsStyle').bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p= document.add_paragraph("")

                    p.add_run("\n"+roletoprint, style = 'CommentsStyle').bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p= document.add_paragraph("")

                    p.add_run("\n"+committee, style = 'CommentsStyle').bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p= document.add_paragraph("")

                    p.add_run("\n\n\n\n\n\n\n\n\n\n\nRome, Italy, "+timezone.now().strftime('%Y'), style = 'CommentsStyle1').bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    #p= document.add_paragraph("")

                    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = 'attachment; filename=' + doc_title
                    strfpath1=os.path.join(certificatedir_new,doc_title+'.docx')
                    document.save(strfpath1)
                    zip_all.write(strfpath1, doc_title)
    

            zip_all.close()
            zip_all=zip_all1
            shutil.rmtree(certificatedir_new)
             
            info(request, _("Certificates generated."))
            return redirect("b-certificatestool-detail",new_certificatestool.id)
        else:
             return render_to_response('certificates/b_certificatestool_create.html', {'form': form,'emailgroups':g_set,'img':img_template_example,'zip_all':zip_all,},
             context_instance=RequestContext(request))
    else:
        form = B_CertificatesToolForm(instance=CertificatesTool())
        
      
    return render_to_response('certificates/b_certificatestool_create.html', {'form': form,'emailgroups':g_set,'zip_all':zip_all,'img':img_template_example ,},
        context_instance=RequestContext(request))

   

@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def generate_b_certificates(request):
    """ Create generate_certificates """
    form =B_CertificatesToolForm(request.POST)
    img_template_example="TemplateCertificate2017_Xia_example.jpg"
    certificatedir_template = MEDIA_ROOT+'/certificate_template'
    img_template=MEDIA_ROOT+'certificate_template/TemplateCertificate2017_Xia.png'
    fontfile =MEDIA_ROOT+"certificate_template/times.ttf"
  
    g_set=[]
    for g in Group.objects.filter():
        users = g.user_set.all()
        users_all=[]
        users_all.append(str(g))
        users_all.append(str(g.id))
        for u in users:
           users_u=[]
           user_obj=User.objects.get(username=u)
           if user_obj.is_active:
            userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
            users_u.append((unicode(userippc.first_name)))
            users_u.append((unicode(userippc.last_name)))
            users_u.append((user_obj.id))
            users_all.append(users_u)
        g_set.append(users_all)
        

    zip_all=''
    zip_all_s=''
    users_participants=[]       
    users_address=[]
    if request.method == "POST":
        if form.is_valid():
            committee = str(request.POST['text3'])
            role = str(request.POST['role'])
            username = str(request.POST['user_name'])
            
            usersfrom_g=[]
            for g in Group.objects.filter():
                for u_id in request.POST.getlist('user_'+str(g.id)+'_0'):
                    usersfrom_g.append(str(u_id))
                    
            new_certificatestool = form.save(commit=False)
            new_certificatestool.creation_date = timezone.now()
            new_certificatestool.author =  request.user
            
            date = timezone.now().strftime('%Y%m%d%H%M%S')
            certificatedir_new = MEDIA_ROOT+'files/b_certificates/'+date
            zip_all1 ="/static/media/files/b_certificates/certificates_"+ date+".zip"
            zip_all = zipfile.ZipFile(MEDIA_ROOT+"/files/b_certificates/certificates_"+ date+".zip", "w")
            
       
            new_certificatestool.filezip =zip_all1
            
            form.save()
            
            try: 
                os.makedirs(certificatedir_new)
            except OSError:
                if not os.path.isdir(certificatedir_new):
                    raise
                
            if username!='' :
                name= username
                roletoprint=role
                
                
                lenghtofname=len(name)
                xcoord_name=1760-(lenghtofname/2)*50

                lenghtofrole=len(roletoprint)
                xcoordrole=1760-(lenghtofrole/2)*50

                lenghtofcommittee=len(committee)
                xcoordcommittee=1760-(lenghtofcommittee/2)*50


                #font = ImageFont.truetype(fontfile, 100) #Arial.ttf", 48)
                #font1= ImageFont.truetype(fontfile, 50) #Arial.ttf", 48)
                font= ImageFont.load_default()
                img = Image.open(img_template)
                draw = ImageDraw.Draw(img)
                font.setsize=100
                # font = ImageFont.truetype(<font-file>, <font-size>)
                # draw.text((x, y),"Sample Text",(r,g,b))
                
                draw.text((xcoord_name, 929),name,(0,119,67),font=font)
                draw.text((xcoordrole, 1165),roletoprint,(0,119,67))#,font=font)
                draw.text((xcoordcommittee, 1400),committee,(0,119,67))#),font=font)
                draw.text((1605, 2134),"Rome, Italy, "+timezone.now().strftime('%Y')+" ",(0,0,0))#,font=font1)


                img.save(os.path.join(certificatedir_new,str(slugify(name))+'.png'))
                strfpath=os.path.join(certificatedir_new,str(slugify(name))+'.png')
                zip_all.write(strfpath, str(slugify(name))+'.png')
                
            else:
                for u in usersfrom_g:
                    user_obj=User.objects.get(id=u)
                    userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                    name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
                    roletoprint=role

                    lenghtofname=len(name)
                    xcoord_name=1760-(lenghtofname/2)*50

                    lenghtofrole=len(roletoprint)
                    xcoordrole=1760-(lenghtofrole/2)*50

                    lenghtofcommittee=len(committee)
                    xcoordcommittee=1760-(lenghtofcommittee/2)*50


                    font = ImageFont.truetype(fontfile, 100) #Arial.ttf", 48)
                    font1= ImageFont.truetype(fontfile, 50) #Arial.ttf", 48)

                    img = Image.open(img_template)
                    draw = ImageDraw.Draw(img)
                    # font = ImageFont.truetype(<font-file>, <font-size>)
                    # draw.text((x, y),"Sample Text",(r,g,b))
                    draw.text((xcoord_name, 929),name,(0,119,67),font=font)
                    draw.text((xcoordrole, 1165),roletoprint,(0,119,67),font=font)
                    draw.text((xcoordcommittee, 1400),committee,(0,119,67),font=font)
                    draw.text((1605, 2134),"Rome, Italy, "+timezone.now().strftime('%Y')+" ",(0,0,0),font=font1)

                  #  draw.text((3210, 4250),"Rome, Italy, "+timezone.now().strftime('%Y')+" ",(0,0,0),font=font1)
                    img.save(os.path.join(certificatedir_new,str(user_obj.username)+'.png'))
                    strfpath=os.path.join(certificatedir_new,str(user_obj.username)+'.png')
                    zip_all.write(strfpath, str(user_obj.username)+'.png')

            zip_all.close()
            zip_all=zip_all1
            shutil.rmtree(certificatedir_new)
             
            info(request, _("Certificates generated."))
            return redirect("b-certificatestool-detail",new_certificatestool.id)
        else:
             return render_to_response('certificates/b_certificatestool_create.html', {'form': form,'emailgroups':g_set,'img':img_template_example,'zip_all':zip_all,},
             context_instance=RequestContext(request))
    else:
        form = B_CertificatesToolForm(instance=CertificatesTool())
        
      
    return render_to_response('certificates/b_certificatestool_create.html', {'form': form,'emailgroups':g_set,'zip_all':zip_all,'img':img_template_example ,},
        context_instance=RequestContext(request))

class WorkshopCertificatesToolListView(ListView):
    """    WorkshopCertificatesTool List view """
    context_object_name = 'latest'
    model = WorkshopCertificatesTool
    date_field = 'creation_date'
    template_name = 'certificates/w_certificatestool_list.html'
    queryset = WorkshopCertificatesTool.objects.all().order_by('-creation_date')
   
       
class WorkshopCertificatesToolDetailView(DetailView):
    """ WorkshopCertificatesTool detail page """
    model = WorkshopCertificatesTool
    context_object_name = 'workshopcertificatestool'
    template_name = 'certificates/w_certificatestool_detail.html'
    queryset = WorkshopCertificatesTool.objects.filter()

import shutil


@login_required
@permission_required('ippc.certificatestool', login_url="/accounts/login/")
def generate_workshopcertificates(request):
    """ Create generate_certificates """
    form = WorkshopCertificatesToolForm(request.POST)
    certificatedir_template = MEDIA_ROOT+'\\certificate_template' 
    fontfile =MEDIA_ROOT+"\\static\\media\\certificate_template\\times.ttf"
  
    imgA=os.path.join(certificatedir_template,"template_a.jpg")
    imgB=os.path.join(certificatedir_template,"template_b.jpg")
   
    events_set= Event.objects.filter()
     

    zip_all=''
    zip_all_s=''
    users_selected=[]       
    if request.method == "POST":
        if form.is_valid():
            workshoptitle = str(request.POST['title'])
            eventid = Event.objects.filter(id=request.POST['id_event'])
            eventParticipants=EventParticipants.objects.filter(event_id=eventid)
    
            imgid =request.POST['id_template']
            #img_template="TemplateCertificate2017_Xia.jpg"#template_"+str(imgid)+".jpg"
            img_template="template_"+str(imgid)+".jpg"
            
            new_wcertificatestool = form.save(commit=False)
            new_wcertificatestool.creation_date = timezone.now()
            new_wcertificatestool.author =  request.user
            
          
            date = timezone.now().strftime('%Y%m%d%H%M%S')
            certificatedir_new = MEDIA_ROOT+'\\files\\w_certificates\\'+date
            zip_all1 ="/static/media/files/w_certificates/w_certificates_"+ date+".zip"
            zip_all = zipfile.ZipFile(MEDIA_ROOT+"/files/w_certificates/"+"/w_certificates_"+ date+".zip", "w")
            new_wcertificatestool.filezip =zip_all1
            
            form.save()

            try: 
                os.makedirs(certificatedir_new)
            except OSError:
                if not os.path.isdir(certificatedir_new):
                    raise


            for u in eventParticipants:
                user_obj=User.objects.get(username=u.user)
                userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
                name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()

                role=u.role
                roletoprint=''
                for r in u.role.all():
                    role_=str(r)
                   #print(role_)
                    roletoprint=roletoprint+role_


             
                lenghtofname=len(name)
                xcoord_name=3507-(lenghtofname/2)*100
                  
                lenghtofrole=len(roletoprint)
                xcoordrole=3507-(lenghtofrole/2)*100
                  
                lenghtofworkshoptitle=len(workshoptitle)
                xcoordworkshoptitle=3507-(lenghtofworkshoptitle/2)*100
                fontfile =MEDIA_ROOT+"\\static\\media\\certificate_template\\times.ttf"
  
                font = ImageFont.truetype(fontfile, 200) #Arial.ttf", 48)
                font1= ImageFont.truetype(fontfile, 100) #Arial.ttf", 48)
                  
                img = Image.open(os.path.join(certificatedir_template,img_template))
                draw = ImageDraw.Draw(img)
                # font = ImageFont.truetype(<font-file>, <font-size>)
                # draw.text((x, y),"Sample Text",(r,g,b))
                draw.text((xcoord_name, 1863),name,(0,119,67),font=font)
                draw.text((xcoordrole, 2280),roletoprint,(0,119,67),font=font)
                draw.text((xcoordworkshoptitle, 2745),workshoptitle,(0,119,67),font=font)
                draw.text((3150, 4250),"Rome, Italy, "+timezone.now().strftime('%Y')+" ",(0,0,0),font=font1)
                img.save(os.path.join(certificatedir_new,str(u.user)+'.jpg'))
                strfpath=os.path.join(certificatedir_new,str(u.user)+'.jpg')
                zip_all.write(strfpath, str(u.user)+'.jpg')

            

            zip_all.close()
            zip_all=zip_all1
            
            shutil.rmtree(certificatedir_new)
           
            info(request, _("Workshop Certificates generated."))
            return redirect("w-certificatestool-detail",new_wcertificatestool.id)
        else:
             return render_to_response('certificates/w_certificatestool_create.html', {'form': form,'events_set':events_set,'imga':imgA,'imgb':imgB,'zip_all':zip_all,'fontfile':fontfile},
             context_instance=RequestContext(request))
    else:
        form = WorkshopCertificatesToolForm(instance=WorkshopCertificatesTool())
        
      
    return render_to_response('certificates/w_certificatestool_create.html', {'form': form,'zip_all':zip_all,'events_set':events_set,'imga':imgA,'imgb':imgB,},
        context_instance=RequestContext(request))

class MembershipListView(ListView):
    """    Membership List view """
    context_object_name = 'latest'
    model = Group
    date_field = 'date'
    template_name = 'certificates/membershiptool_list.html'
    queryset = Group.objects.all()

class MembershipShortListView(ListView):
    """    Membership List view """
    context_object_name = 'latest'
    model = Group
    date_field = 'date'
    template_name = 'certificates/membershiptool_shortlist.html'
    queryset = Group.objects.all()
    
    
from django.db.models import Q

class TopicListView(ListView):
    """
   Topic
       
    """
    context_object_name = 'latest'
    model = Topic
    date_field = 'publish_date'
    template_name = 'topic/topic_list.html'
    queryset = Topic.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
   # paginate_by = 500

    def get_queryset(self):
        """ only return a from the  """
        return Topic.objects.filter()
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(TopicListView, self).get_context_data(**kwargs)
        
       
        topic_table1=[]
        topic_table2=[]
        topic_table3=[]
        topic_table4=[]
        topic_table5=[]
        #queryset0=Topic.objects.filter(Q( topic_type = 1)|Q(topic_type = 2)).order_by('id')
        queryset0=Topic.objects.filter( topic_type = 0,is_version=False).order_by('id')
            
        queryset1=Topic.objects.filter(topic_type = 1,is_version=False).order_by('priority', 'topicstatus')
        queryset2=Topic.objects.filter(topic_type = 2,is_version=False).order_by('priority', 'topicstatus')
                   
        for t in queryset0:
            topic_table1.append(t)
              
        for t in queryset1:
            drafting_body=t.drafting_body
            for d in drafting_body.all():
                if d.draftingbody =='EWG' or  d.draftingbody == 'TPFF' or d.draftingbody == 'TPFQ'or d.draftingbody == 'TPPT' or  d.draftingbody == 'N/A':
                    if t not in topic_table2:
                        topic_table2.append(t)
                
        for t in queryset2:
            drafting_body=t.drafting_body
            for d in drafting_body.all():
                if d.draftingbody == 'TPDP'    :
                   topic_table3.append(t)
        
                elif d.draftingbody == 'TPPT':
                    topic_table4.append(t)
        
                elif d.draftingbody == 'TPG':
                    topic_table5.append(t)
   
        context['topic_table1']=    topic_table1
        context['topic_table2']=    topic_table2
        context['topic_table3']=    topic_table3
        context['topic_table4']=    topic_table4
        context['topic_table5']=    topic_table5
        context['12col']=    1
        
        return context

class TopicDetailView(DetailView):
    """ Topic detail page """
    model = Topic
    context_object_name = 'topic'
    template_name = 'topic/topic_detail.html'
    queryset = Topic.objects.filter()
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(TopicDetailView, self).get_context_data(**kwargs)
      
        p = get_object_or_404(Topic, slug=self.kwargs['slug'])
        
        versions= Topic.objects.filter(status=CONTENT_STATUS_PUBLISHED, is_version=True, parent_id=p.id).order_by('-modify_date')
        context['versions'] = versions
        context['8col'] = 1
        
        return context
  
    
   
     
@login_required
@permission_required('ippc.add_publication', login_url="/accounts/login/")
def topic_create(request):
    """ Create topic """
    user = request.user
    author = user
    form = TopicForm(request.POST)
    tl_form = TopicLeadsFormSet(request.POST or None)
    ta_form = TopicAssistantsFormSet(request.POST or None)
        
    if request.method == "POST":
        if form.is_valid() and tl_form.is_valid() and ta_form.is_valid():
            new_topic = form.save(commit=False)
            new_topic.author = request.user
            new_topic.author_id = author.id
            form.save()
            tl_form.instance = new_topic
            tl_form.save()
            ta_form.instance = new_topic
            ta_form.save()
            info(request, _("Successfully created Topic."))
            return redirect("topic-detail",  slug=new_topic.slug)
        else:
             return render_to_response('topic/topic_create.html', {'form': form,'tl_form':tl_form,'ta_form':ta_form},
             context_instance=RequestContext(request))
    else:
        form = TopicForm(instance=Topic())
        tl_form = TopicLeadsFormSet()
        ta_form = TopicAssistantsFormSet()
  
      
    return render_to_response('topic/topic_create.html', {'form': form,'tl_form':tl_form,'ta_form':ta_form},
        context_instance=RequestContext(request))
		


@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def topic_edit(request,  id=None, template_name='topic/topic_edit.html'):
    """ Edit topic """
    user = request.user
    author = user
    if id:
        topic = get_object_or_404(Topic,  pk=id)
        old_topic = get_object_or_404(Topic,  pk=id)
    else:
        topic = Topic(author=request.user)

    if request.POST:
        form = TopicForm(request.POST,    instance=topic)
        tl_form = TopicLeadsFormSet(request.POST, instance=topic)
        ta_form = TopicAssistantsFormSet(request.POST,  instance=topic)
       
        draftingbody=topic.drafting_body
        arraydraftingbody=[]
        for d in draftingbody.all():
               arraydraftingbody.append(d.id)
        str_obj=topic.strategicobj
        arraystr_obj=[]
        for s in str_obj.all():
               arraystr_obj.append(s.id)
        topiclead=TopicLeads.objects.filter(topic_id=old_topic.id)
        arrayleads=[]
        for d in topiclead.all():
            arrayl=[]
           #print(d)
           #print(d.user_id)
            arrayl.append(d.user_id)
            arrayl.append(d.representing_country_id)
            arrayl.append(d.meetingassistantassigned)
            arrayleads.append(arrayl)
        topicassistants=TopicAssistants.objects.filter(topic_id=old_topic.id)
        arrayassistants=[]
        for d in topicassistants.all():
            arrayl=[]
           #print(d)
           #print(d.user_id)
            arrayl.append(d.user_id)
            arrayl.append(d.representing_country_id)
            arrayl.append(d.meetingassistantassigned)
            arrayassistants.append(arrayl)
           
        if form.is_valid() and tl_form.is_valid() and ta_form.is_valid():
            form.save()
           
            tl_form.instance = topic
            tl_form.save()
            ta_form.instance = topic
            ta_form.save()

            old_topic.pk = None
            old_topic.is_version = True
            old_topic.topicnumber_version = topic.topicnumber
            old_topic.parent_id = id
            versions= Topic.objects.filter( is_version=True, parent_id=id).count()
            slug1 = versions+1
            old_topic.topicnumber = topic.topicnumber+str('-'+str(slug1))
            old_topic.slug= old_topic.slug+'-'+str(slug1)
            old_topic.save()
            db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
            cursor = db.cursor()
            
            str_obj=old_topic.strategicobj
           #print(str_obj)
            for d in arraydraftingbody:
                sql = "INSERT INTO ippc_topic_drafting_body(topic_id,draftingbodytype_id) VALUES ("+str(old_topic.id)+", '"+str(d)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()
            for s in arraystr_obj:
                sql = "INSERT INTO ippc_topic_strategicobj(topic_id,stratigicobjective_id) VALUES ("+str(old_topic.id)+", '"+str(s)+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()        
            for d in arrayleads:
               #print(d)
                sql = "INSERT INTO ippc_topicleads(topic_id,user_id,representing_country_id,meetingassistantassigned) VALUES ("+str(old_topic.id)+", '"+str(d[0])+"', '"+str(d[1])+"', '"+str(d[2])+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()     
            for d in arrayassistants:
               #print(d)
                sql = "INSERT INTO ippc_topicassistants(topic_id,user_id,representing_country_id,meetingassistantassigned) VALUES ("+str(old_topic.id)+", '"+str(d[0])+"', '"+str(d[1])+"', '"+str(d[2])+"')"
                #print(sql)
                try:
                    cursor.execute(sql)
                    db.commit()
                except:
                    db.rollback()     

            db.close()
        
            # If the save was successful, success message and redirect to another page
            info(request, _("Successfully updated Topic."))
            return redirect("topic-detail", slug=topic.slug)
                             
    else:
        form = TopicForm(instance=topic)
        tl_form = TopicLeadsFormSet(instance=topic)
        ta_form = TopicAssistantsFormSet(instance=topic)
  
       
    return render_to_response(template_name, {
        'form': form,'tl_form':tl_form,'ta_form':ta_form, "topic": topic,
    }, context_instance=RequestContext(request))
		
      

# http://stackoverflow.com/a/1854453/412329
@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def topic_translate(request,lang, id=None, template_name='topic/topic_translate.html'):
    """ translate topic """
    user = request.user
    author = user
    if id:
        topic = get_object_or_404(Topic, id=id)
        try:
            t_topic = get_object_or_404(TransTopic, translation_id=id,lang=lang)
        except:
            t_topic = TransTopic(lang=lang,translation_id=id)        
    

    if request.POST:
        transform = TransTopicForm(request.POST, instance=t_topic)
        if transform.is_valid():
            transform.save()
            info(request, _("Successfully translated Topic."))
            return redirect("topic-detail", slug=topic.slug)
    else:
        transform = TransTopicForm(instance=t_topic)
        
    return render_to_response(template_name, {
        'transform':transform,'lang':lang,"topic":topic,"t_topic": t_topic,
    }, context_instance=RequestContext(request))
#
#
#
     
from docx import Document
from docx.shared import Inches,Pt,Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.utils.translation import ugettext, ugettext_lazy as _
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
  
from docx.enum.section import WD_ORIENT

def get_gender(val):
    if val!= None:
        if val == 1:
            return "Mr."
        elif val==2: 
            return "Ms."
        elif val==3:    
            return "Mrs."
        elif val==4:    
            return "Professor."
        elif val==5:    
            return "M."
        elif val==6:    
            return "Mme."
        elif val==7:    
            return "Dr."
        elif val==8:    
            return "Sr."
        elif val==9:    
            return "Sra."
    else:    
        return ""     
    

def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width

from docx.shared import RGBColor
from docx.enum.section import WD_ORIENT

@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def generate_topiclist(request,lang):
    """ generate Topic List  """
    updateddate=timezone.now().strftime('%d-%m-%Y')
    
    #document = Document()
   #windows template_path = MEDIA_ROOT+'\\certificate_template\\init_landscape.docx'
    template_path = MEDIA_ROOT+'/certificate_template/init_landscape.docx'
   
    document = Document(template_path)
      
    docx_title="ListOfTopics_"+lang+"_"+updateddate+".docx"
    #docHeader_image_path = MEDIA_ROOT+'\\certificate_template\\banner_landscape.jpg'
    #UNIXdocHeader_image_path = MEDIA_ROOT+'/certificate_template/header.jpg'
   
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.25)
        section.right_margin = Cm(1.25)
    
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Arial'
    
      
    #TITLE and sub-titles
    p= document.add_paragraph("")
    p.add_run(_("List of Topics for IPPC Standards"), style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p= document.add_paragraph("")
    p= document.add_paragraph("")
   # p.add_run(_("(Updated "+updateddate+")"), style = 'CommentsStyle').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p.add_run(_("This List of topics for IPPC standards was last updated on May 2019 and reflects the modifications adopted by the CPM or approved by the SC."), style = 'CommentsStyle').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    
    p.add_run(_("Table 1: "), style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(_("Technical panels and topics for the Technical Panel on Diagnostic Protocols (TPDP), the Technical Panel for the Glossary (TPG) and the Technical Panel on Phytosanitary Treatments (TPPT)"), style = 'CommentsStyle')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p= document.add_paragraph("")
  
    p.add_run(_("Table 2: "), style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(_("Topics for the Expert Working Groups (EWGs), Technical Panel on Forest Quarantine (TPFQ), and Technical Panel on Phytosanitary Treatments (TPPT)"), style = 'CommentsStyle')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p= document.add_paragraph("")
  
    p.add_run(_("Table 3: "), style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(_("Subjects for the Technical Panel on Diagnostic Protocols (TPDP)"), style = 'CommentsStyle')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p= document.add_paragraph("")
  
    p.add_run(_("Table 4: "), style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(_("Subjects for the Technical Panel on Phytosanitary Treatments (TPPT)"), style = 'CommentsStyle')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p= document.add_paragraph("")
  
    p.add_run(_("Table 5: "), style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(_("Subjects for the Technical Panel for the Glossary (TPG)"), style = 'CommentsStyle')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    
   
    p.add_run(_("IPPC Strategic Objectives"), style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p= document.add_paragraph()
    p.add_run(_("A: Food Security\nB: Environmental Protection\nC: Trade Facilitation\nD: Capacity Development"), style = 'CommentsStyle')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    
    p.add_run(_("Priority"), style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p= document.add_paragraph()
    p.add_run(_("Priority 1 to 4 (with 1 being of high priority and 4 being of low priority)"), style = 'CommentsStyle')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    
 
    p.add_run(_("Notes: "), style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(_("Country names and dates are in ISO format (respectively: ISO 3166-1-alpha-2 code and YYYY-MM)."), style = 'CommentsStyle').italic = True
    p.add_run(_("The List of topics is presented in order of priority, as requested by CPM-7 (2012)."), style = 'CommentsStyle').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    

      
    main_tables_array=[]
    topic_table1=[]
    topic_table2=[]
    topic_table3=[]
    topic_table4=[]
    topic_table5=[]
    queryset0=Topic.objects.filter(topic_type = 0,is_version=False).order_by('id')
    queryset1=Topic.objects.filter(topic_type = 1,is_version=False).order_by('priority', 'topicstatus')
    queryset2=Topic.objects.filter(topic_type = 2,is_version=False).order_by('priority', 'topicstatus')
    for t in queryset0:
        topic_table1.append(t)
    for t in queryset1:
        drafting_body=t.drafting_body
        for d in drafting_body.all():
            if d.draftingbody =='EWG' or  d.draftingbody == 'TPFF' or d.draftingbody == 'TPFQ'or d.draftingbody == 'TPPT':
                topic_table2.append(t)
    for t in queryset2:
        drafting_body=t.drafting_body
        for d in drafting_body.all():
            if d.draftingbody == 'TPDP'    :
               topic_table3.append(t)
            elif d.draftingbody == 'TPPT':
                topic_table4.append(t)
            elif d.draftingbody == 'TPG':
                topic_table5.append(t)
    
    table_array1=[]            
    table_array1.append(8)#numcol
    table_array1.append('0000FF')#bgcolorHeader
    table_array1.append('8DB3E2')#bgcolorrows
    table_array1.append(_("Table 1: Technical panels and topics for TPDP, TPG and TPPT"))
    array_titles1= [_("Topic N."),
                    _("Current Title"),
                    _("Drafting Body"),
                   _("Topic under technical area (if applicable)"),
                   _("Added to the list"),
                   _("Lead Steward / TP Lead (Country, Meeting assigned)"),
                   _("Assistant Stewards (Country, Meeting assigned)"),
                   _("Spec No")]
    table_array1.append(array_titles1)    
    cols_width1= [2,4.5,1.2,3,3,5,5,2]
    table_array1.append(topic_table1)      
    table_array1.append(cols_width1)      
    
    table_array2=[]            
    table_array2.append(10)#numcol
    table_array2.append('008000')#bgcolorHeader
    table_array2.append('C2D69B')#bgcolorrows
    table_array2.append(_("Table 2: Topics for EWGs, TPFQ and TPPT (sorted by priority, drafting body, then status)"))
    array_titles2= [_("Topic N."),
                    _("Current Title"),
                    _("Priority"),
                   _("Strategic objective"),
                    _("Drafting Body"),
                    _("Added to the list"),
                    _("Lead Steward / TP Lead (Country, Meeting assigned)"),
                    _("Assistant Stewards (Country, Meeting assigned)"),
                    _("Spec No"),
                    _("Status")]
    table_array2.append(array_titles2)      
    cols_width2= [1.8,4.2,1,1,1.2,1.7,5,5,1,4]
    table_array2.append(topic_table2)      
    table_array2.append( cols_width2)
    
    table_array3=[]            
    table_array3.append(9)#numcol
    table_array3.append('7030A0')#bgcolorHeader
    table_array3.append('CCC0D9')#bgcolorrows
    table_array3.append(_("Table 3: Subjects for TPDP (sorted by priority, topic under, then status)"))
    array_titles3= [ _("Topic N."),
                    _("Current Title"),
                    _("Priority"),
                    _("Strategic objective"),
                    _("Topic under technical area (if applicable)"),
                    _("Added to the list"),
                    _("Discipline Lead (Country)"),
                    _("Referee"),
                    _("Status")]
    table_array3.append(array_titles3)      
    table_array3.append(topic_table3)      
    cols_width3= [1.8,4.7,1,1,1.5,1.5,4.5,5,4]
    table_array3.append( cols_width3)
   
    table_array4=[]            
    table_array4.append(9)#numcol
    table_array4.append('E36C0A')#bgcolorHeader
    table_array4.append('FBD4B4')#bgcolorrows
    table_array4.append(_("Table 4:  Subjects for TPPT (sorted by priority, status, then topic number)"))
    array_titles4= [ _("Topic N."),
                    _("Current Title"),
                    _("Priority"),
                    _("Strategic objective"),
                    _("Topic under technical area (if applicable)"),
                    _("Added to the list"),
                    _("Treatment Lead (Country,Meeting assigned)"),
                    _("Assistant  Lead (Country,Meeting assigned)"),
                    _("Status")]
    table_array4.append(array_titles4)    
    table_array4.append(topic_table4)     
    cols_width4= [1.8,5.2,1,1,1.7,5,5,5]
    table_array4.append( cols_width4)
  
    
    table_array5=[]            
    table_array5.append(6)#numcol
    table_array5.append('C00000')#bgcolorHeader
    table_array5.append('E5B8B7')#bgcolorrows
    table_array5.append(_("Table 5: Subjects for TPG (sorted in English alphabetical order)"))
    array_titles5= [ _("Topic N."),
                    _("Current Title"),
                    _("Drafting body"),
                    _("Topic under technical area (if applicable)"),
                    _("Added to the list"),
                    _("Status")]
    table_array5.append(array_titles5) 
    table_array5.append(topic_table5)      
    cols_width5= [1.8,6,1,8,1.7,6]
    table_array5.append( cols_width5)
  
    mains_array=[]
    mains_array.append(table_array1)  
    mains_array.append(table_array2)  
    mains_array.append(table_array3)  
    mains_array.append(table_array4)  
    mains_array.append(table_array5)
   #print('len='+str(len(mains_array)))
    ii=1

    for tableT in mains_array:
       #print('ii='+str(ii))
        if ii==1:
           sizesplit=14 
        if ii==2:
           sizesplit=10 
        if ii==3:
           sizesplit=11
        if ii==4:
           sizesplit=14 
        if ii==5:
           sizesplit=20 
        
         
        num_col = tableT[0]
        head_color = tableT[1]
        row_color = tableT[2]
        t_title = tableT[3]
       #print('title=:'+str(t_title))
        array_titles =  tableT[4]      
        tabletopisc = tableT[5]
        cols_width= tableT[6]
        
        tabletopisc_split=[]
       #print('len='+str(len(tabletopisc)))
        if len(tabletopisc)>=sizesplit:
             tabletopisc_split = split(tabletopisc,sizesplit) 
            #print(sizesplit)
        else:
             tabletopisc_split = split(tabletopisc,len(tabletopisc)) 
        
           
        for tabletopisc in tabletopisc_split:
            document.add_page_break()    

            p= document.add_paragraph("")
            p= document.add_paragraph("")
            table = document.add_table(rows=2, cols=num_col)
            table.style = document.styles['Table Grid']
            hdr_cells0 = table.rows[0].cells
            hdr_cells  = table.rows[1].cells
            font.size = Pt(8)

            #run = hdr_cells0[0].paragraphs[0].add_run('')
            
            run = hdr_cells0[0].paragraphs[0].add_run(t_title)#.bold #= True
            font = run.font
            font.color.rgb = RGBColor(255, 255, 255)
            run.bold=True
            table.cell(0,0).merge(table.cell(0,num_col-1))

            j=0
            for tt in array_titles:
                run = hdr_cells[j].paragraphs[0].add_run(tt)#.bold #= True
                font = run.font
                font.color.rgb = RGBColor(255,255,255)
                run.bold=True
                j=j+1
            #num_line=1    
            for t in tabletopisc:
                row_cells = table.add_row().cells
              #  run = row_cells[0].paragraphs[0].add_run(str(ii)+'.'+str(num_line))
             #   font = run.font
             #   run.bold=True
             #   font.color.rgb = RGBColor(0, 0, 255)
                    
             #   num_line=num_line+1  
                drafting_body=''
                for d in t.drafting_body.all():
                    dd=''
                    if lang=='en':
                        dd=d.draftingbody
                    elif lang=='es':
                        dd=d.draftingbody_es
                        if dd =='':
                            dd=d.draftingbody
                    elif lang=='fr':
                        dd=d.draftingbody_fr
                        if dd =='':
                            dd=d.draftingbody
                    elif lang=='ar':
                        dd=d.draftingbody_ar
                        if dd =='':
                            dd=d.draftingbody
                    elif lang=='ru':
                        dd=d.draftingbody_ru
                        if dd =='':
                            dd=d.draftingbody
                    elif lang=='zh':
                        dd=d.draftingbody_zh
                        if dd =='':
                            dd=d.draftingbody
                    drafting_body=drafting_body+dd+'\n'
                        
                    
                leads=''
                topiclead=TopicLeads.objects.filter(topic_id=t.id)
               #print(t.id)
                for d in topiclead.all():
                    #print('LEAD')
                    #print(d.user_id)
                    if d.user_id!=None:
                        userippc = get_object_or_404(IppcUserProfile, user_id=d.user_id)
                    
                        leads=leads+(unicode(userippc.first_name))+' '+(unicode(userippc.last_name))
                    if d.representing_country!=None:
                        leads=leads+' ('+d.representing_country.iso+', '
                    if d.meetingassistantassigned!=None:
                        leads=leads+d.meetingassistantassigned+')\n'
                assistants=''
                topicassistants=TopicAssistants.objects.filter(topic_id=t.id)
                for d in topicassistants.all():
                    #print('ASSSS')
                    #print(d.user_id)
                    #print(d.user_id)
                    if d.user_id!=None: 
                        userippc = get_object_or_404(IppcUserProfile, user_id=d.user_id)
                        assistants=assistants+(unicode(userippc.first_name))+' '+(unicode(userippc.last_name))
                    if d.representing_country!=None:
                        assistants=assistants+' ('+d.representing_country.iso+', '
                    if d.meetingassistantassigned!=None:
                        assistants=assistants+d.meetingassistantassigned+')\n'
                addedtolist=''
                if t.addedtolist != 0:
                     addedtolist=addedtolist+ugettext(dict(CPMS)[t.addedtolist]+' ('+str(t.addedtolist)+')')+'\n'
                if t.addedtolist_sc != 0:
                    addedtolist=addedtolist+ugettext(dict(SC_TYPE_CHOICES)[t.addedtolist_sc])
                status=''
              
                if t.topicstatus!=None:
                    status=  ugettext(dict(TOPIC_STATUS_CHOICES)[t.topicstatus])
                priority=  ugettext(dict(TOPIC_PRIORITY_CHOICES)[t.priority])
                str_obj=''
                for d in t.strategicobj.all():
                    str_obj=str_obj+d.strategicobj+'\n'

                n=0
                t_topic=None
                #print('...........1....................')
                #print(t.id)
                t_topics=TransTopic.objects.filter(translation_id=t.id,lang=lang)
                if t_topics.count()>0 and lang!='en':
                    t_topic= get_object_or_404(TransTopic, translation_id=t.id,lang=lang)
                    #print('............2...................')
                    
                for tt in array_titles:
                    field=''
                    if tt==_("Topic N."):
                        field=t.topicnumber
                    elif  tt==_("Current Title"):
                        if t_topic!=None:
                           field= t_topic.title
                        else:
                           field= t.title
                    elif  tt==_("Drafting Body"):
                      field=drafting_body
                    elif  tt==_("Topic under technical area (if applicable)"):
                       if t_topic!=None:
                           field= t_topic.topic_under
                       else:
                           field= t.topic_under
                    elif  tt==_("Added to the list"):
                      field=addedtolist
                    elif  tt==_("Lead Steward / TP Lead (Country, Meeting assigned)") or tt==_("Discipline Lead (Country)") or tt==_("Treatment Lead (Country,Meeting assigned)"):
                      field=leads
                    elif  tt==_("Assistant Stewards (Country, Meeting assigned)") or tt==_("Referee")  or tt==_("Assistant  Lead (Country,Meeting assigned)"):
                      field=assistants
                    elif  tt==_("Spec No"):
                        field=t.specification_number
                    elif  tt==_("Priority"):
                        field=priority
                    elif  tt==_("Status"):
                        field=status
                    elif  tt==_("Strategic objective"):
                        field=str_obj

                    run = row_cells[n].paragraphs[0].add_run(field)
                    font = run.font
                    font.color.rgb = RGBColor(0, 0, 0)
                    n=n+1 

            range1=range(0,num_col)
            k=0
            for row in table.rows:
                if k ==0 or k==1:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "500")
                    trHeight.set(qn('w:hRule'), "atLeast")
                    trPr.append(trHeight)
                    for i in range1 :
                        tc = row.cells[i]._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcVAlign = OxmlElement('w:vAlign')
                        tcVAlign.set(qn('w:val'), "center")
                        tcColor = OxmlElement('w:shd')
                        tcColor.set(qn('w:fill'), head_color)
                        tcPr.append(tcColor)
                        tcPr.append(tcVAlign)
                if k >1:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "400")
                    trHeight.set(qn('w:hRule'), "atLeast")
                    trPr.append(trHeight)
                    for i in range1 :
                        tc = row.cells[i]._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcVAlign = OxmlElement('w:vAlign')
                        tcVAlign.set(qn('w:val'), "center")
                        tcHeight = OxmlElement('w:tcHeight')
                        tcHeight.set(qn('w:val'), "400")
                        tcHeight.set(qn('w:hRule'), "atLeast")
                        tcColor = OxmlElement('w:shd')

                        tcColor.set(qn('w:fill'), row_color)
                        tcPr.append(tcColor)
                        tcPr.append(tcVAlign)
                        tcPr.append(tcHeight)
                        tcHeight
                k=k+1

            s=0
            for c_w in cols_width:
                set_column_width(table.columns[s], Cm(c_w))
                s=s+1
        ii=ii+1   
           
        
           

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    document.save(response)
    return response 



@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def generate_listNEW(request, id=None):
    """ generate Participants/membership List  """
    updateddate=timezone.now().strftime('%d-%m-%Y')
    
    document_title=''
    doc_title=""
    template_path = MEDIA_ROOT+'/certificate_template/init_portrait_MembList.docx'
    numcols=5
    
    all_users=[]
    
    group=None
    
          
       
        
    if id:
        group=Group.objects.get(id=id)
        all_users = group.user_set.all()
        document_title=(group.name).upper()
        grouptitle=slugify(group.name)
        doc_title="Membership_List_"+str(grouptitle)+".docx"
    #windows template_path = MEDIA_ROOT+'\\certificate_template\\init_portrait_MembList.docx'
    
    array_user_r=[]
    for u in all_users:
        regionname=''
        role_region=''
        name=''
        name_title_address=''
        array_u=[]
        country=''
        countryobj=None
        user_obj=u
        
        userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
        splitaddress=userippc.address2.splitlines()
        address2f=''
        for s in splitaddress:
            if s!='':
                address2f+='\r'+s  
        for s in reversed(splitaddress):
            if s!='' and s!=' ':
                country=s
                break
        country=country.strip()
        if country!='':
            countryobj=CountryPage.objects.filter(name=country)
        region=-1    
        if countryobj:
           region=countryobj[0].region
           regionname=ugettext(dict(REGIONS)[region]) 
        else:
            if country=='Israel':
               region=3
               regionname=ugettext(dict(REGIONS)[region]) 

        name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
        
        if userippc.title!=None and  userippc.title!='':
           name_title_address=name_title_address+'\r'+userippc.title
        if userippc.address1!='':
            name_title_address=name_title_address+userippc.address1
        if address2f!='':
           name_title_address=name_title_address+address2f
        if userippc.phone!='':
           name_title_address=name_title_address+'\rTel:'+userippc.phone
        if userippc.mobile!='':
            name_title_address=name_title_address+'\r'+'Mobile:'+userippc.mobile
        if userippc.fax!='':
            name_title_address=name_title_address+'\r'+'Fax:'+userippc.fax

        email=user_obj.email
        if userippc.email_address_alt!='':
            email+=';\r'+userippc.email_address_alt

        membership=''
        cpm=''
        term=''
        termexpires=''
        termbegins=''
        termends=''
        term1=''
        funding=''
        membership=UserMembershipHistory.objects.filter(user_id=userippc.user_id)
        
        role_region=   regionname+'\rMember'
        if membership.count()>0:
            for g in user_obj.groups.all():
                if g.id==group.id:
                    membership1=UserMembershipHistory.objects.filter(user_id=userippc.user_id, group_id=g.id)
                    for m in membership1:
                        funding=m.funding
                        m_sdate=int(m.start_date.year)
                        m_edate=int(m.end_date.year)
                        termbegins=''
                        termexpires=m.end_date.year
                        if m_edate-m_sdate<=3:
                            term='1st term / 3 years'
                        elif m_edate-m_sdate>3:
                            term='2nd term / 3 years'

                        if group.id ==4:#SC
                            termbegins=m.start_date.strftime('%Y')
                            termends=m.end_date.strftime('%Y')
                            if m_edate-m_sdate<=3:
                                cpm=ugettext(dict(CPMS)[m_sdate])+'('+str(m_sdate)+')'
                            elif m_edate-m_sdate>3:
                                cpm=ugettext(dict(CPMS)[m_sdate])+'('+str(m_sdate)+')'+ '\r'+ugettext(dict(CPMS)[m_sdate+3])+'('+str(m_sdate+3)+')'
                            term1=cpm+'\r\r'+term
                            termexpires=termexpires

                        else:#if group.id ==6 or  group.id ==7 or group.id ==8 or group.id ==9 or group.id ==10:#TPFF,etc
                            termbegins=m.start_date.strftime('%b-%Y')
                            termends=m.end_date.strftime('%b-%Y')
                            term1=termbegins+'\r\r'+term
                            termexpires=termends  
        if funding!='':
            term1= term1+ '\r\r('+str(funding)+')'
       
        array_u.append(country)
        array_u.append(role_region)
        array_u.append(name)
        array_u.append(name_title_address)
        array_u.append(email)
        array_u.append(term1)
        array_u.append(termexpires)
        array_u.append(regionname)
       
        array_user_r.append(array_u)
                
    new_list = sorted(array_user_r, key=lambda x: x[7])
    finalusers=[]
    for uu in new_list:
        finalusers.append(uu)  
          
    firstpage=[]
    otherpages=[]
    k=0
    for users in finalusers:#all_users:
        if k<3:
            firstpage.append(users)
        else:
            otherpages.append(users)
        k+=1    
    users_split2=[]
    users_split2.append(firstpage)
    users_split = split(otherpages,4)   
    for arrayuser in users_split:
        users_split2.append(arrayuser)  
    xx=0   
    range1=range(0,numcols)         
    
    #DOC        
    document = Document(template_path)
   
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'
    obj_charstyle1 = obj_styles.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
    obj_font1 = obj_charstyle1.font
    obj_font1.size = Pt(11)
    obj_font1.name = 'Times New Roman'
    p= document.add_paragraph("")

    p.add_run(document_title, style = 'CommentsStyle').bold = True
    
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    p.add_run(" MEMBERSHIP LIST", style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p= document.add_paragraph("")

    p.add_run(ugettext("The numbers in parenthesis refers to FAO travel funding assistance. (0) No funding; (1) Airfare funding; (2) Airfare and DSA funding."), style = 'CommentsStyle1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p= document.add_paragraph("")
    p.add_run("(Updated "+updateddate+")", style = 'CommentsStyle1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    for user in users_split2:
            if xx>0:
                document.add_page_break()  
            xx=xx+1    
            p= document.add_paragraph("")

            table = document.add_table(rows=1, cols=numcols)
            table.style = document.styles['Table Grid']
            hdr_cells = table.rows[0].cells
            font.size = Pt(9)
            kk=0
            run = hdr_cells[kk].paragraphs[0].add_run('Region/\nRole').bold = True
            kk+=1
            run = hdr_cells[kk].paragraphs[0].add_run('Name, mailing, address, telephone').bold = True
            kk+=1
            run = hdr_cells[kk].paragraphs[0].add_run('Email address').bold = True
            kk+=1
            run = hdr_cells[kk].paragraphs[0].add_run('Membership Confirmed/Term begins').bold = True
            kk+=1
            run = hdr_cells[kk].paragraphs[0].add_run('Term expires').bold = True
            kk+=1  

            for u in user:
                row_cells = table.add_row().cells
               
                j=0
                run = row_cells[j].paragraphs[0].add_run(u[1])
                j+=1
                run = row_cells[j].paragraphs[0].add_run(u[2]).bold = True
                run = row_cells[j].paragraphs[0].add_run(u[3]) 
                row_cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
                j+=1
                row_cells[j].text = u[4]
                j+=1
                row_cells[j].paragraphs[0].add_run(u[5])
                j+=1    
                row_cells[j].text = str(u[6])

                k=0
                for row in table.rows:
                    if k ==0:
                        tr = row._tr
                        trPr = tr.get_or_add_trPr()
                        trHeight = OxmlElement('w:trHeight')
                        trHeight.set(qn('w:val'), "500")
                        trHeight.set(qn('w:hRule'), "atLeast")
                        trPr.append(trHeight)
                        for i in range1 :
                            tc = row.cells[i]._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcVAlign = OxmlElement('w:vAlign')
                            tcVAlign.set(qn('w:val'), "center")
                            tcColor = OxmlElement('w:shd')
                            tcColor.set(qn('w:fill'), 'e6e6e6')
                            tcPr.append(tcColor)
                            tcPr.append(tcVAlign)
                    if k >0:
                        tr = row._tr
                        trPr = tr.get_or_add_trPr()
                        trHeight = OxmlElement('w:trHeight')
                        trHeight.set(qn('w:val'), "2500")
                        trHeight.set(qn('w:hRule'), "atLeast")
                        trPr.append(trHeight)

                        for i in range1 :
                            tc = row.cells[i]._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcVAlign = OxmlElement('w:vAlign')
                            tcVAlign.set(qn('w:val'), "top")
                            tcHeight = OxmlElement('w:tcHeight')
                            tcHeight.set(qn('w:val'), "1500")
                            tcHeight.set(qn('w:hRule'), "atLeast")
                            tcPr.append(tcVAlign)
                            tcPr.append(tcHeight)
                    k=k+1
                f=0
                set_column_width(table.columns[f], Cm(3.0))
                f+=1
                set_column_width(table.columns[f], Cm(6.5))
                f+=1
                set_column_width(table.columns[f], Cm(5.0))
                f+=1
                set_column_width(table.columns[f], Cm(2.0))
                f+=1
                set_column_width(table.columns[f], Cm(1.5))

    
    
    
   
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + doc_title
    document.save(response)
    return response 

@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def generate_list(request, id=None, type=None):
    """ generate Participants/membership List  """
    updateddate=timezone.now().strftime('%d-%m-%Y')
    
    event=None
    eventParticipants=None
    document_title=''
    meeting_date= ''
    location=''
    all_users=[]
    numcols=5#3
    doc_title=''
    show_extracols=True#False #for everygroup
    group=None
    template_path=''
    
    if type=='participant':#participant list
        #winodws template_path = MEDIA_ROOT+'\\certificate_template\\init_portrait_ParticipantList.docx'
        template_path = MEDIA_ROOT+'/certificate_template/init_portrait_ParticipantList.docx'
  
        if id:
            event = get_object_or_404(Event, id=id)
            eventtitle=slugify(event.title)
            doc_title="Participants_List_"+str(eventtitle)+".docx"
         
            eventParticipants=EventParticipants.objects.filter(event_id=id)
            all_other_eventParticipants=[]
            all_observers_eventParticipants=[]
            all_secretariat_eventParticipants=[]
            for u in eventParticipants:
                if u.attended:
                    if u.role.count()>0:
                        for r in u.role.all():
                            if r.role == 'Observer':
                                all_observers_eventParticipants.append(u)
                            elif r.role == 'IPPC Secretariat':
                                all_secretariat_eventParticipants.append(u)
                            else:
                                all_other_eventParticipants.append(u)
                    else:            
                        all_other_eventParticipants.append(u)
                      
            for u in all_other_eventParticipants:
                all_users.append(u)
            for u in all_observers_eventParticipants:
                all_users.append(u)
            for u in all_secretariat_eventParticipants:
                all_users.append(u)
#            for g in event.groups.all():
#                if g.id==4 or g.id==6 or g.id==7 or g.id==8  or g.id==9 or g.id==10 :#4=SC 6=TPDP TPFQ=9  TPPT=7 TPG=8 TPFF=10
#                    show_extracols=True
#                    numcols=5
#                         
                
            document_title=(event.title).upper()
            if event.start.month == event.end.month:
                meeting_date=str(event.start.day)+'-'+event.end.strftime('%d %B %Y')
            else:
                meeting_date=event.start.strftime('%d %B')+'-'+event.end.strftime('%d %B %Y')
            if event.location!='':
               location+=str(event.location)+" - "
            if event.venuecity!='':
               location+=event.venuecity+", "
            if event.venuecountry!='':
               location+=event.venuecountry
               
#    elif type=='membership':
#        doc_title="Membership_List.docx"
#        #numcols=3
#        if id:
#            group=Group.objects.get(id=id)
##            if id=='4' or id=='6' or id=='7'  or id=='8'  or id=='9' or id=='10' or id=='3' or id=='28' :#4=SC 6=TPDP TPFQ=9  TPPT=7 TPG=8 TPFF=10 Bureau=3 FC=28
##                show_extracols=True
##                numcols=5
#            all_users = group.user_set.all()
#            document_title=(group.name).upper()
#            grouptitle=slugify(group.name)
#            doc_title="Membership_List_"+str(grouptitle)+".docx"
#        #windows template_path = MEDIA_ROOT+'\\certificate_template\\init_portrait_MembList.docx'
#        template_path = MEDIA_ROOT+'/certificate_template/init_portrait_MembList.docx'
      
    range1=range(0,numcols)         
    #DOC        
    document = Document(template_path)
    #docHeader_image_path = MEDIA_ROOT+'\\certificate_template\\banner_landscape.jpg'
    #UNIXdocHeader_image_path = MEDIA_ROOT+'/certificate_template/header.jpg'
   
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'
    obj_charstyle1 = obj_styles.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
    obj_font1 = obj_charstyle1.font
    obj_font1.size = Pt(11)
    obj_font1.name = 'Times New Roman'
    p= document.add_paragraph("")

    p.add_run(document_title, style = 'CommentsStyle').bold = True
    
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if type=='participant':
        p= document.add_paragraph("")
        p= document.add_paragraph("")
       
        p.add_run(meeting_date+'\r'+location, style = 'CommentsStyle').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p= document.add_paragraph("")
        p.add_run("PARTICIPANTS LIST", style = 'CommentsStyle').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #obj_font.size = Pt(11)  
        p= document.add_paragraph("")
        #p.add_run(ugettext("A check ("+u"\u2713"+") in column 1 indicates confirmed attendance at the meeting."+'\r'+"Members not attending have been taken off the list."+'\r'), style = 'CommentsStyle')#"The numbers in parenthesis refers to FAO travel funding assistance. (0) No funding; (1) Airfare funding; (2) Airfare and DSA funding."
        p.add_run(ugettext("List of members attending the meeting."), style = 'CommentsStyle1')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#    elif type=='membership':
#       #p= document.add_paragraph("")
#       p.add_run(" MEMBERSHIP LIST", style = 'CommentsStyle').bold = True
#       p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#       p= document.add_paragraph("")
#       
#       if show_extracols:
#            p.add_run(ugettext("The numbers in parenthesis refers to FAO travel funding assistance. (0) No funding; (1) Airfare funding; (2) Airfare and DSA funding."), style = 'CommentsStyle1')
#       p.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
    p= document.add_paragraph("")
    p.add_run("(Updated "+updateddate+")", style = 'CommentsStyle1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    
    firstpage=[]
    otherpages=[]
    k=0
    for users in all_users:
        if k<3:
            firstpage.append(users)
        else:
            otherpages.append(users)
        k+=1    
    users_split2=[]
    users_split2.append(firstpage)
    users_split = split(otherpages,4)   
    for arrayuser in users_split:
        users_split2.append(arrayuser)  
    xx=0    
    for user in users_split2:
        if xx>0:
            document.add_page_break()  
        xx=xx+1    
        p= document.add_paragraph("")
        
        table = document.add_table(rows=1, cols=numcols)
        table.style = document.styles['Table Grid']
        hdr_cells = table.rows[0].cells
        font.size = Pt(9)
        kk=0
#        if type=='participant':
#            run = hdr_cells[kk].paragraphs[0].add_run(' ').bold = True
#            kk+=1
        run = hdr_cells[kk].paragraphs[0].add_run('Region/\nRole').bold = True
        kk+=1
        run = hdr_cells[kk].paragraphs[0].add_run('Name, mailing, address, telephone').bold = True
        kk+=1
        run = hdr_cells[kk].paragraphs[0].add_run('Email address').bold = True
        kk+=1
        if show_extracols:
            run = hdr_cells[kk].paragraphs[0].add_run('Membership Confirmed/Term begins').bold = True
            kk+=1
            run = hdr_cells[kk].paragraphs[0].add_run('Term expires').bold = True
            kk+=1  
                 
        for u in user:
            row_cells = table.add_row().cells
            if type=='participant':
                if u.attended == 0:
                    break
            if type=='participant':
                user_obj=User.objects.get(username=u.user)
            #elif type=='membership':
            #    user_obj=u
            
            userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
            
            #attended=''
            role_region=''
            name=''
            name_title_address=''
            address2f=''
            splitaddress=userippc.address2.splitlines()
            for s in splitaddress:
               if s!='':
                 address2f+='\r'+s  
            if type=='participant':
#                if u.attended == 1:
#                    attended =u"\u2713"
                if u.representing_region>0:
                    if u.representing_region==1:
                        region="Africa"
                    elif  u.representing_region==2:
                        region="Asia"
                    elif  u.representing_region==3:
                        region="Europe"
                    elif  u.representing_region==4: 
                        region="Latin America and Caribbean"
                    elif  u.representing_region==5: 
                        region="Near East"
                    elif  u.representing_region==6: 
                        region="North America"
                    elif  u.representing_region==7: 
                        region="South West Pacific"  
                    role_region=role_region+region.upper()    
                if u.representing_organization!=None:
                    role_region=role_region+'\r'+str(u.representing_organization)
                if u.representing_country!=None:
                    role_region=role_region+'\r'+str(u.representing_country)
                if u.role.count()>0:
                    for r in u.role.all():
                        role_region=role_region+'\r'+str(r)

#            elif type=='membership':
#                v=0
#                country=''
#                regionname=''
#                for s in splitaddress:
#                    if s!='':
#                        if v==len(splitaddress)-1:
#                            country=s
#                    v+=1 
#                country=country.strip()
#                countryobj=CountryPage.objects.filter(name=country)
#                if countryobj:
#                   region=countryobj[0].region
#                   regionname=ugettext(dict(REGIONS)[region])
#                role_region=   regionname+'\rMember'
            
            name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
            
            if userippc.title!=None and  userippc.title!='':
               name_title_address=name_title_address+'\r'+userippc.title
            if userippc.address1!='':
                name_title_address=name_title_address+userippc.address1
            if address2f!='':
               name_title_address=name_title_address+address2f
            if userippc.phone!='':
               name_title_address=name_title_address+'\rTel:'+userippc.phone
            if userippc.mobile!='':
                name_title_address=name_title_address+'\r'+'Mobile:'+userippc.mobile
            if userippc.fax!='':
                name_title_address=name_title_address+'\r'+'Fax:'+userippc.fax

            email=user_obj.email
            if userippc.email_address_alt!='':
                email+=';\r'+userippc.email_address_alt
                
            membership=''
            cpm=''
            term=''
            termexpires=''
            termbegins=''
            termends=''
            term1=''
            funding=''
            membership=UserMembershipHistory.objects.filter(user_id=userippc.user_id)
            if type=='participant':
                    if membership.count()>0:
                        for g in user_obj.groups.all():
                            if g in event.groups.all():
                                if g.id==4 or g.id==6 or g.id==7 or g.id==8  or g.id==9 or g.id==10 :#4=SC 6=TPDP TPFQ=9  TPPT=7 TPG=8 TPFF=10
                                    membership1=UserMembershipHistory.objects.filter(user_id=userippc.user_id, group_id=g.id)
                                    for m in membership1:
                                        m_sdate=int(m.start_date.year)
                                        m_edate=int(m.end_date.year)
                                        termbegins=''
                                        termexpires=m.end_date.year
                                        if m_edate-m_sdate<=3:
                                            term='1st term / 3 years'
                                        elif m_edate-m_sdate>3:
                                            term='2nd term / 3 years'


                                        if g.id ==4  or g.id ==3 or g.id ==28:#SC
                                            termbegins=m.start_date.strftime('%Y')
                                            termends=m.end_date.strftime('%Y')
                                            if m_edate-m_sdate<=3:
                                                cpm=ugettext(dict(CPMS)[m_sdate])+'('+str(m_sdate)+')'
                                            elif m_edate-m_sdate>3:
                                                cpm=ugettext(dict(CPMS)[m_sdate])+'('+str(m_sdate)+')'+ '\r'+ugettext(dict(CPMS)[m_sdate+3])+'('+str(m_sdate+3)+')'
                                            term1=cpm+'\r\r'+term
                                            termexpires=termexpires

                                        else:#if g.id ==6 or  g.id ==7 or g.id ==8 or g.id ==9 or g.id ==10:#TPFF,etc
                                            termbegins=m.start_date.strftime('%b-%Y')
                                            termends=m.end_date.strftime('%b-%Y')
                                            term1=termbegins+'\r\r'+term
                                            termexpires=termends
#            elif type=='membership':
#                if membership.count()>0:
#                    for g in user_obj.groups.all():
#                        
#                        if g.id==group.id:
#                            membership1=UserMembershipHistory.objects.filter(user_id=userippc.user_id, group_id=g.id)
#                            for m in membership1:
#                                funding=m.funding
#                                m_sdate=int(m.start_date.year)
#                                m_edate=int(m.end_date.year)
#                                termbegins=''
#                                termexpires=m.end_date.year
#                                if m_edate-m_sdate<=3:
#                                    term='1st term / 3 years'
#                                elif m_edate-m_sdate>3:
#                                    term='2nd term / 3 years'
#                                    
#                                
#                                if group.id ==4:#SC
#                                    termbegins=m.start_date.strftime('%Y')
#                                    termends=m.end_date.strftime('%Y')
#                                    if m_edate-m_sdate<=3:
#                                        cpm=ugettext(dict(CPMS)[m_sdate])+'('+str(m_sdate)+')'
#                                    elif m_edate-m_sdate>3:
#                                        cpm=ugettext(dict(CPMS)[m_sdate])+'('+str(m_sdate)+')'+ '\r'+ugettext(dict(CPMS)[m_sdate+3])+'('+str(m_sdate+3)+')'
#                                    term1=cpm+'\r\r'+term
#                                    termexpires=termexpires
#   
#                                elif group.id ==6 or  group.id ==7 or group.id ==8 or group.id ==9 or group.id ==10:#TPFF,etc
#                                    termbegins=m.start_date.strftime('%b-%Y')
#                                    termends=m.end_date.strftime('%b-%Y')
#                                    term1=termbegins+'\r\r'+term
#                                    termexpires=termends
    
                                
                               

            j=0
#            if type=='participant':
#                 row_cells[j].text=attended
#                 j+=1

            run = row_cells[j].paragraphs[0].add_run(role_region)
            j+=1
            run = row_cells[j].paragraphs[0].add_run(name).bold = True
            run = row_cells[j].paragraphs[0].add_run(name_title_address) 
            row_cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
            j+=1
            row_cells[j].text = email
            j+=1
            
           # if type=='membership':
           #     if funding!='':
           #         term1= term1+ '\r\r('+str(funding)+')'
            if show_extracols:
                row_cells[j].paragraphs[0].add_run(term1)
                j+=1    
                row_cells[j].text = str(termexpires)

            k=0
            for row in table.rows:
                if k ==0:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "500")
                    trHeight.set(qn('w:hRule'), "atLeast")
                    trPr.append(trHeight)
                    for i in range1 :
                        tc = row.cells[i]._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcVAlign = OxmlElement('w:vAlign')
                        tcVAlign.set(qn('w:val'), "center")
                        tcColor = OxmlElement('w:shd')
                        tcColor.set(qn('w:fill'), 'e6e6e6')
                        tcPr.append(tcColor)
                        tcPr.append(tcVAlign)
                if k >0:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "2500")
                    trHeight.set(qn('w:hRule'), "atLeast")
                    trPr.append(trHeight)

                    for i in range1 :
                        tc = row.cells[i]._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcVAlign = OxmlElement('w:vAlign')
                        tcVAlign.set(qn('w:val'), "top")
                        tcHeight = OxmlElement('w:tcHeight')
                        tcHeight.set(qn('w:val'), "1500")
                        tcHeight.set(qn('w:hRule'), "atLeast")
                        tcPr.append(tcVAlign)
                        tcPr.append(tcHeight)
                k=k+1
            f=0
#            if type=='participant':
#               set_column_width(table.columns[f], Cm(0.5))
#               f+=1
            set_column_width(table.columns[f], Cm(3.0))
            f+=1
            set_column_width(table.columns[f], Cm(6.5))
            f+=1
            set_column_width(table.columns[f], Cm(5.0))
            f+=1
            if show_extracols:
                set_column_width(table.columns[f], Cm(2.0))
                f+=1
                set_column_width(table.columns[f], Cm(1.5))
            
             
         
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + doc_title
    document.save(response)
    return response 

@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def generate_listNEW1(request, id=None):
    """ generate Participants/membership List  """
    updateddate=timezone.now().strftime('%d-%m-%Y')
    
    event=None
    eventParticipants=None
    document_title=''
    meeting_date= ''
    location=''
    all_users=[]
    numcols=5
    doc_title=''
    
    template_path = MEDIA_ROOT+'/certificate_template/init_portrait_ParticipantList.docx'
  
        
    if id:
        event = get_object_or_404(Event, id=id)
        eventtitle=slugify(event.title)
        doc_title="Participants_List_"+str(eventtitle)+".docx"

        eventParticipants=EventParticipants.objects.filter(event_id=id)
        all_other_eventParticipants=[]
        all_observers_eventParticipants=[]
        all_secretariat_eventParticipants=[]
        for u in eventParticipants:
            if u.attended:
                if u.role.count()>0:
                    for r in u.role.all():
                        if r.role == 'Observer':
                            all_observers_eventParticipants.append(u)
                        elif r.role == 'IPPC Secretariat':
                            all_secretariat_eventParticipants.append(u)
                        else:
                            all_other_eventParticipants.append(u)
                else:            
                    all_other_eventParticipants.append(u)

        for u in all_other_eventParticipants:
            all_users.append(u)
        for u in all_observers_eventParticipants:
            all_users.append(u)
        for u in all_secretariat_eventParticipants:
            all_users.append(u)

        document_title=(event.title).upper()
        if event.start.month == event.end.month:
            meeting_date=str(event.start.day)+'-'+event.end.strftime('%d %B %Y')
        else:
            meeting_date=event.start.strftime('%d %B')+'-'+event.end.strftime('%d %B %Y')
        if event.location!='':
           location+=str(event.location)+" - "
        if event.venuecity!='':
           location+=event.venuecity+", "
        if event.venuecountry!='':
           location+=event.venuecountry
    
    array_user_r=[]
    for u in all_users:
        regionname=''
        role_region=''
        name=''
        name_title_address=''
        array_u=[]
        country=''
        
        #user_obj=u
        user_obj=User.objects.get(username=u.user)
        userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
        splitaddress=userippc.address2.splitlines()
        address2f=''
        for s in splitaddress:
            if s!='':
                address2f+='\r'+s  
       
               
        name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
        
        if userippc.title!=None and  userippc.title!='':
           name_title_address=name_title_address+'\r'+userippc.title
        if userippc.address1!='':
            name_title_address=name_title_address+userippc.address1
        if address2f!='':
           name_title_address=name_title_address+address2f
        if userippc.phone!='':
           name_title_address=name_title_address+'\rTel:'+userippc.phone
        if userippc.mobile!='':
            name_title_address=name_title_address+'\r'+'Mobile:'+userippc.mobile
        if userippc.fax!='':
            name_title_address=name_title_address+'\r'+'Fax:'+userippc.fax

        regionname=''
        role_region=''
        print(u.representing_region)
        if u.representing_region>0:
            region=''
            if u.representing_region==1:
                region="Africa"
            elif  u.representing_region==2:
                region="Asia"
            elif  u.representing_region==3:
                region="Europe"
            elif  u.representing_region==4: 
                region="Latin America and Caribbean"
            elif  u.representing_region==5: 
                region="Near East"
            elif  u.representing_region==6: 
                region="North America"
            elif  u.representing_region==7: 
                region="South West Pacific"  
            regionname=region
            role_region=role_region+region.upper()    
        if u.representing_organization!=None:
            partnerpage = get_object_or_404(PartnersPage, name=u.representing_organization)
            page = get_object_or_404(Page, id=partnerpage.page_ptr_id)
            role_region=role_region+'\r'+str(page.title)
            regionname=page.title
            
        if u.representing_country!=None:
            role_region=role_region+'\r'+str(u.representing_country)
            regionname=u.representing_country
     
        if u.role.count()>0:
            
            for r in u.role.all():
                role_region=role_region+'\r'+str(r)  
        
        email=user_obj.email
        if userippc.email_address_alt!='':
            email+=';\r'+userippc.email_address_alt

        membership=''
        cpm=''
        term=''
        termexpires=''
        termbegins=''
        termends=''
        term1=''
        funding=''
        membership=UserMembershipHistory.objects.filter(user_id=userippc.user_id)
        
        #role_region+=   str(regionname)+' **Member'
        if membership.count()>0:
            for g in user_obj.groups.all():
                if g in event.groups.all():
                    if g.id==4 or g.id==6 or g.id==7 or g.id==8  or g.id==9 or g.id==10 :#4=SC 6=TPDP TPFQ=9  TPPT=7 TPG=8 TPFF=10
                        membership1=UserMembershipHistory.objects.filter(user_id=userippc.user_id, group_id=g.id)
                        for m in membership1:
                            m_sdate=int(m.start_date.year)
                            m_edate=int(m.end_date.year)
                            termbegins=''
                            termexpires=m.end_date.year
                            if m_edate-m_sdate<=3:
                                term='1st term / 3 years'
                            elif m_edate-m_sdate>3:
                                term='2nd term / 3 years'


                            if g.id ==4  or g.id ==3 or g.id ==28:#SC
                                termbegins=m.start_date.strftime('%Y')
                                termends=m.end_date.strftime('%Y')
                                if m_edate-m_sdate<=3:
                                    cpm=ugettext(dict(CPMS)[m_sdate])+'('+str(m_sdate)+')'
                                elif m_edate-m_sdate>3:
                                    cpm=ugettext(dict(CPMS)[m_sdate])+'('+str(m_sdate)+')'+ '\r'+ugettext(dict(CPMS)[m_sdate+3])+'('+str(m_sdate+3)+')'
                                term1=cpm+'\r\r'+term
                                termexpires=termexpires

                            else:#if g.id ==6 or  g.id ==7 or g.id ==8 or g.id ==9 or g.id ==10:#TPFF,etc
                                termbegins=m.start_date.strftime('%b-%Y')
                                termends=m.end_date.strftime('%b-%Y')
                                term1=termbegins+'\r\r'+term
                                termexpires=termends
        if funding!='':
            term1= term1+ '\r\r('+str(funding)+')'
       
        array_u.append(country)
        array_u.append(role_region)
        array_u.append(name)
        array_u.append(name_title_address)
        array_u.append(email)
        array_u.append(term1)
        array_u.append(termexpires)
        print(regionname)
        array_u.append(regionname)
       
        array_user_r.append(array_u)
                
    new_list = sorted(array_user_r, key=lambda x: x[7])
    finalusers=[]
    for uu in new_list:
        finalusers.append(uu)  
    
    firstpage=[]
    otherpages=[]
    k=0
    for users in finalusers:#all_users:
        if k<3:
            firstpage.append(users)
        else:
            otherpages.append(users)
        k+=1    
    users_split2=[]
    users_split2.append(firstpage)
    users_split = split(otherpages,4)   
    for arrayuser in users_split:
        users_split2.append(arrayuser)  
        
   
    range1=range(0,numcols)         
    #DOC        
    document = Document(template_path)
    #docHeader_image_path = MEDIA_ROOT+'\\certificate_template\\banner_landscape.jpg'
    #UNIXdocHeader_image_path = MEDIA_ROOT+'/certificate_template/header.jpg'
   
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'
    obj_charstyle1 = obj_styles.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
    obj_font1 = obj_charstyle1.font
    obj_font1.size = Pt(11)
    obj_font1.name = 'Times New Roman'
    p= document.add_paragraph("")

    p.add_run(document_title, style = 'CommentsStyle').bold = True
    
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p= document.add_paragraph("")
    p= document.add_paragraph("")

    p.add_run(meeting_date+'\r'+location, style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p= document.add_paragraph("")
    p.add_run("PARTICIPANTS LIST", style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #obj_font.size = Pt(11)  
    p= document.add_paragraph("")
    #p.add_run(ugettext("A check ("+u"\u2713"+") in column 1 indicates confirmed attendance at the meeting."+'\r'+"Members not attending have been taken off the list."+'\r'), style = 'CommentsStyle')#"The numbers in parenthesis refers to FAO travel funding assistance. (0) No funding; (1) Airfare funding; (2) Airfare and DSA funding."
    p.add_run(ugettext("List of members attending the meeting."), style = 'CommentsStyle1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p= document.add_paragraph("")
    p.add_run("(Updated "+updateddate+")", style = 'CommentsStyle1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    
#    firstpage=[]
#    otherpages=[]
#    k=0
#    for users in all_users:
#        if k<3:
#            firstpage.append(users)
#        else:
#            otherpages.append(users)
#        k+=1    
#    users_split2=[]
#    users_split2.append(firstpage)
#    users_split = split(otherpages,4)   
#    for arrayuser in users_split:
#        users_split2.append(arrayuser)  
    xx=0    
    for user in users_split2:
        if xx>0:
            document.add_page_break()  
        xx=xx+1    
        p= document.add_paragraph("")
        
        table = document.add_table(rows=1, cols=numcols)
        table.style = document.styles['Table Grid']
        hdr_cells = table.rows[0].cells
        font.size = Pt(9)
        kk=0
        run = hdr_cells[kk].paragraphs[0].add_run('Region/\nRole').bold = True
        kk+=1
        run = hdr_cells[kk].paragraphs[0].add_run('Name, mailing, address, telephone').bold = True
        kk+=1
        run = hdr_cells[kk].paragraphs[0].add_run('Email address').bold = True
        kk+=1
        run = hdr_cells[kk].paragraphs[0].add_run('Membership Confirmed/Term begins').bold = True
        kk+=1
        run = hdr_cells[kk].paragraphs[0].add_run('Term expires').bold = True
        kk+=1  
                 
        for u in user:
            row_cells = table.add_row().cells
#            if u.attended == 0:
#                break
#            user_obj=User.objects.get(username=u.user)
#            userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
#            
#            role_region=''
#            name=''
#            name_title_address=''
#            address2f=''
#            splitaddress=userippc.address2.splitlines()
#            for s in splitaddress:
#               if s!='':
#                 address2f+='\r'+s  
#            
#            if u.representing_region>0:
#                 if u.representing_region==1:
#                     region="Africa"
#                 elif  u.representing_region==2:
#                     region="Asia"
#                 elif  u.representing_region==3:
#                     region="Europe"
#                 elif  u.representing_region==4: 
#                     region="Latin America and Caribbean"
#                 elif  u.representing_region==5: 
#                     region="Near East"
#                 elif  u.representing_region==6: 
#                     region="North America"
#                 elif  u.representing_region==7: 
#                     region="South West Pacific"  
#                 role_region=role_region+region.upper()    
#            if u.representing_organization!=None:
#                 role_region=role_region+'\r'+str(u.representing_organization)
#            if u.representing_country!=None:
#                 role_region=role_region+'\r'+str(u.representing_country)
#            if u.role.count()>0:
#                 for r in u.role.all():
#                     role_region=role_region+'\r'+str(r)
#   
#            name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
#            
#            if userippc.title!=None and  userippc.title!='':
#               name_title_address=name_title_address+'\r'+userippc.title
#            if userippc.address1!='':
#                name_title_address=name_title_address+userippc.address1
#            if address2f!='':
#               name_title_address=name_title_address+address2f
#            if userippc.phone!='':
#               name_title_address=name_title_address+'\rTel:'+userippc.phone
#            if userippc.mobile!='':
#                name_title_address=name_title_address+'\r'+'Mobile:'+userippc.mobile
#            if userippc.fax!='':
#                name_title_address=name_title_address+'\r'+'Fax:'+userippc.fax
#
#            email=user_obj.email
#            if userippc.email_address_alt!='':
#                email+=';\r'+userippc.email_address_alt
#                
#            membership=''
#            cpm=''
#            term=''
#            termexpires=''
#            termbegins=''
#            termends=''
#            term1=''
#            funding=''
#            membership=UserMembershipHistory.objects.filter(user_id=userippc.user_id)
#           
#            if membership.count()>0:
#                for g in user_obj.groups.all():
#                    if g in event.groups.all():
#                        if g.id==4 or g.id==6 or g.id==7 or g.id==8  or g.id==9 or g.id==10 :#4=SC 6=TPDP TPFQ=9  TPPT=7 TPG=8 TPFF=10
#                            membership1=UserMembershipHistory.objects.filter(user_id=userippc.user_id, group_id=g.id)
#                            for m in membership1:
#                                m_sdate=int(m.start_date.year)
#                                m_edate=int(m.end_date.year)
#                                termbegins=''
#                                termexpires=m.end_date.year
#                                if m_edate-m_sdate<=3:
#                                    term='1st term / 3 years'
#                                elif m_edate-m_sdate>3:
#                                    term='2nd term / 3 years'
#
#
#                                if g.id ==4  or g.id ==3 or g.id ==28:#SC
#                                    termbegins=m.start_date.strftime('%Y')
#                                    termends=m.end_date.strftime('%Y')
#                                    if m_edate-m_sdate<=3:
#                                        cpm=ugettext(dict(CPMS)[m_sdate])+'('+str(m_sdate)+')'
#                                    elif m_edate-m_sdate>3:
#                                        cpm=ugettext(dict(CPMS)[m_sdate])+'('+str(m_sdate)+')'+ '\r'+ugettext(dict(CPMS)[m_sdate+3])+'('+str(m_sdate+3)+')'
#                                    term1=cpm+'\r\r'+term
#                                    termexpires=termexpires
#
#                                else:#if g.id ==6 or  g.id ==7 or g.id ==8 or g.id ==9 or g.id ==10:#TPFF,etc
#                                    termbegins=m.start_date.strftime('%b-%Y')
#                                    termends=m.end_date.strftime('%b-%Y')
#                                    term1=termbegins+'\r\r'+term
#                                    termexpires=termends
#                        
                               

            j=0
            run = row_cells[j].paragraphs[0].add_run(u[1])
            j+=1
            run = row_cells[j].paragraphs[0].add_run(u[2]).bold = True
            run = row_cells[j].paragraphs[0].add_run(u[3]) 
            row_cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
            j+=1
            row_cells[j].text = u[4]
            j+=1
            row_cells[j].paragraphs[0].add_run(u[5])
            j+=1    
            row_cells[j].text = str(u[6])

            k=0
            for row in table.rows:
                if k ==0:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "500")
                    trHeight.set(qn('w:hRule'), "atLeast")
                    trPr.append(trHeight)
                    for i in range1 :
                        tc = row.cells[i]._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcVAlign = OxmlElement('w:vAlign')
                        tcVAlign.set(qn('w:val'), "center")
                        tcColor = OxmlElement('w:shd')
                        tcColor.set(qn('w:fill'), 'e6e6e6')
                        tcPr.append(tcColor)
                        tcPr.append(tcVAlign)
                if k >0:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "2500")
                    trHeight.set(qn('w:hRule'), "atLeast")
                    trPr.append(trHeight)

                    for i in range1 :
                        tc = row.cells[i]._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcVAlign = OxmlElement('w:vAlign')
                        tcVAlign.set(qn('w:val'), "top")
                        tcHeight = OxmlElement('w:tcHeight')
                        tcHeight.set(qn('w:val'), "1500")
                        tcHeight.set(qn('w:hRule'), "atLeast")
                        tcPr.append(tcVAlign)
                        tcPr.append(tcHeight)
                k=k+1
            f=0
            set_column_width(table.columns[f], Cm(3.0))
            f+=1
            set_column_width(table.columns[f], Cm(6.5))
            f+=1
            set_column_width(table.columns[f], Cm(5.0))
            f+=1
            set_column_width(table.columns[f], Cm(2.0))
            f+=1
            set_column_width(table.columns[f], Cm(1.5))
         
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + doc_title
    document.save(response)
    return response 




@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def generate_shortlist(request, id=None, type=type):
    """ generate Short List membership List  """
    updateddate=timezone.now().strftime('%d-%m-%Y')
    
    all_users=[]
    numcols=4
    range1=range(0,numcols)
    document_title=''
    doc_title=''
    
    if type=='membership':
        group=Group.objects.get(id=id)#PAOLAnew
        all_users = group.user_set.all()
        gname=group.name
        doc_title="Membership_"+gname+"_ShortList_"+updateddate+".docx"
        document_title=gname+" Members"
    else:
        if id!=None:
            event = get_object_or_404(Event, id=id)
            eventtitle=slugify(event.title)
            doc_title="Participants_ShortList_"+str(eventtitle)+"_"+updateddate+".docx"
            location=''
            meeting_date=''
            eventParticipants=EventParticipants.objects.filter(event_id=id)
            all_other_eventParticipants=[]
            all_observers_eventParticipants=[]
            for u in eventParticipants:
                if u.attended:
                    all_users.append(u)

            document_title=(event.title).upper()
            if event.start.month == event.end.month:
                meeting_date=str(event.start.day)+'-'+event.end.strftime('%d %B %Y')
            else:
                meeting_date=event.start.strftime('%d %B')+'-'+event.end.strftime('%d %B %Y')
            if event.location!='':
               location+=str(event.location)+" - "
            if event.venuecity!='':
               location+=event.venuecity+", "
            if event.venuecountry!='':
               location+=event.venuecountry

            document_title =document_title+"\n"+meeting_date+"\n"+location
        
    array_region=[]
    array_regionmain=[]
    
    for k,r in REGIONS:
        reg = r+''#.lower()
        array_user_r=[]
                   
        for u in all_users:
            if type=='participant':
                user_obj=User.objects.get(username=u.user)
            elif type=='membership':
                user_obj=u
            
            userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
            array_u=[]
            country=''
            splitaddress=userippc.address2.splitlines()
            
            for s in reversed(splitaddress):
                if s!='' and s!=' ':
                    country=s
                    break    
            country=country.strip()
            if country!='':
                countryobj=CountryPage.objects.filter(name=country)
#            else:
#                v=0
#                for s in splitaddress:
#                    if s!='':
#                        if v==len(splitaddress)-2:
#                            country=s
#                    v+=1 
#                country=country.strip()
#                countryobj=CountryPage.objects.filter(name=country)
              
            region=-1    
            if countryobj:
               region=countryobj[0].region
            else:
               if country=='Israel':
                   region=3
            
            if region == k:
                name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
                term=''
                membership=UserMembershipHistory.objects.filter(user_id=userippc.user_id, group_id=4)
                if membership.count()>0:
                    for m in membership:
                        term=str(m.end_date.year)
                array_u.append(country)
                array_u.append(name)
                array_u.append(term)        
                array_user_r.append(array_u)
        arraya=[]
        
        arraya.append(reg)
        arraya.append(array_user_r)
        array_region.append(arraya)
    
    #DOC        
    #windows 
    #template_path = MEDIA_ROOT+'\\certificate_template\\init_portrait_ShortList.docx'
    template_path = MEDIA_ROOT+'/certificate_template/init_portrait_ShortList.docx'
    document = Document(template_path)
     
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)
    
  
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'
    obj_charstyle1 = obj_styles.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
    obj_font1 = obj_charstyle1.font
    obj_font1.size = Pt(11)
    obj_font1.name = 'Times New Roman'
   
 

    p= document.add_paragraph("")
    p.add_run(document_title, style = 'CommentsStyle').bold = True
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p.add_run("  (as of "+updateddate+")", style = 'CommentsStyle1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    
    p= document.add_paragraph("")
    p.add_run('Regional representation: ', style = 'CommentsStyle1')
    for rr in array_region:
        region=str(rr[0])
        userscount=len(rr[1])
        
        p.add_run(' '+region+'-'+str(userscount)+',', style = 'CommentsStyle1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
   
        
    
    table = document.add_table(rows=1, cols=numcols)
    table.style = document.styles['Table Grid']
    hdr_cells = table.rows[0].cells
    font.size = Pt(9)

    run = hdr_cells[0].paragraphs[0].add_run('FAO Region').bold = True
    run = hdr_cells[1].paragraphs[0].add_run('Country').bold = True
    run = hdr_cells[2].paragraphs[0].add_run('Name').bold = True
    run = hdr_cells[3].paragraphs[0].add_run('Term ends').bold = True
    
    for rr in array_region:
         
        users=rr[1]
        i=0
        for u in users: 
           
            row_cells = table.add_row().cells
            regionanme=''
        
            if i==0:
                regionanme= rr[0]
            i=i+1  
            region= regionanme   
            country=u[0]
            name=u[1]
            term=u[2]
            

            run = row_cells[0].paragraphs[0].add_run(region)
            run = row_cells[1].paragraphs[0].add_run(country)
            run = row_cells[2].paragraphs[0].add_run(name) 
            row_cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
         
            run = row_cells[3].paragraphs[0].add_run(term) 
           
    k=0
    for row in table.rows:
        if k ==0:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), "200")
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)
            for i in range1 :
                tc = row.cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "center")
                tcColor = OxmlElement('w:shd')
                tcColor.set(qn('w:fill'), 'e6e6e6')
                tcPr.append(tcColor)
                tcPr.append(tcVAlign)
        if k >0:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), "400")
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)

            for i in range1 :
                tc = row.cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "top")
                tcHeight = OxmlElement('w:tcHeight')
                tcHeight.set(qn('w:val'), "400")
                tcHeight.set(qn('w:hRule'), "atLeast")
                tcPr.append(tcVAlign)
                tcPr.append(tcHeight)
        k=k+1
    set_column_width(table.columns[0], Cm(4.0))
    set_column_width(table.columns[1], Cm(4.0))
    set_column_width(table.columns[2], Cm(6.0))
    set_column_width(table.columns[3], Cm(3.0))
            
    indx=1
    index2=0
    for rr in array_region:
        userscount=len(rr[1]) 
       #print(userscount)
        if userscount>0:
            index2=indx+userscount-1
            table.cell(indx,0).merge(table.cell(index2,0))
            indx=indx+userscount
       
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + doc_title
    document.save(response)
    return response 


@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def generate_shortlistparticipant(request, id=None):
    """ generate Short List membership List  """
    updateddate=timezone.now().strftime('%d-%m-%Y')
    
    all_users=[]
    numcols=3
    range1=range(0,numcols)
    document_title=''
    doc_title=''
    array_user_r=[]
    array_region=[]
    
    
    if id!=None:
        event = get_object_or_404(Event, id=id)
        eventtitle=slugify(event.title)
        doc_title="Participants_ShortList_"+str(eventtitle)+"_"+updateddate+".docx"
        location=''
        meeting_date=''
        eventParticipants=EventParticipants.objects.filter(event_id=id)
        all_other_eventParticipants=[]
        all_observers_eventParticipants=[]
        for u in eventParticipants:
            if u.attended:
                all_users.append(u)

        document_title=(event.title).upper()
        if event.start.month == event.end.month:
            meeting_date=str(event.start.day)+'-'+event.end.strftime('%d %B %Y')
        else:
            meeting_date=event.start.strftime('%d %B')+'-'+event.end.strftime('%d %B %Y')
        if event.location!='':
           location+=str(event.location)+" - "
        if event.venuecity!='':
           location+=event.venuecity+", "
        if event.venuecountry!='':
           location+=event.venuecountry

        document_title =document_title+"\n"+meeting_date+"\n"+location
      
        
       
        for u in all_users:
            user_obj=User.objects.get(username=u.user)
            userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
            role_region=''
            if u.representing_region>0:
                    if u.representing_region==1:
                        region="Africa"
                    elif  u.representing_region==2:
                        region="Asia"
                    elif  u.representing_region==3:
                        region="Europe"
                    elif  u.representing_region==4: 
                        region="Latin America and Caribbean"
                    elif  u.representing_region==5: 
                        region="Near East"
                    elif  u.representing_region==6: 
                        region="North America"
                    elif  u.representing_region==7: 
                        region="South West Pacific"  
                    role_region=role_region+region.upper()    
            if u.representing_organization!=None:
                partnerpage = get_object_or_404(PartnersPage, name=u.representing_organization)
                page = get_object_or_404(Page, id=partnerpage.page_ptr_id)
                role_region=role_region+'\r'+str(page.title)
            if u.representing_country!=None:
                role_region=role_region+'\r'+str(u.representing_country)
            if u.role.count()>0:
                for r in u.role.all():
                    role_region=role_region+'\r'+str(r)  
            
            array_u=[]
            
            name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
            term=''
            membership=UserMembershipHistory.objects.filter(user_id=userippc.user_id, group_id=4)
            if membership.count()>0:
               for m in membership:
                    term=str(m.end_date.year)
            array_u.append(role_region)
            array_u.append(name)
            array_u.append(term)      
            
            array_user_r.append(array_u)
      
    
    #DOC        
    #windows 
    #template_path = MEDIA_ROOT+'\\certificate_template\\init_portrait_ShortList.docx'
    template_path = MEDIA_ROOT+'/certificate_template/init_portrait_ShortList.docx'
    document = Document(template_path)
     
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)
    
  
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'
    obj_charstyle1 = obj_styles.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
    obj_font1 = obj_charstyle1.font
    obj_font1.size = Pt(11)
    obj_font1.name = 'Times New Roman'
   
 

    p= document.add_paragraph("")
    p.add_run(document_title, style = 'CommentsStyle').bold = True
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p.add_run("  (as of "+updateddate+")", style = 'CommentsStyle1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    
    p= document.add_paragraph("")
    
    table = document.add_table(rows=1, cols=3)
    table.style = document.styles['Table Grid']
    hdr_cells = table.rows[0].cells
    font.size = Pt(9)

    run = hdr_cells[0].paragraphs[0].add_run('Region/Role').bold = True
    run = hdr_cells[1].paragraphs[0].add_run('Name').bold = True
    run = hdr_cells[2].paragraphs[0].add_run('Term ends').bold = True
    
    for rr in array_user_r:
        row_cells = table.add_row().cells
        region= rr[0]   
        name=rr[1]
        term=rr[2]
        run = row_cells[0].paragraphs[0].add_run(region)
        run = row_cells[1].paragraphs[0].add_run(name)
        row_cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
        run = row_cells[2].paragraphs[0].add_run(term) 
           
    k=0
    for row in table.rows:
        if k ==0:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), "200")
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)
            for i in range1 :
                tc = row.cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "center")
                tcColor = OxmlElement('w:shd')
                tcColor.set(qn('w:fill'), 'e6e6e6')
                tcPr.append(tcColor)
                tcPr.append(tcVAlign)
        if k >0:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), "400")
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)

            for i in range1 :
                tc = row.cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "top")
                tcHeight = OxmlElement('w:tcHeight')
                tcHeight.set(qn('w:val'), "400")
                tcHeight.set(qn('w:hRule'), "atLeast")
                tcPr.append(tcVAlign)
                tcPr.append(tcHeight)
        k=k+1
    set_column_width(table.columns[0], Cm(6.0))
    set_column_width(table.columns[1], Cm(6.0))
    set_column_width(table.columns[2], Cm(6.0))
            
#    indx=1
#    index2=0
#    for rr in array_region:
#        userscount=len(rr[1]) 
#      
#        if userscount>0:
#            index2=indx+userscount-1
#            table.cell(indx,0).merge(table.cell(index2,0))
#            indx=indx+userscount
#       
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + doc_title
    document.save(response)
    return response 


@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def generate_shortlistOLD(request, id=None, type=type):
    """ generate Short List membership List  """
    updateddate=timezone.now().strftime('%d-%m-%Y')
    
    all_users=[]
    numcols=4
    range1=range(0,numcols)
    document_title=''
    doc_title=''
    
    if type=='membership':
        group=Group.objects.get(id=id)#PAOLAnew
        all_users = group.user_set.all()
        gname=group.name
        doc_title="Membership_"+gname+"_ShortList_"+updateddate+".docx"
        document_title=gname+" Members"
    else:
        if id!=None:
            event = get_object_or_404(Event, id=id)
            eventtitle=slugify(event.title)
            doc_title="Participants_ShortList_"+str(eventtitle)+"_"+updateddate+".docx"
            location=''
            meeting_date=''
            eventParticipants=EventParticipants.objects.filter(event_id=id)
            all_other_eventParticipants=[]
            all_observers_eventParticipants=[]
            for u in eventParticipants:
                if u.attended:
                    all_users.append(u)

            document_title=(event.title).upper()
            if event.start.month == event.end.month:
                meeting_date=str(event.start.day)+'-'+event.end.strftime('%d %B %Y')
            else:
                meeting_date=event.start.strftime('%d %B')+'-'+event.end.strftime('%d %B %Y')
            if event.location!='':
               location+=str(event.location)+" - "
            if event.venuecity!='':
               location+=event.venuecity+", "
            if event.venuecountry!='':
               location+=event.venuecountry

            document_title =document_title+"\n"+meeting_date+"\n"+location
        
    array_region=[]
    array_regionmain=[]
    
    for k,r in REGIONS:
        reg = r+''#.lower()
        array_user_r=[]
                   
        for u in all_users:
            if type=='participant':
                user_obj=User.objects.get(username=u.user)
            elif type=='membership':
                user_obj=u
            
            userippc = get_object_or_404(IppcUserProfile, user_id=user_obj.id)
           #print("")
           #print("")
           #print(userippc)
            array_u=[]
            country=''
            #userippc = get_object_or_404(IppcUserProfile, user_id=u.id)
            splitaddress=userippc.address2.splitlines()
            
            v=0
           #print(splitaddress)
            for s in splitaddress:
                if s!='':
                    if v==len(splitaddress)-1:
                        country=s
                       #print(country)
                v+=1 
            country=country.strip()
            if country!='':
                countryobj=CountryPage.objects.filter(name=country)
               #print(country)
               #print(countryobj)
            else:
                v=0
                for s in splitaddress:
                    if s!='':
                        if v==len(splitaddress)-2:
                            country=s
                           #print(country)
                    v+=1 
                country=country.strip()
                countryobj=CountryPage.objects.filter(name=country)
               #print(country)
               #print(countryobj)
            region=-1    
            if countryobj:
               region=countryobj[0].region
            else:
               if country=='Israel':
                   region=3
            
            if region == k:
                name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
                term=''
                membership=UserMembershipHistory.objects.filter(user_id=userippc.user_id, group_id=4)
                if membership.count()>0:
                    for m in membership:
                        term=str(m.end_date.year)
                array_u.append(country)
                array_u.append(name)
                array_u.append(term)        
                array_user_r.append(array_u)
        arraya=[]
        
        arraya.append(reg)
        arraya.append(array_user_r)
        array_region.append(arraya)
    
    #DOC        
    #windows 
    #template_path = MEDIA_ROOT+'\\certificate_template\\init_portrait_ShortList.docx'
    template_path = MEDIA_ROOT+'/certificate_template/init_portrait_ShortList.docx'
    document = Document(template_path)
     
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)
    
  
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'
    obj_charstyle1 = obj_styles.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
    obj_font1 = obj_charstyle1.font
    obj_font1.size = Pt(11)
    obj_font1.name = 'Times New Roman'
   
 

    p= document.add_paragraph("")
    p.add_run(document_title, style = 'CommentsStyle').bold = True
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p.add_run("  (as of "+updateddate+")", style = 'CommentsStyle1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    
    p= document.add_paragraph("")
    p.add_run('Regional representation: ', style = 'CommentsStyle1')
    for rr in array_region:
        region=str(rr[0])
        userscount=len(rr[1])
        
        p.add_run(' '+region+'-'+str(userscount)+',', style = 'CommentsStyle1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
   
        
    
    table = document.add_table(rows=1, cols=numcols)
    table.style = document.styles['Table Grid']
    hdr_cells = table.rows[0].cells
    font.size = Pt(9)

    run = hdr_cells[0].paragraphs[0].add_run('FAO Region').bold = True
    run = hdr_cells[1].paragraphs[0].add_run('Country').bold = True
    run = hdr_cells[2].paragraphs[0].add_run('Name').bold = True
    run = hdr_cells[3].paragraphs[0].add_run('Term ends').bold = True
    
    for rr in array_region:
         
        users=rr[1]
        i=0
        for u in users: 
           
            row_cells = table.add_row().cells
            regionanme=''
        
            if i==0:
                regionanme= rr[0]
            i=i+1  
            region= regionanme   
            country=u[0]
            name=u[1]
            term=u[2]
            

            run = row_cells[0].paragraphs[0].add_run(region)
            run = row_cells[1].paragraphs[0].add_run(country)
            run = row_cells[2].paragraphs[0].add_run(name) 
            row_cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
         
            run = row_cells[3].paragraphs[0].add_run(term) 
           
    k=0
    for row in table.rows:
        if k ==0:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), "200")
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)
            for i in range1 :
                tc = row.cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "center")
                tcColor = OxmlElement('w:shd')
                tcColor.set(qn('w:fill'), 'e6e6e6')
                tcPr.append(tcColor)
                tcPr.append(tcVAlign)
        if k >0:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), "400")
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)

            for i in range1 :
                tc = row.cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "top")
                tcHeight = OxmlElement('w:tcHeight')
                tcHeight.set(qn('w:val'), "400")
                tcHeight.set(qn('w:hRule'), "atLeast")
                tcPr.append(tcVAlign)
                tcPr.append(tcHeight)
        k=k+1
    set_column_width(table.columns[0], Cm(4.0))
    set_column_width(table.columns[1], Cm(4.0))
    set_column_width(table.columns[2], Cm(6.0))
    set_column_width(table.columns[3], Cm(3.0))
            
    indx=1
    index2=0
    for rr in array_region:
        userscount=len(rr[1]) 
       #print(userscount)
        if userscount>0:
            index2=indx+userscount-1
            table.cell(indx,0).merge(table.cell(index2,0))
            indx=indx+userscount
       
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + doc_title
    document.save(response)
    return response 

#@login_required
#@permission_required('ippc.change_publication', login_url="/accounts/login/")
#def generate_shortlist(request):
#    """ generate Short List SC membership List  """
#    updateddate=timezone.now().strftime('%d-%m-%Y')
#    
#    all_users=[]
#    numcols=4
#    range1=range(0,numcols)  
#    
#    doc_title="Membership_SC_ShortList_"+updateddate+".docx"
#    document_title='Standards Committee Members'
#    group=Group.objects.get(id=4)
#    all_users = group.user_set.all()
#    
#        
#    array_region=[]
#    array_regionmain=[]
#    
#    for k,r in REGIONS:
#        reg = r+''#.lower()
#        array_user_r=[]
#                   
#        for u in all_users:
#            array_u=[]
#            country=''
#            userippc = get_object_or_404(IppcUserProfile, user_id=u.id)
#            splitaddress=userippc.address2.splitlines()
#            
#            v=0
#            print(splitaddress)
#            for s in splitaddress:
#                if s!='':
#                    if v==len(splitaddress)-1:
#                        country=s
#                v+=1 
#            country=country.strip()
#            countryobj=CountryPage.objects.filter(name=country)
#            if countryobj:
#               region=countryobj[0].region
#            else:
#               if country=='Israel':
#                   region=3
#            
#            if region == k:
#                name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
#                term=''
#                membership=UserMembershipHistory.objects.filter(user_id=userippc.user_id, group_id=4)
#                if membership.count()>0:
#                    for m in membership:
#                        term=str(m.end_date.year)
#                array_u.append(country)
#                array_u.append(name)
#                array_u.append(term)        
#                array_user_r.append(array_u)
#        arraya=[]
#        
#        arraya.append(reg)
#        arraya.append(array_user_r)
#        array_region.append(arraya)
#   
#        
#        
#
#           
#    
#    #DOC        
#    #windows 
#    #template_path = MEDIA_ROOT+'\\certificate_template\\init_portrait.docx'
#    template_path = MEDIA_ROOT+'/certificate_template/init_portrait_ShortList.docx'
#    document = Document(template_path)
#     
#    style = document.styles['Normal']
#    font = style.font
#    font.name = 'Arial'
#    font.size = Pt(9)
#    
#  
#    obj_styles = document.styles
#    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
#    obj_font = obj_charstyle.font
#    obj_font.size = Pt(12)
#    obj_font.name = 'Times New Roman'
#    obj_charstyle1 = obj_styles.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
#    obj_font1 = obj_charstyle1.font
#    obj_font1.size = Pt(11)
#    obj_font1.name = 'Times New Roman'
#   
# 
#
#    p= document.add_paragraph("")
#    p.add_run(document_title, style = 'CommentsStyle').bold = True
#    p.style = document.styles['Normal']
#    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#    
#    p.add_run("  (as of "+updateddate+")", style = 'CommentsStyle1')
#    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#    document.add_paragraph()
#    
#    p= document.add_paragraph("")
#    p.add_run('Regional representation: ', style = 'CommentsStyle1')
#    for rr in array_region:
#        region=str(rr[0])
#        userscount=len(rr[1])
#        
#        p.add_run(' '+region+'-'+str(userscount)+',', style = 'CommentsStyle1')
#    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#    document.add_paragraph()
#   
#        
#    
#    table = document.add_table(rows=1, cols=numcols)
#    table.style = document.styles['Table Grid']
#    hdr_cells = table.rows[0].cells
#    font.size = Pt(9)
#
#    run = hdr_cells[0].paragraphs[0].add_run('FAO Region').bold = True
#    run = hdr_cells[1].paragraphs[0].add_run('Country').bold = True
#    run = hdr_cells[2].paragraphs[0].add_run('Name').bold = True
#    run = hdr_cells[3].paragraphs[0].add_run('Term ends').bold = True
#    
#    for rr in array_region:
#         
#        users=rr[1]
#        i=0
#        for u in users: 
#           
#            row_cells = table.add_row().cells
#            regionanme=''
#        
#            if i==0:
#                regionanme= rr[0]
#            i=i+1  
#            region= regionanme   
#            country=u[0]
#            name=u[1]
#            term=u[2]
#            
#
#            run = row_cells[0].paragraphs[0].add_run(region)
#            run = row_cells[1].paragraphs[0].add_run(country)
#            run = row_cells[2].paragraphs[0].add_run(name) 
#            row_cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
#         
#            run = row_cells[3].paragraphs[0].add_run(term) 
#           
#    k=0
#    for row in table.rows:
#        if k ==0:
#            tr = row._tr
#            trPr = tr.get_or_add_trPr()
#            trHeight = OxmlElement('w:trHeight')
#            trHeight.set(qn('w:val'), "200")
#            trHeight.set(qn('w:hRule'), "atLeast")
#            trPr.append(trHeight)
#            for i in range1 :
#                tc = row.cells[i]._tc
#                tcPr = tc.get_or_add_tcPr()
#                tcVAlign = OxmlElement('w:vAlign')
#                tcVAlign.set(qn('w:val'), "center")
#                tcColor = OxmlElement('w:shd')
#                tcColor.set(qn('w:fill'), 'e6e6e6')
#                tcPr.append(tcColor)
#                tcPr.append(tcVAlign)
#        if k >0:
#            tr = row._tr
#            trPr = tr.get_or_add_trPr()
#            trHeight = OxmlElement('w:trHeight')
#            trHeight.set(qn('w:val'), "400")
#            trHeight.set(qn('w:hRule'), "atLeast")
#            trPr.append(trHeight)
#
#            for i in range1 :
#                tc = row.cells[i]._tc
#                tcPr = tc.get_or_add_tcPr()
#                tcVAlign = OxmlElement('w:vAlign')
#                tcVAlign.set(qn('w:val'), "top")
#                tcHeight = OxmlElement('w:tcHeight')
#                tcHeight.set(qn('w:val'), "400")
#                tcHeight.set(qn('w:hRule'), "atLeast")
#                tcPr.append(tcVAlign)
#                tcPr.append(tcHeight)
#        k=k+1
#    set_column_width(table.columns[0], Cm(4.0))
#    set_column_width(table.columns[1], Cm(4.0))
#    set_column_width(table.columns[2], Cm(6.0))
#    set_column_width(table.columns[3], Cm(3.0))
#            
#    indx=1
#    index2=0
#    for rr in array_region:
#        userscount=len(rr[1]) 
#        print(userscount)
#        if userscount>0:
#            index2=indx+userscount-1
#            table.cell(indx,0).merge(table.cell(index2,0))
#            indx=indx+userscount
#       
#    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#    response['Content-Disposition'] = 'attachment; filename=' + doc_title
#    document.save(response)
#    return response 
#

@login_required
@permission_required('ippc.change_publication', login_url="/accounts/login/")
def generate_replacementlist(request):
    """ generate Replacement List SC   """
    updateddate=timezone.now().strftime('%d-%m-%Y')
    
    all_users=[]
    numcols=5
    range1=range(0,numcols)  
    
    doc_title="Membership_SCReplacements_ContactInfo_"+updateddate+".docx"
    document_title='SUBSIDIARY BODY OF\nTHE COMMISSION ON PHYTOSANITARY MEASURES\n\nSTANDARDS COMMITTEE POTENTIAL REPLACEMENTS'

    group=Group.objects.get(id=84)
    all_users = group.user_set.all()
    
        
    array_region=[]
    array_regionmain=[]
    
    for k,r in REGIONS:
        reg = r+''
        array_user_r=[]
        z=1;           
        for u in all_users:
            array_u=[]
            
            country=''
            region=0
            userippc = get_object_or_404(IppcUserProfile, user_id=u.id)
            name_title_address=''
            term1=''
            term2=''
            email=''
            address2f=''
            splitaddress=userippc.address2.splitlines()
            v=0
            for s in splitaddress:
                if s!='':
                    address2f+='\r'+s  
                    for s in splitaddress:
                        if s!='':
                            if v==len(splitaddress)-1:
                                country=s
                        v+=1 
                country=country.strip()
                countryobj=CountryPage.objects.filter(name=country)
                if countryobj:
                   region=countryobj[0].region
           
            
            if region == k:
                replace=''
                if z==1:
                    replace='(Replacement 1)'
                else:
                    replace='(Replacement 2)'
                name=get_gender(userippc.gender)+' '+userippc.first_name+' '+(userippc.last_name).upper()
                if userippc.title!=None and  userippc.title!='':
                    name_title_address=name_title_address+'\r'+userippc.title
                if userippc.address1!='':
                    name_title_address=name_title_address+userippc.address1
                if address2f!='':
                    name_title_address=name_title_address+address2f
                if userippc.phone!='':
                    name_title_address=name_title_address+'\rTel:'+userippc.phone
                if userippc.mobile!='':
                    name_title_address=name_title_address+'\r'+'Mobile:'+userippc.mobile
                if userippc.fax!='':
                    name_title_address=name_title_address+'\r'+'Fax:'+userippc.fax

                email=u.email
                if userippc.email_address_alt!='':
                    email+=';\r'+userippc.email_address_alt    
                
                membership1=UserMembershipHistory.objects.filter(user_id=userippc.user_id, group_id=80)
                for m in membership1:
                    m_sdate=int(m.start_date.year)
                    m_edate=int(m.end_date.year)
                    termbegins=''
                    term2=m.end_date.year
                    term=''
                    if m_edate-m_sdate<=3:
                        term='1st term / 3 years'
                    elif m_edate-m_sdate>3:
                        term='2nd term / 3 years'
                    termbegins=m.start_date.strftime('%Y')
                    termends=m.end_date.strftime('%Y')
                    if m_edate-m_sdate<=3:
                        cpm=ugettext(dict(CPMS)[m_sdate])+'('+str(m_sdate)+')'
                    elif m_edate-m_sdate>3:
                        cpm=ugettext(dict(CPMS)[m_sdate])+'('+str(m_sdate)+')'+ '\r'+ugettext(dict(CPMS)[m_sdate+3])+'('+str(m_sdate+3)+')'
                    term1=cpm+'\r\r'+term
                    term2=termexpires

                  
                array_u.append(replace)
                array_u.append(name)
                array_u.append(name_title_address)
                array_u.append(email)
                array_u.append(term1) 
                array_u.append(term2) 
                array_user_r.append(array_u)
                z+=1
                 
        arraya=[]
        
        arraya.append(reg)
        if len(array_user_r)==0:
            array_u1=[]
            array_u1.append('(Replacement 1)')
            array_u1.append('VACANT')
            array_u1.append('')
            array_u1.append('')
            array_u1.append('') 
            array_u1.append('') 
            array_u2=[]
            array_u2.append('(Replacement 2)')
            array_u2.append('VACANT')
            array_u2.append('')
            array_u2.append('')
            array_u2.append('') 
            array_u2.append('') 
            array_user_r.append(array_u1)
            array_user_r.append(array_u2)
        elif len(array_user_r)==1:
            array_u2=[]
            array_u2.append('(Replacement 1)')
            array_u2.append('VACANT')
            array_u2.append('')
            array_u2.append('')
            array_u2.append('') 
            array_u2.append('') 
            array_user_r.append(array_u2)    
        elif len(array_user_r)>1:
           print('')     
        arraya.append(array_user_r)
        array_region.append(arraya)
   
        
        
   
    
    #DOC        
    #windows 
   # template_path = MEDIA_ROOT+'\\certificate_template\\init_portrait.docx'
    #UNIX 
    template_path = MEDIA_ROOT+'/certificate_template/init_portrait_SC_replaceList.docx'
   
    document = Document(template_path)
     
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)
    
    
    
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'
    obj_charstyle1 = obj_styles.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
    obj_font1 = obj_charstyle1.font
    obj_font1.size = Pt(11)
    obj_font1.name = 'Times New Roman'
    p= document.add_paragraph("")
    p= document.add_paragraph("")
    p.add_run(document_title, style = 'CommentsStyle').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p.add_run("  (as of "+updateddate+")", style = 'CommentsStyle1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    
    p= document.add_paragraph("")
    
    table = document.add_table(rows=1, cols=numcols)
    table.style = document.styles['Table Grid']
    hdr_cells = table.rows[0].cells
    font.size = Pt(9)

    run = hdr_cells[0].paragraphs[0].add_run('FAO Region').bold = True
    run = hdr_cells[1].paragraphs[0].add_run('Name, mailing address, telephone').bold = True
    run = hdr_cells[2].paragraphs[0].add_run('Email address').bold = True
    run = hdr_cells[3].paragraphs[0].add_run('Membership Confirmed').bold = True
    run = hdr_cells[4].paragraphs[0].add_run('Term expires').bold = True
    
    for rr in array_region:
        i=0 
        users=rr[1]
        for u in users: 
            row_cells = table.add_row().cells
            regionanme=''
            replac=''
            if i==0:
                regionanme= rr[0]
            i=i+1  
            
            region = regionanme   
            replace = u[0]
            name = u[1]
            name_title_address = u[2]
            email = u[3]
            term1 = u[4]
            term2 = u[5]
            
            run = row_cells[0].paragraphs[0].add_run(region)
            run = row_cells[1].paragraphs[0].add_run(replace+'\n')
            run = row_cells[1].paragraphs[0].add_run(name+'\n').bold = True
            run = row_cells[1].paragraphs[0].add_run(name_title_address) 
            row_cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
         
            run = row_cells[2].paragraphs[0].add_run(email) 
            run = row_cells[3].paragraphs[0].add_run(term1) 
            run = row_cells[4].paragraphs[0].add_run(term2) 
    height="500"
    bgcolor='ffffff'
    k=0
    for row in table.rows:
        if k ==0:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), "200")
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)
            for i in range1 :
                tc = row.cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "center")
                tcColor = OxmlElement('w:shd')
                tcColor.set(qn('w:fill'), 'e6e6e6')
                tcPr.append(tcColor)
                tcPr.append(tcVAlign)
        if k >0:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), height)
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)

            for i in range1 :
                tc = row.cells[i]._tc
                vacant=False;
                if i==1:
                    splittext=[]
                    for paragraph in row.cells[i].paragraphs:
                        text=paragraph.text
                        splittext=text.splitlines()
                    for sp in splittext:
                        if sp!='' and sp=='VACANT':
                            vacant=True

                if vacant:
                    height="500"
                    bgcolor='e6e6e6'
                else:  
                     height="1500"
                     bgcolor='ffffff'
                
                tcPr = tc.get_or_add_tcPr()
                
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "top")
                tcHeight = OxmlElement('w:tcHeight')
                tcHeight.set(qn('w:val'), height)
                tcHeight.set(qn('w:hRule'), "atLeast")
                tcColor = OxmlElement('w:shd')
                tcColor.set(qn('w:fill'),bgcolor )
                tcPr.append(tcColor)
                tcPr.append(tcVAlign)
                tcPr.append(tcHeight)
               
        k=k+1
    set_column_width(table.columns[0], Cm(4.0))
    set_column_width(table.columns[1], Cm(6.0))
    set_column_width(table.columns[2], Cm(6.0))
    set_column_width(table.columns[3], Cm(3.0))
    set_column_width(table.columns[4], Cm(3.0))

    indx=1
    index2=0
    for rr in array_region:
        userscount=len(rr[1]) 
        if userscount>0:
            index2=indx+userscount-1
            table.cell(indx,0).merge(table.cell(index2,0))
            indx=indx+userscount
       
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + doc_title
    document.save(response)
    return response 

@login_required
@permission_required('ippc.delete_publication', login_url="/accounts/login/")
def my_tool(request):
    form = MyToolForm(request.POST, request.FILES)
    if request.method == "POST":
         if form.is_valid():
            
            new_mytool = form.save(commit=False)
             
            
            form.save()
           
            info(request, _("Successfully created entry"))
            
            return redirect("my_toolres", pk=new_mytool.id)
         else:
             return render_to_response('certificates/mytool.html', {'form': form,},
             context_instance=RequestContext(request))
       
    else:
        form = MyToolForm()
    return render_to_response('certificates/mytool.html', {'form': form},
        context_instance=RequestContext(request))

from django.core.files.storage import FileSystemStorage
import sys
@login_required
@permission_required('ippc.delete_publication', login_url="/accounts/login/")
def nro_stats_files(request):
    prj_dir =PROJECT_ROOT
    
    list_files=None
    list_files1=None
    list_files12=None
    deleted=''
    msg=''
    read=''
    
    if request.method == 'POST':
        #list
        if request.POST['path2']:
            path2 = request.POST['path2']
            list_files = os.listdir(path2)
            msg='path2'+path2
            info(request, _("Successfully diplaying list of files!"+msg))
        
        elif request.POST['path3'] and request.POST['filenametoremove']:
            path3 = request.POST['path3']
            filenametoremove = request.POST['filenametoremove']
            if os.path.isfile(path3+'/'+filenametoremove)      :
                msg=msg+'--->File '+path3+'/'+filenametoremove+' exist! - '
                
                try:
                    os.remove(path3+'/'+filenametoremove)
                    deleted='deleted'
                except OSError:
                    deleted='NO-deleted'
                    pass
                msg= msg+' Deleted: '+str(deleted)+'<br>'
            
            info(request, _("Successfully deleted file!  "+msg))
        
        elif request.POST['path5'] and request.POST['filetoread']:
            path5 = request.POST['path5']
            filetoread = request.POST['filetoread']
            readl=[]
              
            if os.path.isfile(path5+'/'+filetoread)      :
                msg=msg+'--->File '+path5+'/'+filetoread+' exist! - '
                
                try:
                    fd = open(path5+'/'+filetoread)
                    readl = fd.readlines()
                    # Close opened file
                    #os.close(fd)
                    fd.close()
                except OSError:
                    deleted='NO-red'
                    pass
                msg= msg+' read: <br>'
            for ll in readl:
                read += str(ll)+'<br>'
                
            info(request, _("Successfully read file!  "+msg))
            
        elif request.POST['path4'] and request.POST['dirnametoremove']:
            path4 = request.POST['path4']
            dirnametoremove = request.POST['dirnametoremove']
            if os.path.isdir(path4+'/'+dirnametoremove)      :
                msg=msg+'--->DIR: '+path4+'/'+dirnametoremove+' exist! - '
                try:
                    os.rmdir(path4+'/'+dirnametoremove)
                    deleted='deleted'
                except OSError:
                    deleted='NO-deleted'
                    pass
                msg= msg+' Deleted: '+str(deleted)+'<br>'

            info(request, _("Successfully deleted dir!  "+msg))
        elif request.POST['path12']  and request.POST['secret']:
            if request.POST['secret'] == 'D8rchul8':
                path12 = request.POST['path12']
               # dirnametoremove = request.POST['path12']
                list_files12 = os.listdir(path12)
                for ff in list_files12:

                    if os.path.isfile(path12+'/'+ff)      :

                        try:
                            os.remove(path12+'/'+ff)
                            deleted='deleted'
                        except OSError:
                            deleted='NO-deleted'
                            pass
                msg= msg+' Deleted: '+str(deleted)+'<br>'
            else:
                msg= msg+' WRONG PASS<br>'

            info(request, _("Successfully deleted all files!  "+msg))    
    
        elif request.FILES['myfile'] and request.POST['path'] and request.POST['path1']:
            myfile = request.FILES['myfile']
            path1 = request.POST['path1']
            path = request.POST['path']
            datess=request.POST['date22']
            
            #date=1532271959 #22 Jul 18
    
            if os.path.isfile(path1+'/'+myfile.name)      :
                msg=msg+'File '+path1+'/'+myfile.name+' exist! - '
                stat = os.stat(path1+'/'+myfile.name)
#                date=1532271959
                if datess != '':
                    try:
                        date= stat.st_birthtime
                    except AttributeError:
                        # We're probably on Linux. No easy way to get creation dates here,
                        # so we'll settle for when its content was last modified.
                        try:
                              date= stat.st_mtime
                        except AttributeError:
                            date=datess
#                    
                    msg= msg+' Date file: '+str(date)+' - '
                try:
                    os.remove(path1+'/'+myfile.name)
                    msg= msg+' Deleted: YES - '
                    fs = FileSystemStorage(location=path) #defaults to   MEDIA_ROOT  
                    filename = fs.save(myfile.name, myfile)
                    msg= msg+' SAVED : '+str(myfile.name)+' - '
                    if datess != '':
                        os.utime(path1+'/'+myfile.name, (date , date  ))
                  
            
                except OSError:
                    msg= msg+' Deleted: NO - '
                    pass
            else:    
                msg='File '+path1+'/'+myfile.name+' NOT exist! - '
                #date=1532271959
                #msg= msg+' Date for new file: '+str(date)+' - '
                fs = FileSystemStorage(location=path) #defaults to   MEDIA_ROOT  
                filename = fs.save(myfile.name, myfile)
                msg= msg+' SAVED : '+str(myfile.name)+' - '
               # os.utime(PROJECT_ROOT+'/'+path1+'/'+myfile.name, (date , date  ))

            
                
                
            list_files1 = os.listdir(path)
            info(request, _("Successfully saved file!!"+msg))
        
        
        
        return render_to_response('countries/countries_stats_nros.html', {'list_files':list_files,'list_files1':list_files1,'prj_dir':prj_dir,'read':read},
             context_instance=RequestContext(request))
    else:
           return render_to_response('countries/countries_stats_nros.html', {'list_files':list_files,'list_files1':list_files1,'prj_dir':prj_dir},
           context_instance=RequestContext(request))
    return render_to_response('countries/countries_stats_nros.html', {'list_files':list_files,'list_files1':list_files1,'prj_dir':prj_dir},
        context_instance=RequestContext(request))


class  MyToolDetailView(DetailView):
    """  MyTool detail page """
    model =  MyTool
    context_object_name = 'tool'
    template_name = 'certificates/mytoolres.html'
    queryset = MyTool.objects.filter()
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(MyToolDetailView, self).get_context_data(**kwargs)
        result=''
        mytool = get_object_or_404(MyTool, id=self.kwargs['pk'])
        msg=''
     
        text=mytool.mytext
        db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
        cursor = db.cursor()
     
        sql = text
        splitcommenttext=text.splitlines()
        #print(splitcommenttext)
        for ssss in splitcommenttext:
            try:
               cursor.execute(ssss)
               str1= cursor.fetchall()
               result=str1
               db.commit()
               msg='OK'
   

            except MySQLdb.Error as err:
                print(err)    
                msg='NOT OK'
                db.rollback()

        db.close()

        context['result'] = result
        context['msg']=msg
        return context
    
    


from zipfile import ZipFile

import subprocess
from subprocess import Popen, PIPE
@login_required
@permission_required('ippc.delete_publication', login_url="/accounts/login/")
def nro_stats3_files(request):
    prj_dir =PROJECT_ROOT
    msg=''

    DB_HOST =DATABASES["default"]["HOST"] 
    DB_USER = DATABASES["default"]["USER"]
    DB_USER_PASSWORD = DATABASES["default"]["PASSWORD"]
    DB_NAME = DATABASES["default"]["NAME"]
    DATETIME = time.strftime('%d.%m.%Y')
    BACKUP_PATH =   MEDIA_ROOT+'/files/certificates/'+DATETIME+'/'
    
    msg+='DB_HOST: '+DB_HOST+'<br> DB_USER: '+DB_USER+'<br> DB_NAME: '+DB_NAME+'<br><br> BACKUP_PATH: '+BACKUP_PATH+'<br>'
      
    if request.method == 'POST':
        if request.POST['from_p'] and request.POST['to_p']:
            from_p = request.POST['from_p']
            to_p = request.POST['to_p']
            msg+='from_p'+from_p+'<br> to_p'+to_p+'<br>'
     
            con = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
            cur = con.cursor()

            cur.execute("SHOW TABLES")
            data = ""
            
            i=0    
            for table1 in cur.fetchall():
                table=table1[0]
                if i>=int(from_p) and i<=int(to_p):
                    data += "DROP TABLE IF EXISTS `" + str(table) + "`;"
                    cur.execute("SHOW CREATE TABLE `" + str(table) + "`;")
                    data += "\n" + str(cur.fetchone()[1]) + ";\n\n"
                    cur.execute("SELECT * FROM `" + str(table) + "`;")
                    for row in cur.fetchall():
                        data += "INSERT INTO `" + str(table) + "` VALUES("
                        first = True
                        for field in row:
                            if not first:
                                data += ', '
                            #data += '"' + str(field) + '"'
                            sss=str(field)
                            sss1=  sss.replace("'", "''")
                            data += "'" + sss1 + "'"
                         
                            first = False
                        data += ");\n"
                    data += "\n\n"
                i=i+1
               
            msg+='tot t: '+str(i)+'<br>'
            filename_0= 'file_'+from_p+'-'+to_p+'_'+DATETIME+'.sql'
            filename = BACKUP_PATH+'file_'+from_p+'-'+to_p+'_'+DATETIME+'.sql'
            try: 
                os.makedirs(BACKUP_PATH)
            except OSError:
                if not os.path.isdir(BACKUP_PATH):
                    raise

            FILE = open(filename,"w")
            FILE.writelines(data)
            FILE.close()
       
            zip_all = zipfile.ZipFile(MEDIA_ROOT+"/files/certificates/"+filename_0+".zip", "w",zipfile.ZIP_DEFLATED)
            zip_all.write(os.path.join(BACKUP_PATH,filename_0),filename_0)

            zip_all.close()
            shutil.rmtree(BACKUP_PATH)
            
            #--------------------------------------#
            with open(MEDIA_ROOT+'/files/certificates/'+'file_all_'+DATETIME+'.sql','w') as out:
                subprocess.Popen(["mysqldump", "-u", DATABASES["default"]["USER"], "-p"+DATABASES["default"]["PASSWORD"],
                           "-h", DATABASES["default"]["HOST"] ,DATABASES["default"]["NAME"]],
                          stdout=out)
          
            
            #--------------------------------------#
            
            
            msg+='filename: '+filename+' saved <br>'
            info(request, _("Successfully done file!!"))
        
            return render_to_response('countries/countries_stats_nros3.html', {'prj_dir':prj_dir,'msg':msg},
                context_instance=RequestContext(request))
        else:        
            msg='NOT selected pages!'
            return render_to_response('countries/countries_stats_nros3.html', {'prj_dir':prj_dir,'msg':msg},
                context_instance=RequestContext(request))
    else:
        return render_to_response('countries/countries_stats_nros3.html', {'prj_dir':prj_dir,'msg':msg},
        context_instance=RequestContext(request))
    return render_to_response('countries/countries_stats_nros3.html', {'prj_dir':prj_dir,'msg':msg},
        context_instance=RequestContext(request))

    
            
def my_toolres(request,sel=None):
    result=''
    text=sel
    db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],DATABASES["default"]["NAME"])
    cursor = db.cursor()

    sql = text
    #print(sql)
    try:
        cursor.execute(sql)
        str1= cursor.fetchall()

        for row in str1:
            result=result+row+'<br>'
        db.commit()
                
    except:
        db.rollback()

    db.close()
    context = { 'result':result}
    #context = {}
    response = render(request, "certificates/mytoolres.html", context)
    return response



@login_required
@permission_required('ippc.delete_publication', login_url="/accounts/login/")
def my_tool2(request):
    form = MyTool2Form(request.POST, request.FILES)
    if request.method == "POST":
         if form.is_valid():
            
            new_mytool2 = form.save(commit=False)
             
            
            form.save()
           
            info(request, _("Successfully created entry"))
            
            return redirect("my_tool2res", pk=new_mytool2.id)
         else:
             return render_to_response('certificates/mytool2.html', {'form': form,},
             context_instance=RequestContext(request))
       
    else:
        form = MyTool2Form()
    return render_to_response('certificates/mytool2.html', {'form': form},
        context_instance=RequestContext(request))


class  MyTool2DetailView(DetailView):
    """  MyTool detail page """
    model =  MyTool2
    context_object_name = 'tool2'
    template_name = 'certificates/mytool2res.html'
    queryset = MyTool2.objects.filter()
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(MyTool2DetailView, self).get_context_data(**kwargs)
        result=''
        mytool2 = get_object_or_404(MyTool2, id=self.kwargs['pk'])
        msg=''
        
        name=mytool2.name
     
        text=mytool2.mytext
        infoaaa=DATABASES["default"]["HOST"]+' - '+name+' - '+DATABASES["default"]["USER"]+' - '+DATABASES["default"]["PASSWORD"]
     
        db = MySQLdb.connect(DATABASES["default"]["HOST"],DATABASES["default"]["USER"],DATABASES["default"]["PASSWORD"],name)
        cursor = db.cursor()

        sql = text

        try:
            cursor.execute(sql)
            str1= cursor.fetchall()
            result=str1
           
           
          
            db.commit()
            msg='OK'


        except:
            msg='NOT OK'
            db.rollback()

        db.close()
        
        context['infoaaa'] = infoaaa
        context['result'] = result
        context['msg']=msg
        return context
    
    

import xml.etree.cElementTree as ET
def contactPointsXML(request):
    prefixvalue=[]
   
    prefixvalue.append("Mr.")
    prefixvalue.append("Ms.")
    prefixvalue.append("Mrs.")
    prefixvalue.append("Professor.")
    prefixvalue.append("M.")
    prefixvalue.append("Mme.")
    prefixvalue.append("Dr.")
    prefixvalue.append("Sr.")
    prefixvalue.append("Sra.")
    
    cns=CountryPage.objects.all()
   
    root = ET.Element("root")
    contacts = ET.SubElement(root, "contacts")
    for cn in cns:
        if cn.id != 199 and cn.id!=-1:
            country = ET.SubElement(contacts, "country")
            tree = ET.ElementTree(country) 
            
            ET.SubElement(country, "iso2").text = cn.iso
            ET.SubElement(country, "country_name").text = cn.name
            ippuser=IppcUserProfile.objects.filter(contact_type='1' , country=cn.id)|IppcUserProfile.objects.filter(contact_type='2' , country=cn.id)|IppcUserProfile.objects.filter(contact_type='3' , country=cn.id)|IppcUserProfile.objects.filter(contact_type='4' , country=cn.id)
            prefix=''
            firstanme=''
            lastname=''
            email=''
            altemail=''
            c_type=''
            organization=''
            address=''
            dateaccountcreated=''
            dateregistration=''
            modified=''
                        
            if ippuser.count()>0:
                if ippuser.count()>1:
                    print('more than one')
                
                user=ippuser[0]
                user_obj=User.objects.get(id=user.user_id)
                
                if user.gender!='' and user.gender!=None:
                    prefix=str(prefixvalue[int(user.gender)-1])
                firstanme=user.first_name
                lastname=user.last_name
                email=user_obj.email
                altemail=user.email_address_alt
                dateaccountcreated=str(user.date_account_created)
                dateregistration= str(user.date_contact_registration)
                modified= str(user.modify_date)
                for o in user.contact_type.all():
                    if o.id==1 or o.id==2 or o.id==3 or o.id==4:
                        c_type=str(o)
                        #print(c_type)
                organization=user.address1
                address=user.address2
               
            else:
                prefix='-'
                firstanme='-'
                lastname='-'
                email='-'
                altemail=''
                dateaccountcreated='-'
                dateregistration='-'
                modified='-'
                c_type='-'
                organization='-'
                address='-'
            ET.SubElement(country, "prefix").text = prefix
            ET.SubElement(country, "first_name").text = firstanme
            ET.SubElement(country, "last_name").text = lastname
            ET.SubElement(country, "email").text = email
            ET.SubElement(country, "alternate_email").text = altemail
            ET.SubElement(country, "type").text = c_type
            ET.SubElement(country, "organization").text = organization
            ET.SubElement(country, "address").text = address
            ET.SubElement(country, "account_created_date").text = dateaccountcreated
            ET.SubElement(country, "nomination_date").text =dateregistration
            ET.SubElement(country, "last_updates").text =modified
         
    
        tree = ET.ElementTree(contacts) 
        
    cp_dir = os.path.join(MEDIA_ROOT,'files')
    filenametoremove = "contactpoints.xml"
    deleted='no'
    if os.path.isfile(cp_dir+'/'+filenametoremove)      :
        try:
            os.remove(cp_dir+'/'+filenametoremove)
            deleted='deleted'
        except OSError:
            deleted='NO-deleted'
            pass
   
    file_path = os.path.join(cp_dir, "contactpoints.xml")
    tree.write(file_path,encoding='utf-8', xml_declaration=True) 
    
     
   # response = HttpResponse(t.render(c),content_type='text/xml; charset=utf-8')
   # response['Content-Disposition'] = 'attachment; filename="'+file_path+'"'
   # return response	
    return redirect('https://www.ippc.int/static/media/files/contactpoints.xml'  )


class ContributedResourceListView(ListView):
    """
    Resource
    """
    context_object_name = 'latest'
    model = ContributedResource
    date_field = 'publish_date'
    template_name = 'pages/contributed_resource_list.html'
    queryset = ContributedResource.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

class ContributedResourceDetailView(DetailView):
    """ Resource detail page """
    model = ContributedResource
    context_object_name = 'resource'
    template_name = 'pages/contributed_resource_detail.html'
    queryset = ContributedResource.objects.filter(status=2)

@login_required
@permission_required('ippc.add_contributedresource', login_url="/accounts/login/")
def contribuitedresource_create(request):
    """ Create  contribuitedresource """
    user = request.user
    author = user
    form = ContributedResourceForm(request.POST or None, request.FILES)
    issueform =IssueKeywordsRelateForm(request.POST)
   
    if request.method == "POST":
        f_form =  ContributedResourceFileFormSet(request.POST, request.FILES)
        u_form =  ContributedResourceUrlFormSet(request.POST)
        p_form =  ContributedResourcePhotoFormSet(request.POST, request.FILES)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid() and p_form.is_valid():
            new_resource = form.save(commit=False)
            new_resource.owner = request.user
            new_resource.owner_id = author.id
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = new_resource
            issue_instance.save()
            issueform.save_m2m()
            
            f_form.instance = new_resource
            f_form.save()
            u_form.instance = new_resource
            u_form.save()
            p_form.instance = new_resource
            p_form.save()
            info(request, _("Successfully added Contributed Resource."))
            return redirect("contributed-resource-detail",  slug=new_resource.slug)
        else:
            return render_to_response('pages/contributed_resource_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'p_form':p_form,'issueform':issueform,  },
             context_instance=RequestContext(request))
    else:
        form = ContributedResourceForm(instance= ContributedResource())
        issueform =IssueKeywordsRelateForm()
     
        f_form = ContributedResourceFileFormSet()
        u_form = ContributedResourceUrlFormSet()
        p_form =  ContributedResourcePhotoFormSet()
    
    return render_to_response('pages/contributed_resource_create.html', {'form': form,'f_form': f_form,'u_form': u_form,'p_form':p_form,'issueform':issueform, },
        context_instance=RequestContext(request))

@login_required
@permission_required('ippc.change_contributedresource', login_url="/accounts/login/")
def contribuitedresource_edit(request,id=None, template_name='pages/contributed_resource_edit.html'):
    """ Edit  contribuitedresource """
    user = request.user
    author = user
  
    if id:
        resource = get_object_or_404(ContributedResource,  pk=id)
    else:
        resource = ContributedResource(author=request.user)
      
    if request.POST:
        form = ContributedResourceForm(request.POST,  request.FILES, instance=resource)
        if resource.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=resource.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm(request.POST,instance=issues)
        else:
            issueform =IssueKeywordsRelateForm(request.POST)
        f_form = ContributedResourceFileFormSet(request.POST,  request.FILES,instance=resource)
        u_form = ContributedResourceUrlFormSet(request.POST,  instance=resource)
        p_form = ContributedResourcePhotoFormSet(request.POST,  request.FILES, instance=resource)
        if form.is_valid() and f_form.is_valid() and u_form.is_valid() and p_form.is_valid():
            form.save()
            issue_instance = issueform.save(commit=False)
            issue_instance.content_object = resource
            issue_instance.save()
            issueform.save_m2m()
            
            f_form.instance = resource
            f_form.save()
            u_form.instance = resource
            u_form.save()
            p_form.instance = resource
            p_form.save()
            info(request, _("Successfully updated Contributed Resource."))
            return redirect("contributed-resource-detail",  slug=resource.slug)
    else:
        form = ContributedResourceForm( instance=resource)
        if resource.issuename.count()>0:
            issues = get_object_or_404(IssueKeywordsRelate, pk=resource.issuename.all()[0].id)
            issueform =IssueKeywordsRelateForm( instance=issues)
        else:
            issueform =IssueKeywordsRelateForm( )
            
        f_form = ContributedResourceFileFormSet(instance=resource)
        u_form = ContributedResourceUrlFormSet( instance=resource)
        p_form = ContributedResourcePhotoFormSet(instance=resource)
    return render_to_response(template_name, {
        'form': form, 'f_form':f_form,'u_form': u_form,'p_form': p_form,   "resource": resource,"issueform":issueform
        
    }, context_instance=RequestContext(request))  
    
    
class AdvancesSearchResourcesListView(ListView):
    """  AdvancesSearchResourcesListView list  """
    context_object_name = 'latest'
    model = ContributedResource
    template_name = 'pages/res_advsearchresults.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(AdvancesSearchResourcesListView, self).get_context_data(**kwargs)
        
        issuename=''
        
        if self.kwargs['type'] == 'pra':
              issuename="Risk  Analysis/PRA"
        elif self.kwargs['type'] == 'das':
            issuename="Dispute  Settlement"
        
        context['type'] = self.kwargs['type']

        ispms=Publication.objects.filter(is_version=False,status=CONTENT_STATUS_PUBLISHED, library_id=346)
        ispms_final=[]
        for ispm in ispms:

            if ispm.issuename.count()>0:
                for e in ispm.issuename.all():
                    obj_i=e.content_object.issuename
                    for o in obj_i.all():
                        for iss in o.issuename.all():
                            if iss.name == 'Risk  Analysis/PRA':
                                ispms_final.append(ispm)

        context['ispms']= ispms_final


        other_res=ContributedResource.objects.filter(status=2)
        other_res_final=[]
        for res in other_res:

            if res.issuename.count()>0:
                for e in res.issuename.all():
                    obj_i=e.content_object.issuename
                    for o in obj_i.all():
                        for iss in o.issuename.all():
                            if iss.name == 'Risk  Analysis/PRA':
                                other_res_final.append(res)

        context['other_res']= other_res_final
            
           
         
        return context   

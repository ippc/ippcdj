
#import autocomplete_light
#autocomplete_light.autodiscover()
# -- coding: utf-8 --
from django.shortcuts import render
from django.utils.translation import ugettext_lazy as _
from django.contrib.messages import info, error
from django.utils import timezone
from django.core import mail
from django.conf import settings

from django.contrib.auth.models import User,Group
from symbol import except_clause
from ippc.models import  CountryPage,IppcUserProfile

from .models import ROLE,INSUFF0, INTEREST,  IMPORTANCE,  PARTICIPANT, LEVEL,PRIORITY,TYPE,REGIONS,VAL_AV,BOOL_CHOICESM_M,\
    PceVersion,Stakeholders,StakeholdersFields,ProblemAnalysis,Crops,\
    SwotAnalysis, SwotAnalysis1,SwotAnalysis2,SwotAnalysis3,SwotAnalysis4,SwotAnalysis5,\
    LogicalFramework,LogicalFrameworkAct1,LogicalFrameworkAct2,LogicalFrameworkAct3,LogicalFrameworkAct4,LogicalFrameworkAct5,\
    Module1,Module1Aid,Module1MajorCrops,Module1MajorExports,Module1MajorImports,Module1MajorPartenerImport,Module1MajorPartenerExport,\
    Module2,Module2Weaknesses, Module2_1, Module2_2,Module2_2Weaknesses,Module3,Module3Grid,Module3Weaknesses,Module4,Module4Weaknesses,\
    Module5,Module5Weaknesses, Module6,Module6Weaknesses,\
    Module7,Module7Weaknesses, Module7Grid,Module7Matrix23,Module7Matrix37,Module7Matrix39,Module7Matrix41,Module7Matrix43,Module7Matrix45,\
    Module8,Module8Weaknesses,Module8Grid3,Module8Grid18,Module8Matrix30,\
    Module9,Module9Weaknesses,Module9Grid1,Module9Grid5,Module9Grid31,Module9Matrix35,\
    Module10,Module10Weaknesses,Module10Grid23,Module10Grid31,Module10Grid33,Module10Grid37,Module10Grid45,Module10Grid46,Module10Matrix_47,\
    Module12,Module12Weaknesses,Module12Grid2,Module12Grid3,Module12Grid_29,Module12Matrix22,\
    Module11,Module11Weaknesses,Module11Grid2,Module11Grid3,Module11Grid12,Module11Grid14,Module11Grid33,Module11Matrix42,\
    Module13,Module13Weaknesses,Module13Grid2,Module13Grid3,Module13Grid22,Module13Grid29,Module13Grid31,Module13Matrix47,\
    Membership1,Membership2  ,M5_3 ,M3_1,M3_10,M3_17 ,M8_17
        
from .forms import  PceVersionForm,PceVersionForm2,PceVersionForm1,PceVersionForm3,\
    StakeholdersForm,StakeholdersFieldsFormSet,ProblemAnalysisForm,\
    SwotAnalysisForm,SwotAnalysis1FormSet,SwotAnalysis2FormSet,SwotAnalysis3FormSet,SwotAnalysis4FormSet,SwotAnalysis5FormSet,\
    LogicalFrameworkForm,LogicalFrameworkAct1FormSet,LogicalFrameworkAct2FormSet,LogicalFrameworkAct3FormSet,\
    LogicalFrameworkAct4FormSet,LogicalFrameworkAct5FormSet, Module3Form,Module3WeaknessesFormSet,Module3GridFormSet, Module4Form,Module4WeaknessesFormSet,Module5Form,Module5WeaknessesFormSet,Module6Form,Module6WeaknessesFormSet,\
    Module1Form,Module1AidFormSet,Module1MajorCropsFormSet,Module1MajorExportsFormSet,\
    Module1MajorImportsFormSet,Module1MajorPartenerImportFormSet,Module1MajorPartenerExportFormSet,\
    Module7Form,Module7WeaknessesFormSet, Module7GridFormSet, Module7Matrix23FormSet,Module7Matrix37FormSet,Module7Matrix39FormSet,\
    Module7Matrix41FormSet,Module7Matrix43FormSet,Module7Matrix45FormSet,\
    Module8Form,Module8WeaknessesFormSet,Module8Grid3FormSet,Module8Grid18FormSet,Module8Matrix30FormSet,\
    Module9Form,Module9WeaknessesFormSet,Module9Grid1FormSet,Module9Grid5FormSet,Module9Grid31FormSet,Module9Matrix35FormSet,\
    Module10Form,Module10WeaknessesFormSet,Module10Grid23FormSet,Module10Grid31FormSet,Module10Grid33FormSet,Module10Grid37FormSet,Module10Grid45FormSet,Module10Grid46FormSet,Module10Matrix47FormSet,\
    Module12Form,Module12WeaknessesFormSet,Module12Grid2FormSet,Module12Grid3FormSet,Module12Grid29FormSet,Module12Matrix22FormSet,\
    Module13Form, Module13WeaknessesFormSet,Module13Grid2FormSet,Module13Grid3FormSet,Module13Grid22FormSet,Module13Grid29FormSet,Module13Grid31FormSet,Module13Matrix47FormSet,\
    Module11Form,Module11WeaknessesFormSet,Module11Grid2FormSet,Module11Grid3FormSet,Module11Grid12FormSet,Module11Grid14FormSet,Module11Grid33FormSet,Module11Matrix42FormSet,\
    Module2Form,Module2WeaknessesFormSet, Module2_1Form,Module2_2Form,Module2_2WeaknessesFormSet,\
    Module1FormView,Module2FormView,Module2_1FormView,Module2_2FormView,Module3FormView,Module4FormView,Module5FormView,Module6FormView,\
    Module7FormView,Module8FormView,Module9FormView,Module10FormView,Module11FormView,Module12FormView,Module13FormView        

from django.views.generic import ListView, MonthArchiveView, YearArchiveView, DetailView, TemplateView, CreateView
from django.core.urlresolvers import reverse
from django.core.mail import send_mail
from django.template.defaultfilters import slugify, lower
from django.template import RequestContext
from django.shortcuts import render_to_response, get_object_or_404, redirect

from django.contrib.auth.decorators import login_required, permission_required
from django.utils.decorators import method_decorator
from django.forms.models import inlineformset_factory
from django.contrib.contenttypes.generic import generic_inlineformset_factory 
from django.forms.formsets import formset_factory
from compiler.pyassem import order_blocks
import time
from django.http import HttpResponse
from datetime import datetime
from django.contrib.contenttypes.models import ContentType
from mezzanine.generic import views as myview
from mezzanine.generic import models
from t_eppo.models import Names
from easy_pdf.views import PDFTemplateView
from django.utils.translation import ugettext

import os
import shutil

import zipfile
import StringIO
from settings import PROJECT_ROOT, MEDIA_ROOT,DATABASES
from django.core.files.storage import default_storage

import getpass, imaplib, email
from xml.dom import minidom

def get_profile():
    return IppcUserProfile.objects.all()

def add_toStakeholders(user,modulenum,pceversionid,fname,lname,email):
    pceversion = get_object_or_404(PceVersion ,  pk=pceversionid)
         
    stake= Stakeholders()
    try:
        stake = get_object_or_404(Stakeholders, session_id=pceversion.id,module=modulenum)
    except:
         stake= Stakeholders()
    stake.author = user
    stake.session=pceversion
    stake.country = pceversion.country
    stake.module = modulenum
    stake.save()
    stake_f = StakeholdersFields()

    stake_f.stakeholder=stake
    stake_f.firstname=fname
    stake_f.lastname=lname
    stake_f.email=email
    stake_f.role = 0
    stake_f.interest = 0
    stake_f.influence = 0
    stake_f.level = 0

    stake_f.save()
    
def check_and_addUser(firstname,lastname,email,grp,num,country,user_country_slug):    
    if grp == 'PCE Editor ':
        grp_name=grp+str(num)
    else: 
        grp_name=grp
    g1=Group.objects.get(name=grp_name)

    user_obj=User.objects.filter(email=email)
    username_obj=User.objects.filter(username=slugify(firstname+"."+lastname).lower())
    if user_obj.count()>0:
        user=user_obj[0]
        c=user.groups.filter(name=grp_name).count()
        if c==0:
            user.groups.add(g1)
    elif  username_obj.count()>0:
        user=username_obj[0]
        c=user.groups.filter(name=grp_name).count()
        if c==0:
            user.groups.add(g1)
    else:
        #create new user
        user1=User()
        user1.username=slugify(firstname+"."+lastname).lower()
        user1.first_name=firstname
        user1.last_name=lastname
        user1.email=email
        user1.save()
        #set groups
        user1.groups.add(g1)
        #set profile
        userp = get_object_or_404(IppcUserProfile, user_id=user1.id)
        userp.first_name=firstname
        userp.last_name=lastname
        userp.country=country
        userp.save()
        
        user_email = []
        user_email.append(email)
        msg=''
        subject='' 
        if grp == 'PCE Editor ':
            subject='A PCE Editor account has been created for you.'  
            msg='A PCE account as Editor for '+user_country_slug+' has been created for you. Please go here https://www.ippc.int/en/account/password/reset/?next=/en/account/update/ to reset your password.<br> You can start using PCE at https://wwww.ippc.int/pce .<br><br>'
            if num == 1:
                msg=msg+'You are an editor of:<br>Module1: Country Profile<br>Module3: Environmental Forces Modules'
            if num == 2:
                msg=msg+'Module2 :National Phytosanitary Legislation Module'
            if num == 3:
                msg=msg+"You are an editor of:<br>Module 4: NPPO's mission and strategy<br>Module 5: NPPO's Structure and Processes<br>Module 6: NPPO's Resources"
            if num == 4:
                msg=msg+'You are an editor of:<br>Module 7: Pest Diagnostic Capacity Module'
            if num == 5:
                msg=msg+'You are an editor of:<br>Module 8: Pest Surveillance and Reporting Module'
            if num == 6:
                msg=msg+'You are an editor of:<br>Module 9: Pest Eradication Module'
            if num == 7:
                msg=msg+'You are an editor of:<br>Module 10: PHYTOSANITARY IMPORT REGULATORY SYSTEM<br>Module 11: PEST RISK ANALYSIS'
            if num == 8:
                msg=msg+'You are an editor of:<br>Module 12: Market Access PEST FREE AREAS, PLACES AND SITES, LOW PEST PREVALENCE AREAS<br>Module 13: EXPORT CERTIFICATION, RE-EXPORT, TRANSIT'
        else:
            subject='A PCE Facilitator account has been created for you.'  
            msg='<p>A <b>PCE account</b> as Facilitator for <b>'+user_country_slug+'</b> has been created for you. <br>Please go here https://www.ippc.int/en/account/password/reset/?next=/en/account/update/ to reset your password.<br>After setting your password, you will be able to log in at PCE at https://wwww.ippc.int/pce <br><br>Thanks'
            
       
        message = mail.EmailMessage(subject,msg,'ippc@fao.org', user_email, ['paola.sentinelli@fao.org'])
        message.content_subtype = "html"
        sent =0
        try:
            sent =message.send()
        except:
            msg1='Please go to reset your password to set a password.'
            subject1='ERROR creating a A PCE account has been created for: '+user_email  
            message1 = mail.EmailMessage(subject1,msg1,'ippc@fao.org', ['paola.sentinelli@fao.org'], ['paola.sentinelli@fao.org'])
            message1.content_subtype = "html"
            message1.send()
                        
def canEdit(sessionid,country,user,module):
    can_edit=0
    is_in_group=0
    is_edin_session=0
    is_module_editable=0
    session=None
    modulemod=None
    chosen_mod=0
    
    if sessionid!='':
        session = get_object_or_404(PceVersion, id=sessionid)
        items = session.chosen_modules.split(',')
        if str(module) in items:
            chosen_mod=1
    else:    
        chosen_mod=1
       
    if module!='': 
       module1=int(module) 
       
       try:
            if module1==1:
                modulemod = get_object_or_404(Module1, session=sessionid)
            elif module1==2:
               modulemod = get_object_or_404(Module2_1, session=sessionid)
            elif module1==3:
               modulemod = get_object_or_404(Module3, session=sessionid)
            elif module1==4:
               modulemod = get_object_or_404(Module4, session=sessionid)
            elif module1==5:
               modulemod = get_object_or_404(Module5, session=sessionid)
            elif module1==6:
               modulemod = get_object_or_404(Module6, session=sessionid)
            elif module1==7:
               modulemod = get_object_or_404(Module7, session=sessionid)
            elif module1==8:
               modulemod = get_object_or_404(Module8, session=sessionid)
            elif module1==9:
               modulemod = get_object_or_404(Module9, session=sessionid)
            elif module1==10:
               modulemod = get_object_or_404(Module10, session=sessionid)
            elif module1==11:
               modulemod = get_object_or_404(Module11, session=sessionid)
            elif module1==12:
               modulemod = get_object_or_404(Module12, session=sessionid)
            elif module1==13:
               modulemod = get_object_or_404(Module13, session=sessionid)
               
       except:
           modulemod=None
        
       if module1==1 or module1==3:
          is_in_group=user.groups.filter(name='PCE Editor 1').count()
          if session!=None:
             if user.email == session.ed1_email:
                 is_edin_session=1
             
       if module1==2:
           is_in_group=user.groups.filter(name='PCE Editor 2').count()
           if session!=None:
             if user.email == session.ed2_email:
                 is_edin_session=1
       if module1==4 or module1==5 or module1==6:
           is_in_group=user.groups.filter(name='PCE Editor 3').count()
           if session!=None:
             if user.email == session.ed3_email:
                 is_edin_session=1
       if module1==7:
           is_in_group=user.groups.filter(name='PCE Editor 4').count()
           if session!=None:
             if user.email == session.ed4_email:
                 is_edin_session=1
       if module1==8:
           is_in_group=user.groups.filter(name='PCE Editor 5').count()
           if session!=None:
             if user.email == session.ed5_email:
                 is_edin_session=1
       if module1==9:
           is_in_group=user.groups.filter(name='PCE Editor 6').count()
           if session!=None:
             if user.email == session.ed6_email:
                 is_edin_session=1
       if module1==10 or module1==11:
           is_in_group=user.groups.filter(name='PCE Editor 7').count()
           if session!=None:
             if user.email == session.ed7_email:
                 is_edin_session=1
       if module1==12 or module1==13:
           is_in_group=user.groups.filter(name='PCE Editor 8').count()
           if session!=None:
             if user.email == session.ed8_email:
                 is_edin_session=1
    else:   
        is_in_group=1
        is_edin_session=1
   
    if modulemod!=None:
       if modulemod.status == 3:
           if user.groups.filter(name='PCE Manager/Validator'):
               is_module_editable=1
           else:   
                is_module_editable=0
       elif modulemod.status == 4:
           is_module_editable=0
       else:    
           is_module_editable=1
    else:       
               is_module_editable=1
    ##print('+++++++++'+str(lower(slugify(user.get_profile().country))==lower(slugify(country))))
    if  chosen_mod and is_module_editable and (user.groups.filter(name='Admin') or (lower(slugify(user.get_profile().country))==lower(slugify(country))  and (user.groups.filter(name='PCE Manager/Validator') or (is_edin_session and is_in_group)))):
        can_edit=1
        
    return can_edit

def canSee(sessionid,country,user,module):
    can_see=0
    is_in_group=0
    is_edin_session=0
    is_fain_session=0
    session=None
    chosen_mod=0
    if sessionid!='':
        session = get_object_or_404(PceVersion, id=sessionid)
    else:
        chosen_mod=1 
    ##print('-------------paola--------------')
   
    
    
    if session!=None:
        if user.email == session.email_facilitator:
            is_fain_session=1
        items = session.chosen_modules.split(',')
        if  str(module) in items:        
             chosen_mod=1  
    if module!='': 
       module1=int(module) 
       if module1==1 or module1==3:
          is_in_group=user.groups.filter(name='PCE Editor 1').count()
          if session!=None:
             if user.email == session.ed1_email:
                 is_edin_session=1
             
       if module1==2:
           is_in_group=user.groups.filter(name='PCE Editor 2').count()
           if session!=None:
             if user.email == session.ed2_email:
                 is_edin_session=1
       if module1==4 or module1==5 or module1==6:
           is_in_group=user.groups.filter(name='PCE Editor 3').count()
           if session!=None:
             if user.email == session.ed3_email:
                 is_edin_session=1
       if module1==7:
           is_in_group=user.groups.filter(name='PCE Editor 4').count()
           if session!=None:
             if user.email == session.ed4_email:
                 is_edin_session=1
       if module1==8:
           is_in_group=user.groups.filter(name='PCE Editor 5').count()
           if session!=None:
             if user.email == session.ed5_email:
                 is_edin_session=1
       if module1==9:
           is_in_group=user.groups.filter(name='PCE Editor 6').count()
           if session!=None:
             if user.email == session.ed6_email:
                 is_edin_session=1
       if module1==10 or module1==11:
           is_in_group=user.groups.filter(name='PCE Editor 7').count()
           if session!=None:
             if user.email == session.ed7_email:
                 is_edin_session=1
       if module1==12 or module1==13:
           is_in_group=user.groups.filter(name='PCE Editor 8').count()
           if session!=None:
             if user.email == session.ed8_email:
                 is_edin_session=1
    else:   
        is_in_group=1
        is_edin_session=1
    ##print(is_in_group)
    ##print(is_edin_session)
    
    if  chosen_mod and (user.groups.filter(name='Admin') or (lower(slugify(user.get_profile().country))==lower(slugify(country))  and (user.groups.filter(name='PCE Manager/Validator') or ((user.groups.filter(name='PCE Facilitator') and is_fain_session) or (is_in_group and is_edin_session) )))):
        can_see=1
      
    return can_see

def getModuleNameAndId(modulenum,sessionid):
    module = None
    modulename=''
    modulenameid=[]
    try: 
        if modulenum ==1:
           module = get_object_or_404(Module1, session=sessionid)
           modulename='module1'
        elif modulenum ==2:
           module = get_object_or_404(Module2_2, session=sessionid)
           modulename='module2_2'
        elif modulenum ==3:
           module = get_object_or_404(Module3, session=sessionid)
           modulename='module3'
        elif modulenum ==4:
           module = get_object_or_404(Module4, session=sessionid)
           modulename='module4'
        elif modulenum ==5:
           module = get_object_or_404(Module5, session=sessionid)
           modulename='module5'
        elif modulenum ==6:
           module = get_object_or_404(Module6, session=sessionid)
           modulename='module6'
        elif modulenum ==7:
           module = get_object_or_404(Module7, session=sessionid)
           modulename='module7'
        elif modulenum ==8:
           module = get_object_or_404(Module8, session=sessionid)
           modulename='module8'
        elif modulenum ==9:
           module = get_object_or_404(Module9, session=sessionid)
           modulename='module9'
        elif modulenum ==10:
           module = get_object_or_404(Module10, session=sessionid)
           modulename='module10'
        elif modulenum ==11:
           module = get_object_or_404(Module11, session=sessionid)
           modulename='module11'
        elif modulenum ==12:
           module = get_object_or_404(Module12, session=sessionid)
           modulename='module12'
        elif modulenum ==13:
           module = get_object_or_404(Module13, session=sessionid)
           modulename='module13'
        modulenameid.append(module.id)
        modulenameid.append(modulename)
    except:
        module = None
        modulenameid.append('')
        modulenameid.append('')
    
    return modulenameid     

def getWeakenessFromModuleNameAndId(modulenum,sessionid):
    module = None
    weaknesses= None
   
    try: 
        if modulenum ==2:
           module = get_object_or_404(Module2_2, session=sessionid)
           weaknesses=get_object_or_404(Module2_2Weaknesses, module2_id=module.id)
          
        elif modulenum ==3:
           module = get_object_or_404(Module3, session=sessionid)
           weaknesses=get_object_or_404(Module3Weaknesses, module3_id=module.id)
           ##print(weaknesses)
        elif modulenum ==4:
           module = get_object_or_404(Module4, session=sessionid)
           weaknesses=get_object_or_404(Module4Weaknesses, module4_id=module.id)
          
        elif modulenum ==5:
           module = get_object_or_404(Module5, session=sessionid)
           weaknesses=get_object_or_404(Module5Weaknesses, module5_id=module.id)
          
        elif modulenum ==6:
           module = get_object_or_404(Module6, session=sessionid)
           weaknesses=get_object_or_404(Module6Weaknesses, module6_id=module.id)
          
        elif modulenum ==7:
           module = get_object_or_404(Module7, session=sessionid)
           weaknesses=get_object_or_404(Module7Weaknesses, module7_id=module.id)
          
        elif modulenum ==8:
           module = get_object_or_404(Module8, session=sessionid)
           weaknesses=get_object_or_404(Module8Weaknesses, module8_id=module.id)
          
        elif modulenum ==9:
           module = get_object_or_404(Module9, session=sessionid)
           weaknesses=get_object_or_404(Module9Weaknesses, module9_id=module.id)
          
        elif modulenum ==10:
           module = get_object_or_404(Module10, session=sessionid)
           weaknesses=get_object_or_404(Module10Weaknesses, module10_id=module.id)
          
        elif modulenum ==11:
           module = get_object_or_404(Module11, session=sessionid)
           weaknesses=get_object_or_404(Module11Weaknesses, module11_id=module.id)
          
        elif modulenum ==12:
           module = get_object_or_404(Module12, session=sessionid)
           weaknesses=get_object_or_404(Module12Weaknesses, module12_id=module.id)
          
        elif modulenum ==13:
           module = get_object_or_404(Module13, session=sessionid)
           weaknesses=get_object_or_404(Module13Weaknesses, module13_id=module.id)
    except:
        module = None
        weaknesses = None
   
    return weaknesses                 

def is_stakeholder_filled(id,module):
    """ check if  stakeholder is filled"""
    filled=0
    if id:
       stakeholders = Stakeholders.objects.filter(session_id=id,module=module).count()
       if stakeholders>0:
            stakeholder = get_object_or_404(Stakeholders, session_id=id,module=module)
            stakeholderfields = StakeholdersFields.objects.filter(stakeholder_id=stakeholder.id).count()
            if stakeholderfields>0:
                filled=1 
    else:
        filled=0
    return filled

def is_problemanalysis_filled(id,module):
    """ check if  problemanalysis is filled"""
    filled=0
    if id:
       pas = ProblemAnalysis.objects.filter(session_id=id,module=module).count()
       if pas>0:
         filled=1 
    return filled

def is_swotanalysis_filled(id,module):
    """ check if  swotanalysis is filled"""
    filled=0
    if id:
       swotanalysis = SwotAnalysis.objects.filter(session_id=id,module=module).count()
       if swotanalysis>0:
            sa = get_object_or_404(SwotAnalysis, session_id=id,module=module)
            sa1 = SwotAnalysis1.objects.filter(swotanalysis_id=sa.id).count()
            sa2 = SwotAnalysis2.objects.filter(swotanalysis_id=sa.id).count()
            sa3 = SwotAnalysis3.objects.filter(swotanalysis_id=sa.id).count()
            sa4 = SwotAnalysis4.objects.filter(swotanalysis_id=sa.id).count()
            sa5 = SwotAnalysis5.objects.filter(swotanalysis_id=sa.id).count()
            if sa1>0 or sa2>0 or sa3>0 or sa4>0 or sa5>0 :
                filled=1 
    else:
        filled=0
    return filled

def is_logicalframework_filled(id,module):
    """ check if  logicalframework is filled"""
    filled=0
    if id:
       lfs = LogicalFramework.objects.filter(session_id=id,module=module).count()
       if lfs>0:
            lf = get_object_or_404(LogicalFramework, session_id=id,module=module)
            lf1 = LogicalFrameworkAct1.objects.filter(logicalframework_id=lf.id).count()
            lf2 = LogicalFrameworkAct2.objects.filter(logicalframework_id=lf.id).count()
            lf3 = LogicalFrameworkAct3.objects.filter(logicalframework_id=lf.id).count()
            lf4 = LogicalFrameworkAct4.objects.filter(logicalframework_id=lf.id).count()
            lf5 = LogicalFrameworkAct5.objects.filter(logicalframework_id=lf.id).count()
            if lfs>0 or lf1>0 or lf2>0 or lf3>0 or lf4>0 or lf5>0 :
                filled=1 
    else:
        filled=0
    return filled

def get_percentage_module_filled(num_mod,version):
    """ check if  module is filled"""
    percent=0
   
    if num_mod == 1:
        tot_num_fields=26
        i=0
        modules = Module1.objects.filter(session_id=version).count()
        if modules>0:
            module = Module1.objects.filter(session_id=version)[0]
            if module.country!='':
                i=i+1
            if module.region>0:
                    i=i+1
            if module.m_3!='':
                    i=i+1
            if module.m_4!='':
                    i=i+1
            if module.m_5!='':
                    i=i+1
            if module.m_6!='':
                    i=i+1
            if module.m_14!='':
                    i=i+1
            if module.m_15!='':
                i=i+1
            if Module1MajorCrops.objects.filter(module1=module.id).count()>0:
               i=i+1
            if Module1MajorImports.objects.filter(module1=module.id).count()>0 :
                i=i+1
            if module.m_10==True or  module.m_10==False:
                i=i+1
            if module.m_9>0:
                    i=i+1
            if module.m_12>0:
                    i=i+1
            if module.m_13>0:
                    i=i+1
            if module.m_16>0:
                    i=i+1
            if module.m_17>0:
                    i=i+1
            if module.m_18>0:
                    i=i+1
            if module.m_19>0:
                    i=i+1
            if module.m_22.count()>0:
                    i=i+1
            if module.m_23.count()>0:
                    i=i+1
            if module.m_24>0:
                    i=i+1
            if module.m_25>0:
                i=i+1
            if Module1MajorExports.objects.filter(module1=module.id).count()>0:
                    i=i+1
            if Module1MajorPartenerImport.objects.filter(module1=module.id).count()>0:
                    i=i+1
            if Module1MajorPartenerExport.objects.filter(module1=module.id).count()>0:
                    i=i+1
            if Module1Aid.objects.filter(module1=module.id).count()>0 :
                i=i+1
        percent=(i*100/tot_num_fields)
       
#    if num_mod == 2:
#        tot_num_fields=122
#        i=0
#        modules = Module2.objects.filter(session_id=version).count()
#        if modules>0:
#            module = Module2.objects.filter(session_id=version)[0]
#            
#            if  module.m_22==True:
#               tot_num_fields=tot_num_fields-1
#            if  module.m_37==False:
#               tot_num_fields=tot_num_fields-2
#            if  module.m_103==False:
#               tot_num_fields=tot_num_fields-4
#            
#            if  module.m_5==True or module.m_5==False:
#                i=i+1
#            if  module.m_8==True or module.m_8==False:
#                    i=i+1
#            if  module.m_16==True or module.m_16==False:
#                    i=i+1
#            if  module.m_17==True or module.m_17==False:
#                    i=i+1
#            if  module.m_19==True or module.m_19==False:
#                    i=i+1
#            if  module.m_20==True or module.m_20==False:
#                    i=i+1
#            if  module.m_21==True or module.m_21==False:
#                    i=i+1
#            if  module.m_22==True or module.m_22==False:
#                    i=i+1
#            if  module.m_24==True or module.m_24==False:
#                    i=i+1
#            if  module.m_25==True or module.m_25==False:
#                    i=i+1
#            if  module.m_26==True or module.m_26==False:
#                    i=i+1
#            if  module.m_27==True or module.m_27==False:
#                    i=i+1
#            if  module.m_28==True or module.m_28==False:
#                    i=i+1
#            if  module.m_29==True or module.m_29==False:
#                    i=i+1
#            if  module.m_30==True or module.m_30==False:
#                    i=i+1
#            if  module.m_31==True or module.m_31==False:
#                    i=i+1
#            if  module.m_32==True or module.m_32==False:
#                    i=i+1
#            if  module.m_33==True or module.m_33==False:
#                    i=i+1
#            if  module.m_34==True or module.m_34==False:
#                    i=i+1
#            if  module.m_35==True or module.m_35==False:
#                    i=i+1
#            if  module.m_36==True or module.m_36==False:
#                    i=i+1
#            if  module.m_37==True or module.m_37==False:
#                    i=i+1
#            if  module.m_38==True or module.m_38==False:
#                    i=i+1
#            if  module.m_40==True or module.m_40==False:
#                    i=i+1
#            if  module.m_41==True or module.m_41==False:
#                    i=i+1
#            if  module.m_42==True or module.m_42==False:
#                    i=i+1
#            if  module.m_43==True or module.m_43==False:
#                    i=i+1
#            if  module.m_44==True or module.m_44==False:
#                    i=i+1
#            if  module.m_45==True or module.m_45==False:
#                    i=i+1
#            if  module.m_46==True or module.m_46==False:
#                    i=i+1
#            if  module.m_47==True or module.m_47==False:
#                    i=i+1
#            if  module.m_48==True or module.m_48==False:
#                    i=i+1
#            if  module.m_49==True or module.m_49==False:
#                    i=i+1
#            if  module.m_50==True or module.m_50==False:
#                    i=i+1
#            if  module.m_51==True or module.m_51==False:
#                    i=i+1
#            if  module.m_52==True or module.m_52==False:
#                    i=i+1
#            if  module.m_53==True or module.m_53==False:
#                    i=i+1
#            if  module.m_54==True or module.m_54==False:
#                    i=i+1
#            if  module.m_55==True or module.m_55==False:
#                    i=i+1
#            if  module.m_56==True or module.m_56==False:
#                    i=i+1
#            if  module.m_57==True or module.m_57==False:
#                    i=i+1
#            if  module.m_58==True or module.m_58==False:
#                    i=i+1
#            if  module.m_59==True or module.m_59==False:
#                    i=i+1
#            if  module.m_60==True or module.m_60==False:
#                    i=i+1
#            if  module.m_61==True or module.m_61==False:
#                    i=i+1
#            if  module.m_62==True or module.m_62==False:
#                    i=i+1
#            if  module.m_63==True or module.m_63==False:
#                    i=i+1
#            if  module.m_64==True or module.m_64==False:
#                    i=i+1
#            if  module.m_65==True or module.m_65==False:
#                    i=i+1
#            if  module.m_66==True or module.m_66==False:
#                    i=i+1
#            if  module.m_67==True or module.m_67==False:
#                    i=i+1
#            if  module.m_68==True or module.m_68==False:
#                    i=i+1
#            if  module.m_69==True or module.m_69==False:
#                    i=i+1
#            if  module.m_70==True or module.m_70==False:
#                    i=i+1
#            if  module.m_71==True or module.m_71==False:
#                    i=i+1
#            if  module.m_72==True or module.m_72==False:
#                    i=i+1
#            if  module.m_73==True or module.m_73==False:
#                    i=i+1
#            if  module.m_74==True or module.m_74==False:
#                    i=i+1
#            if  module.m_75==True or module.m_75==False:
#                    i=i+1
#            if  module.m_76==True or module.m_76==False:
#                    i=i+1
#            if  module.m_77==True or module.m_77==False:
#                    i=i+1
#            if  module.m_78==True or module.m_78==False:
#                    i=i+1
#            if  module.m_79==True or module.m_79==False:
#                    i=i+1
#            if  module.m_80==True or module.m_80==False:
#                    i=i+1
#            if  module.m_81==True or module.m_81==False:
#                    i=i+1
#            if  module.m_82==True or module.m_82==False:
#                    i=i+1
#            if  module.m_83==True or module.m_83==False:
#                    i=i+1
#            if  module.m_84==True or module.m_84==False:
#                    i=i+1
#            if  module.m_85==True or module.m_85==False:
#                    i=i+1
#            if  module.m_86==True or module.m_86==False:
#                    i=i+1
#            if  module.m_87==True or module.m_87==False:
#                    i=i+1
#            if  module.m_88==True or module.m_88==False:
#                    i=i+1
#            if  module.m_89==True or module.m_89==False:
#                    i=i+1
#            if  module.m_90==True or module.m_90==False:
#                    i=i+1
#            if  module.m_91==True or module.m_91==False:
#                    i=i+1
#            if  module.m_92==True or module.m_92==False:
#                    i=i+1
#            if  module.m_93==True or module.m_93==False:
#                    i=i+1
#            if  module.m_94==True or module.m_94==False:
#                    i=i+1
#            if  module.m_95==True or module.m_95==False:
#                    i=i+1
#            if  module.m_96==True or module.m_96==False:
#                    i=i+1
#            if  module.m_97==True or module.m_97==False:
#                    i=i+1
#            if  module.m_98==True or module.m_98==False:
#                    i=i+1
#            if  module.m_99==True or module.m_99==False:
#                    i=i+1
#            if  module.m_100==True or module.m_100==False:
#                    i=i+1
#            if  module.m_101==True or module.m_101==False:
#                    i=i+1
#            if  module.m_102==True or module.m_102==False:
#                    i=i+1
#            if  module.m_103==True or module.m_103==False:
#                    i=i+1
#            if  module.m_104==True or module.m_104==False:
#                    i=i+1
#            if  module.m_105==True or module.m_105==False:
#                    i=i+1
#            if  module.m_106==True or module.m_106==False:
#                    i=i+1
#            if  module.m_107==True or module.m_107==False:
#                    i=i+1
#            if  module.m_108==True or module.m_108==False:
#                    i=i+1
#            if  module.m_109==True or module.m_109==False:
#                    i=i+1
#            if  module.m_110==True or module.m_110==False:
#                    i=i+1
#            if  module.m_111==True or module.m_111==False:
#                    i=i+1
#            if  module.m_112==True or module.m_112==False:
#                    i=i+1
#            if  module.m_113==True or module.m_113==False:
#                    i=i+1
#            if  module.m_114==True or module.m_114==False:
#                    i=i+1
#            if  module.m_115==True or module.m_115==False:
#                    i=i+1
#            if  module.m_116==True or module.m_116==False:
#                    i=i+1
#            if  module.m_117==True or module.m_117==False:
#                    i=i+1
#            if  module.m_118==True or module.m_118==False:
#                    i=i+1
#            if  module.m_119==True or module.m_119==False:
#                    i=i+1
#            if  module.m_120==True or module.m_120==False:
#                    i=i+1
#            if  module.m_121==True or module.m_121==False:
#                    i=i+1
#            if  module.m_122==True or module.m_122==False:
#                    i=i+1
#            if module.m_23!='':
#                    i=i+1
#            if module.m_39!='':
#                    i=i+1
#            if module.m_14!='':
#                    i=i+1
#            if module.m_2!='':
#                    i=i+1
#            if module.m_3!='':
#                    i=i+1
#            if module.m_4!='':
#                    i=i+1
#            if module.m_6!='':
#                    i=i+1
#            if module.m_7!='':
#                    i=i+1
#            if module.m_9!='':
#                    i=i+1
#            if module.m_10!='':
#                i=i+1
#            if module.m_1>0:
#                    i=i+1
#            if module.m_11>0:
#                    i=i+1
#            if module.m_12>0:
#                    i=i+1
#            if module.m_13>0:
#                    i=i+1
#            if module.m_15>0:
#                    i=i+1
#            if module.m_18>0 :
#                    i=i+1
#            if Module2Weaknesses.objects.filter(module2=module.id).count()>0 :
#                    i=i+1
#           
#        percent=(i*100/tot_num_fields)
    if num_mod == 2:
        tot_num_fields=122
        i=0
        module2_1=None
        module2_2=None
        module2_1s = Module2_1.objects.filter(session_id=version).count()
        if module2_1s>0:
            module2_1 = Module2_1.objects.filter(session_id=version)[0]
        module2_2s = Module2_2.objects.filter(session_id=version).count()
        if module2_2s>0:
            module2_2 = Module2_2.objects.filter(session_id=version)[0]
            
            if  module2_1  != None and module2_1.m_22==True:
               tot_num_fields=tot_num_fields-1
            if  module2_1  != None and module2_1.m_37==False:
               tot_num_fields=tot_num_fields-2
            if module2_2 != None and module2_2.m_103==False:
               tot_num_fields=tot_num_fields-4
            
            if module2_1  != None and (module2_1.m_5==True or module2_1.m_5==False):
                i=i+1
            if module2_1  != None and (module2_1.m_8==True or module2_1.m_8==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_16==True or module2_1.m_16==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_17==True or module2_1.m_17==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_19==True or module2_1.m_19==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_20==True or module2_1.m_20==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_21==True or module2_1.m_21==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_22==True or module2_1.m_22==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_24==True or module2_1.m_24==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_25==True or module2_1.m_25==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_26==True or module2_1.m_26==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_27==True or module2_1.m_27==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_28==True or module2_1.m_28==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_29==True or module2_1.m_29==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_30==True or module2_1.m_30==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_31==True or module2_1.m_31==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_32==True or module2_1.m_32==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_33==True or module2_1.m_33==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_34==True or module2_1.m_34==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_35==True or module2_1.m_35==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_36==True or module2_1.m_36==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_37==True or module2_1.m_37==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_38==True or module2_1.m_38==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_40==True or module2_1.m_40==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_41==True or module2_1.m_41==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_42==True or module2_1.m_42==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_43==True or module2_1.m_43==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_44==True or module2_1.m_44==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_45==True or module2_1.m_45==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_46==True or module2_1.m_46==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_47==True or module2_1.m_47==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_48==True or module2_1.m_48==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_49==True or module2_1.m_49==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_50==True or module2_1.m_50==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_51==True or module2_1.m_51==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_52==True or module2_1.m_52==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_53==True or module2_1.m_53==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_54==True or module2_1.m_54==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_55==True or module2_1.m_55==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_56==True or module2_1.m_56==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_57==True or module2_1.m_57==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_58==True or module2_1.m_58==False):
                    i=i+1
            if module2_1  != None and (module2_1.m_59==True or module2_1.m_59==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_60==True or module2_2.m_60==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_61==True or module2_2.m_61==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_62==True or module2_2.m_62==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_63==True or module2_2.m_63==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_64==True or module2_2.m_64==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_65==True or module2_2.m_65==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_66==True or module2_2.m_66==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_67==True or module2_2.m_67==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_68==True or module2_2.m_68==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_69==True or module2_2.m_69==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_70==True or module2_2.m_70==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_71==True or module2_2.m_71==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_72==True or module2_2.m_72==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_73==True or module2_2.m_73==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_74==True or module2_2.m_74==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_75==True or module2_2.m_75==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_76==True or module2_2.m_76==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_77==True or module2_2.m_77==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_78==True or module2_2.m_78==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_79==True or module2_2.m_79==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_80==True or module2_2.m_80==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_81==True or module2_2.m_81==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_82==True or module2_2.m_82==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_83==True or module2_2.m_83==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_84==True or module2_2.m_84==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_85==True or module2_2.m_85==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_86==True or module2_2.m_86==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_87==True or module2_2.m_87==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_88==True or module2_2.m_88==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_89==True or module2_2.m_89==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_90==True or module2_2.m_90==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_91==True or module2_2.m_91==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_92==True or module2_2.m_92==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_93==True or module2_2.m_93==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_94==True or module2_2.m_94==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_95==True or module2_2.m_95==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_96==True or module2_2.m_96==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_97==True or module2_2.m_97==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_98==True or module2_2.m_98==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_99==True or module2_2.m_99==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_100==True or module2_2.m_100==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_101==True or module2_2.m_101==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_102==True or module2_2.m_102==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_103==True or module2_2.m_103==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_104==True or module2_2.m_104==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_105==True or module2_2.m_105==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_106==True or module2_2.m_106==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_107==True or module2_2.m_107==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_108==True or module2_2.m_108==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_109==True or module2_2.m_109==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_110==True or module2_2.m_110==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_111==True or module2_2.m_111==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_112==True or module2_2.m_112==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_113==True or module2_2.m_113==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_114==True or module2_2.m_114==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_115==True or module2_2.m_115==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_116==True or module2_2.m_116==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_117==True or module2_2.m_117==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_118==True or module2_2.m_118==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_119==True or module2_2.m_119==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_120==True or module2_2.m_120==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_121==True or module2_2.m_121==False):
                    i=i+1
            if module2_2  != None and (module2_2.m_122==True or module2_2.m_122==False):
                    i=i+1
            if module2_1  != None and module2_1.m_23!='':
                    i=i+1
            if module2_1  != None and module2_1.m_39!='':
                    i=i+1
            if module2_1  != None and module2_1.m_14!='':
                    i=i+1
            if module2_1  != None and  module2_1.m_2!='':
                    i=i+1
            if module2_1  != None and module2_1.m_3!='':
                    i=i+1
            if module2_1  != None and module2_1.m_4!='':
                    i=i+1
            if module2_1  != None and module2_1.m_6!='':
                    i=i+1
            if module2_1  != None and module2_1.m_7!='':
                    i=i+1
            if module2_1  != None and module2_1.m_9!='':
                    i=i+1
            if module2_1  != None and module2_1.m_10!='':
                i=i+1
            if module2_1  != None and module2_1.m_1>0:
                    i=i+1
            if module2_1  != None and module2_1.m_11>0:
                    i=i+1
            if module2_1  != None and module2_1.m_12>0:
                    i=i+1
            if module2_1  != None and module2_1.m_13>0:
                    i=i+1
            if module2_1  != None and module2_1.m_15>0:
                    i=i+1
            if module2_1  != None and module2_1.m_18>0 :
                    i=i+1
            if module2_2  != None and Module2_2Weaknesses.objects.filter(module2=module2_2.id).count()>0 :
                    i=i+1
           
        percent=(i*100/tot_num_fields)   
    if num_mod == 3:
        tot_num_fields=33
        i=0
        modules = Module3.objects.filter(session_id=version).count()
        if modules>0:
            module = Module3.objects.filter(session_id=version)[0]
            if module.m_5==False:
               tot_num_fields=tot_num_fields-2
            if module.m_29==False:
               tot_num_fields=tot_num_fields-1
            
            
            if module.m_4==True or module.m_4==False:
                    i=i+1
            if module.m_5==True or module.m_5==False:
                    i=i+1
            if module.m_21==True or module.m_21==False:
                    i=i+1
            if module.m_25==True or module.m_25==False:
                    i=i+1
            if module.m_26==True or module.m_26==False:
                    i=i+1
            if module.m_27==True or module.m_27==False:
                    i=i+1
            if module.m_28==True or module.m_28==False:
                    i=i+1
            if module.m_29==True or module.m_29==False:
                    i=i+1
            if module.m_6!='':
                    i=i+1
            if module.m_7!='':
                    i=i+1
            if module.m_8!='':
                    i=i+1
            if module.m_12!='':
                    i=i+1
            if module.m_30!='':
                    i=i+1
            if module.m_32!='':
                    i=i+1
            if module.m_1.count()>0:
                    i=i+1
            if module.m_3>0:
                    i=i+1
            if module.m_9>0:
                    i=i+1
            if module.m_10.count()>0:
                    i=i+1
            if module.m_14>0:
                    i=i+1
            if module.m_15>0:
                    i=i+1
            if module.m_16>0:
                    i=i+1
            if module.m_17.count()>0:
                    i=i+1
            if module.m_2>0:
                    i=i+1
            if module.m_11>0:
                    i=i+1
            if module.m_13>0:
                    i=i+1
            if module.m_18>0:
                    i=i+1
            if module.m_18>0:
                    i=i+1
            if module.m_20>0:
                    i=i+1
            if module.m_22>0:
                    i=i+1
            if module.m_23>0:
                    i=i+1
            if module.m_24>0:
                    i=i+1
            if Module3Grid.objects.filter(module3=module.id).count()>0:
                    i=i+1
            if Module3Weaknesses.objects.filter(module3=module.id).count()>0:
                    i=i+1

        percent=(i*100/tot_num_fields)
        
    if num_mod == 4:
        tot_num_fields=34
        i=0
        modules = Module4.objects.filter(session_id=version).count()
        if modules>0:
            module = Module4.objects.filter(session_id=version)[0]
            if module.m_18==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_20==False:
                 tot_num_fields=tot_num_fields-3
            if module.m_24==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_26==False:
                 tot_num_fields=tot_num_fields-3
            if module.m_31==False:
                 tot_num_fields=tot_num_fields-2
                 
            if module.m_2==True or module.m_2==False:
                    i=i+1
            if module.m_7==True or module.m_7==False:
                    i=i+1
            if module.m_9==True or module.m_9==False:
                    i=i+1
            if module.m_14==True or module.m_14==False:
                    i=i+1
            if module.m_15==True or module.m_15==False:
                    i=i+1
            if module.m_16==True or module.m_16==False:
                    i=i+1
            if module.m_18==True or module.m_18==False:
                    i=i+1
            if module.m_19==True or module.m_19==False:
                    i=i+1
            if module.m_20==True or module.m_20==False:
                    i=i+1
            if module.m_21==True or module.m_21==False:
                    i=i+1
            if module.m_22==True or module.m_22==False:
                    i=i+1
            if module.m_23==True or module.m_23==False:
                    i=i+1
            if module.m_24==True or module.m_24==False:
                    i=i+1 
            if module.m_26==True or module.m_26==False:
                    i=i+1
            if module.m_27==True or module.m_27==False:
                    i=i+1
            if module.m_29==True or module.m_29==False:
                    i=i+1
            if module.m_30==True or module.m_30==False:
                    i=i+1
            if module.m_31==True or module.m_31==False:
                    i=i+1
            if module.m_33==True or module.m_33==False:
                    i=i+1
          
           
            if module.m_1>0:
                    i=i+1
            if module.m_3>0:
                    i=i+1
            if module.m_4>0:
                    i=i+1
            if module.m_5>0:
                    i=i+1
            if module.m_6>0:
                    i=i+1
            if module.m_8>0:
                    i=i+1
            if module.m_10>0:
                    i=i+1
            if module.m_11>0:
                    i=i+1
            if module.m_12>0:
                    i=i+1
            if module.m_13>0:
                    i=i+1
            if module.m_17>0:
                    i=i+1
            if module.m_25!='':
                    i=i+1        
            if module.m_28!='':
                    i=i+1        
            if module.m_32!='':
                    i=i+1        
            if Module4Weaknesses.objects.filter(module4=module.id).count()>0:
                    i=i+1

        percent=(i*100/tot_num_fields)
       
    if num_mod == 5:
        tot_num_fields=25
        i=0
        modules = Module5.objects.filter(session_id=version).count()
        if modules>0:
            module = Module5.objects.filter(session_id=version)[0]
            if module.m_1>0:
                    i=i+1
            if  module.m_2>0:
                    i=i+1
            if module.m_3.count()>0 :
               i=i+1
            if  module.m_4==True or module.m_4==False:
                    i=i+1
            if  module.m_5==True or module.m_5==False:
                    i=i+1
            if  module.m_6==True or module.m_6==False:
                    i=i+1
            if  module.m_7==True or module.m_7==False:
                    i=i+1
            if  module.m_8==True or module.m_8==False:
                    i=i+1
            if  module.m_9==True or module.m_9==False:
                    i=i+1
            if  module.m_10==True or module.m_10==False:
                    i=i+1
            if  module.m_11==True or module.m_11==False:
                    i=i+1
            if  module.m_12==True or module.m_12==False:
                    i=i+1
            if  module.m_13==True or module.m_13==False:
                    i=i+1
            if  module.m_14==True or module.m_14==False:
                    i=i+1
            if  module.m_15==True or module.m_15==False:
                    i=i+1
            if  module.m_16>0:
                    i=i+1
            if  module.m_17>0:
                    i=i+1
            if  module.m_18>0:
                    i=i+1
            if  module.m_19>0:
                    i=i+1
            if  module.m_20>0:
                    i=i+1
            if  module.m_21>0:
                i=i+1  
				
	    if  module.m_22==True or module.m_22==False:
                    i=i+1
            if  module.m_23==True or module.m_23==False:
                    i=i+1
            if  module.m_24==True or module.m_24==False:
                    i=i+1
            if Module5Weaknesses.objects.filter(module5=module.id).count()>0 :
                i=i+1
        percent=(i*100/tot_num_fields)
        
    if num_mod == 6:
        tot_num_fields=37
        i=0
        ##print()     
        ##print()     
        ##print()     
        ##print("------------------------------------------------")     
        ##print()     
        ##print("------------        MOD 6              -----------------")     
        ##print("tot_num_fields"+str(tot_num_fields))     
          
        modules = Module6.objects.filter(session_id=version).count()
        if modules>0:
            module = Module6.objects.filter(session_id=version)[0]
            if module.m_1==False:
                 tot_num_fields=tot_num_fields-2
                 ##print("m1 tot_num_fields"+str(tot_num_fields))     
        
            if module.m_3==True:
                 tot_num_fields=tot_num_fields-1
                 ##print("m3 tot_num_fields"+str(tot_num_fields))     
        
            if module.m_6==False:
                 tot_num_fields=tot_num_fields-1
                 ##print("m6 tot_num_fields"+str(tot_num_fields))     
        
            ##print(">>>tot_num_fields"+str(tot_num_fields))     
             
            if module.m_1==True or module.m_1==False:
                    i=i+1
            if module.m_2==True or module.m_2==False:
                    i=i+1
            if module.m_3==True or module.m_3==False:
                    i=i+1
            if module.m_5==True or module.m_5==False:
                    i=i+1
            if module.m_6==True or module.m_6==False:
                    i=i+1
            if module.m_7==True or module.m_7==False:
                    i=i+1
            if module.m_10==True or module.m_10==False:
                    i=i+1
            if module.m_12==True or module.m_12==False:
                    i=i+1
            if module.m_14==True or module.m_14==False:
                    i=i+1
            if module.m_15==True or module.m_15==False:
                    i=i+1
            if module.m_18==True or module.m_18==False:
                    i=i+1
            if module.m_20==True or module.m_20==False:
                    i=i+1
            if module.m_26==True or module.m_26==False:
                    i=i+1
            if module.m_17!='':
                i=i+1
                 
            if module.m_4>0:
                    i=i+1
            if module.m_8>0:
                    i=i+1
            if module.m_9>0:
                    i=i+1
            if module.m_11>0:
                    i=i+1
            if module.m_13>0:
                    i=i+1
            if module.m_16>0:
                    i=i+1
            if module.m_19>0:
                    i=i+1
            if module.m_21>0:
                    i=i+1
            if module.m_22>0:
                    i=i+1
            if module.m_23>0:
                    i=i+1
            if module.m_24>0:
                    i=i+1
            if module.m_25>0:
                    i=i+1
            if module.m_27>0:
                    i=i+1
            if module.m_28>0:
                    i=i+1
            if module.m_29>0:
                    i=i+1
            if module.m_30>0:
                    i=i+1
            if module.m_31>0:
                    i=i+1
            if module.m_32>0:
                    i=i+1
            if module.m_33>0:
                    i=i+1
            if module.m_34>0:
                    i=i+1
            if module.m_35>0:
                    i=i+1
            if module.m_36>0:
                    i=i+1
            if Module6Weaknesses.objects.filter(module6=module.id).count()>0 :
                i=i+1
        ##print(i)        
        percent=(i*100/tot_num_fields)
        
    if num_mod == 7:
        tot_num_fields=69
        i=0
        modules = Module7.objects.filter(session_id=version).count()
        if modules>0:
            module = Module7.objects.filter(session_id=version)[0]
            if module.m_1==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_21==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_29==True:
                 tot_num_fields=tot_num_fields-1
            if module.m_59==False:
                 tot_num_fields=tot_num_fields-1
        
            if module.m_1==True or module.m_1==False:
                    i=i+1
            if module.m_2==True or module.m_2==False:
                    i=i+1
            if module.m_3==True or module.m_3==False:
                    i=i+1
            if module.m_4==True or module.m_4==False:
                    i=i+1
            if module.m_5==True or module.m_5==False:
                    i=i+1
            if module.m_7==True or module.m_7==False:
                    i=i+1
            if module.m_9==True or module.m_9==False:
                    i=i+1
            if module.m_10==True or module.m_10==False:
                    i=i+1
            if module.m_13==True or module.m_13==False:
                    i=i+1
            if module.m_15==True or module.m_15==False:
                    i=i+1
            if module.m_16==True or module.m_16==False:
                    i=i+1
            if module.m_17==True or module.m_17==False:
                    i=i+1
            if module.m_18==True or module.m_18==False:
                    i=i+1
            if module.m_19==True or module.m_19==False:
                    i=i+1
            if module.m_20==True or module.m_20==False:
                    i=i+1
            if module.m_21==True or module.m_21==False:
                    i=i+1
            if module.m_28==True or module.m_28==False:
                    i=i+1
            if module.m_29==True or module.m_29==False:
                    i=i+1
            if module.m_31==True or module.m_31==False:
                    i=i+1
            if module.m_32==True or module.m_32==False:
                    i=i+1
            if module.m_33==True or module.m_33==False:
                    i=i+1
            if module.m_34==True or module.m_34==False:
                    i=i+1
            if module.m_35==True or module.m_35==False:
                    i=i+1
            if module.m_36==True or module.m_36==False:
                    i=i+1
            
            
            if module.m_47==True or module.m_47==False:
                    i=i+1
            if module.m_48==True or module.m_48==False:
                    i=i+1
            if module.m_50==True or module.m_50==False:
                    i=i+1
            if module.m_51==True or module.m_51==False:
                    i=i+1
            if module.m_52==True or module.m_52==False:
                    i=i+1
            if module.m_54==True or module.m_54==False:
                    i=i+1
            if module.m_55==True or module.m_55==False:
                    i=i+1
            
            if module.m_56==True or module.m_56==False:
                    i=i+1
            if module.m_57==True or module.m_57==False:
                    i=i+1
            if module.m_58==True or module.m_58==False:
                    i=i+1
            if module.m_59==True or module.m_59==False:
                    i=i+1
            if module.m_60==True or module.m_60==False:
                    i=i+1
            if module.m_62==True or module.m_62==False:
                    i=i+1
            if module.m_64==True or module.m_64==False:
                    i=i+1
            
            if module.m_6!='':
                    i=i+1
            if module.m_11!='':
                    i=i+1
            if module.m_22!='':
                    i=i+1
            if module.m_30!='':
                    i=i+1
            if module.m_67!='':
                 i=i+1
                 
            if module.m_8>0:
                    i=i+1
            if module.m_12>0:
                    i=i+1
            if module.m_24>0:
                    i=i+1
            if module.m_25>0:
                    i=i+1
            if module.m_26>0:
                    i=i+1
            if module.m_27>0:
                    i=i+1
            if module.m_38>0:
                    i=i+1
            if module.m_40>0:
                    i=i+1
            if module.m_42>0:
                    i=i+1
            if module.m_44>0:
                    i=i+1
            if module.m_46>0:
                    i=i+1
            if module.m_49>0:
                    i=i+1
            if module.m_53>0:
                    i=i+1
          
            if module.m_61>0:
                    i=i+1
            
            if module.m_63>0:
                    i=i+1
            if module.m_65>0:
                    i=i+1
            if module.m_66>0:
                    i=i+1
            if module.m_68>0:
                 i=i+1
            if Module7Grid.objects.filter(module7=module.id).count()>0:
                    i=i+1
            if Module7Matrix23.objects.filter(module7=module.id).count()>0:
                    i=i+1
            if Module7Matrix37.objects.filter(module7=module.id).count()>0:
                    i=i+1
            if Module7Matrix39.objects.filter(module7=module.id).count()>0:
                    i=i+1
            if Module7Matrix41.objects.filter(module7=module.id).count()>0:
                    i=i+1
            if Module7Matrix43.objects.filter(module7=module.id).count()>0:
                    i=i+1
            if Module7Matrix45.objects.filter(module7=module.id).count()>0:                               
                    i=i+1
            if Module7Weaknesses.objects.filter(module7=module.id).count()>0 :
                i=i+1
        percent=(i*100/tot_num_fields)
         
    if num_mod == 8:
        tot_num_fields=45
        ##print(tot_num_fields)
        i=0
        modules = Module8.objects.filter(session_id=version).count()
        if modules>0:
            module = Module8.objects.filter(session_id=version)[0]
            if module.m_14==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_22==False:
                 tot_num_fields=tot_num_fields-3
            ###print(tot_num_fields)
            if module.m_1==True or module.m_1==False:
                 i=i+1        
                 ##print('1:'+str(i))
            if module.m_2==True or module.m_2==False:
                     i=i+1     
                     ##print('2:'+str(i))
            if Module8Grid3.objects.filter(module8=module.id).count()>0:
                    m8g3=Module8Grid3.objects.filter(module8=module.id)[0]
                    if (m8g3.c1== True or m8g3.c1== False) and (m8g3.c2== True or m8g3.c2== False) and (m8g3.c3== True or m8g3.c3== False) and(m8g3.c4== True or m8g3.c4== False)and (m8g3.c5== True or m8g3.c5== False):
                         i=i+1   
                         ##print('3:'+str(i))        
            if module.m_4==True or module.m_4==False:
                     i=i+1     
                     ##print('4:'+str(i))
            if module.m_5==True or module.m_5==False:
                     i=i+1      
                     ##print('5:'+str(i))
            if module.m_6==True or module.m_6==False:
                     i=i+1      
                     ##print('6:'+str(i))
            if module.m_7==True or module.m_7==False:
                     i=i+1        
                     ##print('7:'+str(i))
            
            if module.m_8!='':
                     i=i+1        
                     ##print('8:'+str(i))
            if module.m_9!='':
                     i=i+1         
                     ##print('9:'+str(i))
            
            
            if module.m_10==True or module.m_10==False:
                     i=i+1      
                     ##print('10:'+str(i))
            if module.m_11==True or module.m_11==False:
                     i=i+1      
                     ##print('11:'+str(i))
            if module.m_12==True or module.m_12==False:
                     i=i+1    
                     ##print('12:'+str(i))
            if module.m_13==True or module.m_13==False:
                     i=i+1    
                     ##print('13:'+str(i))
            if module.m_14==True or module.m_14==False:
                     i=i+1     
                     ##print('14:'+str(i))
            if module.m_15!='':
                     i=i+1     
                     ##print('15:'+str(i))
                    
            if module.m_16==True or module.m_16==False:
                     i=i+1      
                     ##print('16:'+str(i))
            if module.m_17.count()>0:
                     i=i+1       
                     ##print('17:'+str(i))
            if Module8Grid18.objects.filter(module8=module.id).count()>0:
                    m8g18=Module8Grid18.objects.filter(module8=module.id)[0]
                    if (m8g18.c1== True or m8g18.c1== False) and (m8g18.c2== True or m8g18.c2== False) and (m8g18.c3== True or m8g18.c3== False) and(m8g18.c4== True or m8g18.c4== False)and (m8g18.c5== True or m8g18.c5== False)and (m8g18.c6== True or m8g18.c6== False):
                         i=i+1    
                         ##print('18:'+str(i))        
            if module.m_19==True or module.m_19==False:
                     i=i+1       
                     ##print('18:'+str(i))
            if module.m_20==True or module.m_20==False:
                     i=i+1     
                     ##print('20:'+str(i))
            if module.m_21==True or module.m_21==False:
                     i=i+1       
                     ##print('21:'+str(i))
            if module.m_22==True or module.m_22==False:
                     i=i+1          
                     ##print('22:'+str(i))
            if module.m_23==True or module.m_23==False:
                     i=i+1          
                     ##print('23:'+str(i))
                    
            if module.m_24>0:
                     i=i+1       
                     ##print('24:'+str(i))
            if module.m_25!='':
                     i=i+1       
                     ##print('25:'+str(i))
            if module.m_26>0:
                     i=i+1        
                     ##print('26:'+str(i))     
            if module.m_27!='':
                
                
                     i=i+1     
                     ##print('27:'+str(i))     
                    
            if module.m_28==True or module.m_28==False:
                     i=i+1     
                     ##print('28:'+str(i))
            if module.m_29==True or module.m_29==False:
                     i=i+1    
                     ##print('29:'+str(i))
            if Module8Matrix30.objects.filter(module8=module.id).count()>0:                               
                 i=i+1         
                 ##print('30:'+str(i))
        
            if module.m_31>0:
                     i=i+1       
                     ##print('31:'+str(i))
            if module.m_32>0:
                     i=i+1      
                     ##print('32:'+str(i))
            if module.m_33>0:
                     i=i+1         
                     
                     ##print('33:'+str(i))
            if module.m_34>0:
                     i=i+1         
                     ##print('34:'+str(i))
            if module.m_35>0:
                     i=i+1            
                     ##print('35:'+str(i))
            
            if module.m_36==True or module.m_36==False:
                     i=i+1       
                     ##print('36:'+str(i))
            if module.m_37>0:
                     i=i+1        
                     ##print('37:'+str(i))
            if module.m_38>0:
                 i=i+1            
                 ##print('38:'+str(i))
            
            if module.m_39==True or module.m_39==False:
                     i=i+1         
                     ##print('39:'+str(i))
            if module.m_40>0:
                     i=i+1         
                     ##print('40:'+str(i))
            
            if module.m_41==True or module.m_41==False:
                     i=i+1        
                     ##print('41:'+str(i))
                    
                 
            if module.m_42>0:
                    i=i+1
                    ##print('42:'+str(i))
            if module.m_43!='':
                    i=i+1
                    ##print('43:'+str(i))

            if module.m_44>0:
                    i=i+1
                    ##print('44:'+str(i))
            
            
           
               
            if Module8Weaknesses.objects.filter(module8=module.id).count()>0 :
                i=i+1
                ##print('45:'+str(i))
                
        ##print('i='+str(i))
        percent=(i*100/tot_num_fields)
       
       
    if num_mod == 9:
        tot_num_fields=47
        print("tot_num_fields="+str(tot_num_fields))
        i=0
        modules = Module9.objects.filter(session_id=version).count()
        if modules>0:
            module = Module9.objects.filter(session_id=version)[0]
            if module.m_15==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_20==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_22==False:
                 tot_num_fields=tot_num_fields-3
            if module.m_28==False:
                 tot_num_fields=tot_num_fields-1
               
            if module.m_2==True or module.m_2==False:
                    i=i+1
                    #print("m2")
            if module.m_3==True or module.m_3==False:
                    i=i+1
                    #print("m3")
            if module.m_4==True or module.m_4==False:
                    i=i+1
                    #print("m4")
            if module.m_6==True or module.m_6==False:
                    i=i+1
                    #print("m6")
            if module.m_7==True or module.m_7==False:
                    i=i+1
                    #print("m7")
            if module.m_8==True or module.m_8==False:
                    i=i+1
                    #print("m8")
            if module.m_9==True or module.m_9==False:
                    i=i+1
                    #print("m9")
            if module.m_11==True or module.m_11==False:
                    i=i+1
                    #print("m11")
            if module.m_12==True or module.m_12==False:
                    i=i+1
                    #print("m12")
            if module.m_13==True or module.m_13==False:
                    i=i+1
                    #print("m13")
            if module.m_14==True or module.m_14==False:
                    i=i+1
                    #print("m14")
            if module.m_15==True or module.m_15==False:
                    i=i+1
                    #print("m15")
            if module.m_16==True or module.m_16==False:
                    i=i+1
                    #print("m16")
            if module.m_17==True or module.m_17==False:
                    i=i+1
                    #print("m17")
            if module.m_18==True or module.m_18==False:
                    i=i+1
                    #print("m18")
            if module.m_19==True or module.m_19==False:
                    i=i+1
                    #print("m19")
            if module.m_20==True or module.m_20==False:
                    i=i+1
                    #print("m20")
            if module.m_22==True or module.m_22==False:
                    i=i+1
                    #print("m22")
            if module.m_23==True or module.m_23==False:
                    i=i+1
                    #print("m23")
            if module.m_26==True or module.m_26==False:
                    i=i+1
                    #print("m26")
            if module.m_27==True or module.m_27==False:
                    i=i+1
                    #print("m27")
            if module.m_28==True or module.m_28==False:
                    i=i+1
                    #print("m28")
            if module.m_29==True or module.m_29==False:
                    i=i+1
                    #print("m29")
            if module.m_30==True or module.m_30==False:
                    i=i+1
                    #print("m30")
            if module.m_32==True or module.m_32==False:
                    i=i+1
                    #print("m32")
            if module.m_33==True or module.m_33==False:
                    i=i+1
                    #print("m33")
            if module.m_34==True or module.m_34==False:
                    i=i+1
                    #print("m34")
            if module.m_41==True or module.m_41==False:
                    i=i+1
                    #print("m41")
            if module.m_10!='':
                    i=i+1
                    #print("m10")
            if module.m_21!='':
                    i=i+1
                    #print("m21")
            if module.m_24!='':
                    i=i+1
                    #print("m24")
            if module.m_25!='':
                    i=i+1
                    #print("m25")

            if module.m_36>0:
                    i=i+1
                    #print("m36")
            if module.m_37>0:
                    i=i+1
                    #print("m37")
            if module.m_38>0:
                    i=i+1
                    #print("m38")
            if module.m_39>0:
                    i=i+1
                    #print("m39")
            if module.m_40>0:
                    i=i+1
                    #print("m40")
            if module.m_42>0:
                    i=i+1
                    #print("m42")
            if module.m_43>0:
                    i=i+1
                    #print("m43")
            if module.m_44>0:
                    i=i+1
                    #print("m44")
            if module.m_45>0:
                    i=i+1
                    #print("m45")
            if module.m_46>0:
                    i=i+1
                    #print("m46")
                    
            if Module9Grid1.objects.filter(module9=module.id).count()>0:
                m9g1=Module9Grid1.objects.filter(module9=module.id)[0]
                if (m9g1.c1== True or m9g1.c1== False) and (m9g1.c2== True or m9g1.c2== False) and (m9g1.c3== True or m9g1.c3== False) and(m9g1.c4== True or m9g1.c4== False):
                    i=i+1
                    #print("m9g1=yes")
            
            if Module9Grid5.objects.filter(module9=module.id).count()>0:
                m9g5=Module9Grid5.objects.filter(module9=module.id)[0]
                if (m9g5.c1== True or m9g5.c1== False) and (m9g5.c2== True or m9g5.c2== False) and (m9g5.c3== True or m9g5.c3== False) and(m9g5.c4== True or m9g5.c4== False)and(m9g5.c5== True or m9g5.c5== False):
                    i=i+1
                    #print("m9g5=yes")
            if Module9Grid31.objects.filter(module9=module.id).count()>0:
                m9g31=Module9Grid31.objects.filter(module9=module.id)[0]
                if (m9g31.c1== True or m9g31.c1== False) and (m9g31.c2== True or m9g31.c2== False) and (m9g31.c3== True or m9g31.c3== False) and(m9g31.c4== True or m9g31.c4== False)and(m9g31.c5== True or m9g31.c5== False):
                    i=i+1
                    #print("m9g31=yes")
            if Module9Matrix35.objects.filter(module9=module.id).count()>0:                               
                i=i+1
                print("Module9Matrix35=yes")
               
         
            if Module9Weaknesses.objects.filter(module9=module.id).count()>0 :
                    i=i+1
                    print("Module9Weaknesses=yes")
            print("=>tot_num_fields="+str(tot_num_fields))
        
            print("Module9")
            print(i)
                
                
   
        percent=(i*100/tot_num_fields)
        
    if num_mod == 10:
        tot_num_fields=61
        i=0
        modules = Module10.objects.filter(session_id=version).count()
        if modules>0:
            module = Module10.objects.filter(session_id=version)[0]
            if module.m_8==True:
                 tot_num_fields=tot_num_fields-1
            if module.m_10==False:
                 tot_num_fields=tot_num_fields-3
            if module.m_32==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_9!='':
                    i=i+1
            if module.m_48>0:
                    i=i+1
            if module.m_49>0:
                    i=i+1
            if module.m_51>0:
                    i=i+1
            if module.m_53>0:
                    i=i+1
            if module.m_54>0:
                    i=i+1
            if module.m_55>0:
                    i=i+1
            if module.m_56>0:
                    i=i+1
            if module.m_57>0:
                    i=i+1
            if module.m_58>0:
                    i=i+1
            if module.m_59>0:
                    i=i+1
            if module.m_60>0:
                    i=i+1
                    
            if module.m_1==True or module.m_1==False:
                    i=i+1
            if module.m_2==True or module.m_2==False:
                    i=i+1
            if module.m_3==True or module.m_3==False:
                    i=i+1
            if module.m_4==True or module.m_4==False:
                    i=i+1
            if module.m_5==True or module.m_5==False:
                    i=i+1
            if module.m_6==True or module.m_6==False:
                    i=i+1
            if module.m_7==True or module.m_7==False:
                    i=i+1
            if module.m_8==True or module.m_8==False:
                    i=i+1
            if module.m_10==True or module.m_10==False:
                    i=i+1
            if module.m_11==True or module.m_11==False:
                    i=i+1
            if module.m_12==True or module.m_12==False:
                    i=i+1
            if module.m_13==True or module.m_13==False:
                    i=i+1
            if module.m_14==True or module.m_14==False:
                    i=i+1
            if module.m_15==True or module.m_15==False:
                    i=i+1
            if module.m_16==True or module.m_16==False:
                    i=i+1
            if module.m_17==True or module.m_17==False:
                    i=i+1
            if module.m_18==True or module.m_18==False:
                    i=i+1
            if module.m_19==True or module.m_19==False:
                    i=i+1
            if module.m_20==True or module.m_20==False:
                    i=i+1
            if module.m_21==True or module.m_21==False:
                    i=i+1
            if module.m_22==True or module.m_22==False:
                    i=i+1
            if module.m_24==True or module.m_24==False:
                    i=i+1
            if module.m_25==True or module.m_25==False:
                    i=i+1
            if module.m_26==True or module.m_26==False:
                    i=i+1
            if module.m_27==True or module.m_27==False:
                    i=i+1
            if module.m_28==True or module.m_28==False:
                    i=i+1
            if module.m_29==True or module.m_29==False:
                    i=i+1
            if module.m_30==True or module.m_30==False:
                    i=i+1
            if module.m_32==True or module.m_32==False:
                    i=i+1
            
            if module.m_34==True or module.m_34==False:
                    i=i+1
            if module.m_35==True or module.m_35==False:
                               i=i+1
            if module.m_36==True or module.m_36==False:
                               i=i+1
            if module.m_38==True or module.m_38==False:
                               i=i+1
            if module.m_39==True or module.m_39==False:
                               i=i+1                    
            if module.m_40==True or module.m_40==False:
                               i=i+1
            if module.m_41==True or module.m_41==False:
                               i=i+1
            if module.m_42==True or module.m_42==False:
                               i=i+1
            if module.m_43==True or module.m_43==False:
                               i=i+1                    
            if module.m_44==True or module.m_44==False:
                               i=i+1
            if module.m_50==True or module.m_50==False:
                               i=i+1
            if module.m_52==True or module.m_52==False:
                               i=i+1
                               
            if Module10Matrix_47.objects.filter(module10=module.id).count()>0:
                    i=i+1
                    #print('Module10Matrix_47 +')
            if Module10Weaknesses.objects.filter(module10=module.id).count()>0:
                    i=i+1
                    #print('Module10Weaknesses +')
            if Module10Grid23.objects.filter(module10=module.id).count()>0:
                m10g23=Module10Grid23.objects.filter(module10=module.id)[0]
                if (m10g23.c1== True or m10g23.c1== False) and (m10g23.c2== True or m10g23.c2== False) and (m10g23.c3== True or m10g23.c3== False) and(m10g23.c4== True or m10g23.c4== False)and(m10g23.c5== True or m10g23.c5== False):
                    i=i+1
                    #print('Module10Grid23 +')
            if Module10Grid33.objects.filter(module10=module.id).count()>0:
                m10g33=Module10Grid33.objects.filter(module10=module.id)[0]
                if (m10g33.c1== True or m10g33.c1== False) and (m10g33.c2== True or m10g33.c2== False) and (m10g33.c3== True or m10g33.c3== False) and(m10g33.c4== True or m10g33.c4== False):
                    i=i+1
                    #print('Module10Grid33 +')
            if Module10Grid45.objects.filter(module10=module.id).count()>0:
                m10g45=Module10Grid45.objects.filter(module10=module.id)[0]
                if (m10g45.c1== True or m10g45.c1== False) and (m10g45.c2== True or m10g45.c2== False) and (m10g45.c3== True or m10g45.c3== False) and(m10g45.c4== True or m10g45.c4== False)and(m10g45.c5== True or m10g45.c5== False):
                    i=i+1
                    #print('Module10Grid45 +')
            if Module10Grid31.objects.filter(module10=module.id).count()>0:
                m10g31=Module10Grid31.objects.filter(module10=module.id)[0]
                if (m10g31.c1== True or m10g31.c1== False) and (m10g31.c2== True or m10g31.c2== False) and (m10g31.c3== True or m10g31.c3== False):
                    i=i+1
                    #print('Module10Grid31 +')
            if Module10Grid46.objects.filter(module10=module.id).count()>0:
                m10g46=Module10Grid46.objects.filter(module10=module.id)[0]
                if (m10g46.c1== True or m10g46.c1== False) and (m10g46.c2== True or m10g46.c2== False) and (m10g46.c3== True or m10g46.c3== False):
                    i=i+1
                    #print('Module10Grid46 +')
            if Module10Grid37.objects.filter(module10=module.id).count()>0:
                m10g37=Module10Grid37.objects.filter(module10=module.id)[0]
                if (m10g37.c1== True or m10g37.c1== False) and (m10g37.c2== True or m10g37.c2== False) and (m10g37.c3== True or m10g37.c3== False) and(m10g37.c4== True or m10g37.c4== False)and(m10g37.c5== True or m10g37.c5== False)and(m10g37.c6== True or m10g37.c6== False)and(m10g37.c7== True or m10g37.c7== False)and(m10g37.c8== True or m10g37.c8== False):
                    i=i+1        
                    #print('Module10Grid37 +')
          
    
   
        percent=(i*100/tot_num_fields)
             
    if num_mod == 11:
        tot_num_fields=68
        i=0
        modules = Module11.objects.filter(session_id=version).count()
        if modules>0:
            module = Module11.objects.filter(session_id=version)[0]
            if module.m_1==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_17==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_26==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_28==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_31==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_57==True:
                 tot_num_fields=tot_num_fields-1
            if module.m_7!='':
                    i=i+1
            if module.m_29!='':
                    i=i+1
            if module.m_27!='':
                i=i+1
            if module.m_21>0:
                    i=i+1
            if module.m_43>0:
                    i=i+1
            if module.m_45>0:
                    i=i+1
            if module.m_44>0:
                    i=i+1
            if module.m_46>0:
                    i=i+1
          
            if module.m_59>0:
                    i=i+1
            if module.m_60>0:
                    i=i+1
            if module.m_62>0:
                    i=i+1
            if module.m_63>0:
                    i=i+1        
            if module.m_64>0:
                    i=i+1        
            if module.m_65>0:
                    i=i+1
            if module.m_66>0:
                    i=i+1
            if module.m_67>0:
                    i=i+1
         
            if module.m_9==True or module.m_9==False:
                    i=i+1
            if module.m_8==True or module.m_8==False:
                    i=i+1
            if module.m_61==True or module.m_61==False:
                    i=i+1
            if module.m_6==True or module.m_6==False:
                    i=i+1
            if module.m_58==True or module.m_58==False:
                    i=i+1
            if module.m_57==True or module.m_57==False:
                    i=i+1
            if module.m_56==True or module.m_56==False:
                    i=i+1
            if module.m_55==True or module.m_55==False:
                    i=i+1
            if module.m_54==True or module.m_54==False:
                    i=i+1
            if module.m_53==True or module.m_53==False:
                    i=i+1
            if module.m_52==True or module.m_52==False:
                    i=i+1
            if module.m_51==True or module.m_51==False:
                    i=i+1
            if module.m_50==True or module.m_50==False:
                    i=i+1
            if module.m_5==True or module.m_5==False:
                    i=i+1
            if module.m_49==True or module.m_49==False:
                    i=i+1
            if module.m_48==True or module.m_48==False:
                    i=i+1
            if module.m_47==True or module.m_47==False:
                    i=i+1
            if module.m_41==True or module.m_41==False:
                    i=i+1
            if module.m_40==True or module.m_40==False:
                    i=i+1
            if module.m_4==True or module.m_4==False:
                    i=i+1
            if module.m_39==True or module.m_39==False:
                    i=i+1
            if module.m_38==True or module.m_38==False:
                    i=i+1
            if module.m_37==True or module.m_37==False:
                    i=i+1
            if module.m_36==True or module.m_36==False:
                    i=i+1
            if module.m_35==True or module.m_35==False:
                    i=i+1
            if module.m_34==True or module.m_34==False:
                    i=i+1
            if module.m_32==True or module.m_32==False:
                    i=i+1
            if module.m_31==True or module.m_31==False:
                    i=i+1
            if module.m_30==True or module.m_30==False:
                    i=i+1
            if module.m_28==True or module.m_28==False:
                    i=i+1
            if module.m_26==True or module.m_26==False:
                    i=i+1
            if module.m_25==True or module.m_25==False:
                    i=i+1
            if module.m_24==True or module.m_24==False:
                    i=i+1
            if module.m_23==True or module.m_23==False:
                    i=i+1
            if module.m_22==True or module.m_22==False:
                    i=i+1
            if module.m_20==True or module.m_20==False:
                    i=i+1
            if module.m_19==True or module.m_19==False:
                    i=i+1
            if module.m_18==True or module.m_18==False:
                    i=i+1
            if module.m_17==True or module.m_17==False:
                    i=i+1
            if module.m_16==True or module.m_16==False:
                    i=i+1
            if module.m_15==True or module.m_15==False:
                    i=i+1
            if module.m_13==True or module.m_13==False:
                    i=i+1
            if module.m_11==True or module.m_11==False:
                    i=i+1
            if module.m_10==True or module.m_10==False:
                    i=i+1
            if module.m_1==True or module.m_1==False:
                    i=i+1
            if Module11Grid2.objects.filter(module11=module.id).count()>0:
                m11g2=Module11Grid2.objects.filter(module11=module.id)[0]
                if (m11g2.c1== True or m11g2.c1== False) and (m11g2.c2== True or m11g2.c2== False) and (m11g2.c3== True or m11g2.c3== False) :
                    i=i+1        
            if Module11Grid3.objects.filter(module11=module.id).count()>0:
                m11g3=Module11Grid3.objects.filter(module11=module.id)[0]
                if (m11g3.c1== True or m11g3.c1== False) and (m11g3.c2== True or m11g3.c2== False) and (m11g3.c3== True or m11g3.c3== False) and (m11g3.c4== True or m11g3.c4== False) and (m11g3.c5== True or m11g3.c5== False):
                    i=i+1        
                    
            if Module11Grid12.objects.filter(module11=module.id).count()>0:
                m11g12=Module11Grid12.objects.filter(module11=module.id)[0]
                if (m11g12.c1== True or m11g12.c1== False) and (m11g12.c2== True or m11g12.c2== False) and (m11g12.c3== True or m11g12.c3== False) and (m11g12.c4== True or m11g12.c4== False) and (m11g12.c5== True or m11g12.c5== False):
                    i=i+1        
            if Module11Grid14.objects.filter(module11=module.id).count()>0:
                m11g14=Module11Grid14.objects.filter(module11=module.id)[0]
                if (m11g14.c1== True or m11g14.c1== False) and (m11g14.c2== True or m11g14.c2== False) and (m11g14.c3== True or m11g14.c3== False) and (m11g14.c4== True or m11g14.c4== False) and (m11g14.c5== True or m11g14.c5== False)and (m11g14.c6== True or m11g14.c6== False):
                    i=i+1   
            if Module11Grid33.objects.filter(module11=module.id).count()>0:
                m11g33=Module11Grid33.objects.filter(module11=module.id)[0]
                if (m11g33.c1== True or m11g33.c1== False) and (m11g33.c2== True or m11g33.c2== False) and (m11g33.c3== True or m11g33.c3== False) and (m11g33.c4== True or m11g33.c4== False):
                    i=i+1   
            
          
            if Module11Matrix42.objects.filter(module11=module.id).count()>0:
                    i=i+1
            if Module11Weaknesses.objects.filter(module11=module.id).count()>0:
                i=i+1
        percent=(i*100/tot_num_fields)    
    if num_mod == 12:
        tot_num_fields=34
        i=0
        modules = Module12.objects.filter(session_id=version).count()
        if modules>0:
            module = Module12.objects.filter(session_id=version)[0]
            #print('************   Module12  *****************')
            if module.m_1==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_12==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_18==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_1==True or module.m_1==False:
                    i=i+1
            if module.m_10==True or module.m_10==False:
                    i=i+1
            if module.m_11==True or module.m_11==False:
                    i=i+1
            if module.m_12==True or module.m_12==False:
                    i=i+1
            if module.m_13==True or module.m_13==False:
                    i=i+1
            if module.m_14==True or module.m_14==False:
                    i=i+1
            if module.m_15==True or module.m_15==False:
                    i=i+1
            if module.m_16==True or module.m_16==False:
                    i=i+1
            if module.m_17==True or module.m_17==False:
                    i=i+1
            if module.m_18==True or module.m_18==False:
                    i=i+1
            if module.m_20==True or module.m_20==False:
                    i=i+1
            if module.m_21==True or module.m_21==False:
                    i=i+1
            if module.m_28==True or module.m_28==False:
                    i=i+1
            if module.m_4==True or module.m_4==False:
                    i=i+1
            if module.m_5==True or module.m_5==False:
                    i=i+1
            if module.m_6==True or module.m_6==False:
                    i=i+1
            if module.m_8==True or module.m_8==False:
                    i=i+1
            if module.m_9==True or module.m_9==False:
                    i=i+1
            if module.m_30>0:
                    i=i+1
            if module.m_31>0:
                    i=i+1
            if module.m_23>0:
                    i=i+1
            if module.m_24>0:
                    i=i+1
            if module.m_25>0:
                    i=i+1
            if module.m_26>0:
                    i=i+1
            if module.m_27>0:
                    i=i+1
            if module.m_32>0:
                    i=i+1
            if module.m_33>0:
                     i=i+1
            if module.m_7!='':
                    i=i+1
            if module.m_19!='':
                     i=i+1
                     
            if Module12Grid2.objects.filter(module12=module.id).count()>0:
                m12g2=Module12Grid2.objects.filter(module12=module.id)[0]
                if (m12g2.c1== True or m12g2.c1== False) and (m12g2.c2== True or m12g2.c2== False) and (m12g2.c3== True or m12g2.c3== False):
                    i=i+1   
            if Module12Grid3.objects.filter(module12=module.id).count()>0:
                m12g3=Module12Grid3.objects.filter(module12=module.id)[0]
                if (m12g3.c1== True or m12g3.c1== False) and (m12g3.c2== True or m12g3.c2== False) and (m12g3.c3== True or m12g3.c3== False)and (m12g3.c4== True or m12g3.c4== False)and (m12g3.c5== True or m12g3.c5== False):
                    i=i+1   
            if Module12Grid_29.objects.filter(module12=module.id).count()>0:
                m12g29=Module12Grid_29.objects.filter(module12=module.id)[0]
                if (m12g29.c1>0) and (m12g29.c2>0)and (m12g29.c3>0)and (m12g29.c4>0):
                    i=i+1   
             
            
            if Module12Matrix22.objects.filter(module12=module.id).count()>0:
                    i=i+1
            if Module12Weaknesses.objects.filter(module12=module.id).count()>0:
                i=i+1
        percent=(i*100/tot_num_fields)    
    if num_mod == 13:
        tot_num_fields=65#66
        i=0
        modules = Module13.objects.filter(session_id=version).count()
        if modules>0:
            module = Module13.objects.filter(session_id=version)[0]
            if module.m_1==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_13==False:
                 tot_num_fields=tot_num_fields-2
            if module.m_14==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_20==False:
                 tot_num_fields=tot_num_fields-2
            if module.m_23==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_27==False:
                 tot_num_fields=tot_num_fields-1
            if module.m_30==False:
                 tot_num_fields=tot_num_fields-2
            if module.m_40==False:
                 tot_num_fields=tot_num_fields-1
           
            
            if module.m_1==True or module.m_1==False:
                    i=i+1
            if Module13Grid2.objects.filter(module13=module.id).count()>0:
                m13g2=Module13Grid2.objects.filter(module13=module.id)[0]
                if (m13g2.c1== True or m13g2.c1== False) and (m13g2.c2== True or m13g2.c2== False) and (m13g2.c3== True or m13g2.c3== False) and (m13g2.c4== True or m13g2.c4== False):
                    i=i+1  
            if Module13Grid3.objects.filter(module13=module.id).count()>0:
                m13g3=Module13Grid3.objects.filter(module13=module.id)[0]
                if (m13g3.c1== True or m13g3.c1== False) and (m13g3.c2== True or m13g3.c2== False) and (m13g3.c3== True or m13g3.c3== False) and (m13g3.c4== True or m13g3.c4== False):
                    i=i+1  
            if module.m_4==True or module.m_4==False:
                    i=i+1
            if module.m_5==True or module.m_5==False:
                    i=i+1
            if module.m_6==True or module.m_6==False:
                    i=i+1
            if module.m_7==True or module.m_7==False:
                    i=i+1
            if module.m_8!='':
                    i=i+1
            if module.m_9==True or module.m_9==False:
                    i=i+1
            if module.m_10==True or module.m_10==False:
                    i=i+1
            if module.m_11==True or module.m_11==False:
                    i=i+1
            if module.m_12==True or module.m_12==False:
                    i=i+1
            if module.m_13==True or module.m_13==False:
                    i=i+1
            if module.m_14==True or module.m_14==False:
                    i=i+1
            if module.m_15==True or module.m_15==False:
                    i=i+1
            if module.m_16==True or module.m_16==False:
                    i=i+1
            if module.m_17==True or module.m_17==False:
                    i=i+1
            if module.m_18==True or module.m_18==False:
                    i=i+1
            if module.m_19==True or module.m_19==False:
                    i=i+1
            if module.m_20==True or module.m_20==False:
                    i=i+1
            if module.m_21!='':
                    i=i+1
                    
                    
            if Module13Grid22.objects.filter(module13=module.id).count()>0:
                m13g22=Module13Grid22.objects.filter(module13=module.id)[0]
                if (m13g22.c1== True or m13g22.c1== False) and (m13g22.c2== True or m13g22.c2== False) and (m13g22.c3== True or m13g22.c3== False):
                    i=i+1  
            if module.m_23==True or module.m_23==False:
                    i=i+1
            if module.m_24>0:
                    i=i+1
            if module.m_25==True or module.m_25==False:
                    i=i+1
            if module.m_26==True or module.m_26==False:
                    i=i+1
            if module.m_27==True or module.m_27==False:
                    i=i+1
            if module.m_28==True or module.m_28==False:
                    i=i+1
            if Module13Grid29.objects.filter(module13=module.id).count()>0:
                m13g29=Module13Grid29.objects.filter(module13=module.id)[0]
                if (m13g29.c1== True or m13g29.c1== False) and (m13g29.c2== True or m13g29.c2== False) and (m13g29.c3== True or m13g29.c3== False) and (m13g29.c4== True or m13g29.c4== False) and (m13g29.c5== True or m13g29.c5== False) and (m13g29.c6== True or m13g29.c6== False)  and (m13g29.c7== True or m13g29.c7== False) and (m13g29.c8== True or m13g29.c8== False) and (m13g29.c9== True or m13g29.c9== False) and (m13g29.c10== True or m13g29.c10== False) and (m13g29.c11== True or m13g29.c11== False) and (m13g29.c12== True or m13g29.c12== False):
                    i=i+1  
            if module.m_30==True or module.m_30==False:
                    i=i+1
  
          
            if Module13Grid31.objects.filter(module13=module.id).count()>0:
                m13g31=Module13Grid31.objects.filter(module13=module.id)[0]
                if (m13g31.c1== True or m13g31.c1== False) and (m13g31.c2== True or m13g31.c2== False) and (m13g31.c3== True or m13g31.c3== False) and (m13g31.c4== True or m13g31.c4== False) and (m13g31.c5== True or m13g31.c5== False):
                    i=i+1  
            if module.m_32==True or module.m_32==False:
                    i=i+1
            if module.m_33==True or module.m_33==False:
                    i=i+1
            if module.m_34==True or module.m_34==False:
                    i=i+1
            if module.m_35==True or module.m_35==False:
                    i=i+1
            if module.m_36==True or module.m_36==False:
                    i=i+1
            if module.m_37==True or module.m_37==False:
                    i=i+1
            if module.m_38==True or module.m_38==False:
                    i=i+1
            if module.m_39==True or module.m_39==False:
                    i=i+1
            if module.m_40==True or module.m_40==False:
                    i=i+1
            if module.m_41==True or module.m_41==False:
                    i=i+1
            if module.m_42>0:
                    i=i+1
            if module.m_43>0:
                    i=i+1
            if module.m_44==True or module.m_44==False:
                    i=i+1
            if module.m_45==True or module.m_45==False:
                    i=i+1
            if module.m_46>0:
                    i=i+1
            if Module13Matrix47.objects.filter(module13=module.id).count()>0:
                    i=i+1
            if module.m_48>0:
                    i=i+1
            if module.m_49>0:
                    i=i+1
            if module.m_50>0:
                    i=i+1
            if module.m_51>0:
                    i=i+1
            if module.m_52>0:
                    i=i+1
            if module.m_53==True or module.m_53==False:
                    i=i+1
            if module.m_54>0:
                    i=i+1
            if module.m_55>0:
                    i=i+1
            if module.m_56>0:
                    i=i+1
            if module.m_57>0:
                    i=i+1
            if module.m_58>0:
                    i=i+1
            if module.m_59==True or module.m_59==False:
                    i=i+1
            if module.m_60 >0:
                     i=i+1
            if module.m_61==True or module.m_61==False:
                    i=i+1
            #if module.m_62>0:
            #        i=i+1
            if module.m_63>0:
                    i=i+1
            if module.m_64!='':
                    i=i+1
            if module.m_65>0:
                    i=i+1
            if Module13Weaknesses.objects.filter(module13=module.id).count()>0:
                i=i+1
                                            
                
        percent=(i*100/tot_num_fields)
          
    return percent

def get_percentage_module(sessionid,moduleid):
    fullepce=0
    percent_m=get_percentage_module_filled(moduleid,sessionid)
    if moduleid==1: 
        if is_stakeholder_filled(sessionid,moduleid):
            fullepce=(percent_m+100)*100/200
        else:
            fullepce=percent_m
        return fullepce    
    else:
        t=0
        if is_stakeholder_filled(sessionid,moduleid):
            t=t+1
        if is_problemanalysis_filled(sessionid,moduleid):
            t=t+1
        if is_swotanalysis_filled(sessionid,moduleid):
            t=t+1
        if is_problemanalysis_filled(sessionid,moduleid):
            t=t+1
        fullepce=(percent_m+(t*100))/(5)
       
        return fullepce

def get_tot_percentage(sessionid):
    """ check if  module is filled"""
    h=0
    perc=0
    tot_percentage=0
    try:
        session = get_object_or_404(PceVersion, id=sessionid)
        items = session.chosen_modules.split(',')
        
        for i in items:
            if i!='':
               y=int(i)
               perc= perc + get_percentage_module(sessionid,y)
              
               h=h+1
        tot_percentage=  (perc/float(h))
        tot_percentage=format(tot_percentage, '.2f')
    except:
        tot_percentage=0
    return tot_percentage



class ModuleListView(ListView):
    """
    ModuleListView
        http://stackoverflow.com/questions/8547880/listing-object-with-specific-tag-using-django-taggit
        http://stackoverflow.com/a/7382708/412329
    """
    context_object_name = 'latest'
    model = PceVersion
    date_field = 'publish_date'
    template_name = 'pce/module_list.html'
    queryset =  PceVersion.objects.all().order_by('-id')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Modules from the specific country and version"""
        self.country = self.kwargs['country']
        self.id = self.kwargs['id'] 
        return PceVersion.objects.filter(country__country_slug=self.country).order_by('-id')
       
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(ModuleListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        user = self.request.user
        
        can_see=0
        can_edit=0
        can_validate=0
        can_sendtovalidator=0
        
        if not user.id == None:
            country = user.get_profile().country
            user_country_slug = lower(slugify(country))
            session = get_object_or_404(PceVersion, id=self.kwargs['id'])
            if session:
                if self.request.user.groups.filter(name='Admin'):
                  can_see=1
                  can_edit=1
                  can_validate=1
                  can_sendtovalidator=0
        
                if self.kwargs['country'] == user_country_slug:
                    if self.request.user.groups.filter(name='PCE Manager/Validator') :
                        can_see=1
                        can_edit=1
                        can_validate=1
                        can_sendtovalidator=0
                    elif self.request.user.groups.filter(name__contains='PCE Editor'):
                        can_see=1
                        can_edit=1
                        can_validate=0
                        can_sendtovalidator=1
                    elif self.request.user.groups.filter(name__contains='PCE Facilitator'):
                        can_see=1
                        can_edit=0
                        can_validate=0
                        can_sendtovalidator=0    
                    
                context['version_number'] = session.version_number            
                context['sessionstatus'] = session.status
                context['session_id'] = session.id
                
                items = session.chosen_modules.split(',')
                items2 = []
                chosen_modules_array = []
                filled_m_array=[]
                filled_stakeh_array=[]
                stake_ids={}
                filled_pa_array=[]
                pa_ids={}
                filled_sa_array=[]
                sa_ids={}
                filled_lf_array=[]
                lf_ids={}
                full_pce={}
                m_ids={}
                m_status={}
                m2_ids={}
                m2_status={}
                m_names={}
                modifydate_ids={}
                for i in items:
                        
                    if i!='':
                        y=int(i)
                        items2.append(y)
                        chosen_modules_array.append(y)
                        percentage_module=0
                        module=None
                       
                        percentage_module=get_percentage_module_filled(y,self.kwargs['id'])
                        if  percentage_module==100:  
                                filled_m_array.append(y)        
                        try:
                            if y==1:
                                module = get_object_or_404(Module1, session=self.kwargs['id'])
                            elif y==2:
                               module = get_object_or_404(Module2, session=self.kwargs['id'])
                            elif y==3:
                               module = get_object_or_404(Module3, session=self.kwargs['id'])
                            elif y==4:
                               module = get_object_or_404(Module4, session=self.kwargs['id'])
                            elif y==5:
                               module = get_object_or_404(Module5, session=self.kwargs['id'])
                            elif y==6:
                               module = get_object_or_404(Module6, session=self.kwargs['id'])
                            elif y==7:
                               module = get_object_or_404(Module7, session=self.kwargs['id'])
                            elif y==8:
                               module = get_object_or_404(Module8, session=self.kwargs['id'])
                            elif y==9:
                               module = get_object_or_404(Module9, session=self.kwargs['id'])
                            elif y==10:
                               module = get_object_or_404(Module10, session=self.kwargs['id'])
                            elif y==11:
                               module = get_object_or_404(Module11, session=self.kwargs['id'])
                            elif y==12:
                               module = get_object_or_404(Module12, session=self.kwargs['id'])
                            elif y==13:
                               module = get_object_or_404(Module13, session=self.kwargs['id'])
                        except:
                            module=None
                            percentage_module=0          
                                
                        if module!=None:
                            m_ids[y] = module.id 
                            m_status[y] = module.status
                            modifydate_ids[y]= module.modify_date
                        else:
                            m_ids[y] =''
                            m_status[y] =''
                        if y == 1:
                            m_names[y]=_("COUNTRY PROFILE")
                        elif y == 2:
                            m_names[y]=_("National phytosanitary legislation")
                        elif y == 3:
                            m_names[y]=_("Environmental forces assessment")
                        elif y == 4:
                            m_names[y]=_("NPPO's mission and strategy")
                        elif y == 5:
                            m_names[y]=_("NPPO's structure and processes")
                        elif y == 6:
                            m_names[y]=_("NPPO's Resources")
                        elif y == 7:
                            m_names[y]=_("Pest diagnostic capacity")
                        elif y == 8:
                            m_names[y]=_("NPPO pest surveillance and pest reporting capacity")
                        elif y == 9:
                            m_names[y]=_("Pest eradication capacity")
                        elif y == 10:
                            m_names[y]=_("Phytosanitary import regulatory system")
                        elif y == 11:
                            m_names[y]=_("Pest risk analysis")
                        elif y == 12:
                            m_names[y]=_("Pest free areas, places and sites, low pest prevalence areas")
                        elif y ==  13:
                            m_names[y]=_("Export certification, re-export, transit")    
                        
                          
                        #stakeholders
                        is_st_filled=is_stakeholder_filled(session.id,y)
                        if is_st_filled:
                            filled_stakeh_array.append(y)
                            try: 
                                stakeholder = get_object_or_404(Stakeholders, session_id=session.id,module=y)
                                stake_ids[y] = stakeholder.id
                            except:
                                stake_ids[y] = ''
                        else:
                            stake_ids[y] = ''
                        #problemanalysis
                        is_pa_filled=is_problemanalysis_filled(session.id,y)
                        ##print('>>>>:'+str(y))
                        ##print(is_pa_filled)
                        if is_pa_filled:
                            filled_pa_array.append(y)
                            try: 
                                problemanalysis = get_object_or_404(ProblemAnalysis, session_id=session.id,module=y)
                                pa_ids[y] = problemanalysis.id
                            except:
                                pa_ids[y] = ''
                        else:
                            pa_ids[y] = ''   
                        #swotmanalysis
                        is_sa_filled=is_swotanalysis_filled(session.id,y)
                        if is_sa_filled:
                            filled_sa_array.append(y)
                            try: 
                                swotanalysis = get_object_or_404(SwotAnalysis, session_id=session.id,module=y)
                                sa_ids[y] = swotanalysis.id
                            except:
                                sa_ids[y] = ''
                        else:
                            sa_ids[y] = ''   
                        #problemanalysis
                        is_lf_filled=is_logicalframework_filled(session.id,y)
                        if is_lf_filled:
                            filled_lf_array.append(y)
                            try: 
                                logicalframework = get_object_or_404(LogicalFramework, session_id=session.id,module=y)
                                lf_ids[y] = logicalframework.id
                            except:
                                lf_ids[y] = ''
                        else:
                            lf_ids[y] = ''
                        #fullpce    
                        full_pce[y] =  get_percentage_module(session.id,y)
                 ####MOD2
                module2_1=None
                module2_2=None 
                for i in items:
                    if i!='' and i=='2':
                        y=int(i)
                        try:
                            if y==2:
                                module2_1 = get_object_or_404(Module2_1, session=self.kwargs['id'])
                                module2_2 = get_object_or_404(Module2_2, session=self.kwargs['id'])
                        except:
                            module=None 
                            percentage_module=0
                            
                        if module2_1!=None:
                            m2_ids[0] =module2_1.id
                            m2_status[0] =module2_1.status
                        if module2_2!=None:
                           m2_ids[1] =module2_2.id
                           m2_status[1] =module2_2.status

                        
                               
                          
                       
                context['chosen_modules'] = chosen_modules_array
                context['filled_stakeh_array'] = filled_stakeh_array
                context['stake_ids'] = stake_ids
                context['filled_pa_array'] = filled_pa_array
                context['pa_ids'] = pa_ids
                context['filled_sa_array'] = filled_sa_array
                context['sa_ids'] = sa_ids
                context['filled_lf_array'] = filled_lf_array
                context['lf_ids'] = lf_ids
                context['filled_m_array'] = filled_m_array
                context['m_ids'] = m_ids
                context['m_status'] = m_status
                context['m2_ids'] = m2_ids
                context['m2_status'] = m2_status
                context['modifydate_ids'] = modifydate_ids
                context['sessionid'] = session.id
                context['m_names'] = m_names
                
                context['full_pce'] =full_pce
                context['tot_percentage'] = get_tot_percentage(session.id)                        
                  
                context['chosen_modules'] = items2
                context['range1']= range(1,14)
              
            
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['can_validate'] = can_validate
        context['can_sendtovalidator'] = can_sendtovalidator
        return context

class PceSessionListView(ListView):
    """
    PceVersionListView
    """
    context_object_name = 'latest'
    model = PceVersion
    date_field = 'publish_date'
    template_name = 'pce/session_list.html'
    queryset = PceVersion.objects.all().order_by('-id')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return PceVersion from the specific country """
        self.country = self.kwargs['country']
        return PceVersion.objects.filter(country__country_slug=self.country).order_by('-id')
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PceSessionListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        user = self.request.user
        
        empty=0
        can_see=0
        can_edit=0
        latest_closed=0
        
        if not user.id == None:
            country = user.get_profile().country
            user_country_slug = lower(slugify(country))
            sessions= PceVersion.objects.filter(country__country_slug=self.country).order_by('-id')
            if sessions.count()>0:
                session= sessions[0]
                context['version_number'] = session.version_number            
                context['sessionstatus'] = session.status
                if session.status == 2  or session.status == 3:
                    latest_closed = 1
                else:
                    latest_closed = 0
            else:
                empty=1
                
              
            #self.request.user.groups.filter(name='Admin') or (self.kwargs['country'] == user_country_slug and self.request.user.groups.filter(name='PCE Manager/Validator') and self.request.user.groups.filter(name='PCE Facilitator')) :
            if canSee('',self.kwargs['country'],self.request.user,''):
               can_see=1
            if canEdit('',self.kwargs['country'],self.request.user,''):
               can_edit=1
        
        context['latest_closed'] = latest_closed
        context['empty'] = empty
        context['can_see'] = can_see
        context['can_edit'] = can_edit
       
        return context


class PceDashboardListView(ListView):
    """
    PceDashboardListView
        http://stackoverflow.com/questions/8547880/listing-object-with-specific-tag-using-django-taggit
        http://stackoverflow.com/a/7382708/412329
    """
    context_object_name = 'latest'
    model = PceVersion
    date_field = 'publish_date'
    template_name = 'pce/dashboard.html'
    queryset = PceVersion.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return PceVersion from the specific country """
        self.country = self.kwargs['country']
        self.id = self.kwargs['id']
        return PceVersion.objects.filter(country__country_slug=self.country,id=self.id )
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PceDashboardListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] =  self.kwargs['id']
        
        user = self.request.user
        
        can_see=0
        can_edit=0
        can_validate=0
        can_close=0
        can_configure=0
        can_sendtovalidator=0
        
        if not user.id == None:
            country = user.get_profile().country
            user_country_slug = lower(slugify(country))
            session = get_object_or_404(PceVersion, id=self.kwargs['id'])
          
            if session:
                if self.request.user.groups.filter(name='Admin'):
                  can_see=1
                  can_edit=1
                  can_validate=1
                  can_close=1
                  can_configure=1
                  can_sendtovalidator=0
        
                if self.kwargs['country'] == user_country_slug:
                    if self.request.user.groups.filter(name='PCE Manager/Validator') :
                        can_see=1
                        can_edit=1
                        can_validate=1
                        can_close=1
                        can_configure=1
                        can_sendtovalidator=0
                    elif self.request.user.groups.filter(name__contains='PCE Editor'):
                        can_see=1
                        can_edit=1
                        can_validate=0
                        can_close=0
                        can_configure=0
                        can_sendtovalidator=1
                    elif self.request.user.groups.filter(name__contains='PCE Facilitator'):
                        can_see=1
                        can_edit=0
                        can_validate=0
                        can_close=0
                        can_configure=0
                        can_sendtovalidator=0
                        
           
                context['range']= range(1,14)
                context['version_number'] = session.version_number
                context['sessionstatus'] = session.status
                chosen_modules_array = []
                filled_m_array=[]
                filled_stakeh_array=[]
                stake_ids={}
                filled_pa_array=[]
                pa_ids={}
                filled_sa_array=[]
                sa_ids={}
                filled_lf_array=[]
                lf_ids={}
                full_pce={}
                m_ids={}
                m2_ids={}
                m_status={}
                m2_status={}
               
                
                module=None
                module2_1=None
                module2_2=None
                items=[]
                if session.chosen_modules!=None:
                    items = session.chosen_modules.split(',')
                for i in items:
                    if i!='':
                        y=int(i)
                  
                        chosen_modules_array.append(y)
                        percentage_module=0
                        percentage_module=get_percentage_module_filled(y,self.kwargs['id'])
                        if  percentage_module==100:  
                           filled_m_array.append(y)
                        try:
                            print(y)
                            if y==1:
                                module = get_object_or_404(Module1, session=self.kwargs['id'])
                            elif y==2:
                                 module = get_object_or_404(Module2, session=self.kwargs['id'])
                            elif y==3:
                                module = get_object_or_404(Module3, session=self.kwargs['id'])
                            elif y==4:
                                module = get_object_or_404(Module4, session=self.kwargs['id'])
                            elif y==5:
                                module = get_object_or_404(Module5, session=self.kwargs['id'])
                            elif y==6:
                                module = get_object_or_404(Module6, session=self.kwargs['id'])
                            elif y==7:
                                module = get_object_or_404(Module7, session=self.kwargs['id'])
                            elif y==8:
                                module = get_object_or_404(Module8, session=self.kwargs['id'])
                            elif y==9:
                                module = get_object_or_404(Module9, session=self.kwargs['id'])
                            elif y==10:
                                module = get_object_or_404(Module10, session=self.kwargs['id'])
                            elif y==11:
                                module = get_object_or_404(Module11, session=self.kwargs['id'])
                            elif y==12:
                                module = get_object_or_404(Module12, session=self.kwargs['id'])
                            elif y==13:
                                module = get_object_or_404(Module13, session=self.kwargs['id'])
                        except:
                            module=None 
                            percentage_module=0
                        
                            
                        if module!=None:
                            m_ids[y] =module.id
                            m_status[y] =module.status
                        else:
                            m_ids[y] =''
                            m_status[y] =''
                        
                        #stakeholders
                        is_st_filled=is_stakeholder_filled(session.id,y)
                        if is_st_filled:
                            filled_stakeh_array.append(y)
                            try: 
                                stakeholder = get_object_or_404(Stakeholders, session_id=session.id,module=y)
                                stake_ids[y] = stakeholder.id
                            except:
                                stake_ids[y] = ''
                        else:
                            stake_ids[y] = ''
                        #problemanalysis
                        is_pa_filled=is_problemanalysis_filled(session.id,y)
                        if is_pa_filled:
                            filled_pa_array.append(y)
                            try: 
                                problemanalysis = get_object_or_404(ProblemAnalysis, session_id=session.id,module=y)
                                pa_ids[y] = problemanalysis.id
                            except:
                                pa_ids[y] = ''
                        else:
                            pa_ids[y] = ''   
                        #swotmanalysis
                        is_sa_filled=is_swotanalysis_filled(session.id,y)
                        if is_sa_filled:
                            filled_sa_array.append(y)
                            try: 
                                swotanalysis = get_object_or_404(SwotAnalysis, session_id=session.id,module=y)
                                sa_ids[y] = swotanalysis.id
                            except:
                                sa_ids[y] = ''
                        else:
                            sa_ids[y] = ''   
                        #problemanalysis
                        is_lf_filled=is_logicalframework_filled(session.id,y)
                        if is_lf_filled:
                            filled_lf_array.append(y)
                            try: 
                                logicalframework = get_object_or_404(LogicalFramework, session_id=session.id,module=y)
                                lf_ids[y] = logicalframework.id
                            except:
                                lf_ids[y] = ''
                        else:
                            lf_ids[y] = ''
                        #fullpce  
                        #print('               DAHS    get percentage mod')
                        full_pce[y] =  get_percentage_module(session.id,y)
                ####MOD2
                for i in items:
                    if i!='' and i=='2':
                        y=int(i)
                        try:
                            if y==2:
                                module2_1 = get_object_or_404(Module2_1, session=self.kwargs['id'])
                                module2_2 = get_object_or_404(Module2_2, session=self.kwargs['id'])
                        except:
                            module=None 
                            percentage_module=0
                            
                        if module2_1!=None:
                            m2_ids[0] =module2_1.id
                            m2_status[0] =module2_1.status
                        if module2_2!=None:
                           m2_ids[1] =module2_2.id
                           m2_status[1] =module2_2.status

                        
                        
                       
                                    
                context['chosen_modules'] = chosen_modules_array
                context['filled_stakeh_array'] = filled_stakeh_array
                context['stake_ids'] = stake_ids
                context['filled_pa_array'] = filled_pa_array
                context['pa_ids'] = pa_ids
                context['filled_sa_array'] = filled_sa_array
                context['sa_ids'] = sa_ids
                context['filled_lf_array'] = filled_lf_array
                context['lf_ids'] = lf_ids
                context['filled_m_array'] = filled_m_array
                context['m_ids'] = m_ids
                context['m_status'] = m_status
                context['m2_ids'] = m2_ids
                context['m2_status'] = m2_status
                context['full_pce'] =full_pce
                #print('               DAHS    get_tot_percentage')
                context['tot_percentage'] = get_tot_percentage(session.id)
                    
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['can_validate'] = can_validate
        context['can_close'] = can_close
        context['can_configure'] = can_configure
        context['can_sendtovalidator'] = can_sendtovalidator
        return context



class PceVersionDetailView(DetailView):
    """ Pest report detail page """
    model = PceVersion
    context_object_name = 'pceversion'
    template_name = 'pce/pceversion_detail.html'
    queryset = PceVersion.objects.filter()
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(PceVersionDetailView, self).get_context_data(**kwargs)
      
        p = get_object_or_404(PceVersion, id=self.kwargs['pk'])
        
        versions= PceVersion.objects.filter(country__country_slug=self.kwargs['country']).order_by('-modify_date')
        context['versions'] = versions
        return context
    

 
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def pceversion_create(request, country):
    """ Create pceversion """
    user = request.user
    author = user
    country1=country
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))

    sessions= PceVersion.objects.filter(country__country_slug=user_country_slug).order_by('-id')
        
    numberR=PceVersion.objects.filter(country_id=country.id).count()
    numberR=numberR+1
    pceumber=str(numberR)
   
    can_see=0
    can_edit=0
    version_number=''
    version_id=''

    if user.groups.filter(name='Admin') or (country1 == user_country_slug and user.groups.filter(name='PCE Manager/Validator')) :
        can_see=1
        if sessions.count()>0:
            session= sessions[0]
            version_number = session.version_number
            version_id=session.id
            if session.status==2 or session.status==3:
                can_edit=1
            else:    
                can_edit=0        
        else:    
              version_number = 1
              can_edit=1   
    
    if request.method == "POST":
         form = PceVersionForm(request.POST, request.FILES)
         if form.is_valid():
            new_pceversion = form.save(commit=False)
            new_pceversion.author = request.user
            new_pceversion.version_number=pceumber
            new_pceversion.author_id = author.id
            form.save()
           
            g1=Group.objects.get(name='PCE Facilitator')
            if request.POST['is_facilitated'] == 'True':
                firstname_facilitator=request.POST['firstname_facilitator']
                lastname_facilitator=request.POST['lastname_facilitator']
                email_facilitator=request.POST['email_facilitator']
                check_and_addUser(firstname_facilitator,lastname_facilitator,email_facilitator,'PCE Facilitator',0,country,user_country_slug)
                for i in range(1,14):
                    add_toStakeholders(request.user,str(i),new_pceversion.id,firstname_facilitator,lastname_facilitator,email_facilitator)

            info(request, _("Successfully saved configuration data - step 1."))
            return redirect("pceversion-edit-2", country=user_country_slug, id=new_pceversion.id)
         else:
             return render_to_response('pce/pceversion_create.html', {'form': form,'can_edit':can_edit,'can_see':can_see,'version_number':version_number,'version_id':version_id},
             context_instance=RequestContext(request))
    else:
        form = PceVersionForm(initial={'country': country}, instance=PceVersion())
     
    return render_to_response('pce/pceversion_create.html', {'form': form,'can_edit':can_edit,'can_see':can_see,'version_number':version_number,'version_id':version_id},
        context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def pceversion_edit_step1(request, country, id=None ):
    """ Create pceversion step 2"""
    user = request.user
    author = user
    country1=country
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    
    pceversion=None
    can_see=0
    can_edit=0
    version_number=''
    version_id=''
    if id:
        if user.groups.filter(name='Admin') or (country1 == user_country_slug and user.groups.filter(name='PCE Manager/Validator')) :
            can_see=1
            can_edit=1
            pceversion = get_object_or_404(PceVersion, country=country, pk=id)
            version_number = pceversion.version_number
            version_id=pceversion.id
    else: 
        version_number = 1
    
    if request.method == "POST":
        #print('aaaaaaaaaaaaaaaaaaaaaaa')
        form = PceVersionForm1(request.POST, instance=pceversion)
      #  if form.is_valid():
       # form.save()
        info(request, _("Successfully saved session configuration - step 1."))
        return redirect("pceversion-edit-2", country=user_country_slug,  id=pceversion.id)
       # else:
       #     return render_to_response('pce/pceversion_edit_step1.html', {'form': form,'pceversion':pceversion,'can_edit':can_edit,'can_see':can_see,'version_number':version_number,'version_id':version_id},
       #     context_instance=RequestContext(request))
    else:
        form = PceVersionForm1(instance=pceversion )
     
    return render_to_response('pce/pceversion_edit_step1.html', {'form': form,'pceversion':pceversion,'can_edit':can_edit,'can_see':can_see,'version_number':version_number,'version_id':version_id},
        context_instance=RequestContext(request))
 
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def pceversion_edit_step2(request, country, id=None ):
    """ Create pceversion step 2"""
    user = request.user
    author = user
    country1=country
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion=None
    can_see=0
    can_edit=0
    version_number=''
    version_id=''
    if user.groups.filter(name='Admin') or (country1 == user_country_slug and user.groups.filter(name='PCE Manager/Validator')) :
        can_see=1
        can_edit=1
        
    if id:
        pceversion = get_object_or_404(PceVersion, country=country, pk=id)
        version_number = pceversion.version_number
        version_id=pceversion.id
       
    else: 
        version_number = 1
        
        
    if request.method == "POST":
            
        form = PceVersionForm2(request.POST, request.FILES,instance=pceversion)
        ##print(form)
        if form.is_valid():
            modules=''
            for u in request.POST:
                if u == 'full_1':
                    modules=modules+str(1)+','
                if u == 'full_2':
                    modules=modules+str(2)+','
                if u == 'full_3':
                    modules=modules+str(3)+','
                if u == 'full_4':
                    modules=modules+str(4)+','                    
                if u == 'full_5':
                    modules=modules+str(5)+','
                if u == 'full_6':
                    modules=modules+str(6)+','                    
                if u == 'full_7':
                    modules=modules+str(7)+','
                if u == 'full_8':
                    modules=modules+str(8)+','                    
                if u == 'full_9':
                    modules=modules+str(9)+','
                if u == 'full_10':
                    modules=modules+str(10)+','                    
                if u == 'full_11':
                    modules=modules+str(11)+','
                if u == 'full_12':
                    modules=modules+str(12)+','                    
                if u == 'full_13':
                    modules=modules+str(13)+','
            pceversion.chosen_modules=modules
            pceversion.modify_date=timezone.now()
            form.save()
           
            
            info(request, _("Successfully saved session configuration - step 2."))
            
            return redirect("pceversion-edit-3", country=user_country_slug, id=pceversion.id)
        else:
            
            return render_to_response('pce/pceversion_edit_step2.html', {'form': form,'pceversion':pceversion,'can_edit':can_edit,'can_see':can_see,'version_number':version_number,'version_id':version_id},
             context_instance=RequestContext(request))
       
    else:
        form = PceVersionForm2(instance=pceversion )
     
    return render_to_response('pce/pceversion_edit_step2.html', {'form': form,'pceversion':pceversion,'can_edit':can_edit,'can_see':can_see,'version_number':version_number,'version_id':version_id},
        context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def pceversion_edit_step3(request, country, id=None ):
    """ Create pceversion step3"""
    user = request.user
    author = user
    country1=country
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion=None
    can_see=0
    can_edit=0
    version_number=''
    version_id=''
    if user.groups.filter(name='Admin') or (country1 == user_country_slug and user.groups.filter(name='PCE Manager/Validator')) :
        can_see=1
        can_edit=1
    #print('')    
    #print(' >>>> STEP 3 <<<< ')    
    #print(id)    
    if id:
        pceversion = get_object_or_404(PceVersion, country=country, pk=id)
        version_number = pceversion.version_number
        version_id=pceversion.id
    else: 
        version_number = 1
    #print(pceversion.id)
    if request.method == "POST":
        form = PceVersionForm3(request.POST, request.FILES,instance=pceversion)
        
        if form.is_valid():
            for i in range(1,9):
                firstname=request.POST['ed'+str(i)+'_firstname']
                lastname=request.POST['ed'+str(i)+'_lastname']
                email=request.POST['ed'+str(i)+'_email']
                if email!='':
                    check_and_addUser(firstname,lastname,email,'PCE Editor ',i,pceversion.country,user_country_slug)
                    if i==1:
                        add_toStakeholders(request.user,1,pceversion.id,firstname,lastname,email)
                        add_toStakeholders(request.user,3,pceversion.id,firstname,lastname,email)
                    if i==2:
                        add_toStakeholders(request.user,2,pceversion.id,firstname,lastname,email)
                    if i==3:
                        add_toStakeholders(request.user,4,pceversion.id,firstname,lastname,email)
                        add_toStakeholders(request.user,5,pceversion.id,firstname,lastname,email)
                        add_toStakeholders(request.user,6,pceversion.id,firstname,lastname,email)
                    if i==4:
                        add_toStakeholders(request.user,7,pceversion.id,firstname,lastname,email)
                    if i==5:
                        add_toStakeholders(request.user,8,pceversion.id,firstname,lastname,email)
                    if i==6:
                        add_toStakeholders(request.user,9,pceversion.id,firstname,lastname,email)
                    if i==7:
                        add_toStakeholders(request.user,10,pceversion.id,firstname,lastname,email)
                        add_toStakeholders(request.user,11,pceversion.id,firstname,lastname,email)
                    if i==8:
                        add_toStakeholders(request.user,12,pceversion.id,firstname,lastname,email)
                        add_toStakeholders(request.user,13,pceversion.id,firstname,lastname,email)
            pceversion.modify_date=timezone.now()
            form.save()
            info(request, _("Successfully saved session configuration - step 3."))
            return redirect("pceversion-edit-4", country=user_country_slug, id=pceversion.id)
        else:
             return render_to_response('pce/pceversion_edit_step3.html', {'form': form,'pceversion':pceversion,'can_edit':can_edit,'can_see':can_see,'version_number':version_number,'version_id':version_id},
             context_instance=RequestContext(request))
    else:
        form = PceVersionForm3(instance=pceversion )
        ##print(pceversion.ed1_firstname)
    return render_to_response('pce/pceversion_edit_step3.html', {'form': form,'pceversion':pceversion,'can_edit':can_edit,'can_see':can_see,'version_number':version_number,'version_id':version_id},
        context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def pceversion_edit_step4(request, country, id=None ):
    """ Create pceversion step 4"""
    user = request.user
    author = user
    country1=country
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion=None
    can_see=0
    can_edit=0
    version_number=''
    version_id=''
    if user.groups.filter(name='Admin') or (country1 == user_country_slug and user.groups.filter(name='PCE Manager/Validator')) :
        can_see=1
        can_edit=1
    if id:
        pceversion = get_object_or_404(PceVersion, country=country, pk=id)
        version_number = pceversion.version_number
        version_id=pceversion.id
    else: 
        version_number = 1
     
    return render_to_response('pce/pceversion_edit_step4.html', {'can_edit':can_edit,'can_see':can_see,'version_number':version_number,'version_id':version_id,'sessionid':version_id},
        context_instance=RequestContext(request)) 
        
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def pceversion_close(request, country, id=None ):
    """ CLOSE/CANCEL pceversion session"""
    user = request.user
    author = user
    country1=country
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    if id:
        if user.groups.filter(name='Admin') or (country1 == user_country_slug and user.groups.filter(name='PCE Manager/Validator')) :
            pceversion = get_object_or_404(PceVersion, country=country, pk=id)
            pceversion.status=2
            pceversion.modify_date=timezone.now()
            pceversion.completed_date=timezone.now()
            pceversion.save()

            info(request, _("Successfully Closed session."))
            return redirect("pceversion-list", country=user_country_slug,  id=pceversion.id)
        
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def pceversion_completed(request, country, id=None ):
    """ completed pceversion session"""
    user = request.user
    author = user
    country1=country
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    if id:
        if user.groups.filter(name='Admin') or (country1 == user_country_slug and user.groups.filter(name='PCE Manager/Validator')) :
            pceversion = get_object_or_404(PceVersion, country=country, pk=id)
            pceversion.status=3
            pceversion.modify_date=timezone.now()
            pceversion.completed_date=timezone.now()
            pceversion.save()

            info(request, _("Successfully Completed session."))
            return redirect("pceversion-list", country=user_country_slug,  id=pceversion.id)
           
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module_validate(request, country, sessionid=None,module=None,id=None ):
    """ VALIDATE Module """
    user = request.user
    author = user
    
    country1=country
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    modnum='module'+str(module)
    print(modnum)
    if id and module:
        
         if user.groups.filter(name='Admin') or (country1 == user_country_slug and user.groups.filter(name='PCE Manager/Validator')) :
            moduletovalidate=None
            
            moduletovalidate2_1=None
            moduletovalidate2_2=None
            
            percentage_module=get_percentage_module(sessionid,int(module))
            
            try:
                if modnum == 'module1':
                    moduletovalidate = get_object_or_404(Module1, id=id)
                elif modnum=='module2':
                   moduletovalidate2_1 = get_object_or_404(Module2_1, id=id)
                   moduletovalidate2_2 = get_object_or_404(Module2_2, id=id)
                elif modnum=='module3':
                    moduletovalidate = get_object_or_404(Module3, id=id)
                elif modnum=='module4':
                    moduletovalidate = get_object_or_404(Module4, id=id)
                elif modnum=='module5':
                    moduletovalidate = get_object_or_404(Module5, id=id)
                elif modnum=='module6':
                    moduletovalidate = get_object_or_404(Module6, id=id)
                elif modnum=='module7':
                    moduletovalidate = get_object_or_404(Module7, id=id)
                elif modnum=='module8':
                    moduletovalidate = get_object_or_404(Module8, id=id)
                elif modnum=='module9':
                    moduletovalidate = get_object_or_404(Module9, id=id)
                elif modnum=='module10':
                    moduletovalidate = get_object_or_404(Module10, id=id)
                elif modnum=='module11':
                    moduletovalidate = get_object_or_404(Module11, id=id)
                elif modnum=='module12':
                    moduletovalidate = get_object_or_404(Module12, id=id)
                elif  modnum=='module13':
                    moduletovalidate = get_object_or_404(Module13, id=id)
            except:
                print('--------------ERROR --')
            
            if percentage_module == 100:
               
                if modnum=='module2':
                   moduletovalidate2_1.status=4
                   moduletovalidate2_1.save()
                   moduletovalidate2_2.status=4
                   moduletovalidate2_2.save()
                else:   
                    moduletovalidate.status=4
                    moduletovalidate.save()
                info(request, _("Successfully VALIDATED Module: "+str(module)+"."))
                return redirect("pceversion-list", country=user_country_slug,  id=sessionid)
       
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module_unvalidate(request, country, sessionid=None,module=None,id=None ):
    """ UN-VALIDATE Module """
    user = request.user
    author = user
    country1=country
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    modnum='module'+str(module)
    
    if id and module:
         if user.groups.filter(name='Admin') or (country1 == user_country_slug and user.groups.filter(name='PCE Manager/Validator')) :
            moduletovalidate=None
            moduletovalidate2_1=None
            moduletovalidate2_2=None
            percentage_module=get_percentage_module(sessionid,int(module))
            try:
                if modnum == 'module1':
                    moduletovalidate = get_object_or_404(Module1, id=id)
                elif modnum=='module2':
                    moduletovalidate2_1 = get_object_or_404(Module2_1, id=id)
                    moduletovalidate2_2 = get_object_or_404(Module2_2, id=id)
                elif modnum=='module3':
                    moduletovalidate = get_object_or_404(Module3, id=id)
                elif modnum=='module4':
                    moduletovalidate = get_object_or_404(Module4, id=id)
                elif modnum=='module5':
                    moduletovalidate = get_object_or_404(Module5, id=id)
                elif modnum=='module6':
                    moduletovalidate = get_object_or_404(Module6, id=id)
                elif modnum=='module7':
                    moduletovalidate = get_object_or_404(Module7, id=id)
                elif modnum=='module8':
                    moduletovalidate = get_object_or_404(Module8, id=id)
                elif modnum=='module9':
                    moduletovalidate = get_object_or_404(Module9, id=id)
                elif modnum=='module10':
                    moduletovalidate = get_object_or_404(Module10, id=id)
                elif modnum=='module11':
                    moduletovalidate = get_object_or_404(Module11, id=id)
                elif modnum=='module12':
                    moduletovalidate = get_object_or_404(Module12, id=id)
                elif  modnum=='module13':
                    moduletovalidate = get_object_or_404(Module13, id=id)
            except:
                print('--------------ERROR --')
            
            if percentage_module == 100:
                if modnum=='module2':
                   moduletovalidate2_1.status=2
                   moduletovalidate2_1.save()
                   moduletovalidate2_2.status=2
                   moduletovalidate2_2.save()
                else:   
                    moduletovalidate.status=2
                    moduletovalidate.save()
                info(request, _("Successfully UN-VALIDATED Module: "+str(module)+"."))
                return redirect("pceversion-list", country=user_country_slug,  id=sessionid)

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module_sendtovalidator(request, country, sessionid=None,module=None,id=None ):
    """ SEND TO VALIDATOR Module """
    user = request.user
    author = user
    country1=country
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    modnum='module'+str(module)
    if id and module:
         if canEdit(sessionid,country,user,module):
            moduletovalidate=None
            moduletovalidate2_1=None
            moduletovalidate2_2=None
            percentage_module=get_percentage_module(sessionid,int(module))
            try:
                if modnum == 'module1':
                    moduletovalidate = get_object_or_404(Module1, id=id)
                elif modnum=='module2':
                    moduletovalidate2_1 = get_object_or_404(Module2_1, id=id)
                    moduletovalidate2_2 = get_object_or_404(Module2_2, id=id)
                elif modnum=='module3':
                    moduletovalidate = get_object_or_404(Module3, id=id)
                elif modnum=='module4':
                    moduletovalidate = get_object_or_404(Module4, id=id)
                elif modnum=='module5':
                    moduletovalidate = get_object_or_404(Module5, id=id)
                elif modnum=='module6':
                    moduletovalidate = get_object_or_404(Module6, id=id)
                elif modnum=='module7':
                    moduletovalidate = get_object_or_404(Module7, id=id)
                elif modnum=='module8':
                    moduletovalidate = get_object_or_404(Module8, id=id)
                elif modnum=='module9':
                    moduletovalidate = get_object_or_404(Module9, id=id)
                elif modnum=='module10':
                    moduletovalidate = get_object_or_404(Module10, id=id)
                elif modnum=='module11':
                    moduletovalidate = get_object_or_404(Module11, id=id)
                elif modnum=='module12':
                    moduletovalidate = get_object_or_404(Module12, id=id)
                elif  modnum=='module13':
                    moduletovalidate = get_object_or_404(Module13, id=id)
            except:
                print('--------------ERROR --')
            if percentage_module == 100:
                if modnum=='module2':
                   moduletovalidate2_1.status=3
                   moduletovalidate2_1.save()
                   moduletovalidate2_2.status=3
                   moduletovalidate2_2.save()
                else:  
                    moduletovalidate.status=3
                    moduletovalidate.save()
                idcountry=user.get_profile().country_id
                pce_validators=IppcUserProfile.objects.filter(country=idcountry)
                validator_email = []
         
                for u in pce_validators:
                    user_obj=User.objects.get(id=u.user_id)
                    if user_obj.groups.filter(name='PCE Manager/Validator'):
                        validator_email.append(user_obj.email)
      
                msg='Dear '+str(country1)+' PCE Manager,<br><br>this message is to notify that the PCE Editor completed the module: '+str(module)+'.<br><br>You can now VALIDATE the module going to the <a href="www.ippc.int/pce/'+str(user_country_slug)+'/session/dashboard/'+str(sessionid)+'">PCE Dashbaord<strong></strong></a>.'
                subject='PCE NOTIFICATION ['+str(country)+']'  
                notifificationmessage = mail.EmailMessage(subject,msg,'ippc@fao.org', validator_email, ['paola.sentinelli@fao.org'])
                notifificationmessage.content_subtype = "html"
                try:
                    sent =notifificationmessage.send()
                    
                except:
                    print('ERROR sending')

                info(request, _("Successfully SENT Module "+str(module)+" to PCE Manager for validation."))
    return redirect("pceversion-list", country=user_country_slug,  id=sessionid)

        
class Module1ListView(ListView):
    context_object_name = 'latest'
    model = Module1
    date_field = 'publish_date'
    template_name = 'pce/module_1.html'#'pce/module_1_view.html'
    queryset = Module1.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module1 from the specific id """
        
        self.id =None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module1.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module1ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        context['regions'] =REGIONS
        #crops = Crops.objects.all().order_by()
       
        #context['crops'] =crops
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id
      
        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'1'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'1') and session.status==1:
            can_edit=1
        if  can_see or can_edit:  
            module= 1
            module1=None
            form   = Module1FormView()
            form7  = Module1MajorCropsFormSet()
            form8  = Module1MajorImportsFormSet()
            form11 = Module1MajorExportsFormSet()
            form20 =  Module1MajorPartenerImportFormSet()
            form21 = Module1MajorPartenerExportFormSet()
            form26 = Module1AidFormSet()
            if id!='':
                module1 = get_object_or_404(Module1,  id=self.kwargs['id'])
                form   = Module1FormView(instance=module1)
                form7  = Module1MajorCropsFormSet(instance=module1)
                form8  = Module1MajorImportsFormSet(instance=module1)
                form11 = Module1MajorExportsFormSet(instance=module1)
                form20 =  Module1MajorPartenerImportFormSet(instance=module1)
                form21 = Module1MajorPartenerExportFormSet(instance=module1)
                form26 = Module1AidFormSet(instance=module1)
          
            context['module1']=module1
            context['form']=form
            context['form26']=form26
            context['form7']=form7
            context['form8']=form8
            context['form11']=form11
            context['form20']=form20
            context['form21']=form21
            context['form26']=form26
            
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(1,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],1)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],1)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],1)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],1)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=1).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=1).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=1).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=1).id

        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
class Module1ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module1
    date_field = 'publish_date'
    template_name = 'pce/module_1.html'#'pce/module_1_pdf.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module1ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'
       
        REGIONS = (
            (0, _("--- Please select ---")),
            (1, _("Africa")),
            (2, _("Asia")),
            (3, _("Europe")),
            (4, _("Latin America and Caribbean")),
            (5, _("Near East")),
            (6, _("North America")),
            (7, _("South West Pacific")),
        )
       
        VAL_IMP = (
            (0, _("--- Please select ---")),
            (1, _("Not known")),
            (2, _("0 to $100,000")),
            (3, _("$100,000 to $500,000")),
            (4, _("$500,000 to $1M")),
            (5, _("$1M to $10M")),
            (6, _("$10M to $25M")),
            (7, _("$25M to $50M")),
            (8, _("$50M to $100M")),
            (9, _("Greater than $100M ")),
        )

       
        VAL_EXP = (
             (0, _("--- Please select ---")),
             (1, _("Unknown")),
            (2, _("0 to $100,000")),
            (3, _("$100,000 to $500,000")),
            (4, _("$500,000 to $1M")),
            (5, _("$1M to $10M")),
            (6, _("$10M to $50M")),
            (7, _("$25M to $50M")),
            (8, _("greater than $50M")),
        )

    
        VAL_PERCENT = (
            (0, _("--- Please select ---")),
            (1, _("0")),
            (2, _("10")),
            (3, _("20")),
            (4, _("30")),
            (5, _("40")),
            (6, _("50")),
            (7, _("60")),
            (8, _("70")),
            (9, _("80")),
            (10, _("90")),
            (11, _("100")),
        )

       
       
        NUM_BILATERAL = (
            (0, _("--- Please select ---")),
            (1, _("1-3")),
            (2, _("4-6")),
            (3, _("7-10")),
            (4, _("greater than 10")),
        )
        context['regions'] =REGIONS
        context['VAL_IMP'] =VAL_IMP
        context['VAL_EXP'] =VAL_EXP
        context['VAL_PERCENT'] =VAL_PERCENT
        context['regions'] =REGIONS
      
        context['NUM_BILATERAL'] =NUM_BILATERAL
      
        
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module1.objects.filter(id=id)
        else:
            context['latest']= ''
        context['id'] = id
        
        can_see=0
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'1'):
            can_see=1
        if  can_see:  
            module= 1
            module1=None
            form   = Module1FormView()
            form7  = Module1MajorCropsFormSet()
            form8  = Module1MajorImportsFormSet()
            form11 = Module1MajorExportsFormSet()
            form20 =  Module1MajorPartenerImportFormSet()
            form21 = Module1MajorPartenerExportFormSet()
            form26 = Module1AidFormSet()
            if id!='':
                module1 = get_object_or_404(Module1,  id=self.kwargs['id'])
                form   = Module1FormView(instance=module1)
                form7  = Module1MajorCropsFormSet(instance=module1)
                form8  = Module1MajorImportsFormSet(instance=module1)
                form11 = Module1MajorExportsFormSet(instance=module1)
                form20 =  Module1MajorPartenerImportFormSet(instance=module1)
                form21 = Module1MajorPartenerExportFormSet(instance=module1)
                form26 = Module1AidFormSet(instance=module1)
          
            context['module1']=module1
            context['form']=form
            context['form26']=form26
            context['form7']=form7
            context['form8']=form8
            context['form11']=form11
            context['form20']=form20
            context['form21']=form21
            context['form26']=form26
            context['membership1']=  Membership1.objects.filter()
            context['membership2']=  Membership2.objects.filter()
    
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(1,self.kwargs['sessionid'])
         
        context['can_see'] = can_see
        
        return context
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module1_create(request, country,sessionid=None):
    """ Create module1  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module1_count= Module1.objects.filter(country__country_slug=country,session=sessionid).order_by('-modify_date').count()
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    is_st_filled = is_stakeholder_filled(sessionid,1)
    is_pa_filled = is_problemanalysis_filled(sessionid,1)
    is_sa_filled = is_swotanalysis_filled(sessionid,1)
    is_lf_filled = is_logicalframework_filled(sessionid,1)
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    
    if module1_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,1):
                can_edit=1
              
                m_percentage=get_percentage_module_filled(1,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=1).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=1).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=1).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=1).id


            
    if request.method == "POST":
        form = Module1Form(request.POST, request.FILES)
        form26 = Module1AidFormSet(request.POST)
        form7 = Module1MajorCropsFormSet(request.POST)
        form8 = Module1MajorImportsFormSet(request.POST)
        form11 = Module1MajorExportsFormSet(request.POST)
        form20 =  Module1MajorPartenerImportFormSet(request.POST)
        form21 = Module1MajorPartenerExportFormSet(request.POST)

        if form.is_valid() and form26.is_valid()and form7.is_valid()and form8.is_valid()and form11.is_valid()and form20.is_valid()and form21.is_valid():
            new_mod1 = form.save(commit=False)
            new_mod1.author = request.user
            new_mod1.session=pceversion
            new_mod1.country = country
            form.save()

            form7.instance = new_mod1
            form7.save()

            form8.instance = new_mod1
            form8.save()

            form11.instance = new_mod1
            form11.save()

            form20.instance = new_mod1
            form20.save()

            form21.instance = new_mod1
            form21.save()

            form26.instance = new_mod1
            form26.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
            
            info(request, _("Successfully saved Module 1."))
            #return redirect("module1-edit",country=user_country_slug,sessionid=pceversion.id,id=module1.id,)
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
            return render_to_response('pce/module_1.html', {'context' : 'Edit','form': form,'form26': form26,'form7': form7,'form8': form8,'form11': form11,'form20': form20,'form21': form21,
             'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'module':5,
            },
            context_instance=RequestContext(request))

    else:
        form = Module1Form(initial={'country': country,'session': pceversion.id})
        form26 = Module1AidFormSet()
        form7 = Module1MajorCropsFormSet()
        form8 = Module1MajorImportsFormSet()
        form11 = Module1MajorExportsFormSet()
        form20 =  Module1MajorPartenerImportFormSet()
        form21 = Module1MajorPartenerExportFormSet()

    return render_to_response('pce/module_1.html', {'context' : 'Edit','form': form,'form26': form26,'form7': form7,'form8': form8,
            'form11': form11,'form20': form20,'form21': form21,'sessionid':sessionid, 'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'module':1,},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module1_edit(request, country, id=None,sessionid=None, template_name='pce/module_1.html'):
    """ Edit module_1 """
    user = request.user
    author = user
    country=user.get_profile().country
    
       
             
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    stakeholders = Stakeholders.objects.filter(session=sessionid)
    if stakeholders.count()>0:
         stakeid = Stakeholders.objects.filter(session=sessionid)[0].id
   
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    is_st_filled = is_stakeholder_filled(sessionid,1)
    is_pa_filled = is_problemanalysis_filled(sessionid,1)
    is_sa_filled = is_swotanalysis_filled(sessionid,1)
    is_lf_filled = is_logicalframework_filled(sessionid,1)
  
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
  
    if id:
        module1 = get_object_or_404(Module1, country=country, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,1):
            can_edit=1
            m_percentage=get_percentage_module_filled(1,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=1).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=1).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=1).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=1).id
    else:
        module1 = Module1()
  
    if request.POST:
        form = Module1Form(request.POST,  request.FILES, instance=module1)
        form26 = Module1AidFormSet(request.POST,instance=module1)
        form7 = Module1MajorCropsFormSet(request.POST,instance=module1)
        form8 = Module1MajorImportsFormSet(request.POST,instance=module1)
        form11 = Module1MajorExportsFormSet(request.POST,instance=module1)
        form20 =  Module1MajorPartenerImportFormSet(request.POST,instance=module1)
        form21 = Module1MajorPartenerExportFormSet(request.POST,instance=module1)
       
        if form.is_valid() and form26.is_valid()and form7.is_valid()and form8.is_valid()and form11.is_valid()and form20.is_valid()and form21.is_valid():
            form.save()
          
            form7.instance = module1
            form7.save()
            
            form8.instance = module1
            form8.save()
            
            form11.instance = module1
            form11.save()
            
            form20.instance = module1
            form20.save()
            
            form21.instance = module1
            form21.save()
            
            form26.instance = module1
            form26.save()
           
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 1."))
            #if 'save' in request.POST:
            #    return redirect("module1-edit",country=user_country_slug,sessionid=pceversion.id,id=module1.id,)
            #elif 'saveclose' in request.POST:
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module1Form(instance=module1)
        form26 = Module1AidFormSet(instance=module1)
        form7 = Module1MajorCropsFormSet(instance=module1)
        form8 = Module1MajorImportsFormSet(instance=module1)
        form11 = Module1MajorExportsFormSet(instance=module1)
        form20 =  Module1MajorPartenerImportFormSet(instance=module1)
        form21 = Module1MajorPartenerExportFormSet(instance=module1)
        
    return render_to_response(template_name, {'context' : 'Edit',
        'form': form,'form26': form26,'form7': form7,'form8': form8,
        'form11': form11,'form20': form20,'form21': form21,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'sessionid':sessionid,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'module':1,'tot_percentage':tot_percentage},  
        context_instance=RequestContext(request))

#MODULE 2
class Module2ListView(ListView):
    context_object_name = 'latest'
    model = Module2
    date_field = 'publish_date'
    template_name = 'pce/module_2.html'
    queryset = Module2.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module2 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module2.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module2ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id
        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'2'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'2') and session.status==1:
            can_edit=1
        if  can_see or can_edit:  
            module= 2
            form   = Module2FormView()
            form123= Module2WeaknessesFormSet()
            module2=None
            if id!='':
                module2 = get_object_or_404(Module2,  id=self.kwargs['id'])
                form   = Module2FormView(instance=module2)
                form123= Module2WeaknessesFormSet(instance=module2)
      
            context['form']=form
            context['form123']=form123
            context['module2']=module2
           
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(2,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],2)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],2)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],2)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],2)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=2).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=2).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=2).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=2).id

      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
    
    
class Module2ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module1
    date_field = 'publish_date'
    template_name = 'pce/module_2.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module2ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
       
        DEFINITIONS = (
            (0, _("--- Please select ---")),
            (1, _("Not at all")),
            (2, _("Major improvements required")),
            (3, _("Mostly")),
            (4, _("Minor modifications needed")),
            (5, _("Totally")),
        ) 
       
        DEFINITIONS1 = (
            (0, _("--- Please select ---")),
            (1, _("Not at all")),
            (2, _("Major improvements required")),
            (3, _("Mostly")),
            (4, _("Minor improvements needed")),
            (5, _("Totally")),
        ) 
     
        ACT = (
            (0, _("--- Please select ---")),
            (1, _("Not updated")),
            (2, _("Under revision")),
            (3, _("For Cabinet consideration")),
            (4, _("For consideration at Parliament")),
            (5, _("Updated to IPPC 1997")),
        )
    
        CIVIL = (
            (0, _("--- Please select ---")),
            (1, _("Civil law system")),
            (2, _("Common law system")),
            (3, _("Religious law")),
            (44, _("Pluralistic system")),
        )
        
      
        YEAR = (
        (0, _("--Please Select--")),
        (1, _("1901")),
        (2, _("1902")),
        (3, _("1903")),
        (4, _("1904")),
        (5, _("1905")),
        (6, _("1906")),
        (7, _("1907")),
        (8, _("1908")),
        (9, _("1909")),
        (10, _("1910")),
        (11, _("1911")),
        (12, _("1912")),
        (13, _("1913")),
        (14, _("1914")),
        (15, _("1915")),
        (16, _("1916")),
        (17, _("1917")),
        (18, _("1918")),
        (19, _("1919")),
        (20, _("1920")),
        (21, _("1921")),
        (22, _("1922")),
        (23, _("1923")),
        (24, _("1924")),
        (25, _("1925")),
        (26, _("1926")),
        (27, _("1927")),
        (28, _("1928")),
        (29, _("1929")),
        (30, _("1930")),
        (31, _("1931")),
        (32, _("1932")),
        (33, _("1933")),
        (34, _("1934")),
        (35, _("1935")),
        (36, _("1936")),
        (37, _("1937")),
        (38, _("1938")),
        (39, _("1939")),
        (40, _("1940")),
        (41, _("1941")),
        (42, _("1942")),
        (43, _("1943")),
        (44, _("1944")),
        (45, _("1945")),
        (46, _("1946")),
        (47, _("1947")),
        (48, _("1948")),
        (49, _("1949")),
        (50, _("1950")),
        (51, _("1951")),
        (52, _("1952")),
        (53, _("1953")),
        (54, _("1954")),
        (55, _("1955")),
        (56, _("1956")),
        (57, _("1957")),
        (58, _("1958")),
        (59, _("1959")),
        (60, _("1960")),
        (61, _("1961")),
        (62, _("1962")),
        (63, _("1963")),
        (64, _("1964")),
        (65, _("1965")),
        (66, _("1966")),
        (67, _("1967")),
        (68, _("1968")),
        (69, _("1969")),
        (70, _("1970")),
        (71, _("1971")),
        (72, _("1972")),
        (73, _("1973")),
        (74, _("1974")),
        (75, _("1975")),
        (76, _("1976")),
        (77, _("1977")),
        (78, _("1978")),
        (79, _("1979")),
        (80, _("1980")),
        (81, _("1981")),
        (82, _("1982")),
        (83, _("1983")),
        (84, _("1984")),
        (85, _("1985")),
        (86, _("1986")),
        (87, _("1987")),
        (88, _("1988")),
        (89, _("1989")),
        (90, _("1990")),
        (91, _("1991")),
        (92, _("1992")),
        (93, _("1993")),
        (94, _("1994")),
        (95, _("1995")),
        (96, _("1996")),
        (97, _("1997")),
        (98, _("1998")),
        (99, _("1999")),
        (100, _("2000")),
        (101, _("2001")),
        (102, _("2002")),
        (103, _("2003")),
        (104, _("2004")),
        (105, _("2005")),
        (106, _("2006")),
        (107, _("2007")),
        (108, _("2008")),
        (109, _("2009")),
        (110, _("2010")),
        (111, _("2011")),
        (112, _("2012")),
        (113, _("2013")),
        (114, _("2014")),
        (115, _("2015")),
        )
        THEM1 = (
            (0, _("--- Please select ---")),
            (1,_("None of them")),
            (2,_("One of them")),
            (3,_("Two of them")),
            (4,_("All of them")),
        )
        context['DEFINITIONS'] =DEFINITIONS
        context['DEFINITIONS1'] =DEFINITIONS1
        context['ACT'] =ACT
        context['CIVIL'] =CIVIL
        context['YEAR'] =YEAR
        context['THEM1'] =THEM1
        
        id=''
        context['latest']= ''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module2.objects.filter(id=id)
    
        context['id'] = id
    
        can_see=0
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'2'):
            can_see=1
        if  can_see:  
            module= 2
            form   = Module2FormView()
            form123= Module2WeaknessesFormSet()
            module2=None
            if id!='':
                module2 = get_object_or_404(Module2,  id=self.kwargs['id'])
                form   = Module2FormView(instance=module2)
                form123= Module2WeaknessesFormSet(instance=module2)
      
            context['module2']=module2
            context['form']=form
            context['form123']=form123
            context['module2']=module2
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(2,self.kwargs['sessionid'])
         
        context['can_see'] = can_see
        
        return context
    
    
    
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module2_create(request, country,sessionid=None):
    """ Create module2  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module2_count= Module2.objects.filter(session=sessionid).order_by('-modify_date').count()
    
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,2)
    is_pa_filled = is_problemanalysis_filled(sessionid,2)
    is_sa_filled = is_swotanalysis_filled(sessionid,2)
    is_lf_filled = is_logicalframework_filled(sessionid,2)

    if module2_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,2):   
                can_edit=1
                m_percentage=get_percentage_module_filled(2,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=2).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=2).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=2).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=2).id
            
    if request.method == "POST":
        form = Module2Form(request.POST, request.FILES)
        form123= Module2WeaknessesFormSet(request.POST)
      
        if form.is_valid() and form123.is_valid():
            new_mod2 = form.save(commit=False)
            new_mod2.author = request.user
            new_mod2.session=pceversion
            form.save()
            
            form123.instance = new_mod2
            form123.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
                 
            info(request, _("Successfully saved Module 2."))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_2.html', {'context':'Edit','form': form,'form123': form123,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':2,},
             context_instance=RequestContext(request))
    else:
        form = Module2Form(initial={'country': country,'session': pceversion.id})
        form123= Module2WeaknessesFormSet()

    return render_to_response('pce/module_2.html', {'context':'Edit','form': form,'form123': form123,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':2,},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module2_edit(request, country, id=None,sessionid=None, template_name='pce/module_2.html'):
    """ Edit module_2 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)

    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,2)
    is_pa_filled = is_problemanalysis_filled(sessionid,2)
    is_sa_filled = is_swotanalysis_filled(sessionid,2)
    is_lf_filled = is_logicalframework_filled(sessionid,2)          
    
    if id:
        module2 = get_object_or_404(Module2, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,2):
            can_edit=1
            m_percentage=get_percentage_module_filled(2,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=2).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=2).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=2).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=2).id
    else:
        module2 = Module2()
  
    if request.POST:
        form = Module2Form(request.POST,  request.FILES, instance=module2)
        form123= Module2WeaknessesFormSet(request.POST,  request.FILES, instance=module2)
        if form.is_valid() and form123.is_valid():    
            form.save()
          
            form123.instance = module2
            form123.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 2"))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module2Form(instance=module2)
        form123= Module2WeaknessesFormSet(instance=module2)
      
    return render_to_response(template_name, {
      'context':'Edit','form': form,'form123': form123,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':2,

    }, context_instance=RequestContext(request))

### #MODULE 2 part1:
class Module2_1ListView(ListView):
    context_object_name = 'latest'
    model = Module2_1
    date_field = 'publish_date'
    template_name = 'pce/module_2_1.html'
    queryset = Module2_1.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module2_1 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module2_1.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module2_1ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id
        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'2'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'2') and session.status==1:
            can_edit=1
        if  can_see or can_edit:  
            module= 2
            form   = Module2_1FormView()
            module2=None
            if id!='':
                module2 = get_object_or_404(Module2_1,  id=self.kwargs['id'])
                form   = Module2FormView(instance=module2)
               
            context['form']=form
            context['module2']=module2
           
            context['module']='2'
            context['modulestr']='2 part 1'
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(2,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],2)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],2)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],2)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],2)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=2).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=2).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=2).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=2).id

      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
    
    
class Module2_1ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module2_1
    date_field = 'publish_date'
    template_name = 'pce/module_2_1.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module2_1ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
     
        DEFINITIONS = (
            (0, _("--- Please select ---")),
            (1, _("Not at all")),
            (2, _("Major improvements required")),
            (3, _("Mostly")),
            (4, _("Minor modifications needed")),
            (5, _("Totally")),
        ) 
       
        DEFINITIONS1 = (
            (0, _("--- Please select ---")),
            (1, _("Not at all")),
            (2, _("Major improvements required")),
            (3, _("Mostly")),
            (4, _("Minor improvements needed")),
            (5, _("Totally")),
        ) 
     
        ACT = (
            (0, _("--- Please select ---")),
            (1, _("Not updated")),
            (2, _("Under revision")),
            (3, _("For Cabinet consideration")),
            (4, _("For consideration at Parliament")),
            (5, _("Updated to IPPC 1997")),
        )
    
        CIVIL = (
            (0, _("--- Please select ---")),
            (1, _("Civil law system")),
            (2, _("Common law system")),
            (3, _("Religious law")),
            (44, _("Pluralistic system")),
        )
        
      
        YEAR = (
        (0, _("--Please Select--")),
        (1, _("1901")),
        (2, _("1902")),
        (3, _("1903")),
        (4, _("1904")),
        (5, _("1905")),
        (6, _("1906")),
        (7, _("1907")),
        (8, _("1908")),
        (9, _("1909")),
        (10, _("1910")),
        (11, _("1911")),
        (12, _("1912")),
        (13, _("1913")),
        (14, _("1914")),
        (15, _("1915")),
        (16, _("1916")),
        (17, _("1917")),
        (18, _("1918")),
        (19, _("1919")),
        (20, _("1920")),
        (21, _("1921")),
        (22, _("1922")),
        (23, _("1923")),
        (24, _("1924")),
        (25, _("1925")),
        (26, _("1926")),
        (27, _("1927")),
        (28, _("1928")),
        (29, _("1929")),
        (30, _("1930")),
        (31, _("1931")),
        (32, _("1932")),
        (33, _("1933")),
        (34, _("1934")),
        (35, _("1935")),
        (36, _("1936")),
        (37, _("1937")),
        (38, _("1938")),
        (39, _("1939")),
        (40, _("1940")),
        (41, _("1941")),
        (42, _("1942")),
        (43, _("1943")),
        (44, _("1944")),
        (45, _("1945")),
        (46, _("1946")),
        (47, _("1947")),
        (48, _("1948")),
        (49, _("1949")),
        (50, _("1950")),
        (51, _("1951")),
        (52, _("1952")),
        (53, _("1953")),
        (54, _("1954")),
        (55, _("1955")),
        (56, _("1956")),
        (57, _("1957")),
        (58, _("1958")),
        (59, _("1959")),
        (60, _("1960")),
        (61, _("1961")),
        (62, _("1962")),
        (63, _("1963")),
        (64, _("1964")),
        (65, _("1965")),
        (66, _("1966")),
        (67, _("1967")),
        (68, _("1968")),
        (69, _("1969")),
        (70, _("1970")),
        (71, _("1971")),
        (72, _("1972")),
        (73, _("1973")),
        (74, _("1974")),
        (75, _("1975")),
        (76, _("1976")),
        (77, _("1977")),
        (78, _("1978")),
        (79, _("1979")),
        (80, _("1980")),
        (81, _("1981")),
        (82, _("1982")),
        (83, _("1983")),
        (84, _("1984")),
        (85, _("1985")),
        (86, _("1986")),
        (87, _("1987")),
        (88, _("1988")),
        (89, _("1989")),
        (90, _("1990")),
        (91, _("1991")),
        (92, _("1992")),
        (93, _("1993")),
        (94, _("1994")),
        (95, _("1995")),
        (96, _("1996")),
        (97, _("1997")),
        (98, _("1998")),
        (99, _("1999")),
        (100, _("2000")),
        (101, _("2001")),
        (102, _("2002")),
        (103, _("2003")),
        (104, _("2004")),
        (105, _("2005")),
        (106, _("2006")),
        (107, _("2007")),
        (108, _("2008")),
        (109, _("2009")),
        (110, _("2010")),
        (111, _("2011")),
        (112, _("2012")),
        (113, _("2013")),
        (114, _("2014")),
        (115, _("2015")),
        )
        THEM1 = (
            (0, _("--- Please select ---")),
            (1,_("None of them")),
            (2,_("One of them")),
            (3,_("Two of them")),
            (4,_("All of them")),
        )
        context['DEFINITIONS'] =DEFINITIONS
        context['DEFINITIONS1'] =DEFINITIONS1
        context['ACT'] =ACT
        context['CIVIL'] =CIVIL
        context['YEAR'] =YEAR
        context['THEM1'] =THEM1
        
        id=''
        context['latest']= ''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module2_1.objects.filter(id=id)
    
        context['id'] = id
    
        can_see=0
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'2'):
            can_see=1
        print('can see')    
        if  can_see:  
            module= 2
            form   = Module2_1FormView()
            module2=Module2_1
            if id!='':
                module2 = get_object_or_404(Module2_1,  id=self.kwargs['id'])
                form   = Module2_1FormView(instance=module2)
            
            context['module2']=module2
            context['form']=form
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(2,self.kwargs['sessionid'])
         
        context['can_see'] = can_see
        
        return context
    
    
    
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module2_1_create(request, country,sessionid=None):
    """ Create module2  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module2_count= Module2_1.objects.filter(session=sessionid).order_by('-modify_date').count()
    
    
    ###ID of MOD2 part 2
    idmod2_2=-1
    module2_2= Module2_2.objects.filter(session=sessionid).order_by('-modify_date')
    if module2_2.count()>0:
        idmod2_2=module2_2[0].id
        
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,2)
    is_pa_filled = is_problemanalysis_filled(sessionid,2)
    is_sa_filled = is_swotanalysis_filled(sessionid,2)
    is_lf_filled = is_logicalframework_filled(sessionid,2)

    if module2_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,2):   
                can_edit=1
                m_percentage=get_percentage_module_filled(2,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=2).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=2).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=2).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=2).id
            
    if request.method == "POST":
        form = Module2_1Form(request.POST, request.FILES)
        form2 = Module2_2Form(initial={'country': country,'session': pceversion.id})
        form_2_123= Module2_2WeaknessesFormSet()
        
        if form.is_valid() :
            new_mod2 = form.save(commit=False)
            new_mod2.author = request.user
            new_mod2.session=pceversion
            if idmod2_2>0:
                new_mod2.id= idmod2_2
                new_mod2.site_id= 1
            form.save()
            
           
            pceversion.modify_date=timezone.now()
            pceversion.save()
                 
            info(request, _("Successfully saved Module 2 Part 1."))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_2_1.html', {'context':'Edit','form': form,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':'2','modulestr':'2 part 1'},
             context_instance=RequestContext(request))
    else:
        form = Module2_1Form(initial={'country': country,'session': pceversion.id})
      
    return render_to_response('pce/module_2_1.html', {'context':'Edit','form': form,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':'2','modulestr':'2 part 1',},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module2_1_edit(request, country, id=None,sessionid=None, template_name='pce/module_2_1.html'):
    """ Edit module_2_1 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)

    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,2)
    is_pa_filled = is_problemanalysis_filled(sessionid,2)
    is_sa_filled = is_swotanalysis_filled(sessionid,2)
    is_lf_filled = is_logicalframework_filled(sessionid,2)          
    
    if id:
        module2 = get_object_or_404(Module2_1, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,2):
            can_edit=1
            m_percentage=get_percentage_module_filled(2,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=2).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=2).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=2).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=2).id
    else:
        module2 = Module2_1()
  
    if request.POST:
        form = Module2_1Form(request.POST,  request.FILES, instance=module2)
        if form.is_valid():    
            form.save()
          
           
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 2 Part 1"))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module2_1Form(instance=module2)
      
    return render_to_response(template_name, {
      'context':'Edit','form': form,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':'2','modulestr':'2 part 1',

    }, context_instance=RequestContext(request))



### #MODULE 2 part2:
class Module2_2ListView(ListView):
    context_object_name = 'latest'
    model = Module2_2
    date_field = 'publish_date'
    template_name = 'pce/module_2_2.html'
    queryset = Module2_2.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module2_2 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module2_2.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module2_2ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id
        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'2'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'2') and session.status==1:
            can_edit=1
        if  can_see or can_edit:  
            module= 2
            form   = Module2_2FormView()
            form123= Module2_2WeaknessesFormSet()
            module2=None
            if id!='':
                module2 = get_object_or_404(Module2_2,  id=self.kwargs['id'])
                form   = Module2_2FormView(instance=module2)
                form123= Module2_2WeaknessesFormSet(instance=module2)
      
            context['form']=form
            context['form123']=form123
            context['module2']=module2
           
            context['module']='2'
            modulestr['module']='2 part 2'
     
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(2,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],2)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],2)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],2)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],2)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=2).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=2).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=2).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=2).id

      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
    
    
class Module2_2ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module2_2
    date_field = 'publish_date'
    template_name = 'pce/module_2_2.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module2_2ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
        THEM1 = (
            (0, _("--- Please select ---")),
            (1,_("None of them")),
            (2,_("One of them")),
            (3,_("Two of them")),
            (4,_("All of them")),
        )
        
        context['THEM1'] =THEM1
        
        id=''
        context['latest']= ''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module2_2.objects.filter(id=id)
    
        context['id'] = id
    
        can_see=0
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'2'):
            can_see=1
        if  can_see:  
            module= 2
            form   = Module2_2FormView()
            form123= Module2_2WeaknessesFormSet()
            module2=None
            if id!='':
                module2 = get_object_or_404(Module2_2,  id=self.kwargs['id'])
                form   = Module2_2FormView(instance=module2)
                form123= Module2_2WeaknessesFormSet(instance=module2)
      
            context['module2']=module2
            context['form']=form
            context['form123']=form123
            context['module2']=module2
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(2,self.kwargs['sessionid'])
         
        context['can_see'] = can_see
        
        return context
    
    
    
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module2_2_create(request, country,sessionid=None):
    """ Create module2  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module2_count= Module2_2.objects.filter(session=sessionid).order_by('-modify_date').count()
    
    
    
    ###ID of MOD2 part 1
    idmod2_1=-1
    module2_1= Module2_1.objects.filter(session=sessionid).order_by('-modify_date')
    if module2_1.count()>0:
        idmod2_1=module2_1[0].id
    
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,2)
    is_pa_filled = is_problemanalysis_filled(sessionid,2)
    is_sa_filled = is_swotanalysis_filled(sessionid,2)
    is_lf_filled = is_logicalframework_filled(sessionid,2)

    if module2_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,2):   
                can_edit=1
                m_percentage=get_percentage_module_filled(2,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=2).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=2).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=2).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=2).id
            
    if request.method == "POST":
        form = Module2_2Form(request.POST, request.FILES)
        form123= Module2_2WeaknessesFormSet(request.POST)
      
        if form.is_valid() and form123.is_valid():
            new_mod2 = form.save(commit=False)
            new_mod2.author = request.user
            new_mod2.session=pceversion
            if idmod2_1>0:
                new_mod2.id= idmod2_1
                new_mod2.site_id= 1
            form.save()
            
            form123.instance = new_mod2
            form123.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
                 
            info(request, _("Successfully saved Module 2 part 2."))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_2_2.html', {'context':'Edit','form': form,'form123': form123,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,       'module':'2','modulestr':'2 part 2'},
             context_instance=RequestContext(request))
    else:
        form = Module2_2Form(initial={'country': country,'session': pceversion.id})
        form123= Module2_2WeaknessesFormSet()

    return render_to_response('pce/module_2_2.html', {'context':'Edit','form': form,'form123': form123,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,  'module':'2','modulestr':'2 part 2',},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module2_2_edit(request, country, id=None,sessionid=None, template_name='pce/module_2_2.html'):
    """ Edit module_2_2 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)

    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,2)
    is_pa_filled = is_problemanalysis_filled(sessionid,2)
    is_sa_filled = is_swotanalysis_filled(sessionid,2)
    is_lf_filled = is_logicalframework_filled(sessionid,2)          
    
    if id:
        module2 = get_object_or_404(Module2_2, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,2):
            can_edit=1
            m_percentage=get_percentage_module_filled(2,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=2).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=2).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=2).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=2).id
    else:
        module2 = Module2_2()
  
    if request.POST:
        form = Module2_2Form(request.POST,  request.FILES, instance=module2)
        form123= Module2_2WeaknessesFormSet(request.POST,  request.FILES, instance=module2)
        if form.is_valid() and form123.is_valid():    
            form.save()
          
            form123.instance = module2
            form123.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 2 part 2"))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module2_2Form(instance=module2)
        form123= Module2_2WeaknessesFormSet(instance=module2)
      
    return render_to_response(template_name, {
      'context':'Edit','form': form,'form123': form123,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,  'module':'2','modulestr':'2 part 2',

    }, context_instance=RequestContext(request))




#MODULE 3
class Module3ListView(ListView):
    context_object_name = 'latest'
    model = Module3
    date_field = 'publish_date'
    template_name = 'pce/module_3.html'
    queryset = Module3.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module3 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module3.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module3ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id
        
        VAL_M3_3 = (
            (0, _("--- Please select ---")),
            (1, _("No involvement")),
            (2, _("Partial involvement (provide comments)")),
            (3, _("Consulted only on phytosanitary issues")),
            (4, _("Consulted generally on all matters (as a partner)")),
            (5, _("Active stakeholder in the process")),

        )

        VAL_M3_9 = (
            (0, _("--- Please select ---")),
            (1, _("absent from the process")),
            (2, _("limited involvement")),
            (3, _("mostly")),
            (4, _("involved")),
            (5, _("Highly involved")),

        )
      
        VAL_M3_14 = (
            (0, _("--- Please select ---")),
            (1, _("Severe limitations")),
            (2, _("Limited")),
            (3, _("Marginally adequate")),
            (4, _("Good")),
            (5, _("Excellent")),

        )
       
        VAL_M3_15 = (
            (0, _("--- Please select ---")),
            (1, _("Very")),
            (2, _("High")),
            (3, _("Medium")),
            (4, _("Low")),

        )

        
        VAL_M3_16 = (
            (0, _("--- Please select ---")),
            (1, _("Not at all")),
            (2, _("Very slightly")),
            (3, _("Slightly")),
            (4, _("Supportive")),
            (5, _("Very supportive")),

        )

        STABLE = (
            (0, _("--- Please select ---")),
            (1, _("Very unstable")),
            (2, _("Unstable")),
            (3, _("Slightly Stable")),
            (4, _("Stable")),
            (5, _("Very Stable")),
        ) 


        CONDITIONS = (
            (0, _("--- Please select ---")),
            (1, _("Not at all (poor conditions and salaries)")),
            (2, _("Somewhat adequate (conditions fine but not salaries)")),
            (3, _("Almost (salaries fine but not conditions)")),
            (4, _("Adequate (reasonable salaries and conditions acceptable to the standard of living)")),
            (5, _("Totally (salaries and conditions are competitive)")),
        ) 

        RATE = (
            (0, _("--- Please select ---")),
            (1, _("None")),
            (2, _("Insufficient")),
            (3, _("Moderate")),
            (4, _("Good")),
            (5, _("Excellent")),
        ) 

        SUPPORT = (
            (0, _("--- Please select ---")),
            (1, _("Not at all")),
            (2, _("Minimal support")),
            (3, _("Supportive")),
            (4, _("Moderately supportive")),
            (5, _("Very supportive")),
        ) 


        PARTIAL = (
            (0, _("--- Please select ---")),
            (1, _("Not at all")),
            (2, _("Partially")),
            (3, _("Selectively")),
            (4, _("Mostly")),
            (5, _("All")),
        ) 
        RATE1 = (
        (0, _("--- Please select ---")),
        (1, _("Capacity None existent")),
        (2, _("Insufficient")),
        (3, _("Moderate")),
        (4, _("Good")),
        (5, _("Excellent")),
        )
        SERVICE = (
            (0, _("--- Please select ---")),
            (1, _("None existent")),
            (2, _("Formative stages")),
            (3, _("Limited to a specific sub-area")),
            (4, _("Developed but targeted")),
            (5, _("Well developed and broad")),
        ) 
     #context['regions'] =REGIONS
        context['VAL_M3_3'] =VAL_M3_3
        context['VAL_M3_9'] =VAL_M3_9
        context['VAL_M3_14'] =VAL_M3_14
        context['VAL_M3_15'] =VAL_M3_15
        context['VAL_M3_16'] =VAL_M3_16
        context['STABLE'] =STABLE
        context['CONDITIONS'] =CONDITIONS
        context['RATE'] =RATE
        context['RATE1'] =RATE1
        context['SUPPORT'] =SUPPORT
        context['SERVICE'] =SERVICE
        context['PARTIAL'] =PARTIAL
        context['PRIORITY'] =PRIORITY
        context['M3_1']=   M3_1.objects.filter()
        context['M3_10']=  M3_10.objects.filter()
        context['M3_17']=  M3_17.objects.filter()
        
        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'3'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'3') and session.status==1:
            can_edit=1
        if  can_see or can_edit:    
    
            module= 3
            module3=None
            form   = Module3FormView()
            form31 = Module3GridFormSet()
            form33 = Module3WeaknessesFormSet()
            if id!='':
                module3 = get_object_or_404(Module3,  id=self.kwargs['id'])
                form   = Module3FormView(instance=module3)
                form31 = Module3GridFormSet(instance=module3)
                form33 = Module3WeaknessesFormSet(instance=module3)
     
            context['form']=form
            context['form31']=form31
            context['form33']=form33
            context['module3']=module3
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(3,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],3)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],3)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],3)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],3)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=3).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=3).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=3).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=3).id
      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
class Module3ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module3
    date_field = 'publish_date'
    template_name = 'pce/module_3.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module3ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
        
        VAL_M3_3 = (
            (0, _("--- Please select ---")),
            (1, _("No involvement")),
            (2, _("Partial involvement (provide comments)")),
            (3, _("Consulted only on phytosanitary issues")),
            (4, _("Consulted generally on all matters (as a partner)")),
            (5, _("Active stakeholder in the process")),

        )

        VAL_M3_9 = (
            (0, _("--- Please select ---")),
            (1, _("absent from the process")),
            (2, _("limited involvement")),
            (3, _("mostly")),
            (4, _("involved")),
            (5, _("Highly involved")),

        )
      
        VAL_M3_14 = (
            (0, _("--- Please select ---")),
            (1, _("Severe limitations")),
            (2, _("Limited")),
            (3, _("Marginally adequate")),
            (4, _("Good")),
            (5, _("Excellent")),

        )
       
        VAL_M3_15 = (
            (0, _("--- Please select ---")),
            (1, _("Very")),
            (2, _("High")),
            (3, _("Medium")),
            (4, _("Low")),

        )

        
        VAL_M3_16 = (
            (0, _("--- Please select ---")),
            (1, _("Not at all")),
            (2, _("Very slightly")),
            (3, _("Slightly")),
            (4, _("Supportive")),
            (5, _("Very supportive")),

        )

        STABLE = (
            (0, _("--- Please select ---")),
            (1, _("Very unstable")),
            (2, _("Unstable")),
            (3, _("Slightly Stable")),
            (4, _("Stable")),
            (5, _("Very Stable")),
        ) 


        CONDITIONS = (
            (0, _("--- Please select ---")),
            (1, _("Not at all (poor conditions and salaries)")),
            (2, _("Somewhat adequate (conditions fine but not salaries)")),
            (3, _("Almost (salaries fine but not conditions)")),
            (4, _("Adequate (reasonable salaries and conditions acceptable to the standard of living)")),
            (5, _("Totally (salaries and conditions are competitive)")),
        ) 

        RATE = (
            (0, _("--- Please select ---")),
            (1, _("None")),
            (2, _("Insufficient")),
            (3, _("Moderate")),
            (4, _("Good")),
            (5, _("Excellent")),
        ) 

        SUPPORT = (
            (0, _("--- Please select ---")),
            (1, _("Not at all")),
            (2, _("Minimal support")),
            (3, _("Supportive")),
            (4, _("Moderately supportive")),
            (5, _("Very supportive")),
        ) 


        PARTIAL = (
            (0, _("--- Please select ---")),
            (1, _("Not at all")),
            (2, _("Partially")),
            (3, _("Selectively")),
            (4, _("Mostly")),
            (5, _("All")),
        ) 
        RATE1 = (
        (0, _("--- Please select ---")),
        (1, _("Capacity None existent")),
        (2, _("Insufficient")),
        (3, _("Moderate")),
        (4, _("Good")),
        (5, _("Excellent")),
        )
        SERVICE = (
            (0, _("--- Please select ---")),
            (1, _("None existent")),
            (2, _("Formative stages")),
            (3, _("Limited to a specific sub-area")),
            (4, _("Developed but targeted")),
            (5, _("Well developed and broad")),
        ) 
     #context['regions'] =REGIONS
        context['VAL_M3_3'] =VAL_M3_3
        context['VAL_M3_9'] =VAL_M3_9
        context['VAL_M3_14'] =VAL_M3_14
        context['VAL_M3_15'] =VAL_M3_15
        context['VAL_M3_16'] =VAL_M3_16
        context['STABLE'] =STABLE
        context['CONDITIONS'] =CONDITIONS
        context['RATE'] =RATE
        context['RATE1'] =RATE1
        context['SUPPORT'] =SUPPORT
        context['SERVICE'] =SERVICE
        context['PARTIAL'] =PARTIAL
        context['PRIORITY'] =PRIORITY
        context['M3_1']=   M3_1.objects.filter()
        context['M3_10']=  M3_10.objects.filter()
        context['M3_17']=  M3_17.objects.filter()
                  
        
        id=''
        context['latest']=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module3.objects.filter(id=id)
        context['id'] = id
     
        can_see=0
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'3'):
            can_see=1
        if  can_see:  
            module= 3
            module3=None
            form   = Module3FormView()
            form31 = Module3GridFormSet()
            form33 = Module3WeaknessesFormSet()
            if id!='':
                module3 = get_object_or_404(Module3,  id=self.kwargs['id'])
                form   = Module3FormView(instance=module3)
                form31 = Module3GridFormSet(instance=module3)
                form33 = Module3WeaknessesFormSet(instance=module3)
     
            context['form']=form
            context['form31']=form31
            context['form33']=form33
            context['module3']=module3
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(3,self.kwargs['sessionid'])
         
        context['can_see'] = can_see
        
        return context
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module3_create(request, country,sessionid=None):
    """ Create module1  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module3_count= Module3.objects.filter(session=sessionid).order_by('-modify_date').count()
   
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,3)
    is_pa_filled = is_problemanalysis_filled(sessionid,3)
    is_sa_filled = is_swotanalysis_filled(sessionid,3)
    is_lf_filled = is_logicalframework_filled(sessionid,3)
    
    can_edit=0
    if module3_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,3):
                can_edit=1
                m_percentage=get_percentage_module_filled(3,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=3).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=3).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=3).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=3).id
    if request.method == "POST":
        form = Module3Form(request.POST, request.FILES)
        form31 = Module3GridFormSet(request.POST)
        form33 = Module3WeaknessesFormSet(request.POST)
        
        if form.is_valid() and form31.is_valid() and form33.is_valid():
            new_mod3 = form.save(commit=False)
            new_mod3.author = request.user
            new_mod3.session=pceversion
            form.save()
            
            form31.instance = new_mod3
            form31.save()
            
            form33.instance = new_mod3
            form33.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 3."))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_3.html', {'context':'Edit',  'M3_1': M3_1.objects.filter(),'M3_10': M3_10.objects.filter(),'M3_17': M3_17.objects.filter(),  'form': form,'form31': form31,'form33': form33,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':3,},
             context_instance=RequestContext(request))

    else:
        form = Module3Form(initial={'country': country,'session': pceversion.id})
        form31 = Module3GridFormSet()
        form33 = Module3WeaknessesFormSet()
    
    return render_to_response('pce/module_3.html', {'context':'Edit', 'M3_1': M3_1.objects.filter(),'M3_10': M3_10.objects.filter(),'M3_17': M3_17.objects.filter(),'form': form,'form31': form31,'form33': form33,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':3,},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module3_edit(request, country, id=None,sessionid=None, template_name='pce/module_3.html'):
    """ Edit module_3 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,3)
    is_pa_filled = is_problemanalysis_filled(sessionid,3)
    is_sa_filled = is_swotanalysis_filled(sessionid,3)
    is_lf_filled = is_logicalframework_filled(sessionid,3)
              
    
    if id:
        module3 = get_object_or_404(Module3, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,3):
            can_edit=1
            m_percentage=get_percentage_module_filled(3,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=3).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=3).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=3).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=3).id
				
    else:
        module3 = Module3()
  
    if request.POST:
        form = Module3Form(request.POST,  request.FILES, instance=module3)
        form31= Module3GridFormSet(request.POST,instance=module3)
        form33 = Module3WeaknessesFormSet(request.POST,instance=module3)
    
        
            
        if form.is_valid()and form31.is_valid() and form33.is_valid() :
            form.save()
            
            form31.instance = module3
            form31.save()
            
            form33.instance = module3
            form33.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 3."))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module3Form(instance=module3)
        form31= Module3GridFormSet(instance=module3)
        form33 = Module3WeaknessesFormSet(instance=module3)
    
    return render_to_response(template_name, {
        'context':'Edit','M3_1': M3_1.objects.filter(),'M3_10': M3_10.objects.filter(),'M3_17': M3_17.objects.filter(),'form': form,'form31': form31,'form33': form33,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':3,
    }, context_instance=RequestContext(request))

#MODULE 4

class Module4ListView(ListView):
    context_object_name = 'latest'
    model = Module4
    date_field = 'publish_date'
    template_name = 'pce/module_4.html'
    queryset = Module4.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module4 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module4.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module4ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id

        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'4'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'4') and session.status==1:
            can_edit=1
        if  can_see or can_edit: 
            module= 4
            
            form   = Module4FormView()
            form34 = Module4WeaknessesFormSet()
            module4=None
            if id!='':
                module4 = get_object_or_404(Module4,  id=self.kwargs['id'])
                form   = Module4FormView(instance=module4)
                form34 = Module4WeaknessesFormSet(instance=module4)
               
     
            context['form']=form
            context['form34']=form34
            context['module4']=module4
                    
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(4,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],4)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],4)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],4)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],4)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=4).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=4).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=4).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=4).id
      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context    
   
   
class Module4ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module4
    date_field = 'publish_date'
    template_name = 'pce/module_4.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module4ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
        
       
        STATEMENT = (
            (0,_("--- Please select ---")),
            (1,_("No statement exists")),
                (2,_("Under consideration")),
                (3,_("In a draft document")),
                (4,_("In an internal strategic plan")),
                (5,_("In a published strategic plan ")),
        )

        MODERATE = (
            (0,_("--- Please select ---")),
                (1,_("Not at all")),
                (2,_("Slightly")),
                (3,_("Moderately")),
                (4,_("Very much")),
                (5,_("Completely")),
        )

        HQ = (
                (0,_("--- Please select ---")),
                (1,_("Not at all")),
                (2,_("Very few")),
                (3,_("Only staff at HQ")),
                (4,_("HQ and some in the field")),
                (5,_("Yes, all staff")),
        )

      
        WRITTEN = (
        (0,_("--- Please select ---")),
        (1,_("Not at all")),
        (2,_("Yes, informal")),
        (3,_("Yes, written")),
        )

        INPUTSTAKE = (
                (0,_("--- Please select ---")),
                (1,_("No input by stakeholders")),
                (2,_("Stakeholders comment on plans submitted by NPPO after they are developed")),
                (3,_("Stakeholders actively involved in the planning process")),
        )
       
        THEM = (
                (0,_("--- Please select ---")),
                (1,_("None at all")),
                (2,_("A few of them")),
                (3,_("Some of them")),
                (4,_("Most of them")),
                (5,_("All of them")),
        )
       
        DEGREE = (
        (0, _("--- Please select ---")),
        (1,_("Not at all")),
        (2,_("To a small degree")),
        (3,_("To a medium degree")),
        (4,_("To a large degree")),
        (5,_("Completely")),
        )
       
        TERM = (
                (0, _("--- Please select ---")),
                (1,_("Never")),
                (2,_("Rarely")),
                (3,_("Sometimes, no fixed term ")),
                (4,_("Accordingly with an established term ")),
                (5,_("Annually")),
        )
        context['STATEMENT'] =STATEMENT
        context['MODERATE'] =MODERATE
        context['DEGREE'] =DEGREE
        context['THEM'] =THEM
        context['TERM'] =TERM
        context['WRITTEN'] =WRITTEN
        context['INPUTSTAKE'] =INPUTSTAKE
        context['HQ'] =HQ
       
        id=''
        context['latest']=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module4.objects.filter(id=id)
        context['id'] = id
    
        can_see=0
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'4'):
            can_see=1
        if  can_see:  
            module= 4
            
            form   = Module4FormView()
            form34 = Module4WeaknessesFormSet()
            module4=None
            if id!='':
                module4 = get_object_or_404(Module4,  id=self.kwargs['id'])
                form   = Module4FormView(instance=module4)
                form34 = Module4WeaknessesFormSet(instance=module4)
               
     
            context['form']=form
            context['form34']=form34
            context['module4']=module4
                    
            context['module']=module
    
      
          
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(4,self.kwargs['sessionid'])
         
        context['can_see'] = can_see
        
        return context
    



@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module4_create(request, country,sessionid=None):
    """ Create module4  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module4_count= Module4.objects.filter(session=sessionid).order_by('-modify_date').count()
   
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,4)
    is_pa_filled = is_problemanalysis_filled(sessionid,4)
    is_sa_filled = is_swotanalysis_filled(sessionid,4)
    is_lf_filled = is_logicalframework_filled(sessionid,4)

  
    
    can_edit=0
    if module4_count>0:
        can_edit=0
    else:
        if is_st_filled:
             if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,4):
                can_edit=1
                m_percentage=get_percentage_module_filled(4,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=4).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=4).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=4).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=4).id
    if request.method == "POST":
        form = Module4Form(request.POST, request.FILES)
        form34 = Module4WeaknessesFormSet(request.POST)
        

        if form.is_valid() and form34.is_valid():
            new_mod4 = form.save(commit=False)
            new_mod4.author = request.user
            new_mod4.session=pceversion
            form.save()
            
            form34.instance = new_mod4
            form34.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 4."))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_4.html', {'context':'Edit','form': form,'form34': form34,'sessionid':sessionid,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':4,},
             context_instance=RequestContext(request))
    else:
        form = Module4Form(initial={'country': country,'session': pceversion.id})
        form34 = Module4WeaknessesFormSet()
    
    return render_to_response('pce/module_4.html', {'context':'Edit','form': form,'form34': form34,'sessionid':sessionid,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':4,},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module4_edit(request, country, id=None,sessionid=None, template_name='pce/module_4.html'):
    """ Edit module_4 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,4)
    is_pa_filled = is_problemanalysis_filled(sessionid,4)
    is_sa_filled = is_swotanalysis_filled(sessionid,4)
    is_lf_filled = is_logicalframework_filled(sessionid,4)
    ##print(sessionid)
    if id:
        module4 = get_object_or_404(Module4, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,4):
            can_edit=1
            m_percentage=get_percentage_module_filled(4,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=4).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=4).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=4).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=4).id
    else:
        module4 = Module4()
  
    if request.POST:
        form = Module4Form(request.POST,  request.FILES, instance=module4)
        form34 = Module4WeaknessesFormSet(request.POST,  request.FILES, instance=module4)
       
        if form.is_valid() and form34.is_valid():
            form.save()
          
            form34.instance = module4
            form34.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 4."))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module4Form(instance=module4)
        form34 = Module4WeaknessesFormSet(instance=module4)
      
    return render_to_response(template_name, {
        'context':'Edit','form': form, 'form34': form34,'sessionid':sessionid,	'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':4,
         
    }, context_instance=RequestContext(request))
#MODULE 5
class Module5ListView(ListView):
    context_object_name = 'latest'
    model = Module5
    date_field = 'publish_date'
    template_name = 'pce/module_5.html'
    queryset = Module5.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module5 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module5.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module5ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
          
        ACHIEVE = (
            (0, _("--- Please select ---")),
            (1,_("Very difficult")),
            (2,_("Somewhat difficult")),
            (3,_("Easy")),
            (4,_("Very easy")),
        )
        CARRY_AC = (
                (0, _("--- Please select ---")),
                (1,_("None at all")),
                (2,_("Lacking in many areas")),
                (3,_("Partially")),
                (4,_("Mostly")),
                (5,_("Completely")),
        )

        DEFINED = (
                (0, _("--- Please select ---")),
                (1,_("No roles definition")),
                (2,_("Poorly defined")),
                (3,_("Defined satisfactorily")),
                (4,_("Well defined")),
                (5,_("Well defined and flexible")),
        )
        CLEAR = (
                (0, _("--- Please select ---")),
                (1,_("Not at all")),
                (2,_("Unsatisfactory")),
                (3,_("With difficulty")),
                (4,_("Satisfactory")),
                (5,_("Very clear and at all levels")),
        )

        EXPEDI = (
                (0, _("--- Please select ---")),
                (1,_("Not at all")),
                (2,_("Poorly")),
                (3,_("Satisfactorily")),
                (4,_("Good")),
                (5,_("Fully")),
        )
        LINKAGE = (
                (0, _("--- Please select ---")),
                (1,_("Not at all")),
                (2,_("Unsatisfactory")),
                (3,_("With difficulty")),
                (4,_("Satisfactory")),
                (5,_("Easily")),
        )
        POLICY = (
                (0, _("--- Please select ---")),
                (1,_("Not at all")),
                (2,_("Somewhat")),
                (3,_("Substantive")),
                (4,_("Totally")),
        )
        THEM = (
                (0,("--- Please select ---")),
                (1,("None at all")),
                (2,("A few of them")),
                (3,("Some of them")),
                (4,("Most of them")),
                (5,("All of them")),
        )

        
        context['ACHIEVE'] =ACHIEVE
        context['CARRY_AC'] =CARRY_AC
       
        context['DEFINED'] =DEFINED
        context['CLEAR'] =CLEAR
        context['EXPEDI'] =EXPEDI
        context['LINKAGE'] =LINKAGE
        context['POLICY'] =POLICY
        context['THEM'] =THEM
        
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id

        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        
        if canSee(session.id,session.country,self.request.user,'5'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'5') and session.status==1:
            can_edit=1
        if  can_see or can_edit:   
    
            module= 5
            module5=None
            form   = Module5FormView()
            form25 = Module5WeaknessesFormSet()
            if id!='':
                module5 = get_object_or_404(Module5,  id=self.kwargs['id'])
                form   = Module5FormView(instance=module5)
                form25 = Module5WeaknessesFormSet(instance=module5)
    
            context['form']=form
            context['form25']=form25
            context['module5']=module5
            context['M5_3']=  M5_3.objects.filter()
         
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(5,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],5)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],5)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],5)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],5)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=5).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=5).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=5).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=5).id
      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
 

 
    
class Module5ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module5
    date_field = 'publish_date'
    template_name = 'pce/module_5.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module5ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
        
        ACHIEVE = (
            (0, _("--- Please select ---")),
            (1,_("Very difficult")),
            (2,_("Somewhat difficult")),
            (3,_("Easy")),
            (4,_("Very easy")),
        )
        CARRY_AC = (
                (0, _("--- Please select ---")),
                (1,_("None at all")),
                (2,_("Lacking in many areas")),
                (3,_("Partially")),
                (4,_("Mostly")),
                (5,_("Completely")),
        )

        DEFINED = (
                (0, _("--- Please select ---")),
                (1,_("No roles definition")),
                (2,_("Poorly defined")),
                (3,_("Defined satisfactorily")),
                (4,_("Well defined")),
                (5,_("Well defined and flexible")),
        )
        CLEAR = (
                (0, _("--- Please select ---")),
                (1,_("Not at all")),
                (2,_("Unsatisfactory")),
                (3,_("With difficulty")),
                (4,_("Satisfactory")),
                (5,_("Very clear and at all levels")),
        )

        EXPEDI = (
                (0, _("--- Please select ---")),
                (1,_("Not at all")),
                (2,_("Poorly")),
                (3,_("Satisfactorily")),
                (4,_("Good")),
                (5,_("Fully")),
        )
        LINKAGE = (
                (0, _("--- Please select ---")),
                (1,_("Not at all")),
                (2,_("Unsatisfactory")),
                (3,_("With difficulty")),
                (4,_("Satisfactory")),
                (5,_("Easily")),
        )
        POLICY = (
                (0, _("--- Please select ---")),
                (1,_("Not at all")),
                (2,_("Somewhat")),
                (3,_("Substantive")),
                (4,_("Totally")),
        )
        THEM = (
                (0,("--- Please select ---")),
                (1,("None at all")),
                (2,("A few of them")),
                (3,("Some of them")),
                (4,("Most of them")),
                (5,("All of them")),
        )

        
        context['ACHIEVE'] =ACHIEVE
        context['CARRY_AC'] =CARRY_AC
       
        context['DEFINED'] =DEFINED
        context['CLEAR'] =CLEAR
        context['EXPEDI'] =EXPEDI
        context['LINKAGE'] =LINKAGE
        context['POLICY'] =POLICY
        context['THEM'] =THEM
        
        id=''
        context['latest']=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module5.objects.filter(id=id)
        context['id'] = id
    
        can_see=0
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'5'):
            can_see=1
        if  can_see:  
            module= 5
            module5=None
            form   = Module5FormView()
            form25 = Module5WeaknessesFormSet()
            if id!='':
                module5 = get_object_or_404(Module5,  id=self.kwargs['id'])
                form   = Module5FormView(instance=module5)
                form25 = Module5WeaknessesFormSet(instance=module5)
    
            context['form']=form
            context['form25']=form25
            context['module5']=module5
         
            context['module']=module
            context['M5_3']=  M5_3.objects.filter()
         
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(5,self.kwargs['sessionid'])
         
        context['can_see'] = can_see
        
        return context
    
    
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module5_create(request, country,sessionid=None):
    """ Create module5  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module5_count= Module5.objects.filter(session=sessionid).order_by('-modify_date').count()
   
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,5)
    is_pa_filled = is_problemanalysis_filled(sessionid,5)
    is_sa_filled = is_swotanalysis_filled(sessionid,5)
    is_lf_filled = is_logicalframework_filled(sessionid,5)

    
  
    if module5_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,5):
                   can_edit=1
                   m_percentage=get_percentage_module_filled(5,sessionid)
                   tot_percentage=get_tot_percentage(sessionid)
                   if is_st_filled:
                       st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=5).id
                   if is_pa_filled:
                       pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=5).id
                   if is_sa_filled:
                       sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=5).id
                   if is_lf_filled:
                       lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=5).id
    if request.method == "POST":
        form = Module5Form(request.POST, request.FILES)
        form25 = Module5WeaknessesFormSet(request.POST)
        

        if form.is_valid() and form25.is_valid():
            new_mod5 = form.save(commit=False)
            new_mod5.author = request.user
            new_mod5.session=pceversion
            form.save()
            
            form25.instance = new_mod5
            form25.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 5."))

            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_5.html', {'context':'Edit','M5_3': M5_3.objects.filter(),'form': form,'form25': form25,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':5,},
             context_instance=RequestContext(request))

    else:
        form = Module5Form(initial={'country': country,'session': pceversion.id})
        form25 = Module5WeaknessesFormSet()
        

    return render_to_response('pce/module_5.html', {'context':'Edit','M5_3': M5_3.objects.filter(),'form': form,'form25': form25,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':5,},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module5_edit(request, country, id=None,sessionid=None, template_name='pce/module_5.html'):
    """ Edit module_5 """
    user = request.user
    
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)

    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,5)
    is_pa_filled = is_problemanalysis_filled(sessionid,5)
    is_sa_filled = is_swotanalysis_filled(sessionid,5)
    is_lf_filled = is_logicalframework_filled(sessionid,5)
       
    
    if id:
        module5 = get_object_or_404(Module5, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,5):
            can_edit=1
            m_percentage=get_percentage_module_filled(5,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=5).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=5).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=5).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=5).id
	
    else:
        module5 = Module5()
  
    if request.POST:
        form = Module5Form(request.POST,  request.FILES, instance=module5)
        form25 = Module5WeaknessesFormSet(request.POST,  request.FILES, instance=module5)
        
       
        if form.is_valid() and form25.is_valid():
            form.save()
          
            form25.instance = module5
            form25.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 4."))
            
            
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module5Form(instance=module5)
        form25 = Module5WeaknessesFormSet(instance=module5)
      
    return render_to_response(template_name, {
        'context':'Edit','M5_3': M5_3.objects.filter(),'form': form, 'form25': form25,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':5,'module':5,
    }, context_instance=RequestContext(request))


#MODULE 5
class Module6ListView(ListView):
    context_object_name = 'latest'
    model = Module6
    date_field = 'publish_date'
    template_name = 'pce/module_6.html'
    queryset = Module6.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module6 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module6.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module6ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id

        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'6'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'6') and session.status==1:
            can_edit=1
        if  can_see or can_edit:    
    
            module= 6
            module6=None
            form   = Module6FormView()
            form37 = Module6WeaknessesFormSet()
            if id!='':
                module6 = get_object_or_404(Module6,  id=self.kwargs['id'])
                form   = Module6FormView(instance=module6)
                form37 = Module6WeaknessesFormSet(instance=module6)
        
            context['form']=form
            context['form37']=form37
            context['module6']=module6
          
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(6,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],6)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],6)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],6)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],6)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=6).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=6).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=6).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=6).id
      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
    
    
class Module6ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module6
    date_field = 'publish_date'
    template_name = 'pce/module_6.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module6ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
        
        BUDGET = (
        (0, _("--- Please select ---")),
        (1, _("less than 10 % of the NPPO budget")),
        (2, _("between 11 - 20 %")),
        (3, _("between 21 - 40 %")),
        (4, _("between 41 - 50 %")),
        (5, _("More than 51 %")),
        )
        FUNDING = (
        (0, _("--- Please select ---")),
        (1, _("Depends on Ministerial allocation ")),
        (2, _("Specific line in the national budget ")),
        (3, _("Specific line in the national budget plus fees and fines")),
        )
        ACQUIRE = (
        (0, _("--- Please select ---")),
        (1, _("no effort")),
        (2, _("International Grants")),
        (3, _("International Grants and National Resources")),
        (4, _("International Loans")),
        (5, _("National Resources")),
        )
        ADEQUATELY = (
        (0, _("--- Please select ---")),
        (1, _("Not at all")),
        (2, _("marginally")),
        (3, _("sufficiently")),
        (4, _("well staff")),
        )
        APPOINTED = (
        (0, _("--- Please select ---")),
        (1, _("All the staff appointed are not competitively sourced")),
        (2, _("The Head and managers are not competitively sourced ")),
        (3, _("Only the Head is appointed without being competitively sourced")),
        (4, _("All management positions are competitively sourced and appointed")),
        (5, _("All the positions in the NPPO are competitively sourced and appointed")),
        )
        TURNOVER = (
        (0, _("--- Please select ---")),
        (1, _("very low")),
        (2, _("low")),
        (3, _("intermediate")),
        (4, _("high")),
        (5, _("very high")),
        )
        TRAINING = (
        (0, _("--- Please select ---")),
        (1, _("Not at all")),
        (2, _("Training is done only when there are external project resources to do so")),
        (3, _("Training is done adhoc/when needed")),
        (4, _("Limited or targeted training is done")),
        (5, _("In all areas")),
        )

        PROPORTION = (
        (0, _("--- Please select ---")),
        (1, _("Less than 10 %")),
        (2, _("10 - 25 %")),
        (3, _("26 - 50 %")),
        (4, _("51 - 75 %")),
        (5, _("More than 75%")),
        )
        
        COMM_SKILLS  = (
        (0, _("--- Please select ---")),
        (1, _("very poor ")),
        (2, _("poor")),
        (3, _("average")),
        (4, _("good")),
        (5, _("very good")),
        )
        
        LANG_SKILLS = (
        (0, _("--- Please select ---")),
        (1, _("very poor")),
        (2, _("basic")),
        (3, _("average")),
        (4, _("good")),
        (5, _("advanced")),
        )
       
        RESOURCES = (
        (0, _("--- Please select ---")),
        (1, _("Not at all")),
        (2, _("unsatisfactory resource allocation")),
        (3, _("insufficient resource allocation")),
        (4, _("adequate resource allocation")),
        (5, _("More than adequatate")),
        )
       
        RECORD = (
        (0, _("--- Please select ---")),
        (1, _("No record keeping system")),
        (2, _("unsatisfactory system (incongruous)")),
        (3, _("satisfactory system (mostly hardcopy based)")),
        (4, _("very good system (combination electronic and hardcopy)")),
        (5, _("Computerized record keeping, central dbase, on line system")),
        )
       
        CAPACITY = (
        (0, _("--- Please select ---")),
        (1, _("very low")),
        (2, _("low")),
        (3, _("intermediate")),
        (4, _("high")),
        (5, _("very high")),
        )

        ACCESS = (
        (0, _("--- Please select ---")),
        (1, _("No access at all")),
        (2, _("Only when external resources are available (e.g. projects)")),
        (3, _("Limited access to paper based and internet resources")),
        (4, _("Good access to paper based and internet sources of information")),
        (4, _("Excellent, no limitations")),
        )
        BAD = (
        (0, _("--- Please select ---")),
        (1, _("Very bad")),
        (2, _("Bad")),
        (3, _("Not so bad")),
        (4, _("Good")),
        (5, _("Very good")),
        )

                     
        context['BUDGET'] =BUDGET
        context['FUNDING'] =FUNDING
        context['ACQUIRE'] =ACQUIRE
        context['ADEQUATELY'] =ADEQUATELY
        context['APPOINTED'] =APPOINTED
        context['TURNOVER'] =TURNOVER
        context['TRAINING'] =TRAINING
        context['PROPORTION'] =PROPORTION
        context['COMM_SKILLS'] =COMM_SKILLS
        context['LANG_SKILLS'] =LANG_SKILLS
        context['RESOURCES'] =RESOURCES
        context['RECORD'] =RECORD
        context['CAPACITY'] =CAPACITY
        context['ACCESS'] =ACCESS
        context['BAD'] =BAD

        
       
        id=''
        context['latest']=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module6.objects.filter(id=id)
        context['id'] = id
    
        can_see=0
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'6'):
            can_see=1
        if  can_see:  
            module= 6
            module6=None
            form   = Module6FormView()
            form37 = Module6WeaknessesFormSet()
            if id!='':
                module6 = get_object_or_404(Module6,  id=self.kwargs['id'])
                form   = Module6FormView(instance=module6)
                form37 = Module6WeaknessesFormSet(instance=module6)

            context['form']=form
            context['form37']=form37
            context['module6']=module6

            context['module']=module

            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(6,self.kwargs['sessionid'])

        context['can_see'] = can_see
        
        return context
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module6_create(request, country,sessionid=None):
    """ Create module6  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module6_count= Module6.objects.filter(session=sessionid).order_by('-modify_date').count()
   
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,6)
    is_pa_filled = is_problemanalysis_filled(sessionid,6)
    is_sa_filled = is_swotanalysis_filled(sessionid,6)
    is_lf_filled = is_logicalframework_filled(sessionid,6)

  
    if module6_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,6):
                can_edit=1
                m_percentage=get_percentage_module_filled(6,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=6).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=6).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=6).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=6).id

				
    if request.method == "POST":
        form = Module6Form(request.POST, request.FILES)
        form37 = Module6WeaknessesFormSet(request.POST)
        

        if form.is_valid() and form37.is_valid():
            new_mod6 = form.save(commit=False)
            new_mod6.author = request.user
            new_mod6.session=pceversion
            form.save()
            
            form37.instance = new_mod6
            form37.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 6."))

            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_6.html', {'context':'Edit','form': form,'form37': form37,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':6,},
             context_instance=RequestContext(request))

    else:
        form = Module6Form(initial={'country': country,'session': pceversion.id})
        form37 = Module6WeaknessesFormSet()
        

    return render_to_response('pce/module_6.html', {'context':'Edit','form': form,'form37': form37,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':6,},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module6_edit(request, country, id=None,sessionid=None, template_name='pce/module_6.html'):
    """ Edit module_6 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)

    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,6)
    is_pa_filled = is_problemanalysis_filled(sessionid,6)
    is_sa_filled = is_swotanalysis_filled(sessionid,6)
    is_lf_filled = is_logicalframework_filled(sessionid,6)
          
    
    if id:
        module6 = get_object_or_404(Module6, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,6):
            can_edit=1
            m_percentage=get_percentage_module_filled(6,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=6).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=6).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=6).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=6).id
    else:
        module6 = Module6()
  
    if request.POST:
        form = Module6Form(request.POST,  request.FILES, instance=module6)
        form37 = Module6WeaknessesFormSet(request.POST,  request.FILES, instance=module6)
        
       
        if form.is_valid() and form37.is_valid():
            form.save()
          
            form37.instance = module6
            form37.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 6."))
            
            
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module6Form(instance=module6)
        form37 = Module6WeaknessesFormSet(instance=module6)
      
    return render_to_response(template_name, {
        'context':'Edit','form': form, 'form37': form37,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':6,

    }, context_instance=RequestContext(request))

#MODULE 5
class Module7ListView(ListView):
    context_object_name = 'latest'
    model = Module7
    date_field = 'publish_date'
    template_name = 'pce/module_7.html'
    queryset = Module7.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module7 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module7.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module7ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
   

        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id

        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        
        if canSee(session.id,session.country,self.request.user,'7'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'7') and session.status==1:
            can_edit=1
        if  can_see or can_edit: 
            module= 7
            module7=None
            form   = Module7FormView()
            form14 = Module7GridFormSet()
            form23 = Module7Matrix23FormSet()
            form37 = Module7Matrix37FormSet()
            form39 = Module7Matrix39FormSet()
            form41 = Module7Matrix41FormSet()
            form43 = Module7Matrix43FormSet()
            form45 = Module7Matrix45FormSet()
            form69 = Module7WeaknessesFormSet()
            if id!='':
                module7 = get_object_or_404(Module7,  id=self.kwargs['id'])
                form   = Module7FormView(instance=module7)
                form14 = Module7GridFormSet(instance=module7)
                form23 = Module7Matrix23FormSet(instance=module7)
                form37 = Module7Matrix37FormSet(instance=module7)
                form39 = Module7Matrix39FormSet(instance=module7)
                form41 = Module7Matrix41FormSet(instance=module7)
                form43 = Module7Matrix43FormSet(instance=module7)
                form45 = Module7Matrix45FormSet(instance=module7)
                form69 = Module7WeaknessesFormSet(instance=module7)
     
            context['form']=form
            context['form14']=form14
            context['form23']=form23
            context['form37']=form37
            context['form39']=form39
            context['form41']=form41
            context['form43']=form43
            context['form45']=form45
            context['form69']=form69
            
            context['module7']=module7
         
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(7,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],7)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],7)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],7)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],7)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=7).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=7).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=7).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=7).id
      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
   
    
class Module7ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module7
    date_field = 'publish_date'
    template_name = 'pce/module_7.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module7ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
        context['context'] = 'Pdf'   
        
        GEO =(
        (0, _("--- Please select ---")),
        (1, _("Not at all")),
        (2, _("Poor geographic coverage of services")),
        (3, _("Most areas not serviced adequately")),
        (4, _("Most areas adequately serviced")),
        (5, _("Totally")),
        )

        POOR  = (
        (0, _("--- Please select ---")),
        (1, _("very poor ")),
        (2, _("poor")),
        (3, _("adequate")),
        (4, _("good")),
        (5, _("very good")),
        )
        BAD1 = (
        (0, _("--- Please select ---")),
        (1, _("Very bad")),
        (2, _("Bad")),
        (3, _("Satisfactory")),
        (4, _("Adequate")),
        (5, _("Very good")),
        )
        SUFFICIENT=(
        (0, _("--- Please select ---")),
        (1, _("Not at all sufficient")),
        (2, _("Lacking a lot")),
        (3, _("Neither insufficient or sufficient")),
        (4, _("Needs to be improved")),
        (5, _("Very sufficient")),

        )
        INSUFF=(
        (0, _("--- Please select ---")),
        (1, _("Insufficient ")),
        (2, _("Lacking in most areas")),
        (3, _("Satisfactory (key areas addressed)")),
        (4, _("Lacking in areas for comprehensive coverage")),
        (5, _("Sufficient")),
        )
        INSUFF1=(
        (0, _("--- Please select ---")),
        (1, _("Insufficient")),
        (2, _("Some basic skills")),
        (3, _("Basic skills")),
        (4, _("Adequate training in key areas")),
        (5, _("Very sufficient")),
        )

        TRAINING1=(
        (0, _("--- Please select ---")),
        (1, _("No training at all")),
        (2, _("Internal training")),
        (3, _("External training")),
        (4, _("Regular training programmes")),
        (5, _("Special training programmes")),
        )


        EQUIP=(
        (0, _("--- Please select ---")),
        (1, _("Not at all for both")),
        (2, _("We maintain equipment internally")),
        (3, _("We outsource maintenance (no internal capacity)")),
        (4, _("We have capacity for both (internal and outsource)")),
        )


        CONDICIVE=(
        (0, _("--- Please select ---")),
        (1, _("Totally not conducive")),
        (2, _("Unsatisfactory")),
        (3, _("Adequate")),
        (4, _("Good")),
        (5, _("Very good")),
        )

        WEAK=(
        (0, _("--- Please select ---")),
        (1, _("Very weak")),
        (2, _("Weak")),
        (3, _("Intermediate")),
        (4, _("Good")),
        (5, _("Very good")),
        )

        WEAK1=(
        (0, _("--- Please select ---")),
        (1, _("Very weak relevance")),
        (2, _("Weak")),
        (3, _("Intermediate")),
        (4, _("Good relevance")),
        (5, _("Excellent")),
        )

        WEAK2=(
        (0, _("--- Please select ---")),
        (1, _("Very weak perfomance")),
        (2, _("Weak")),
        (3, _("Intermediate")),
        (4, _("Good performance")),
        (5, _("Excellent")),
        )
                     
     
        context['GEO'] =GEO
        context['POOR'] =POOR
        context['BAD1'] =BAD1
        context['SUFFICIENT'] =SUFFICIENT
        context['INSUFF'] =INSUFF
        context['INSUFF1'] =INSUFF1
        context['TRAINING1'] =TRAINING1
        context['EQUIP'] =EQUIP
        context['CONDICIVE'] =CONDICIVE
        context['WEAK'] =WEAK
        context['WEAK1'] =WEAK1
        context['WEAK2'] =WEAK2

        
        id=''
        context['latest']=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module7.objects.filter(id=id)
        context['id'] = id
    
        can_see=0
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'7'):
            can_see=1
        if  can_see:  
            module= 7
            module7=None
            form   = Module7FormView()
            form14 = Module7GridFormSet()
            form23 = Module7Matrix23FormSet()
            form37 = Module7Matrix37FormSet()
            form39 = Module7Matrix39FormSet()
            form41 = Module7Matrix41FormSet()
            form43 = Module7Matrix43FormSet()
            form45 = Module7Matrix45FormSet()
            form69 = Module7WeaknessesFormSet()
            if id!='':
                module7 = get_object_or_404(Module7,  id=self.kwargs['id'])
                form   = Module7FormView(instance=module7)
                form14 = Module7GridFormSet(instance=module7)
                form23 = Module7Matrix23FormSet(instance=module7)
                form37 = Module7Matrix37FormSet(instance=module7)
                form39 = Module7Matrix39FormSet(instance=module7)
                form41 = Module7Matrix41FormSet(instance=module7)
                form43 = Module7Matrix43FormSet(instance=module7)
                form45 = Module7Matrix45FormSet(instance=module7)
                form69 = Module7WeaknessesFormSet(instance=module7)
     
            context['form']=form
            context['form14']=form14
            context['form23']=form23
            context['form37']=form37
            context['form39']=form39
            context['form41']=form41
            context['form43']=form43
            context['form45']=form45
            context['form69']=form69

            context['module7']=module7

            context['module']=module
         
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(7,self.kwargs['sessionid'])

        context['can_see'] = can_see
        
        return context
    
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module7_create(request, country,sessionid=None):
    """ Create module7  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module7_count= Module7.objects.filter(session=sessionid).order_by('-modify_date').count()
   
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,7)
    is_pa_filled = is_problemanalysis_filled(sessionid,7)
    is_sa_filled = is_swotanalysis_filled(sessionid,7)
    is_lf_filled = is_logicalframework_filled(sessionid,7)

  
    if module7_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,7):
                can_edit=1
                m_percentage=get_percentage_module_filled(7,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=7).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=7).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=7).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=7).id
				
    if request.method == "POST":
        form = Module7Form(request.POST, request.FILES)
        form14 = Module7GridFormSet(request.POST)
        form23 = Module7Matrix23FormSet(request.POST)
        form37 = Module7Matrix37FormSet(request.POST)
        form39 = Module7Matrix39FormSet(request.POST)
        form41= Module7Matrix41FormSet(request.POST)
        form43 = Module7Matrix43FormSet(request.POST)
        form45 = Module7Matrix45FormSet(request.POST)
        form69 = Module7WeaknessesFormSet(request.POST)
      
        if form.is_valid() and form14.is_valid() and  form23.is_valid() and  form37.is_valid() and form39.is_valid() and form41.is_valid() and form43.is_valid() and form45.is_valid() and form69.is_valid():
            new_mod7 = form.save(commit=False)
            new_mod7.author = request.user
            new_mod7.session=pceversion
            form.save()
            
            form14.instance = new_mod7
            form14.save()
            form23.instance = new_mod7
            form23.save()
            form37.instance = new_mod7
            form37.save()
            form39.instance = new_mod7
            form39.save()
            form41.instance = new_mod7
            form41.save()
            form43.instance = new_mod7
            form43.save()
            form45.instance = new_mod7
            form45.save()
            form69.instance = new_mod7
            form69.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 7."))

            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_7.html', {'context':'Edit','form': form,'form14': form14,'form23': form23,'form37': form37,'form39': form39,'form41': form41,'form43': form43,'form45': form45,'form69': form69,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':7,},
             context_instance=RequestContext(request))

    else:
        form = Module7Form(initial={'country': country,'session': pceversion.id})
        form14 = Module7GridFormSet()
        form23 = Module7Matrix23FormSet()
        form37 = Module7Matrix37FormSet()
        form39 = Module7Matrix39FormSet()
        form41 = Module7Matrix41FormSet()
        form43 = Module7Matrix43FormSet()
        form45 = Module7Matrix45FormSet()
        form69 = Module7WeaknessesFormSet()
      

    return render_to_response('pce/module_7.html', {'context':'Edit','form': form,'form14': form14,'form23': form23,'form37': form37,'form39': form39,'form41': form41,'form43': form43,'form45': form45,'form69': form69,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':7,},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module7_edit(request, country, id=None,sessionid=None, template_name='pce/module_7.html'):
    """ Edit module_7 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,7)
    is_pa_filled = is_problemanalysis_filled(sessionid,7)
    is_sa_filled = is_swotanalysis_filled(sessionid,7)
    is_lf_filled = is_logicalframework_filled(sessionid,7)
              
    
    if id:
        module7 = get_object_or_404(Module7, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,7):
            can_edit=1
            m_percentage=get_percentage_module_filled(7,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=7).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=7).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=7).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=7).id
    else:
        module7 = Module7()
  
    if request.POST:
        form = Module7Form(request.POST,  request.FILES, instance=module7)
        form14 = Module7GridFormSet(request.POST,  request.FILES, instance=module7)
        form23 = Module7Matrix23FormSet(request.POST,  request.FILES, instance=module7)
        form37 = Module7Matrix37FormSet(request.POST,  request.FILES, instance=module7)
        form39 = Module7Matrix39FormSet(request.POST,  request.FILES, instance=module7)
        form41 = Module7Matrix41FormSet(request.POST,  request.FILES, instance=module7)
        form43 = Module7Matrix43FormSet(request.POST,  request.FILES, instance=module7)
        form45 = Module7Matrix45FormSet(request.POST,  request.FILES, instance=module7)
        form69 = Module7WeaknessesFormSet(request.POST,  request.FILES, instance=module7)
    
       
        if form.is_valid() and form14.is_valid() and  form23.is_valid() and  form37.is_valid() and form39.is_valid() and form41.is_valid() and form43.is_valid() and form45.is_valid() and form69.is_valid():
            form.save()
          
            form14.instance = module7
            form14.save()
            form23.instance = module7
            form23.save()
            form37.instance = module7
            form37.save()
            form39.instance = module7
            form39.save()
            form41.instance = module7
            form41.save()
            form43.instance = module7
            form43.save()
            form45.instance = module7
            form45.save()
            form69.instance = module7
            form69.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
               
    
            info(request, _("Successfully saved Module 7."))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module7Form(instance=module7)
        form14 = Module7GridFormSet(instance=module7)
        form23 = Module7Matrix23FormSet(instance=module7)
        form37 = Module7Matrix37FormSet(instance=module7)
        form39 = Module7Matrix39FormSet(instance=module7)
        form41 = Module7Matrix41FormSet(instance=module7)
        form43 = Module7Matrix43FormSet(instance=module7)
        form45 = Module7Matrix45FormSet(instance=module7)
        form69 = Module7WeaknessesFormSet(instance=module7)
    
    return render_to_response(template_name, {
        'context':'Edit','form': form, 'form14': form14,'form23': form23,'form37': form37,'form39': form39,'form41': form41,'form43': form43,'form45': form45,'form69': form69,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':7,
    }, context_instance=RequestContext(request))

#MODULE 8
class Module8ListView(ListView):
    context_object_name = 'latest'
    model = Module8
    date_field = 'publish_date'
    template_name = 'pce/module_8.html'
    queryset = Module8.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module8 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module8.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module8ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
        context['M8_17']=  M8_17.objects.filter()
     
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id

        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        
        if canSee(session.id,session.country,self.request.user,'8'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'8') and session.status==1:
            can_edit=1
        if  can_see or can_edit:    
    
            module= 8
            module8=None
            form   = Module8FormView()
            form3 = Module8Grid3FormSet()
            form18 = Module8Grid18FormSet()
            form30 = Module8Matrix30FormSet()
            form45 = Module8WeaknessesFormSet()
      
            if id!='':
                module8 = get_object_or_404(Module8,  id=self.kwargs['id'])
                form   = Module8FormView(instance=module8)
                form3 = Module8Grid3FormSet(instance=module8)
                form18 = Module8Grid18FormSet(instance=module8)
                form30 = Module8Matrix30FormSet(instance=module8)
                form45 = Module8WeaknessesFormSet(instance=module8)
      
            context['module8']=module8
         
            context['form']=form
            context['form3']=form3
            context['form18']=form18
            context['form30']=form30
            context['form45']=form45
            
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(8,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],8)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],8)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],8)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],8)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=8).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=8).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=8).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=8).id
      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context

class Module8ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module8
    date_field = 'publish_date'
    template_name = 'pce/module_8.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module8ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
        
        RANGE1=(
        (0, _("--- Please select ---")),
        (1,_("0")),
        (2, _("1-25")),
        (3,_("26-50")),
        (4, _("51-75")),
        (5, _(">75")),
        )


        NUMPLANT = (
            (0, _("--- Please select ---")),
            (1, _("0")),
            (2, _("1")),
            (3, _("2")),
            (4, _("3")),
            (5, _("4")),
            (6, _("5")),
            (7, _("10 +")),
            (8, _("20 +")),
            (9, _("50 +")),
        )
        WEAK3=(
        (0, _("--- Please select ---")),
        (1, _("Very weak")),
        (2, _("Weak")),
        (3, _("Average")),
        (4, _("Good")),
        (5, _("Very strong")),
        )

        INSUFF3=(
        (0, _("--- Please select ---")),
        (1,_("Not at all")),
        (2,_("Insufficient")),
        (3,_("With difficulty")),
        (4,_("Almost")),
        (5,_("Completely")),
        )
        PERC0=(
        (0, _("--- Please select ---")),
        (1, _("None")),
        (2, _("25%")),
        (3, _("50%")),
        (4, _("75%")),
        (5, _("All")),
        )

        PROGRAMMED=(
        (0, _("--- Please select ---")),
        (1,_("No programmed training")),
        (2,_("Once every five years")),
        (3,_("Once every three years")),
        (4,_("Once every two years")),
        (5,_("At least once per year")),

        )
        WEAK=(
        (0, _("--- Please select ---")),
        (1, _("Very weak")),
        (2, _("Weak")),
        (3, _("Intermediate")),
        (4, _("Good")),
        (5, _("Very good")),
        )
        NOTATALL=(
        (0, _("--- Please select ---")),
        (1,_("Not at all")),
        (2,_("0.25")),
        (3,_("0.5")),
        (4,_("0.75")),
        (5,_("1")),
        )
        NOTATALL1=(
        (0, _("--- Please select ---")),
        (1,_("Not at all")),
        (2,_("Marginally")),
        (3,_("Intermediate")),
        (4,_("Sufficient")),
        (5,_("More than Sufficient")),
        )
        WEAK4=(
        (0, _("--- Please select ---")),
        (1,_("Very weak")),
        (2,_("0.25")),
        (3,_("0.5")),
        (4,_("0.75")),
        (5,_("1")),
        )

     
        
        context['RANGE1'] =RANGE1
        context['NUMPLANT'] =NUMPLANT
        context['WEAK3'] =WEAK3
        context['INSUFF3'] =INSUFF3
        context['PERC0'] =PERC0 
        context['PROGRAMMED'] =PROGRAMMED
        
        context['NOTATALL'] =NOTATALL
        context['NOTATALL1'] =NOTATALL1
        context['WEAK'] =WEAK
        context['WEAK4'] =WEAK4
        context['M8_17']=  M8_17.objects.filter()
     
        
        id=''
        context['latest']=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module8.objects.filter(id=id)
        context['id'] = id
    
        can_see=0
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'8'):
            can_see=1
        if  can_see:  
            module= 8
            module8=None
            form   = Module8FormView()
            form3 = Module8Grid3FormSet()
            form18 = Module8Grid18FormSet()
            form30 = Module8Matrix30FormSet()
            form45 = Module8WeaknessesFormSet()
      
            if id!='':
                module8 = get_object_or_404(Module8,  id=self.kwargs['id'])
                form   = Module8FormView(instance=module8)
                form3 = Module8Grid3FormSet(instance=module8)
                form18 = Module8Grid18FormSet(instance=module8)
                form30 = Module8Matrix30FormSet(instance=module8)
                form45 = Module8WeaknessesFormSet(instance=module8)
      
            context['module8']=module8
         
            context['form']=form
            context['form3']=form3
            context['form18']=form18
            context['form30']=form30
            context['form45']=form45
            
            context['module']=module
         
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(8,self.kwargs['sessionid'])

        context['can_see'] = can_see
        
        return context
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module8_create(request, country,sessionid=None):
    """ Create module8  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module8_count= Module8.objects.filter(session=sessionid).order_by('-modify_date').count()
   
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,8)
    is_pa_filled = is_problemanalysis_filled(sessionid,8)
    is_sa_filled = is_swotanalysis_filled(sessionid,8)
    is_lf_filled = is_logicalframework_filled(sessionid,8)

    if module8_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,8):
                can_edit=1
                m_percentage=get_percentage_module_filled(8,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=8).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=8).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=8).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=8).id


    if request.method == "POST":
        form = Module8Form(request.POST, request.FILES)
        form3 = Module8Grid3FormSet(request.POST)
        form18 = Module8Grid18FormSet(request.POST)
        form30 = Module8Matrix30FormSet(request.POST)
        form45 = Module8WeaknessesFormSet(request.POST)
        #print('form.is_valid()'+str(form.is_valid()))
        #print('form3.is_valid()'+str(form3.is_valid()))
        #print('form18.is_valid()'+str(form18.is_valid()))
        #print('form30.is_valid()'+str(form30.is_valid()))
        #print('form45.is_valid()'+str(form45.is_valid()))
        
        if form.is_valid() and form3.is_valid() and  form18.is_valid() and  form30.is_valid() and form45.is_valid():
            new_mod8 = form.save(commit=False)
            new_mod8.author = request.user
            new_mod8.session=pceversion
            form.save()
            
            form3.instance = new_mod8
            form3.save()
            form18.instance = new_mod8
            form18.save()
            form30.instance = new_mod8
            form30.save()
            form45.instance = new_mod8
            form45.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
           
            info(request, _("Successfully saved Module 8."))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
            info(request, _("AAAAAAAAAAAAAAAAAA"))
            return render_to_response('pce/module_8.html', {'context':'Edit','M8_17': M8_17.objects.filter(),'form': form,'form3': form3,'form18': form18,'form30': form30,'form45': form45,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':8,},
            context_instance=RequestContext(request))

    else:
        form = Module8Form(initial={'country': country,'session': pceversion.id})
        form3 = Module8Grid3FormSet()
        form18 = Module8Grid18FormSet()
        form30 = Module8Matrix30FormSet()
        form45 = Module8WeaknessesFormSet()
      
 
      

    return render_to_response('pce/module_8.html', {'context':'Edit','M8_17': M8_17.objects.filter(),'form': form,'form3': form3,'form18': form18,'form30': form30,'form45': form45,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':8,},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module8_edit(request, country, id=None,sessionid=None, template_name='pce/module_8.html'):
    """ Edit module_8 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)

    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,8)
    is_pa_filled = is_problemanalysis_filled(sessionid,8)
    is_sa_filled = is_swotanalysis_filled(sessionid,8)
    is_lf_filled = is_logicalframework_filled(sessionid,8)         
    
    if id:
        module8 = get_object_or_404(Module8, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,8):
            can_edit=1
            m_percentage=get_percentage_module_filled(8,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=8).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=8).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=8).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=8).id
				
    else:
        module8 = Module8()
  
    if request.POST:
        form = Module8Form(request.POST,  request.FILES, instance=module8)
        form3 = Module8Grid3FormSet(request.POST,  request.FILES, instance=module8)
        form18 = Module8Grid18FormSet(request.POST,  request.FILES, instance=module8)
        form30 = Module8Matrix30FormSet(request.POST,  request.FILES, instance=module8)
        form45 = Module8WeaknessesFormSet(request.POST,  request.FILES, instance=module8)
        #print('form.is_valid()'+str(form.is_valid()))
        #print('form3.is_valid()'+str(form3.is_valid()))
        #print('form18.is_valid()'+str(form18.is_valid()))
        #print('form30.is_valid()'+str(form30.is_valid()))
        #print('form45.is_valid()'+str(form45.is_valid()))
        
   
       
        if form.is_valid()  and form3.is_valid() and  form18.is_valid() and  form30.is_valid() and form45.is_valid():
            form.save()
          
            form3.instance = module8
            form3.save()
            form18.instance = module8
            form18.save()
            form30.instance = module8
            form30.save()
            form45.instance = module8
            form45.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 8."))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module8Form(instance=module8)
        form3 = Module8Grid3FormSet(instance=module8)
        form18 = Module8Grid18FormSet(instance=module8)
        form30 = Module8Matrix30FormSet(instance=module8)
        form45 = Module8WeaknessesFormSet(instance=module8)
     
    return render_to_response(template_name, {
        'context':'Edit','id':id,'M8_17': M8_17.objects.filter(),'form': form, 'form3': form3,'form18': form18,'form30': form30,'form45': form45,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':8,
    }, context_instance=RequestContext(request))

#MODULE 9
class Module9ListView(ListView):
    context_object_name = 'latest'
    model = Module9
    date_field = 'publish_date'
    template_name = 'pce/module_9.html'
    queryset = Module9.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module9 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module9.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module9ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id

        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'9'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'9') and session.status==1:
            can_edit=1
        if  can_see or can_edit:    
    
            module= 9
            module9=None
            form   = Module9FormView()
            form1 = Module9Grid1FormSet()
            form5 = Module9Grid5FormSet()
            form31 = Module9Grid31FormSet()
            form35 = Module9Matrix35FormSet()
            form47 = Module9WeaknessesFormSet()
      
    
            if id!='':
                module9 = get_object_or_404(Module9,  id=self.kwargs['id'])
                form   = Module9FormView(instance=module9)
                form1 = Module9Grid1FormSet(instance=module9)
                form5 = Module9Grid5FormSet(instance=module9)
                form31 = Module9Grid31FormSet(instance=module9)
                form35 = Module9Matrix35FormSet(instance=module9)
                form47 = Module9WeaknessesFormSet(instance=module9)
     
            context['form']=form
            context['form1']=form1
            context['form5']=form5
            context['form31']=form31
            context['form35']=form35   
            context['form47']=form47
            context['module9']=module9
         
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(9,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],9)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],9)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],9)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],9)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=9).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=9).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=9).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=9).id
      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context

class Module9ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module9
    date_field = 'publish_date'
    template_name = 'pce/module_9.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module9ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
        
        PERC2=(
        (0, _("--- Please select ---")),
        (1,_("None")),
        (2, _("20-30%")),
        (3,_("31-50%")),
        (4, _("51-75%")),
        (5, _(">75%")),
        )
        WEAK=(
        (0, _("--- Please select ---")),
        (1, _("Very weak")),
        (2, _("Weak")),
        (3, _("Intermediate")),
        (4, _("Good")),
        (5, _("Very good")),
        )
        WEAK5=(
        (0, _("--- Please select ---")),
        (1,_("Very weak")),
        (2,_("0.25")),
        (3,_("0.5")),
        (4,_("0.75")),
        (5,_("Very strong")),
        )
        TRAIN=(
        (0, _("--- Please select ---")),
        (1,_("No regular training")),
        (2, _("Once every 5 years")),
        (3,_("Once every 3 years")),
        (4, _("Once every 2 years")),
        (5, _("At least once per year")),
        )
        SUFF1=(
        (0, _("--- Please select ---")),
        (1,_("Not at all")),
        (2,_("Marginally sufficient")),
        (3,_("Somewhat sufficient")),
        (4,_("Sufficient")),
        (5,_("More than sufficient")),
        )
        NOTATALL2=(
        (0, _("--- Please select ---")),
        (1,_("Not at all")),
        (2,_("0.25")),
        (3,_("0.5")),
        (4,_("0.75")),
        (5,_("Totally")),
        )

        NOTATALL3=(
        (0, _("--- Please select ---")),
        (1,_("Totally inefficient")),
        (2,_("0.25")),
        (3,_("0.5")),
        (4,_("0.75")),
        (5,_("Very efficient")),
        )
        INSUFF4=(
        (0, _("--- Please select ---")),
        (1, _("Totally insufficient")),
        (2, _("Weak")),
        (3, _("Intermediate")),
        (4, _("Strong")),
        (5, _("Completely sufficient")),

        )
        INSUFF1=(
        (0, _("--- Please select ---")),
        (1, _("Insufficient")),
        (2, _("Some basic skills")),
        (3, _("Basic skills")),
        (4, _("Adequate training in key areas")),
        (5, _("Very sufficient")),
        )

        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
     
        
    
        context['PERC2'] =PERC2
        context['WEAK5'] =WEAK5
        context['WEAK'] =WEAK
        context['TRAIN'] =TRAIN 
        context['SUFF1'] =SUFF1
        
        context['NOTATALL2'] =NOTATALL2
        context['NOTATALL3'] =NOTATALL3
        context['INSUFF4'] =INSUFF4
        context['INSUFF1'] =INSUFF1

        
        id=''
        context['latest']=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module9.objects.filter(id=id)
        context['id'] = id
    
        can_see=0
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'9'):
            can_see=1
        if  can_see:  
            module= 9
            module9=None
            form   = Module9FormView()
            form1 = Module9Grid1FormSet()
            form5 = Module9Grid5FormSet()
            form31 = Module9Grid31FormSet()
            form35 = Module9Matrix35FormSet()
            form47 = Module9WeaknessesFormSet()
      
    
            if id!='':
                module9 = get_object_or_404(Module9,  id=self.kwargs['id'])
                form   = Module9FormView(instance=module9)
                form1 = Module9Grid1FormSet(instance=module9)
                form5 = Module9Grid5FormSet(instance=module9)
                form31 = Module9Grid31FormSet(instance=module9)
                form35 = Module9Matrix35FormSet(instance=module9)
                form47 = Module9WeaknessesFormSet(instance=module9)
     
            context['form']=form
            context['form1']=form1
            context['form5']=form5
            context['form31']=form31
            context['form35']=form35   
            context['form47']=form47
            context['module9']=module9
         
            context['module']=module
         
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(9,self.kwargs['sessionid'])

        context['can_see'] = can_see
        
        return context
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module9_create(request, country,sessionid=None):
    """ Create module9  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module9_count= Module9.objects.filter(session=sessionid).order_by('-modify_date').count()
   
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,9)
    is_pa_filled = is_problemanalysis_filled(sessionid,9)
    is_sa_filled = is_swotanalysis_filled(sessionid,9)
    is_lf_filled = is_logicalframework_filled(sessionid,9)

    
    can_edit=0
    if module9_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,9):
                can_edit=1
                m_percentage=get_percentage_module_filled(9,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=9).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=9).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=9).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=9).id
    if request.method == "POST":
        form = Module9Form(request.POST, request.FILES)
        form1 = Module9Grid1FormSet(request.POST)
        form5 = Module9Grid5FormSet(request.POST)
        form31 = Module9Grid31FormSet(request.POST)
        form35 = Module9Matrix35FormSet(request.POST)
        form47 = Module9WeaknessesFormSet(request.POST)
      
        if form.is_valid() and form1.is_valid() and  form5.is_valid() and  form31.is_valid() and form35.is_valid() and form47.is_valid():
            new_mod9 = form.save(commit=False)
            new_mod9.author = request.user
            new_mod9.session=pceversion
            form.save()
            
            form1.instance = new_mod9
            form1.save()
            form5.instance = new_mod9
            form5.save()
            form31.instance = new_mod9
            form31.save()
            form35.instance = new_mod9
            form35.save()
            form47.instance = new_mod9
            form47.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
           
            info(request, _("Successfully saved Module 9."))

            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_9.html', {'context':'Edit','form': form,'form1': form1,'form5': form5,'form31': form31,'form35': form35,'form47': form47,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':9,},
             context_instance=RequestContext(request))

    else:
        form = Module9Form(initial={'country': country,'session': pceversion.id})
        form1 = Module9Grid1FormSet()
        form5 = Module9Grid5FormSet()
        form31 = Module9Grid31FormSet()
        form35 = Module9Matrix35FormSet()
        form47 = Module9WeaknessesFormSet()
      
 
      

    return render_to_response('pce/module_9.html', {'context':'Edit','form': form,'form1': form1,'form5': form5,'form31': form31,'form35': form35,'form47': form47,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':9,},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module9_edit(request, country, id=None,sessionid=None, template_name='pce/module_9.html'):
    """ Edit module_9 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,9)
    is_pa_filled = is_problemanalysis_filled(sessionid,9)
    is_sa_filled = is_swotanalysis_filled(sessionid,9)
    is_lf_filled = is_logicalframework_filled(sessionid,9)
    
    if id:
        module9 = get_object_or_404(Module9, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,9):
            can_edit=1
            m_percentage=get_percentage_module_filled(9,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=9).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=9).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=9).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=9).id
    else:
        module9 = Module9()
  
    if request.POST:
        form = Module9Form(request.POST,  request.FILES, instance=module9)
        form1 = Module9Grid1FormSet(request.POST,  request.FILES, instance=module9)
        form5 = Module9Grid5FormSet(request.POST,  request.FILES, instance=module9)
        form31 = Module9Grid31FormSet(request.POST,  request.FILES, instance=module9)
        form35 = Module9Matrix35FormSet(request.POST,  request.FILES, instance=module9)
        form47 = Module9WeaknessesFormSet(request.POST,  request.FILES, instance=module9)
      
       
        if form.is_valid()  and form1.is_valid() and  form5.is_valid() and  form31.is_valid() and form35.is_valid() and form47.is_valid():
            form.save()
          
            form1.instance = module9
            form1.save()
            form5.instance = module9
            form5.save()
            form31.instance = module9
            form31.save()
            form35.instance = module9
            form35.save()
            form47.instance = module9
            form47.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 9."))
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module9Form(instance=module9)
        form1 = Module9Grid1FormSet(instance=module9)
        form5 = Module9Grid5FormSet(instance=module9)
        form31 = Module9Grid31FormSet(instance=module9)
        form35 = Module9Matrix35FormSet(instance=module9)
        form47 = Module9WeaknessesFormSet( instance=module9)
      
    return render_to_response(template_name, {
        'context':'Edit','form': form, 'form1': form1,'form5': form5,'form31': form31,'form35': form35,'form47': form47,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':9,
    }, context_instance=RequestContext(request))

#MODULE 10
class Module10ListView(ListView):
    context_object_name = 'latest'
    model = Module10
    date_field = 'publish_date'
    template_name = 'pce/module_10.html'
    queryset = Module10.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module10 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module10.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/11010110220
        context = super(Module10ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id

        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'10'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'10') and session.status==1:
            can_edit=1
        if  can_see or can_edit: 
            module= 10
            module10=None
            form   = Module10FormView()
            form23 = Module10Grid23FormSet()
            form31 = Module10Grid31FormSet()
            form33 = Module10Grid33FormSet()
            form37 = Module10Grid37FormSet()
            form45 = Module10Grid45FormSet()
            form46 = Module10Grid46FormSet()
            form47 = Module10Matrix47FormSet()
            form61 = Module10WeaknessesFormSet()

            if id!='':
                module10 = get_object_or_404(Module10,  id=self.kwargs['id'])
                form   = Module10FormView(instance=module10)
                form23 = Module10Grid23FormSet(instance=module10)
                form31 = Module10Grid31FormSet(instance=module10)
                form33 = Module10Grid33FormSet(instance=module10)
                form37 = Module10Grid37FormSet(instance=module10)
                form45 = Module10Grid45FormSet(instance=module10)
                form46 = Module10Grid46FormSet(instance=module10)
                form47 = Module10Matrix47FormSet(instance=module10)
                form61= Module10WeaknessesFormSet(instance=module10)
            context['module10']=module10
         
            context['form']=form
            context['form23']=form23
            context['form31']=form31
            context['form33']=form33
            context['form37']=form37
            context['form45']=form45
            context['form46']=form46
            context['form47']=form47
            context['form61']=form61
            
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(10,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],10)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],10)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],10)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],10)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=10).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=10).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=10).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=10).id
      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
class Module10ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module10
    date_field = 'publish_date'
    template_name = 'pce/module_10.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module10ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M



        INSUFF5=(
        (0, _("--- Please select ---")),
        (1, _("Totally insufficient")),
        (2, _("Insufficient")),
        (3, _("Not enough")),
        (4, _("Enough")),
        (5, _("More than enough")),
        )
        INSUFF6=(
        (0, _("--- Please select ---")),
        (1, _("Totally insufficient")),
        (2, _("insufficient")),
        (3, _("Not enough")),
        (4, _("Not so bad")),
        (5, _("Totally")),
        )

        BAD = (
        (0, _("--- Please select ---")),
        (1, _("Very bad")),
        (2, _("Bad")),
        (3, _("Not so bad")),
        (4, _("Good")),
        (5, _("Very good")),
        )
        TRAIN2=(
        (0, _("--- Please select ---")),
        (1,_("No programmed training")),
        (2, _("Once in 5 years")),
        (3,_("Once in 3 years")),
        (4, _("Once in 2 years")),
        (5, _("Once per year")),
        )
        EFF=(
        (0, _("--- Please select ---")),
        (1,_("Very ineffective")),
        (2,_("Ineffective")),
        (3,_("Intermediate")),
        (4,_("Effective")),
        (5,_("Very effective")),
        )
        CAPACITY = (
        (0, _("--- Please select ---")),
        (1, _("very low")),
        (2, _("low")),
        (3, _("intermediate")),
        (4, _("high")),
        (5, _("very high")),
        )

    
 
        context['INSUFF5'] =INSUFF5 
        context['INSUFF6'] =INSUFF6
        context['EFF'] =EFF
        
        context['BAD'] =BAD
        context['TRAIN2'] =TRAIN2
        context['EFF'] =EFF
        context['CAPACITY'] =CAPACITY

        
        id=''
        context['latest']=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module10.objects.filter(id=id)
        context['id'] = id
    
        can_see=0
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'10'):
            can_see=1
        if  can_see:  
            module= 10
            module10=None
            form   = Module10FormView()
            form23 = Module10Grid23FormSet()
            form31 = Module10Grid31FormSet()
            form33 = Module10Grid33FormSet()
            form37 = Module10Grid37FormSet()
            form45 = Module10Grid45FormSet()
            form46 = Module10Grid46FormSet()
            form47 = Module10Matrix47FormSet()
            form61 = Module10WeaknessesFormSet()

            if id!='':
                module10 = get_object_or_404(Module10,  id=self.kwargs['id'])
                form   = Module10FormView(instance=module10)
                form23 = Module10Grid23FormSet(instance=module10)
                form31 = Module10Grid31FormSet(instance=module10)
                form33 = Module10Grid33FormSet(instance=module10)
                form37 = Module10Grid37FormSet(instance=module10)
                form45 = Module10Grid45FormSet(instance=module10)
                form46 = Module10Grid46FormSet(instance=module10)
                form47 = Module10Matrix47FormSet(instance=module10)
                form61= Module10WeaknessesFormSet(instance=module10)
            context['module10']=module10
         
            context['form']=form
            context['form23']=form23
            context['form31']=form31
            context['form33']=form33
            context['form37']=form37
            context['form45']=form45
            context['form46']=form46
            context['form47']=form47
            context['form61']=form61
            
            context['module']=module
         
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(10,self.kwargs['sessionid'])

        context['can_see'] = can_see
        
        return context
    
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module10_create(request, country,sessionid=None):
    """ Create module10  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module10_count= Module10.objects.filter(session=sessionid).order_by('-modify_date').count()
   
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,10)
    is_pa_filled = is_problemanalysis_filled(sessionid,10)
    is_sa_filled = is_swotanalysis_filled(sessionid,10)
    is_lf_filled = is_logicalframework_filled(sessionid,10)

  
    if module10_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,10):
                can_edit=1
                m_percentage=get_percentage_module_filled(10,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=10).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=10).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=10).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=10).id

    if request.method == "POST":
        form = Module10Form(request.POST, request.FILES)
        form23 = Module10Grid23FormSet(request.POST)
        form31 = Module10Grid31FormSet(request.POST)
        form33 = Module10Grid33FormSet(request.POST)
        form37 = Module10Grid37FormSet(request.POST)
        form45 = Module10Grid45FormSet(request.POST)
        form46 = Module10Grid46FormSet(request.POST)
        form47 = Module10Matrix47FormSet(request.POST)
        form61= Module10WeaknessesFormSet(request.POST)
      
        if form.is_valid() and  form23.is_valid() and  form31.is_valid() and form33.is_valid() and form37.is_valid() and form45.is_valid() and form46.is_valid() and form47.is_valid() and form61.is_valid():
            new_mod10 = form.save(commit=False)
            new_mod10.author = request.user
            new_mod10.session=pceversion
            form.save()
            
            form23.instance = new_mod10
            form23.save()
            form31.instance = new_mod10
            form31.save()
            form33.instance = new_mod10
            form33.save()
            form37.instance = new_mod10
            form37.save()
            form45.instance = new_mod10
            form45.save()
            form46.instance = new_mod10
            form46.save()
            form47.instance = new_mod10
            form47.save()
            form61.instance = new_mod10
            form61.save()
           
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 10."))

            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_10.html', {'context':'Edit','form': form,'form23': form23,'form31': form31,'form33': form33,'form37': form37,'form45': form45,'form46': form46,'form47': form47,'form61': form61,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':10,},
             context_instance=RequestContext(request))

    else:
        form = Module10Form(initial={'country': country,'session': pceversion.id})
        form23 = Module10Grid23FormSet()
        form31 = Module10Grid31FormSet()
        form33 = Module10Grid33FormSet()
        form37 = Module10Grid37FormSet()
        form45 = Module10Grid45FormSet()
        form46 = Module10Grid46FormSet()
        form47 = Module10Matrix47FormSet()
        form61= Module10WeaknessesFormSet()
      
 
      

    return render_to_response('pce/module_10.html', {'context':'Edit','form': form,'form23': form23,'form31': form31,'form33': form33,'form37': form37,'form45': form45,'form46': form46,'form47': form47,'form61': form61,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':10,},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module10_edit(request, country, id=None,sessionid=None, template_name='pce/module_10.html'):
    """ Edit module_10 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,10)
    is_pa_filled = is_problemanalysis_filled(sessionid,10)
    is_sa_filled = is_swotanalysis_filled(sessionid,10)
    is_lf_filled = is_logicalframework_filled(sessionid,10)
              
    
    if id:
        module10 = get_object_or_404(Module10, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,10):
            can_edit=1
            m_percentage=get_percentage_module_filled(10,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=10).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=10).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=10).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=10).id
    else:
        module10 = Module10()
  
    if request.POST:
        form = Module10Form(request.POST,  request.FILES, instance=module10)
        form23 = Module10Grid23FormSet(request.POST,  request.FILES, instance=module10)
        form31 = Module10Grid31FormSet(request.POST,  request.FILES, instance=module10)
        form33 = Module10Grid33FormSet(request.POST,  request.FILES, instance=module10)
        form37 = Module10Grid37FormSet(request.POST,  request.FILES, instance=module10)
        form45 = Module10Grid45FormSet(request.POST,  request.FILES, instance=module10)
        form46 = Module10Grid46FormSet(request.POST,  request.FILES, instance=module10)
        form47 = Module10Matrix47FormSet(request.POST,  request.FILES, instance=module10)
        form61 = Module10WeaknessesFormSet(request.POST,  request.FILES, instance=module10)
      
   
        if form.is_valid() and  form23.is_valid() and  form31.is_valid() and form33.is_valid() and form37.is_valid() and form45.is_valid() and form46.is_valid()  and form47.is_valid() and form61.is_valid():
            form.save()
          
        
            form23.instance = module10
            form23.save()
            form31.instance = module10
            form31.save()
            form33.instance = module10
            form33.save()
            form37.instance = module10
            form37.save()
            form45.instance = module10
            form45.save()
            form46.instance = module10
            form46.save()
            form47.instance = module10
            form47.save()
            form61.instance = module10
            form61.save()
            
            
        
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 10."))
            
            
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module10Form(instance=module10)
        form23 = Module10Grid23FormSet(instance=module10)
        form31 = Module10Grid31FormSet(instance=module10)
        form33 = Module10Grid33FormSet(instance=module10)
        form37 = Module10Grid37FormSet(instance=module10)
        form45 = Module10Grid45FormSet(instance=module10)
        form46 = Module10Grid46FormSet(instance=module10)
        form47 = Module10Matrix47FormSet(instance=module10)
        form61 = Module10WeaknessesFormSet( instance=module10)
      
    return render_to_response(template_name, {
        'context':'Edit','form': form,'form23': form23,'form31': form31,'form33': form33,'form37': form37,'form45': form45,'form46': form46,'form47': form47,'form61': form61,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':10,
    }, context_instance=RequestContext(request))
#MODULE 11
class Module11ListView(ListView):
    context_object_name = 'latest'
    model = Module11
    date_field = 'publish_date'
    template_name = 'pce/module_11.html'
    queryset = Module11.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module11 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module11.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module11ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
        id=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id

        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'11'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'11') and session.status==1:
            can_edit=1
        if  can_see or can_edit: 
    
            module= 11
            module11=None
            form   = Module11FormView()
            form2 = Module11Grid2FormSet()
            form3 = Module11Grid3FormSet()
            form12 = Module11Grid12FormSet()
            form14 = Module11Grid14FormSet()
            form33 = Module11Grid33FormSet()
            form42 = Module11Matrix42FormSet()
            form66= Module11WeaknessesFormSet()
      
            if id!='':
                module11 = get_object_or_404(Module11,  id=self.kwargs['id'])
                form   = Module11FormView(instance=module11)
                form2 = Module11Grid2FormSet(instance=module11)
                form3 = Module11Grid3FormSet(instance=module11)
                form12 = Module11Grid12FormSet(instance=module11)
                form14 = Module11Grid14FormSet(instance=module11)
                form33 = Module11Grid33FormSet(instance=module11)
                form42 = Module11Matrix42FormSet(instance=module11)
                form66= Module11WeaknessesFormSet(instance=module11)
     
            context['form']=form
            context['form2']=form2
            context['form3']=form3
            context['form12']=form12
            context['form14']=form14
            context['form33']=form33
            context['form42']=form42
            context['form66']=form66
            context['module11']=module11
            
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(11,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],11)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],11)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],11)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],11)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=11).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=11).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=11).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=11).id
      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
class Module11ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module11
    date_field = 'publish_date'
    template_name = 'pce/module_11.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module11ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
        context['VAL_AV'] =VAL_AV
        
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
        RATHER = (
            (0, _("--- Please select ---")),
            (1, _("Not at all")),
            (2, _("Rather low")),
            (3, _("Intermediate")),
            (4, _("Much")),
            (5, _("Very much so")),
        ) 

        BAD2 = (
        (0, _("--- Please select ---")),
        (1, _("Very badly")),
        (2, _("Badly")),
        (3, _("Intermediate")),
        (4, _("Well")),
        (5, _("Very well")),
        )
        EFF=(
        (0, _("--- Please select ---")),
        (1,_("Very ineffective")),
        (2,_("Ineffective")),
        (3,_("Intermediate")),
        (4,_("Effective")),
        (5,_("Very effective")),
        )
        INSUFF7=(
        (0, _("--- Please select ---")),
        (1, _("Totally insufficient")),
        (2, _("insufficient")),
        (3, _("Huge limitations")),
        (4, _("Minor limitations")),
        (5, _("Without limitations")),
        )
        WEAK=(
        (0, _("--- Please select ---")),
        (1, _("Very weak")),
        (2, _("Weak")),
        (3, _("Intermediate")),
        (4, _("Good")),
        (5, _("Very good")),
        )
        OUT=(
        (0, _("--- Please select ---")),
        (1,_("all outsourced")),
        (2,_("Partially outsourced")),
        (3,_("All done in house")),
        )
        LIM=(
        (0, _("--- Please select ---")),
        (1,_("Not at all")),
        (2,_("Severe limitations")),
        (3,_("Limited")),
        (4,_("Few limitations")),
        (5,_("Completely sufficient")),
        )
        PERC1=(
        (0, _("--- Please select ---")),
        (1,_("0")),
        (2, _("1-25")),
        (3,_("26-50")),
        (4, _("51-75")),
        (5, _(">75")),
        )
        THEM = (
                (0, _("--- Please select ---")),
                (1,_("None at all")),
                (2,_("A few of them")),
                (3,_("Some of them")),
                (4,_("Most of them")),
                (5,_("All of them")),
        )
        TRAIN2 =(
        (0, _("--- Please select ---")),
        (1,_("No programmed training")),
        (2, _("Once in 5 years")),
        (3,_("Once in 3 years")),
        (4, _("Once in 2 years")),
        (5, _("Once per year")),
        )

 
                   
        context['RATHER'] =RATHER
        context['BAD2'] =BAD2
        context['EFF'] =EFF
        context['INSUFF7'] =INSUFF7
        context['WEAK'] =WEAK
        context['OUT'] =OUT
        context['LIM'] =LIM
        context['PERC1'] =PERC1
        context['TRAIN2'] =TRAIN2
        context['THEM'] =THEM
       
        
        id=''
        context['latest']=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module11.objects.filter(id=id)
        context['id'] = id
    
        can_see=0
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'11'):
            can_see=1
        if  can_see:  
            module= 11
            module11=None
            form   = Module11FormView()
            form2 = Module11Grid2FormSet()
            form3 = Module11Grid3FormSet()
            form12 = Module11Grid12FormSet()
            form14 = Module11Grid14FormSet()
            form33 = Module11Grid33FormSet()
            form42 = Module11Matrix42FormSet()
            form66= Module11WeaknessesFormSet()
      
            if id!='':
                module11 = get_object_or_404(Module11,  id=self.kwargs['id'])
                form   = Module11FormView(instance=module11)
                form2 = Module11Grid2FormSet(instance=module11)
                form3 = Module11Grid3FormSet(instance=module11)
                form12 = Module11Grid12FormSet(instance=module11)
                form14 = Module11Grid14FormSet(instance=module11)
                form33 = Module11Grid33FormSet(instance=module11)
                form42 = Module11Matrix42FormSet(instance=module11)
                form66= Module11WeaknessesFormSet(instance=module11)
     
            context['form']=form
            context['form2']=form2
            context['form3']=form3
            context['form12']=form12
            context['form14']=form14
            context['form33']=form33
            context['form42']=form42
            context['form66']=form66
            context['module11']=module11
            
            context['module']=module
         
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(11,self.kwargs['sessionid'])

        context['can_see'] = can_see
        
        return context
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module11_create(request, country,sessionid=None):
    """ Create module11  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module11_count= Module11.objects.filter(session=sessionid).order_by('-modify_date').count()
   
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,11)
    is_pa_filled = is_problemanalysis_filled(sessionid,11)
    is_sa_filled = is_swotanalysis_filled(sessionid,11)
    is_lf_filled = is_logicalframework_filled(sessionid,11)

  
    if module11_count>0:
        can_edit=0
    else:
        if is_st_filled:
             if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,11):
                can_edit=1
                m_percentage=get_percentage_module_filled(11,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=11).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=11).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=11).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=11).id
    if request.method == "POST":
        form = Module11Form(request.POST, request.FILES)
        form2 = Module11Grid2FormSet(request.POST)
        form3 = Module11Grid3FormSet(request.POST)
        form12 = Module11Grid12FormSet(request.POST)
        form14 = Module11Grid14FormSet(request.POST)
        form33 = Module11Grid33FormSet(request.POST)
        form42 = Module11Matrix42FormSet(request.POST)
        form66= Module11WeaknessesFormSet(request.POST)
      
        if form.is_valid() and form2.is_valid() and  form3.is_valid() and  form12.is_valid() and form14.is_valid() and form33.is_valid()and form42.is_valid() and form66.is_valid():
            new_mod11 = form.save(commit=False)
            new_mod11.author = request.user
            new_mod11.session=pceversion
            form.save()
            
            form2.instance = new_mod11
            form2.save()
            form3.instance = new_mod11
            form3.save()
            form12.instance = new_mod11
            form12.save()
            form14.instance = new_mod11
            form14.save()
            form33.instance = new_mod11
            form33.save()
            form42.instance = new_mod11
            form42.save()
            form66.instance = new_mod11
            form66.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
           
            info(request, _("Successfully saved Module 11."))

            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_11.html', {'context':'Edit','form': form,'form2': form2,'form3': form3,'form12': form12,'form14': form14,'form33': form33,'form42': form42,'form66': form66,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':11,},
             context_instance=RequestContext(request))

    else:
        form = Module11Form(initial={'country': country,'session': pceversion.id})
        form2 = Module11Grid2FormSet()
        form3 = Module11Grid3FormSet()
        form12 = Module11Grid12FormSet()
        form14 = Module11Grid14FormSet()
        form33 = Module11Grid33FormSet()
        form42 = Module11Matrix42FormSet()
        form66= Module11WeaknessesFormSet()
      
      

    return render_to_response('pce/module_11.html', {'context':'Edit','form': form,'form2': form2,'form3': form3,'form12': form12,'form14': form14,'form33': form33,'form42': form42,'form66': form66,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':11,},
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module11_edit(request, country, id=None,sessionid=None, template_name='pce/module_11.html'):
    """ Edit module_11 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)

    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,11)
    is_pa_filled = is_problemanalysis_filled(sessionid,11)
    is_sa_filled = is_swotanalysis_filled(sessionid,11)
    is_lf_filled = is_logicalframework_filled(sessionid,11)         
    
    if id:
        module11 = get_object_or_404(Module11, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,11):
            can_edit=1
            m_percentage=get_percentage_module_filled(11,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=11).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=11).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=11).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=11).id
    else:
        module11 = Module11()
  
    if request.POST:
        form = Module11Form(request.POST,  request.FILES, instance=module11)
        form2 = Module11Grid2FormSet(request.POST,  request.FILES, instance=module11)
        form3 = Module11Grid3FormSet(request.POST,  request.FILES, instance=module11)
        form12 = Module11Grid12FormSet(request.POST,  request.FILES, instance=module11)
        form14 = Module11Grid14FormSet(request.POST,  request.FILES, instance=module11)
        form33 = Module11Grid33FormSet(request.POST,  request.FILES, instance=module11)
        form42 = Module11Matrix42FormSet(request.POST,  request.FILES, instance=module11)
        form66= Module11WeaknessesFormSet(request.POST,  request.FILES, instance=module11)
        #print('form'+str(form.is_valid()))
        #print('form2'+str(form2.is_valid()))
        #print('form3'+str(form3.is_valid()))
        #print('form12'+str(form12.is_valid()))
        #print('form14'+str(form14.is_valid()))
        #print('form33'+str(form33.is_valid()))
        #print('form42'+str(form42.is_valid()))
        #print('form66'+str(form66.is_valid()))    
        if form.is_valid() and form2.is_valid() and  form3.is_valid() and  form12.is_valid() and form14.is_valid() and form33.is_valid() and form42.is_valid() and  form66.is_valid():
            form.save()
          
        
            form2.instance = module11
            form2.save()
            form3.instance = module11
            form3.save()
            form12.instance = module11
            form12.save()
            form14.instance = module11
            form14.save()
            form33.instance = module11
            form33.save()
            form42.instance = module11
            form42.save()
            form66.instance = module11
            form66.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
                
            info(request, _("Successfully saved Module 11"))
            
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module11Form(instance=module11)
        form2 = Module11Grid2FormSet(instance=module11)
        form3 = Module11Grid3FormSet(instance=module11)
        form12 = Module11Grid12FormSet(instance=module11)
        form14 = Module11Grid14FormSet(instance=module11)
        form33 = Module11Grid33FormSet(instance=module11)
        form42 = Module11Matrix42FormSet(instance=module11)
        form66= Module11WeaknessesFormSet(instance=module11)
      
    return render_to_response(template_name, {
       'context':'Edit','form': form,'form2': form2,'form3': form3,'form12': form12,'form14': form14,'form33': form33,'form42': form42,'form66': form66,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':11,
    }, context_instance=RequestContext(request))
#MODULE 12
class Module12ListView(ListView):
    context_object_name = 'latest'
    model = Module12
    date_field = 'publish_date'
    template_name = 'pce/module_12.html'
    queryset = Module12.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module12 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module12.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module12ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        context['vals'] =INSUFF0
        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
        id=''
      
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id

        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'12'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'12') and session.status==1:
            can_edit=1
        if  can_see or can_edit: 
            module= 12
            module12=None
            form   = Module12FormView()
            form2 = Module12Grid2FormSet()
            form3 = Module12Grid3FormSet()
            form29 = Module12Grid29FormSet()
            form22 = Module12Matrix22FormSet()
            form34= Module12WeaknessesFormSet()
            if id!='':
                module12 = get_object_or_404(Module12,  id=self.kwargs['id'])
                form   = Module12FormView(instance=module12)
                form2 = Module12Grid2FormSet(instance=module12)
                form3 = Module12Grid3FormSet(instance=module12)
                form29 = Module12Grid29FormSet(instance=module12)
                form22 = Module12Matrix22FormSet(instance=module12)
                form34= Module12WeaknessesFormSet(instance=module12)
            context['module12']=module12
            context['form']=form
            context['form2']=form2
            context['form3']=form3
            context['form29']=form29
            context['form22']=form22
            context['form34']=form34
            
            
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(12,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],12)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],12)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],12)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],12)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=12).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=12).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=12).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=12).id
      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
class Module12ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module12
    date_field = 'publish_date'
    template_name = 'pce/module_12.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module12ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
        context['vals'] =INSUFF0
        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
        RATHER = (
            (0, _("--- Please select ---")),
            (1, _("Not at all")),
            (2, _("Rather low")),
            (3, _("Intermediate")),
            (4, _("Much")),
            (5, _("Very much so")),
        ) 

        BAD2 = (
        (0, _("--- Please select ---")),
        (1, _("Very badly")),
        (2, _("Badly")),
        (3, _("Intermediate")),
        (4, _("Well")),
        (5, _("Very well")),
        )
        EFF=(
        (0, _("--- Please select ---")),
        (1,_("Very ineffective")),
        (2,_("Ineffective")),
        (3,_("Intermediate")),
        (4,_("Effective")),
        (5,_("Very effective")),
        )
     
        WEAK=(
        (0, _("--- Please select ---")),
        (1, _("Very weak")),
        (2, _("Weak")),
        (3, _("Intermediate")),
        (4, _("Good")),
        (5, _("Very good")),
        )
       
    
        THEM = (
                (0, _("--- Please select ---")),
                (1,_("None at all")),
                (2,_("A few of them")),
                (3,_("Some of them")),
                (4,_("Most of them")),
                (5,_("All of them")),
        )
        TRAIN2 =(
        (0, _("--- Please select ---")),
        (1,_("No programmed training")),
        (2, _("Once in 5 years")),
        (3,_("Once in 3 years")),
        (4, _("Once in 2 years")),
        (5, _("Once per year")),
        )
        INSUFF8=(
        (0, _("--- Please select ---")),
        (1, _("Totally insufficient")),
        (2, _("Insufficient")),
        (3, _("Intermediate")),
        (4, _("Not enough")),
        (5, _("Very sufficient")),
        )
 
        PERC3=(
        (0, _("--- Please select ---")),
        (1,_("None")),
        (2, _("5%")),
        (3,_("20%")),
        (4, _("40%")),
        (5, _(">40%")),
        )          
        context['RATHER'] =RATHER
        context['BAD2'] =BAD2
        context['EFF'] =EFF
        context['INSUFF8'] =INSUFF8
        context['WEAK'] =WEAK
        context['PERC3'] =PERC3
        context['TRAIN2'] =TRAIN2
        context['THEM'] =THEM
       
        

        
        
        id=''
        context['latest']=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module12.objects.filter(id=id)
        context['id'] = id
    
        can_see=0
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'12'):
            can_see=1
        if  can_see:  
            module= 12
            module12=None
            form   = Module12FormView()
            form2 = Module12Grid2FormSet()
            form3 = Module12Grid3FormSet()
            form29 = Module12Grid29FormSet()
            form22 = Module12Matrix22FormSet()
            form34= Module12WeaknessesFormSet()
            if id!='':
                module12 = get_object_or_404(Module12,  id=self.kwargs['id'])
                form   = Module12FormView(instance=module12)
                form2 = Module12Grid2FormSet(instance=module12)
                form3 = Module12Grid3FormSet(instance=module12)
                form29 = Module12Grid29FormSet(instance=module12)
                form22 = Module12Matrix22FormSet(instance=module12)
                form34= Module12WeaknessesFormSet(instance=module12)
            context['module12']=module12
            context['form']=form
            context['form2']=form2
            context['form3']=form3
            context['form29']=form29
            context['form22']=form22
            context['form34']=form34
            
            
            context['module']=module
         
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(12,self.kwargs['sessionid'])

        context['can_see'] = can_see
        
        return context
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module12_create(request, country,sessionid=None):
    """ Create module12  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module12_count= Module12.objects.filter(session=sessionid).order_by('-modify_date').count()
   
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,12)
    is_pa_filled = is_problemanalysis_filled(sessionid,12)
    is_sa_filled = is_swotanalysis_filled(sessionid,12)
    is_lf_filled = is_logicalframework_filled(sessionid,12)

  
    if module12_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,12):
                can_edit=1
                m_percentage=get_percentage_module_filled(12,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=12).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=12).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=12).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=12).id
			
    if request.method == "POST":
        form = Module12Form(request.POST, request.FILES)
        form2 = Module12Grid2FormSet(request.POST)
        form3 = Module12Grid3FormSet(request.POST)
        form29 = Module12Grid29FormSet(request.POST)
        form22 = Module12Matrix22FormSet(request.POST)
        form34= Module12WeaknessesFormSet(request.POST)
      
        if form.is_valid() and form2.is_valid() and  form3.is_valid() and  form29.is_valid() and form22.is_valid() and form34.is_valid():
            new_mod12 = form.save(commit=False)
            new_mod12.author = request.user
            new_mod12.session=pceversion
            form.save()
            
            form2.instance = new_mod12
            form2.save()
            form3.instance = new_mod12
            form3.save()
            form22.instance = new_mod12
            form22.save()
            form29.instance = new_mod12
            form29.save()
            form34.instance = new_mod12
            form34.save()
          
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 12."))

            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_12.html', {'context':'Edit','form': form,'form2': form2,'form3': form3,'form22': form22,'form29': form29,'form34': form34,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':12,  },
             context_instance=RequestContext(request))

    else:
        form = Module12Form(initial={'country': country,'session': pceversion.id})
        form2 = Module12Grid2FormSet()
        form3 = Module12Grid3FormSet()
        form29 = Module12Grid29FormSet()
        form22 = Module12Matrix22FormSet()
        form34= Module12WeaknessesFormSet()
      

    return render_to_response('pce/module_12.html', {'context':'Edit','form': form,'form2': form2,'form3': form3,'form22': form22,'form29': form29,'form34': form34,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':12,  },
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module12_edit(request, country, id=None,sessionid=None, template_name='pce/module_12.html'):
    """ Edit module_12 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)

    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,12)
    is_pa_filled = is_problemanalysis_filled(sessionid,12)
    is_sa_filled = is_swotanalysis_filled(sessionid,12)
    is_lf_filled = is_logicalframework_filled(sessionid,12)

           
    
    if id:
        module12 = get_object_or_404(Module12, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,12):
            can_edit=1
            m_percentage=get_percentage_module_filled(12,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=12).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=12).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=12).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=12).id
			
    else:
        module12 = Module12()
  
    if request.POST:
        form = Module12Form(request.POST,  request.FILES, instance=module12)
        form2 = Module12Grid2FormSet(request.POST,  request.FILES, instance=module12)
        form3 = Module12Grid3FormSet(request.POST,  request.FILES, instance=module12)
        form29 = Module12Grid29FormSet(request.POST,  request.FILES, instance=module12)
        form22 = Module12Matrix22FormSet(request.POST,  request.FILES, instance=module12)
        form34= Module12WeaknessesFormSet(request.POST,  request.FILES, instance=module12)
     
        if form.is_valid() and form2.is_valid() and  form3.is_valid() and  form29.is_valid() and form22.is_valid() and form34.is_valid():
            form.save()
          
        
            form2.instance = module12
            form2.save()
            form3.instance = module12
            form3.save()
            form29.instance = module12
            form29.save()
            form22.instance = module12
            form22.save()
            form34.instance = module12
            form34.save()
            
            pceversion.modify_date=timezone.now()
            pceversion.save()
            
        
    
            info(request, _("Successfully saved Module 12"))
            
            
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module12Form(instance=module12)
        form2 = Module12Grid2FormSet(instance=module12)
        form3 = Module12Grid3FormSet(instance=module12)
        form29 = Module12Grid29FormSet(instance=module12)
        form22 = Module12Matrix22FormSet(instance=module12)
        form34= Module12WeaknessesFormSet(instance=module12)
     
      
    return render_to_response(template_name, {
        'context':'Edit','form': form,'form2': form2,'form3': form3,'form22': form22,'form29': form29,'form34': form34,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':12,  
    }, context_instance=RequestContext(request))

#MODULE 13
class Module13ListView(ListView):
    context_object_name = 'latest'
    model = Module13
    date_field = 'publish_date'
    template_name = 'pce/module_13.html'
    queryset = Module13.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Module13 from the specific id """
        self.id=None
        if  'id' in self.kwargs:
            self.id = self.kwargs['id']
        return Module13.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module13ListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
        id=''
      
        if  'id' in self.kwargs:
            id = self.kwargs['id']
        context['id'] = id

        can_edit=0
        can_see=0
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
       
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        module=13
        if canSee(session.id,session.country,self.request.user,'13'):
            can_see=1
        if canEdit(session.id,session.country,self.request.user,'13') and session.status==1:
            can_edit=1
        if  can_see or can_edit: 
            module13=None
            form   = Module13FormView()
            form2 = Module13Grid2FormSet()
            form3 = Module13Grid3FormSet()
            form22 = Module13Grid22FormSet()
            form29 = Module13Grid29FormSet()
            form31 = Module13Grid31FormSet()
            form47 = Module13Matrix47FormSet()
            form66= Module13WeaknessesFormSet()
            if id!='':
                module13 = get_object_or_404(Module13,  id=self.kwargs['id'])
                form   = Module13FormView(instance=module13)
                form2 = Module13Grid2FormSet(instance=module13)
                form3 = Module13Grid3FormSet(instance=module13)
                form22 = Module13Grid22FormSet(instance=module13)
                form29 = Module13Grid29FormSet(instance=module13)
                form31 = Module13Grid31FormSet(instance=module13)
                form47 = Module13Matrix47FormSet(instance=module13)
                form66= Module13WeaknessesFormSet(instance=module13)
      
            context['module13']=module13
            context['form']=form
            context['form2']=form2
            context['form3']=form3
            context['form22']=form22
            context['form29']=form29
            context['form31']=form31
            context['form47']=form47
            context['form66']=form66
            
            
            context['module']=module
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(13,self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],13)
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],13)
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],13)
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],13)
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=13).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=13).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=13).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=13).id
      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
class Module13ListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Module13
    date_field = 'publish_date'
    template_name = 'pce/module_13.html'
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(Module13ListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'   
        context['VAL_AV'] =VAL_AV
        context['BOOL_CHOICESM_M'] =BOOL_CHOICESM_M
    

        LIM1=(
          (0, _("--- Please select ---")),
          (1,_("Totally insufficient")),
          (2,_("Insufficient")),
          (3,_("With strong limitations")),
          (4,_("With some limitations")),
          (5,_("Without limitations")),
          )

        WEAK=(
        (0, _("--- Please select ---")),
        (1, _("Very weak")),
        (2, _("Weak")),
        (3, _("Intermediate")),
        (4, _("Good")),
        (5, _("Very good")),
        )
       
    
        TRAIN2 =(
        (0, _("--- Please select ---")),
        (1,_("No programmed training")),
        (2, _("Once in 5 years")),
        (3,_("Once in 3 years")),
        (4, _("Once in 2 years")),
        (5, _("Once per year")),
        )
                
       
        INSP=(
        (0, _("--- Please select ---")),
        (1,_("Pest diagnostic")),
        (2,_("Treatment")),
        (3,_("Field inspection")),
        (4,_("Packing inspection")),
        (5,_("Inspection and storage facilities")),
        )   
        FEW=(
        (0, _("--- Please select ---")),
        (1,_("None")),
        (2,_("Few")),
        (3,_("Some")),
        (4,_("Most")),
        (5,_("All")),
        )
        PHY=(
        (0, _("--- Please select ---")),
        (1,_("None")),
        (2,_("Phytosanitary certificate")),
        (3,_("Re-export certificate")),
        )      
        context['WEAK'] =WEAK
        context['TRAIN2'] =TRAIN2
        context['LIM1'] =LIM1
        context['FEW'] =FEW
        context['PHY'] =PHY
        context['INSP'] =INSP
  
        
        id=''
        context['latest']=''
        if  'id' in self.kwargs:
            id = self.kwargs['id']
            context['latest']=  Module13.objects.filter(id=id)
        context['id'] = id
        
        module=13
        can_see=0
        session = get_object_or_404(PceVersion,  pk= self.kwargs['sessionid'])
    
        if canSee(session.id,session.country,self.request.user,'13'):
            can_see=1
        if  can_see:  
            module13=None
            form   = Module13FormView()
            form2 = Module13Grid2FormSet()
            form3 = Module13Grid3FormSet()
            form22 = Module13Grid22FormSet()
            form29 = Module13Grid29FormSet()
            form31 = Module13Grid31FormSet()
            form47 = Module13Matrix47FormSet()
            form66= Module13WeaknessesFormSet()
            if id!='':
                module13 = get_object_or_404(Module13,  id=self.kwargs['id'])
                form   = Module13FormView(instance=module13)
                form2 = Module13Grid2FormSet(instance=module13)
                form3 = Module13Grid3FormSet(instance=module13)
                form22 = Module13Grid22FormSet(instance=module13)
                form29 = Module13Grid29FormSet(instance=module13)
                form31 = Module13Grid31FormSet(instance=module13)
                form47 = Module13Matrix47FormSet(instance=module13)
                form66= Module13WeaknessesFormSet(instance=module13)
      
            context['module13']=module13
            context['form']=form
            context['form2']=form2
            context['form3']=form3
            context['form22']=form22
            context['form29']=form29
            context['form31']=form31
            context['form47']=form47
            context['form66']=form66
            
            
            context['module']=module
         
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(13,self.kwargs['sessionid'])

        context['can_see'] = can_see
        
        return context
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module13_create(request, country,sessionid=None):
    """ Create module13  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    module13_count= Module13.objects.filter(session=sessionid).order_by('-modify_date').count()
   
    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,13)
    is_pa_filled = is_problemanalysis_filled(sessionid,13)
    is_sa_filled = is_swotanalysis_filled(sessionid,13)
    is_lf_filled = is_logicalframework_filled(sessionid,13)

  
    if module13_count>0:
        can_edit=0
    else:
        if is_st_filled:
            if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,13):
                can_edit=1
                m_percentage=get_percentage_module_filled(13,sessionid)
                tot_percentage=get_tot_percentage(sessionid)
                if is_st_filled:
                    st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=13).id
                if is_pa_filled:
                    pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=13).id
                if is_sa_filled:
                    sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=13).id
                if is_lf_filled:
                    lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=13).id
			
    if request.method == "POST":
        form = Module13Form(request.POST, request.FILES)
        form2 = Module13Grid2FormSet(request.POST)
        form3 = Module13Grid3FormSet(request.POST)
        form22 = Module13Grid22FormSet(request.POST)
        form29 = Module13Grid29FormSet(request.POST)
        form31 = Module13Grid31FormSet(request.POST)
        form47 = Module13Matrix47FormSet(request.POST)
        form66= Module13WeaknessesFormSet(request.POST)
      
        if form.is_valid() and form2.is_valid() and  form3.is_valid() and  form29.is_valid() and form22.is_valid() and form31.is_valid() and form47.is_valid() and form66.is_valid():
            new_mod13 = form.save(commit=False)
            new_mod13.author = request.user
            new_mod13.session=pceversion
            form.save()
            
            form2.instance = new_mod13
            form2.save()
            form3.instance = new_mod13
            form3.save()
            form22.instance = new_mod13
            form22.save()
            form29.instance = new_mod13
            form29.save()
            form31.instance = new_mod13
            form31.save()
            form47.instance = new_mod13
            form47.save()
            form66.instance = new_mod13
            form66.save()
          
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Module 13."))

            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
        else:
             return render_to_response('pce/module_13.html', {'context':'Edit','form': form,'form2': form2,'form3': form3,'form22': form22,'form29': form29,'form31': form31,'form47': form47,'form66': form66,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':13,  },
             context_instance=RequestContext(request))

    else:
        form = Module13Form(initial={'country': country,'session': pceversion.id})
        form2 = Module13Grid2FormSet()
        form3 = Module13Grid3FormSet()
        form22 = Module13Grid22FormSet()
        form29 = Module13Grid29FormSet()
        form31 = Module13Grid31FormSet()
        form47 = Module13Matrix47FormSet()
        form66= Module13WeaknessesFormSet()
      
      

    return render_to_response('pce/module_13.html', {'context':'Edit','form': form,'form2': form2,'form3': form3,'form22': form22,'form29': form29,'form31': form31,'form47': form47,'form66': form66,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':13,  },
            context_instance=RequestContext(request))

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def module13_edit(request, country, id=None,sessionid=None, template_name='pce/module_13.html'):
    """ Edit module_13 """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)

    can_edit=0
    m_percentage=0   
    tot_percentage=''
    st_id=''
    pa_id=''
    sa_id=''
    lf_id=''
    is_st_filled = is_stakeholder_filled(sessionid,13)
    is_pa_filled = is_problemanalysis_filled(sessionid,13)
    is_sa_filled = is_swotanalysis_filled(sessionid,13)
    is_lf_filled = is_logicalframework_filled(sessionid,13)          
    
    if id:
        module13 = get_object_or_404(Module13, pk=id)
        if pceversion.status==1 and  canEdit(sessionid,pceversion.country,user,13):
            can_edit=1
            m_percentage=get_percentage_module_filled(13,sessionid)
            tot_percentage=get_tot_percentage(sessionid)
            if is_st_filled:
                st_id =  get_object_or_404(Stakeholders, session_id=sessionid,module=13).id
            if is_pa_filled:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=sessionid,module=13).id
            if is_sa_filled:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=sessionid,module=13).id
            if is_lf_filled:
                lf_id =  get_object_or_404(LogicalFramework, session_id=sessionid,module=13).id
			
    else:
        module13 = Module13()
  
    if request.POST:
        form = Module13Form(request.POST,  request.FILES, instance=module13)
        form2 = Module13Grid2FormSet(request.POST,  request.FILES, instance=module13)
        form3 = Module13Grid3FormSet(request.POST,  request.FILES, instance=module13)
        form22 = Module13Grid22FormSet(request.POST,  request.FILES, instance=module13)
        form29 = Module13Grid29FormSet(request.POST,  request.FILES, instance=module13)
        form31 = Module13Grid31FormSet(request.POST,  request.FILES, instance=module13)
        form47 = Module13Matrix47FormSet(request.POST,  request.FILES, instance=module13)
        form66= Module13WeaknessesFormSet(request.POST,  request.FILES, instance=module13)
      
        #print('form:'+str(form.is_valid()))
        #print('form2:'+str(form2.is_valid()))
        #print('form3:'+str(form3.is_valid()))
        #print('form22:'+str(form22.is_valid()))
        #print('form29:'+str(form29.is_valid()))
        #print('form31:'+str(form31.is_valid()))
        #print('form47:'+str(form47.is_valid()))
        #print('form66:'+str(form66.is_valid()))
        
        if form.is_valid() and form2.is_valid() and  form3.is_valid() and  form29.is_valid() and form22.is_valid() and form31.is_valid() and form47.is_valid() and form66.is_valid():
            form.save()
          
        
            form2.instance = module13
            form2.save()
            form3.instance = module13
            form3.save()
            form29.instance = module13
            form29.save()
            form22.instance = module13
            form22.save()
            form31.instance = module13
            form31.save()
            form47.instance = module13
            form47.save()
            form66.instance = module13
            form66.save()
            
            
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
    
            info(request, _("Successfully saved Module 13"))
            
            
            return redirect("module-list",country=user_country_slug,id=pceversion.id,)
    else:
        form = Module13Form(instance=module13)
        form2 = Module13Grid2FormSet(instance=module13)
        form3 = Module13Grid3FormSet(instance=module13)
        form22 = Module13Grid22FormSet(instance=module13)
        form29 = Module13Grid29FormSet(instance=module13)
        form31 = Module13Grid31FormSet(instance=module13)
        form47 = Module13Matrix47FormSet(instance=module13)
        form66= Module13WeaknessesFormSet(instance=module13)
      
    return render_to_response(template_name, {
       'context':'Edit','form': form,'form2': form2,'form3': form3,'form22': form22,'form29': form29,'form31': form31,'form47': form47,'form66': form66,'sessionid':sessionid,'can_edit':can_edit,'m_percentage':m_percentage,'tot_percentage':tot_percentage,'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,'is_lf_filled':is_lf_filled,'st_id':st_id,'pa_id':pa_id,'sa_id':sa_id,'lf_id':lf_id,'module':13,  
    }, context_instance=RequestContext(request))


class StakeholdersListView(ListView):
    context_object_name = 'latest'
    model = Stakeholders
    date_field = 'publish_date'
    template_name = 'pce/stakeholders.html'
    queryset = StakeholdersFields.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return Stakeholders from the specific country """
        self.id = None
        if 'id' in self.kwargs:
            self.id = self.kwargs['id']
        return StakeholdersFields.objects.filter(stakeholder_id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(StakeholdersListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['module'] = self.kwargs['module']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        context['roles'] =ROLE
        context['interests'] =INTEREST
        context['influences'] =INTEREST
        context['levels'] =LEVEL
        
        stakeid=''
        if 'id' in self.kwargs:
            stakeid= self.kwargs['id']
        context['stakeid'] = stakeid
      
        pa_id=''
        st_id=''
        sa_id=''
        lf_id=''
        
        can_edit=0
        can_see=0
        session = get_object_or_404(PceVersion,  pk=self.kwargs['sessionid'])
       
        if canSee(self.kwargs['sessionid'],session.country,self.request.user,str(self.kwargs['module'])):
            can_see=1
        if canEdit(self.kwargs['sessionid'],session.country,self.request.user,str(self.kwargs['module'])) and session.status==1:
            can_edit=1
        if  can_see or can_edit: 
            context['moduleid']=getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[0]
            if context['moduleid']!= '':
                 context['modulename'] = getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[1]
            else:   
                context['modulename'] = 'module'+str(self.kwargs['module'])
                
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(int(self.kwargs['module']),self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],int(self.kwargs['module']))
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],int(self.kwargs['module']))
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],int(self.kwargs['module']))
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],int(self.kwargs['module']))
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
        
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        return context



class StakeholdersListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = Stakeholders
    date_field = 'publish_date'
    template_name = "pce/stakeholders.html"
     
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(StakeholdersListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['module'] = self.kwargs['module']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'
        ROLE_ = (
            (0, _("--- Please select ---")),
            (1, _("Group leader")),
            (2, _("Client")),
            (3, _("Oversight")),
            (4, _("Advocate")),
        )
     
        INTEREST_ = (
            (0, _("--- Please select ---")),
            (1, _("Low")),
            (2, _("High")),

        )            

        IMPORTANCE_ = (
            (0, _("--- Please select ---")),
            (1, _("0")),
            (2, _("1")),
            (3, _("2")),
            (4, _("3")),
            (5, _("4")),
        )   

       
        LEVEL_ = (
            (0, _("--- Please select ---")),
            (1, _("WORKSHOP PARTICIPANT")),
            (2, _("SURVEY PARTICIPANT")),
            (3, _("FOCUS GROUP MEMBER")),
            (4, _("KEEP INFORMED")),
        )  

        context['roles'] =ROLE_
        context['interests'] =INTEREST_
        context['influences'] =INTEREST_
        context['levels'] =LEVEL_
        context['latest'] =StakeholdersFields.objects.filter(stakeholder_id=  self.kwargs['id'])
    

        stakeid=''
        if 'id' in self.kwargs:
            stakeid= self.kwargs['id']
        context['stakeid'] = stakeid
      
        can_see=0
        session = get_object_or_404(PceVersion,  pk=self.kwargs['sessionid'])
       
        if canSee(self.kwargs['sessionid'],session.country,self.request.user,str(self.kwargs['module'])):
            can_see=1
        if  can_see: 
            context['moduleid']=getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[0]
            if context['moduleid']!= '':
                 context['modulename'] = getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[1]
            else:   
                context['modulename'] = 'module'+str(self.kwargs['module'])
                
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(int(self.kwargs['module']),self.kwargs['sessionid'])
            
        context['can_see'] = can_see
        context['loop_times'] = range(1, 10)

        return context
    
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def stakeholders_create(request, country,sessionid=None,module=None):
    """ Create stakeholders  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    #print("-------------ppppp-------------------------------")
    ##print(pceversion)
    
    percentage_module=get_percentage_module_filled(int(module),sessionid)
  
    is_st_filled=( is_stakeholder_filled(sessionid,module)>0)
    is_pa_filled= ( is_problemanalysis_filled(sessionid,module)>0)
    is_sa_filled= ( is_swotanalysis_filled(sessionid,module)>0)
    is_lf_filled=(  is_logicalframework_filled(sessionid,module)>0)

    stakeholders_count = Stakeholders.objects.filter(session=sessionid,module=module).count()
    #print('sessionid'+str(sessionid))
    #print('module1'+str(module))
    

    can_edit=0
    if stakeholders_count>0:
        can_edit=0
        #print('qui')
    else: 
        #print('qui 1')
        if pceversion.status==1 and canEdit(sessionid,pceversion.country,user,str(module)):
            can_edit=1
            #print('qui 1 edit')
            
    if request.method == "POST":
        form = StakeholdersForm(request.POST, request.FILES)
        form2 = StakeholdersFieldsFormSet(request.POST)
      
        if form.is_valid() and form2.is_valid():
            new_stake = form.save(commit=False)
            new_stake.author = request.user
            new_stake.session=pceversion
            new_stake.country = country
            new_stake.module = module
            form.save()
           
            form2.instance = new_stake
            form2.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Stakeholders for Module "+str(module)))
            return redirect("stakeholders-list", country=user_country_slug, sessionid=sessionid, module=module,id=new_stake.id)
    
        else:
             return render_to_response('pce/stakeholders.html', {'context': 'Edit','form': form,'form2': form2,'sessionid':sessionid,'can_edit':can_edit,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'module':module,'sessionid':pceversion.id, 'm_percentage':percentage_module,  'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled, },
             context_instance=RequestContext(request))
    else:
        form = StakeholdersForm(initial={'country': country,'session': pceversion.id,'module': module})
        form2 = StakeholdersFieldsFormSet()
       

    return render_to_response('pce/stakeholders.html',{'context': 'Edit','form': form,'form2': form2,'can_edit':can_edit,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'module':module,'sessionid':pceversion.id, 'm_percentage':percentage_module,  'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled, },
        context_instance=RequestContext(request))


        
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def stakeholders_edit(request, country, id=None,sessionid=None,module=None, template_name='pce/stakeholders.html'):
    """ Edit stakeholders """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    percentage_module=get_percentage_module_filled(int(module),sessionid)
  
    is_st_filled=( is_stakeholder_filled(sessionid,module)>0)
    is_pa_filled= ( is_problemanalysis_filled(sessionid,module)>0)
    is_sa_filled= ( is_swotanalysis_filled(sessionid,module)>0)
    is_lf_filled=(  is_logicalframework_filled(sessionid,module)>0)

 
    can_edit=0
    if id:
        stakeholder = get_object_or_404(Stakeholders, pk=id)
        if pceversion.status==1 and canEdit(sessionid,pceversion.country,user,str(module)):    
            can_edit=1
    else:
        stakeholder = Stakeholders()
      
    if request.method == "POST":
        form = StakeholdersForm(request.POST, request.FILES, instance=stakeholder)
        form2 = StakeholdersFieldsFormSet(request.POST, instance=stakeholder)
      
        if form.is_valid() and form2.is_valid():
            form.save()
           
            form2.instance = stakeholder
            form2.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Stakeholders for Module "+str(module)))
            return redirect("stakeholders-list", country=user_country_slug, sessionid=sessionid, module=module,id=id)
        else:
             return render_to_response('pce/stakeholders.html', {'context': 'Edit','form': form,'form2': form2,'can_edit':can_edit,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'module':module,'sessionid':pceversion.id,'m_percentage':percentage_module,  'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,},
             context_instance=RequestContext(request))
    else:
        form = StakeholdersForm(instance=stakeholder)
        form2 = StakeholdersFieldsFormSet(instance=stakeholder)
      
    return render_to_response(template_name, {
        'context': 'Edit','form': form,'form2': form2,'can_edit':can_edit,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'module':module,'sessionid':pceversion.id,'m_percentage':percentage_module,  'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,
    }, context_instance=RequestContext(request))
    
    
 
       
    
    
    
class ProblemAnalysisListView(ListView):
    context_object_name = 'latest'
    model = ProblemAnalysis
    date_field = 'publish_date'
    template_name = 'pce/problemanalysis.html'
    queryset = ProblemAnalysis.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return ProblemAnalysis from the specific country """
        self.id =None
        if 'id' in self.kwargs:
            self.id = self.kwargs['id']
        return ProblemAnalysis.objects.filter(id = self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(ProblemAnalysisListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['module'] = self.kwargs['module']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        pa_id=''
        st_id=''
        sa_id=''
        lf_id=''

        if 'id' in self.kwargs:
            pa_id = self.kwargs['id']
        context['pa_id'] = pa_id
       
        can_edit=0
        can_see=0
        session = get_object_or_404(PceVersion,  pk=self.kwargs['sessionid'])
        #print('check')
        if canSee(self.kwargs['sessionid'],session.country,self.request.user,str(self.kwargs['module'])):
            can_see=1
        if canEdit(self.kwargs['sessionid'],session.country,self.request.user,str(self.kwargs['module'])) and session.status==1:
            can_edit=1
        if  can_see or can_edit: 
            
            context['weakeness']=getWeakenessFromModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))
            
            
            context['moduleid']=getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[0]
            context['modulename'] = getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[1]
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(int(self.kwargs['module']),self.kwargs['sessionid'])
            context['is_st_filled'] = is_stakeholder_filled(self.kwargs['sessionid'],int(self.kwargs['module']))
            context['is_pa_filled'] = is_problemanalysis_filled(self.kwargs['sessionid'],int(self.kwargs['module']))
            context['is_sa_filled'] = is_swotanalysis_filled(self.kwargs['sessionid'],int(self.kwargs['module']))
            context['is_lf_filled'] = is_logicalframework_filled(self.kwargs['sessionid'],int(self.kwargs['module']))
      
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
        
        return context
    
class ProblemAnalysisListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = ProblemAnalysis
    date_field = 'publish_date'
    template_name = 'pce/problemanalysis.html'
    
  
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(ProblemAnalysisListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['module'] = self.kwargs['module']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] =  'Pdf'
        pa_id=''
        if 'id' in self.kwargs:
            pa_id = self.kwargs['id']
        context['pa_id'] = pa_id
        context['latest'] =  ProblemAnalysis.objects.filter(id = pa_id)
        
       
        can_see=0
        session = get_object_or_404(PceVersion,  pk=self.kwargs['sessionid'])
        if canSee(self.kwargs['sessionid'],session.country,self.request.user,str(self.kwargs['module'])):
            can_see=1
        if  can_see : 
            context['weakeness']=getWeakenessFromModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))
            
            context['moduleid']=getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[0]
            context['modulename'] = getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[1]
            context['version_number'] = session.version_number
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(self.kwargs['sessionid'])
            context['m_percentage'] = get_percentage_module_filled(int(self.kwargs['module']),self.kwargs['sessionid'])
         
        context['can_see'] = can_see
        context['pa_id'] = pa_id
        
        return context    

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def problemanalysis_create(request, country,sessionid=None,module=None):
    """ Create problemanalysis  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    
    percentage_module=get_percentage_module_filled(int(module),sessionid)
  
    is_st_filled=( is_stakeholder_filled(sessionid,module)>0)
    is_pa_filled= ( is_problemanalysis_filled(sessionid,module)>0)
    is_sa_filled= ( is_swotanalysis_filled(sessionid,module)>0)
    is_lf_filled=(  is_logicalframework_filled(sessionid,module)>0)
   
    can_edit=0
    weakeness=None

    if is_pa_filled:
        can_edit = 0     
    else:
        if is_st_filled and percentage_module==100:
          if pceversion.status==1 and canEdit(sessionid,pceversion.country,user,str(module)):
            can_edit=1
            weakeness=getWeakenessFromModuleNameAndId(int(module),sessionid)
    #print("weakeness")    
    ##print(weakeness)    
    if request.method == "POST":
        form = ProblemAnalysisForm(request.POST, request.FILES)
        
        if form.is_valid() :
            new_pa = form.save(commit=False)
            new_pa.author = request.user
            new_pa.session=pceversion
            new_pa.country = country
            new_pa.module = module
            form.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Problem Analysis for Module "+str(module)))
            return redirect("problemanalysis-list", country=user_country_slug, sessionid=sessionid,module=module,id=new_pa.id)
        else:
             return render_to_response('pce/problemanalysis.html', {'context': 'Edit','form': form,'can_edit':can_edit,'sessionid':sessionid,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'m_percentage':percentage_module,'module':module, 'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,'weakeness':weakeness},
             context_instance=RequestContext(request))
    else:
        form = ProblemAnalysisForm(initial={'country': country,'sessionid':sessionid,'module': module})
    
    return render_to_response('pce/problemanalysis.html',{'context': 'Edit','form': form,'can_edit':can_edit,'sessionid':sessionid,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'m_percentage':percentage_module,'module':module, 'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,'weakeness':weakeness},
        context_instance=RequestContext(request))

        
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def problemanalysis_edit(request, country, id=None,sessionid=None,module=None, template_name='pce/problemanalysis.html'):
    """ Edit problemanalysis """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    can_edit=0
    can_see=0
    weakeness=None
    if id:
        pa = get_object_or_404(ProblemAnalysis, pk=id)
        if pceversion.status==1 and canEdit(sessionid,pceversion.country,user,str(module)):
            can_edit=1
            weakeness=getWeakenessFromModuleNameAndId(int(module),sessionid)
            ##print("weakeness")    
            ##print(weakeness.w2)
    else:
        pa = ProblemAnalysis()
    
    percentage_module=get_percentage_module_filled(int(module),sessionid)
  
    is_st_filled=( is_stakeholder_filled(sessionid,module)>0)
    is_pa_filled= ( is_problemanalysis_filled(sessionid,module)>0)
    is_sa_filled= ( is_swotanalysis_filled(sessionid,module)>0)
    is_lf_filled=(  is_logicalframework_filled(sessionid,module)>0)
   
    if request.method == "POST":
        form = ProblemAnalysisForm(request.POST, request.FILES, instance=pa)
        
        if form.is_valid() :
            form.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved Problem Analysis for Module "+str(module)))
            
            return redirect("problemanalysis-list", country=user_country_slug, sessionid=sessionid,module=module,id=id)
        else:
             return render_to_response('pce/problemanalysis.html', {'context': 'Edit','form': form,'can_edit':can_edit,'can_see':can_see,'sessionid':sessionid,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'m_percentage':percentage_module,'module':module, 'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,'weakeness':weakeness},
             context_instance=RequestContext(request))
       
    else:
        form = ProblemAnalysisForm(instance=pa)
        
      
      
    return render_to_response(template_name, {
       'context': 'Edit', 'form': form,'can_edit':can_edit,'can_see':can_see,'sessionid':sessionid,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'m_percentage':percentage_module,'module':module, 'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,'weakeness':weakeness},
        context_instance=RequestContext(request))
    
    
  
class SwotAnalysisListView(ListView):
    context_object_name = 'latest'
    model = SwotAnalysis
    date_field = 'publish_date'
    template_name = 'pce/swotanalysis.html'
    queryset = SwotAnalysis.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return SwotAnalysis from the specific country """
        self.id = None
        if 'id' in self.kwargs:
            self.id = self.kwargs['id']
        return SwotAnalysis.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(SwotAnalysisListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['module'] = self.kwargs['module']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        #context['sa_id'] = self.kwargs['id']
       
      
         
        latest1= SwotAnalysis1.objects.filter(swotanalysis_id= self.id)
        latest2= SwotAnalysis2.objects.filter(swotanalysis_id= self.id)
        latest3= SwotAnalysis3.objects.filter(swotanalysis_id= self.id)
        latest4= SwotAnalysis4.objects.filter(swotanalysis_id= self.id)
        latest5= SwotAnalysis5.objects.filter(swotanalysis_id= self.id)
        context['latest1'] = latest1
        context['latest2'] = latest2
        context['latest3'] = latest3
        context['latest4'] = latest4
        context['latest5'] = latest5
        context['priority'] =PRIORITY
        context['type'] =TYPE
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
        can_edit=0
        can_see=0
        session = get_object_or_404(PceVersion,  pk=self.kwargs['sessionid'])
   
        if canSee(self.kwargs['sessionid'],session.country,self.request.user,str(self.kwargs['module'])):
            can_see=1
        if canEdit(self.kwargs['sessionid'],session.country,self.request.user,str(self.kwargs['module'])) and session.status==1:
            can_edit=1
        if  can_see or can_edit: 
            
            context['weakeness']=getWeakenessFromModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))
            
            context['problemanalysis'] = get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=self.kwargs['module'])
            context['version_number'] = session.version_number            
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(session.id)
            context['m_percentage'] = get_percentage_module_filled(int(self.kwargs['module']), int(self.kwargs['sessionid']))
            context['is_st_filled'] = is_stakeholder_filled(int(self.kwargs['sessionid']),int(self.kwargs['module']))
            context['is_pa_filled'] = is_problemanalysis_filled(int(self.kwargs['sessionid']),int(self.kwargs['module']))
            context['is_sa_filled'] = is_swotanalysis_filled(int(self.kwargs['sessionid']),int(self.kwargs['module']))
            context['is_lf_filled'] = is_logicalframework_filled(int(self.kwargs['sessionid']),int(self.kwargs['module']))
            context['moduleid']=getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[0]
            context['modulename'] = getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[1]
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id

      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
       
        return context
    
class SwotAnalysisListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = SwotAnalysis
    date_field = 'publish_date'
    template_name = 'pce/swotanalysis.html'
    
  
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(SwotAnalysisListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['module'] = self.kwargs['module']
        context['sessionid'] = self.kwargs['sessionid']
        sa_id=''
        if 'id' in self.kwargs:
            sa_id = self.kwargs['id']
        context['sa_id'] = sa_id
        context['latest'] =  SwotAnalysis.objects.filter(id= sa_id)
        context['context'] = 'Pdf'
        context['weakeness']=getWeakenessFromModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))
            
         
        latest1= SwotAnalysis1.objects.filter(swotanalysis_id= sa_id)
        latest2= SwotAnalysis2.objects.filter(swotanalysis_id= sa_id)
        latest3= SwotAnalysis3.objects.filter(swotanalysis_id= sa_id)
        latest4= SwotAnalysis4.objects.filter(swotanalysis_id= sa_id)
        latest5= SwotAnalysis5.objects.filter(swotanalysis_id= sa_id)
        context['latest1'] = latest1
        context['latest2'] = latest2
        context['latest3'] = latest3
        context['latest4'] = latest4
        context['latest5'] = latest5
     
        PRIORITY_ = (
            (PRIORITY_0, _("--- Please select ---")),
            (PRIORITY_1, _("Very Low")),
            (PRIORITY_2, _("Low")),
            (PRIORITY_3, _("Medium")),
            (PRIORITY_4, _("High")),
            (PRIORITY_5, _("Very high")),
        )
    
        TYPE_ = (
            (TYPE_0, _("--- Please select ---")),
            (TYPE_1, _("A) National coordination and political willigness")),
            (TYPE_2, _("B) Like A plus small technical assisstance")),
            (TYPE_3, _("C) Like A plus significant investments")),
        )        

        context['priority'] =PRIORITY_
        context['type'] =TYPE_
        sa_id=''
       
        can_see=0
        session = get_object_or_404(PceVersion,  pk=self.kwargs['sessionid'])
   
        if canSee(self.kwargs['sessionid'],session.country,self.request.user,str(self.kwargs['module'])):
            can_see=1
        if  can_see : 
            
            
            context['problemanalysis'] = get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=self.kwargs['module'])
            context['version_number'] = session.version_number            
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(session.id)
            context['m_percentage'] = get_percentage_module_filled(int(self.kwargs['module']), int(self.kwargs['sessionid']))
            context['moduleid']=getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[0]
            context['modulename'] = getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[1]
          
      
        context['can_see'] = can_see
     
       
        
        return context    

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def swotanalysis_create(request, country,sessionid=None,module=None):
    """ Create SwotAnalysis  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)

    percentage_module=get_percentage_module_filled(int(module),sessionid)
    
    is_st_filled=( is_stakeholder_filled(sessionid,module)>0)
    is_pa_filled= ( is_problemanalysis_filled(sessionid,module)>0)
    is_sa_filled= ( is_swotanalysis_filled(sessionid,module)>0)
    is_lf_filled=(  is_logicalframework_filled(sessionid,module)>0)
    can_edit=0
    
    problemanalysis=None
    if is_sa_filled:
        can_edit = 0     
    else:
        if is_st_filled and is_pa_filled and percentage_module==100:
           if pceversion.status==1 and canEdit(sessionid,pceversion.country,user,str(module)):
            can_edit=1    
            problemanalysis = get_object_or_404(ProblemAnalysis, session_id=sessionid,module=module)
            weakeness=getWeakenessFromModuleNameAndId(int(module),sessionid)
            
    if request.method == "POST":
        form = SwotAnalysisForm(request.POST, request.FILES)
        form1 = SwotAnalysis1FormSet(request.POST)
        form2 = SwotAnalysis2FormSet(request.POST)
        form3 = SwotAnalysis3FormSet(request.POST)
        form4 = SwotAnalysis4FormSet(request.POST)
        form5 = SwotAnalysis5FormSet(request.POST)
        
        if form.is_valid() and form1.is_valid() and form2.is_valid() and form3.is_valid() and form4.is_valid() and form5.is_valid():
            new_sa = form.save(commit=False)
            new_sa.author = request.user
            new_sa.session=pceversion
            new_sa.country = country
            new_sa.module = module
            form.save()
           
            form1.instance = new_sa
            form1.save()
            form2.instance = new_sa
            form2.save()
            form3.instance = new_sa
            form3.save()
            form4.instance = new_sa
            form4.save()
            form5.instance = new_sa
            form5.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved SwotAnalysis for Module "+str(module)))
            return redirect("swotanalysis-list",country=user_country_slug, sessionid=sessionid,module=module,id=new_sa.id)
        else:
             return render_to_response('pce/swotanalysis.html', {'context':'Edit','form': form,'form1': form1,'form2': form2,'form3': form3,'form4': form4,'form5': form5,'can_edit':can_edit,'sessionid':sessionid,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'m_percentage':percentage_module,'module':module, 'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,'problemanalysis':problemanalysis,'weakeness':weakeness},
             context_instance=RequestContext(request))
    else:
        form = SwotAnalysisForm(initial={'country': country,'session': pceversion.id,'module': module})
        form1 = SwotAnalysis1FormSet()
        form2 = SwotAnalysis2FormSet()
        form3 = SwotAnalysis3FormSet()
        form4 = SwotAnalysis4FormSet()
        form5 = SwotAnalysis5FormSet()
    
    return render_to_response('pce/swotanalysis.html',{'context':'Edit','form': form,'form1': form1,'form2': form2,'form3': form3,'form4': form4,'form5': form5,'can_edit':can_edit,'sessionid':sessionid,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'m_percentage':percentage_module,'module':module, 'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,'problemanalysis':problemanalysis,'weakeness':weakeness},
        context_instance=RequestContext(request))

        
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def swotanalysis_edit(request, country, id=None,sessionid=None,module=None, template_name='pce/swotanalysis.html'):
    """ Edit SwotAnalysis """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    percentage_module=get_percentage_module_filled(int(module),sessionid)
    
    is_st_filled=( is_stakeholder_filled(sessionid,module)>0)
    is_pa_filled= ( is_problemanalysis_filled(sessionid,module)>0)
    is_sa_filled= ( is_swotanalysis_filled(sessionid,module)>0)
    is_lf_filled=(  is_logicalframework_filled(sessionid,module)>0)
    can_edit=0
    
    problemanalysis=None
    if id:
        sa = get_object_or_404(SwotAnalysis, pk=id)
        if pceversion.status==1 and canEdit(sessionid,pceversion.country,user,str(module)):
            can_edit=1
            problemanalysis = get_object_or_404(ProblemAnalysis, session_id=sessionid,module=module)
            weakeness=getWeakenessFromModuleNameAndId(int(module),sessionid)
    else:
        sa = SwotAnalysis()
         
    if request.method == "POST":
        form = SwotAnalysisForm(request.POST, request.FILES, instance=sa)
        form1 = SwotAnalysis1FormSet(request.POST, instance=sa)
        form2 = SwotAnalysis2FormSet(request.POST, instance=sa)
        form3 = SwotAnalysis3FormSet(request.POST, instance=sa)
        form4 = SwotAnalysis4FormSet(request.POST, instance=sa)
        form5 = SwotAnalysis5FormSet(request.POST, instance=sa)
        
        if form.is_valid() and form1.is_valid() and form2.is_valid() and form3.is_valid() and form4.is_valid() and form5.is_valid():
            form.save()
           
            form1.instance = sa
            form1.save()
            form2.instance = sa
            form2.save()
            form3.instance = sa
            form3.save()
            form4.instance = sa
            form4.save()
            form5.instance = sa
            form5.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved SwotAnalysis for Module "+str(module)))
            return redirect("swotanalysis-list",country=user_country_slug, sessionid=sessionid,module=module,id=id)
        else:
             return render_to_response('pce/swotanalysis.html', {'context':'Edit','form': form,'form1': form1,'form2': form2,'form3': form3,'form4': form4,'form5': form5,'can_edit':can_edit,'sessionid':sessionid,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'m_percentage':percentage_module,'module':module, 'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,'problemanalysis':problemanalysis,'weakeness':weakeness},
             context_instance=RequestContext(request))
    else:
        form = SwotAnalysisForm(instance=sa)
        form1 = SwotAnalysis1FormSet(instance=sa)
        form2 = SwotAnalysis2FormSet(instance=sa)
        form3 = SwotAnalysis3FormSet(instance=sa)
        form4 = SwotAnalysis4FormSet(instance=sa)
        form5 = SwotAnalysis5FormSet(instance=sa)
      
    return render_to_response(template_name, {
        'context':'Edit','form': form,'form1': form1,'form2': form2,'form3': form3,'form4': form4,'form5': form5,'can_edit':can_edit,'sessionid':sessionid,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'m_percentage':percentage_module,'module':module, 'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,'problemanalysis':problemanalysis,'weakeness':weakeness
    }, context_instance=RequestContext(request))
   
  
class LogicalFrameworkListView(ListView):
    context_object_name = 'latest'
    model = LogicalFramework
    date_field = 'publish_date'
    template_name = 'pce/logicalframework.html'
    queryset = LogicalFramework.objects.all().order_by('-modify_date', 'title')
    allow_future = False
    allow_empty = True
    paginate_by = 500

    def get_queryset(self):
        """ only return LogicalFramework from the specific country """
        self.id =None
        if 'id' in self.kwargs:
             self.id =self.kwargs['id']
        return LogicalFramework.objects.filter(id= self.id)
    
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(LogicalFrameworkListView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['module'] = self.kwargs['module']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'View'
        
        st_id=''
        pa_id=''
        sa_id=''
        lf_id=''
        can_edit=0
        can_see=0
        session = get_object_or_404(PceVersion,  pk=self.kwargs['sessionid'])
   
        if canSee(self.kwargs['sessionid'],session.country,self.request.user,str(self.kwargs['module'])):
            can_see=1
        if canEdit(self.kwargs['sessionid'],session.country,self.request.user,str(self.kwargs['module'])) and session.status==1:
            can_edit=1
        if  can_see or can_edit: 
            latest1= LogicalFrameworkAct1.objects.filter(logicalframework_id= self.id)
            latest2= LogicalFrameworkAct2.objects.filter(logicalframework_id= self.id)
            latest3= LogicalFrameworkAct3.objects.filter(logicalframework_id= self.id)
            latest4= LogicalFrameworkAct4.objects.filter(logicalframework_id= self.id)
            latest5= LogicalFrameworkAct5.objects.filter(logicalframework_id= self.id)
            context['problemanalysis']    = get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=self.kwargs['module'])
            context['weakeness']=getWeakenessFromModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))
       
            context['latest1'] = latest1
            context['latest2'] = latest2
            context['latest3'] = latest3
            context['latest4'] = latest4
            context['latest5'] = latest5
            context['version_number'] = session.version_number            
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(session.id)
            context['m_percentage'] = get_percentage_module_filled(int(self.kwargs['module']), int(self.kwargs['sessionid']))
            context['is_st_filled'] = is_stakeholder_filled(int(self.kwargs['sessionid']),int(self.kwargs['module']))
            context['is_pa_filled'] = is_problemanalysis_filled(int(self.kwargs['sessionid']),int(self.kwargs['module']))
            context['is_sa_filled'] = is_swotanalysis_filled(int(self.kwargs['sessionid']),int(self.kwargs['module']))
            context['is_lf_filled'] = is_logicalframework_filled(int(self.kwargs['sessionid']),int(self.kwargs['module']))
            context['moduleid']=getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[0]
            context['modulename'] = getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[1]
            if context['is_st_filled']:
                st_id =  get_object_or_404(Stakeholders, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
            if context['is_pa_filled']:
                pa_id =  get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
            if context['is_sa_filled']:
                sa_id =  get_object_or_404(SwotAnalysis, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id
            if context['is_lf_filled']:
                lf_id =  get_object_or_404(LogicalFramework, session_id=self.kwargs['sessionid'],module=self.kwargs['module']).id

      
        context['can_see'] = can_see
        context['can_edit'] = can_edit
        context['st_id'] = st_id
        context['pa_id'] = pa_id
        context['sa_id'] = sa_id
        context['lf_id'] = lf_id
      
        return context
    
class LogicalFrameworkListPDFView(PDFTemplateView):
    context_object_name = 'latest'
    model = LogicalFramework
    date_field = 'publish_date'
    template_name = 'pce/logicalframework.html'
 
    def get_context_data(self, **kwargs): # http://stackoverflow.com/a/15515220
        context = super(LogicalFrameworkListPDFView, self).get_context_data(**kwargs)
        context['country'] = self.kwargs['country']
        context['module'] = self.kwargs['module']
        context['sessionid'] = self.kwargs['sessionid']
        context['context'] = 'Pdf'
       
        context['lf_id'] = self.kwargs['id']
        context['latest'] =     LogicalFramework.objects.filter(id= self.kwargs['id'])
  
        can_see=0
        session = get_object_or_404(PceVersion,  pk=self.kwargs['sessionid'])
   
        if canSee(self.kwargs['sessionid'],session.country,self.request.user,str(self.kwargs['module'])):
            can_see=1
        if  can_see : 
            latest1= LogicalFrameworkAct1.objects.filter(logicalframework_id= self.kwargs['id'])
            latest2= LogicalFrameworkAct2.objects.filter(logicalframework_id= self.kwargs['id'])
            latest3= LogicalFrameworkAct3.objects.filter(logicalframework_id= self.kwargs['id'])
            latest4= LogicalFrameworkAct4.objects.filter(logicalframework_id= self.kwargs['id'])
            latest5= LogicalFrameworkAct5.objects.filter(logicalframework_id= self.kwargs['id'])
            context['problemanalysis']    = get_object_or_404(ProblemAnalysis, session_id=self.kwargs['sessionid'],module=self.kwargs['module'])
            context['weakeness']=getWeakenessFromModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))
       
            context['latest1'] = latest1
            context['latest2'] = latest2
            context['latest3'] = latest3
            context['latest4'] = latest4
            context['latest5'] = latest5
            context['version_number'] = session.version_number            
            context['sessionstatus'] = session.status
            context['tot_percentage'] = get_tot_percentage(session.id)
            context['m_percentage'] = get_percentage_module_filled(int(self.kwargs['module']), int(self.kwargs['sessionid']))
            context['moduleid']=getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[0]
            context['modulename'] = getModuleNameAndId(int(self.kwargs['module']),int(self.kwargs['sessionid']))[1]
        context['can_see'] = can_see
       
        return context
    
    
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def logicalframework_create(request, country,sessionid=None,module=None):
    """ Create LogicalFramework  """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)

    percentage_module=get_percentage_module_filled(int(module),sessionid)
    
    is_st_filled=( is_stakeholder_filled(sessionid,module)>0)
    is_pa_filled= ( is_problemanalysis_filled(sessionid,module)>0)
    is_sa_filled= ( is_swotanalysis_filled(sessionid,module)>0)
    is_lf_filled=(  is_logicalframework_filled(sessionid,module)>0)
    weakeness=None
  
    can_edit=0
    problemanalysis=None
    if is_lf_filled:
        can_edit = 0     
    else:
        if is_st_filled and is_pa_filled and is_sa_filled and percentage_module==100:
          if pceversion.status==1 and canEdit(sessionid,pceversion.country,user,str(module)):
            can_edit=1
            problemanalysis    = get_object_or_404(ProblemAnalysis, session_id=sessionid,module=module)
            weakeness=getWeakenessFromModuleNameAndId(int(module),sessionid)
       
           
    if request.method == "POST":
        form = LogicalFrameworkForm(request.POST, request.FILES)
        form1 = LogicalFrameworkAct1FormSet(request.POST)
        form2 = LogicalFrameworkAct2FormSet(request.POST)
        form3 = LogicalFrameworkAct3FormSet(request.POST)
        form4 = LogicalFrameworkAct4FormSet(request.POST)
        form5 = LogicalFrameworkAct5FormSet(request.POST)
        
        if form.is_valid() and form1.is_valid() and form2.is_valid() and form3.is_valid() and form4.is_valid() and form5.is_valid():
            new_lf = form.save(commit=False)
            new_lf.author = request.user
            new_lf.session=pceversion
            new_lf.country = country
            new_lf.module = module
            form.save()
           
            form1.instance = new_lf
            form1.save()
            form2.instance = new_lf
            form2.save()
            form3.instance = new_lf
            form3.save()
            form4.instance = new_lf
            form4.save()
            form5.instance = new_lf
            form5.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
           
            info(request, _("Successfully saved LogicalFramework for Module "+str(module)))
            
            return redirect("logicalframework-list", country=user_country_slug, sessionid=sessionid,module=module,id=new_lf.id)
        else:
             return render_to_response('pce/logicalframework.html', {'context': 'Edit','form': form,'form1': form1,'form2': form2,'form3': form3,'form4': form4,'form5': form5,'can_edit':can_edit,'sessionid':sessionid,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'m_percentage':percentage_module,'module':module, 'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,'problemanalysis':problemanalysis,'weakeness':weakeness},
             context_instance=RequestContext(request))
       
    else:
        form = LogicalFrameworkForm(initial={'country': country,'session': pceversion.id,'module': module})
        form1 = LogicalFrameworkAct1FormSet()
        form2 = LogicalFrameworkAct2FormSet()
        form3 = LogicalFrameworkAct3FormSet()
        form4 = LogicalFrameworkAct4FormSet()
        form5 = LogicalFrameworkAct5FormSet()
        
       

    return render_to_response('pce/logicalframework.html',{'context': 'Edit','form': form,'form1': form1,'form2': form2,'form3': form3,'form4': form4,'form5': form5,'can_edit':can_edit,'sessionid':sessionid,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'m_percentage':percentage_module,'module':module, 'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,'problemanalysis':problemanalysis,'weakeness':weakeness},
        context_instance=RequestContext(request))

        
@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def logicalframework_edit(request, country, id=None,sessionid=None,module=None, template_name='pce/logicalframework.html'):
    """ Edit LogicalFramework """
    user = request.user
    author = user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    pceversion = get_object_or_404(PceVersion,  pk=sessionid)
    percentage_module=get_percentage_module_filled(int(module),sessionid)
    
    is_st_filled=( is_stakeholder_filled(sessionid,module)>0)
    is_pa_filled= ( is_problemanalysis_filled(sessionid,module)>0)
    is_sa_filled= ( is_swotanalysis_filled(sessionid,module)>0)
    is_lf_filled=(  is_logicalframework_filled(sessionid,module)>0)
    weakeness=None
  
    can_edit=0
    problemanalysis=None
    if id:
        lf = get_object_or_404(LogicalFramework, pk=id)
        if pceversion.status==1 and canEdit(sessionid,pceversion.country,user,str(module)):
            can_edit=1
            problemanalysis    = get_object_or_404(ProblemAnalysis, session_id=sessionid,module=module)
            weakeness=getWeakenessFromModuleNameAndId(int(module),sessionid)
    else:
        lf = LogicalFramework()
   
             
    if request.method == "POST":
        form = LogicalFrameworkForm(request.POST, request.FILES, instance=lf)
        form1 = LogicalFrameworkAct1FormSet(request.POST, instance=lf)
        form2 = LogicalFrameworkAct2FormSet(request.POST, instance=lf)
        form3 = LogicalFrameworkAct3FormSet(request.POST, instance=lf)
        form4 = LogicalFrameworkAct4FormSet(request.POST, instance=lf)
        form5 = LogicalFrameworkAct5FormSet(request.POST, instance=lf)
        
        if form.is_valid() and form1.is_valid() and form2.is_valid() and form3.is_valid() and form4.is_valid() and form5.is_valid():
            form.save()
           
            form1.instance = lf
            form1.save()
            form2.instance = lf
            form2.save()
            form3.instance = lf
            form3.save()
            form4.instance = lf
            form4.save()
            form5.instance = lf
            form5.save()
            pceversion.modify_date=timezone.now()
            pceversion.save()
            
            info(request, _("Successfully saved LogicalFramework for Module "+str(module)))
          
            return redirect("logicalframework-list", country=user_country_slug, sessionid=sessionid,module=module,id=id)
        else:
             return render_to_response('pce/logicalframework.html', {'context': 'Edit','form': form,'form1': form1,'form2': form2,'form3': form3,'form4': form4,'form5': form5,'can_edit':can_edit,'sessionid':sessionid,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'m_percentage':percentage_module,'module':module, 'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,'problemanalysis':problemanalysis,'weakeness':weakeness},
             context_instance=RequestContext(request))
       
    else:
        form = LogicalFrameworkForm(instance=lf)
  
        form1 = LogicalFrameworkAct1FormSet(instance=lf)
        form2 = LogicalFrameworkAct2FormSet(instance=lf)
        form3 = LogicalFrameworkAct3FormSet(instance=lf)
        form4 = LogicalFrameworkAct4FormSet(instance=lf)
        form5 = LogicalFrameworkAct5FormSet(instance=lf)
       
      
    return render_to_response(template_name, {
      'context': 'Edit','form': form,'form1': form1,'form2': form2,'form3': form3,'form4': form4,'form5': form5,'can_edit':can_edit,'sessionid':sessionid,'version_number':pceversion.version_number,'tot_percentage':get_tot_percentage(sessionid),'m_percentage':percentage_module,'module':module, 'is_st_filled':is_st_filled,'is_pa_filled':is_pa_filled,'is_sa_filled':is_sa_filled,    'is_lf_filled':is_lf_filled,'problemanalysis':problemanalysis,'weakeness':weakeness},
  context_instance=RequestContext(request))
  
from docx import Document
from docx.shared import Inches

@login_required
@permission_required('pce.add_pceversion', login_url="/accounts/login/")
def generate_report(request, country,sessionid=None):
    """ generate report  """
    user = request.user
    country=user.get_profile().country
    user_country_slug = lower(slugify(country))
    document = Document()
    docx_title="PCE_Report.docx"
    p= document.add_paragraph("")
    #p.add_run(_("1. Introduction:"+str(sessionid))).bold = True

    if sessionid:
        pceversion = get_object_or_404(PceVersion,  pk=sessionid)
        #p.add_run(_("1. Introduction:"+str(sessionid))).bold = True
        #p.add_run(_("-"+str(lower(pceversion.country)))).bold = True
        #p.add_run(_("-"+str(lower(user_country_slug)))).bold = True
        #p.add_run(_("-"+str(user.groups.filter(name='PCE Manager/Validator')))).bold = True

        if user.groups.filter(name='Admin') or (lower(slugify(pceversion.country)) == user_country_slug and user.groups.filter(name='PCE Manager/Validator')) :
            #p.add_run(_("1. Introduction:"+str(lower(pceversion.country) == user_country_slug))).bold = True
            m_names={}
            w2=None
            w3=None
            w4=None
            w5=None
            w6=None
            w7=None
            w8=None
            w9=None
            w10=None
            w11=None
            w12=None
            w13=None
            m1m7=''
            m1m8=''    
            m1m20=''
            m1m21=''
            m1m11=''
            m1m26=''
            m1m22=''
            m1m23=''
            module1=None
            items = pceversion.chosen_modules.split(',')
            for i in items:
                if i!='':
                    y=int(i)
                    if y == 1:
                        m_names[y]=_("COUNTRY PROFILE")
                    elif y == 2:
                        m_names[y]=_("National phytosanitary legislation")
                    elif y == 3:
                        m_names[y]=_("Environmental forces assessment")
                    elif y == 4:
                        m_names[y]=_("NPPO's mission and strategy")
                    elif y == 5:
                        m_names[y]=_("NPPO's structure and processes")
                    elif y == 6:
                        m_names[y]=_("NPPO's Resources")
                    elif y == 7:
                        m_names[y]=_("Pest diagnostic capacity")
                    elif y == 8:
                        m_names[y]=_("NPPO pest surveillance and pest reporting capacity")
                    elif y == 9:
                        m_names[y]=_("Pest eradication capacity")
                    elif y == 10:
                        m_names[y]=_("Phytosanitary import regulatory system")
                    elif y == 11:
                        m_names[y]=_("Pest risk analysis")
                    elif y == 12:
                        m_names[y]=_("Pest free areas, places and sites, low pest prevalence areas")
                    elif y ==  13:
                        m_names[y]=_("Export certification, re-export, transit")    
                    try:
                        if y==1:
                            module1 = get_object_or_404(Module1, session_id=sessionid)
                        elif y==2:
                            module2_1 = get_object_or_404(Module2_1, session_id=sessionid)
                            module2_2 = get_object_or_404(Module2_2, session_id=sessionid)
                            w2=getWeakenessFromModuleNameAndId(2,sessionid)
                        elif y==3:
                            module3 = get_object_or_404(Module3, session_id=sessionid)
                            w3=getWeakenessFromModuleNameAndId(3,sessionid)
                        elif y==4:
                            module4 = get_object_or_404(Module4, session_id=sessionid)
                            w4=getWeakenessFromModuleNameAndId(4,sessionid)
                        elif y==5:
                            module5 = get_object_or_404(Module5, session_id=sessionid)
                            w5=getWeakenessFromModuleNameAndId(5,sessionid)
                        elif y==6:
                            module6 = get_object_or_404(Module6, session_id=sessionid)
                            w6=getWeakenessFromModuleNameAndId(6,sessionid)
                        elif y==7:
                            module7 = get_object_or_404(Module7, session_id=sessionid)
                            w7=getWeakenessFromModuleNameAndId(7,sessionid)
                        elif y==8:
                            module8 = get_object_or_404(Module8, session_id=sessionid)
                            w8=getWeakenessFromModuleNameAndId(8,sessionid)
                        elif y==9:
                            module9 = get_object_or_404(Module9, session_id=sessionid)
                            w9=getWeakenessFromModuleNameAndId(9,sessionid)
                        elif y==10:
                            module10 = get_object_or_404(Module10, session_id=sessionid)
                            w10=getWeakenessFromModuleNameAndId(10,sessionid)
                        elif y==11:
                            module11 = get_object_or_404(Module11, session_id=sessionid)
                            w11=getWeakenessFromModuleNameAndId(11,sessionid)
                        elif y==12:
                            module12 = get_object_or_404(Module12, session_id=sessionid)
                            w12=getWeakenessFromModuleNameAndId(12,sessionid)
                        elif y==13:
                            module13 = get_object_or_404(Module13, session_id=sessionid)
                            w13=getWeakenessFromModuleNameAndId(13,sessionid)
                    except:
                        module1=None
                        
                        module2=None
                        module3=None
                        module4=None
                        module5=None
                        module6=None
                        module7=None
                        module8=None
                        module9=None
                        module10=None
                        module11=None
                        module12=None
                        module13=None
              
                      
            md1m7 = Module1MajorCrops.objects.filter(module1_id=module1.id)
            for o in md1m7:
                cc=o.crops#get_object_or_404(Crops, id=o.crops)
                m1m7=m1m7+cc+','#+cc.crop+','
            m1m7=m1m7[:-1]    
            md1m8 = Module1MajorImports.objects.filter(module1_id=module1.id)
            for o in md1m8:
                cc=o.crops#get_object_or_404(Crops, id=o.crops)
                m1m8=m1m8+cc+','#+cc.crop+','
            m1m8=m1m8[:-1]    
            md1m11 = Module1MajorExports.objects.filter(module1_id=module1.id)
            for o in md1m11:
                cc=o.crops#get_object_or_404(Crops, id=o.crops)
                m1m11=m1m11+cc+','#+cc.crop+','
            m1m11=m1m11[:-1]    
            md1m20 = Module1MajorPartenerImport.objects.filter(module1_id=module1.id)
            for o in md1m20:
                m1m20=m1m20+o.country+','
            m1m20=m1m20[:-1]    
            md1m21 = Module1MajorPartenerImport.objects.filter(module1_id=module1.id)
            for o in md1m21:
                m1m21=m1m21+o.country+','    
            m1m21=m1m21[:-1]    
            md1m26 = Module1Aid.objects.filter(module1_id=module1.id)
            for o in md1m26:
                m1m26=m1m26+o.titleprj+',' 
            m1m26=m1m26[:-1]    
            md1m22 = module1.m_22
            for o in md1m22.all():
                m1m22=str(o)+',' 
            m1m22=m1m22[:-1]  
            md1m23 = module1.m_23
            for o in md1m22.all():
                m1m23=str(o)+',' 
            m1m23=m1m23[:-1]  
            
         
            VAL_IMP = (
                (0, _("--- Please select ---")),
                (1, _("Not known")),
                (2, _("0 to $100,000")),
                (3, _("$100,000 to $500,000")),
                (4, _("$500,000 to $1M")),
                (5, _("$1M to $10M")),
                (6, _("$10M to $25M")),
                (7, _("$25M to $50M")),
                (8, _("$50M to $100M")),
                (9, _("Greater than $100M ")),
            )

            VAL_EXP = (
                (0, _("--- Please select ---")),
                (1, _("Unknown")),
                (2, _("0 to $100,000")),
                (3, _("$100,000 to $500,000")),
                (4, _("$500,000 to $1M")),
                (5, _("$1M to $10M")),
                (6, _("$10M to $50M")),
                (7, _("$25M to $50M")),
                (8, _("greater than $50M")),
            )

      
            VAL_PERCENT = (
                (0, _("--- Please select ---")),
                (1, _("0")),
                (2, _("10")),
                (3, _("20")),
                (4, _("30")),
                (5, _("40")),
                (6, _("50")),
                (7, _("60")),
                (8, _("70")),
                (9, _("80")),
                (10, _("90")),
                (11, _("100")),
            )

            NUM_BILATERAL = (
                (0, _("--- Please select ---")),
                (1, _("1-3")),
                (2, _("4-6")),
                (3, _("7-10")),
                (4, _("greater than 10")),
            )  
            #  ---- Cover Letter ----
            document.add_heading(_("Phytosanitary Capacity Evaluation (PCE) - REPORT"), 0)
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            p= document.add_paragraph("")
            p.add_run(_("1. Introduction:")).bold = True
           
            REGIONS_LABELS = (
            (1, _("Africa")),
            (2, _("Asia")),
            (3, _("Europe")),
            (4, _("Latin America and Caribbean")),
            (5, _("Near East")),
            (6, _("North America")),
            (7, _("South West Pacific")),
            )
            document.add_paragraph()
            
            countryname=str(user.get_profile().country)
            regionname=ugettext(dict(REGIONS_LABELS)[module1.region])
            #regionname=str(ugettext(dict(REGIONS_LABELS)[module1.region]))#.encode('utf-8')
            
            str_m_3=module1.m_3
            str_m_4=module1.m_4
            str_m_5=module1.m_5
            str_m_6=module1.m_6
            str_m1m7=m1m7
            str_m1m8=m1m8
            str_m_9=ugettext(dict(VAL_IMP)[module1.m_9])
            str_m1m20=m1m20
            str_m1m11=m1m11
            str_m_12=ugettext(dict(VAL_EXP)[module1.m_12])
            str_m1m21=m1m21
            str_m_13=ugettext(dict(VAL_PERCENT)[module1.m_13])
            str_m_14=module1.m_14
            str_m_15=module1.m_15
            str_m_16=ugettext(dict(VAL_PERCENT)[module1.m_16])
            str_m_17=ugettext(dict(VAL_PERCENT)[module1.m_17])
            str_m_18=ugettext(dict(VAL_PERCENT)[module1.m_18])
            str_m_19=ugettext(dict(VAL_PERCENT)[module1.m_19])
           # #print(m1m22)
            ##print(ugettext(str(m1m22)))
            str_m1m23=str(m1m23)
            str_m1m22=str(m1m22)
            
            str_m_24=ugettext(dict(NUM_BILATERAL)[module1.m_24])
            str_m_25=ugettext(dict(NUM_BILATERAL)[module1.m_25])
            str_m1m26=m1m26
# str_m_3=module1.m_3.encode('utf-8')
#            str_m_4=module1.m_4.encode('utf-8')
#            str_m_5=module1.m_5.encode('utf-8')
#            str_m_6=module1.m_6.encode('utf-8')
#            str_m1m7=m1m7.encode('utf-8')
#            str_m1m8=m1m8.encode('utf-8')
#            str_m_9=ugettext(dict(VAL_IMP)[module1.m_9])
#            str_m1m20=m1m20.encode('utf-8')
#            str_m1m11=m1m11.encode('utf-8')
#            str_m_12=ugettext(dict(VAL_EXP)[module1.m_12])
#            str_m1m21=m1m21.encode('utf-8')
#            str_m_13=ugettext(dict(VAL_PERCENT)[module1.m_13])
#            str_m_14=module1.m_14.encode('utf-8')
#            str_m_15=module1.m_15.encode('utf-8')
#            str_m_16=ugettext(dict(VAL_PERCENT)[module1.m_16])
#            str_m_17=ugettext(dict(VAL_PERCENT)[module1.m_17])
#            str_m_18=ugettext(dict(VAL_PERCENT)[module1.m_18])
#            str_m_19=ugettext(dict(VAL_PERCENT)[module1.m_19])
#           # #print(m1m22)
#            ##print(ugettext(str(m1m22)))
#            str_m1m23=str(m1m23)
#            str_m1m22=str(m1m22)
#            
#            str_m_24=ugettext(dict(NUM_BILATERAL)[module1.m_24])
#            str_m_25=ugettext(dict(NUM_BILATERAL)[module1.m_25])
#            str_m1m26=m1m26.encode('utf-8')


            str_1=_(" is a country situated in ")+""
            str_2 =_(".  It has population of about ")+""
            str_3 =_(". Total land area is ")+""
            str_4 =_(" sq.km. with a total arable land area of ")+""
            str_5 =_(" sq.km. Total natural vegetation occupies ")+""
            str_6 =_(" sq.km.\n\nThe major crops grown in the country are: ")+""
            str_7 =_(".\n\nTen major imports of plant and plant products are: ")+""
            str_8 =_(". Total value of imports of plant and plant products (including forestry products) amounted to ")+""
            #str_9 =_(" in ")+""
            str_9 =""
            str_10 =_("'s  major trading partners in plants and plant products imports are ")+""
            str_11 =_(".\n\nTen major exports of plant and plant products are: ")+""
            str_12 =_(". Total value of exports of plant and plant products (including forestry products) amounted to ")+""
            #str_13 =_(" in ")+""
            str_13 =""
            str_14 =_("'s major trading partners in plants and plant products exports are ")+""
            str_15 =_(". ")+""
            str_16 =_(" % of total exports (includes Forestry) are re-export consignments.\n\nThe Gross National Income (GNI) per capita is estimated at ")+""
            str_17 =_(" US $; latest GDP in US $ (World Bank) is ")+""
            str_18 =_(". Percentage contribution of agriculture (including forestry) to GDP is about ")+""
            str_19 =_(" %; with about ")+""
            str_20 =_(" % for plants and plant products (including forestry). The agricultural labour force (including forestry) as a percentage of total labour force is ")+""
            str_21 =_(" %. ")+""
            str_22 =_(" % of the agricultural labour force is directly employed in the production of plant and plant products (including forestry).\n\n")+""
            str_23 =_("  has membership of or is signatory to, the following organizations/ conventions: ")+""
            str_24 =_(". It is a member of the following Regional economic integration/ co-operation organizations: ")+""
            str_25 =_(". Currently there are ")+""
            str_26 =_(" bilateral phytosanitary arrangements in operation and ")+""
            str_27 =_("  more negotiations are in progress.\n\nDuring the last few years, major aid programs that have significantly contributed to phytosanitary capacity development or strengthening in the country include : ")+""


           # #print('------------------------------------------------------------aaa----------')
           # #print(countryname+str_1+regionname+str_2+str_m_3+str_3+str_m_4+str_4+str_m_5+str_5+str_m_6+str_6+str_m1m7+str_7+str_m1m8+str_8+str_m_9+str_9+countryname+str_10+str_m1m20+str_11+str_m1m11+str_12+str_m_12+str_13+countryname+str_14+str_m1m21+str_15+str_m_13+str_16+str_m_14+str_17+str_m_15+str_18+str_m_16+str_19+str_m_17+str_20+str_m_18+str_21+str_m_19+str_22+countryname+str_23+str_m1m22+str_24+str_m1m23+str_25+str_m_24+str_26+str_m_25+str_27+str_m1m26)

             
            #p=document.add_paragraph(""+countryname+str(_(" is a country situated in "))+regionname+_(".  It has population of about ")+str_m_3+_(". Total land area is ")+str_m_4+_(" sq.km. with a total arable land area of ")+str_m_5+_(" sq.km. Total natural vegetation occupies ")+str_m_6+_(" sq.km.\n\nThe major crops grown in the country are: ")+str_m1m7+_(".\n\nTen major imports of plant and plant products are: ")+str_m1m8+_(". Total value of imports of plant and plant products (including forestry products) amounted to ")+str_m_9+_(" in ")+countryname+_("'s  major trading partners in plants and plant products imports are ")+str_m1m20+_(".\n\nTen major exports of plant and plant products are: ")+str_m1m11+_(". Total value of exports of plant and plant products (including forestry products) amounted to ")+str_m_12+_(" in ")+countryname+_("''s major trading partners in plants and plant products exports are ")+str_m1m21+_(". ")+str_m_13+_(" % of total exports (includes Forestry) are re-export consignments.\n\nThe Gross National Income (GNI) per capita is estimated at ")+str_m_14+_(" US $; latest GDP in US $ (World Bank) is ")+str_m_15+_(". Percentage contribution of agriculture (including forestry) to GDP is about ")+str_m_16+_(" %; with about ")+str_m_17+_(" % for plants and plant products (including forestry). The agricultural labour force (including forestry) as a percentage of total labour force is ")+str_m_18+_(" %. ")+str_m_19+_(" % of the agricultural labour force is directly employed in the production of plant and plant products (including forestry).\n\n")+countryname+_("  has membership of or is signatory to, the following organizations/ conventions: ")+str_m1m22+_(". It is a member of the following Regional economic integration/ co-operation organizations: ")+str_m1m23+_(". Currently there are ")+str_m_24+_(" bilateral phytosanitary arrangements in operation and ")+str_m_25+_("  more negotiations are in progress.\n\nDuring the last few years, major aid programs that have significantly contributed to phytosanitary capacity development or strengthening in the country include : ")+str_m1m26)
            p=document.add_paragraph(countryname+str_1+regionname+str_2+str_m_3+str_3+str_m_4+str_4+str_m_5+str_5+str_m_6+str_6+str_m1m7+str_7+str_m1m8+str_8+str_m_9+str_9+countryname+str_10+str_m1m20+str_11+str_m1m11+str_12+str_m_12+str_13+countryname+str_14+str_m1m21+str_15+str_m_13+str_16+str_m_14+str_17+str_m_15+str_18+str_m_16+str_19+str_m_17+str_20+str_m_18+str_21+str_m_19+str_22+countryname+str_23+str_m1m22+str_24+str_m1m23+str_25+str_m_24+str_26+str_m_25+str_27+str_m1m26)
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
         
            document.add_paragraph()
            document.add_paragraph()
#
#
            p=document.add_paragraph("")
            p.add_run(_("2. Summary of Findings/Recommendations:")).bold = True
            
            p=document.add_paragraph("")
            p.add_run(_("2.1. Key issues or points identified:.")).bold = True
            table = document.add_table(rows=1, cols=1)
            table.style = document.styles['Table Grid']
            
            document.add_paragraph()
            p=document.add_paragraph("")
            p.add_run(_("2.2. Identified Weakness:")).bold = True
            ##print(items)
            if '2' in items or '3' in items:
                p=document.add_paragraph("")
                p.add_run(_("Legislation and policy:")).bold = True
                document.add_paragraph()
                table = document.add_table(rows=1, cols=2)
                table.style = document.styles['Table Grid']
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = _("Module")
                hdr_cells[1].text = _("Weakness/Priority (highest priority 1)")
                
                for h in range(2,4):
                    i=str(h)

                    if i == '2' and '2' in items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Module")+' 2 -' + ugettext(m_names[2])
                        row_cells[1].text =  '1. '+w2.w1+'\n'+'2. '+w2.w2+'\n'+'3. '+w2.w3+'\n'+'4. '+w2.w4+'\n'+'5. '+w2.w5
                    if i == '3' and '3' in items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Module")+' 3 -' + ugettext(m_names[3])
                        row_cells[1].text =  '1. '+w3.w1+'\n'+'2. '+w3.w2+'\n'+'3. '+w3.w3+'\n'+'4. '+w3.w4+'\n'+'5. '+w3.w5
#
            if '4' in items or '5' in items or '6' in items:
                document.add_paragraph()
                p=document.add_paragraph("")
                p.add_run(_("NPPO operations:")).bold = True
                document.add_paragraph()
                table = document.add_table(rows=1, cols=2)
                table.style = document.styles['Table Grid']
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = _("Module")
                hdr_cells[1].text = _("Weakness/Priority (highest priority 1)")
                
                for h in range(4,7):
                    i=str(h)
                    if  i == '4' and '4' in items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Module")+' 4 -' +ugettext(m_names[4])
                        row_cells[1].text =  '1. '+w4.w1+'\n'+'2. '+w4.w2+'\n'+'3. '+w4.w3+'\n'+'4. '+w4.w4+'\n'+'5. '+w4.w5
                    if i == '5'  and '5' in items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Module")+' 5 -' + ugettext(m_names[5])
                        row_cells[1].text =  '1. '+w5.w1+'\n'+'2. '+w5.w2+'\n'+'3. '+w5.w3+'\n'+'4. '+w5.w4+'\n'+'5. '+w5.w5
                    if  i == '6' and '6' in items :
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Module")+' 6 -' + ugettext(m_names[6])
                        row_cells[1].text =  '1. '+w6.w1+'\n'+'2. '+w6.w2+'\n'+'3. '+w6.w3+'\n'+'4. '+w6.w4+'\n'+'5. '+w6.w5

           
            if '7' in items or '8' in items or '9' in items or '10' in items or '11' in items or '12' in items or '13' in items:
                document.add_paragraph()
                p=document.add_paragraph("")
                p.add_run(_("Technical:")).bold = True
                document.add_paragraph()
                table = document.add_table(rows=1, cols=2)
                table.style = document.styles['Table Grid']
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = _("Module")
                hdr_cells[1].text = _("Weakness/Priority (highest priority 1)")
                
                for h in range(7,14):
                    i=str(h)
                    if i == '7' and '7' in items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Module")+' 7 -' +ugettext(m_names[7])
                        row_cells[1].text =  '1. '+w7.w1+'\n'+'2. '+w7.w2+'\n'+'3. '+w7.w3+'\n'+'4. '+w7.w4+'\n'+'5. '+w7.w5
                    if i == '8' and '8' in items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Module")+' 8 -' +ugettext(m_names[8])
                        row_cells[1].text =  '1. '+w8.w1+'\n'+'2. '+w8.w2+'\n'+'3. '+w8.w3+'\n'+'4. '+w8.w4+'\n'+'5. '+w8.w5
                    if i == '9' and '9' in items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Module")+' 9 -' + ugettext(m_names[9])
                        row_cells[1].text =  '1. '+w9.w1+'\n'+'2. '+w9.w2+'\n'+'3. '+w9.w3+'\n'+'4. '+w9.w4+'\n'+'5. '+w9.w5
                    if i == '10' and '10' in items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Module")+' 10 -' +ugettext(m_names[10])
                        row_cells[1].text =  '1. '+w10.w1+'\n'+'2. '+w10.w2+'\n'+'3. '+w10.w3+'\n'+'4. '+w10.w4+'\n'+'5. '+w10.w5
                    if i == '11' and '11' in items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Module")+' 11 -' +ugettext(m_names[11])
                        row_cells[1].text =  '1. '+w11.w1+'\n'+'2. '+w11.w2+'\n'+'3. '+w11.w3+'\n'+'4. '+w11.w4+'\n'+'5. '+w11.w5
                    if i == '12' and '12' in items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Module")+' 12 -' +ugettext(m_names[12])
                        row_cells[1].text =  '1. '+w12.w1+'\n'+'2. '+w12.w2+'\n'+'3. '+w12.w3+'\n'+'4. '+w12.w4+'\n'+'5. '+w12.w5
                    if i == '13' and '13' in items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Module")+' 13 -' +ugettext(m_names[13])
                        row_cells[1].text =  '1. '+w13.w1+'\n'+'2. '+w13.w2+'\n'+'3. '+w13.w3+'\n'+'4. '+w13.w4+'\n'+'5. '+w13.w5


            document.add_paragraph()
            p=document.add_paragraph("")
            p.add_run(_("3. Methodology:")).bold = True
            table = document.add_table(rows=1, cols=1)
            table.style = document.styles['Table Grid']
            
            document.add_paragraph()
            p=document.add_paragraph(_("Information on the stakeholders involved in the PCE application is available in the Annex 2. "))
            
            document.add_paragraph()
            p=document.add_paragraph("")
            p.add_run(_("4. Findings:")).bold = True
            document.add_paragraph()
            p=document.add_paragraph("")
            p.add_run(_("4.1 Introduction:")).bold = True
            document.add_paragraph()
            document.add_paragraph(_("The 'New Revised Text' (IPPC, 1997) sets out clear obligations for the NPPO to address plant protection issues and/or broaden its scope of operations and establish systems to address IPPC requirements and responsibilities.\n\nThe objectives of an NPPO within the context of national development plans may be translated into three broad areas of responsibility:\n\t\t- to protect plant resources (including cultivated, wild and aquatic plants)  through implementation of appropriate phytosanitary measures\n\t\t- to support national food security and a healthy environment through effective pest exclusion procedures\n\t\t- to facilitate market access and safe international trade in agricultural commodities by establishing effective phytosanitary certification systems and procedures.\n\nTo fulfil all of these objectives an NPPO needs sustainable financing, planning for long-term staffing arrangements, having contingency plans in place for changes in political contexts and planning for natural disasters, among other areas, to ensure the organization remains sustainable and adaptable over the long term. At the same time, a well-organized, fully functional an NPPO should have appropriate national, regional and international networks."))
            document.add_paragraph()
            p=document.add_paragraph("")
            p.add_run(_("4.2 Causes hindering improvement of the NPPO")).bold = True
         
            document.add_paragraph(_("There are several causes which are hindering the normal activity and improvement of the NPPO which can be separated into primary and secondary ones.\n\nThe principal primary causes are:\n\n\n\nThe secondary ones are : \n\n\n\n"))
            document.add_paragraph()
            p=document.add_paragraph("")
            p.add_run(_("4.3 Impacts:")).bold = True
         
            document.add_paragraph(_("The factors mentioned above generate some undesirable consequences for the NPPO, which are also classified into primary and secondary ones.\n\nPrimary consequences are :\n\n\n\n\nSecondary consequences are:\n\n\n\n\n"))
            document.add_paragraph()
            document.add_paragraph()
            p=document.add_paragraph("")
            p.add_run(_("4.4 SWOT Analysis :")).bold = True
         
            document.add_paragraph(_("SWOT (Strength, Weakness, Opportunities and Threats) -  under this section relevant internal strengths, as well as external opportunities and threats can be found. Opportunities and strength can enable the recovery of each of the weakness."))
            document.add_paragraph()
            
            for i in range(1,14) :
                
                if  i>1 and  str(i) in items:
                    table = document.add_table(rows=1, cols=2)
                    table.style = document.styles['Table Grid']
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = _("Module ")+str(i)+' - '+ugettext(m_names[int(i)])
                    
                    swotanalysis = SwotAnalysis.objects.filter(session_id=sessionid,module=int(i)).count()
                    if swotanalysis>0:
                        sa = get_object_or_404(SwotAnalysis, session_id=sessionid,module=int(i))
                        sa1 = SwotAnalysis1.objects.filter(swotanalysis_id=sa.id)
                        sa2 = SwotAnalysis2.objects.filter(swotanalysis_id=sa.id)
                        sa3 = SwotAnalysis3.objects.filter(swotanalysis_id=sa.id)
                        sa4 = SwotAnalysis4.objects.filter(swotanalysis_id=sa.id)
                        sa5 = SwotAnalysis5.objects.filter(swotanalysis_id=sa.id)
                        sa1_all_s = ''
                        sa1_all_o = ''
                        sa1_all_t = ''
                        v=0
                        for o in sa1:
                            oo=''
                            if v>0:
                              oo=str(v)+'.'
                            if o.strengths !='':  
                                sa1_all_s=sa1_all_s+'1.'+oo+' '+o.strengths 
                            if o.opportunities !='':  
                                sa1_all_o=sa1_all_o+'1.'+oo+' '+o.opportunities+',' 
                            if o.threats !='':  
                                sa1_all_t=sa1_all_t+'1.'+oo+' '+o.threats+',' 
                            v=v+1
                        
                        
                        sa2_all_s = ''
                        sa2_all_o = ''
                        sa2_all_t = ''
                        v=0
                        for o in sa2:
                            oo=''
                            if v>0:
                              oo=str(v)+'.'
                            if o.strengths !='':  
                                sa2_all_s=sa2_all_s+'2.'+oo+' '+o.strengths 
                            if o.opportunities !='':  
                                sa2_all_o=sa2_all_o+'2.'+oo+' '+o.opportunities+',' 
                            if o.threats !='':  
                                sa2_all_t=sa2_all_t+'2.'+oo+' '+o.threats+',' 
                            v=v+1
        
                        sa3_all_s = ''
                        sa3_all_o = ''
                        sa3_all_t = ''
                        v=0
                        for o in sa3:
                            oo=''
                            if v>0:
                              oo=str(v)+'.'
                            if o.strengths !='':  
                                sa3_all_s=sa3_all_s+'3.'+oo+' '+o.strengths 
                            if o.opportunities !='':  
                                sa3_all_o=sa3_all_o+'3.'+oo+' '+o.opportunities+',' 
                            if o.threats !='':  
                                sa3_all_t=sa3_all_t+'3.'+oo+' '+o.threats+',' 
                            v=v+1
        
                        sa4_all_s = ''
                        sa4_all_o = ''
                        sa4_all_t = ''
                        v=0
                        for o in sa4:
                            oo=''
                            if v>0:
                              oo=str(v)+'.'
                            if o.strengths !='':  
                                sa4_all_s=sa4_all_s+'4.'+oo+' '+o.strengths 
                            if o.opportunities !='':  
                                sa4_all_o=sa4_all_o+'4.'+oo+' '+o.opportunities+',' 
                            if o.threats !='':  
                                sa4_all_t=sa4_all_t+'4.'+oo+' '+o.threats+',' 
                            v=v+1
        
                        sa5_all_s = ''
                        sa5_all_o = ''
                        sa5_all_t = ''
                        v=0
                        for o in sa5:
                            oo=''
                            if v>0:
                              oo='\n'+str(v)+'.'
                            if o.strengths !='':  
                                sa5_all_s=sa5_all_s+'5.'+oo+' '+o.strengths 
                            if o.opportunities !='':  
                                sa5_all_o=sa5_all_o+'5.'+oo+' '+o.opportunities+',' 
                            if o.threats !='':  
                                sa5_all_t=sa5_all_t+'5.'+oo+' '+o.threats+',' 
                            v=v+1
        
                        if int(i) == 2:
                            w=w2
                        if int(i) == 3:
                            w=w3
                        if int(i) == 4:
                            w=w4
                        if int(i) == 5:
                            w=w5
                        if int(i) == 6:
                            w=w6
                        if int(i) == 7:
                            w=w7
                        if int(i) == 8:
                            w=w8
                        if int(i) == 9:
                            w=w9
                        if int(i) == 10:
                            w=w10
                        if int(i) == 11:
                            w=w11
                        if int(i) == 12:
                            w=w12
                        if int(i) == 13:
                            w=w13
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Strength")+'\n\n'+sa1_all_s+'\n'+sa2_all_s+'\n'+sa3_all_s+'\n'+sa4_all_s+'\n'+sa5_all_s
                        row_cells[1].text = _("Weakness")+'\n\n1. '+w.w1+'\n'+'2. '+w.w2+'\n'+'3. '+w.w3+'\n'+'4. '+w.w4+'\n'+'5. '+w.w5
                        row_cells = table.add_row().cells
                        row_cells[0].text = _("Opportunities")+'\n\n'+sa1_all_o+'\n'+sa2_all_o+'\n'+sa3_all_o+'\n'+sa4_all_o+'\n'+sa5_all_o
                        row_cells[1].text = _("Threats")+'\n\n'+sa1_all_t+'\n'+sa2_all_t+'\n'+sa3_all_t+'\n'+sa4_all_t+'\n'+sa5_all_t  
#
          
            document.add_page_break()
            document.add_paragraph()
            p=document.add_paragraph("")
            p.add_run(_("5. Recommendations:")).bold = True
            document.add_paragraph()
            p=document.add_paragraph("")
            p.add_run(_("5.1 Programme overview :")).bold = True
            document.add_paragraph(_("The programme overview table constitutes the development strategy for the NPPO during the .... year period 20...- 20...\nIt summarizes the following criteria:\n- Development focus area (i.e. the different modules)\n- Overall objective/Goal of the whole plan\n- Specific objective/purpose related to the overall objectives/goal for each module\n- Outputs expected to result from each activity\nAn estimate of the funds required to realize the activities of the programme is ....\n\nA log frame matrix and indicative work plan for each module can be seen in Annex 1"))
            p=document.add_paragraph("")
            p.add_run(_("Programme overview :")).bold = True
          
            table = document.add_table(rows=1, cols=5)
            table.style = document.styles['Table Grid']
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text =_("Development Focus Area")
            hdr_cells[1].text = _("Overall objectives/Goal")
            hdr_cells[2].text = _("Specific objectives/Purpose")
            hdr_cells[3].text = _("Outputs")
            hdr_cells[4].text = _("Indicative Cost USD")
            for i in range(1,14) :
                if  i>1 and  str(i) in items:
                    row_cells = table.add_row().cells
                    row_cells[0].text=_("Module ")+str(i)+' - '+ugettext(m_names[int(i)])
                    logicalframework = LogicalFramework.objects.filter(session_id=sessionid,module=int(i)).count()
                    count1=0
                    if logicalframework>0:
                        lf = get_object_or_404(LogicalFramework, session_id=sessionid,module=int(i))
                        if LogicalFrameworkAct1.objects.filter(logicalframework_id=lf.id).count()>0:
                            lf1 = LogicalFrameworkAct1.objects.filter(logicalframework_id=lf.id)[0]
                            if lf1.cost!='' and lf1.cost.isdigit():
                                count1=count1+int(lf1.cost)
                        if LogicalFrameworkAct2.objects.filter(logicalframework_id=lf.id).count()>0:
                            lf2 = LogicalFrameworkAct2.objects.filter(logicalframework_id=lf.id)[0]
                            if lf2.cost!='' and lf2.cost.isdigit():
                                count1=count1+int(lf2.cost)
                        if LogicalFrameworkAct3.objects.filter(logicalframework_id=lf.id).count()>0:
                            lf3 = LogicalFrameworkAct3.objects.filter(logicalframework_id=lf.id)[0]
                            if lf3.cost!='' and lf3.cost.isdigit():
                                count1=count1+int(lf3.cost)
                        if LogicalFrameworkAct4.objects.filter(logicalframework_id=lf.id).count()>0:
                            lf4 = LogicalFrameworkAct4.objects.filter(logicalframework_id=lf.id)[0]
                            if lf4.cost!='' and lf4.cost.isdigit():
                                count1=count1+int(lf4.cost)
                        if LogicalFrameworkAct5.objects.filter(logicalframework_id=lf.id).count()>0:
                            lf5 = LogicalFrameworkAct5.objects.filter(logicalframework_id=lf.id)[0]
                            if lf5.cost!='' and lf5.cost.isdigit():
                                count1=count1+int(lf5.cost)
                            
                        
                        row_cells[1].text = lf.overobjective
                        row_cells[2].text = lf.objective
                        row_cells[3].text = lf.output1+'\n'+lf.output2+'\n'+lf.output3+'\n'+lf.output4+'\n'+lf.output5
                        row_cells[4].text = str(count1)
                        
            document.add_paragraph()
            p=document.add_paragraph("")
            p.add_run(_("5.2 Priority Actions:")).bold = True
            document.add_paragraph()
            table = document.add_table(rows=1, cols=2)
            table.style = document.styles['Table Grid']
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text =_("Module")
            hdr_cells[1].text = _("Priority Activities")
            
            for i in range(1,14) :
                if  i>1 and  str(i) in items:
                    row_cells = table.add_row().cells
                    row_cells[0].text=_("Module ")+str(i)+' - '+ugettext(m_names[int(i)])
                    row_cells[1].text=''
           
            document.add_paragraph()
            p=document.add_paragraph("")
            p.add_run(_("Annex 1 - Strategic framework")).bold = True
            document.add_paragraph()
            document.add_paragraph(_("Strategic framework consists of Logical frameworks and associated work plans for each applied module "))
            document.add_paragraph()
            for i in range(1,14) :
                if  i>1 and  str(i) in items:
                    document.add_paragraph("")  
                    p=document.add_paragraph("")
                    p.add_run(m_names[int(i)]).bold = True
         
                    table = document.add_table(rows=1, cols=4)
                    table.style = document.styles['Table Grid']
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = _("Overall Objective")
                    hdr_cells[1].text = _("Key indicators")
                    hdr_cells[2].text = _("Means of Verification")
                    hdr_cells[3].text = _("Assumptions/Risk")
                    logicalframework = LogicalFramework.objects.filter(session_id=sessionid,module=int(i)).count()
                    if logicalframework>0:
                        lf = get_object_or_404(LogicalFramework, session_id=sessionid,module=int(i))
                        row_cells = table.add_row().cells
                        row_cells[0].text=lf.overobjective
                        row_cells[1].text=lf.keyindicator0
                        row_cells[2].text=lf.verification0
                        row_cells[3].text=lf.assumptions0
                    document.add_paragraph("")   
                   
                    table = document.add_table(rows=1, cols=4)
                    table.style = document.styles['Table Grid']
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = _("Immediate Objective (purpose)")
                    hdr_cells[1].text = _("Key indicators")
                    hdr_cells[2].text = _("Means of Verification")
                    hdr_cells[3].text = _("Assumptions/Risk")
                    logicalframework = LogicalFramework.objects.filter(session_id=sessionid,module=int(i)).count()
                    if logicalframework>0:
                        lf = get_object_or_404(LogicalFramework, session_id=sessionid,module=int(i))
                        row_cells = table.add_row().cells
                        row_cells[0].text=lf.objective
                        row_cells[1].text=lf.keyindicator
                        row_cells[2].text=lf.verification
                        row_cells[3].text=lf.assumptions
                    
                    document.add_paragraph("")   
                    
                    table = document.add_table(rows=1, cols=4)
                    table.style = document.styles['Table Grid']
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text =_("Outputs")
                    hdr_cells[1].text = _("Key indicators")
                    hdr_cells[2].text = _("Means of Verification")
                    hdr_cells[3].text = _("Assumptions/Risk")
                    logicalframework = LogicalFramework.objects.filter(session_id=sessionid,module=int(i)).count()
                    if logicalframework>0:
                        lf = get_object_or_404(LogicalFramework, session_id=sessionid,module=int(i))
                        row_cells = table.add_row().cells
                        row_cells[0].text=lf.output1
                        row_cells[1].text=lf.keyindicator1
                        row_cells[2].text=lf.verification1
                        row_cells[3].text=lf.assumptions1
                  
                        row_cells = table.add_row().cells
                        row_cells[0].text=lf.output2
                        row_cells[1].text=lf.keyindicator2
                        row_cells[2].text=lf.verification2
                        row_cells[3].text=lf.assumptions2
                        row_cells = table.add_row().cells
                        row_cells[0].text=lf.output3
                        row_cells[1].text=lf.keyindicator3
                        row_cells[2].text=lf.verification3
                        row_cells[3].text=lf.assumptions3
                        row_cells = table.add_row().cells
                        row_cells[0].text=lf.output4
                        row_cells[1].text=lf.keyindicator4
                        row_cells[2].text=lf.verification4
                        row_cells[3].text=lf.assumptions4
                        row_cells = table.add_row().cells
                        row_cells[0].text=lf.output5
                        row_cells[1].text=lf.keyindicator5
                        row_cells[2].text=lf.verification5
                        row_cells[3].text=lf.assumptions5
                    document.add_paragraph("") 
                    document.add_paragraph("Indicative Work Plan") 
                    table = document.add_table(rows=1, cols=4)
                    table.style = document.styles['Table Grid']
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text =_("Activities") 
                    hdr_cells[1].text = _("Estimated costs") 
                    hdr_cells[2].text = _("Responsible Person") 
                    hdr_cells[3].text = _("deadline") 
                    #hdr_cells[4].text = 'Target'
                    #hdr_cells[5].text = 'Source'
                    #hdr_cells[6].text = 'External factors'
                    logicalframework = LogicalFramework.objects.filter(session_id=sessionid,module=int(i)).count()
                    if logicalframework>0:
                        lf = get_object_or_404(LogicalFramework, session_id=sessionid,module=int(i))
                        if LogicalFrameworkAct1.objects.filter(logicalframework_id=lf.id).count()>0:
#                            lf1 = LogicalFrameworkAct1.objects.filter(logicalframework_id=lf.id)[0]
#                            row_cells = table.add_row().cells
#                            row_cells[0].text=lf1.activity1
#                            row_cells[1].text=lf1.cost
#                            row_cells[2].text=lf1.responsible
#                            row_cells[3].text=lf1.deadline
                            lf1 = LogicalFrameworkAct1.objects.filter(logicalframework_id=lf.id)
                            counter1 =1
                            for lf1_i in lf1:
                                row_cells = table.add_row().cells
                                row_cells[0].text='1.'+str(counter1)+')'+lf1_i.activity1
                                row_cells[1].text=lf1_i.cost
                                row_cells[2].text=lf1_i.responsible
                                row_cells[3].text=lf1_i.deadline
                                counter1=counter1+1
                               # row_cells = table.add_row().cells
                            #row_cells = table.add_row().cells
                            lf2 = LogicalFrameworkAct2.objects.filter(logicalframework_id=lf.id)
                            counter2 =1
                            for lf2_i in lf2:
                                row_cells = table.add_row().cells
                                row_cells[0].text='2.'+str(counter2)+')'+lf2_i.activity2
                                row_cells[1].text=lf2_i.cost
                                row_cells[2].text=lf2_i.responsible
                                row_cells[3].text=lf2_i.deadline
                                counter2=counter2+1   
                            lf3 = LogicalFrameworkAct3.objects.filter(logicalframework_id=lf.id)
                            counter3 =1
                            for lf3_i in lf3:
                                row_cells = table.add_row().cells
                                row_cells[0].text='3.'+str(counter3)+')'+lf3_i.activity3
                                row_cells[1].text=lf3_i.cost
                                row_cells[2].text=lf3_i.responsible
                                row_cells[3].text=lf3_i.deadline
                                counter3=counter3+1       
                            lf4 = LogicalFrameworkAct4.objects.filter(logicalframework_id=lf.id)
                            counter4 =1
                            for lf4_i in lf4:
                                row_cells = table.add_row().cells
                                row_cells[0].text='4.'+str(counter4)+')'+lf4_i.activity4
                                row_cells[1].text=lf4_i.cost
                                row_cells[2].text=lf4_i.responsible
                                row_cells[3].text=lf4_i.deadline
                                counter4=counter4+1  
                            lf5 = LogicalFrameworkAct5.objects.filter(logicalframework_id=lf.id)
                            counter5 =1
                            for lf5_i in lf5:
                                row_cells = table.add_row().cells
                                row_cells[0].text='5.'+str(counter5)+')'+lf5_i.activity5
                                row_cells[1].text=lf5_i.cost
                                row_cells[2].text=lf5_i.responsible
                                row_cells[3].text=lf5_i.deadline
                                counter5=counter5+1  
                            row_cells = table.add_row().cells
                    
                        
                        
                        
                       
                        
            document.add_paragraph("")
            document.add_paragraph("")
          
            p=document.add_paragraph("")
            p.add_run(_("Annex 2:")).bold = True
            p=document.add_paragraph("")
            p.add_run(_("List of stakeholders:")).bold = True
         
       
            
            table = document.add_table(rows=1, cols=3)
            table.style = document.styles['Table Grid']
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = _("Name")
            hdr_cells[1].text = _("Organization/Division")
            hdr_cells[2].text = _("Email")
            
            stakeholders = Stakeholders.objects.filter(session_id=sessionid)
            ##print(stakeholders) 
            ss=[]
            for s in stakeholders:
                stakes=StakeholdersFields.objects.filter(stakeholder_id= s.id)
                for si in stakes:
                    
                    if si.email not in ss:
                        row_cells = table.add_row().cells
                        ss.append(si.email)
                        row_cells[0].text = si.firstname+' '+ si.lastname
                        row_cells[1].text = str(si.organisation)
                        row_cells[2].text = str(si.email)
    
            document.add_paragraph("")
            document.add_paragraph("")
          
            VAL_AV = (
                (0, (" ")),
                (1, ("0-5")),
                (2, ("5-10")),
                (3, ("10-20")),
                (4, (">20")),
             )
            BOOL_CHOICESM_M = (
                (0, (" ")),
                (1, ("Yes")),
                (2, ("No")),
            )  
         
            p=document.add_paragraph("")
            p.add_run(_("Annex 3:")).bold = True
            p=document.add_paragraph("")
            p.add_run(_("Human Resources")).bold = True
            p=document.add_paragraph("")
            
            if '7' in items:
                p.add_run(_("Module")+" 7 - "+ugettext(m_names[7])).bold = True
                p=document.add_paragraph("")
                p.add_run(_("Pest diagnostic laboratory current human resources")).bold = True

                table = document.add_table(rows=1, cols=7)
                table.style = document.styles['Table Grid']


                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = ''
                hdr_cells[1].text = ''
                hdr_cells[2].text = ''
                hdr_cells[3].text = ''
                hdr_cells[4].text = ''
                hdr_cells[5].text = ''
                hdr_cells[6].text = ''

                hdr_cells[1].merge(hdr_cells[4])
                hdr_cells[5].merge(hdr_cells[6])
                hdr_cells[1].text = _("Current")
                hdr_cells[5].text = _("Required")
                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[2].text = ''
                row_cells[4].text = ''
                row_cells[5].text = ''
                row_cells[6].text = ''
                row_cells[1].merge(row_cells[2])
                row_cells[3].merge(row_cells[4])
                row_cells[5].merge(row_cells[6])
                row_cells[1].text = _("PEST DIAGNOSTIC AND SUPPORT STAFF")
                row_cells[3].text = _("LABORATORY MANAGERS")

                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[1].text = _("No. of Staff")
                row_cells[2].text = _("Average years of experience")
                row_cells[3].text = _("No. of Staff")
                row_cells[4].text = _("Average years of experience")
                row_cells[5].text = _("Diagnostic /support")
                row_cells[6].text = _("Managers")

                mod7 = get_object_or_404(Module7, session_id=sessionid)
                q23 = Module7Matrix23.objects.filter(module7=mod7.id)


                bb=0
                for qq in q23:

                   if bb == 0:
                        tt=_("Mycology")
                   if bb == 4:
                        tt=_("Virology")
                   if bb == 8:
                        tt=_("Nematology")
                   if bb == 12:
                        tt=_("Weed science")
                   if bb == 16:
                        tt=_("Entomology")
                   if bb == 20:
                        tt=_("LMOs")
                   if bb == 24:
                        tt=_("BCAs")
                   if bb == 28:
                        tt=_("Plants For Planting")
                   if bb == 32:
                        tt=_("Treatments")
                   if bb == 36:
                        tt=_("Economist")
                   if bb == 40:
                        tt=_("Staticians")
                   if bb == 44:
                        tt=_("Crop specialists")
                   if bb == 48:
                        tt=_("Technical support and administrative staff")
                   if bb == 0 or  bb == 4 or  bb == 8 or  bb == 12 or  bb == 16 or  bb == 20 or  bb == 24 or  bb == 28 or  bb == 32 or  bb == 36 or  bb == 40 or  bb == 44 or  bb == 48:
                       aa=_("Doctoral equivalent")
                   if bb == 1 or  bb == 5 or  bb == 9 or  bb == 13 or  bb == 17 or  bb == 21 or  bb == 25 or  bb == 29 or  bb == 33 or  bb == 37 or  bb == 41 or  bb == 45 or  bb == 49:
                       aa=_("Master equivalent")
                   if bb == 2 or  bb == 6 or  bb == 10 or  bb == 13 or  bb == 18 or  bb == 22 or  bb == 26 or  bb == 30 or  bb == 34 or  bb == 38 or  bb == 42 or  bb == 46 or  bb == 50:
                       aa=_("Bachelor equivalent")
                   if bb == 3 or  bb == 7 or  bb == 11 or  bb == 15 or  bb == 19 or  bb == 23 or  bb == 27 or  bb == 31 or  bb == 35 or  bb == 39 or  bb == 43 or  bb == 47 or  bb == 51:
                       aa=_("Lower than bachelor level")

                   if bb == 0 or  bb == 4 or  bb == 8 or  bb == 12 or  bb == 16 or  bb == 20 or  bb == 24 or  bb == 28 or  bb == 32 or  bb == 36 or  bb == 40 or  bb == 44 or  bb == 48:
                        row_cells = table.add_row().cells
                        row_cells[1].text = ''
                        row_cells[2].text = ''
                        row_cells[3].text = ''
                        row_cells[4].text = ''
                        row_cells[5].text = ''
                        row_cells[6].text = ''
                        row_cells[0].merge(row_cells[6])
                        row_cells[0].text = tt

                   row_cells = table.add_row().cells
                   row_cells[0].text = aa
                   row_cells[1].text = str(qq.nstaff)
                   row_cells[2].text = str((dict(VAL_AV)[qq.average]))
                   row_cells[3].text = str(qq.nstafflab)
                   row_cells[4].text = str((dict(VAL_AV)[qq.averagelab]))
                  
                   row_cells[5].text = str((dict(BOOL_CHOICESM_M)[qq.support]))
                   row_cells[6].text = str((dict(BOOL_CHOICESM_M)[qq.managers]))


                   bb=bb+1 

                p=document.add_paragraph("")
                p=document.add_paragraph("")
                p=document.add_paragraph("")
            if '8' in items:
                p.add_run(_("Module ")+" 8 - "+ugettext(m_names[8])).bold = True
                p=document.add_paragraph("")
                p.add_run(_("The human resources of the pest surveillance activities")).bold = True



                table = document.add_table(rows=1, cols=7)
                table.style = document.styles['Table Grid']

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = ''
                hdr_cells[1].text = ''
                hdr_cells[2].text = ''
                hdr_cells[3].text = ''
                hdr_cells[4].text = ''
                hdr_cells[5].text = ''
                hdr_cells[6].text = ''

                hdr_cells[1].merge(hdr_cells[4])
                hdr_cells[5].merge(hdr_cells[6])
                hdr_cells[1].text = _("Current")
                hdr_cells[5].text = _("Required")




                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[2].text = ''
                row_cells[4].text = ''
                row_cells[5].text = ''
                row_cells[6].text = ''

                row_cells[1].merge(row_cells[2])
                row_cells[3].merge(row_cells[4])
                row_cells[5].merge(row_cells[6])
                row_cells[1].text = _("PEST SURVEILLANCE AND SUPPORT STAFF")
                row_cells[3].text = _("SURVEILLANCE MANAGERS")

                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[1].text = _("No. of Staff")
                row_cells[2].text = _("Average years of experience")
                row_cells[3].text = _("No. of Staff")
                row_cells[4].text = _("Average years of experience")
                row_cells[5].text = _("Surveillance/support")
                row_cells[6].text = _("Managers")

                mod8 = get_object_or_404(Module8, session_id=sessionid)
                q29 = Module8Matrix30.objects.filter(module8=mod8.id)


                bb=0
                for qq in q29:

                   if bb == 0:
                        tt=_("General surveillance")
                   if bb == 4:
                        tt=_("Specific surveillance")

                   if bb == 8:
                        tt=_("Total")

                   if bb == 0 or  bb == 4 or  bb == 8:
                       aa=_("Doctoral equivalent")
                   if bb == 1 or  bb == 5 or  bb == 9: 
                       aa=_("Master equivalent")
                   if bb == 2 or  bb == 6 or  bb == 10:
                       aa=_("Bachelor equivalent")
                   if bb == 3 or  bb == 7 or  bb == 11:
                       aa=_("Lower than bachelor level")

                   if bb == 0 or  bb == 4 or  bb == 8:
                        row_cells = table.add_row().cells
                        row_cells[1].text = ''
                        row_cells[2].text = ''
                        row_cells[3].text = ''
                        row_cells[4].text = ''
                        row_cells[5].text = ''
                        row_cells[6].text = ''
                        row_cells[0].merge(row_cells[6])
                        row_cells[0].text = tt

                   row_cells = table.add_row().cells
                   row_cells[0].text = aa
                   row_cells[1].text = str(qq.nstaff)
                   row_cells[2].text = str(dict(VAL_AV)[qq.average])
                   row_cells[3].text = str(qq.nstafflab)
                   row_cells[4].text = str(dict(VAL_AV)[qq.averagelab])
                   row_cells[5].text = str((dict(BOOL_CHOICESM_M)[qq.support]))
                   row_cells[6].text = str((dict(BOOL_CHOICESM_M)[qq.managers]))


                   bb=bb+1   

                p=document.add_paragraph("")
                p=document.add_paragraph("")
                p=document.add_paragraph("")
            if '9' in items:
                p.add_run(_("Module ")+" 9 - "+ugettext(m_names[9])).bold = True
                p=document.add_paragraph("")
                p.add_run(_("The human resources involved in eradication programmes")).bold = True



                table = document.add_table(rows=1, cols=7)
                table.style = document.styles['Table Grid']

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = ''
                hdr_cells[1].text = ''
                hdr_cells[2].text = ''
                hdr_cells[3].text = ''
                hdr_cells[4].text = ''
                hdr_cells[5].text = ''
                hdr_cells[6].text = ''

                hdr_cells[1].merge(hdr_cells[4])
                hdr_cells[5].merge(hdr_cells[6])
                hdr_cells[1].text = _("Current")
                hdr_cells[5].text = _("Required")




                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[2].text = ''
                row_cells[4].text = ''
                row_cells[5].text = ''
                row_cells[6].text = ''

                row_cells[1].merge(row_cells[2])
                row_cells[3].merge(row_cells[4])
                row_cells[5].merge(row_cells[6])
                row_cells[1].text = _("Field and support Staff")
                row_cells[3].text = _("Eradication managers")

                row_cells = table.add_row().cells
                row_cells[0].text =''
                row_cells[1].text = _("No. of Staff")
                row_cells[2].text = _("Average years of experience")
                row_cells[3].text = _("No. of Staff")
                row_cells[4].text = _("Average years of experience")
                row_cells[5].text = _("Field/support Stafft")
                row_cells[6].text = _("Managers")

                mod9 = get_object_or_404(Module9, session_id=sessionid)
                q35 = Module9Matrix35.objects.filter(module9=mod9.id)


                bb=0
                for qq in q35:

                   if bb == 0 :
                       aa=_("Doctoral equivalent")
                   if bb == 1 : 
                       aa=_("Master equivalent")
                   if bb == 2:
                       aa=_("Bachelor equivalent")
                   if bb == 3:
                       aa=_("Lower than bachelor level")


                   row_cells = table.add_row().cells
                   row_cells[0].text = aa
                   row_cells[1].text = str(qq.nstaff)
                   row_cells[2].text = str((dict(VAL_AV)[qq.average]))
                   row_cells[3].text = str(qq.nstafflab)
                   row_cells[4].text = str((dict(VAL_AV)[qq.averagelab]))
                   row_cells[5].text = str((dict(BOOL_CHOICESM_M)[qq.support]))
                   row_cells[6].text = str((dict(BOOL_CHOICESM_M)[qq.managers]))


                   bb=bb+1   

                p=document.add_paragraph("")
                p=document.add_paragraph("")
                p=document.add_paragraph("")
            if '10' in items:
                p.add_run(_("Module ")+" 10 - "+ugettext(m_names[10])).bold = True
                p=document.add_paragraph("")
                p.add_run(_("The import regulatory system human resour")).bold = True



                table = document.add_table(rows=1, cols=9)
                table.style = document.styles['Table Grid']

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = ''
                hdr_cells[1].text = ''
                hdr_cells[2].text = ''
                hdr_cells[3].text = ''
                hdr_cells[4].text = ''
                hdr_cells[5].text = ''
                hdr_cells[6].text = ''
                hdr_cells[7].text = ''
                hdr_cells[8].text = ''

                hdr_cells[1].merge(hdr_cells[5])
                hdr_cells[6].merge(hdr_cells[8])
                hdr_cells[1].text = _("Current")
                hdr_cells[6].text = _("Required")




                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[1].text = ''
                row_cells[2].text = ''
                row_cells[3].text = ''
                row_cells[4].text = ''
                row_cells[5].text = ''
                hdr_cells[6].text = ''
                hdr_cells[7].text = ''
                hdr_cells[8].text = ''

                row_cells[1].merge(row_cells[2])
                row_cells[3].merge(row_cells[4])
                #row_cells[5].merge(row_cells[6])
                row_cells[6].merge(row_cells[8])
                row_cells[1].text = _("TECHNICAL")
                row_cells[3].text = _("IMPORT MANAGERS")
                row_cells[5 ].text = _("SUPPORT STAFF")

                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[1].text = _("No. of Staff")
                row_cells[2].text = _("Average years of experience")
                row_cells[3].text = _("No. of Staff")
                row_cells[4].text = _("Average years of experience")
                row_cells[5].text = _("No. of Staff")
                row_cells[6].text = _("Technical")
                row_cells[7].text = _("Managers")
                row_cells[8].text = _("Support")

                mod10 = get_object_or_404(Module10, session_id=sessionid)
                q47 = Module10Matrix_47.objects.filter(module10=mod10.id)


                bb=0
                for qq in q47:

                   if bb == 0 :
                       aa=_("Doctoral equivalent")
                   if bb == 1 : 
                       aa=_("Master equivalent")
                   if bb == 2:
                       aa=_("Bachelor equivalent")
                   if bb == 3:
                       aa=_("Lower than bachelor level")

                   row_cells = table.add_row().cells
                   row_cells[0].text = aa
                   row_cells[1].text = str(qq.nstaff)
                   row_cells[2].text = str((dict(VAL_AV)[qq.average]))
                   row_cells[3].text = str(qq.nstafflab)
                   row_cells[4].text = str((dict(VAL_AV)[qq.averagelab]))
                   row_cells[5].text = str(qq.supstafflab)
                   row_cells[6].text = str((dict(BOOL_CHOICESM_M)[qq.technical]))
                   row_cells[7].text = str((dict(BOOL_CHOICESM_M)[qq.managers]))
                   row_cells[8].text = str((dict(BOOL_CHOICESM_M)[qq.support]))


                   bb=bb+1   
                p=document.add_paragraph("")
                p=document.add_paragraph("")
                p=document.add_paragraph("")
            if '11' in items:
                p.add_run(_("Module ")+" 11 - "+ugettext(m_names[11])).bold = True
                p=document.add_paragraph("")
                p.add_run(_("The human resources of the NPPO PRA programme")).bold = True

                table = document.add_table(rows=1, cols=5)
                table.style = document.styles['Table Grid']


                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = ''
                hdr_cells[1].text = ''
                hdr_cells[2].text = ''
                hdr_cells[3].text = ''
                hdr_cells[4].text = ''


                hdr_cells[1].merge(hdr_cells[3])
                hdr_cells[1].text = _("Current")
                hdr_cells[4].text = _("Required")
                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[1].text = ''
                row_cells[2].text = ''
                row_cells[3].text = ''
                row_cells[4].text = ''

                row_cells[1].merge(row_cells[2])
                row_cells[1].text = _("Full Time Staff")
                row_cells[3].text = _("Ad hoc Staff")

                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[1].text = _("No. of Staff")
                row_cells[2].text = _("Average years of experience")
                row_cells[3].text = _("No. of Staff")
                row_cells[4].text = _("Required")

                mod11 = get_object_or_404(Module11, session_id=sessionid)
                q42 = Module11Matrix42.objects.filter(module11=mod11.id)


                bb=0
                for qq in q42:

                   if bb == 0:
                        tt=_("Mycology")
                   if bb == 4:
                        tt=_("Virology")
                   if bb == 8:
                        tt=_("Nematology")
                   if bb == 12:
                        tt=_("Weed science")
                   if bb == 16:
                        tt=_("Entomology")
                   if bb == 20:
                        tt=_("LMOs")
                   if bb == 24:
                        tt=_("BCAs")
                   if bb == 28:
                        tt=_("Plants For Planting")
                   if bb == 32:
                        tt=_("Treatments")
                   if bb == 36:
                        tt=_("Economist")
                   if bb == 40:
                        tt=_("Staticians")
                   if bb == 44:
                        tt=_("Crop specialists")
                   if bb == 48:
                        tt=_("Technical support and administrative staff")
                   if bb == 0 or  bb == 4 or  bb == 8 or  bb == 12 or  bb == 16 or  bb == 20 or  bb == 24 or  bb == 28 or  bb == 32 or  bb == 36 or  bb == 40 or  bb == 44 or  bb == 48:
                       aa=_("Doctoral equivalent")
                   if bb == 1 or  bb == 5 or  bb == 9 or  bb == 13 or  bb == 17 or  bb == 21 or  bb == 25 or  bb == 29 or  bb == 33 or  bb == 37 or  bb == 41 or  bb == 45 or  bb == 49:
                       aa=_("Master equivalent")
                   if bb == 2 or  bb == 6 or  bb == 10 or  bb == 13 or  bb == 18 or  bb == 22 or  bb == 26 or  bb == 30 or  bb == 34 or  bb == 38 or  bb == 42 or  bb == 46 or  bb == 50:
                       aa=_("Bachelor equivalent")
                   if bb == 3 or  bb == 7 or  bb == 11 or  bb == 15 or  bb == 19 or  bb == 23 or  bb == 27 or  bb == 31 or  bb == 35 or  bb == 39 or  bb == 43 or  bb == 47 or  bb == 51:
                       aa=_("Lower than bachelor level")

                   if bb == 0 or  bb == 4 or  bb == 8 or  bb == 12 or  bb == 16 or  bb == 20 or  bb == 24 or  bb == 28 or  bb == 32 or  bb == 36 or  bb == 40 or  bb == 44 or  bb == 48:
                        row_cells = table.add_row().cells
                        row_cells[1].text = ''
                        row_cells[2].text = ''
                        row_cells[3].text = ''
                        row_cells[4].text = ''
                        row_cells[0].merge(row_cells[4])
                        row_cells[0].text = tt

                   row_cells = table.add_row().cells
                   row_cells[0].text = aa
                   row_cells[1].text = str(qq.nstaff)
                   row_cells[2].text = str((dict(VAL_AV)[qq.average]))
                   row_cells[3].text = str(qq.nstafflab)

                   row_cells[4].text = str((dict(BOOL_CHOICESM_M)[qq.support]))


                   bb=bb+1    

                p=document.add_paragraph("")
                p=document.add_paragraph("")
                p=document.add_paragraph("")
            if '12' in items:
                p.add_run(_("Module ")+" 12 - "+ugettext(m_names[12])).bold = True
                p=document.add_paragraph("")
                p.add_run(_("PFA/ALPP/PFPP/PFPS - human resources")).bold = True



                table = document.add_table(rows=1, cols=7)
                table.style = document.styles['Table Grid']

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = ''
                hdr_cells[1].text = ''
                hdr_cells[2].text = ''
                hdr_cells[3].text = ''
                hdr_cells[4].text = ''
                hdr_cells[5].text = ''
                hdr_cells[6].text = ''

                hdr_cells[1].merge(hdr_cells[4])
                hdr_cells[5].merge(hdr_cells[6])
                hdr_cells[1].text = _("Current")
                hdr_cells[5].text = _("Required")




                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[2].text = ''
                row_cells[4].text = ''
                row_cells[5].text = ''
                row_cells[6].text = ''

                row_cells[1].merge(row_cells[2])
                row_cells[3].merge(row_cells[4])
                row_cells[5].merge(row_cells[6])
                row_cells[1].text = _("Field and support Staff")
                row_cells[3].text = _("Eradication managers")

                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[1].text = _("No. of Staff")
                row_cells[2].text = _("Average years of experience")
                row_cells[3].text = _("No. of Staff")
                row_cells[4].text = _("Average years of experience")
                row_cells[5].text = _("Field/support Staff")
                row_cells[6].text = _("Support")

                mod12 = get_object_or_404(Module12, session_id=sessionid)
                q22 = Module12Matrix22.objects.filter(module12=mod12.id)


                bb=0
                for qq in q22:

                   if bb == 0 or  bb == 4 or  bb == 8:
                       aa=_("Doctoral equivalent")
                   if bb == 1 or  bb == 5 or  bb == 9: 
                       aa=_("Master equivalent")
                   if bb == 2 or  bb == 6 or  bb == 10:
                       aa=_("Bachelor equivalent")
                   if bb == 3 or  bb == 7 or  bb == 11:
                       aa=_("Lower than bachelor level")


                   row_cells = table.add_row().cells
                   row_cells[0].text = aa
                   row_cells[1].text = str(qq.nstaff)
                   row_cells[2].text = str((dict(VAL_AV)[qq.average]))
                   row_cells[3].text = str(qq.nstafflab)
                   row_cells[4].text = str((dict(VAL_AV)[qq.averagelab]))
                   row_cells[5].text = str((dict(BOOL_CHOICESM_M)[qq.support]))
                   row_cells[6].text = str((dict(BOOL_CHOICESM_M)[qq.managers]))


                   bb=bb+1 


                p=document.add_paragraph("")
                p=document.add_paragraph("")
                p=document.add_paragraph("")
            if '13' in items:
                p.add_run( _("Module ")+" 13 - "+ugettext(m_names[13])).bold = True
                p=document.add_paragraph("")
                p.add_run( _("The human resources in the NPPO export certification program")).bold = True



                table = document.add_table(rows=1, cols=7)
                table.style = document.styles['Table Grid']

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = ''
                hdr_cells[1].text = ''
                hdr_cells[2].text = ''
                hdr_cells[3].text = ''
                hdr_cells[4].text = ''
                hdr_cells[5].text = ''
                hdr_cells[6].text = ''

                hdr_cells[1].merge(hdr_cells[4])
                hdr_cells[5].merge(hdr_cells[6])
                hdr_cells[1].text =  _("Current")
                hdr_cells[5].text =  _("Required")




                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[2].text = ''
                row_cells[4].text = ''
                row_cells[5].text = ''
                row_cells[6].text = ''

                row_cells[1].merge(row_cells[2])
                row_cells[3].merge(row_cells[4])
                row_cells[5].merge(row_cells[6])
                row_cells[1].text = _("Inspector and support Staff")
                row_cells[3].text = _("Export managers")

                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[1].text = _("No. of Staff")
                row_cells[2].text = _("Average years of experience")
                row_cells[3].text = _("No. of Staff")
                row_cells[4].text = _("Average years of experience")
                row_cells[5].text = _("Inspector/support Staff")
                row_cells[6].text = _("Export managers")

                mod13 = get_object_or_404(Module13, session_id=sessionid)
                q47 = Module13Matrix47.objects.filter(module13=mod13.id)


                bb=0
                for qq in q47:
                   if bb == 0 or  bb == 4 or  bb == 8:
                       aa=_("Doctoral equivalent")
                   if bb == 1 or  bb == 5 or  bb == 9: 
                       aa=_("Master equivalent")
                   if bb == 2 or  bb == 6 or  bb == 10:
                       aa=_("Bachelor equivalent")
                   if bb == 3 or  bb == 7 or  bb == 11:
                       aa=_("Lower than bachelor level")

                   if bb == 0 or  bb == 4 or  bb == 8:
                        row_cells = table.add_row().cells
                        row_cells[1].text = ''
                        row_cells[2].text = ''
                        row_cells[3].text = ''
                        row_cells[4].text = ''
                        row_cells[5].text = ''
                        row_cells[6].text = ''
                        row_cells[0].merge(row_cells[6])
                        row_cells[0].text = tt

                   row_cells = table.add_row().cells
                   row_cells[0].text = aa
                   row_cells[1].text = str(qq.nstaff)
                   row_cells[2].text = str((dict(VAL_AV)[qq.average]))
                   row_cells[3].text = str(qq.nstafflab)
                   row_cells[4].text = str((dict(VAL_AV)[qq.averagelab]))
                   row_cells[5].text = str((dict(BOOL_CHOICESM_M)[qq.support]))
                   row_cells[6].text = str((dict(BOOL_CHOICESM_M)[qq.managers]))


                   bb=bb+1 

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    document.save(response)
    return response 
